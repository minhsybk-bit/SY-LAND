# -*- coding: utf-8 -*-
"""
PHẦN MỀM ĐỐI CHIẾU / ĐỔI TÊN / TÁCH GỘP PDF / XUẤT WORD - SỸ LAND (v4)
--------------------------------------------------------------------------
1) Quét PDF trong 1-2 thư mục, đối chiếu với Excel theo tiêu chí TÙY CHỌN:
   Mã xã / Số tờ / Số thửa / Diện tích / Tên chủ (có nút chọn nhanh tổ hợp)
2) File khớp -> copy sang thư mục đích. File KHÔNG khớp -> copy sang thư mục riêng.
3) Đổi tên hàng loạt PDF theo MẪU TÊN TỰ NHẬP, dựa vào Excel.
4) MỚI: Đổi tên hàng loạt bằng cách ĐỌC TRỰC TIẾP NỘI DUNG file PDF (không cần Excel)
   - lấy Thửa đất số / Tờ bản đồ số / Địa chỉ / Tên xã ngay trong file, tự OCR nếu
   là bản scan - rồi tạo file lệnh .bat với các dòng `ren "tên cũ" "tên mới"`.
5) MỚI: Tách trang PDF (trang 1 / trang 2 / trang 1+2 / tùy chọn) và Gộp nhiều PDF.
6) MỚI: Xuất Word hàng loạt - mỗi PDF ra 1 file .docx riêng, hoặc gộp tất cả vào
   1 file .docx duy nhất, giữ nguyên bố cục nhiều nhất có thể.

Cài đặt (1 lần) - mở Command Prompt gõ:
   pip install openpyxl pypdf pymupdf pytesseract pillow pdf2docx docxcompose python-docx

Riêng chức năng OCR (đọc file PDF dạng scan) cần cài thêm Tesseract-OCR (chương
trình riêng, không phải gói pip): tải tại
https://github.com/UB-Mannheim/tesseract/wiki (bản Windows), khi cài nhớ tick
thêm ngôn ngữ "Vietnamese". Không có Tesseract thì các chức năng khác vẫn
chạy bình thường, chỉ riêng phần đọc file PDF scan sẽ báo lỗi.

Chạy: double-click file này (hoặc chuột phải > Open with > Python)
"""

import json
import os
import random
import re
import io
import shutil
import string
import subprocess
import sys
import time
import datetime
import unicodedata
import threading
import urllib.request
import urllib.error
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from tkinter import font as tkfont

try:
    import openpyxl
    from openpyxl.styles import Font
except ImportError:
    raise SystemExit(
        "Thieu thu vien openpyxl.\n"
        "Hay mo Command Prompt va go: pip install openpyxl\n"
        "Roi chay lai file nay."
    )

try:
    from pypdf import PdfReader, PdfWriter
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import fitz  # PyMuPDF - dùng để render trang PDF thành ảnh cho OCR
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    import pytesseract
    from PIL import Image, ImageOps, ImageFilter
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

if HAS_OCR:
    def _configure_bundled_tesseract():
        """
        Nếu phần mềm được đóng gói kèm Tesseract OCR portable (thư mục tools/tesseract/tesseract.exe
        cạnh phần mềm - xem build/build_exe.bat), tự động dùng bản portable này thay vì trông chờ
        Tesseract đã cài + có sẵn trong PATH hệ thống. Giúp người dùng KHÔNG cần tự cài Tesseract
        hay tự thêm vào PATH thủ công. Nếu không có bản portable kèm theo (VD khi chạy trực tiếp
        file .py lúc phát triển trên máy đã cài Tesseract hệ thống), giữ nguyên hành vi cũ - dùng
        Tesseract từ PATH hệ thống như bình thường.
        """
        try:
            bundled_dir = os.path.join(get_base_dir(), "tools", "tesseract")
            bundled_exe = os.path.join(bundled_dir, "tesseract.exe" if sys.platform == "win32" else "tesseract")
            if os.path.isfile(bundled_exe):
                pytesseract.pytesseract.tesseract_cmd = bundled_exe
                os.environ["TESSDATA_PREFIX"] = os.path.join(bundled_dir, "tessdata")
                return True
        except Exception:
            pass
        return False

try:
    import win32com.client as win32com_client
    HAS_WIN32COM = (sys.platform == "win32")
except ImportError:
    HAS_WIN32COM = False

try:
    from pdf2docx import Converter as PDF2DocxConverter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    import matplotlib
    matplotlib.use("TkAgg")
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    from docx import Document as DocxDocument
    from docxcompose.composer import Composer as DocxComposer
    HAS_DOCXCOMPOSE = True
except ImportError:
    HAS_DOCXCOMPOSE = False

# Số phiên bản - tăng lên mỗi khi có bản vá lỗi hoặc bổ sung tính năng mới.
# Dùng để bạn biết mình đang chạy bản nào (hiện ở tiêu đề cửa sổ).
APP_VERSION = "13.84.0"

# Nhãn phát hành ổn định - đánh dấu mốc "chốt bản" để dùng cho công việc thực tế (theo tài liệu
# "LỆNH TỔNG: CHỐT BẢN ỔN ĐỊNH"). Đi kèm với APP_VERSION (số hiệu build nội bộ, vẫn tăng đều theo
# từng lần sửa lỗi/cải tiến) - không thay thế, chỉ đánh dấu đây là bản đã qua kiểm thử toàn diện,
# đủ ổn định để dùng cho công việc thực tế.
RELEASE_LABEL = "v1.0 Stable"
RELEASE_DATE = "10/07/2026"

# Lịch sử cập nhật (mới nhất ở đầu) - hiển thị trong Tab Cài đặt để biết mỗi bản có gì mới,
# không cần mở file HUONG_DAN_SU_DUNG.md riêng. Chỉ ghi các điểm chính, xem đầy đủ trong
# HUONG_DAN_SU_DUNG.md nếu cần chi tiết hơn.
CHANGELOG = [
    ("13.84.0", "Mở rộng Task Manager thống nhất (mục IV) từ thí điểm Tab 7d sang ĐỦ CẢ 5 luồng "
                "xử lý hàng loạt chính: Tab 8a (Tổng hợp báo cáo Excel), Tab 8c (Điền Excel mẫu từ "
                "PDF), Tab 5 (OCR đổi tên PDF), Tab 6 (Xoay/làm thẳng PDF) - mỗi Tab giờ đều có 1 "
                "TaskInfo theo dõi SONG SONG (tiến độ, số lượng thành công/lỗi/cần kiểm tra, thời "
                "gian), tiếp tục KHÔNG thay thế progress bar/biến đếm hiện có của từng Tab. Đồng bộ "
                "đầy đủ module thật và fallback cho cả 4 Tab. Tab 6 xử lý riêng trạng thái "
                "DA_DUNG_BOI_NGUOI_DUNG khi bị hủy giữa chừng (khác THANH_CONG khi hoàn thành trọn "
                "vẹn) - đúng ý nghĩa chuẩn theo tài liệu. PHÁT HIỆN VÀ TỰ SỬA 1 lỗi thao tác trong "
                "lúc chỉnh sửa fallback Tab 6: str_replace vô tình xóa mất dòng 'if not dry_run:' "
                "khi thay thế đoạn code - bắt được NGAY qua kiểm tra cú pháp thường lệ (IndentationError) "
                "trước khi kịp chạy test, không cần debug sâu - minh chứng giá trị của việc LUÔN "
                "kiểm tra cú pháp ngay sau mỗi lần sửa dù là thay đổi nhỏ. Mở rộng kịch bản test 25 "
                "xác nhận CẢ 4 Tab mới (không chỉ Tab 7d) đã tích hợp đúng ở cả 2 môi trường. 25/25 "
                "bộ test tự động PASS, không hồi quy. Mục IV tài liệu Release Candidate coi như "
                "HOÀN THÀNH Ở MỨC 'từ vựng trạng thái + theo dõi thống nhất song song' cho TOÀN BỘ "
                "5 luồng hàng loạt chính - còn lại (nếu muốn đi xa hơn) là cân nhắc có cần 1 class "
                "TaskManager TRUNG TÂM quản lý nhiều tác vụ cùng lúc hay không, hiện mỗi Tab tự giữ "
                "TaskInfo riêng đã đủ đáp ứng tinh thần tài liệu với rủi ro thấp nhất."),
    ("13.83.0", "Mục IV tài liệu Release Candidate - TASK MANAGER THỐNG NHẤT (khởi đầu): module "
                "mới app/services/task_manager_service.py chuẩn hóa 16 trạng thái tác vụ (CHO_XU_"
                "LY, DANG_KHOI_TAO, DANG_DOC_FILE, DANG_PHAN_TICH, DANG_OCR, DANG_DOI_CHIEU, "
                "DANG_CAP_NHAT, DANG_XUAT_FILE, DANG_LUU, TAM_DUNG, DANG_DUNG_AN_TOAN, THANH_CONG, "
                "CAN_KIEM_TRA, LOI, BO_QUA, DA_DUNG_BOI_NGUOI_DUNG - đúng danh sách tài liệu) và "
                "dataclass TaskInfo (tiến độ, số lượng thành công/lỗi/cần kiểm tra, thời gian đã "
                "chạy/dự kiến còn lại, phần trăm hoàn thành). QUYẾT ĐỊNH KIẾN TRÚC QUAN TRỌNG: đây "
                "là lớp BỔ SUNG song song, KHÔNG viết lại/thay thế TaskControl (pause/resume/cancel) "
                "đã hoạt động ổn định qua 85 giai đoạn - đúng nguyên tắc mục XXXV.1 'Không viết lại "
                "toàn bộ dự án khi chưa có test bảo vệ'. THÍ ĐIỂM áp dụng vào Tab 7d (Xuất Word "
                "hàng loạt) - TaskInfo theo dõi SONG SONG với progress bar/biến đếm hiện có, không "
                "thay thế logic đã kiểm chứng. Đồng bộ đầy đủ module thật và fallback. Test module "
                "lõi kỹ (khởi tạo, cập nhật tiến độ, ghi nhận kết quả đa dạng tên trạng thái cũ, "
                "trường hợp biên tong_so_luong=0 không chia-cho-0). Bổ sung kịch bản test 25 mới. "
                "25/25 bộ test tự động PASS ở cả 2 môi trường, không hồi quy. CÒN LẠI của mục IV: "
                "mở rộng TaskInfo sang Tab 5/6/8a/8c (hiện chỉ 7d), và cân nhắc có nên tái cấu "
                "trúc SÂU HƠN (VD 1 class TaskManager trung tâm quản lý TẤT CẢ tác vụ đang chạy) "
                "hay giữ mô hình 'mỗi Tab tự giữ 1 TaskInfo' như hiện tại (đơn giản hơn, rủi ro "
                "thấp hơn, có thể ĐÃ ĐỦ đáp ứng tinh thần mục IV mà không cần class trung tâm)."),
    ("13.82.0", "Mục XXII tài liệu Release Candidate - XUẤT GÓI CHẨN ĐOÁN: mục mới trong Tab "
                "Cài đặt cho phép xuất 1 file ZIP duy nhất (phiên bản phần mềm, hệ điều hành, "
                "trạng thái Python/Tesseract, cấu hình hiệu năng, trạng thái bản quyền, thống kê "
                "CPU/RAM, danh sách module đã tải, log lỗi gần nhất) để gửi Nguyễn Sỹ hỗ trợ kỹ "
                "thuật từ xa - không cần người dùng chụp ảnh màn hình hay mô tả bằng lời. AN TOÀN "
                "THÔNG TIN nghiêm ngặt (đúng yêu cầu tài liệu): key bản quyền LUÔN ở dạng che "
                "(SYLB-****-****-5678), KHÔNG BAO GIỜ đưa mật khẩu/token/key đầy đủ vào gói, KHÔNG "
                "kèm bất kỳ file hồ sơ đất đai nào của người dùng - chỉ chứa file .json/.txt mô tả "
                "môi trường/cấu hình. Module mới app/services/diagnostic_service.py. PHÁT HIỆN VÀ "
                "SỬA 1 LỖI THẬT ngay khi viết xong (thiếu import time gây NameError) - nhờ CHẠY "
                "SỚM kịch bản kiểm tra tĩnh 14 (thay vì chỉ test chức năng trước) nên phát hiện "
                "và sửa gần như ngay lập tức, rút ra bài học: luôn chạy các kịch bản kiểm tra tĩnh "
                "nhanh (14, 22) NGAY sau khi viết code UI mới, trước khi test chức năng đầy đủ. "
                "Bổ sung kịch bản test 24 mới - đặc biệt kiểm tra BẢO MẬT bằng cách quét TOÀN BỘ "
                "nội dung mọi file trong gói ZIP tìm chuỗi key đầy đủ (không chỉ kiểm tra 1 file "
                "cụ thể), và xác nhận không có phần mở rộng file hồ sơ nào (pdf/docx/xlsx) lọt "
                "vào gói. 24/24 bộ test tự động PASS ở cả 2 môi trường, không hồi quy."),
    ("13.81.0", "Hoàn thiện Checkpoint/Resume (mục XVII) - mở rộng sang Tab 5 (OCR đổi tên, phát "
                "hiện đã được viết sẵn từ trước, xác nhận hoạt động đúng) và Tab 6 (xoay/làm thẳng "
                "PDF, kiến trúc khác - vòng lặp trực tiếp trong UI worker thay vì hàm backend "
                "riêng). PHÁT HIỆN VÀ SỬA 1 LỖI THẬT KHÁC khi mở rộng sang Tab 6: vòng lặp xử lý "
                "từng file dùng except Exception CHUNG, khiến TaskCancelled (khi người dùng bấm "
                "Hủy) bị coi là LỖI của file hiện tại thay vì DỪNG HẲN - làm nút Hủy ở Tab 6 KHÔNG "
                "hoạt động đúng (âm thầm tiếp tục xử lý các file còn lại thay vì dừng ngay theo "
                "yêu cầu người dùng). Đã sửa bằng cách bắt TaskCancelled RIÊNG TRƯỚC except "
                "Exception chung, ở cả module thật (tab_pdfedit.py) và fallback. Giờ Tab 7d, 8a, "
                "8c, 5, 6 đều có Checkpoint/Resume đầy đủ. Bổ sung kịch bản test 23 tái hiện đúng "
                "lỗi TaskCancelled + xác nhận checkpoint/resume Tab 6 hoạt động đúng. Tận dụng "
                "kịch bản 22 (kiểm tra tĩnh fallback parity, phát hiện có sẵn từ trước) để xác "
                "nhận NHANH toàn bộ thay đổi trong đợt này không làm thiếu method nào ở fallback - "
                "không cần chờ chạy GUI đầy đủ. 23/23 bộ test tự động PASS ở cả 2 môi trường, "
                "không hồi quy. Checkpoint/Resume giờ bao phủ toàn bộ 5 luồng xử lý hàng loạt "
                "chính (7d/8a/8c/5/6) - còn lại Task Manager thống nhất (mục IV, quy mô lớn hơn "
                "nhiều) là phần chính chưa làm của Giai đoạn 3 tài liệu Release Candidate. LƯU Ý "
                "THÊM: kịch bản test 22 (kiểm tra tĩnh fallback, có sẵn từ trước) chỉ hoạt động "
                "trong môi trường phát triển - hard-code tên file 'doi_chieu_pdf_v3.py', KHÔNG tồn "
                "tại trong gói đóng gói thật (tên 'doi_chieu_pdf.py') - chỉ phát hiện khi CHẠY THỬ "
                "TRONG GÓI THẬT (đúng quy trình luôn kiểm tra cả 2 môi trường), đã sửa để tự động "
                "phát hiện đúng tên file ở cả 2 nơi."),
    ("13.80.0", "Mở rộng Checkpoint/Resume (mục XVII) sang Tab 8a (Tổng hợp báo cáo ra Excel) và "
                "Tab 8c (Điền Excel MẪU từ PDF) - theo ĐÚNG mẫu đã kiểm chứng ở Tab 7d (giai đoạn "
                "trước): lưu tiến độ định kỳ, hỏi resume khi phát hiện tác vụ dở dang, CHỈ xóa "
                "checkpoint khi THỰC SỰ hoàn thành (dùng cờ bi_huy_giua_chung - đúng bài học từ lỗi "
                "thật đã sửa ở 7d, áp dụng cẩn thận ngay từ đầu nên KHÔNG lặp lại lỗi tương tự lần "
                "này). Test resume Tab 8a/8c bằng control giả lập hủy giữa chừng xác nhận đúng "
                "ngay lần đầu. Bổ sung test Tab 8a vào kịch bản 21. 21/21 bộ test tự động PASS ở "
                "cả 2 môi trường, không hồi quy. Còn lại: Tab 5 (OCR đổi tên) và Tab 6 (xoay PDF) "
                "chưa có checkpoint - quy mô nhỏ hơn, có thể làm ở đợt sau."),
    ("13.79.0", "Mục XVII tài liệu Release Candidate - CHECKPOINT/RESUME: với tác vụ dài (Tab "
                "7d xuất Word hàng loạt, thí điểm trước 1 luồng để chứng minh cơ chế đúng trước "
                "khi mở rộng), lưu tiến độ định kỳ ra đĩa (Data/checkpoints/) - nếu phần mềm bị "
                "đóng đột ngột (mất điện/crash/tắt máy) giữa lúc đang xử lý, LẦN SAU bấm nút Chạy "
                "lại sẽ TỰ PHÁT HIỆN và HỎI người dùng có muốn TIẾP TỤC từ chỗ dở dang (bỏ qua các "
                "dòng đã xong) hay BẮT ĐẦU LẠI TỪ ĐẦU - không tự động chạy lại toàn bộ, không ghi "
                "đè file đã hoàn thành. Module mới app/services/checkpoint_service.py: task_id ổn "
                "định tính từ (loại tác vụ, nguồn, đích) để nhận diện đúng tác vụ giữa các lần mở "
                "phần mềm; ghi checkpoint NGUYÊN TỬ (dùng chung file_safety_service vừa làm). "
                "PHÁT HIỆN VÀ SỬA 1 LỖI THẬT QUAN TRỌNG trong lúc tích hợp: đoạn code xóa "
                "checkpoint ở CUỐI HÀM chạy VÔ ĐIỀU KIỆN, kể cả khi vòng lặp bị 'break' giữa "
                "chừng do người dùng hủy - khiến checkpoint bị XÓA NGAY LẬP TỨC dù CHƯA hoàn "
                "thành, làm toàn bộ tính năng resume VÔ DỤNG (không bao giờ có checkpoint để đọc "
                "lại). Chỉ phát hiện được qua test TÍCH HỢP THẬT (gọi đúng hàm với control giả "
                "lập hủy giữa chừng, kiểm tra checkpoint còn tồn tại sau đó) - không phải test đơn "
                "vị cô lập từng hàm nhỏ (các hàm con đều đúng riêng lẻ, lỗi nằm ở LOGIC ĐIỀU KIỆN "
                "kết hợp giữa chúng). Đã sửa bằng cờ bi_huy_giua_chung, chỉ xóa checkpoint khi "
                "THỰC SỰ xử lý hết toàn bộ. Bổ sung kịch bản test 21 tái hiện ĐÚNG lỗi đã gặp để "
                "chống hồi quy. 21/21 bộ test tự động PASS ở cả 2 môi trường."),
    ("13.78.0", "Tiếp tục mục XXI tài liệu Release Candidate - GHI FILE NGUYÊN TỬ cho TOÀN BỘ 8 "
                "điểm ghi file đầu ra CỐT LÕI (Excel: đối chiếu Tab 3/4, tổng hợp 8a/8c, đối chiếu "
                "8d; Word: mail-merge 7c x2 nơi, gộp 7b x2 nơi, xuất Word từ Excel 8b, gộp GT+TBXN "
                "7d; PDF: xoay/làm thẳng Tab 6) - trước đây ghi TRỰC TIẾP ra file đích, nếu phần "
                "mềm bị tắt đột ngột (mất điện/crash) NGAY LÚC đang ghi có thể để lại file ĐẦU RA "
                "HỎNG/DỞ DANG. Module mới app/services/file_safety_service.py: ghi vào file tạm "
                "CÙNG THƯ MỤC đích (đảm bảo rename luôn cùng ổ đĩa), kiểm tra file tạm hợp lệ "
                "(tồn tại + không rỗng), os.replace() NGUYÊN TỬ sang tên chính thức - tại MỌI THỜI "
                "ĐIỂM file đích hoặc là bản CŨ HOÀN CHỈNH hoặc bản MỚI HOÀN CHỈNH, không có trạng "
                "thái cắt cụt giữa chừng; nếu lỗi thì XÓA file tạm và NÉM LẠI lỗi gốc (không nuốt "
                "lỗi, không để rác). PHÁT HIỆN QUAN TRỌNG khi rà soát: 1 trong 8 điểm "
                "(mail_merge_docx_template) có CẢ fallback trong file chính LẪN hàm thật riêng "
                "trong app/services/word_service.py - ban đầu chỉ sửa fallback, rà soát kỹ lại mới "
                "phát hiện và sửa đủ CẢ HAI. Đã xác nhận 6/7 hàm còn lại KHÔNG có fallback pattern "
                "(chỉ tồn tại 1 nơi trong file chính) VÀ được giao diện thật gọi trực tiếp (rút "
                "kinh nghiệm từ bài học Tab 6 ở đợt trước - luôn xác minh UI THẬT gọi đúng hàm đã "
                "sửa, không chỉ tin vào tên hàm). Bổ sung kịch bản test 20 mới: ghi thành công đúng "
                "nội dung không sót file tạm, lỗi giữa chừng không tạo file đích không sót file "
                "tạm và ném lại đúng lỗi gốc, ghi đè nguyên tử đúng, áp dụng thật với openpyxl, và "
                "RÀ SOÁT MÃ NGUỒN xác nhận đủ ≥8 điểm đã áp dụng (chống hồi quy nếu sau này ai vô "
                "tình sửa lại thành ghi trực tiếp). 20/20 bộ test tự động PASS ở cả 2 môi trường."),
    ("13.77.0", "GIAI ĐOẠN 1-2 tài liệu 'RÀ SOÁT, KIỂM THỬ VÀ TỐI ƯU TOÀN BỘ PHẦN MỀM TRƯỚC KHI "
                "PHÁT HÀNH': đo baseline (BASELINE_TRUOC_TOI_UU.md) rồi MỞ RỘNG cơ chế Resource "
                "Governor (word_perf_service, đã kiểm chứng ở mục 7d) sang Tab 5 (OCR đổi tên), "
                "Tab 6 (xoay/làm thẳng PDF), Tab 8a/8c (tổng hợp/điền báo cáo từ PDF) và mục 7c "
                "(mail-merge từ PDF) - toàn bộ trước đây chạy OCR hàng loạt không nghỉ, cùng lớp "
                "rủi ro treo Windows đã xử lý cho Word. CHỈ throttle khi có TỪ 2 FILE TRỞ LÊN (an "
                "toàn batch nhỏ). PHÁT HIỆN QUAN TRỌNG trong lúc sửa Tab 6: hàm batch_auto_rotate_"
                "deskew_pdf() có sẵn nhưng KHÔNG được giao diện thật gọi - Tab 6 dùng vòng lặp "
                "RIÊNG trong tab_pdfedit.py gọi hàm đơn lẻ auto_rotate_deskew_pdf() - nếu không rà "
                "soát kỹ, tối ưu sẽ vô tác dụng trên giao diện thật; đã sửa ĐÚNG vị trí. Log hiệu "
                "năng dùng CHUNG Logs/LOG_HIEU_NANG_XU_LY.csv cho mọi loại tác vụ. Đo benchmark "
                "trước/sau xác nhận: throttle 'cân bằng' KHÔNG ảnh hưởng đáng kể thời gian với "
                "batch nhỏ (mức nghỉ nhỏ so với thời gian OCR thật), kết quả xử lý GIỐNG HỆT tuyệt "
                "đối trước/sau (2 dòng OK, 5 lỗi - không đổi). Bổ sung kịch bản test 19 mới. Xuất "
                "BAO_CAO_TOI_UU_RELEASE_CANDIDATE.md nêu rõ phạm vi đã làm (GĐ1-2/8) và CÒN THIẾU "
                "trung thực (Task Manager thống nhất, kiểm thử tải lớn, build EXE/Installer Windows "
                "- không khả thi trên môi trường phát triển Linux hiện tại). 19/19 bộ test tự động "
                "PASS ở cả 2 môi trường, không hồi quy."),
    ("13.76.0", "TÍNH NĂNG LỚN: đăng ký tài khoản + kích hoạt bản quyền NGOẠI TUYẾN - đúng Phần II "
                "tài liệu 'LỆNH BỔ SUNG PHẦN CÀI ĐẶT VÀ QUẢN LÝ BẢN QUYỀN'. Vì không thể tự vận "
                "hành máy chủ Internet thật (đã trao đổi rõ với người dùng), triển khai theo mô "
                "hình KÍCH HOẠT NGOẠI TUYẾN bằng CHỮ KÝ SỐ Ed25519 (không chỉ so sánh chuỗi key "
                "đơn giản - đúng yêu cầu mục XV): người dùng tạo 'mã yêu cầu kích hoạt' (không "
                "chứa gì bí mật) → gửi cho Nguyễn Sỹ qua Zalo/điện thoại → Nguyễn Sỹ dùng công "
                "cụ riêng 'SỸ LAND License Manager' (license_manager.py, KHÔNG đóng gói gửi người "
                "dùng) để tạo 'mã phản hồi' đã KÝ SỐ bằng khóa BÍ MẬT chỉ Nguyễn Sỹ giữ → người "
                "dùng dán mã phản hồi, phần mềm XÁC MINH chữ ký bằng khóa CÔNG KHAI đã nhúng sẵn "
                "(không bí mật, không thể dùng để tạo chữ ký giả). Module mới app/services/"
                "license_service.py: hash mật khẩu bằng Scrypt (Argon2/bcrypt không có sẵn trong "
                "môi trường, Scrypt cùng họ thuật toán memory-hard), mã thiết bị dạng hash rút gọn "
                "(không lộ phần cứng thật), token lưu cục bộ MÃ HÓA (Fernet, khóa suy từ máy hiện "
                "tại - copy sang máy khác sẽ KHÔNG đọc được). Thêm banner cảnh báo + KHÓA/MỞ các "
                "Tab nghiệp vụ chính theo trạng thái bản quyền (Tab Cài đặt LUÔN dùng được để đăng "
                "ký/kích hoạt/kiểm tra môi trường/liên hệ hỗ trợ - đúng yêu cầu 'không được làm "
                "mất khả năng mở phần mềm để kích hoạt'). Tạo thêm công cụ độc lập license_manager."
                "py cho Nguyễn Sỹ: tạo cặp khóa quản trị (1 lần duy nhất), đọc mã yêu cầu, chọn "
                "gói/hạn/số thiết bị, tạo mã phản hồi, tự động ghi lại danh sách key đã cấp. Đã "
                "test BẢO MẬT kỹ lưỡng: mã phản hồi bị sửa đổi (giả mạo) BỊ TỪ CHỐI; dùng SAI khóa "
                "công khai (không tương ứng) BỊ TỪ CHỐI; test END-TO-END dùng THẬT cả License "
                "Manager + phần mềm chính xác nhận hoạt động đúng. Thêm cryptography vào "
                "requirements.txt và danh sách thư viện kiểm tra (env_setup_service). 18/18 bộ "
                "test tự động PASS ở cả 2 môi trường."),
    ("13.75.0", "Tiếp tục hoàn thiện tài liệu 'TỐI ƯU HIỆU NĂNG XUẤT WORD' - 2 mục còn lại từ bản "
                "trước: (1) BẢNG THEO DÕI HIỆU NĂNG trên giao diện (mục XV) - nối perf_stats_cb "
                "đã có sẵn hạ tầng vào 1 nhãn mới hiển thị CPU%/RAM/số tiến trình WINWORD.EXE ngay "
                "dưới thanh tiến độ mục 7d, cập nhật realtime trong lúc xuất. (2) KIỂM CHỨNG khả "
                "năng tự phục hồi khi Word gặp sự cố (mục XVII) - test bằng mô phỏng Word 'chết' "
                "LIÊN TỤC qua nhiều file: xác nhận kiến trúc hiện có (mail_merge_smart tự rơi về "
                "phương án dự phòng python-docx mỗi khi COM lỗi) ĐÃ ĐÁP ỨNG đúng tinh thần yêu cầu "
                "(không hủy cả đợt, tiếp tục xử lý) - thậm chí mạnh hơn cách 'khởi động lại Word' "
                "literal vì không cần chờ Word khởi động lại. Bổ sung THÊM: dòng CẢNH BÁO TỔNG KẾT "
                "khi số file phải dùng phương án dự phòng đạt ngưỡng (≥3 file) - dấu hiệu Word có "
                "thể đã gặp sự cố thật, giúp người dùng biết mà không cần đọc từng dòng log riêng "
                "lẻ; đã test xác nhận KHÔNG báo động giả khi chỉ có 1 lỗi đơn lẻ. 17/17 bộ test tự "
                "động PASS ở cả 2 môi trường (bổ sung 2 assertion mới vào kịch bản 17)."),
    ("13.74.0", "TỐI ƯU HIỆU NĂNG XUẤT WORD hàng loạt (mục 7d) - đúng tài liệu 'KHÔNG LÀM TREO "
                "WINDOWS': trước đây xuất hàng loạt (dùng Word COM thật) chạy hết tốc độ liên tục "
                "không nghỉ, có thể khiến Windows Explorer/Excel/PDF không phản hồi, CPU/Disk I/O "
                "lên cao trong lúc xuất nhiều file. Rà soát xác nhận phần LỚN kiến trúc đã ĐÚNG "
                "HƯỚNG từ trước (chỉ 1 tiến trình WINWORD.EXE dùng chung toàn bộ đợt qua "
                "WordCOMSession, đóng Document ngay sau mỗi file, đã có Tạm dừng/Tiếp tục/Hủy) - "
                "PHẦN THIẾU THẬT SỰ là KHÔNG CÓ CHỖ NGHỈ giữa các file. Module mới app/services/"
                "word_perf_service.py: 3 mức tốc độ (Nhanh/Cân bằng/Tiết kiệm) tương ứng mức nghỉ "
                "ngắn sau mỗi file; nghỉ dài hơn sau mỗi đợt (mặc định 100 file/đợt); hạ độ ưu "
                "tiên tiến trình xuống Below Normal trong lúc xuất (khôi phục lại sau); gc.collect() "
                "sau mỗi file; ghi Logs/LOG_HIEU_NANG_WORD.csv (CPU/RAM/thời gian mỗi file, dùng "
                "psutil); cảnh báo mềm khi vượt giới hạn RAM tùy chọn. Thêm mục Cài đặt mới '⚡ "
                "Hiệu năng xuất Word' cho phép người dùng tự chỉnh mọi tham số trên. CHỈ áp dụng "
                "khi ĐANG DÙNG Word COM thật (word_session khác None) và KHÔNG phải chạy thử - "
                "đường cũ (dry_run, không có Word COM) giữ nguyên 100% hành vi, không có độ trễ "
                "nào phát sinh (đã test đo bằng đồng hồ thật xác nhận). Phát hiện và sửa 1 lỗi "
                "thật trong lúc phát triển: write_perf_log_row() lỗi khi đường dẫn log là tên file "
                "trần (không có thư mục cha) - os.makedirs('') ném lỗi - sửa bằng cách chỉ tạo thư "
                "mục khi có phần thư mục thật. Test đo THỜI GIAN THẬT (không giả lập) xác nhận mức "
                "nghỉ đúng theo từng tốc độ đã chọn và đúng số lần nghỉ giữa đợt. 17/17 bộ test tự "
                "động PASS ở cả 2 môi trường (thêm kịch bản test 17 mới)."),
    ("13.73.0", "HOÀN THIỆN VÀ ĐÓNG GÓI Phần I tài liệu 'LỆNH BỔ SUNG PHẦN CÀI ĐẶT VÀ QUẢN LÝ "
                "BẢN QUYỀN': nhóm '🐍 Cài đặt Python và OCR/Tesseract' trong Tab Cài đặt (module "
                "mới app/services/env_setup_service.py) - phát hiện từ phiên làm việc trước bị "
                "NGẮT QUÃNG (đã viết đầy đủ backend + UI + kịch bản test 16 nhưng CHƯA TỪNG được "
                "đóng gói/ghi CHANGELOG). Đã rà soát kỹ, xác nhận hoạt động đúng qua 16/16 bộ test "
                "tự động, bổ sung 1 khoảng trống an toàn quan trọng còn thiếu: nút '🔄 Khôi phục "
                "PATH từ bản sao lưu...' (mục VII tài liệu yêu cầu PATH luôn sao lưu trước khi "
                "sửa - backend restore_path_backup()/list_path_backups() đã có sẵn nhưng chưa nối "
                "vào UI, giờ đã có cửa sổ chọn bản sao lưu + xác nhận trước khi khôi phục). Tính "
                "năng bao gồm: phát hiện Python/Tesseract hiện tại (phiên bản, đường dẫn, trạng "
                "thái pip/PATH), tự động tìm nhiều bản cài đặt, chọn thủ công, dùng bản Tesseract "
                "portable đã đóng gói sẵn, thêm an toàn vào User PATH (sao lưu tự động, không ghi "
                "đè PATH khác), kiểm tra/cài 11 thư viện Python theo đúng danh sách tài liệu, chạy "
                "thử OCR trực tiếp trên ảnh. 16/16 bộ test tự động PASS ở cả 2 môi trường."),
    ("13.72.0", "Bổ sung theo phản hồi người dùng: mục 7c/7d giờ hỗ trợ 2 CỘT ĐỊA CHỈ RIÊNG "
                "BIỆT (thay vì chỉ 1 ô 'Địa chỉ' chung) - file Excel thực tế thường có cả 'Địa "
                "chỉ thường trú' và 'Địa chỉ thửa đất' là 2 dữ liệu khác nhau (đúng file người "
                "dùng gửi ở bản trước: cột L=thường trú, cột X=thửa đất). Thêm ô cấu hình mới "
                "'Địa chỉ thửa đất (tùy chọn)' + token mail-merge {diachithua} mới dùng được trong "
                "mẫu Word, TÁCH RIÊNG khỏi token {location} hiện có (địa chỉ thường trú). Mặc định "
                "trường mới ĐỂ TRỐNG (opt-in an toàn) - không đổi hành vi hiện có cho ai đang dùng "
                "tốt với chỉ 1 địa chỉ. Test xác nhận 2 token tách biệt đúng, dùng chính dữ liệu "
                "người dùng gửi ở lỗi trước. 15/15 bộ test tự động PASS ở cả 2 môi trường."),
    ("13.71.0", "SỬA LỖI THẬT do người dùng báo cáo: mục 7d xuất tên file TBXN sai - vị trí Số "
                "thửa hiện ra ĐỊA CHỈ THỬA ĐẤT thay vì số (VD 'CHUACOGIAY_02146_165_Đồng Lũng "
                "Cậu, Thôn Kim Vân...'). Điều tra bằng chính file Excel người dùng gửi: xác nhận "
                "với cấu hình cột MẶC ĐỊNH cho file đó (Tờ=V, Thửa=W) dữ liệu đọc ĐÚNG hoàn toàn "
                "(to=70, thua=165) - nguyên nhân là CẤU HÌNH CỘT 'Tờ'/'Thửa' trong giao diện Tab "
                "7d bị lệch 1 cột (trỏ nhầm sang cột kế bên, vô tình lấy đúng cột Địa chỉ) - không "
                "phải lỗi thuật toán mà là cấu hình sai (2 ô này KHÔNG được lưu giữa các phiên, "
                "luôn về mặc định V/W khi mở lại phần mềm). ĐÃ THÊM LỚP BẢO VỆ: run_word_batch_"
                "export() giờ kiểm tra 'Số tờ'/'Số thửa' có TRÔNG GIỐNG ĐỊA CHỈ không (có dấu phẩy "
                "hoặc quá dài) - nếu có, CHẶN LẠI với trạng thái CẦN KIỂM TRA kèm ghi chú rõ ràng "
                "chỉ đúng nguyên nhân ('có thể sai cấu hình cột'), thay vì lặng lẽ tạo ra tên file "
                "sai như trước. Test tái hiện ĐÚNG kịch bản lỗi bằng chính dữ liệu người dùng gửi "
                "- xác nhận bắt được lỗi; test dữ liệu bình thường xác nhận KHÔNG báo nhầm (không "
                "có dương tính giả). 15/15 bộ test tự động PASS, thêm 3 assertion mới vào kịch bản 7."),
    ("13.70.0", "Áp dụng CÙNG cách tối ưu Giai đoạn 69 cho đường dự phòng còn lại: đọc file .doc "
                "CŨ khi tên file không đủ khóa Mã xã/Tờ/Thửa (dùng trong đối chiếu 2 thư mục "
                "Word/PDF) - trước đây mỗi file .doc gọi 1 lệnh soffice riêng để chuyển tạm sang "
                ".docx đọc nội dung. Giờ QUÉT TRƯỚC toàn bộ file .doc tên không rõ khóa trong cả 2 "
                "thư mục, GỘP chuyển đổi 1 lần bằng convert_word_to_pdf_libreoffice_batch() (tổng "
                "quát hóa thêm tham số target_format, dùng chung được cho cả pdf/docx), rồi các "
                "lần đọc nội dung sau chỉ cần dùng file đã chuyển sẵn (doc_docx_cache). Test bằng "
                "file .doc thật (tạo qua LibreOffice) xác nhận đúng luồng gộp được kích hoạt, kết "
                "quả đúng. Phát hiện thêm 1 gap kiểm thử: cả batch_word_to_pdf() lẫn compare_"
                "folders_word_pdf() trước đây CHƯA CÓ test tự động riêng dù tồn tại lâu - đã bổ "
                "sung đầy đủ vào kịch bản 15. 15/15 bộ test tự động PASS ở cả 2 môi trường."),
    ("13.69.0", "TĂNG TỐC LỚN mục 7e (Chuyển Word sang PDF): phát hiện mỗi file gọi 1 TIẾN TRÌNH "
                "LibreOffice RIÊNG (subprocess.run) khi không có Word COM - LibreOffice tốn chi phí "
                "khởi động ứng dụng rất lớn, trả giá LẶP LẠI cho MỖI file. Đo thực tế: 3 file riêng "
                "lẻ 6.43s (2.14s/file) → gộp 1 lần gọi 1.83s (0.61s/file, nhanh hơn 3.5 lần); 6 "
                "file riêng lẻ 11.43s (1.9s/file) → gộp 2.55s (0.43s/file, nhanh hơn 4.5 lần) - "
                "CÀNG NHIỀU FILE CÀNG LỢI vì chi phí khởi động chỉ trả 1 LẦN cho cả lô. Thêm hàm "
                "convert_word_to_pdf_libreoffice_batch() gộp nhiều file thành 1 lệnh soffice DUY "
                "NHẤT, viết lại batch_word_to_pdf() nhóm file theo thư mục đích rồi gộp gọi (đường "
                "Word COM sẵn có giữ nguyên tuần tự - đã đủ hiệu quả nhờ session dùng chung). Đo "
                "Tesseract (dùng cho Tab 5/8a/8c) KHÔNG có vấn đề tương tự - overhead khởi động chỉ "
                "~0.12s so với ~8s xử lý thật (~1.5%), không đáng để tối ưu theo hướng này; tốc độ "
                "OCR chậm ở đó là CHI PHÍ THẬT của việc phân tích ảnh, cache đã xử lý phần lặp lại "
                "(Giai đoạn 67), phần xử lý LẦN ĐẦU cho file MỚI vẫn cần thời gian OCR thật, không "
                "rút ngắn được an toàn mà không đánh đổi độ chính xác (dữ liệu pháp lý đất đai). "
                "Test kỹ các trường hợp đặc biệt: bỏ qua file đã tồn tại, chạy thử (dry_run), giữ "
                "cấu trúc thư mục con (nhiều nhóm đích khác nhau) - đều đúng. Bổ sung test hồi quy "
                "đầy đủ cho batch_word_to_pdf() (trước đây CHƯA có test tự động nào cho hàm này). "
                "15/15 bộ test tự động PASS."),
    ("13.68.0", "TĂNG TỐC THÊM mục 8d: cache Excel tổng trong bộ nhớ (read_excel_tong_rows()) - "
                "tránh đọc lại file Excel tổng (thực tế đo được: 6015 dòng mất ~2-4s) khi bấm "
                "'Chạy thử' rồi 'Tổng hợp báo cáo' liền sau trong CÙNG 1 phiên làm việc. Đo thực "
                "tế: 2.15s→0.1s (~21 lần) cho riêng bước đọc, toàn bộ 8d: 2.22s→0.09s (~24 lần) "
                "khi kết hợp với cache OCR đã có từ v13.67. Tự phát hiện và sửa NGAY 1 lỗi cú "
                "pháp tự gây ra khi chèn code (vô tình đóng docstring sớm khiến phần còn lại bị "
                "hiểu nhầm thành code) - py_compile bắt được ngay, sửa trong vài giây. Mặc định "
                "AN TOÀN (use_memory_cache=False giữ nguyên hành vi cũ), test xác nhận cache trả "
                "về đúng data, không cache thì luôn đọc lại từ đĩa. 15/15 bộ test tự động PASS."),
    ("13.67.0", "TĂNG TỐC ĐÁNG KỂ: thêm CACHE OCR vào extract_fields_from_pdf() (hàm lõi dùng ở "
                "TOÀN BỘ Tab 8 - 8a/8c/8d) - phát hiện hàm này dùng ocr_pdf_text() KHÔNG có cache "
                "dù hạ tầng cache OCR đã có sẵn từ trước cho 1 hàm OCR khác. Đo thực tế với dữ liệu "
                "thật: 8a nhanh hơn 267 LẦN khi xử lý lại cùng file (37.36s→0.14s cho 3 file), 8c "
                "nhanh hơn ~1759 lần (17.59s→0.01s), 8d nhanh hơn ~4.6 lần. Hưởng lợi trực tiếp: "
                "nút '🔁 Chạy lại file lỗi' (Giai đoạn 58), và khi chạy 'Chạy thử' rồi 'Tổng hợp "
                "báo cáo' liền sau (2 lần đọc cùng dữ liệu). Mặc định AN TOÀN TUYỆT ĐỐI - tham số "
                "mới ocr_cache_dir=None giữ nguyên 100% hành vi cũ nếu không truyền; đã test xác "
                "nhận kết quả GIỐNG HỆT NHAU dù dùng cache hay không."),
    ("13.66.0", "CHỐT BẢN PHÁT HÀNH ỔN ĐỊNH sau 65 giai đoạn phát triển liên tục. Kiểm tra toàn "
                "diện lần cuối: 14/14 bộ test tự động PASS ở CẢ 2 môi trường (module thật lẫn gói "
                "đóng gói phát hành thật, không chỉ môi trường phát triển); xác nhận đồng bộ TUYỆT "
                "ĐỐI giữa file .py (14 module UI + toàn bộ app/core, app/services) và gói đóng gói "
                "- không lệch 1 dòng nào; xác nhận GUI khởi động ổn định ở cả chế độ Cơ bản/Nâng "
                "cao. Không có thay đổi tính năng - đây là bản ĐÓNG BĂNG (freeze) để phát hành."),
    ("13.65.0", "Tiếp tục mở rộng 'Chế độ Cơ bản/Nâng cao' sang Tab 6: đăng ký TOÀN BỘ khung "
                "'Cải thiện ảnh PDF scan' (chuyển ảnh xám/tăng tương phản/làm nét/khử nhiễu/nhị "
                "phân hóa/tự sửa nghiêng/OCR lại) - tính năng xử lý ảnh chuyên sâu, ít người dùng "
                "phổ thông cần đến, có cảnh báo rủi ro sẵn trong tiêu đề khung ('sẽ làm mất lớp "
                "chữ gốc'). Rà soát Tab 9 (cũ) không tìm thấy tùy chọn kỹ thuật đáng ẩn - các Tab "
                "này vốn đã đơn giản theo thiết kế. Tổng 7 widget/khung nâng cao đã đăng ký."),
    ("13.64.0", "Tiếp tục mở rộng 'Chế độ Cơ bản/Nâng cao' sang Tab 7/8: đăng ký 'Sai số Diện "
                "tích' mục 8d (tham số điều chỉnh kỹ thuật thuần túy), và TOÀN BỘ khung '🏷 Cấu "
                "hình hậu tố tên file' mục 7d (tính năng tùy chọn nâng cao, đa số dùng mặc định "
                "GT/TBXN, đã có checkbox bật/tắt riêng nhưng trước đây LUÔN hiện dù tắt). Giữ "
                "nguyên hiện luôn checkbox 'Xóa dữ liệu CŨ' mục 8d (liên quan an toàn dữ liệu, "
                "không nên ẩn dù là tùy chọn ít dùng). Tổng 6 widget/khung nâng cao đã đăng ký."),
    ("13.63.0", "Tiếp tục mở rộng 'Chế độ Cơ bản/Nâng cao' sang Tab 6 (Tự nhận diện xoay PDF) - "
                "đăng ký hàng chứa 2 ngưỡng tin cậy (tự động xử lý/cần kiểm tra) là nâng cao. Đã "
                "tách riêng khỏi 'Hậu tố tên file' (giữ hiện luôn vì người dùng thường cần chỉnh) "
                "- không ẩn nhầm cả hàng khi chỉ 1 phần thuộc diện nâng cao. Tổng 4 widget nâng "
                "cao đã đăng ký (Tab 5: 2, Tab 6: 1 hàng gộp 2 trường)."),
    ("13.62.0", "Tiếp tục mở rộng 'Chế độ Cơ bản/Nâng cao': đăng ký thêm 2 widget nâng cao ở "
                "Tab 5 (Đổi tên OCR) - 'Lưu ảnh debug + text OCR đầy đủ' và 'Hiện văn bản đã đọc "
                "được khi LỖI'. Tổng 3/nhiều widget đã đăng ký. Phát hiện và tự sửa NGAY 1 lỗi lặp "
                "lại (gọi _main(self).method() thay vì self.method() - đúng lỗi đã gặp ở giai đoạn "
                "trước, nhưng lần này bắt được ngay trong lúc viết code, chưa kịp gây crash)."),
    ("13.61.0", "Bắt đầu tính năng 'Chế độ Cơ bản/Nâng cao' (mục 3 danh sách còn lại) - thêm "
                "che_do_nang_cao vào AppConfig, checkbox trong Tab Cài đặt, cơ chế dùng chung "
                "_register_advanced_widget()/_apply_advanced_mode_visibility() để ẩn/hiện widget "
                "đánh dấu 'nâng cao' TRÊN TOÀN BỘ phần mềm (áp dụng ngay không cần khởi động lại). "
                "Thí điểm trên 1 widget đầu tiên (checkbox 'Xuất text debug' mục 8c) để kiểm chứng "
                "cơ chế hoạt động đúng trước khi mở rộng ra nhiều Tab khác. Mặc định TẮT (Cơ bản) "
                "- không ảnh hưởng gì đến giao diện hiện có của người dùng đang dùng bình thường."),
    ("13.60.0", "THỬ NGHIỆM xử lý song song nhiều file cho Tab 6 (tự động xoay PDF) bằng "
                "ThreadPoolExecutor - PHÁT HIỆN DEADLOCK THẬT khi gọi Tesseract OCR đồng thời từ "
                "nhiều luồng (2 luồng cùng lúc không hoàn thành sau 35s, trong khi 1 luồng chỉ mất "
                "~15s - không phải chậm do GIL mà TREO THẬT). Đã HOÀN TÁC về xử lý tuần tự để đảm "
                "bảo an toàn - xác nhận đúng rủi ro tranh chấp tài nguyên OCR đã cảnh báo từ trước "
                "trong README, không phải lý thuyết suông. Không có thay đổi hành vi cho người "
                "dùng cuối (đã hoàn tác về đúng như trước khi thử nghiệm)."),
    ("13.59.0", "Nối log CHUẨN (12 cột, dùng chung run_id) vào mục 7a (Chuyển đổi Word/PDF hàng "
                "loạt) - mục còn thiếu cuối cùng trong danh sách cũ Giai đoạn 11 (khi rà soát lại "
                "phát hiện 7b/7d/7e đã có log chuẩn từ trước, chỉ 7a và Tab 8 còn thiếu; Tab 8 đã "
                "được thay thế hợp lý bằng log gộp riêng theo Tab từ Giai đoạn 56 nên không cần "
                "thêm log chuẩn nữa). Song song với log CSV riêng đã có (LOG_CHUYEN_DOI_WORD_PDF), "
                "không thay thế."),
    ("13.58.0", "Thêm nút '🔁 Chạy lại file lỗi/cần kiểm tra' cho mục 8c (Điền Excel MẪU từ PDF) "
                "- đúng mục XI tài liệu 'CHỈNH SỬA GIAO DIỆN TAB 7, TAB 8...'. Sau mỗi lần chạy, "
                "tự động lưu kết quả và bật nút nếu có file lỗi/cần kiểm tra; bấm nút sẽ chỉ chạy "
                "lại ĐÚNG các file đó (dùng lại toàn bộ cấu hình hiện tại - template, cột ánh xạ) "
                "thay vì phải chạy lại toàn bộ từ đầu."),
    ("13.57.0", "Thêm MÀU theo trạng thái cho toàn bộ bảng kết quả Tab 7/8 (mục V/VIII tài liệu "
                "'CHỈNH SỬA GIAO DIỆN TAB 7, TAB 8...') - Xanh lá=thành công, Vàng=cần kiểm tra, "
                "Đỏ=lỗi, Xám=bỏ qua. Hàm dùng chung _insert_row_colored() nhận diện trạng thái "
                "theo chuỗi con (khớp được mọi kiểu đặt tên trạng thái khác nhau trong phần mềm), "
                "áp dụng cho cả 4 bảng kết quả hiện có: mục 7d (Xuất Word Excel tổng), 7e (Word→ "
                "PDF), 8c (Điền Excel MẪU), 8d (Đối chiếu Excel tổng - nhận diện thêm GIONG_NHAU/ "
                "KHAC/THIEU_DU_LIEU riêng của mục So_sánh)."),
    ("13.56.0", "Thêm 2 file log GỘP theo Tab đúng mục XII tài liệu 'CHỈNH SỬA GIAO DIỆN TAB 7, "
                "TAB 8...': Logs/LOG_TAB_7_CHUYEN_DOI_XUAT_HO_SO.csv và Logs/LOG_TAB_8_TONG_HOP_"
                "BAO_CAO.csv - tổng hợp hoạt động từ TẤT CẢ 9 mục con (7a-7e, 8a-8d), SONG SONG "
                "với các log CSV riêng theo từng chức năng cụ thể đã có sẵn (không thay thế). "
                "append_tab_log_row() mới, nối vào cả 11 điểm progress_cb đã có (7c có 2 luồng "
                "con). Test xác nhận đúng cột/đúng tên file theo tài liệu."),
    ("13.55.0", "HOÀN THÀNH đầy đủ thanh tiến độ cho TOÀN BỘ 9/9 mục con Tab 7+8 (7a-7e, 8a-8d) - "
                "thêm nốt mục 8b (Thống kê tiến độ - analyze_excel_progress() thêm progress_cb, "
                "test với Excel thật 7000+ dòng) và mục 8d (Đối chiếu Excel tổng - reconcile_excel_"
                "tong_with_pdfs() thêm progress_cb). Đúng tiêu chí hoàn thành mục III/IV tài liệu "
                "'CHỈNH SỬA GIAO DIỆN TAB 7, TAB 8 VÀ BỔ SUNG THANH TIẾN ĐỘ' - mỗi mục xử lý hàng "
                "loạt đều có thanh tiến độ + tên file/dòng đang xử lý + thời gian đã chạy."),
    ("13.54.0", "Tiếp tục bổ sung thanh tiến độ Tab 7: mục 7c (Xuất Word THEO MẪU CÓ SẴN, mail- "
                "merge) - có 2 nguồn dữ liệu (từ PDF/từ Excel) nên thêm 2 thanh tiến độ riêng. "
                "Backend batch_mail_merge_from_pdf_or_excel() thêm tham số progress_cb (nguồn PDF "
                "còn tận dụng progress_cb sẵn có của compile_report_from_pdfs cho bước đọc PDF). "
                "Đã có tiến độ ở 7/9 mục con Tab 7+8 (7a, 7b, 7c, 7d, 7e, 8a, 8c) - còn 8b (Thống "
                "kê)/8d (chưa có Progressbar riêng)."),
    ("13.53.0", "Tiếp tục bổ sung thanh tiến độ Tab 7: mục 7d (Xuất Word hàng loạt — Excel tổng + "
                "mẫu Đơn đăng ký/TBXN) giờ có thanh tiến độ. Backend run_word_batch_export() thêm "
                "tham số progress_cb. Đã có tiến độ ở 6/9 mục con Tab 7+8 (7a, 7b, 7d, 7e, 8a, 8c) "
                "- còn 7c (mail-merge)/8b (Thống kê)/8d (chưa có Progressbar riêng)."),
    ("13.52.0", "Tiếp tục bổ sung thanh tiến độ Tab 7: mục 7b (Gộp file Word hàng loạt) giờ có "
                "thanh tiến độ. Backend batch_merge_word_folder() thêm tham số progress_cb. Phát "
                "hiện và sửa thêm 1 lỗi thiếu import (ttk) trong tab_merge_word.py TRƯỚC KHI xảy "
                "ra sự cố thật - nhờ bộ kiểm tra tĩnh kịch bản 14 đã nâng cấp trước đó. Đã có "
                "tiến độ ở 5/9 mục con Tab 7+8 (7a, 7b, 7e, 8a, 8c)."),
    ("13.51.0", "Tiếp tục bổ sung thanh tiến độ Tab 7: mục 7e (Chuyển Word sang PDF hàng loạt) "
                "giờ có thanh tiến độ (đang xử lý X/Y, tên file, thời gian đã chạy). Backend "
                "batch_word_to_pdf() thêm tham số progress_cb. Đã có tiến độ ở 4/9 mục con Tab "
                "7+8 (7a, 7e, 8a, 8c) - còn 7b/7c/7d và 8b/8d-chi-tiết."),
    ("13.50.0", "Đổi tên hiển thị Tab 7/8 ngắn gọn: '7. Chuyển đổi và xuất hồ sơ', '8. Tổng hợp "
                "và hoàn thiện báo cáo' (bỏ chữ 'Tab' phía trước). Thêm thanh tiến độ cho 2 chức "
                "năng cốt lõi nhất Tab 8: mục 8a (Tổng hợp báo cáo ra Excel) và 8c (Điền Excel "
                "MẪU từ PDF) - hiện đang đọc file X/Y, tên file, thời gian đã chạy. Backend "
                "compile_report_from_pdfs()/fill_excel_template_from_pdfs() thêm tham số "
                "progress_cb. CÒN LẠI: 7b/7c/7d/7e và 8b/8d-chi-tiết vẫn chưa có thanh tiến độ "
                "riêng (quy mô lớn, để làm dần)."),
    ("13.49.0", "Tab 8 mục 8d: thêm giao diện DUYỆT KHO LƯU TRỮ báo cáo (nút '📂 Xem lịch sử báo "
                "cáo đã lưu trữ') - danh sách báo cáo đã xuất, tìm kiếm theo tên/ngày, Mở file/Mở "
                "thư mục/Sao chép đường dẫn. Đồng thời SỬA THÊM 2 LỖI cùng lớp 'thiếu import' "
                "(subprocess chưa import trong tab_report.py và tab_content.py, ảnh hưởng mở file/ "
                "thư mục kết quả). Nâng cấp bộ kiểm tra tĩnh (kịch bản 14) để bắt được cả lỗi dạng "
                "'module.ham()' (trước chỉ bắt 'ham()' trực tiếp) - lỗ hổng khiến 2 lỗi trên sót lại."),
    ("13.48.0", "Tab 6 - Tự nhận diện xoay/làm thẳng PDF: thêm xem trước bằng HÌNH ẢNH trực quan "
                "(2 khung Trước/Sau) - chọn 1 dòng trong bảng xem trước theo trang, tự động hiển "
                "thị ảnh trang gốc và ảnh dự kiến sau khi xoay (áp dụng đúng góc đã ghi đè nếu "
                "có). Hoàn thành mục cuối cùng còn thiếu của tài liệu Tab 6 (before/after preview)."),
    ("13.47.0", "Tab 6 - Tự nhận diện xoay/làm thẳng PDF: bổ sung điều khiển TỪNG TRANG trong "
                "bảng xem trước - nút 'Bỏ qua trang này', 'Xoay thêm +90°', 'Đặt lại về hướng gốc', "
                "'Xóa ghi đè'. Người dùng có thể ghi đè quyết định tự động cho từng trang cụ thể "
                "trước khi xử lý thật, đúng mục VII tài liệu (không còn chỉ có ngưỡng tin cậy tự "
                "động là lựa chọn duy nhất)."),
    ("13.46.0", "SỬA LỖI NGHIÊM TRỌNG: Tab 6 không mở được file PDF khi chọn trong danh sách - "
                "do hàm get_pdf_page_count() bị gọi trực tiếp trong tab_pdfedit.py nhưng CHƯA "
                "IMPORT (NameError bị except nuốt rồi treo ở messagebox.showerror, hiện với "
                "người dùng như 'phần mềm đơ' khi chọn file PDF). Đã sửa + phát hiện thêm 1 lỗi "
                "tương tự (get_base_dir() trong tab_word.py, ảnh hưởng nút 'Dùng mẫu có sẵn'). "
                "Thêm bộ kiểm tra TĨNH bằng AST (kịch bản 14) quét toàn bộ app/ui/*.py để tự "
                "động phát hiện lớp lỗi này trong tương lai, không cần chờ người dùng báo lỗi."),
    ("13.45.0", "Thêm chế độ 'Quét và phân tích' (Xem trước) cho tính năng tự xoay/làm thẳng PDF "
                "- chỉ phân tích, KHÔNG ghi file, hiển thị bảng dự kiến kết quả từng trang (Rotation "
                "gốc/Góc đề xuất/Góc nghiêng/Điểm tin cậy/Nguồn nhận diện/Trạng thái) để người dùng "
                "xem trước rồi mới bấm 'XỬ LÝ TOÀN BỘ' nếu đồng ý - đúng quy trình xác nhận trước "
                "khi ghi thật."),
    ("13.44.0", "Hoàn thiện mục 'Tự nhận diện, xoay và làm thẳng PDF': thêm nút Mở thư mục kết "
                "quả/Mở log, bảng kết quả từng file (tổng trang/đã đúng/xoay theo từng góc/làm "
                "thẳng/cần kiểm tra/trạng thái) cập nhật trực tiếp trong lúc xử lý."),
    ("13.43.0", "Tab 6: thêm mục mới 'Tự nhận diện, xoay và làm thẳng PDF' - tự động phát hiện "
                "trang bị xoay 90/180/270 độ (kết hợp Tesseract OSD + đếm ký tự OCR 4 hướng để "
                "tăng độ tin cậy) và làm thẳng góc nghiêng nhỏ do scan (OpenCV Hough Line Transform "
                "- test chính xác tuyệt đối với góc nghiêng đã biết trước). Xử lý hàng loạt nhiều "
                "file, 2 thanh tiến độ (file + trang), 2 log CSV chi tiết. Test bằng PDF có 4 trang "
                "xoay 4 hướng khác nhau - nhận diện đúng 100%, không sửa file gốc."),
    ("13.42.0", "SỬA LỖI NGHIÊM TRỌNG mục 8d: file mẫu báo cáo .xlsm có macro (VBA) trước đây bị "
                "MẤT MACRO âm thầm khi ghi báo cáo mới (thiếu keep_vba=True) - đã sửa, tự nhận "
                "diện .xlsm và giữ đúng đuôi file + macro. Thêm tùy chọn 'Xóa dữ liệu cũ' an "
                "toàn trước khi ghi - chỉ xóa giá trị ở đúng cột đã ánh xạ, TUYỆT ĐỐI không đụng "
                "công thức/macro/định dạng/merge cell."),
    ("13.41.0", "Hoàn thiện mục 8d: cấu hình sai số Diện tích qua giao diện (trước cố định 1.0 "
                "m²), thêm kho lưu trữ báo cáo nội bộ theo năm/tháng (Output/Bao_cao_da_hoan_thien/) "
                "kèm lịch sử tra cứu (Logs/LICH_SU_BAO_CAO.csv) - có checkbox bật/tắt."),
    ("13.40.0", "Tab 7a (Chuyển đổi PDF sang Word): thêm khả năng chọn NHIỀU FILE cụ thể (không "
                "chỉ thư mục) qua SourcePicker, và thanh tiến độ trực quan (số file/tên file đang "
                "xử lý/thời gian đã chạy/ước tính còn lại) - giống mẫu đã dùng ở Tab 5."),
    ("13.39.0", "Đánh lại số thứ tự Tab liên tục 1-9 (bỏ khoảng trống do gộp Tab 7+9 trước đây), "
                "đổi tên mục cũ '9. Gộp file Word' thành '7b', Tab 8 đổi tên đúng 'Tổng hợp và "
                "hoàn thiện báo cáo'. SỬA LỖI QUAN TRỌNG mục 8d: đối chiếu Excel tổng giờ thử ĐỦ "
                "CÁC BƯỚC dự phòng (chuẩn hóa, đọc nội dung PDF, đối chiếu theo Tờ+Thửa, theo Diện "
                "tích) trước khi kết luận không tìm thấy - không còn báo sai 'không có trong Excel "
                "tổng' khi dữ liệu thực ra CÓ nhưng cần chuẩn hóa. Phát hiện đúng trùng khóa, không "
                "tự chọn. Thêm chế độ nhanh (bỏ qua so sánh chi tiết) - test thực tế 322 file PDF "
                "thật: từ ước tính 30+ phút xuống còn 2.9 giây."),
    ("13.38.0", "Đổi tên Tab '7+9' thành 'Tab 7 - Chuyển đổi và xuất hồ sơ', tiêu đề mục 7a tô "
                "đậm nền nhẹ dễ nhận biết. Bổ sung 3 chế độ chuyển PDF sang Word: Nhanh/Giữ bố "
                "cục (giữ bảng biểu thật qua pdf2docx)/OCR bản quét (PDF scan giờ cho ra văn bản "
                "CÓ THỂ CHỈNH SỬA thay vì chỉ chèn ảnh trang). Có kiểm tra chất lượng sau chuyển "
                "đổi, hậu tố tên file, không ghi đè, log chi tiết 15 cột."),
    ("13.37.0", "Củng cố kiểm tra môi trường khi khởi động: thêm kiểm tra tessdata tiếng Việt, "
                "thư viện Excel/PDF/nén file (hiển thị rõ trong Tab Quy trình xử lý). Chuẩn bị "
                "sẵn hỗ trợ giải nén .rar qua 7-Zip portable (đặt tools/7zip/7z.exe khi đóng gói "
                "trên Windows) - mục 8d giờ có nút chọn cả .zip và .rar."),
    ("13.36.0", "Mục 8d: hỗ trợ file nén .zip cho nguồn PDF (giải nén tự động vào Temp, không "
                "sửa file nén gốc, file lỗi không làm treo phần mềm). Bổ sung đủ 4/4 log CSV theo "
                "tài liệu (log ghi báo cáo, danh sách cần kiểm tra)."),
    ("13.35.0", "Thêm mục 8d Tab Tổng hợp báo cáo: đối chiếu Excel TỔNG đã chỉnh sửa với PDF "
                "theo Mã xã+Số tờ+Số thửa (ưu tiên tách khóa từ TÊN FILE, nhanh, không OCR tràn "
                "lan) - dữ liệu ghi báo cáo lấy từ Excel tổng, kèm sheet So_sanh phát hiện trường "
                "nào đã được chỉnh sửa so với PDF gốc. Tự động áp dụng chuẩn hóa số tờ lâm nghiệp "
                "thống nhất cả 2 phía."),
    ("13.34.0", "Hoàn thiện nốt tổng hợp báo cáo PDF ra Excel: log CSV mở rộng đủ 27 cột (Xã "
                "cũ/Thôn/Số tờ gốc-chuẩn/Loại đất/Nguồn dữ liệu/Điểm tin cậy...), bảng xem trước "
                "Tab 8c mở rộng 14 cột, ngưỡng điểm tin cậy tự động xử lý/cần kiểm tra tùy chỉnh "
                "được qua Tab Cài đặt (mặc định 90%/70% như trước)."),
    ("13.33.0", "Bổ sung KIỂM TRA CHÉO dữ liệu tên file với nội dung PDF khi tổng hợp báo cáo "
                "(Tab 8c): nếu Mã xã/Số tờ/Số thửa mâu thuẫn thật sự, KHÔNG ghi tự động vào Excel "
                "- đưa vào trạng thái DU_LIEU_MAU_THUAN để tự kiểm tra. Hiểu đúng trường hợp tên "
                "file đã có Số tờ chuẩn lâm nghiệp (110000) khớp với nội dung PDF ghi số hiệu gốc "
                "(1) - không báo lỗi giả."),
    ("13.32.0", "Làm sâu tổng hợp báo cáo PDF ra Excel: mở rộng nhận diện theo nhãn (Người sử "
                "dụng đất/Chủ sử dụng đất/Hộ ông/Hộ bà, Căn cước công dân/CMND/Số định danh, Tờ "
                "số/Thửa số...). Thêm xuất text debug riêng từng file (Tab 8c) để chẩn đoán khi "
                "tổng hợp Excel sai."),
    ("13.31.0", "Tự động chuẩn hóa Số tờ bản đồ LÂM NGHIỆP (đất RSX/RPH/RDD): nếu PDF ghi tờ "
                "1/2/3 thì tự chuyển sang số tờ chuẩn 110000/210000/310000 khi đặt tên file và "
                "ghi Excel - tránh trùng với bản đồ địa chính nông nghiệp. Có cấu hình tùy chỉnh "
                "qua sheet Quy_tac_so_to, không hard-code."),
    ("13.30.0", "Hoàn thiện bộ phát hành: VERSION.txt, CHANGELOG.txt, HUONG_DAN_SU_DUNG_NHANH.pdf "
                "(4 trang, có hình minh họa bảng lỗi thường gặp), PHIEU_KIEM_TRA_NGHIEM_THU_SY_LAND.xlsx "
                "(5 sheet, tiền điền kết quả kiểm thử tự động)."),
    ("13.29.0", "Thêm thanh tiến độ trực quan cho Tab 5 (OCR đổi tên) - hiện số file/tên file "
                "đang xử lý/thời gian đã chạy/ước tính còn lại. SỬA LỖI: nút OCR báo lỗi "
                "NameError do thiếu kiểm tra thư viện đúng cách sau khi tách module - đã sửa."),
    ("13.28.0", "Chốt bản v1.0 Stable: khóa nhãn phát hành, VERSION.txt, mở rộng test_data/, "
                "báo cáo nghiệm thu 4 sheet. SỬA LỖI: 2 nhóm Cache OCR/Kiểm thử hệ thống bị lồng "
                "nhầm, có thể biến mất trong 1 số kịch bản đóng gói thiếu module - đã sửa."),
    ("13.27.0", "Thêm nút '🧪 Chạy kiểm thử hệ thống' trong Tab Cài đặt - tự kiểm tra 10 luồng "
                "chính bằng dữ liệu mẫu, xuất báo cáo Excel, không cần dùng dòng lệnh."),
    ("13.26.0", "Thêm CACHE OCR - OCR lại cùng 1 file không đổi (VD chạy thử rồi chạy thật) giờ "
                "dùng lại kết quả cũ thay vì OCR lại, nhanh hơn rất nhiều. Có nút Xóa cache OCR "
                "trong Tab Cài đặt."),
    ("13.25.0", "HOÀN THÀNH nối log chuẩn dùng chung (Logs/) cho TẤT CẢ các Tab: Tách/Gộp PDF, "
                "Xuất Word hàng loạt, Chuyển Word sang PDF, Gộp file Word - không thay log đặc "
                "thù hiện có của từng Tab, chỉ bổ sung song song."),
    ("13.24.0", "HOÀN THÀNH tài liệu cấu hình địa bàn/số tờ: thêm sheet Quy_tac_so_to (đất lâm "
                "nghiệp RSX/RPH/RDD) + Lich_su_sap_nhap vào file cấu hình mẫu; thêm hàm ghi log "
                "17 cột chi tiết đầy đủ."),
    ("13.23.0", "Thêm hàm dùng chung parse_land_key_from_filename() - trả đầy đủ Mã xã/Số tờ/Số "
                "thửa/Hậu tố/Khóa đối chiếu/Trạng thái/Ghi chú từ tên file, dùng thống nhất cho "
                "nhiều module (chuẩn bị nền tảng cho các cải tiến sau)."),
    ("13.22.0", "SỬA LỖI NGHIÊM TRỌNG: chuẩn hóa Mã xã 4 chữ số (VD 2146 -> 02146); Số tờ đọc "
                "từ NỘI DUNG PDF (không chỉ tên file) không còn bị cắt cụt với bản đồ lâm nghiệp "
                "6 chữ số; số đọc từ Excel dạng 210000.0 được xử lý đúng, không làm sai dữ liệu."),
    ("13.20.0", "SỬA LỖI: Tab 10 mục A (Lọc file trùng) không nhận diện được Mã xã/Tờ/Thửa với "
                "file có Số tờ dài 5-6 chữ số (VD Số tờ 210000) - giờ nhận diện đúng."),
    ("13.19.0", "HOÀN THÀNH việc tách toàn bộ giao diện thành 9 file riêng biệt (app/ui/*.py) - "
                "không thay đổi gì về cách dùng, giúp mã nguồn dễ bảo trì hơn lâu dài."),
    ("13.18.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (Tab Tổng hợp báo cáo) - không "
                "thay đổi gì về cách dùng."),
    ("13.17.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (mục Chỉnh sửa PDF) - không thay "
                "đổi gì về cách dùng."),
    ("13.16.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (Tab Đổi tên theo nội dung PDF) - "
                "không thay đổi gì về cách dùng."),
    ("13.15.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (Tab Quy trình xử lý hồ sơ) - "
                "không thay đổi gì về cách dùng."),
    ("13.14.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (mục Lọc Excel & So sánh thư mục "
                "Word/PDF) - không thay đổi gì về cách dùng."),
    ("13.13.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (Tab Đối chiếu & đổi tên theo "
                "Excel) - không thay đổi gì về cách dùng."),
    ("13.12.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (mục Tách/Gộp PDF) - không thay "
                "đổi gì về cách dùng."),
    ("13.11.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (mục Gộp file Word) - không thay "
                "đổi gì về cách dùng."),
    ("13.10.0", "Tiếp tục tách giao diện thành nhiều file nhỏ (mục Lọc file trùng) - không thay "
                "đổi gì về cách dùng."),
    ("13.9.0", "Thêm 'Hướng dẫn sử dụng nhanh' và 'Có gì mới' ngay trong Tab Cài đặt - xem lịch "
               "sử cập nhật mà không cần mở file riêng."),
    ("13.8.0", "Tesseract OCR portable đóng gói kèm phần mềm — không cần tự cài, không cần tự "
               "thêm PATH thủ công. Tự động phát hiện và dùng bản kèm theo nếu có."),
    ("13.7.0", "Bắt đầu tách giao diện thành nhiều file nhỏ (app/ui/) để dễ bảo trì lâu dài - "
               "không thay đổi gì về cách dùng, chỉ ở phía kỹ thuật bên trong."),
    ("13.6.0", "Chuẩn hóa thêm 1 lớp log dùng chung cho Tab OCR đổi tên, Chỉnh sửa PDF, Đối chiếu "
               "báo cáo - tập trung tại 1 thư mục Logs, không thay log quen thuộc của từng Tab."),
    ("13.5.0", "Cài đặt 'Tự động sao lưu' và 'Bắt buộc chạy thử' giờ có tác dụng thật (trước đây "
               "chỉ lưu mà chưa áp dụng)."),
    ("13.4.0", "Xuất Word hàng loạt: hậu tố tên file không còn cố định GT/TBXN - tùy chọn hậu tố "
               "bất kỳ, lấy từ cột Excel, hoặc mẫu tên file tùy chỉnh."),
    ("13.3.0", "Cấu hình địa bàn LINH ĐỘNG: 1 xã hiện tại có thể gồm nhiều xã cũ, nhiều thôn cũ "
               "gộp thành 1 thôn mới - tự phát hiện mã xã đúng hoặc đưa vào cần kiểm tra."),
    ("13.2.0", "Tô đậm toàn bộ tiêu đề nhóm chức năng trên giao diện và trong Excel báo cáo xuất ra."),
    ("13.1.0", "Thêm mục 'Khối lượng và tiến độ thực hiện' - nhập/tự đếm hồ sơ đã làm, biểu đồ "
               "tròn + cột, xuất báo cáo tiến độ riêng."),
    ("13.0.0", "Log tổng hợp báo cáo từ PDF chuẩn hóa đủ 17 cột và các trạng thái rõ ràng hơn."),
    ("12.9.0", "Gộp Tab 'Gộp file Word' và 'Xuất Word' thành 1 Tab; bỏ mục xuất Word từ Excel tự "
               "do (ít dùng), dồn tính năng hữu ích sang mục mail-merge chính."),
    ("12.8.0", "Tab tổng hợp báo cáo: tùy chỉnh được cột Excel muốn ghi (trước đây cố định cứng)."),
    ("12.7.0", "Mẫu Word có thể chèn thêm thông tin thôn/xã cũ-mới khi có cấu hình địa bàn."),
    ("12.6.0", "Cấu hình địa bàn mở rộng - dùng được cho bất kỳ xã/phường/thôn nào, không chỉ "
               "riêng Văn Lang."),
    ("12.5.0", "Thêm Tab '⚙ Cài đặt' - chỉnh chế độ OCR mặc định, tự động backup, bắt buộc chạy thử..."),
    ("12.0.0 - 12.4.0", "Tái cấu trúc nền tảng mã nguồn theo hướng module hóa (không ảnh hưởng "
                        "cách dùng) - làm nền tảng ổn định cho các tính năng mới về sau."),
]

# Bảng tra Tên xã -> Mã xã (có thể chỉnh sửa trực tiếp trên giao diện)
DEFAULT_XA_MAPPING_TEXT = (
    "Văn Lang cũ=02140\n"
    "Lương Thượng=02143\n"
    "Kim Hỷ=02146\n"
)

DEFAULT_WX_TEMPLATE = (
    "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
    "Độc lập - Tự do - Hạnh phúc\n"
    "\n"
    "THÔNG TIN THỬA ĐẤT SỐ {stt}\n"
    "\n"
    "Mã xã: {maxa}\n"
    "Tờ bản đồ số: {to}\n"
    "Thửa đất số: {thua}\n"
    "Tên chủ sử dụng đất: {ten}\n"
    "Địa chỉ: {dia_chi}\n"
    "Diện tích: {dt} m2\n"
)

# Bảng tra Tên thôn -> Mã xã (ưu tiên cao hơn tên xã, vì thôn không đổi khi
# sáp nhập địa giới hành chính - PDF có thể ghi tên xã cũ nhưng tên thôn vẫn giữ nguyên)
DEFAULT_THON_MAPPING_TEXT = (
    "Khuổi Phầy=02146\n"
    "Kim Vân=02146\n"
    "Quốc Tuấn=02146\n"
    "Nà Mỏ=02146\n"
    "Khuổi Hát=02146\n"
    "Nà Làng=02143\n"
    "Bản Giang=02143\n"
    "Pàn Xả=02143\n"
    "Vằng Khít=02143\n"
    "Khuổi Nộc=02143\n"
    "Bản Kén=02140\n"
    "Nặm Cà=02140\n"
    "Chợ Mới=02140\n"
    "Tân An=02140\n"
    "Nà Diệc=02140\n"
    "Bản Sảng=02140\n"
    "Nà Lẹng=02140\n"
    "Cốc Phia=02140\n"
    "Nà Dường=02140\n"
)


def get_base_dir():
    """Trả về thư mục chứa tài nguyên (assets) - đúng cả khi chạy .py lẫn khi đã đóng gói .exe (PyInstaller)."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


# Tự động dùng Tesseract OCR portable kèm theo phần mềm (nếu có, xem tools/tesseract/) - phải gọi
# SAU khi get_base_dir() đã được định nghĩa ở trên. Nếu không có bản portable kèm theo, giữ
# nguyên hành vi cũ (dùng Tesseract từ PATH hệ thống).
HAS_BUNDLED_TESSERACT = _configure_bundled_tesseract() if HAS_OCR else False


def get_app_dir():
    """
    Trả về thư mục chứa file .exe (khi đã đóng gói) hoặc file .py (khi chạy trực tiếp) ĐANG CHẠY.
    Khác với get_base_dir(): khi đóng gói .exe, get_base_dir() trỏ vào thư mục tạm giải nén (mất khi
    thoát chương trình), còn get_app_dir() trỏ đúng vào nơi đặt file .exe/.py - dùng để lưu cấu hình
    cập nhật (update_config.json) và để tự ghi đè khi cập nhật.
    """
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))


def get_app_data_dir():
    """
    Trả về thư mục dữ liệu làm việc của phần mềm (Output/Logs/Temp/Backup/Undo/Can_kiem_tra),
    ĐẶT TRONG %LocalAppData%\\SY LAND trên Windows (không ghi vào Program Files - tránh lỗi
    quyền ghi file). Trên Linux/Mac (môi trường phát triển/thử nghiệm) dùng thư mục home tương ứng.
    """
    if sys.platform == "win32":
        base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
        return os.path.join(base, "SY LAND")
    return os.path.join(os.path.expanduser("~"), ".sy_land")


def append_tab_log_row(tab_number, chuc_nang, ten_file, buoc="", tien_do="", trang_thai="",
                        ket_qua="", loi="", ghi_chu=""):
    """
    Ghi 1 dòng vào log GỘP theo Tab - đúng mục XII tài liệu "CHỈNH SỬA GIAO DIỆN TAB 7, TAB 8 VÀ
    BỔ SUNG THANH TIẾN ĐỘ": mỗi Tab 7/8 có 1 file log DUY NHẤT tổng hợp hoạt động từ TẤT CẢ các
    mục con (7a-7e hoặc 8a-8d), khác với các log CSV riêng theo TỪNG chức năng cụ thể đã có sẵn
    trước đó (VD LOG_XUAT_WORD.csv) - 2 loại log này SONG SONG tồn tại, không thay thế nhau.

    tab_number: 7 hoặc 8 (số Tab).
    chuc_nang: tên mục con đang chạy (VD "7a. Chuyển đổi Word/PDF", "8c. Điền Excel MẪU từ PDF").
    ten_file: tên file/dòng/bản ghi đang xử lý tại thời điểm ghi log.
    Không raise lỗi nếu ghi log thất bại (best-effort) - không được làm gián đoạn luồng chính chỉ
    vì việc ghi log phụ trợ này gặp sự cố (VD ổ đĩa đầy, quyền ghi...).
    """
    try:
        log_dir = os.path.join(get_app_data_dir(), "Logs")
        os.makedirs(log_dir, exist_ok=True)
        fname = ("LOG_TAB_7_CHUYEN_DOI_XUAT_HO_SO.csv" if str(tab_number) == "7"
                 else "LOG_TAB_8_TONG_HOP_BAO_CAO.csv")
        log_path = os.path.join(log_dir, fname)
        is_new = not os.path.isfile(log_path)
        import csv
        with open(log_path, "a", encoding="utf-8-sig", newline="") as f:
            w = csv.writer(f)
            if is_new:
                w.writerow(["Thời gian", "Chức năng", "Tên file hoặc bản ghi", "Bước xử lý",
                           "Tiến độ", "Trạng thái", "Kết quả", "Lỗi", "Ghi chú"])
            w.writerow([datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), chuc_nang, ten_file,
                       buoc, tien_do, trang_thai, ket_qua, loi, ghi_chu])
    except Exception:
        pass  # ghi log phu tro - khong lam gian doan luong chinh neu that bai


def ensure_standard_app_folders(log_cb=None):
    """
    Tự tạo (nếu chưa có) các thư mục làm việc chuẩn: Output, Logs, Temp, Backup, Undo, Can_kiem_tra,
    nằm trong get_app_data_dir(). KHÔNG ảnh hưởng tới việc người dùng tự chọn thư mục riêng ở từng
    Tab - đây chỉ là thư mục MẶC ĐỊNH/dự phòng để phần mềm luôn có chỗ ghi log/backup an toàn.
    Trả về dict {tên_thư_mục: đường_dẫn_đầy_đủ}.
    """
    base = get_app_data_dir()
    folder_names = ["Output", "Logs", "Temp", "Backup", "Undo", "Can_kiem_tra"]
    paths = {}
    for name in folder_names:
        p = os.path.join(base, name)
        try:
            os.makedirs(p, exist_ok=True)
            paths[name] = p
        except Exception as e:
            if log_cb:
                log_cb(f"⚠ Không tạo được thư mục {name}: {e}")
            paths[name] = None
    return paths


def check_tesseract_available():
    """Kiểm tra Tesseract OCR có sẵn trên máy không (qua pytesseract), không raise lỗi."""
    if not HAS_OCR:
        return False
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def check_tessdata_vie():
    """
    Kiểm tra riêng gói ngôn ngữ tiếng Việt (vie.traineddata) của Tesseract - mục IV.7 tài liệu
    "BỔ SUNG TÍNH NĂNG MỚI + SỬA LỖI ĐÓNG GÓI". Có Tesseract chạy được KHÔNG đồng nghĩa có đủ
    tessdata tiếng Việt - đây là lỗi hay gặp khi Tesseract cài thiếu ngôn ngữ.
    """
    if not HAS_OCR:
        return False
    try:
        tessdata_dir = os.environ.get("TESSDATA_PREFIX")
        if tessdata_dir and os.path.isfile(os.path.join(tessdata_dir, "vie.traineddata")):
            return True
        # Thử qua get_languages() của pytesseract (đọc theo cấu hình Tesseract đang dùng)
        langs = pytesseract.get_languages(config="")
        return "vie" in langs
    except Exception:
        return False


def check_word_available():
    """Kiểm tra Microsoft Word có khả năng dùng được qua COM không (chỉ đúng nghĩa trên Windows)."""
    return bool(HAS_WIN32COM)


def check_excel_available():
    """Kiểm tra Microsoft Excel có khả năng dùng được qua COM không (chỉ đúng nghĩa trên Windows).
    Lưu ý: dùng chung cờ HAS_WIN32COM vì việc phát hiện CHÍNH XÁC từng ứng dụng Office đã cài
    riêng lẻ cần thử Dispatch thật (tốn thời gian mở tiến trình) - ở đây chỉ kiểm tra ĐIỀU KIỆN
    CẦN (có pywin32 + Windows), việc thực sự mở Excel COM vẫn có thể lỗi riêng nếu chưa cài Excel."""
    return bool(HAS_WIN32COM)


def check_openpyxl_available():
    """Kiểm tra thư viện đọc/ghi Excel (openpyxl) - mục IV.7 tài liệu. Đây là thư viện Python
    thuần (không cần cài Excel thật), dùng cho MỌI thao tác đọc/ghi Excel trong phần mềm."""
    try:
        import openpyxl  # noqa: F401
        return True
    except ImportError:
        return False


def check_pdf_library_available():
    """Kiểm tra thư viện đọc PDF (pypdf và/hoặc PyMuPDF/fitz) - mục IV.7 tài liệu."""
    return bool(HAS_PYPDF or HAS_FITZ)


def check_zip_support_available():
    """Kiểm tra khả năng giải nén .zip (thư viện chuẩn Python zipfile - luôn có sẵn)."""
    try:
        import zipfile  # noqa: F401
        return True
    except ImportError:
        return False


def check_rar_support_available():
    """
    Kiểm tra khả năng giải nén .rar - mục IV.8 tài liệu muốn đóng gói kèm 7z.exe/unrar.exe
    portable. HIỆN CHƯA đóng gói kèm công cụ này (cần môi trường Windows để chuẩn bị/kiểm tra) -
    hàm này luôn trả về False cho đến khi bổ sung, để giao diện hiển thị đúng trạng thái thay vì
    im lặng coi như đã hỗ trợ.
    """
    bundled_7z = os.path.join(get_base_dir(), "tools", "7zip", "7z.exe")
    return os.path.isfile(bundled_7z)


def check_environment(log_cb=None):
    """
    Kiểm tra môi trường máy tính khi khởi động phần mềm: Microsoft Word, Microsoft Excel,
    Tesseract OCR, các thư mục làm việc chuẩn, file cấu hình mã xã/thôn, file logo/icon.
    KHÔNG BAO GIỜ raise lỗi ra ngoài - trả về dict trạng thái để giao diện tự quyết định
    hiển thị cảnh báo phù hợp, không làm phần mềm dừng lại hay mở không được.
    """
    result = {}
    try:
        result["word"] = check_word_available()
    except Exception:
        result["word"] = False
    try:
        result["excel"] = check_excel_available()
    except Exception:
        result["excel"] = False
    try:
        result["tesseract"] = check_tesseract_available()
        result["tesseract_bundled"] = HAS_BUNDLED_TESSERACT
    except Exception:
        result["tesseract"] = False
        result["tesseract_bundled"] = False
    try:
        result["tessdata_vie"] = check_tessdata_vie()
    except Exception:
        result["tessdata_vie"] = False
    try:
        result["openpyxl"] = check_openpyxl_available()
    except Exception:
        result["openpyxl"] = False
    try:
        result["pdf_library"] = check_pdf_library_available()
    except Exception:
        result["pdf_library"] = False
    try:
        result["zip_support"] = check_zip_support_available()
    except Exception:
        result["zip_support"] = False
    try:
        result["rar_support"] = check_rar_support_available()
    except Exception:
        result["rar_support"] = False
    try:
        result["folders"] = ensure_standard_app_folders(log_cb=log_cb)
    except Exception:
        result["folders"] = {}
    try:
        result["config_xa_thon"] = os.path.isfile(get_xa_thon_config_path())
    except Exception:
        result["config_xa_thon"] = False
    try:
        result["logo"] = os.path.isfile(os.path.join(get_base_dir(), "assets", "app_icon.ico"))
    except Exception:
        result["logo"] = False

    if log_cb:
        log_cb("Kiểm tra môi trường:")
        log_cb(f"  Microsoft Word (COM): {'Đã có' if result['word'] else 'Chưa có'}")
        log_cb(f"  Microsoft Excel (COM): {'Đã có' if result['excel'] else 'Chưa có'}")
        log_cb(f"  Tesseract OCR: {'Đã có' if result['tesseract'] else 'Chưa có'}"
              + (" (dùng bản portable kèm phần mềm)" if result.get("tesseract_bundled") else ""))
        log_cb(f"  Tessdata tiếng Việt: {'Đã có' if result['tessdata_vie'] else 'Chưa có/thiếu'}")
        log_cb(f"  Thư viện đọc/ghi Excel (openpyxl): {'Đã có' if result['openpyxl'] else 'Chưa có'}")
        log_cb(f"  Thư viện đọc PDF: {'Đã có' if result['pdf_library'] else 'Chưa có'}")
        log_cb(f"  Hỗ trợ giải nén .zip: {'Đã có' if result['zip_support'] else 'Chưa có'}")
        log_cb(f"  Hỗ trợ giải nén .rar: {'Đã có' if result['rar_support'] else 'Chưa có (chưa đóng gói kèm 7z.exe)'}")
        log_cb(f"  Cấu hình mã xã/thôn: {'Đã có' if result['config_xa_thon'] else 'Chưa có (dùng mặc định)'}")
        log_cb(f"  Logo/icon: {'Đã có' if result['logo'] else 'Chưa có'}")

    return result


UPDATE_CHECK_CONFIG_FILENAME = "update_check_config.json"
LAST_UPDATE_CHECK_FILENAME = ".last_update_check.json"
UPDATE_CHECK_TIMEOUT_SECONDS = 5


def get_update_check_url():
    """
    Đọc URL kiểm tra phiên bản mới (trỏ tới file latest_version.json online) từ
    update_check_config.json đặt cạnh phần mềm. Trả về None nếu chưa cấu hình.
    """
    path = os.path.join(get_app_dir(), UPDATE_CHECK_CONFIG_FILENAME)
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        url = (data.get("check_url") or "").strip()
        return url or None
    except Exception:
        return None


def parse_version_tuple(v):
    parts = []
    for p in re.split(r"[.\-]", str(v).strip()):
        try:
            parts.append(int(p))
        except ValueError:
            parts.append(0)
    return tuple(parts)


def should_check_update_today():
    """Trả về True nếu HÔM NAY chưa kiểm tra cập nhật lần nào (giới hạn tối đa 1 lần/ngày lúc khởi động)."""
    path = os.path.join(get_app_dir(), LAST_UPDATE_CHECK_FILENAME)
    today = datetime.date.today().isoformat()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("last_check_date") != today
    except Exception:
        return True


def mark_update_checked_today():
    path = os.path.join(get_app_dir(), LAST_UPDATE_CHECK_FILENAME)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"last_check_date": datetime.date.today().isoformat()}, f)
    except Exception:
        pass  # không quan trọng, bỏ qua im lặng - không ảnh hưởng phần mềm chính


def fetch_latest_version_info(check_url, timeout=UPDATE_CHECK_TIMEOUT_SECONDS):
    """
    Tải file latest_version.json từ check_url (có timeout, mặc định 5 giây).
    Trả về dict {app_name, latest_version, release_date, download_url, release_notes, force_update}
    hoặc None nếu lỗi/mất mạng - LUÔN im lặng bỏ qua, không được làm phần mềm treo hay lỗi.
    """
    try:
        with urllib.request.urlopen(check_url, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except Exception:
        return None


# ------------------------- TIỆN ÍCH DÙNG CHUNG -------------------------

def strip_diacritics(s):
    if s is None:
        return ""
    s = str(s).replace("đ", "d").replace("Đ", "D")
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


def name_tokens(s):
    return set(re.findall(r"[A-Z0-9]+", strip_diacritics(s).upper()))


def norm_num(s):
    """Chuẩn hóa số: bỏ số 0 thừa ở đầu (VD '013' -> '13'), nhưng giữ nguyên nếu toàn số 0.
    Nếu đọc từ Excel ra dạng số thực (VD '210000.0' do Excel tự hiểu ô là number), bỏ đuôi '.0'
    trước - KHÔNG được làm tròn/cắt bớt chữ số có nghĩa, chỉ bỏ đúng phần thập phân .0 thừa."""
    s = str(s).strip()
    if re.fullmatch(r"\d+\.0+", s):
        s = s.split(".")[0]
    s = s.lstrip("0")
    return s if s != "" else "0"


def extract_digit_groups(filename):
    name = os.path.splitext(filename)[0]
    return re.findall(r"\d+", name)


def extract_numeric_candidates(filename):
    name = os.path.splitext(filename)[0]
    decimals = re.findall(r"\d+[.,]\d+", name)
    vals = [float(d.replace(",", ".")) for d in decimals]
    vals += [float(d) for d in re.findall(r"\d+", name)]
    return vals


def detect_loai(fname):
    up = fname.upper()
    if "TBXN" in up:
        return "TBXN"
    if "GT" in up:
        return "GT"
    return "KHAC"


INVALID_FILENAME_CHARS = r'\/:*?"<>|'


def sanitize_filename_part(s):
    s = str(s) if s is not None else ""
    for ch in INVALID_FILENAME_CHARS:
        s = s.replace(ch, "")
    return s.strip()


# ------------------------- CẤU HÌNH HẬU TỐ TÊN FILE TÙY CHỌN -------------------------

SUFFIX_PRESETS = ["GT", "DDK", "TBXN", "BBKT", "CV", "KH", "DON", "PHIEU", "BIENBAN",
                  "XACMINH", "BOSUNG", "KHAC"]

DEFAULT_FILENAME_TEMPLATE = "CHUACOGIAY_{maxa}_{soto}_{sothua}_{hautofile}"

FILENAME_TEMPLATE_VARS = ["maxa", "soto", "sothua", "hoten", "cccd", "thon", "xacu", "hautofile", "ngay", "stt"]


def normalize_suffix(text, uppercase=False, remove_diacritics_opt=False):
    """
    Chuẩn hóa hậu tố tên file trước khi ghép vào tên file:
    1. Bỏ khoảng trắng thừa ở 2 đầu.
    2. Chuyển khoảng trắng giữa các từ thành dấu gạch dưới.
    3. Loại bỏ ký tự không hợp lệ trong tên file Windows (\\ / : * ? " < > |).
    4. Tùy chọn viết hoa toàn bộ.
    5. Tùy chọn bỏ dấu tiếng Việt.
    """
    if not text:
        return ""
    s = str(text).strip()
    if remove_diacritics_opt:
        s = strip_diacritics(s)
    s = re.sub(r"\s+", "_", s)
    for ch in INVALID_FILENAME_CHARS:
        s = s.replace(ch, "")
    if uppercase:
        s = s.upper()
    return s.strip("_")


def resolve_file_suffix(default_suffix, manual_suffix="", combobox_suffix="", excel_value="",
                         use_excel_col=False, prioritize_manual=False,
                         uppercase=False, remove_diacritics_opt=False):
    """
    Xác định hậu tố tên file THEO ĐÚNG THỨ TỰ ƯU TIÊN:
    1. Cột Excel (nếu bật + dòng có dữ liệu).
    2. Hậu tố nhập tay (nếu bật 'Ưu tiên hậu tố nhập tay' + có nhập).
    3. Hậu tố chọn ở combobox.
    4. Hậu tố nhập tay (nếu có nhập nhưng KHÔNG bật ưu tiên - vẫn dùng nếu combobox trống).
    5. Hậu tố mặc định (default_suffix - VD "GT"/"TBXN" theo mẫu đang xuất) nếu không có gì khác.
    Trả về (hau_to_da_chuan_hoa, nguon) - nguon in {TU_COT_EXCEL, TU_NHAP_TAY, TU_COMBOBOX, MAC_DINH, KHONG_CO_HAU_TO}.
    Tuyệt đối KHÔNG ép cứng về "GT"/"TBXN" nếu người dùng đã chọn nguồn khác.
    """
    excel_value = str(excel_value or "").strip()
    manual_suffix = str(manual_suffix or "").strip()
    combobox_suffix = str(combobox_suffix or "").strip()

    if use_excel_col and excel_value:
        raw, source = excel_value, "TU_COT_EXCEL"
    elif prioritize_manual and manual_suffix:
        raw, source = manual_suffix, "TU_NHAP_TAY"
    elif combobox_suffix and combobox_suffix != "KHAC":
        raw, source = combobox_suffix, "TU_COMBOBOX"
    elif manual_suffix:
        raw, source = manual_suffix, "TU_NHAP_TAY"
    elif default_suffix:
        raw, source = default_suffix, "MAC_DINH"
    else:
        return "", "KHONG_CO_HAU_TO"

    return normalize_suffix(raw, uppercase=uppercase, remove_diacritics_opt=remove_diacritics_opt), source


def render_output_filename(template_str, maxa="", to="", thua="", ten="", cccd="", thon="", xacu="",
                            hauto="", stt=""):
    """
    Render mẫu tên file tùy chỉnh, hỗ trợ biến: {maxa} {soto} {sothua} {hoten} {cccd} {thon}
    {xacu} {hautofile} {ngay} {stt}. Tự dọn dấu gạch dưới dư khi 1 biến rỗng (VD "..__GT" -> ".._GT").
    """
    mapping = {
        "maxa": maxa or "", "soto": to or "", "sothua": thua or "",
        "hoten": sanitize_filename_part(ten or ""), "cccd": cccd or "",
        "thon": sanitize_filename_part(thon or ""), "xacu": sanitize_filename_part(xacu or ""),
        "hautofile": hauto or "", "ngay": datetime.datetime.now().strftime("%Y%m%d"), "stt": stt or "",
    }
    try:
        name = (template_str or DEFAULT_FILENAME_TEMPLATE).format(**mapping)
    except (KeyError, IndexError) as e:
        raise ValueError(f"Mẫu tên file chứa placeholder không hợp lệ: {{{e}}}")
    name = re.sub(r"_+", "_", name).strip("_ ")
    if not name:
        raise ValueError("Mẫu tên file tạo ra tên rỗng.")
    return sanitize_filename_part(name)


# ------------------------- ĐỌC DỮ LIỆU EXCEL -------------------------

def load_excel_data(excel_path, sheet_name, header_row,
                     col_stt, col_maxa, col_to, col_thua, col_ten, col_files, col_dt):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    rows = []
    for r in range(header_row, ws.max_row + 1):
        maxa = ws[f"{col_maxa}{r}"].value
        to = ws[f"{col_to}{r}"].value
        thua = ws[f"{col_thua}{r}"].value
        if maxa is None or to is None or thua is None:
            continue
        if str(maxa).strip() == "" or str(to).strip() == "" or str(thua).strip() == "":
            continue

        files_raw = ws[f"{col_files}{r}"].value if col_files else None
        expected_files = []
        if files_raw:
            expected_files = [f.strip().lower() for f in str(files_raw).split(",") if f.strip()]

        dt_val = None
        if col_dt:
            raw_dt = ws[f"{col_dt}{r}"].value
            if raw_dt is not None:
                try:
                    dt_val = float(str(raw_dt).replace(",", "."))
                except ValueError:
                    dt_val = None

        ten_val = ws[f"{col_ten}{r}"].value if col_ten else ""

        rows.append({
            "stt": ws[f"{col_stt}{r}"].value if col_stt else r,
            "maxa": str(maxa).strip(),
            "to": str(to).strip(),
            "thua": str(thua).strip(),
            "dt": dt_val,
            "ten": ten_val,
            "ten_tokens": name_tokens(ten_val),
            "expected_files": expected_files,
            "matched": False,
            "matched_files": [],
        })
    return rows


# ------------------------- ĐỐI CHIẾU -------------------------

def find_match(filename, rows, xa_codes_present, criteria, dt_tolerance):
    """
    criteria: subset of {'maxa','to','thua','dt','ten'} - các tiêu chí BẮT BUỘC phải khớp.
    Trả về (row_hoac_list, status) voi status in {'matched','ambiguous','none'}
    """
    lname = filename.lower()

    # Pass 1: khớp trực tiếp với tên file đã ghi sẵn trong Excel (luôn ưu tiên)
    for row in rows:
        if lname in row["expected_files"]:
            return row, "matched"

    digits = extract_digit_groups(filename)
    numeric_needed = bool(criteria & {"maxa", "to", "thua", "dt"})
    if numeric_needed and not digits:
        # Vẫn có thể chỉ cần "ten" nếu nó là tiêu chí duy nhất còn lại; nhưng nếu numeric_needed=True
        # nghĩa là có ít nhất 1 tiêu chí số được yêu cầu -> không có số thì chắc chắn không khớp.
        return None, "none"

    remaining_norm = [norm_num(d) for d in digits]
    numeric_candidates_dt = extract_numeric_candidates(filename)
    file_name_tokens = name_tokens(filename)

    require_maxa = "maxa" in criteria
    require_to = "to" in criteria
    require_thua = "thua" in criteria
    require_dt = "dt" in criteria
    require_ten = "ten" in criteria

    if require_maxa:
        maxa_candidates = {norm_num(d) for d in digits if d in xa_codes_present or norm_num(d) in xa_codes_present}
        if not maxa_candidates:
            return None, "none"

    matches = []
    for row in rows:
        if require_maxa and norm_num(row["maxa"]) not in maxa_candidates:
            continue
        ok = True
        if require_to:
            ok = ok and (norm_num(row["to"]) in remaining_norm)
        if require_thua:
            ok = ok and (norm_num(row["thua"]) in remaining_norm)
        if require_dt:
            if row["dt"] is None:
                ok = False
            else:
                ok = ok and any(abs(row["dt"] - v) <= dt_tolerance for v in numeric_candidates_dt)
        if require_ten:
            ok = ok and bool(row["ten_tokens"]) and row["ten_tokens"].issubset(file_name_tokens)
        if ok:
            matches.append(row)

    if len(matches) == 1:
        return matches[0], "matched"
    elif len(matches) > 1:
        return matches, "ambiguous"
    return None, "none"


def process(folders, excel_path, output_matched, output_unmatched, sheet_name, header_row,
            col_stt, col_maxa, col_to, col_thua, col_ten, col_files, col_dt,
            criteria, dt_tolerance, log_cb):

    rows = load_excel_data(excel_path, sheet_name, header_row,
                            col_stt, col_maxa, col_to, col_thua, col_ten, col_files, col_dt)
    xa_codes_present = set(r["maxa"] for r in rows) | {norm_num(r["maxa"]) for r in rows}

    os.makedirs(output_matched, exist_ok=True)
    os.makedirs(output_unmatched, exist_ok=True)

    matched_count = 0
    unmatched_count = 0
    ambiguous_count = 0
    unmatched_files = []
    seen_matched = set()
    seen_unmatched = set()

    for folder in folders:
        if not folder or not os.path.isdir(folder):
            continue
        for fname in sorted(os.listdir(folder)):
            if not fname.lower().endswith(".pdf"):
                continue
            fpath = os.path.join(folder, fname)
            result, status = find_match(fname, rows, xa_codes_present, criteria, dt_tolerance)

            if status == "matched":
                row = result
                dest_name = fname
                if dest_name in seen_matched:
                    base, ext = os.path.splitext(fname)
                    dest_name = f"{base}__trung_ten{ext}"
                seen_matched.add(dest_name)
                shutil.copy2(fpath, os.path.join(output_matched, dest_name))
                matched_count += 1
                row["matched"] = True
                row["matched_files"].append(fname)
                log_cb(f"✓ KHỚP  | {fname}  →  Mã xã {row['maxa']}, Tờ {row['to']}, Thửa {row['thua']}")
            else:
                dest_name = fname
                if dest_name in seen_unmatched:
                    base, ext = os.path.splitext(fname)
                    dest_name = f"{base}__trung_ten{ext}"
                seen_unmatched.add(dest_name)
                shutil.copy2(fpath, os.path.join(output_unmatched, dest_name))
                unmatched_count += 1
                unmatched_files.append(fname)
                if status == "ambiguous":
                    ambiguous_count += 1
                    log_cb(f"⚠ KHỚP NHIỀU DÒNG (không rõ ràng) | {fname}  →  đã copy sang thư mục KHÔNG khớp để kiểm tra tay")
                else:
                    log_cb(f"✗ Không khớp | {fname}")

    missing_rows = [r for r in rows if not r["matched"]]
    report_path = os.path.join(output_matched, "BAO_CAO_DOI_CHIEU.xlsx")
    export_report(report_path, rows, unmatched_files)

    return matched_count, unmatched_count, ambiguous_count, missing_rows, report_path


def style_excel_header_row(ws, row_num=1, num_cols=None, auto_width=True):
    """
    Tô đậm 1 dòng tiêu đề bảng trong file Excel TỰ TẠO (KHÔNG dùng cho việc ghi vào Excel MẪU có
    sẵn của người dùng - những chỗ đó phải giữ nguyên định dạng gốc, không được gọi hàm này).
    Áp dụng: chữ đậm màu trắng, nền xanh đậm, căn giữa, kẻ viền mảnh. Tự co giãn độ rộng cột theo
    nội dung dài nhất trong cột (kể cả tiêu đề) nếu auto_width=True.
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F4E78")
    center = Alignment(horizontal="center", vertical="center")
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    max_col = num_cols or ws.max_column
    for c in range(1, max_col + 1):
        cell = ws.cell(row=row_num, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    if auto_width:
        for col in ws.columns:
            length = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max(length + 2, 10), 60)


def export_report(report_path, rows, unmatched_files):
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Khớp theo Excel"
    ws1.append(["STT", "Mã xã", "Số tờ", "Số thửa", "Diện tích", "Tên chủ SD", "Trạng thái", "File PDF đã copy"])
    for r in rows:
        trang_thai = "Đã có PDF" if r["matched"] else "CHƯA CÓ PDF"
        ws1.append([r["stt"], r["maxa"], r["to"], r["thua"], r["dt"], r["ten"] or "",
                    trang_thai, ", ".join(r["matched_files"])])

    ws2 = wb.create_sheet("PDF không khớp")
    ws2.append(["Tên file PDF (không tìm thấy trong Excel theo tiêu chí đã chọn)"])
    for f in unmatched_files:
        ws2.append([f])

    for ws in (ws1, ws2):
        style_excel_header_row(ws, row_num=1)
    wb.save(report_path)


# --------------------------- ĐỔI TÊN HÀNG LOẠT (MẪU TÙY CHỈNH) ---------------------------

TEMPLATE_PLACEHOLDERS = ["maxa", "to", "thua", "ten", "loai", "stt", "dt", "original",
                          "id", "mucdich", "core", "year"]


def render_template(template, row, loai, original_name):
    ext = os.path.splitext(original_name)[1] or ".pdf"
    base_original = os.path.splitext(original_name)[0]
    mapping = {
        "maxa": row.get("maxa"),
        "to": row.get("to"),
        "thua": row.get("thua"),
        "ten": sanitize_filename_part(row.get("ten") or ""),
        "loai": loai,
        "stt": row.get("stt"),
        "dt": row.get("dt") if row.get("dt") is not None else "",
        "original": base_original,
        "id": row.get("id") or "",
        "mucdich": sanitize_filename_part(row.get("mucdich") or ""),
        "core": sanitize_filename_part(row.get("core") or ""),
        "year": sanitize_filename_part(row.get("year") or ""),
    }
    try:
        new_base = template.format(**mapping)
    except (KeyError, IndexError) as e:
        raise ValueError(f"Mẫu tên chứa placeholder không hợp lệ: {{{e}}}")
    new_base = sanitize_filename_part(new_base)
    if not new_base:
        raise ValueError("Mẫu tên tạo ra tên rỗng.")
    return new_base + ext


def rename_bulk(folder, excel_path, sheet_name, header_row,
                 col_stt, col_maxa, col_to, col_thua, col_ten, col_files, col_dt,
                 criteria, dt_tolerance, name_template, log_cb,
                 xa_mapping=None, thon_mapping=None, use_ocr_fallback=False, ocr_dpi=300, debug=False,
                 control=None):
    """
    xa_mapping: nếu được truyền vào (khác None), bật cơ chế DỰ PHÒNG - khi 1 file PDF không
    khớp được với Excel theo tên file (số tờ/thửa/mã xã không tách được từ tên file), phần mềm
    sẽ tự ĐỌC TRỰC TIẾP NỘI DUNG file PDF đó (giống Tab 5) để lấy Mã xã/Tờ/Thửa/Tên chủ/Diện tích
    rồi vẫn đổi tên theo mẫu, thay vì bỏ qua.
    """
    rows = load_excel_data(excel_path, sheet_name, header_row,
                            col_stt, col_maxa, col_to, col_thua, col_ten, col_files, col_dt)
    xa_codes_present = set(r["maxa"] for r in rows) | {norm_num(r["maxa"]) for r in rows}

    renamed_count = 0
    skipped_count = 0
    already_ok_count = 0
    content_fallback_count = 0
    rename_log = []

    files = [f for f in sorted(os.listdir(folder)) if f.lower().endswith(".pdf")]
    existing_targets = set(files)

    for fname in files:
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - files.index(fname)} file chưa xử lý).")
                break
        result, status = find_match(fname, rows, xa_codes_present, criteria, dt_tolerance)
        row = None
        loai = None

        if status == "matched":
            row = result
            loai = detect_loai(fname)
        elif xa_mapping is not None:
            # Dự phòng: không khớp Excel theo tên file -> thử đọc trực tiếp nội dung PDF
            fpath_full = os.path.join(folder, fname)
            log_cb(f"   (Không tách được số từ tên file '{fname}', đang thử đọc nội dung PDF...)")
            info = extract_fields_from_pdf(fpath_full, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, debug=debug)
            if "error" not in info:
                row = {"stt": "", "maxa": info["maxa"], "to": info["to"], "thua": info["thua"],
                       "ten": info.get("ten") or "", "dt": info.get("dt")}
                loai = info.get("loai") or detect_loai(fname)
                content_fallback_count += 1
                via_txt = f"Thôn {info.get('thon_name')}" if info.get("nguon_maxa") == "thon" else f"Xã {info.get('xa_name')}"
                log_cb(f"   ✓ Đọc được nội dung PDF: {via_txt} (Mã {info['maxa']}), Tờ {info['to']}, Thửa {info['thua']}")
            else:
                skipped_count += 1
                log_cb(f"⏭ Bỏ qua (không khớp Excel theo tên file, đọc nội dung PDF cũng lỗi: {info['error']}): {fname}")
                if debug and info.get("raw_text"):
                    log_cb("   ---- Văn bản đã đọc được (để chẩn đoán) ----")
                    log_cb("   " + info["raw_text"].replace("\n", "\n   "))
                    log_cb("   ---- Hết văn bản ----")
                continue
        else:
            skipped_count += 1
            log_cb(f"⏭ Bỏ qua (không khớp rõ ràng để đổi tên): {fname}")
            continue

        try:
            new_name = render_template(name_template, row, loai, fname)
        except ValueError as e:
            skipped_count += 1
            log_cb(f"⚠ Bỏ qua ({e}): {fname}")
            continue

        if new_name == fname:
            already_ok_count += 1
            log_cb(f"= Đã đúng chuẩn, giữ nguyên: {fname}")
            continue

        base, ext = os.path.splitext(new_name)
        i = 2
        while new_name in existing_targets and new_name != fname:
            new_name = f"{base}_{i}{ext}"
            i += 1

        old_path = os.path.join(folder, fname)
        new_path = os.path.join(folder, new_name)
        os.rename(old_path, new_path)
        existing_targets.discard(fname)
        existing_targets.add(new_name)
        renamed_count += 1
        rename_log.append((fname, new_name))
        log_cb(f"✎ Đổi tên: {fname}  →  {new_name}")

    report_path = None
    if rename_log:
        report_path = os.path.join(folder, "BAO_CAO_DOI_TEN.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Đổi tên"
        ws.append(["Tên cũ", "Tên mới"])
        for old, new in rename_log:
            ws.append([old, new])
        style_excel_header_row(ws, row_num=1)
        wb.save(report_path)

    return renamed_count, skipped_count, already_ok_count, report_path, content_fallback_count


# ============================================================================
# MỤC 4 (MỚI): ĐỔI TÊN BẰNG CÁCH ĐỌC TRỰC TIẾP NỘI DUNG FILE PDF
# ============================================================================

def parse_xa_mapping(text):
    """Chuyển text nhiều dòng 'Tên xã=Mã xã' thành dict {TEN_XA_KHONG_DAU: ma_xa}."""
    mapping = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or "=" not in line:
            continue
        ten, ma = line.split("=", 1)
        key = strip_diacritics(ten).upper().strip()
        if key:
            mapping[key] = ma.strip()
    return mapping


CONFIG_XA_THON_FILENAME = "Cau_hinh_ma_xa_thon.xlsx"


def get_xa_thon_config_path():
    """Đường dẫn file cấu hình mã xã/thôn, đặt cạnh file .exe/.py đang chạy (không nằm trong thư mục tạm)."""
    return os.path.join(get_app_dir(), CONFIG_XA_THON_FILENAME)


def create_default_xa_thon_config_file(path):
    """Tạo file cấu hình mẫu Cau_hinh_ma_xa_thon.xlsx từ bảng mặc định đang có sẵn trong phần mềm."""
    thon_mapping = parse_xa_mapping(DEFAULT_THON_MAPPING_TEXT)
    # Bảng {TenThon: (XaCu, MaXa)} dựng lại tên có dấu đầy đủ để người dùng dễ đọc/sửa
    rows = [
        ("Khuổi Phầy", "Kim Hỷ", "02146", ""), ("Kim Vân", "Kim Hỷ", "02146", ""),
        ("Quốc Tuấn", "Kim Hỷ", "02146", ""), ("Nà Mỏ", "Kim Hỷ", "02146", ""),
        ("Khuổi Hát", "Kim Hỷ", "02146", ""),
        ("Nà Làng", "Lương Thượng", "02143", ""), ("Bản Giang", "Lương Thượng", "02143", ""),
        ("Pàn Xả", "Lương Thượng", "02143", ""), ("Vằng Khít", "Lương Thượng", "02143", ""),
        ("Khuổi Nộc", "Lương Thượng", "02143", ""),
        ("Bản Kén", "Văn Lang cũ", "02140", ""), ("Nặm Cà", "Văn Lang cũ", "02140", ""),
        ("Chợ Mới", "Văn Lang cũ", "02140", ""), ("Tân An", "Văn Lang cũ", "02140", ""),
        ("Nà Diệc", "Văn Lang cũ", "02140", ""), ("Bản Sảng", "Văn Lang cũ", "02140", ""),
        ("Nà Lẹng", "Văn Lang cũ", "02140", ""), ("Cốc Phia", "Văn Lang cũ", "02140", ""),
        ("Nà Dường", "Văn Lang cũ", "02140", ""),
    ]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "MaXaThon"
    headers = ["TenThon", "XaCu", "MaXa", "GhiChu"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True)
    for r, row in enumerate(rows, start=2):
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)
    for col, width in zip("ABCD", [18, 16, 10, 30]):
        ws.column_dimensions[col].width = width
    dirpath = os.path.dirname(path)
    if dirpath:
        os.makedirs(dirpath, exist_ok=True)
    wb.save(path)


def load_xa_thon_config_from_excel(path, log_cb=None):
    """
    Đọc file cấu hình Cau_hinh_ma_xa_thon.xlsx (cột TenThon, XaCu, MaXa, GhiChu), trả về
    (thon_mapping, xa_mapping) đúng định dạng {TEN_KHONG_DAU_HOA: ma_xa} dùng chung toàn phần mềm.
    Nếu file không tồn tại hoặc lỗi đọc, trả về (None, None) để phần gọi tự dùng cấu hình mặc định.
    """
    if not os.path.isfile(path):
        return None, None
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        thon_mapping = {}
        xa_mapping = {}
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or len(row) < 3:
                continue
            ten_thon, xa_cu, ma_xa = row[0], row[1], row[2]
            if not ma_xa:
                continue
            ma_xa = str(ma_xa).strip()
            if ten_thon:
                key = strip_diacritics(str(ten_thon)).upper().strip()
                if key:
                    thon_mapping[key] = ma_xa
            if xa_cu:
                key = strip_diacritics(str(xa_cu)).upper().strip()
                if key:
                    xa_mapping[key] = ma_xa
        if not thon_mapping and not xa_mapping:
            if log_cb:
                log_cb(f"⚠ File cấu hình {path} không có dữ liệu hợp lệ, dùng cấu hình mặc định.")
            return None, None
        if log_cb:
            log_cb(f"✓ Đã nạp cấu hình mã xã/thôn từ: {path} "
                   f"({len(thon_mapping)} thôn, {len(xa_mapping)} xã cũ)")
        return thon_mapping, xa_mapping
    except Exception as e:
        if log_cb:
            log_cb(f"⚠ Lỗi đọc file cấu hình {path}: {e} - dùng cấu hình mặc định.")
        return None, None


def get_pdf_text_all_pages(reader):
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def ocr_pdf_text(pdf_path, dpi=300, lang="vie", max_pages=3, log_cb=None):
    """OCR các trang đầu của PDF (dùng khi file là bản scan, không có lớp chữ)."""
    if not (HAS_FITZ and HAS_OCR):
        raise RuntimeError(
            "Thiếu thư viện OCR. Cần cài: pip install pymupdf pytesseract pillow "
            "và cài chương trình Tesseract-OCR (kèm gói ngôn ngữ Vietnamese)."
        )
    import tempfile
    doc = fitz.open(pdf_path)
    texts = []
    with tempfile.TemporaryDirectory() as tmpdir:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(dpi=dpi)
            img_path = os.path.join(tmpdir, f"ocr_tmp_{i}.png")
            pix.save(img_path)
            try:
                txt = pytesseract.image_to_string(Image.open(img_path), lang=lang, config="--psm 6")
            except Exception as e:
                if log_cb:
                    log_cb(f"   (OCR lỗi trang {i+1}: {e})")
                txt = ""
            texts.append(txt)
    doc.close()
    return "\n".join(texts)


def ocr_pdf_text_cached(pdf_path, dpi=300, lang="vie", max_pages=3, log_cb=None, cache_dir=None):
    """
    Bọc ngoài ocr_pdf_text() - THÊM CACHE tùy chọn, dùng CHUNG hạ tầng cache với
    ocr_pdf_to_text_cached() (cùng thư mục get_ocr_cache_dir(), cùng nút "Xóa cache OCR" ở Tab Cài
    đặt) nhưng KHÓA CACHE RIÊNG (tiền tố "simple_v1_") để không trùng/lẫn với kết quả OCR của hàm
    ocr_pdf_to_text() khác (2 hàm dùng tham số/cách xử lý ảnh khác nhau nên kết quả OCR khác nhau,
    không thể dùng chung 1 khóa cache).

    Tăng tốc đáng kể khi CÙNG 1 FILE PDF được xử lý NHIỀU LẦN - ví dụ: nút "🔁 Chạy lại file lỗi/
    cần kiểm tra" (mục 8c), hoặc khi người dùng chạy "Chạy thử" rồi chạy "Tổng hợp báo cáo" ngay
    sau (2 lần đọc cùng dữ liệu). cache_dir=None -> hành vi Y HỆT ocr_pdf_text() cũ, không cache
    (mặc định AN TOÀN - không đổi hành vi nếu không truyền tham số mới này).
    """
    if cache_dir is None:
        return ocr_pdf_text(pdf_path, dpi=dpi, lang=lang, max_pages=max_pages, log_cb=log_cb)

    cache_key = compute_ocr_cache_key(pdf_path, f"simple_v1_{dpi}_{lang}_{max_pages}")
    cached_text = load_ocr_cache(cache_dir, cache_key)
    if cached_text is not None:
        if log_cb:
            log_cb("   (dùng lại kết quả OCR đã lưu trước đó - không OCR lại, file chưa đổi)")
        return cached_text

    text = ocr_pdf_text(pdf_path, dpi=dpi, lang=lang, max_pages=max_pages, log_cb=log_cb)
    save_ocr_cache(cache_dir, cache_key, text)
    return text


# Các mẫu biểu thức được viết LỎNG (chấp nhận nhiều cách viết / lỗi OCR phổ biến)
FIELD_TO_RE = re.compile(
    r"(?:T[ờoO0][\.\s]*(?:b[ảaA][nN]?\s*đ[ồoO0]\s*)?s[ốoO0]|T[ờoO0]\s*s[ốoO0])[:\.\s]*([0-9]{1,6})",
    re.IGNORECASE)
FIELD_THUA_RE = re.compile(
    r"(?:Th[ửuU][ăa]?\s*đ[ấaA][ấaA]?[tT]?\s*s[ốoO0]|Th[ửuU][ăa]?\s*s[ốoO0])[:\.\s]*([0-9]{1,6})",
    re.IGNORECASE)
FIELD_DIACHI_RE = re.compile(
    r"(?:Đ[ịiI][aA]\s*ch[ỉiI]|N[ơoO][iI])[^:\n]{0,15}:\s*([^\n]+)", re.IGNORECASE)
FIELD_XA_RE = re.compile(r"[Xx][ãaA]\s+([^,\n]+)")
FIELD_THON_RE = re.compile(r"Th[oôơOÔƠ][nN]\s+([^,\n]+)", re.IGNORECASE)
FIELD_TEN_RE = re.compile(
    r"(?:H[ọoO]\s*v[àaA]\s*t[êeE]n|Ng[ưuU][ờoO][iI]\s*s[ửuU]\s*d[ụuU][nN]g\s*đ[ấaA][ấaA]?[tT]|"
    r"Ch[ủuU]\s*s[ửuU]\s*d[ụuU][nN]g\s*đ[ấaA][ấaA]?[tT]|H[ộôO]\s*[ôoO]ng|H[ộôO]\s*b[àaA]|T[êeE]n)"
    r"[^:\n]{0,25}:\s*([^\n]+)", re.IGNORECASE)
FIELD_DT_RE = re.compile(
    r"Di[ệeE][nN]?\s*t[íiI]ch[^:\n]{0,15}:\s*([0-9]+[.,]?[0-9]*)\s*m", re.IGNORECASE)
FIELD_ID_RE = re.compile(
    r"(?:CCCD|C[ăaA][nN]\s*c[ướuU][ớoO][cC]\s*c[ôoO][nN]g\s*d[âaA][nN]|"
    r"S[ốôO]\s*đ[ịiI][nN]h\s*d[aA][nN]h(?:\s*c[áaA]\s*nh[âaA][nN])?|CMND|"
    r"Gi[ấaA]y\s*t[ờoO]\s*nh[âaA]n\s*th[âaA]n)[^:\n]*:\s*([0-9]{8,15})", re.IGNORECASE)
FIELD_MUCDICH_RE = re.compile(
    r"(?:S[ửuU]\s*d[ụuU][nN]g\s*v[àaA]o\s*m[ụuU][cC]\s*đ[íiI]ch|"
    r"K[ýyY]\s*hi[ệeE]u\s*lo[ạaA][iI]\s*đ[ấaA][ấaA]?[tT]|m[ụuU][cC]\s*đ[íiI]ch)"
    r"[^:\n]{0,10}:\s*([^\n,]+)", re.IGNORECASE)
FIELD_CORE_RE = re.compile(
    r"Ngu[ồôOÔ][nN]\s*g[ốôOÔ][cC]\s*s[ửưƯ]\s*d[ụuU][nN]g\s*đ[ấaA][ấaA]?[tT][^:\n]*:\s*(.*?)(?=\n\s*[a-zđà-ỹ]\)|\Z)",
    re.IGNORECASE | re.DOTALL)
FIELD_YEAR_RE = re.compile(
    r"Th[ờoO][iI]\s*h[ạaA][nN][^:\n]*s[ửưƯ]\s*d[ụuU][nN]g\s*đ[ấaA][ấaA]?[tT][^:\n]*:\s*([^\n]+)", re.IGNORECASE)
TITLE_TBXN_RE = re.compile(r"TH[ÔO]NG\s*B[ÁA]O\s*X[ÁA]C\s*NH[ẬA]N", re.IGNORECASE)
TITLE_GT_RE = re.compile(r"Đ[ƠO]N\s*Đ[ĂA]NG\s*K[ÝY]", re.IGNORECASE)


def clean_multiline_value(raw):
    if not raw:
        return None
    s = re.sub(r"\s*\n\s*", " ", raw)   # nối các dòng bị ngắt giữa chừng (do PDF xuống dòng)
    s = re.sub(r"\s{2,}", " ", s).strip().rstrip(".")
    return s or None


def clean_name_text(raw):
    if not raw:
        return None
    s = re.sub(r"\([^)]*\)", "", raw)          # bỏ chú thích trong ngoặc, VD "(in hoa)"
    s = re.sub(r"^[\s\-–—:.,]+", "", s)         # bỏ ký tự thừa ở đầu do OCR
    s = re.sub(r"\s{2,}", " ", s).strip()
    return s or None


def parse_fields_from_text(text):
    to_m = FIELD_TO_RE.search(text)
    thua_m = FIELD_THUA_RE.search(text)

    diachi_all = list(FIELD_DIACHI_RE.finditer(text))
    target = None
    if thua_m:
        for m in diachi_all:
            if m.start() > thua_m.start():
                target = m
                break
    if target is None and diachi_all:
        target = diachi_all[-1]
    dia_chi = target.group(1).strip() if target else None
    xa_m = FIELD_XA_RE.search(dia_chi) if dia_chi else None
    xa_name = xa_m.group(1).strip().rstrip(".") if xa_m else None
    thon_m = FIELD_THON_RE.search(dia_chi) if dia_chi else None
    thon_name = thon_m.group(1).strip().rstrip(".") if thon_m else None

    # Địa chỉ THƯỜNG TRÚ (mục 1.c) - occurrence xuất hiện TRƯỚC mục 2 (khác với dia_chi ở trên là mục 2.b)
    dia_chi_thuong_tru = None
    for m in diachi_all:
        if not thua_m or m.start() < thua_m.start():
            dia_chi_thuong_tru = m.group(1).strip()
            break
    xa_m_1c = FIELD_XA_RE.search(dia_chi_thuong_tru) if dia_chi_thuong_tru else None
    xa_name_1c = xa_m_1c.group(1).strip().rstrip(".") if xa_m_1c else None
    thon_m_1c = FIELD_THON_RE.search(dia_chi_thuong_tru) if dia_chi_thuong_tru else None
    thon_name_1c = thon_m_1c.group(1).strip().rstrip(".") if thon_m_1c else None

    # Tên chủ sử dụng đất: lấy occurrence xuất hiện TRƯỚC mục 2 (thông tin thửa đất)
    ten_all = list(FIELD_TEN_RE.finditer(text))
    ten_target = None
    for m in ten_all:
        if not thua_m or m.start() < thua_m.start():
            ten_target = m
            break
    ten = clean_name_text(ten_target.group(1)) if ten_target else None

    # Diện tích: ưu tiên occurrence sau mục 2 (diện tích thửa đất, không phải diện tích xây dựng)
    dt_all = list(FIELD_DT_RE.finditer(text))
    dt_target = None
    if thua_m:
        for m in dt_all:
            if m.start() > thua_m.start():
                dt_target = m
                break
    if dt_target is None and dt_all:
        dt_target = dt_all[0]
    dt_val = None
    if dt_target:
        try:
            dt_val = float(dt_target.group(1).replace(",", "."))
        except ValueError:
            dt_val = None

    loai = "KHAC"
    if TITLE_TBXN_RE.search(text):
        loai = "TBXN"
    elif TITLE_GT_RE.search(text):
        loai = "GT"

    id_m = FIELD_ID_RE.search(text)
    mucdich_m = FIELD_MUCDICH_RE.search(text)
    core_m = FIELD_CORE_RE.search(text)
    year_m = FIELD_YEAR_RE.search(text)

    return {
        "to": to_m.group(1) if to_m else None,
        "thua": thua_m.group(1) if thua_m else None,
        "dia_chi": dia_chi,
        "dia_chi_thuong_tru": dia_chi_thuong_tru,
        "xa_name": xa_name,
        "xa_name_1c": xa_name_1c,
        "thon_name": thon_name,
        "thon_name_1c": thon_name_1c,
        "ten": ten,
        "dt": dt_val,
        "loai": loai,
        "id": id_m.group(1).strip() if id_m else None,
        "mucdich": clean_multiline_value(mucdich_m.group(1)) if mucdich_m else None,
        "core": clean_multiline_value(core_m.group(1)) if core_m else None,
        "year": clean_multiline_value(year_m.group(1)) if year_m else None,
    }


EMPTY_FIELDS = {"to": None, "thua": None, "dia_chi": None, "dia_chi_thuong_tru": None, "xa_name": None,
                 "xa_name_1c": None, "thon_name": None, "thon_name_1c": None, "ten": None, "dt": None,
                 "loai": "KHAC", "id": None, "mucdich": None, "core": None, "year": None}


def write_pdf_debug_text_file(debug_text_dir, pdf_path, text, ocr_text, info):
    """
    Ghi file debug "{TenFilePDF}_extracted_text.txt" vào debug_text_dir - dùng để chẩn đoán khi
    tổng hợp Excel sai (mục X tài liệu "LÀM SÂU CHỨC NĂNG TỔNG HỢP PDF RA EXCEL"). Lỗi ghi file chỉ
    bỏ qua (không phải điều kiện bắt buộc), không được làm hỏng luồng đọc PDF chính.
    """
    if not debug_text_dir:
        return
    try:
        os.makedirs(debug_text_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        out_path = os.path.join(debug_text_dir, f"{base_name}_extracted_text.txt")
        lines = [
            f"FILE PDF: {pdf_path}", "=" * 70, "",
            "----- TEXT LAYER (đọc trực tiếp từ PDF) -----",
            text or "(không có / rỗng)", "",
        ]
        if ocr_text:
            lines += ["----- TEXT OCR (chỉ chạy khi text layer thiếu dữ liệu) -----", ocr_text, ""]
        lines += ["=" * 70, "CÁC TRƯỜNG ĐÃ NHẬN DIỆN + NGUỒN", "=" * 70, ""]
        field_labels = {
            "maxa": "Mã xã", "nguon_maxa": "Nguồn Mã xã", "to": "Số tờ (sau chuẩn hóa)",
            "so_to_goc": "Số tờ gốc đọc được", "loai_dat": "Loại đất",
            "forest_map_status": "Trạng thái số tờ lâm nghiệp", "thua": "Số thửa",
            "ten": "Họ tên", "id": "CCCD/số định danh", "dia_chi_thuong_tru": "Địa chỉ thường trú",
            "dia_chi": "Địa chỉ thửa đất", "dt": "Diện tích", "mucdich": "Mục đích sử dụng đất",
            "nguon": "Nguồn dữ liệu (text/ocr)", "diem_tin_cay_ocr": "Điểm tin cậy",
            "muc_do_tin_cay": "Mức độ tin cậy",
        }
        for key, label in field_labels.items():
            if key in info:
                lines.append(f"{label}: {info.get(key)}")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
    except Exception:
        pass  # file debug la tien ich chan doan, khong duoc lam hong luong doc PDF chinh


def extract_fields_from_pdf(pdf_path, xa_mapping, thon_mapping=None, use_ocr_fallback=True,
                             ocr_dpi=300, log_cb=None, debug=False, debug_text_dir=None,
                             threshold_cao=90, threshold_thap=70, ocr_cache_dir=None):
    """
    Đọc file PDF, trả về dict {to, thua, dia_chi, dia_chi_thuong_tru, xa_name, thon_name, ten, dt,
    id, mucdich, core, year, maxa, loai, nguon, raw_text} hoặc dict có 'error' nếu không đọc đủ thông tin
    (vẫn kèm raw_text nếu debug=True để chẩn đoán).

    Xác định Mã xã: ưu tiên tra theo TÊN THÔN trước (chính xác hơn vì thôn không đổi khi sáp nhập
    địa giới hành chính), nếu không khớp thôn nào mới tra theo TÊN XÃ.

    debug_text_dir (tùy chọn): nếu truyền vào 1 đường dẫn thư mục, tự ghi 1 file
    "{TenFilePDF}_extracted_text.txt" chứa text layer/OCR đọc được + toàn bộ trường đã nhận diện +
    nguồn nhận diện - dùng để chẩn đoán khi tổng hợp Excel sai (mục X tài liệu "LÀM SÂU CHỨC NĂNG
    TỔNG HỢP PDF RA EXCEL"). Lỗi ghi file debug chỉ bỏ qua, KHÔNG được làm hỏng luồng đọc PDF chính.

    ocr_cache_dir (tùy chọn): nếu truyền vào, DÙNG LẠI kết quả OCR đã lưu từ lần đọc TRƯỚC ĐÓ của
    CHÍNH file PDF này (nếu file chưa đổi) thay vì OCR lại từ đầu - tăng tốc RẤT NHIỀU (đo thực tế:
    từ ~8s xuống dưới 1ms) khi cùng 1 file được xử lý nhiều lần (VD nút "Chạy lại file lỗi", hoặc
    "Chạy thử" rồi "Tổng hợp báo cáo" liền sau). None = không cache, hành vi Y HỆT trước đây (an
    toàn - không đổi hành vi mặc định).
    """
    thon_mapping = thon_mapping or {}
    if not HAS_PYPDF:
        return {"error": "Thiếu thư viện pypdf. Cài: pip install pypdf"}

    text = ""
    nguon = "text"
    try:
        reader = PdfReader(pdf_path)
        text = get_pdf_text_all_pages(reader)
    except Exception as e:
        text = ""
        if log_cb:
            log_cb(f"   (Lỗi đọc text PDF: {e})")

    # Nếu lớp chữ có sẵn nhưng quá ít ký tự (VD chỉ có chữ ký số overlay trên ảnh scan) -> vẫn coi là cần OCR
    info = parse_fields_from_text(text) if len(text.strip()) > 30 else dict(EMPTY_FIELDS)

    ocr_text = ""
    if (not info["to"] or not info["thua"] or not info["xa_name"]) and use_ocr_fallback:
        nguon = "ocr"
        try:
            ocr_text = ocr_pdf_text_cached(pdf_path, dpi=ocr_dpi, log_cb=log_cb, cache_dir=ocr_cache_dir)
            ocr_info = parse_fields_from_text(ocr_text)
            for k in ("to", "thua", "dia_chi", "dia_chi_thuong_tru", "xa_name", "xa_name_1c",
                      "thon_name", "thon_name_1c", "ten", "dt", "id", "mucdich", "core", "year"):
                if not info.get(k) and ocr_info.get(k):
                    info[k] = ocr_info[k]
            if info.get("loai", "KHAC") == "KHAC" and ocr_info.get("loai", "KHAC") != "KHAC":
                info["loai"] = ocr_info["loai"]
        except Exception as e:
            if log_cb:
                log_cb(f"   (OCR không khả dụng: {e})")

    debug_text = (text or "") + ("\n----- OCR -----\n" + ocr_text if ocr_text else "")

    if not info["to"] or not info["thua"]:
        missing = []
        if not info["to"]:
            missing.append("Tờ bản đồ số")
        if not info["thua"]:
            missing.append("Thửa đất số")
        result = {"error": f"Không tìm thấy: {', '.join(missing)}.", "nguon": nguon, **info}
        if debug:
            result["raw_text"] = debug_text[:1500]
        write_pdf_debug_text_file(debug_text_dir, pdf_path, text, ocr_text, result)
        return result

    # Nhận diện Loại đất (RSX/RPH/RDD/ONT/CLN/HNK/LUC/BHK) và CHUẨN HÓA Số tờ bản đồ LÂM NGHIỆP
    # nếu cần (tờ 1/2/3 -> 110000/210000/310000...) - PHẢI làm TRƯỚC khi Số tờ được dùng để đặt
    # tên file/ghi Excel, để không lẫn với bản đồ địa chính nông nghiệp. Tìm trong "mucdich" trước
    # (đúng vị trí thường ghi loại đất), sau đó thử toàn văn nếu "mucdich" không có.
    loai_dat = detect_loai_dat(info.get("mucdich")) or detect_loai_dat(text) or detect_loai_dat(ocr_text)
    forest_result = normalize_forest_map_sheet_number(info["to"], loai_dat)
    info["loai_dat"] = loai_dat
    info["so_to_goc"] = info["to"]
    info["to"] = forest_result["so_to_chuan"]
    info["forest_map_status"] = forest_result["trang_thai"]
    info["forest_map_note"] = forest_result["ly_do_chuyen_doi"]
    if forest_result["da_chuyen_doi"] and log_cb:
        log_cb(f"   🌲 {forest_result['ly_do_chuyen_doi']}")

    # Xác định Mã xã theo ĐÚNG 4 mức ưu tiên bắt buộc:
    # 1) Tên thôn trong mục 2.b (Địa chỉ thửa đất)
    # 2) Tên thôn trong mục 1.c (Địa chỉ thường trú)
    # 3) Tên xã CŨ rõ ràng (Kim Hỷ / Lương Thượng / Văn Lang cũ) - mục 2.b rồi mục 1.c
    #    (KHÔNG dùng tên xã sau sáp nhập "Văn Lang" một mình để suy ra 02140)
    # 4) Nếu không thỏa mãn điều nào ở trên -> maxa=None -> bắt buộc CẦN KIỂM TRA
    maxa = None
    nguon_maxa = None
    if info.get("thon_name"):
        key = strip_diacritics(info["thon_name"]).upper().strip()
        maxa = thon_mapping.get(key)
        if maxa:
            nguon_maxa = "thôn (mục 2.b - Địa chỉ thửa đất)"
    if not maxa and info.get("thon_name_1c"):
        key = strip_diacritics(info["thon_name_1c"]).upper().strip()
        maxa = thon_mapping.get(key)
        if maxa:
            nguon_maxa = "thôn (mục 1.c - Địa chỉ thường trú)"
    if not maxa and info.get("xa_name"):
        key = strip_diacritics(info["xa_name"]).upper().strip()
        maxa = xa_mapping.get(key)
        if maxa:
            nguon_maxa = "xã (mục 2.b - Địa chỉ thửa đất)"
    if not maxa and info.get("xa_name_1c"):
        key = strip_diacritics(info["xa_name_1c"]).upper().strip()
        maxa = xa_mapping.get(key)
        if maxa:
            nguon_maxa = "xã (mục 1.c - Địa chỉ thường trú)"

    if not maxa:
        thon_hint = info.get("thon_name") or info.get("thon_name_1c")
        xa_hint = info.get("xa_name") or info.get("xa_name_1c")
        if thon_hint:
            goi_y = f"thôn '{thon_hint}' (chưa có trong bảng tra Thôn→Mã xã)"
        elif xa_hint:
            goi_y = (f"chỉ thấy tên xã '{xa_hint}' (có thể là tên xã SAU SÁP NHẬP, không đủ để suy ra "
                    f"mã xã cũ một cách chắc chắn - không tự động gán)")
        else:
            goi_y = "không đọc được tên thôn/xã nào trong Địa chỉ"
        result = {"error": f"Không xác định được Mã xã: {goi_y}. Hãy bổ sung bảng tra hoặc kiểm tra tay.",
                  "nguon": nguon, **info}
        conf = compute_ocr_confidence(result, debug_text, threshold_cao=threshold_cao, threshold_thap=threshold_thap)
        result["diem_tin_cay_ocr"] = conf["diem"]
        result["muc_do_tin_cay"] = conf["muc_do"]
        result["ly_do_can_kiem_tra"] = conf["ly_do"]
        if debug:
            result["raw_text"] = debug_text[:1500]
        write_pdf_debug_text_file(debug_text_dir, pdf_path, text, ocr_text, result)
        return result

    info["maxa"] = maxa
    info["nguon"] = nguon
    info["nguon_maxa"] = nguon_maxa
    conf = compute_ocr_confidence(info, debug_text, threshold_cao=threshold_cao, threshold_thap=threshold_thap)
    info["diem_tin_cay_ocr"] = conf["diem"]
    info["muc_do_tin_cay"] = conf["muc_do"]
    info["ly_do_can_kiem_tra"] = conf["ly_do"]
    if debug:
        info["raw_text"] = debug_text[:1500]
    write_pdf_debug_text_file(debug_text_dir, pdf_path, text, ocr_text, info)
    return info


def compute_ocr_confidence(info, raw_text=None, threshold_cao=90, threshold_thap=70):
    """
    Chấm điểm độ tin cậy OCR/nhận diện cho 1 file, theo 8 tiêu chí: có Mã xã, Số tờ, Số thửa,
    Họ tên, CCCD, Diện tích, Mục đích SD, và text có chứa đủ từ khóa chuẩn hay không.
    Trả về dict {"diem": 0-100, "muc_do": "TIN_CAY_CAO"|"CAN_KIEM_TRA_NHANH"|"KHONG_TU_XU_LY",
    "ly_do": mô_tả_ngắn_các_trường_còn_thiếu}.
    threshold_cao/threshold_thap: ngưỡng % tùy chỉnh được (mặc định 90/70, cấu hình qua
    data/app_settings.json - Tab Cài đặt). >= threshold_cao: TIN_CAY_CAO (xử lý tự động).
    threshold_thap <= điểm < threshold_cao: CAN_KIEM_TRA_NHANH. < threshold_thap: KHONG_TU_XU_LY
    (bắt buộc CẦN KIỂM TRA thủ công).
    """
    checks = {
        "mã xã": bool(info.get("maxa")),
        "số tờ": bool(info.get("to")),
        "số thửa": bool(info.get("thua")),
        "họ tên": bool(info.get("ten")),
        "CCCD": bool(info.get("id")),
        "diện tích": info.get("dt") is not None,
        "mục đích SD": bool(info.get("mucdich")),
    }

    keyword_ok = False
    if raw_text:
        norm = strip_diacritics(raw_text).lower()
        keywords = ["thua dat so", "to ban do so", "dia chi", "dien tich", "su dung vao muc dich"]
        found = sum(1 for kw in keywords if kw in norm)
        keyword_ok = found >= 2
    checks["từ khóa chuẩn trong văn bản"] = keyword_ok

    total = len(checks)
    passed = sum(1 for v in checks.values() if v)
    score = round(passed / total * 100, 1)

    if score >= threshold_cao:
        muc_do = "TIN_CAY_CAO"
    elif score >= threshold_thap:
        muc_do = "CAN_KIEM_TRA_NHANH"
    else:
        muc_do = "KHONG_TU_XU_LY"

    missing = [k for k, v in checks.items() if not v]
    ly_do = ("Đủ các trường chính (mã xã, tờ, thửa, họ tên, CCCD, diện tích, mục đích SD)."
             if not missing else "Thiếu: " + ", ".join(missing) + ".")

    return {"diem": score, "muc_do": muc_do, "ly_do": ly_do}


def list_files_from_source(source, ext):
    """
    source có thể là:
    - str: đường dẫn 1 thư mục -> lấy toàn bộ file đúng đuôi trong thư mục đó
    - list/tuple: danh sách đường dẫn file cụ thể (có thể ở NHIỀU thư mục khác nhau)
    Trả về list đường dẫn TUYỆT ĐỐI, đã sắp xếp.
    """
    ext = ext.lower()
    if isinstance(source, (list, tuple, set)):
        return sorted(os.path.abspath(f) for f in source if f.lower().endswith(ext))
    return sorted(
        os.path.join(source, f) for f in os.listdir(source) if f.lower().endswith(ext)
    )


def extract_pdfs_from_zip(zip_path, extract_to_dir=None, log_cb=None):
    """
    Giải nén file .zip chứa PDF vào thư mục Temp (mục XVI tài liệu "BỔ SUNG TAB 8 MỤC 8C" -
    hiện CHỈ hỗ trợ .zip qua thư viện chuẩn Python; .rar cần công cụ ngoài (unrar/7z) chưa đóng
    gói kèm phần mềm - để triển khai khi có môi trường Windows thật để đóng gói/kiểm tra công cụ
    portable). KHÔNG sửa file nén gốc.

    extract_to_dir: thư mục đích, mặc định tự tạo trong get_app_data_dir()/Temp.
    Trả về (thư_mục_đã_giải_nén, danh_sách_file_pdf_tuyệt_đối). Nếu file nén lỗi/hỏng thì ghi log
    và trả về (None, []) - KHÔNG làm treo phần mềm (mục XVI: "Nếu tệp nén lỗi thì ghi log, không
    làm treo phần mềm").
    """
    import zipfile
    log_cb = log_cb or (lambda x: None)

    if not zipfile.is_zipfile(zip_path):
        log_cb(f"⚠ File nén lỗi hoặc không phải định dạng .zip hợp lệ: {zip_path}")
        return None, []

    if extract_to_dir is None:
        base_name = os.path.splitext(os.path.basename(zip_path))[0]
        extract_to_dir = os.path.join(get_app_data_dir(), "Temp", f"giai_nen_{base_name}_{int(time.time())}")
    os.makedirs(extract_to_dir, exist_ok=True)

    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            pdf_members = [m for m in zf.namelist() if m.lower().endswith(".pdf") and not m.endswith("/")]
            for member in pdf_members:
                # Chỉ lấy tên file, bỏ cấu trúc thư mục con bên trong zip - tránh path traversal
                # và giữ đơn giản (mọi PDF giải nén phẳng vào 1 thư mục).
                safe_name = os.path.basename(member)
                if not safe_name:
                    continue
                with zf.open(member) as src, open(os.path.join(extract_to_dir, safe_name), "wb") as dst:
                    dst.write(src.read())
        pdf_files = sorted(
            os.path.join(extract_to_dir, f) for f in os.listdir(extract_to_dir) if f.lower().endswith(".pdf"))
        log_cb(f"Đã giải nén {len(pdf_files)} file PDF từ {os.path.basename(zip_path)} vào {extract_to_dir}")
        return extract_to_dir, pdf_files
    except Exception as e:
        log_cb(f"⚠ Lỗi giải nén file .zip {zip_path}: {e}")
        return None, []


def extract_pdfs_from_rar(rar_path, extract_to_dir=None, log_cb=None):
    """
    Giải nén file .rar chứa PDF - dùng bản 7-Zip portable đóng gói kèm phần mềm tại
    tools/7zip/7z.exe (mục IV.8 tài liệu "BỔ SUNG TÍNH NĂNG MỚI + SỬA LỖI ĐÓNG GÓI"). Nếu chưa có
    7z.exe (xem tools/7zip/README.md), báo lỗi RÕ RÀNG hướng dẫn cách bổ sung, KHÔNG làm treo
    phần mềm. KHÔNG sửa file nén gốc.

    Trả về (thư_mục_đã_giải_nén, danh_sách_file_pdf_tuyệt_đối) - giống hệt extract_pdfs_from_zip()
    để 2 hàm dùng thay thế nhau được ở nơi gọi.
    """
    import subprocess
    log_cb = log_cb or (lambda x: None)

    if not check_rar_support_available():
        log_cb("⚠ Chưa hỗ trợ giải nén .rar - thiếu tools/7zip/7z.exe. "
              "Xem tools/7zip/README.md để biết cách bổ sung, hoặc dùng file .zip thay thế.")
        return None, []

    if extract_to_dir is None:
        base_name = os.path.splitext(os.path.basename(rar_path))[0]
        extract_to_dir = os.path.join(get_app_data_dir(), "Temp", f"giai_nen_{base_name}_{int(time.time())}")
    os.makedirs(extract_to_dir, exist_ok=True)

    seven_zip_exe = os.path.join(get_base_dir(), "tools", "7zip", "7z.exe")
    try:
        result = subprocess.run(
            [seven_zip_exe, "e", "-y", f"-o{extract_to_dir}", rar_path, "*.pdf", "-r"],
            capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            log_cb(f"⚠ Lỗi giải nén file .rar {rar_path}: {result.stderr.strip() or result.stdout.strip()}")
            return None, []
        pdf_files = sorted(
            os.path.join(extract_to_dir, f) for f in os.listdir(extract_to_dir) if f.lower().endswith(".pdf"))
        log_cb(f"Đã giải nén {len(pdf_files)} file PDF từ {os.path.basename(rar_path)} vào {extract_to_dir}")
        return extract_to_dir, pdf_files
    except subprocess.TimeoutExpired:
        log_cb(f"⚠ Giải nén file .rar {rar_path} quá lâu (>120s) - đã hủy.")
        return None, []
    except Exception as e:
        log_cb(f"⚠ Lỗi giải nén file .rar {rar_path}: {e}")
        return None, []


# ============================================================================
# MODULE OCR-RENAME CHUYÊN DỤNG CHO PDF SCAN/ẢNH (không có text layer)
# Luồng: render trang -> ảnh (matrix 3x) -> OCR vie+eng -> chuẩn hóa text
#        -> trích số thửa/số tờ (mục 2.a) + thôn/xã (mục 2.b) -> đổi tên
#
# GIAI ĐOẠN 1 TÁI KIẾN TRÚC: các hàm nhận diện lõi bên dưới nay được lấy từ module dùng
# chung app/core/land_parser.py (KHÔNG viết logic riêng ở đây nữa). Nếu vì lý do nào đó
# không import được module này (VD thiếu thư mục app/ khi đóng gói), tự động dùng lại
# bản sao lưu trữ tại chỗ (giữ nguyên logic đã kiểm thử) để phần mềm KHÔNG bao giờ vỡ.
# ============================================================================

try:
    from app.core.land_parser import (
        normalize_ocr_text,
        OCR_THUA_RE, OCR_TO_RE, OCR_PAIR_RE,
        _clean_ocr_number,
        ocr_extract_to_thua,
        build_norm_lookup,
        detect_loai_dat,
        normalize_forest_map_sheet_number,
        cross_check_filename_vs_content,
        normalize_ma_xa,
        compare_field_values,
        parse_land_key_from_filename,
        _normalize_area_for_compare,
    )
    from app.core.land_parser import ocr_find_ma_xa as ocr_find_maxa
    _LAND_PARSER_MODULE_OK = True
except Exception:
    _LAND_PARSER_MODULE_OK = False

    def normalize_ocr_text(s):
        """Chuẩn hóa text OCR: 'đ'->'d', bỏ dấu, lower, gộp khoảng trắng."""
        if not s:
            return ""
        s = s.replace("đ", "d").replace("Đ", "D")
        s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
        s = s.lower()
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    OCR_THUA_RE = re.compile(r"(?:thua\s*dat\s*s[o0]|thua\s*s[o0]|s[o0]\s*thua|dat\s*s[o0])\s*:?\s*((?:\d\s*){1,6})")
    OCR_TO_RE = re.compile(r"(?:t[o0]\s*ban\s*d[eo]\s*s[o0]|s[o0]\s*t[o0]\s*ban\s*d[eo]\s*s[o0]|t[o0]\s*ban\s*d[eo]\s*s[o0]?)\s*:?\s*((?:\d\s*){1,6})")
    OCR_PAIR_RE = re.compile(r"s[o0]\s*:?\s*((?:\d\s*){1,6})\s+(?:t[o0]\s*ban\s*d[eo]\s*s[o0])\s*:?\s*((?:\d\s*){1,6})")

    def _clean_ocr_number(s):
        """Ghép lại số bị OCR tách rời bởi khoảng trắng (VD '1 13' -> '113'), bỏ số 0 thừa ở đầu."""
        if s is None:
            return None
        digits = re.sub(r"\s+", "", s)
        digits = digits.lstrip("0") or "0"
        return digits

    def ocr_extract_to_thua(text_norm):
        """Trả về (so_to, so_thua, vi_tri_neo) từ text đã chuẩn hóa. vi_tri_neo dùng để tách văn bản
        thành phần TRƯỚC (mục 1 - thông tin người SDĐ) và phần SAU (mục 2 - thông tin thửa đất)."""
        tm = OCR_THUA_RE.search(text_norm)
        om = OCR_TO_RE.search(text_norm)
        thua = _clean_ocr_number(tm.group(1)) if tm else None
        to = _clean_ocr_number(om.group(1)) if om else None
        anchor_pos = tm.start() if tm else None
        if thua is None:
            pm = OCR_PAIR_RE.search(text_norm)
            if pm:
                thua = _clean_ocr_number(pm.group(1))
                if to is None:
                    to = _clean_ocr_number(pm.group(2))
                anchor_pos = pm.start()
        if anchor_pos is None and om:
            anchor_pos = om.start()
        return to, thua, anchor_pos

    def build_norm_lookup(mapping):
        return {normalize_ocr_text(k): v for k, v in mapping.items()}

    def ocr_find_maxa(text_norm, thon_lookup_norm, xa_lookup_norm, anchor_pos=None):
        """
        Xác định Mã xã theo ĐÚNG 4 mức ưu tiên bắt buộc (dùng anchor_pos - vị trí mục 2.a - để tách
        văn bản thành phần SAU anchor = mục 2.b "Địa chỉ thửa đất", phần TRƯỚC anchor = mục 1.c
        "Địa chỉ thường trú"):
        1) Tên thôn trong phần SAU (mục 2.b)
        2) Tên thôn trong phần TRƯỚC (mục 1.c)
        3) Tên xã CŨ rõ ràng (Kim Hỷ/Lương Thượng/Văn Lang cũ) trong phần SAU rồi phần TRƯỚC
           (bảng tra xã mặc định KHÔNG có "Văn Lang" trơn, nên "xã Văn Lang" một mình sẽ KHÔNG khớp)
        4) Không có gì khớp -> (None, None) -> bắt buộc CẦN KIỂM TRA
        Trả về (maxa, mô_tả_nguồn) hoặc (None, None).
        """
        if anchor_pos is not None:
            after = text_norm[anchor_pos:]
            before = text_norm[:anchor_pos]
        else:
            after = text_norm
            before = ""

        for thon in sorted(thon_lookup_norm.keys(), key=len, reverse=True):
            if thon and re.search(r"\b" + re.escape(thon) + r"\b", after):
                return thon_lookup_norm[thon], f"thôn {thon} (mục 2.b)"
        for thon in sorted(thon_lookup_norm.keys(), key=len, reverse=True):
            if thon and re.search(r"\b" + re.escape(thon) + r"\b", before):
                return thon_lookup_norm[thon], f"thôn {thon} (mục 1.c)"
        for xa in sorted(xa_lookup_norm.keys(), key=len, reverse=True):
            if xa and re.search(r"\b" + re.escape(xa) + r"\b", after):
                return xa_lookup_norm[xa], f"xã {xa} (mục 2.b)"
        for xa in sorted(xa_lookup_norm.keys(), key=len, reverse=True):
            if xa and re.search(r"\b" + re.escape(xa) + r"\b", before):
                return xa_lookup_norm[xa], f"xã {xa} (mục 1.c)"
        return None, None

    LOAI_DAT_LAM_NGHIEP = ("RSX", "RPH", "RDD")
    LOAI_DAT_RE = re.compile(r"\b(RSX|RPH|RDD|ONT|CLN|HNK|LUC|BHK)\b")
    DEFAULT_FOREST_MAP_RULES = [
        {"LoaiDat": "RSX", "SoToGoc": "1", "SoToChuan": "110000", "ChapNhanTuDongChuyen": "YES"},
        {"LoaiDat": "RPH", "SoToGoc": "2", "SoToChuan": "210000", "ChapNhanTuDongChuyen": "YES"},
        {"LoaiDat": "RDD", "SoToGoc": "3", "SoToChuan": "310000", "ChapNhanTuDongChuyen": "YES"},
    ]
    FOREST_MAP_STANDARD_NUMBERS = ("110000", "210000", "310000")

    def detect_loai_dat(text):
        if not text:
            return None
        m = LOAI_DAT_RE.search(text)
        return m.group(1).upper() if m else None

    def normalize_forest_map_sheet_number(so_to_goc, loai_dat, rules=None):
        so_to = norm_num(so_to_goc) if so_to_goc is not None else ""
        loai_dat_norm = str(loai_dat or "").strip().upper()
        if loai_dat_norm not in LOAI_DAT_LAM_NGHIEP:
            return {"so_to_chuan": so_to, "da_chuyen_doi": False,
                    "ly_do_chuyen_doi": "Không phải đất lâm nghiệp RSX/RPH/RDD",
                    "trang_thai": "KHONG_PHAI_DAT_LAM_NGHIEP"}
        if so_to in FOREST_MAP_STANDARD_NUMBERS:
            return {"so_to_chuan": so_to, "da_chuyen_doi": False,
                    "ly_do_chuyen_doi": "Số tờ lâm nghiệp đã đúng định dạng chuẩn, giữ nguyên",
                    "trang_thai": "SO_TO_LAM_NGHIEP_DA_DUNG"}
        active_rules = rules if rules else DEFAULT_FOREST_MAP_RULES
        for rule in active_rules:
            rule_loai = str(rule.get("LoaiDat", "")).strip().upper()
            rule_so_to_goc = norm_num(rule.get("SoToGoc", ""))
            if rule_loai == loai_dat_norm and rule_so_to_goc == so_to:
                chap_nhan = str(rule.get("ChapNhanTuDongChuyen", "")).strip().upper()
                if chap_nhan == "YES":
                    so_to_chuan = norm_num(rule.get("SoToChuan", ""))
                    return {"so_to_chuan": so_to_chuan, "da_chuyen_doi": True,
                            "ly_do_chuyen_doi": f"Tự chuyển theo cấu hình bản đồ lâm nghiệp 1/10000 "
                                                f"({loai_dat_norm}: tờ {so_to} -> {so_to_chuan})",
                            "trang_thai": "DA_CHUYEN_SO_TO_LAM_NGHIEP"}
        return {"so_to_chuan": so_to, "da_chuyen_doi": False,
                "ly_do_chuyen_doi": f"Đất {loai_dat_norm} nhưng không có quy tắc chuyển Số tờ "
                                    f"{so_to} phù hợp trong cấu hình - không tự đoán",
                "trang_thai": "CAN_KIEM_TRA_SO_TO_LAM_NGHIEP"}

    def cross_check_filename_vs_content(filename, content_maxa, content_so_to, content_so_thua,
                                          loai_dat=None, valid_ma_xa_list=None):
        base = os.path.splitext(os.path.basename(filename))[0]
        m3 = KEY_FROM_FILENAME_RE.search(base)
        if not m3:
            return {"co_mau_thuan": False, "chi_tiet": [],
                    "trang_thai": "KHONG_CO_DU_LIEU_TEN_FILE_DE_SO_SANH"}
        ma_xa_ten_file_goc = norm_num(m3.group(1)) if m3.group(1).isdigit() and len(m3.group(1)) == 4 else m3.group(1)
        ma_xa_ten_file = m3.group(1)
        if len(ma_xa_ten_file) == 4:
            candidate = "0" + ma_xa_ten_file
            if not valid_ma_xa_list or candidate in valid_ma_xa_list:
                ma_xa_ten_file = candidate
        so_to_ten_file = norm_num(m3.group(2))
        so_thua_ten_file = norm_num(m3.group(3))

        chi_tiet = []
        if content_maxa and ma_xa_ten_file != content_maxa:
            chi_tiet.append(f"Mã xã: tên file ghi {ma_xa_ten_file}, nội dung PDF đọc được {content_maxa}")
        if content_so_to and so_to_ten_file != content_so_to:
            so_to_ten_file_chuan = normalize_forest_map_sheet_number(so_to_ten_file, loai_dat)["so_to_chuan"]
            if so_to_ten_file_chuan != content_so_to:
                chi_tiet.append(f"Số tờ: tên file ghi {so_to_ten_file}, nội dung PDF đọc được {content_so_to}")
        if content_so_thua and so_thua_ten_file != content_so_thua:
            chi_tiet.append(f"Số thửa: tên file ghi {so_thua_ten_file}, nội dung PDF đọc được {content_so_thua}")
        if chi_tiet:
            return {"co_mau_thuan": True, "chi_tiet": chi_tiet, "trang_thai": "DU_LIEU_MAU_THUAN_CAN_KIEM_TRA"}
        return {"co_mau_thuan": False, "chi_tiet": [], "trang_thai": "KHOP"}

    def normalize_ma_xa(ma_xa, valid_ma_xa_list=None):
        if ma_xa is None:
            return "", ""
        ma_xa_goc = str(ma_xa).strip()
        digits = re.sub(r"\D", "", ma_xa_goc)
        if not digits:
            return "", ma_xa_goc
        if len(digits) == 5:
            return digits, ma_xa_goc
        if len(digits) == 4:
            candidate = "0" + digits
            if valid_ma_xa_list:
                if candidate in valid_ma_xa_list:
                    return candidate, ma_xa_goc
                return digits, ma_xa_goc
            return candidate, ma_xa_goc
        return digits, ma_xa_goc

    def _normalize_text_for_compare(value):
        if value is None:
            return ""
        s = str(value).strip()
        s = re.sub(r"\s+", " ", s)
        s = s.lower()
        s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
        s = s.replace("đ", "d")
        return s

    def _normalize_area_for_compare(value):
        if value is None or value == "":
            return None
        s = str(value).strip().replace(" ", "")
        if re.fullmatch(r"\d{1,3}(\.\d{3})+,\d+", s):
            s = s.replace(".", "").replace(",", ".")
        elif "," in s and "." not in s:
            s = s.replace(",", ".")
        try:
            return float(s)
        except ValueError:
            return None

    def compare_field_values(field_name, pdf_value, excel_value):
        pdf_empty = pdf_value is None or str(pdf_value).strip() == ""
        excel_empty = excel_value is None or str(excel_value).strip() == ""
        if pdf_empty and excel_empty:
            return {"ket_qua": "PDF_VA_EXCEL_DEU_THIEU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
        if pdf_empty:
            return {"ket_qua": "PDF_THIEU_DU_LIEU", "co_chinh_sua": "KHONG",
                    "ghi_chu": "PDF không có dữ liệu để so sánh, không tự kết luận là sai"}
        if excel_empty:
            return {"ket_qua": "EXCEL_TONG_THIEU_DU_LIEU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
        if field_name == "dientich":
            pdf_num = _normalize_area_for_compare(pdf_value)
            excel_num = _normalize_area_for_compare(excel_value)
            if pdf_num is not None and excel_num is not None and abs(pdf_num - excel_num) < 0.01:
                return {"ket_qua": "GIONG_NHAU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
            return {"ket_qua": "KHAC", "co_chinh_sua": "CO",
                    "ghi_chu": "Dữ liệu báo cáo lấy theo Excel tổng đã chỉnh sửa"}
        if field_name in ("so_to", "so_thua", "cccd"):
            pdf_norm = norm_num(pdf_value) if str(pdf_value).strip().replace(".0", "").isdigit() else str(pdf_value).strip()
            excel_norm = norm_num(excel_value) if str(excel_value).strip().replace(".0", "").isdigit() else str(excel_value).strip()
            if pdf_norm == excel_norm:
                return {"ket_qua": "GIONG_NHAU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
            return {"ket_qua": "KHAC", "co_chinh_sua": "CO",
                    "ghi_chu": "Dữ liệu báo cáo lấy theo Excel tổng đã chỉnh sửa"}
        if field_name == "maxa":
            pdf_norm, _ = normalize_ma_xa(pdf_value)
            excel_norm, _ = normalize_ma_xa(excel_value)
            if pdf_norm == excel_norm:
                return {"ket_qua": "GIONG_NHAU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
            return {"ket_qua": "KHAC", "co_chinh_sua": "CO",
                    "ghi_chu": "Dữ liệu báo cáo lấy theo Excel tổng đã chỉnh sửa"}
        pdf_norm = _normalize_text_for_compare(pdf_value)
        excel_norm = _normalize_text_for_compare(excel_value)
        if pdf_norm == excel_norm:
            return {"ket_qua": "GIONG_NHAU", "co_chinh_sua": "KHONG", "ghi_chu": ""}
        return {"ket_qua": "KHAC", "co_chinh_sua": "CO",
                "ghi_chu": "Dữ liệu báo cáo lấy theo Excel tổng đã chỉnh sửa"}

    def parse_land_key_from_filename(filename, valid_ma_xa_list=None):
        fname_only = os.path.basename(filename)
        base = os.path.splitext(fname_only)[0]
        m3 = KEY_FROM_FILENAME_RE.search(base)
        if not m3:
            return {"ma_xa": "", "ma_xa_goc": "", "so_to": "", "so_thua": "", "hau_to": "",
                    "compare_key": "", "status": "KHONG_TACH_DUOC_TU_TEN_FILE",
                    "note": f"Không tách được Mã xã/Tờ/Thửa từ tên file: {fname_only}"}
        ma_xa, ma_xa_goc = normalize_ma_xa(m3.group(1), valid_ma_xa_list)
        so_to = norm_num(m3.group(2))
        so_thua = norm_num(m3.group(3))
        remainder = base[m3.end():].lstrip("_-")
        hau_to = remainder.upper() if remainder else ""
        compare_key = f"{ma_xa}_{so_to}_{so_thua}" if ma_xa else f"{so_to}_{so_thua}"
        note = "Nhận diện Mã xã/Số tờ/Số thửa từ tên file"
        if ma_xa_goc and ma_xa_goc != ma_xa:
            note += f" (đã chuẩn hóa Mã xã {ma_xa_goc} -> {ma_xa})"
        return {"ma_xa": ma_xa, "ma_xa_goc": ma_xa_goc, "so_to": so_to, "so_thua": so_thua,
                "hau_to": hau_to, "compare_key": compare_key, "status": "DA_TACH_TU_TEN_FILE", "note": note}




try:
    from app.services.ocr_service import (
        OCR_MODE_PRESETS,
        DEFAULT_OCR_MODE,
        preprocess_image_for_ocr,
        ocr_image_with_fallback,
        ocr_pdf_to_text,
    )
    _OCR_SERVICE_MODULE_OK = True
except Exception:
    _OCR_SERVICE_MODULE_OK = False

    OCR_MODE_PRESETS = {
        "Nhanh": {"matrix": 2.5, "preprocess": False, "psm_list": [6], "thresholds": None},
        "Chuẩn": {"matrix": 3.5, "preprocess": True, "psm_list": [6, 4], "thresholds": [150]},
        "Kỹ": {"matrix": 4.0, "preprocess": True, "psm_list": [6, 4, 11], "thresholds": [130, 150, 170]},
    }
    DEFAULT_OCR_MODE = "Chuẩn"

    def preprocess_image_for_ocr(img, do_preprocess=True, thresholds=None):
        """
        Tiền xử lý ảnh trước OCR: grayscale -> tăng contrast -> sharpen -> threshold.
        Trả về LIST các biến thể ảnh (để thử nhiều mức threshold ở chế độ "Kỹ").
        do_preprocess=False (chế độ "Nhanh"): chỉ chuyển grayscale, không xử lý sâu (nhanh hơn).
        """
        gray = img.convert("L")
        if not do_preprocess:
            return [gray]
        proc = ImageOps.autocontrast(gray)
        proc = proc.filter(ImageFilter.SHARPEN)
        thresholds = thresholds or [150]
        return [proc.point(lambda p, t=t: 255 if p > t else 0) for t in thresholds]

    def ocr_image_with_fallback(img_variants, psm_list, lang="vie+eng"):
        """
        Thử OCR trên từng biến thể ảnh với từng chế độ --psm trong psm_list, giữ lại kết quả DÀI NHẤT
        (coi là ít khả năng bị OCR bỏ sót/đọc thiếu chữ hơn). Trả về (text, mo_ta_psm_da_dung).
        """
        best_text, best_desc = "", "-"
        for img in img_variants:
            for psm in psm_list:
                try:
                    txt = pytesseract.image_to_string(img, lang=lang, config=f"--psm {psm}")
                except Exception:
                    txt = ""
                if len(txt.strip()) > len(best_text.strip()):
                    best_text, best_desc = txt, f"--psm {psm}"
        return best_text, best_desc

    def ocr_pdf_to_text(pdf_path, matrix_scale=4, lang="vie+eng", log_cb=None, debug_dir=None,
                         mode=None, do_preprocess=None, psm_list=None, thresholds=None):
        """
        Render mọi trang PDF thành ảnh, tiền xử lý, OCR (thử nhiều --psm / nhiều threshold tùy chế độ),
        ghép text các trang. Có 2 cách chọn cấu hình:
        - Truyền `mode` = "Nhanh" / "Chuẩn" / "Kỹ" (dùng preset có sẵn, bỏ qua matrix_scale truyền vào), HOẶC
        - Tự truyền matrix_scale/do_preprocess/psm_list/thresholds riêng.
        Nếu debug_dir được truyền vào: lưu debug_page_N.png + debug_page_N_processed.png cho từng trang.
        Trả về text OCR thô (chưa chuẩn hóa) ghép từ tất cả các trang.
        """
        if not (HAS_FITZ and HAS_OCR):
            raise RuntimeError(
                "Thiếu thư viện OCR. Cần cài: pip install pymupdf pytesseract pillow "
                "và cài chương trình Tesseract-OCR (kèm gói ngôn ngữ Vietnamese - vie)."
            )
        if mode and mode in OCR_MODE_PRESETS:
            preset = OCR_MODE_PRESETS[mode]
            matrix_scale = preset["matrix"]
            do_preprocess = preset["preprocess"]
            psm_list = preset["psm_list"]
            thresholds = preset["thresholds"]
        else:
            do_preprocess = True if do_preprocess is None else do_preprocess
            psm_list = psm_list or [6, 4]
            thresholds = thresholds or [150]

        import tempfile
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        doc = fitz.open(pdf_path)
        texts = []

        def _process_pages(tmp_or_debug_dir, save_debug):
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=fitz.Matrix(matrix_scale, matrix_scale))
                raw_path = os.path.join(tmp_or_debug_dir, f"debug_page_{i+1}.png" if save_debug else f"page_{i}.png")
                pix.save(raw_path)
                img = Image.open(raw_path)
                variants = preprocess_image_for_ocr(img, do_preprocess=do_preprocess, thresholds=thresholds)
                if save_debug:
                    proc_path = os.path.join(tmp_or_debug_dir, f"debug_page_{i+1}_processed.png")
                    variants[0].save(proc_path)
                try:
                    txt, psm_used = ocr_image_with_fallback(variants, psm_list, lang=lang)
                except Exception as e:
                    if log_cb:
                        log_cb(f"   (OCR lỗi trang {i+1}: {e})")
                    txt, psm_used = "", "-"
                if log_cb:
                    log_cb(f"   Trang {i+1}: OCR {psm_used}, {len(txt.strip())} ký tự")
                texts.append(txt)

        if debug_dir:
            page_debug_dir = os.path.join(debug_dir, base_name)
            os.makedirs(page_debug_dir, exist_ok=True)
            _process_pages(page_debug_dir, save_debug=True)
        else:
            with tempfile.TemporaryDirectory() as td:
                _process_pages(td, save_debug=False)

        doc.close()
        full_text = "\n".join(texts)

        if debug_dir:
            txt_path = os.path.join(debug_dir, f"debug_ocr_text__{base_name}.txt")
            try:
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(full_text)
            except Exception as e:
                if log_cb:
                    log_cb(f"   (Không lưu được debug_ocr_text: {e})")

        return full_text


# ------------------------- CACHE OCR (dùng chung, không phụ thuộc module thật/dự phòng) -------------------------
# Bọc ngoài ocr_pdf_to_text() ở TRÊN (dù là bản thật từ app.services.ocr_service hay bản dự phòng
# tại chỗ) - xem giải thích đầy đủ trong app/services/ocr_service.py.

import hashlib as _hashlib_ocr_cache


def compute_ocr_cache_key(pdf_path, mode_or_scale):
    try:
        stat = os.stat(pdf_path)
        size = stat.st_size
        mtime = int(stat.st_mtime)
    except OSError:
        return None
    raw_key = f"{os.path.abspath(pdf_path)}|{size}|{mtime}|{mode_or_scale}"
    return _hashlib_ocr_cache.sha1(raw_key.encode("utf-8")).hexdigest()


def load_ocr_cache(cache_dir, cache_key):
    if not cache_key:
        return None
    cache_path = os.path.join(cache_dir, cache_key + ".txt")
    if os.path.isfile(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                return f.read()
        except OSError:
            return None
    return None


def save_ocr_cache(cache_dir, cache_key, text):
    if not cache_key:
        return
    try:
        os.makedirs(cache_dir, exist_ok=True)
        cache_path = os.path.join(cache_dir, cache_key + ".txt")
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(text)
    except OSError:
        pass


def clear_ocr_cache(cache_dir):
    if not os.path.isdir(cache_dir):
        return 0
    count = 0
    for fname in os.listdir(cache_dir):
        if fname.endswith(".txt"):
            try:
                os.remove(os.path.join(cache_dir, fname))
                count += 1
            except OSError:
                pass
    return count


def ocr_pdf_to_text_cached(pdf_path, cache_dir=None, matrix_scale=4, lang="vie+eng", log_cb=None,
                            debug_dir=None, mode=None, do_preprocess=None, psm_list=None, thresholds=None):
    """Bọc ngoài ocr_pdf_to_text() - THÊM CACHE tùy chọn. cache_dir=None -> hành vi y hệt cũ."""
    if cache_dir is None:
        return ocr_pdf_to_text(pdf_path, matrix_scale=matrix_scale, lang=lang, log_cb=log_cb,
                               debug_dir=debug_dir, mode=mode, do_preprocess=do_preprocess,
                               psm_list=psm_list, thresholds=thresholds)

    cache_key = compute_ocr_cache_key(pdf_path, mode or matrix_scale)
    cached_text = load_ocr_cache(cache_dir, cache_key)
    if cached_text is not None:
        if log_cb:
            log_cb("   (dùng lại kết quả OCR đã lưu trước đó - không OCR lại, file chưa đổi)")
        return cached_text

    text = ocr_pdf_to_text(pdf_path, matrix_scale=matrix_scale, lang=lang, log_cb=log_cb,
                           debug_dir=debug_dir, mode=mode, do_preprocess=do_preprocess,
                           psm_list=psm_list, thresholds=thresholds)
    save_ocr_cache(cache_dir, cache_key, text)
    return text


def get_ocr_cache_dir():
    """Thư mục cache OCR (cache/ocr_text/ trong dữ liệu ứng dụng - cùng chỗ với Logs/Backup)."""
    return os.path.join(get_app_data_dir(), "cache", "ocr_text")


def ocr_rename_scan_pdfs(source, thon_mapping, xa_mapping, matrix_scale, log_cb,
                          name_suffix="GT", do_rename=False, debug_dir=None, control=None, mode=None,
                          use_cache=True, progress_cb=None,
                          toc_do_xu_ly="can_bang", so_file_moi_dot=100, nghi_giua_dot_giay=2.0,
                          perf_log_path=None, checkpoint_task_id=None, files_da_xong=None):
    """
    Đổi tên PDF scan theo luồng OCR. mode: "Nhanh"/"Chuẩn"/"Kỹ" (None = dùng matrix_scale thủ công).
    Trả về {ok: [...], need_check: [...]}. Chỉ đổi tên nếu do_rename=True.
    Mỗi bản ghi trong ok/need_check gồm đủ cột log yêu cầu: old_name, new_name (dự kiến), maxa,
    nguon_maxa, tu_khoa_nhan_dien, to, thua, status, note, ocr_short.
    use_cache=True (mặc định): dùng lại kết quả OCR đã lưu nếu file KHÔNG đổi (kích thước + ngày
    sửa) và CÙNG chế độ OCR - tránh OCR lại khi "Chạy thử" rồi "Chạy thật" trên cùng lô hồ sơ.
    progress_cb(idx, total, filename) nếu có: gọi TRƯỚC khi xử lý mỗi file - dùng để cập nhật
    thanh tiến độ trên giao diện (idx bắt đầu từ 1, không phải 0).

    CHECKPOINT/RESUME (mục XVII) - CHỈ áp dụng khi checkpoint_task_id KHÔNG None VÀ do_rename=True
    (chạy thử không cần checkpoint vì không thay đổi gì thật). files_da_xong: set đường dẫn file
    đã xử lý xong ở lần chạy TRƯỚC (nếu đang resume) - sẽ được BỎ QUA.
    TỐI ƯU HIỆU NĂNG (đúng tài liệu "RÀ SOÁT, KIỂM THỬ VÀ TỐI ƯU TOÀN BỘ PHẦN MỀM" mục V/VI/VII) -
    dùng CHUNG cơ chế đã kiểm chứng ở mục 7d (word_perf_service): nghỉ ngắn sau mỗi file OCR,
    nghỉ dài hơn sau mỗi đợt, hạ độ ưu tiên tiến trình - để không chiếm hết CPU/Disk I/O khi OCR
    hàng loạt nhiều file, người dùng vẫn dùng được Explorer/Excel/PDF khác trong lúc xử lý.
    CHỈ áp dụng khi có TỪ 2 FILE TRỞ LÊN (tránh làm chậm không cần thiết với batch nhỏ).
    """
    files = list_files_from_source(source, ".pdf")
    if not files:
        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")

    from app.services import word_perf_service as _wperf
    ap_dung_toi_uu = len(files) >= 2
    if ap_dung_toi_uu:
        _wperf.apply_low_priority()
        sleep_per_file = _wperf.get_sleep_per_file(toc_do_xu_ly)
    else:
        sleep_per_file = 0.0

    thon_lookup_norm = build_norm_lookup(thon_mapping)
    xa_lookup_norm = build_norm_lookup(xa_mapping)

    ok = []
    need_check = []
    existing_by_dir = {}

    bi_huy_giua_chung = False
    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                bi_huy_giua_chung = True
                break
        if files_da_xong and fpath in files_da_xong:
            continue  # RESUME: file nay da xu ly xong o lan chay TRUOC, khong xu ly lai
        fname = os.path.basename(fpath)
        if progress_cb:
            progress_cb(i + 1, len(files), fname)
        dirpath = os.path.dirname(fpath) or "."
        if dirpath not in existing_by_dir:
            try:
                existing_by_dir[dirpath] = set(os.listdir(dirpath))
            except OSError:
                existing_by_dir[dirpath] = set()

        log_cb(f"→ OCR ({mode or f'{matrix_scale}x'}): {fname}")
        try:
            cache_dir = get_ocr_cache_dir() if use_cache else None
            raw = ocr_pdf_to_text_cached(fpath, cache_dir=cache_dir, matrix_scale=matrix_scale,
                                         log_cb=log_cb, debug_dir=debug_dir, mode=mode)
        except Exception as e:
            need_check.append({"old_name": fname, "new_name": "", "maxa": "", "nguon_maxa": "",
                               "tu_khoa_nhan_dien": "", "to": "", "thua": "",
                               "status": "LOI_OCR", "note": f"Lỗi OCR: {e}", "ocr_short": ""})
            log_cb(f"   ✗ LOI_OCR: {e}")
            continue

        norm = normalize_ocr_text(raw)
        ocr_short = norm[:250]
        to, thua, anchor_pos = ocr_extract_to_thua(norm)
        maxa, nguon = ocr_find_maxa(norm, thon_lookup_norm, xa_lookup_norm, anchor_pos)

        tu_khoa = ""
        if nguon:
            km = re.match(r"^(?:thôn|xã)\s+(.+?)\s+\(", nguon)
            tu_khoa = km.group(1) if km else nguon

        missing = []
        if not maxa:
            missing.append("Mã xã (thôn/xã)")
        if not to:
            missing.append("Số tờ")
        if not thua:
            missing.append("Số thửa")

        if missing:
            reason = "Không nhận diện đủ: " + ", ".join(missing)
            need_check.append({"old_name": fname, "new_name": "", "maxa": maxa or "", "nguon_maxa": nguon or "",
                               "tu_khoa_nhan_dien": tu_khoa, "to": to or "", "thua": thua or "",
                               "status": "CAN_KIEM_TRA", "note": reason, "ocr_short": ocr_short})
            log_cb(f"   ✗ CAN_KIEM_TRA: {reason}")
            if ap_dung_toi_uu:
                _wperf.collect_garbage()
                time.sleep(sleep_per_file)
                if so_file_moi_dot and (i + 1) % so_file_moi_dot == 0 and (i + 1) < len(files):
                    log_cb(f"⏸ Đã xử lý {i + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                    time.sleep(nghi_giua_dot_giay)
            if checkpoint_task_id and do_rename:
                from app.services import checkpoint_service as _ckpt
                da_xu_ly_toi_idx = set(files[:i + 1]) | (files_da_xong or set())
                _ckpt.luu_checkpoint(
                    checkpoint_task_id, "ocr_doi_ten_pdf", files,
                    list(da_xu_ly_toi_idx), [], [], i + 1,
                    {"toc_do_xu_ly": toc_do_xu_ly}, "")
            continue

        new_name = f"CHUACOGIAY_{maxa}_{to}_{thua}_{name_suffix}.pdf"
        existing = existing_by_dir[dirpath]
        base, ext = os.path.splitext(new_name)
        idx = 1
        final_name = new_name
        conflict = False
        while final_name in existing and final_name != fname:
            final_name = f"{base}_{idx}{ext}"
            idx += 1
            conflict = True
        existing.discard(fname)
        existing.add(final_name)

        status = "TRUNG_TEN_FILE" if conflict else ("DA_DOI_TEN" if do_rename else "DA_NHAN_DIEN")
        note = "Đã thêm hậu tố tránh trùng tên" if conflict else ""

        rec = {"old_path": fpath, "old_name": fname, "new_name": final_name,
               "maxa": maxa, "to": to, "thua": thua, "nguon_maxa": nguon or "",
               "tu_khoa_nhan_dien": tu_khoa, "status": status, "note": note, "ocr_short": ocr_short}
        ok.append(rec)
        log_cb(f"   ✓ {status}: {nguon} → Mã {maxa}, Tờ {to}, Thửa {thua} ⇒ {final_name}")

        if do_rename:
            try:
                os.rename(fpath, os.path.join(dirpath, final_name))
            except Exception as e:
                log_cb(f"   ⚠ Lỗi khi đổi tên: {e}")

        if ap_dung_toi_uu:
            _wperf.collect_garbage()
            time.sleep(sleep_per_file)
            if so_file_moi_dot and (i + 1) % so_file_moi_dot == 0 and (i + 1) < len(files):
                log_cb(f"⏸ Đã xử lý {i + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                time.sleep(nghi_giua_dot_giay)
        if checkpoint_task_id and do_rename:
            from app.services import checkpoint_service as _ckpt
            da_xu_ly_toi_idx = set(files[:i + 1]) | (files_da_xong or set())
            _ckpt.luu_checkpoint(
                checkpoint_task_id, "ocr_doi_ten_pdf", files,
                list(da_xu_ly_toi_idx), [], [], i + 1,
                {"toc_do_xu_ly": toc_do_xu_ly}, "")

    if ap_dung_toi_uu:
        _wperf.restore_normal_priority()

    if checkpoint_task_id and do_rename and not bi_huy_giua_chung:
        from app.services import checkpoint_service as _ckpt
        _ckpt.xoa_checkpoint(checkpoint_task_id)

    return {"ok": ok, "need_check": need_check}


def write_ocr_rename_csv(csv_path, result):
    """Xuất CSV: Tên file cũ, Tên file mới dự kiến, Mã xã, Nguồn nhận diện mã xã, Từ khóa nhận diện,
    Số tờ, Số thửa, Trạng thái, Ghi chú, Text OCR rút gọn."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Tên file cũ", "Tên file mới dự kiến", "Mã xã", "Nguồn nhận diện mã xã",
                    "Từ khóa nhận diện", "Số tờ", "Số thửa", "Trạng thái", "Ghi chú", "Text OCR rút gọn"])
        for r in result["ok"]:
            w.writerow([r["old_name"], r["new_name"], r["maxa"], r.get("nguon_maxa", ""),
                        r.get("tu_khoa_nhan_dien", ""), r["to"], r["thua"], r["status"], r["note"], r["ocr_short"]])
        for r in result["need_check"]:
            w.writerow([r["old_name"], r.get("new_name", ""), r.get("maxa", ""), r.get("nguon_maxa", ""),
                        r.get("tu_khoa_nhan_dien", ""), r.get("to", ""), r.get("thua", ""),
                        r["status"], r["note"], r.get("ocr_short", "")])


def write_ocr_rename_cmd(cmd_path, result):
    """Xuất file .bat: các dòng ren "tên cũ" "tên mới" (đường dẫn đầy đủ)."""
    with open(cmd_path, "w", encoding="utf-8-sig") as f:
        f.write("@echo off\n")
        f.write("chcp 65001 >nul\n\n")
        for r in result["ok"]:
            old_full = r["old_path"].replace('"', '')
            new_name = r["new_name"].replace('"', '')
            f.write(f'ren "{old_full}" "{new_name}"\n')
        f.write("\necho Da doi ten xong. Nhan phim bat ky de dong.\npause >nul\n")


try:
    from app.services.backup_service import (
        create_undo_script_rename as _bs_create_undo_script_rename,
        undo_renames_in_app as _bs_undo_renames_in_app,
        create_backup,
    )
    _BACKUP_SERVICE_MODULE_OK = True
except Exception:
    _BACKUP_SERVICE_MODULE_OK = False

    def create_backup(files, backup_dir, run_id):
        """Sao lưu dự phòng tại chỗ nếu thiếu module app/services/backup_service.py - copy các file
        vào backup_dir/BACKUP_<run_id>/, giữ nguyên tên gốc, KHÔNG xóa/sửa file gốc."""
        backup_folder = os.path.join(backup_dir, f"BACKUP_{run_id}")
        os.makedirs(backup_folder, exist_ok=True)
        done, errors = [], []
        for f in files:
            if not os.path.isfile(f):
                errors.append((f, "File không tồn tại"))
                continue
            dest = os.path.join(backup_folder, os.path.basename(f))
            base, ext = os.path.splitext(dest)
            i = 1
            while os.path.isfile(dest):
                dest = f"{base}_{i}{ext}"
                i += 1
            try:
                shutil.copy2(f, dest)
                done.append((f, dest))
            except Exception as e:
                errors.append((f, str(e)))
        return backup_folder, done, errors


try:
    from app.services.log_service import (
        generate_run_id, create_run_log as _create_standard_run_log, LOG_HEADER_STANDARD,
    )
    _LOG_SERVICE_MODULE_OK = True
except Exception:
    _LOG_SERVICE_MODULE_OK = False
    LOG_HEADER_STANDARD = ["Time", "Module", "Action", "SourceFile", "OutputFile",
                           "MaXa", "SoTo", "SoThua", "Key", "Status", "Message", "Note"]

    def generate_run_id():
        return time.strftime("%Y%m%d_%H%M%S")

    def _create_standard_run_log(logs_dir, module_name, run_id=None):
        """Logger dự phòng tại chỗ nếu thiếu module app/services/log_service.py - CÙNG giao diện
        (.add()/.log_error()/.save()) và CÙNG định dạng cột chuẩn để không ảnh hưởng nơi gọi."""
        import csv as _csv

        class _FallbackRunLogger:
            def __init__(self):
                self.run_id = run_id or generate_run_id()
                self.entries = []

            def add(self, action="", source_file="", output_file="", ma_xa=None, so_to=None,
                    so_thua=None, key="", status="", message="", note=""):
                self.entries.append([
                    time.strftime("%Y-%m-%d %H:%M:%S"), module_name, action, source_file, output_file,
                    ma_xa or "", so_to or "", so_thua or "", key, status, message, note,
                ])

            def log_error(self, source_file, message, action="ERROR"):
                self.add(action=action, source_file=source_file, status="LOI", message=message)

            def save(self):
                os.makedirs(logs_dir, exist_ok=True)
                fname = f"{self.run_id}_{module_name.strip().upper().replace(' ', '_')}.csv"
                path = os.path.join(logs_dir, fname)
                with open(path, "w", encoding="utf-8-sig", newline="") as f:
                    w = _csv.writer(f)
                    w.writerow(LOG_HEADER_STANDARD)
                    for row in self.entries:
                        w.writerow(row)
                return path

        return _FallbackRunLogger()


def get_standard_logs_dir():
    """Thư mục Logs chuẩn dùng chung cho log_service (nằm trong dữ liệu ứng dụng, KHÔNG phải nơi
    người dùng chọn xuất kết quả) - để có 1 nơi tập trung xem lại lịch sử thao tác của mọi Tab."""
    return os.path.join(get_app_data_dir(), "Logs")


def write_undo_rename_cmd(cmd_path, result):
    """
    Xuất file UNDO_*.cmd chứa lệnh ren NGƯỢC LẠI (tên mới -> tên cũ), dùng để hoàn tác thủ công
    nếu cần (kể cả khi đã đóng phần mềm). Chỉ tạo lệnh cho các file ĐÃ THỰC SỰ đổi tên thành công.
    Là 1 ADAPTER MỎNG gọi app/services/backup_service.py (Giai đoạn 2 tái kiến trúc) - có fallback
    tại chỗ nếu thiếu module, để không bao giờ làm vỡ chức năng.
    """
    rename_pairs = [(r["old_path"], r["new_name"]) for r in result["ok"]]
    if _BACKUP_SERVICE_MODULE_OK:
        undo_dir = os.path.dirname(cmd_path) or "."
        run_id = os.path.splitext(os.path.basename(cmd_path))[0].replace("UNDO_", "")
        generated_path = _bs_create_undo_script_rename(rename_pairs, undo_dir, run_id)
        if generated_path != cmd_path and os.path.isfile(generated_path):
            os.replace(generated_path, cmd_path)
        return

    with open(cmd_path, "w", encoding="utf-8-sig") as f:
        f.write("@echo off\n")
        f.write("chcp 65001 >nul\n\n")
        f.write("echo HOAN TAC DOI TEN - dua file ve ten cu\n\n")
        for r in result["ok"]:
            dirpath = os.path.dirname(r["old_path"]) or "."
            new_full = os.path.join(dirpath, r["new_name"]).replace('"', '')
            old_name = os.path.basename(r["old_path"]).replace('"', '')
            f.write(f'ren "{new_full}" "{old_name}"\n')
        f.write("\necho Da hoan tac xong. Nhan phim bat ky de dong.\npause >nul\n")


def undo_rename_in_app(result, log_cb):
    """
    Hoàn tác NGAY trong phần mềm (không cần chạy file .bat riêng): đổi tên các file đã đổi tên
    thành công trở về đúng tên cũ ban đầu. Chỉ áp dụng cho các mục có status DA_DOI_TEN (đã đổi
    tên thật) - bỏ qua các mục chỉ ở chế độ xem trước (DA_NHAN_DIEN/CHAY_THU).
    Là 1 ADAPTER MỎNG gọi app/services/backup_service.py (Giai đoạn 2 tái kiến trúc) - có fallback
    tại chỗ nếu thiếu module.
    Trả về (so_da_hoan_tac, danh_sach_loi).
    """
    rename_pairs = [(r["old_path"], r["new_name"]) for r in result["ok"] if r.get("status") == "DA_DOI_TEN"]

    if _BACKUP_SERVICE_MODULE_OK:
        return _bs_undo_renames_in_app(rename_pairs, log_cb)

    undone = 0
    errors = []
    for old_path, new_name in rename_pairs:
        dirpath = os.path.dirname(old_path) or "."
        current_path = os.path.join(dirpath, new_name)
        original_name = os.path.basename(old_path)
        original_path = os.path.join(dirpath, original_name)
        if not os.path.isfile(current_path):
            errors.append(f"{new_name}: không tìm thấy file để hoàn tác (có thể đã bị di chuyển/xóa).")
            continue
        if os.path.isfile(original_path) and os.path.abspath(original_path) != os.path.abspath(current_path):
            errors.append(f"{new_name}: không hoàn tác được vì đã có file trùng tên gốc '{original_name}'.")
            continue
        try:
            os.rename(current_path, original_path)
            undone += 1
            log_cb(f"↩ Đã hoàn tác: {new_name} → {original_name}")
        except Exception as e:
            errors.append(f"{new_name}: lỗi khi hoàn tác - {e}")
    return undone, errors


def build_rename_commands_from_pdf_content(source, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb,
                                            debug=False, name_template="CHUACOGIAY_{maxa}_{to}_{thua}_{loai}",
                                            control=None):
    """
    Đọc nội dung PDF từ `source` (1 thư mục HOẶC danh sách file cụ thể, có thể ở nhiều thư mục khác nhau).
    name_template: mẫu tên mới do người dùng tự đặt, dùng placeholder {maxa} {to} {thua} {loai} {ten}
    {dt} {id} {mucdich} {core} {year} {original}.
    Trả về:
    - ok_list: danh sách (đường_dẫn_đầy_đủ_cũ, tên_mới) cho các file đọc được đủ thông tin
    - err_list: danh sách (tên_file, lý_do_lỗi) cho các file lỗi/thiếu thông tin
    """
    files = list_files_from_source(source, ".pdf")
    ok_list = []
    err_list = []
    existing_by_dir = {}

    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        fname = os.path.basename(fpath)
        dirpath = os.path.dirname(fpath) or "."
        if dirpath not in existing_by_dir:
            try:
                existing_by_dir[dirpath] = set(os.listdir(dirpath))
            except OSError:
                existing_by_dir[dirpath] = set()

        log_cb(f"→ Đang đọc: {fname}")
        info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, debug=debug)
        if "error" in info:
            err_list.append((fname, info["error"]))
            log_cb(f"   ✗ {info['error']}")
            if debug and info.get("raw_text"):
                log_cb("   ---- Văn bản đã đọc được (để chẩn đoán) ----")
                log_cb("   " + info["raw_text"].replace("\n", "\n   "))
                log_cb("   ---- Hết văn bản ----")
            continue

        try:
            new_name = render_template(name_template, info, info["loai"], fname)
        except ValueError as e:
            err_list.append((fname, str(e)))
            log_cb(f"   ✗ {e}")
            continue

        base, ext = os.path.splitext(new_name)
        i = 2
        existing = existing_by_dir[dirpath]
        while new_name in existing and new_name != fname:
            new_name = f"{base}_{i}{ext}"
            i += 1
        existing.discard(fname)
        existing.add(new_name)

        nguon_txt = "OCR (bản scan)" if info.get("nguon") == "ocr" else "chữ có sẵn"
        via_txt = f"Thôn {info.get('thon_name')}" if info.get("nguon_maxa") == "thon" else f"Xã {info.get('xa_name')}"
        log_cb(f"   ✓ {via_txt} (Mã {info['maxa']}) - Tờ {info['to']} - Thửa {info['thua']} "
               f"- Loại {info['loai']} - đọc bằng {nguon_txt}")
        if new_name != fname:
            ok_list.append((fpath, new_name))
        else:
            log_cb("   = Tên đã đúng chuẩn, không cần đổi.")

    return ok_list, err_list


def write_ren_bat_file(ok_list, bat_path):
    """ok_list: danh sách (đường_dẫn_đầy_đủ_cũ, tên_mới). Không cần cd vào thư mục vì dùng đường dẫn đầy đủ."""
    with open(bat_path, "w", encoding="utf-8-sig") as f:
        f.write("@echo off\n")
        f.write("chcp 65001 >nul\n\n")
        for old_full, new_name in ok_list:
            old_escaped = old_full.replace('"', '')
            new_escaped = new_name.replace('"', '')
            f.write(f'ren "{old_escaped}" "{new_escaped}"\n')
        f.write("\necho Da doi ten xong. Nhan phim bat ky de dong.\npause >nul\n")


# ============================================================================
# MỤC 5 (MỚI): TÁCH TRANG / GỘP FILE PDF
# ============================================================================

try:
    from app.services.pdf_service import (
        parse_page_spec,
        split_pdf_pages,
        split_pdf_into_fixed_groups,
        merge_pdfs,
    )
    _PDF_SERVICE_MODULE_OK_2 = True
except Exception:
    _PDF_SERVICE_MODULE_OK_2 = False

    def parse_page_spec(spec, total_pages):
        """
        Chuyển chuỗi trang (VD: '1', '2', '1,2', '1-3', '1,3-5') thành list số trang (1-based).
        """
        pages = []
        spec = spec.strip()
        if not spec:
            raise ValueError("Chưa nhập trang cần lấy.")
        for part in spec.split(","):
            part = part.strip()
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                a, b = int(a), int(b)
                pages.extend(range(a, b + 1))
            else:
                pages.append(int(part))
        pages = [p for p in pages if 1 <= p <= total_pages]
        if not pages:
            raise ValueError(f"Không có trang hợp lệ trong tổng số {total_pages} trang.")
        return pages

    def split_pdf_pages(input_path, page_spec, output_path):
        if not HAS_PYPDF:
            raise RuntimeError("Thiếu thư viện pypdf. Cài: pip install pypdf")
        reader = PdfReader(input_path)
        pages = parse_page_spec(page_spec, len(reader.pages))
        writer = PdfWriter()
        for p in pages:
            writer.add_page(reader.pages[p - 1])
        with open(output_path, "wb") as f:
            writer.write(f)
        return len(pages)

    def split_pdf_into_fixed_groups(input_path, group_size, output_folder, base_name=None, log_cb=None):
        """
        Tách 1 file PDF (có thể nhiều trang) thành NHIỀU file nhỏ, mỗi file gồm `group_size`
        trang liên tiếp (VD group_size=2: trang 1+2 -> file 1, trang 3+4 -> file 2, ...).
        Trả về số file đã tạo.
        """
        if not HAS_PYPDF:
            raise RuntimeError("Thiếu thư viện pypdf. Cài: pip install pypdf")
        if group_size < 1:
            raise ValueError("Số trang mỗi file phải lớn hơn hoặc bằng 1.")

        reader = PdfReader(input_path)
        total = len(reader.pages)
        if total == 0:
            raise RuntimeError("File PDF không có trang nào.")

        os.makedirs(output_folder, exist_ok=True)
        base_name = base_name or os.path.splitext(os.path.basename(input_path))[0]

        count = 0
        for start in range(0, total, group_size):
            end = min(start + group_size, total)
            writer = PdfWriter()
            for p in range(start, end):
                writer.add_page(reader.pages[p])
            count += 1
            out_name = f"{base_name}_phan{count:03d}_trang{start + 1}-{end}.pdf"
            out_path = os.path.join(output_folder, out_name)
            with open(out_path, "wb") as f:
                writer.write(f)
            if log_cb:
                log_cb(f"✓ Đã tạo: {out_name}  (trang {start + 1}-{end})")

        return count

    def merge_pdfs(file_list, output_path):
        if not HAS_PYPDF:
            raise RuntimeError("Thiếu thư viện pypdf. Cài: pip install pypdf")
        writer = PdfWriter()
        for fpath in file_list:
            reader = PdfReader(fpath)
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)


def batch_split_folder(source, page_spec, suffix, output_folder, log_cb, control=None):
    if not HAS_PYPDF:
        raise RuntimeError("Thiếu thư viện pypdf. Cài: pip install pypdf")
    os.makedirs(output_folder, exist_ok=True)
    files = list_files_from_source(source, ".pdf")
    count = 0
    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        fname = os.path.basename(fpath)
        base = os.path.splitext(fname)[0]
        out_path = os.path.join(output_folder, f"{base}{suffix}.pdf")
        try:
            n = split_pdf_pages(fpath, page_spec, out_path)
            log_cb(f"✓ {fname}  →  {os.path.basename(out_path)}  ({n} trang)")
            count += 1
        except Exception as e:
            log_cb(f"✗ {fname}: {e}")
    return count


# ============================================================================
# MỤC 10 (MỚI): LỌC & DI CHUYỂN FILE PDF TRÙNG THEO MÃ XÃ - TỜ - THỬA
# ============================================================================

KEY_FROM_FILENAME_RE = re.compile(r"(\d{4,6})[_\-](\d{1,6})[_\-](\d{1,6})")


try:
    from app.services.file_compare_service import extract_key_from_filename, find_and_move_duplicate_pdfs
    _FILE_COMPARE_SERVICE_MODULE_OK_A = True
except Exception:
    _FILE_COMPARE_SERVICE_MODULE_OK_A = False

    def extract_key_from_filename(fname):
        """
        Tách (Mã xã, Số tờ, Số thửa) từ tên file ĐÃ ĐẶT TÊN ĐÚNG CHUẨN (VD CHUACOGIAY_02140_111_13_GT.pdf),
        bất kể phần đuôi/hậu tố phía sau khác nhau (_GT, _TBXN, hay bất kỳ chữ gì).
        Trả về tuple (maxa, to_norm, thua_norm) hoặc None nếu không tách được.
        """
        m = KEY_FROM_FILENAME_RE.search(os.path.splitext(fname)[0])
        if not m:
            return None
        maxa, to, thua = m.group(1), m.group(2), m.group(3)
        return (maxa, norm_num(to), norm_num(thua))

    def find_and_move_duplicate_pdfs(folder_reference, folder_compare, output_folder, log_cb, control=None):
        """
        So sánh file PDF trong `folder_compare` với `folder_reference` theo khóa (Mã xã, Tờ, Thửa)
        tách từ TÊN FILE (bỏ qua phần hậu tố khác nhau ở cuối tên).
        Nếu tìm thấy khóa trùng, DI CHUYỂN file đó từ folder_compare sang output_folder.

        Trả về dict:
        - moved: [(tên file, khóa, tên file khớp trong thư mục tham chiếu)]
        - unmatched: [tên file trong folder_compare không trùng khóa nào]
        - unparsed_reference: [tên file trong folder_reference không tách được khóa - bị bỏ qua]
        - unparsed_compare: [tên file trong folder_compare không tách được khóa - bị bỏ qua, không di chuyển]
        """
        os.makedirs(output_folder, exist_ok=True)

        ref_files = [f for f in sorted(os.listdir(folder_reference)) if f.lower().endswith(".pdf")]
        cmp_files = [f for f in sorted(os.listdir(folder_compare)) if f.lower().endswith(".pdf")]

        ref_keys = {}
        unparsed_reference = []
        for f in ref_files:
            key = extract_key_from_filename(f)
            if key is None:
                unparsed_reference.append(f)
                continue
            ref_keys.setdefault(key, []).append(f)

        if unparsed_reference:
            log_cb(f"⚠ {len(unparsed_reference)} file trong thư mục tham chiếu không tách được Mã xã/Tờ/Thửa "
                   f"từ tên (bỏ qua khi so sánh): {', '.join(unparsed_reference)}")

        moved = []
        unmatched = []
        unparsed_compare = []

        for i, f in enumerate(cmp_files):
            if control:
                try:
                    control.checkpoint()
                except TaskCancelled:
                    log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(cmp_files) - i} file chưa xử lý).")
                    break
            key = extract_key_from_filename(f)
            if key is None:
                unparsed_compare.append(f)
                log_cb(f"⚠ Không tách được Mã xã/Tờ/Thửa từ tên file, bỏ qua: {f}")
                continue

            if key in ref_keys:
                src = os.path.join(folder_compare, f)
                dest = os.path.join(output_folder, f)
                base, ext = os.path.splitext(dest)
                i = 2
                while os.path.isfile(dest):
                    dest = f"{base}__trung_{i}{ext}"
                    i += 1
                shutil.move(src, dest)
                matched_names = ", ".join(ref_keys[key])
                moved.append((f, key, matched_names))
                log_cb(f"✓ TRÙNG (Mã xã {key[0]}, Tờ {key[1]}, Thửa {key[2]}) — đã di chuyển: {f}  "
                       f"(khớp với: {matched_names})")
            else:
                unmatched.append(f)

        return {
            "moved": moved,
            "unmatched": unmatched,
            "unparsed_reference": unparsed_reference,
            "unparsed_compare": unparsed_compare,
        }

# ============================================================================
# MỤC 11 (MỚI): LỌC EXCEL TỔNG THEO THƯ MỤC PDF & HOÀN THIỆN BÁO CÁO
# ============================================================================

try:
    from app.services.file_compare_service import extract_key_from_filename_v2
    _FILE_COMPARE_SERVICE_MODULE_OK_B = True
except Exception:
    _FILE_COMPARE_SERVICE_MODULE_OK_B = False

    def extract_key_from_filename_v2(fname):
        """
        Tách khóa đối chiếu từ tên file PDF, chấp nhận nhiều dạng tên khác nhau:
        - CHUACOGIAY_02140_29_199_GT.pdf / _TBXN.pdf / (không hậu tố).pdf
        - 02140_29_199.pdf (không tiền tố CHUACOGIAY)
        - 29_199.pdf (KHÔNG có mã xã, chỉ có Số tờ + Số thửa)
        Trả về dict {"maxa": str|None, "to": str, "thua": str} hoặc None nếu không tách được gì.
        """
        base = os.path.splitext(fname)[0]
        m3 = KEY_FROM_FILENAME_RE.search(base)
        if m3:
            return {"maxa": m3.group(1), "to": norm_num(m3.group(2)), "thua": norm_num(m3.group(3))}
        m2 = re.fullmatch(r"[^\d]*(\d{1,6})[_\-](\d{1,6})[^\d]*", base)
        if m2:
            return {"maxa": None, "to": norm_num(m2.group(1)), "thua": norm_num(m2.group(2))}
        return None


RECONCILE_PDF_COLS = ["B", "H", "I", "J", "K", "L", "V", "W", "X", "Y", "Z", "AA"]


def build_pdf_index_for_reconcile(source_pdfs, include_subfolders, xa_mapping, thon_mapping,
                                   use_ocr_fallback, ocr_dpi, log_cb, control=None, need_content=False):
    """
    Đọc toàn bộ PDF (từ thư mục, có thể gồm thư mục con), xác định khóa đối chiếu
    (Mã xã, Tờ, Thửa) theo thứ tự ưu tiên: tên file trước, OCR nội dung sau nếu tên
    file không đủ thông tin. Nếu need_content=True (bật "hoàn thiện dữ liệu từ PDF"),
    LUÔN đọc thêm nội dung PDF để lấy Tên/CCCD/Địa chỉ/Diện tích/Mục đích, kể cả khi
    khóa đối chiếu đã xác định đủ từ tên file.
    Trả về:
    - idx_with_maxa: {(maxa,to,thua): entry}
    - idx_to_thua_only: {(to,thua): entry}  (dùng khi PDF không xác định được mã xã)
    - need_check: [ {file, reason} ] - các PDF không tách được cả tên lẫn OCR
    entry = {"file": tên_file, "path": đường_dẫn, "maxa","to","thua","ten","cccd","dia_chi",
             "dia_chi_thuong_tru","dt","mucdich", "nguon": "ten_file"|"ocr"}
    """
    if include_subfolders:
        files = []
        for root, dirs, fnames in os.walk(source_pdfs):
            for f in fnames:
                if f.lower().endswith(".pdf"):
                    files.append(os.path.join(root, f))
        files.sort()
    else:
        files = list_files_from_source(source_pdfs, ".pdf")

    idx_with_maxa = {}
    idx_to_thua_only = {}
    need_check = []

    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file PDF chưa đọc).")
                break

        fname = os.path.basename(fpath)
        key = extract_key_from_filename_v2(fname)
        entry = {"file": fname, "path": fpath, "maxa": None, "to": None, "thua": None,
                 "ten": None, "cccd": None, "dia_chi": None, "dia_chi_thuong_tru": None,
                 "dt": None, "mucdich": None, "nguon": None}

        if key and key["to"] and key["thua"]:
            entry.update(maxa=key["maxa"], to=key["to"], thua=key["thua"], nguon="tên file")
            log_cb(f"→ {fname}: khóa từ TÊN FILE = (Mã xã {key['maxa']}, Tờ {key['to']}, Thửa {key['thua']})")
            if need_content:
                info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb)
                if "error" not in info:
                    entry.update(ten=info.get("ten"), cccd=info.get("id"),
                                dia_chi=info.get("dia_chi"), dia_chi_thuong_tru=info.get("dia_chi_thuong_tru"),
                                dt=info.get("dt"), mucdich=info.get("mucdich"))
                    if not entry["maxa"]:
                        entry["maxa"] = info.get("maxa")
                    log_cb(f"   ✓ Đã đọc thêm nội dung để hoàn thiện: Tên={entry['ten']}")
                else:
                    log_cb(f"   ⚠ Không đọc được nội dung chi tiết ({info['error']}) - vẫn giữ khóa từ tên file")
        else:
            log_cb(f"→ {fname}: tên file không đủ khóa, đang đọc nội dung (OCR nếu cần)...")
            try:
                info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb)
            except Exception as e:
                need_check.append({"file": fname, "reason": f"Lỗi đọc file PDF: {e}", "status": "LOI_FILE_PDF"})
                log_cb(f"   ✗ LOI_FILE_PDF: {e}")
                continue

            if "error" in info:
                err_text = info["error"]
                if not info.get("to") or not info.get("thua"):
                    status = "CAN_KIEM_TRA_THIEU_SO_TO_SO_THUA"
                elif "mã xã" in err_text.lower():
                    status = "CAN_KIEM_TRA_THIEU_MA_XA"
                else:
                    status = "CAN_KIEM_TRA_OCR_YEU"
                need_check.append({"file": fname, "reason": err_text, "status": status})
                log_cb(f"   ✗ {status}: {err_text}")
                continue
            entry.update(maxa=info.get("maxa"), to=info["to"], thua=info["thua"],
                        ten=info.get("ten"), cccd=info.get("id"),
                        dia_chi=info.get("dia_chi"), dia_chi_thuong_tru=info.get("dia_chi_thuong_tru"),
                        dt=info.get("dt"), mucdich=info.get("mucdich"), nguon="OCR/nội dung")
            log_cb(f"   ✓ Khóa từ NỘI DUNG = (Mã xã {info.get('maxa')}, Tờ {info['to']}, Thửa {info['thua']})")

        if entry["maxa"]:
            idx_with_maxa[(str(entry["maxa"]), str(entry["to"]), str(entry["thua"]))] = entry
        idx_to_thua_only.setdefault((str(entry["to"]), str(entry["thua"])), entry)

    return idx_with_maxa, idx_to_thua_only, need_check


def reconcile_excel_with_pdfs(excel_path, source_pdfs, sheet_name, header_row,
                               col_maxa, col_to, col_thua, col_stt,
                               include_subfolders, complete_from_pdf, allow_overwrite,
                               renumber_stt, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi,
                               dry_run, output_path, log_cb, control=None):
    """
    Đối chiếu từng dòng Excel tổng với dữ liệu PDF trong thư mục:
    - Dòng KHÔNG có PDF tương ứng -> xóa khỏi báo cáo (trạng thái XOA_KHONG_CO_PDF)
    - Dòng CÓ PDF tương ứng -> giữ lại (GIU_LAI_DA_KHOP), nếu complete_from_pdf=True thì điền/hoàn
      thiện dữ liệu vào các cột B,H,I,J,K,L,V,W,X,Y,Z,AA (GIU_LAI_CAP_NHAT). Mặc định KHÔNG ghi đè ô
      đã có dữ liệu trừ khi allow_overwrite=True.
    dry_run=True: không xóa/không lưu file, chỉ trả về log để xem trước.
    Trả về dict: {"log_rows": [...], "deleted_rows": [...], "pdf_need_check": [...],
                  "kept_count": int, "deleted_count": int}
    """
    log_cb("Bước 1/3: Đọc & lập chỉ mục dữ liệu PDF...")
    idx_with_maxa, idx_to_thua_only, pdf_need_check = build_pdf_index_for_reconcile(
        source_pdfs, include_subfolders, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, control,
        need_content=complete_from_pdf)
    log_cb(f"  → {len(idx_with_maxa) + len(idx_to_thua_only)} PDF lập chỉ mục được, "
           f"{len(pdf_need_check)} PDF cần kiểm tra (không đọc được).")

    log_cb("Bước 2/3: Đối chiếu từng dòng Excel...")
    wb = openpyxl.load_workbook(excel_path)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    log_rows = []
    deleted_rows = []
    rows_to_delete = []
    kept_count = 0

    for r in range(header_row, ws.max_row + 1):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu ở dòng {r}.")
                break

        maxa_val = ws[f"{col_maxa}{r}"].value if col_maxa else None
        to_val = ws[f"{col_to}{r}"].value if col_to else None
        thua_val = ws[f"{col_thua}{r}"].value if col_thua else None

        if to_val in (None, "") and thua_val in (None, "") and maxa_val in (None, ""):
            continue  # dòng trống hoàn toàn -> bỏ qua, không tính là "xóa"

        maxa_norm = str(maxa_val).strip() if maxa_val not in (None, "") else None
        to_norm = norm_num(str(to_val).strip()) if to_val not in (None, "") else None
        thua_norm = norm_num(str(thua_val).strip()) if thua_val not in (None, "") else None

        match = None
        warn = ""
        if maxa_norm and to_norm and thua_norm and (maxa_norm, to_norm, thua_norm) in idx_with_maxa:
            match = idx_with_maxa[(maxa_norm, to_norm, thua_norm)]
        elif to_norm and thua_norm and (to_norm, thua_norm) in idx_to_thua_only:
            match = idx_to_thua_only[(to_norm, thua_norm)]
            if match.get("maxa") is None:
                warn = ("Đối chiếu không có mã xã, cần kiểm tra nếu có khả năng trùng số tờ số thửa "
                       "giữa các xã cũ.")

        if not match:
            rows_to_delete.append(r)
            deleted_rows.append({"dong": r, "maxa": maxa_norm or "", "to": to_norm or "", "thua": thua_norm or ""})
            log_rows.append({"dong": r, "maxa": maxa_norm or "", "to": to_norm or "", "thua": thua_norm or "",
                            "status": "XOA_KHONG_CO_PDF", "note": "Không có file PDF tương ứng"})
            continue

        kept_count += 1
        status = "GIU_LAI_DA_KHOP"
        note = warn

        if complete_from_pdf and not dry_run:
            values_map = {
                "B": match.get("maxa"), "H": match.get("ten"), "I": match.get("cccd"),
                "L": match.get("dia_chi_thuong_tru"), "V": match.get("to"), "W": match.get("thua"),
                "X": match.get("dia_chi"), "Y": match.get("dt"), "Z": match.get("mucdich"),
                "AA": match.get("dt"),
            }
            updated_any = False
            for col, new_val in values_map.items():
                if new_val in (None, ""):
                    continue
                cell = ws[f"{col}{r}"]
                if cell.value in (None, "") or allow_overwrite:
                    cell.value = new_val
                    updated_any = True
            if updated_any:
                status = "GIU_LAI_CAP_NHAT"

        log_rows.append({"dong": r, "maxa": maxa_norm or (match.get("maxa") or ""), "to": to_norm or "",
                        "thua": thua_norm or "", "status": status, "note": note or f"Khớp PDF: {match['file']}"})

    if renumber_stt and col_stt and not dry_run:
        new_stt = 1
        for r in range(header_row, ws.max_row + 1):
            if r in rows_to_delete:
                continue
            cell = ws[f"{col_stt}{r}"]
            if cell.value not in (None, ""):
                cell.value = new_stt
                new_stt += 1

    if not dry_run and rows_to_delete:
        log_cb(f"Bước 3/3: Xóa {len(rows_to_delete)} dòng không khớp (từ dưới lên)...")
        for r in sorted(rows_to_delete, reverse=True):
            ws.delete_rows(r, 1)
    else:
        log_cb("Bước 3/3: (Chạy thử - không xóa dòng nào thật sự)")

    if not dry_run:
        try:
            from app.services import file_safety_service as _fsafe
            _fsafe.ghi_nguyen_tu(wb.save, output_path)
            log_cb(f"✓ Đã lưu báo cáo: {output_path}")
        except Exception as e:
            log_cb(f"✗ LOI_GHI_EXCEL: Không lưu được file kết quả: {e}")
            raise RuntimeError(f"LOI_GHI_EXCEL: Không lưu được file kết quả - {e}")

    return {
        "log_rows": log_rows,
        "deleted_rows": deleted_rows,
        "pdf_need_check": pdf_need_check,
        "kept_count": kept_count,
        "deleted_count": len(rows_to_delete),
    }


def write_reconcile_logs(log_dir, stamp, result):
    import csv
    paths = {}

    p1 = os.path.join(log_dir, f"DONG_EXCEL_BI_XOA_{stamp}.csv")
    with open(p1, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Dòng Excel", "Mã xã", "Số tờ", "Số thửa"])
        for r in result["deleted_rows"]:
            w.writerow([r["dong"], r["maxa"], r["to"], r["thua"]])
    paths["deleted_csv"] = p1

    p2 = os.path.join(log_dir, f"PDF_CAN_KIEM_TRA_{stamp}.csv")
    with open(p2, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Tên file PDF", "Trạng thái", "Lý do"])
        for r in result["pdf_need_check"]:
            w.writerow([r["file"], r.get("status", "CAN_KIEM_TRA_OCR_YEU"), r["reason"]])
    paths["pdf_check_csv"] = p2

    p3 = os.path.join(log_dir, f"LOG_HOAN_THIEN_BAO_CAO_{stamp}.csv")
    with open(p3, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Dòng Excel", "Mã xã", "Số tờ", "Số thửa", "Trạng thái", "Ghi chú"])
        for r in result["log_rows"]:
            w.writerow([r["dong"], r["maxa"], r["to"], r["thua"], r["status"], r["note"]])
    paths["full_log_csv"] = p3

    return paths


# ============================================================================
# MỤC 12 (MỚI): CHỈNH SỬA PDF HÀNG LOẠT — XOAY, XEM TRƯỚC, CẢI THIỆN ẢNH SCAN
# ============================================================================

try:
    import cv2
    import numpy as np
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False


# ============================================================================
# TỰ NHẬN DIỆN HƯỚNG TRANG + XOAY + LÀM THẲNG PDF (mục III-V tài liệu "LỆNH PHÁT
# TRIỂN TAB 6 – MỤC XOAY PDF: TỰ NHẬN DIỆN HƯỚNG TRANG, XOAY THẲNG...")
# ============================================================================

def detect_page_rotation_osd(pil_image, log_cb=None):
    """
    Nhận diện góc xoay 0/90/180/270 độ bằng Tesseract OSD (Orientation and Script Detection) -
    mục IV bước 3 tài liệu. KHÔNG chỉ dựa vào rotation metadata sẵn có của PDF (nhiều file scan
    có ảnh đã bị xoay nhưng metadata vẫn bằng 0 - đúng cảnh báo mục III tài liệu).

    Trả về (goc_can_xoay, do_tin_cay_tho) - goc_can_xoay là 0/90/180/270 (độ cần xoay THUẬN CHIỀU
    KIM ĐỒNG HỒ để đưa trang về đúng hướng), hoặc (None, 0) nếu OSD không nhận diện được (VD
    trang quá ít chữ, chủ yếu là ảnh/bản đồ/sơ đồ - mục VIII tài liệu).
    """
    if not HAS_OCR:
        return None, 0
    try:
        osd = pytesseract.image_to_osd(pil_image)
        rotate_match = re.search(r"Rotate:\s*(\d+)", osd)
        conf_match = re.search(r"Orientation confidence:\s*([\d.]+)", osd)
        if not rotate_match:
            return None, 0
        goc = int(rotate_match.group(1))
        conf = float(conf_match.group(1)) if conf_match else 0
        return goc, conf
    except Exception as e:
        log_cb and log_cb(f"   OSD không nhận diện được hướng trang: {e}")
        return None, 0


def detect_skew_angle_hough(pil_image, max_angle=15):
    """
    Nhận diện góc NGHIÊNG NHỎ (deskew, mục V.2 tài liệu - phân biệt với xoay 90/180/270 ở mục
    V.1) bằng Hough Line Transform trên OpenCV - nhắm đúng vào các đường thẳng dòng chữ, KHÔNG
    dùng minAreaRect trên toàn khung ảnh (dễ bị lẫn với viền trang giấy, kém chính xác hơn).

    max_angle: chỉ xét góc nghiêng trong khoảng (-max_angle, +max_angle) độ - đúng phạm vi
    "nghiêng nhỏ do scan" mục V.2, KHÔNG dùng để phát hiện xoay 90 độ (mục đích khác).

    Trả về góc nghiêng phát hiện được (độ, dương = nghiêng theo chiều kim đồng hồ) hoặc None nếu
    không đủ đường thẳng để xác định đáng tin cậy.
    """
    if not HAS_CV2:
        return None
    cv_img = cv2.cvtColor(np.array(pil_image.convert("RGB")), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLines(edges, 1, np.pi / 720, threshold=150)
    if lines is None:
        return None

    angles = []
    for line in lines[:300]:
        rho, theta = line[0]
        angle_deg = (theta * 180 / np.pi) - 90
        if -max_angle < angle_deg < max_angle:
            angles.append(angle_deg)
    if len(angles) < 20:  # qua it duong thang phat hien duoc - khong du tin cay de ket luan
        return None
    return -float(np.median(angles))  # dao dau: goc CAN BU de lam thang, khong phai goc da nghieng


def analyze_page_orientation(doc, page_idx, pdf_path=None, ocr_dpi=200, log_cb=None):
    """
    Phân tích ĐẦY ĐỦ hướng của 1 trang PDF - đúng quy trình 5 bước mục IV tài liệu:
    (1) đọc rotation metadata, (2) kiểm tra text layer, (3) phân tích ảnh trang qua OSD, (4) thử
    các hướng xoay, (5) phát hiện góc nghiêng nhỏ sau khi đã xác định hướng chính.

    Trả về dict: {so_trang, rotation_metadata, co_text_layer, goc_xoay_de_xuat, goc_nghieng,
    diem_tin_cay, trang_thai, ghi_chu}.
    trang_thai thuộc {"TRANG_DA_DUNG_HUONG", "DA_XOAY_90/180/270", "CAN_KIEM_TRA_HUONG",
    "KHONG_CO_DU_TEXT"}.
    """
    page = doc[page_idx]
    rotation_metadata = page.rotation
    text_content = page.get_text().strip()
    co_text_layer = len(text_content) >= 20

    pix = page.get_pixmap(matrix=fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    goc_xoay, osd_conf = detect_page_rotation_osd(img, log_cb)

    # Ket hop voi detect_best_orientation() da co san (dem ky tu OCR theo 4 huong) khi OSD khong
    # chac chan - dung mục XIV tài liệu ("Nên kết hợp: Text direction, OSD, ... Điểm OCR theo
    # từng hướng", không chỉ dựa vào 1 kết quả OCR duy nhất).
    nguon_nhan_dien = "OSD"
    if (goc_xoay is None or osd_conf < 5) and pdf_path:
        try:
            fallback = detect_best_orientation(pdf_path, page_idx)
        except Exception:
            fallback = None
        if fallback and fallback["chac_chan"]:
            goc_xoay = fallback["angle_de_xuat"]
            osd_conf = 25  # gan tuong duong muc tin cay trung binh vi da qua kiem tra cheo
            nguon_nhan_dien = "Đếm ký tự OCR 4 hướng"

    if goc_xoay is None:
        # Khong du can cu (qua it chu, chu yeu la anh/ban do/so do - mục VIII tài liệu)
        return {"so_trang": page_idx + 1, "rotation_metadata": rotation_metadata,
                "co_text_layer": co_text_layer, "goc_xoay_de_xuat": 0, "goc_nghieng": 0,
                "diem_tin_cay": 0, "trang_thai": "CAN_KIEM_TRA_HUONG" if not co_text_layer else "KHONG_CO_DU_TEXT",
                "nguon_nhan_dien": "",
                "ghi_chu": "Không đủ căn cứ nhận diện hướng (quá ít chữ, có thể là ảnh/bản đồ/sơ đồ)"}

    # Diem tin cay: ket hop do tin cay tho cua OSD (thuong 10-30) len thang 0-100, co thuong neu
    # trang co text layer that (dau hieu ro rang hon la chi dua vao OCR anh).
    diem_tin_cay = min(100, osd_conf * 3 + (15 if co_text_layer else 0))

    goc_nghieng = 0
    if goc_xoay == 0:
        # Chi lam thang (deskew) khi trang DA o dung huong tong the (0/90/180/270) - deskew
        # KHONG thay the cho xoay 90 do (dung mục V tài liệu).
        img_for_deskew = img
        goc_nghieng_phat_hien = detect_skew_angle_hough(img_for_deskew)
        if goc_nghieng_phat_hien is not None and abs(goc_nghieng_phat_hien) >= 0.3:
            goc_nghieng = round(goc_nghieng_phat_hien, 2)

    if goc_xoay == 0 and goc_nghieng == 0:
        trang_thai = "TRANG_DA_DUNG_HUONG"
        ghi_chu = ""
    elif goc_xoay == 0:
        trang_thai = "DA_LAM_THANG"
        ghi_chu = f"Nghiêng nhỏ {goc_nghieng}° - đã bù để làm thẳng"
    else:
        trang_thai = f"DA_XOAY_{goc_xoay}"
        ghi_chu = f"Xoay {goc_xoay}°" + (f" + làm thẳng thêm {goc_nghieng}°" if goc_nghieng else "")

    return {"so_trang": page_idx + 1, "rotation_metadata": rotation_metadata,
            "co_text_layer": co_text_layer, "goc_xoay_de_xuat": goc_xoay, "goc_nghieng": goc_nghieng,
            "diem_tin_cay": round(diem_tin_cay, 1), "trang_thai": trang_thai,
            "nguon_nhan_dien": nguon_nhan_dien, "ghi_chu": ghi_chu}


def auto_rotate_deskew_pdf(pdf_path, output_path, log_cb, threshold_auto=60, threshold_warn=40,
                            ocr_dpi=200, only_fix_wrong_pages=True, control=None, page_progress_cb=None,
                            dry_run=False, skip_pages=None, manual_rotations=None):
    """
    Tự động nhận diện và xoay/làm thẳng TOÀN BỘ trang trong 1 file PDF - đúng quy trình mục III-V
    tài liệu "LỆNH PHÁT TRIỂN TAB 6 – MỤC XOAY PDF". KHÔNG sửa file gốc - xuất ra file MỚI.

    threshold_auto: điểm tin cậy (0-100) từ đó tự động xử lý (mục VIII tài liệu).
    threshold_warn: điểm tin cậy dưới mức này thì bắt buộc CẦN KIỂM TRA thủ công, không tự xoay.
    only_fix_wrong_pages: chỉ xoay/làm thẳng các trang PHÁT HIỆN SAI hướng, giữ nguyên trang đã
    đúng (mục VI tài liệu - KHÔNG ép mọi trang phải cùng 1 khổ giấy).
    page_progress_cb(so_trang, tong_so_trang): gọi sau mỗi trang xử lý xong - dùng cho thanh tiến
    độ mục IX tài liệu.
    dry_run=True: CHỈ PHÂN TÍCH, không xoay/làm thẳng/ghi file - dùng cho "Quét và phân tích" (mục
    VII tài liệu - xem trước kết quả dự kiến trước khi người dùng xác nhận xử lý thật).
    skip_pages: set các số trang (1-based, khớp với cột "Số trang" trong bảng xem trước) mà người
    dùng chọn "Bỏ qua trang này" sau khi xem bảng xem trước - đúng mục VII "Bỏ qua một trang".
    manual_rotations: dict {so_trang(1-based): goc_xoay_thu_cong} - người dùng tự ghi đè góc xoay
    đề xuất cho 1 trang cụ thể, đúng mục VII "Xoay thủ công thêm 90°" / "Đặt lại về hướng gốc".

    Trả về dict {"trang_ket_qua": [...], "so_trang_da_dung": n, "so_trang_da_xoay": n,
    "so_trang_da_lam_thang": n, "so_trang_can_kiem_tra": n}.
    """
    skip_pages = skip_pages or set()
    manual_rotations = manual_rotations or {}
    if not HAS_FITZ:
        raise RuntimeError("Thiếu thư viện PyMuPDF (fitz).")
    doc = fitz.open(pdf_path)
    trang_ket_qua = []
    so_da_dung = so_da_xoay = so_da_lam_thang = so_can_kiem_tra = 0

    try:
        for i in range(len(doc)):
            if control:
                try:
                    control.checkpoint()
                except TaskCancelled:
                    log_cb(f"⏹ Đã hủy - còn {len(doc) - i} trang chưa xử lý.")
                    break

            so_trang_1based = i + 1
            if so_trang_1based in skip_pages:
                # Nguoi dung chon "Bo qua trang nay" sau khi xem bang xem truoc - giu nguyen
                # hoan toan, khong phan tich lai, khong xoay (mục VII tài liệu).
                so_da_dung += 1
                trang_ket_qua.append({
                    "so_trang": so_trang_1based, "rotation_metadata": doc[i].rotation,
                    "co_text_layer": None, "goc_xoay_de_xuat": 0, "goc_nghieng": 0,
                    "diem_tin_cay": 100, "trang_thai": "DA_DUNG_BOI_NGUOI_DUNG",
                    "nguon_nhan_dien": "Người dùng chọn bỏ qua", "ghi_chu": "Người dùng chọn giữ nguyên trang này"})
                log_cb(f"   Trang {so_trang_1based}: người dùng chọn BỎ QUA - giữ nguyên.")
                continue

            phan_tich = analyze_page_orientation(doc, i, pdf_path=pdf_path, ocr_dpi=ocr_dpi, log_cb=log_cb)
            page_progress_cb and page_progress_cb(i + 1, len(doc))

            if so_trang_1based in manual_rotations:
                # Nguoi dung tu ghi de goc xoay de xuat (mục VII - "Xoay thủ công thêm"/"Đặt lại
                # về hướng gốc") - GHI ĐÈ ket qua tu dong, khong con la CAN_KIEM_TRA nua vi da co
                # xac nhan ro rang tu con nguoi.
                phan_tich["goc_xoay_de_xuat"] = manual_rotations[so_trang_1based] % 360
                phan_tich["trang_thai"] = f"DA_XOAY_{manual_rotations[so_trang_1based] % 360}" \
                    if manual_rotations[so_trang_1based] % 360 else "DA_DUNG_BOI_NGUOI_DUNG"
                phan_tich["diem_tin_cay"] = 100
                phan_tich["nguon_nhan_dien"] = "Người dùng xoay thủ công"
                phan_tich["ghi_chu"] = f"Người dùng tự chọn xoay {manual_rotations[so_trang_1based] % 360}°"
                log_cb(f"   Trang {so_trang_1based}: người dùng tự xoay thủ công {phan_tich['goc_xoay_de_xuat']}°")

            if phan_tich["trang_thai"] in ("CAN_KIEM_TRA_HUONG", "KHONG_CO_DU_TEXT"):
                so_can_kiem_tra += 1
                trang_ket_qua.append(phan_tich)
                log_cb(f"   Trang {i + 1}: {phan_tich['ghi_chu']} → CẦN KIỂM TRA")
                continue

            if phan_tich["diem_tin_cay"] < threshold_warn:
                phan_tich["trang_thai"] = "CAN_KIEM_TRA_HUONG"
                phan_tich["ghi_chu"] += f" (điểm tin cậy {phan_tich['diem_tin_cay']} dưới ngưỡng {threshold_warn})"
                so_can_kiem_tra += 1
                trang_ket_qua.append(phan_tich)
                log_cb(f"   Trang {i + 1}: điểm tin cậy thấp ({phan_tich['diem_tin_cay']}) → CẦN KIỂM TRA")
                continue

            if phan_tich["diem_tin_cay"] < threshold_auto:
                log_cb(f"   ⚠ Trang {i + 1}: điểm tin cậy trung bình ({phan_tich['diem_tin_cay']}) - "
                      f"vẫn xử lý nhưng cần xem lại")

            if dry_run:
                # Chi PHAN TICH va DU DOAN ket qua se the nao - KHONG dung page.set_rotation() hay
                # ghi anh (mục VII - chỉ xem trước, chưa xử lý thật).
                if phan_tich["goc_xoay_de_xuat"] == 0 and phan_tich["goc_nghieng"] == 0:
                    so_da_dung += 1
                else:
                    if phan_tich["goc_xoay_de_xuat"] != 0:
                        so_da_xoay += 1
                    if phan_tich["goc_nghieng"] != 0 and not phan_tich["co_text_layer"]:
                        so_da_lam_thang += 1
                trang_ket_qua.append(phan_tich)
                continue

            page = doc[i]
            if phan_tich["goc_xoay_de_xuat"] != 0 or not only_fix_wrong_pages:
                page.set_rotation((page.rotation + phan_tich["goc_xoay_de_xuat"]) % 360)
                so_da_xoay += 1

            if phan_tich["goc_nghieng"] != 0:
                # Lam thang (deskew) THAT SU can bien doi noi dung anh - chi ap dung cho trang
                # KHONG co text layer that (dang scan) de tranh lam hong noi dung vector/text
                # (dung nguyen tac mục V - deskew la thao tac anh, khac voi xoay trang kim loai).
                if not phan_tich["co_text_layer"]:
                    pix = page.get_pixmap(matrix=fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    rotated_img = img.rotate(-phan_tich["goc_nghieng"], expand=True, fillcolor=(255, 255, 255))
                    buf = io.BytesIO()
                    rotated_img.save(buf, format="PNG")
                    buf.seek(0)
                    img_bytes = buf.read()
                    rect = page.rect
                    page.clean_contents()
                    page.insert_image(rect, stream=img_bytes, overlay=True)
                    so_da_lam_thang += 1
                else:
                    phan_tich["ghi_chu"] += " (bỏ qua làm thẳng - trang có text layer thật, chỉ xoay)"

            if phan_tich["goc_xoay_de_xuat"] == 0 and phan_tich["goc_nghieng"] == 0:
                so_da_dung += 1

            trang_ket_qua.append(phan_tich)

        if dry_run:
            log_cb("Đã quét và phân tích xong - CHƯA ghi file nào (chế độ xem trước).")
        else:
            out_dir = os.path.dirname(os.path.abspath(output_path))
            if out_dir:
                os.makedirs(out_dir, exist_ok=True)
            from app.services import file_safety_service as _fsafe
            _fsafe.ghi_nguyen_tu(doc.save, output_path)
    finally:
        doc.close()

    return {"trang_ket_qua": trang_ket_qua, "so_trang_da_dung": so_da_dung,
            "so_trang_da_xoay": so_da_xoay, "so_trang_da_lam_thang": so_da_lam_thang,
            "so_trang_can_kiem_tra": so_can_kiem_tra}


def write_log_tu_dong_xoay_pdf(csv_path, log_rows):
    """Ghi Logs/LOG_TU_DONG_XOAY_PDF.csv - đúng mục XIII tài liệu (tổng hợp theo từng FILE)."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file gốc", "Đường dẫn file gốc", "Tổng số trang", "Số trang đã đúng",
                    "Số trang xoay 90°", "Số trang xoay 180°", "Số trang xoay 270°",
                    "Số trang được làm thẳng", "Số trang cần kiểm tra", "Tên file đầu ra",
                    "Đường dẫn đầu ra", "Thời gian xử lý", "Trạng thái", "Ghi chú"])
        for r in log_rows:
            w.writerow([r["stt"], r["file_in"], r["path_in"], r["tong_so_trang"], r["so_trang_da_dung"],
                       r["so_trang_xoay_90"], r["so_trang_xoay_180"], r["so_trang_xoay_270"],
                       r["so_trang_lam_thang"], r["so_trang_can_kiem_tra"], r["file_out"],
                       r["path_out"], r["thoi_gian_xu_ly"], r["status"], r["note"]])


def write_log_chi_tiet_huong_trang(csv_path, chi_tiet_rows):
    """Ghi Logs/LOG_CHI_TIET_HUONG_TRANG.csv - đúng mục XIII tài liệu (chi tiết theo từng TRANG)."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Tên file", "Số trang", "Rotation metadata", "Hướng chữ phát hiện",
                    "Góc nghiêng phát hiện", "Góc xoay đã áp dụng", "Góc deskew đã áp dụng",
                    "Điểm tin cậy", "Nguồn nhận diện", "Trạng thái", "Ghi chú"])
        for r in chi_tiet_rows:
            w.writerow([r["file"], r["so_trang"], r["rotation_metadata"], r["goc_xoay_de_xuat"],
                       r["goc_nghieng"], r["goc_xoay_de_xuat"], r["goc_nghieng"], r["diem_tin_cay"],
                       r["nguon_nhan_dien"], r["trang_thai"], r["ghi_chu"]])


def batch_auto_rotate_deskew_pdf(source, output_folder, log_cb, name_suffix="DA_XOAY_DUNG",
                                  no_overwrite=True, threshold_auto=60, threshold_warn=40,
                                  ocr_dpi=200, control=None, progress_cb=None, dry_run=False,
                                  toc_do_xu_ly="can_bang", so_file_moi_dot=100, nghi_giua_dot_giay=2.0,
                                  perf_log_path=None):
    """
    Xử lý HÀNG LOẠT nhiều file PDF - tự nhận diện và xoay/làm thẳng từng file, đúng mục I-II tài
    liệu "LỆNH PHÁT TRIỂN TAB 6 – MỤC XOAY PDF". KHÔNG sửa file gốc.

    dry_run=True: chỉ "Quét và phân tích" (mục VII) - KHÔNG ghi file kết quả nào, dùng để xem
    trước kết quả dự kiến trước khi người dùng xác nhận xử lý thật.

    QUAN TRỌNG: hàm này CHỈ xử lý TUẦN TỰ (không hỗ trợ song song nhiều file cùng lúc) - đã THỬ
    NGHIỆM xử lý song song bằng ThreadPoolExecutor và PHÁT HIỆN DEADLOCK THẬT khi gọi Tesseract
    OCR đồng thời từ nhiều luồng (2 luồng cùng lúc không hoàn thành sau 35s, trong khi 1 luồng chỉ
    mất ~15s - không phải chậm do tranh chấp GIL bình thường mà là TREO THẬT, khả năng cao do
    pytesseract dùng chung tên file tạm giữa các lần gọi đồng thời). Đã HOÀN TÁC về tuần tự để
    đảm bảo an toàn - xác nhận đúng rủi ro đã cảnh báo trước đó về xử lý song song với OCR.

    Trả về (log_rows, chi_tiet_rows) - log_rows: 1 dòng/file (mục XIII.1), chi_tiet_rows: 1
    dòng/trang trên TẤT CẢ các file (mục XIII.2).
    """
    files = list_files_from_source(source, ".pdf")
    if not files:
        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")

    log_rows = []
    chi_tiet_rows = []
    if not dry_run:
        os.makedirs(output_folder, exist_ok=True)

    from app.services import word_perf_service as _wperf
    ap_dung_toi_uu = len(files) >= 2
    if ap_dung_toi_uu:
        _wperf.apply_low_priority()
        sleep_per_file = _wperf.get_sleep_per_file(toc_do_xu_ly)
    else:
        sleep_per_file = 0.0

    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy - còn {len(files) - i} file chưa xử lý.")
                break
        fname = os.path.basename(fpath)
        if progress_cb:
            progress_cb(i + 1, len(files), fname)
        t0 = time.time()
        log_cb(f"→ Đang xử lý: {fname}")

        base = os.path.splitext(fname)[0]
        suffix_part = f"_{name_suffix}" if name_suffix else "_DA_XOAY_DUNG"
        out_path = os.path.join(output_folder, base + suffix_part + ".pdf")
        if no_overwrite and os.path.isfile(out_path):
            k = 1
            while os.path.isfile(os.path.join(output_folder, f"{base}{suffix_part}_{k}.pdf")):
                k += 1
            out_path = os.path.join(output_folder, f"{base}{suffix_part}_{k}.pdf")

        row = {"stt": i + 1, "file_in": fname, "path_in": fpath, "tong_so_trang": 0,
              "so_trang_da_dung": 0, "so_trang_xoay_90": 0, "so_trang_xoay_180": 0,
              "so_trang_xoay_270": 0, "so_trang_lam_thang": 0, "so_trang_can_kiem_tra": 0,
              "file_out": "", "path_out": "", "thoi_gian_xu_ly": "", "status": "", "note": ""}
        try:
            def page_progress_cb(so_trang, tong):
                log_cb(f"   Đang phân tích trang {so_trang}/{tong}")

            result = auto_rotate_deskew_pdf(
                fpath, out_path, log_cb, threshold_auto=threshold_auto, threshold_warn=threshold_warn,
                ocr_dpi=ocr_dpi, control=control, page_progress_cb=page_progress_cb, dry_run=dry_run)

            row["tong_so_trang"] = len(result["trang_ket_qua"])
            row["so_trang_da_dung"] = result["so_trang_da_dung"]
            row["so_trang_lam_thang"] = result["so_trang_da_lam_thang"]
            row["so_trang_can_kiem_tra"] = result["so_trang_can_kiem_tra"]
            for tr in result["trang_ket_qua"]:
                if tr["trang_thai"] == "DA_XOAY_90":
                    row["so_trang_xoay_90"] += 1
                elif tr["trang_thai"] == "DA_XOAY_180":
                    row["so_trang_xoay_180"] += 1
                elif tr["trang_thai"] == "DA_XOAY_270":
                    row["so_trang_xoay_270"] += 1
                chi_tiet_rows.append({"file": fname, **tr})

            row["file_out"] = os.path.basename(out_path)
            row["path_out"] = out_path
            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)
            row["status"] = "CAN_KIEM_TRA" if result["so_trang_can_kiem_tra"] else "THANH_CONG"
            row["note"] = (f"{result['so_trang_can_kiem_tra']} trang cần kiểm tra thủ công"
                          if result["so_trang_can_kiem_tra"] else "")
            log_cb(f"✓ Đã xử lý: {fname} → {row['file_out']}")
        except Exception as e:
            row["status"] = "LOI_DOC_PDF"
            row["note"] = str(e)
            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)
            log_cb(f"✗ Lỗi xử lý {fname}: {e}")

        log_rows.append(row)

        if ap_dung_toi_uu:
            _wperf.collect_garbage()
            time.sleep(sleep_per_file)
            if perf_log_path:
                kich_thuoc_kb = round(os.path.getsize(fpath) / 1024, 1) if os.path.isfile(fpath) else 0
                stats_now = _wperf.sample_perf_stats() or {}
                _wperf.write_perf_log_row(
                    perf_log_path, time.strftime("%Y-%m-%d %H:%M:%S"),
                    stats_now.get("cpu_percent", ""), stats_now.get("ram_mb", ""),
                    fname, row.get("thoi_gian_xu_ly", ""), kich_thuoc_kb,
                    row.get("status") not in ("LOI_DOC_PDF",), row.get("note", ""))
            if so_file_moi_dot and (i + 1) % so_file_moi_dot == 0 and (i + 1) < len(files):
                log_cb(f"⏸ Đã xử lý {i + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                time.sleep(nghi_giua_dot_giay)

    if ap_dung_toi_uu:
        _wperf.restore_normal_priority()

    return log_rows, chi_tiet_rows


try:
    from app.services.pdf_service import (
        get_pdf_page_count,
        render_pdf_page_to_image,
        apply_pdf_rotations,
        parse_page_range_spec,
        images_to_pdf,
        detect_best_orientation,
    )
    from app.services.pdf_service import improve_scanned_pdf_image as improve_scan_image
    _PDF_SERVICE_MODULE_OK = True
except Exception:
    _PDF_SERVICE_MODULE_OK = False

    def get_pdf_page_count(pdf_path):
        doc = fitz.open(pdf_path)
        try:
            return len(doc)
        finally:
            doc.close()

    def render_pdf_page_to_image(pdf_path, page_index=0, zoom=1.5, extra_rotation=0):
        """
        Render 1 trang PDF thành ảnh PIL để xem trước. extra_rotation (0/90/180/270) được cộng
        thêm vào rotation hiện có của trang (dùng để xem trước hiệu ứng xoay TRƯỚC khi lưu thật).
        """
        doc = fitz.open(pdf_path)
        try:
            if page_index < 0 or page_index >= len(doc):
                raise ValueError(f"Số trang {page_index + 1} vượt quá tổng số trang ({len(doc)}).")
            page = doc[page_index]
            original_rotation = page.rotation
            if extra_rotation:
                page.set_rotation((original_rotation + extra_rotation) % 360)
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            return img
        finally:
            doc.close()

    def apply_pdf_rotations(input_pdf, output_pdf, rotations, default_delta=0):
        """
        Lưu 1 file PDF mới với rotation đã áp dụng, KHÔNG sửa file gốc.
        rotations: dict {page_index(0-based): delta_độ} cho các trang có xoay riêng.
        default_delta: độ xoay áp dụng cho các trang KHÔNG có trong rotations (0 = giữ nguyên).
        Trả về số trang đã bị xoay (khác 0 độ).
        """
        doc = fitz.open(input_pdf)
        changed = 0
        try:
            for i, page in enumerate(doc):
                delta = rotations.get(i, default_delta)
                if delta:
                    page.set_rotation((page.rotation + delta) % 360)
                    changed += 1
            os.makedirs(os.path.dirname(os.path.abspath(output_pdf)), exist_ok=True)
            doc.save(output_pdf)
        finally:
            doc.close()
        return changed

    def parse_page_range_spec(spec, total_pages):
        """Phân tích chuỗi khoảng trang kiểu '1-5,8,10-12' -> set số trang (0-based). Trống -> tất cả."""
        spec = (spec or "").strip()
        if not spec:
            return set(range(total_pages))
        pages = set()
        for part in spec.split(","):
            part = part.strip()
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                a, b = int(a.strip()), int(b.strip())
                for p in range(a, b + 1):
                    if 1 <= p <= total_pages:
                        pages.add(p - 1)
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p - 1)
        return pages

    def improve_scan_image(img, contrast=1.0, sharpen=False, grayscale=False, threshold=None,
                            denoise=False, deskew=False, crop_border=False, normalize_a4=False):
        """
        Cải thiện ảnh scan: tăng tương phản, làm nét, grayscale, nhị phân hóa, khử nhiễu nhẹ, sửa nghiêng,
        cắt viền trắng/đen thừa, chuẩn hóa kích thước về A4. KHÔNG sửa ảnh gốc (trả về bản mới).
        Trả về (ảnh_đã_xử_lý, góc_nghiêng_đã_sửa_hoặc_None).
        """
        from PIL import ImageEnhance
        deskew_angle = None

        if crop_border:
            gray_for_crop = ImageOps.grayscale(img)
            bbox = ImageOps.invert(gray_for_crop).getbbox()
            if bbox:
                w, h = img.size
                pad_x, pad_y = int(w * 0.01), int(h * 0.01)
                l = max(0, bbox[0] - pad_x)
                t = max(0, bbox[1] - pad_y)
                r = min(w, bbox[2] + pad_x)
                b = min(h, bbox[3] + pad_y)
                if r > l and b > t:
                    img = img.crop((l, t, r, b))

        if deskew:
            if HAS_CV2:
                arr = np.array(img.convert("L"))
                arr_inv = cv2.bitwise_not(arr)
                thresh_img = cv2.threshold(arr_inv, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
                coords = cv2.findNonZero(thresh_img)
                if coords is not None:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle
                    if abs(angle) > 0.3:
                        img = img.rotate(angle, expand=True, fillcolor="white", resample=Image.BICUBIC)
                        deskew_angle = round(angle, 2)

        if grayscale:
            img = ImageOps.grayscale(img)
            if img.mode != "RGB":
                img = img.convert("RGB") if threshold is None else img

        if contrast and contrast != 1.0:
            base = img if img.mode in ("L", "RGB") else img.convert("RGB")
            img = ImageEnhance.Contrast(base).enhance(contrast)

        if denoise:
            img = img.filter(ImageFilter.MedianFilter(size=3))

        if sharpen:
            img = img.filter(ImageFilter.SHARPEN)

        if threshold is not None:
            gray = img.convert("L")
            img = gray.point(lambda x: 0 if x < threshold else 255, "L").convert("RGB")

        if normalize_a4:
            w, h = img.size
            if w > h:
                target_w, target_h = 3508, 2480
            else:
                target_w, target_h = 2480, 3508
            ratio = min(target_w / w, target_h / h)
            new_size = (max(1, int(w * ratio)), max(1, int(h * ratio)))
            resized = img.resize(new_size, Image.LANCZOS)
            canvas = Image.new("RGB", (target_w, target_h), "white")
            offset = ((target_w - new_size[0]) // 2, (target_h - new_size[1]) // 2)
            canvas.paste(resized, offset)
            img = canvas

        return img, deskew_angle

    def detect_best_orientation(pdf_path, page_index, lang="vie+eng"):
        """
        Tự phát hiện hướng đúng của 1 trang: OCR nhanh ở 0/90/180/270 độ, chọn hướng có nhiều
        ký tự đọc được nhất (coi là proxy cho độ tin cậy). Trả về dict:
        {"angle_de_xuat": int, "huong_ban_dau": int, "ket_qua_tung_huong": {0:n_ky_tu,...},
         "chac_chan": bool, "ly_do": str}
        chac_chan=False -> nên đưa vào "CẦN KIỂM TRA" thay vì tự động xoay.
        """
        if not (HAS_FITZ and HAS_OCR):
            raise RuntimeError("Cần cài pymupdf + pytesseract + Tesseract-OCR để dùng tính năng này.")

        doc = fitz.open(pdf_path)
        try:
            page = doc[page_index]
            original_rotation = page.rotation
            counts = {}
            for extra in (0, 90, 180, 270):
                page.set_rotation((original_rotation + extra) % 360)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                try:
                    txt = pytesseract.image_to_string(img, lang=lang, config="--psm 6")
                except Exception:
                    txt = ""
                counts[extra] = len(re.sub(r"\s+", "", txt))
            page.set_rotation(original_rotation)

            best_extra = max(counts, key=counts.get)
            best_count = counts[best_extra]
            sorted_counts = sorted(counts.values(), reverse=True)
            chac_chan = best_count >= 15 and (len(sorted_counts) < 2 or best_count > sorted_counts[1] * 1.3)

            return {
                "angle_de_xuat": best_extra,
                "huong_ban_dau": original_rotation,
                "ket_qua_tung_huong": counts,
                "chac_chan": chac_chan,
                "ly_do": (f"Hướng +{best_extra}° đọc được {best_count} ký tự (nhiều nhất)" if chac_chan
                         else "Không đủ chênh lệch giữa các hướng - có thể là trang ảnh/bảng biểu, ít chữ"),
            }
        finally:
            doc.close()

    def images_to_pdf(images, output_pdf):
        """Đóng gói 1 danh sách ảnh PIL thành 1 file PDF mới (dùng khi đã cải thiện ảnh scan)."""
        if not images:
            raise RuntimeError("Không có ảnh nào để đóng gói thành PDF.")
        os.makedirs(os.path.dirname(os.path.abspath(output_pdf)), exist_ok=True)
        first, rest = images[0].convert("RGB"), [im.convert("RGB") for im in images[1:]]
        first.save(output_pdf, save_all=True, append_images=rest)


def edit_pdf_file(input_pdf, output_pdf, rotations, enhance_opts, ocr_after=False,
                   ocr_lang="vie+eng", log_cb=None):
    """
    Xử lý 1 file PDF: áp dụng rotations (dict trang->độ xoay) và/hoặc cải thiện ảnh scan
    (enhance_opts: dict {contrast, sharpen, grayscale, threshold, denoise, deskew} hoặc None/rỗng).
    - Nếu KHÔNG có enhance_opts: chỉ xoay nhanh bằng set_rotation (giữ nguyên text layer/vector).
    - Nếu CÓ enhance_opts: rasterize từng trang -> xử lý ảnh -> đóng gói lại thành PDF mới
      (mất text layer gốc - đây là đánh đổi cố hữu của việc xử lý ảnh scan).
    Trả về dict: {"so_trang": int, "trang_da_xoay": [...], "deskew_angles": {...}}.
    """
    doc = fitz.open(input_pdf)
    total_pages = len(doc)
    rotated_pages = [i + 1 for i, d in rotations.items() if d]

    if not enhance_opts:
        doc.close()
        apply_pdf_rotations(input_pdf, output_pdf, rotations)
        return {"so_trang": total_pages, "trang_da_xoay": rotated_pages, "deskew_angles": {}}

    images = []
    deskew_angles = {}
    try:
        for i, page in enumerate(doc):
            delta = rotations.get(i, 0)
            if delta:
                page.set_rotation((page.rotation + delta) % 360)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            proc_img, deskew_angle = improve_scan_image(
                img,
                contrast=enhance_opts.get("contrast", 1.0),
                sharpen=enhance_opts.get("sharpen", False),
                grayscale=enhance_opts.get("grayscale", False),
                threshold=enhance_opts.get("threshold"),
                denoise=enhance_opts.get("denoise", False),
                deskew=enhance_opts.get("deskew", False),
                crop_border=enhance_opts.get("crop_border", False),
                normalize_a4=enhance_opts.get("normalize_a4", False),
            )
            if deskew_angle is not None:
                deskew_angles[i + 1] = deskew_angle
            images.append(proc_img)
            if log_cb:
                log_cb(f"   Trang {i+1}/{total_pages} đã xử lý ảnh"
                       + (f" (sửa nghiêng {deskew_angle}°)" if deskew_angle is not None else ""))
    finally:
        doc.close()

    images_to_pdf(images, output_pdf)

    if ocr_after:
        if log_cb:
            log_cb("   Đang OCR lại để tạo lớp chữ tìm kiếm được...")
        # Ghi chú: tạo PDF "searchable" đầy đủ (ảnh + lớp text ẩn) cần pipeline riêng
        # (VD ocrmypdf). Ở đây phần mềm chỉ đảm bảo ảnh đã cải thiện được lưu ra PDF mới;
        # việc thêm lớp text ẩn tương thích 100% các trình đọc PDF nằm ngoài phạm vi bản này.
        if log_cb:
            log_cb("   ⚠ Lưu ý: PDF mới là ảnh đã cải thiện, chưa có lớp chữ ẩn để tìm kiếm/copy.")

    return {"so_trang": total_pages, "trang_da_xoay": rotated_pages, "deskew_angles": deskew_angles}


def batch_edit_pdfs(file_list, output_folder, get_rotations_for_file, enhance_opts, dry_run,
                     ocr_after, log_cb, control=None):
    """
    Xử lý hàng loạt file PDF (xoay + tùy chọn cải thiện ảnh).
    get_rotations_for_file(fpath, total_pages) -> dict {page_index: delta_độ} cho riêng file đó
    (cho phép GUI truyền logic áp dụng khác nhau: toàn bộ/trang chẵn/lẻ/khoảng trang/tự phát hiện).
    Trả về list dict log: {stt, file_goc, file_xuat, so_trang, trang_da_xoay, goc_xoay,
    tang_tuong_phan, lam_net, sua_nghieng, status, note}.
    """
    results = []
    if not dry_run:
        os.makedirs(output_folder, exist_ok=True)

    for idx, fpath in enumerate(file_list, start=1):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(file_list) - idx + 1} file chưa xử lý).")
                break

        fname = os.path.basename(fpath)
        log_cb(f"→ Đang xử lý: {fname}")
        try:
            total_pages = get_pdf_page_count(fpath)
        except Exception as e:
            results.append({"stt": idx, "file_goc": fname, "file_xuat": "", "so_trang": "",
                            "trang_da_xoay": "", "goc_xoay": "", "tang_tuong_phan": "", "lam_net": "",
                            "sua_nghieng": "", "status": "LOI_FILE_PDF", "note": str(e)})
            log_cb(f"   ✗ LOI_FILE_PDF: {e}")
            continue

        try:
            rotations = get_rotations_for_file(fpath, total_pages)
        except Exception as e:
            results.append({"stt": idx, "file_goc": fname, "file_xuat": "", "so_trang": total_pages,
                            "trang_da_xoay": "", "goc_xoay": "", "tang_tuong_phan": "", "lam_net": "",
                            "sua_nghieng": "", "status": "CAN_KIEM_TRA", "note": str(e)})
            log_cb(f"   ⚠ CAN_KIEM_TRA: {e}")
            continue

        base = os.path.splitext(fname)[0]
        out_name = f"{base}_chinh_sua.pdf"
        out_path = os.path.join(output_folder, out_name) if output_folder else out_name
        bpath, ext = os.path.splitext(out_path)
        i2 = 1
        while not dry_run and os.path.isfile(out_path):
            out_path = f"{bpath}_{i2}{ext}"
            i2 += 1
        out_name = os.path.basename(out_path)

        if dry_run:
            rotated = [i + 1 for i, d in rotations.items() if d]
            results.append({"stt": idx, "file_goc": fname, "file_xuat": out_name, "so_trang": total_pages,
                            "trang_da_xoay": ",".join(map(str, rotated)) or "(không)",
                            "goc_xoay": ",".join(f"{rotations[i]}°" for i in rotations if rotations[i]) or "",
                            "tang_tuong_phan": "Có" if enhance_opts and enhance_opts.get("contrast", 1.0) != 1.0 else "",
                            "lam_net": "Có" if enhance_opts and enhance_opts.get("sharpen") else "",
                            "sua_nghieng": "Có" if enhance_opts and enhance_opts.get("deskew") else "",
                            "status": "CHAY_THU", "note": ""})
            log_cb(f"   ✓ CHAY_THU: {len(rotated)}/{total_pages} trang sẽ được xoay")
            continue

        try:
            info = edit_pdf_file(fpath, out_path, rotations, enhance_opts, ocr_after=ocr_after, log_cb=log_cb)
            results.append({"stt": idx, "file_goc": fname, "file_xuat": out_name, "so_trang": info["so_trang"],
                            "trang_da_xoay": ",".join(map(str, info["trang_da_xoay"])) or "(không)",
                            "goc_xoay": ",".join(f"{rotations[i]}°" for i in rotations if rotations[i]) or "",
                            "tang_tuong_phan": "Có" if enhance_opts and enhance_opts.get("contrast", 1.0) != 1.0 else "",
                            "lam_net": "Có" if enhance_opts and enhance_opts.get("sharpen") else "",
                            "sua_nghieng": "Có" if info["deskew_angles"] else "",
                            "status": "DA_XU_LY", "note": ""})
            log_cb(f"   ✓ DA_XU_LY: {fname} → {out_name}")
        except Exception as e:
            results.append({"stt": idx, "file_goc": fname, "file_xuat": "", "so_trang": total_pages,
                            "trang_da_xoay": "", "goc_xoay": "", "tang_tuong_phan": "", "lam_net": "",
                            "sua_nghieng": "", "status": "LOI_LUU_FILE", "note": str(e)})
            log_cb(f"   ✗ LOI_LUU_FILE: {e}")

    return results


def write_pdf_edit_log_csv(csv_path, results):
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file gốc", "Tên file xuất", "Tổng số trang", "Trang đã chỉnh sửa",
                    "Góc xoay", "Có tăng tương phản", "Có làm nét", "Có sửa nghiêng", "Trạng thái", "Ghi chú"])
        for r in results:
            w.writerow([r["stt"], r["file_goc"], r["file_xuat"], r["so_trang"], r["trang_da_xoay"],
                        r["goc_xoay"], r["tang_tuong_phan"], r["lam_net"], r["sua_nghieng"],
                        r["status"], r["note"]])


# ============================================================================
# MỤC 10-C (MỚI): SO SÁNH THƯ MỤC WORD/PDF THEO MÃ XÃ-TỜ-THỬA & DI CHUYỂN FILE KHÔNG KHỚP
# ============================================================================

WORD_PDF_COMPARE_EXTS = (".pdf", ".doc", ".docx", ".docm")


try:
    from app.services.word_service import read_word_text
    _WORD_SERVICE_MODULE_OK_C = True
except Exception:
    _WORD_SERVICE_MODULE_OK_C = False

    def read_word_text(word_path, log_cb=None, word_session=None, doc_docx_cache=None):
        """
        Đọc toàn bộ text 1 file Word (.doc/.docx/.docm). Ưu tiên phiên Word COM dùng chung nếu có
        (nhanh, đọc được cả .doc cũ); dự phòng .docx/.docm dùng python-docx trực tiếp; .doc cũ không
        có Word thì thử chuyển tạm qua LibreOffice để đọc.

        doc_docx_cache (tùy chọn): dict {đường_dẫn_.doc_gốc: đường_dẫn_.docx_đã_chuyển_sẵn} - nếu
        word_path CÓ trong dict này, DÙNG LUÔN file đã chuyển sẵn thay vì tự chuyển đổi lại.
        """
        ext = os.path.splitext(word_path)[1].lower()

        if word_session is not None:
            try:
                return word_session.read_text(word_path)
            except Exception as e:
                if log_cb:
                    log_cb(f"   (Word COM lỗi đọc nội dung: {e}, thử phương án khác...)")

        if ext in (".docx", ".docm"):
            doc = DocxDocument(word_path)
            return "\n".join(p.text for p in doc.paragraphs)

        if ext == ".doc":
            if doc_docx_cache and word_path in doc_docx_cache:
                doc = DocxDocument(doc_docx_cache[word_path])
                return "\n".join(p.text for p in doc.paragraphs)
            if HAS_WIN32COM:
                try:
                    with WordCOMSession() as session:
                        return session.read_text(word_path)
                except Exception as e:
                    if log_cb:
                        log_cb(f"   (Word COM lỗi đọc .doc: {e}, thử LibreOffice...)")
            soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
            if soffice_bin:
                import tempfile
                with tempfile.TemporaryDirectory() as td:
                    result = subprocess.run(
                        [soffice_bin, "--headless", "--convert-to", "docx", "--outdir", td,
                         os.path.abspath(word_path)],
                        capture_output=True, text=True, timeout=60,
                    )
                    produced = os.path.join(td, os.path.splitext(os.path.basename(word_path))[0] + ".docx")
                    if os.path.isfile(produced):
                        doc = DocxDocument(produced)
                        return "\n".join(p.text for p in doc.paragraphs)
            raise RuntimeError("Không đọc được nội dung file .doc cũ (cần Microsoft Word hoặc LibreOffice).")

        raise RuntimeError(f"Định dạng không hỗ trợ: {ext}")

def extract_fields_from_word(word_path, xa_mapping, thon_mapping, log_cb=None, word_session=None,
                              doc_docx_cache=None):
    """
    Đọc nội dung file Word, trích Số tờ/Số thửa/Mã xã theo đúng 4 mức ưu tiên (thôn mục 2.b ->
    thôn mục 1.c -> xã cũ rõ ràng -> không gán). Trả về dict giống extract_fields_from_pdf,
    hoặc {'error':...} nếu không đọc đủ thông tin.
    """
    try:
        text = read_word_text(word_path, log_cb=log_cb, word_session=word_session, doc_docx_cache=doc_docx_cache)
    except Exception as e:
        return {"error": str(e)}

    if not text or len(text.strip()) < 10:
        return {"error": "Nội dung file Word trống hoặc quá ngắn để đọc."}

    info = parse_fields_from_text(text)

    if not info.get("to") or not info.get("thua"):
        return {"error": "Không đọc được Số tờ/Số thửa từ nội dung Word.", **info}

    maxa = None
    nguon_maxa = None
    if info.get("thon_name"):
        key = strip_diacritics(info["thon_name"]).upper().strip()
        maxa = thon_mapping.get(key)
        if maxa:
            nguon_maxa = "thôn (mục 2.b)"
    if not maxa and info.get("thon_name_1c"):
        key = strip_diacritics(info["thon_name_1c"]).upper().strip()
        maxa = thon_mapping.get(key)
        if maxa:
            nguon_maxa = "thôn (mục 1.c)"
    if not maxa and info.get("xa_name"):
        key = strip_diacritics(info["xa_name"]).upper().strip()
        maxa = xa_mapping.get(key)
        if maxa:
            nguon_maxa = "xã (mục 2.b)"
    if not maxa and info.get("xa_name_1c"):
        key = strip_diacritics(info["xa_name_1c"]).upper().strip()
        maxa = xa_mapping.get(key)
        if maxa:
            nguon_maxa = "xã (mục 1.c)"

    info["maxa"] = maxa
    info["nguon_maxa"] = nguon_maxa
    return info


def identify_file_key_word_or_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi,
                                   log_cb, word_session=None, doc_docx_cache=None):
    """
    Xác định (Mã xã, Tờ, Thửa) cho 1 file PDF HOẶC Word: ưu tiên tên file trước, không đủ mới
    đọc nội dung (OCR cho PDF scan, đọc Word cho .doc/.docx). Trả về dict {maxa,to,thua,nguon}
    hoặc {'error':...}.
    doc_docx_cache: xem read_word_text() - dict {.doc gốc: .docx đã chuyển sẵn} để tránh khởi động
    LibreOffice riêng cho từng file .doc khi xử lý hàng loạt (người gọi tự chuẩn bị trước).
    """
    fname = os.path.basename(fpath)
    key = extract_key_from_filename_v2(fname)
    if key and key["to"] and key["thua"]:
        return {"maxa": key["maxa"], "to": key["to"], "thua": key["thua"], "nguon": "tên file"}

    ext = os.path.splitext(fname)[1].lower()
    log_cb(f"   Tên file không đủ khóa, đang đọc nội dung {fname}...")
    if ext == ".pdf":
        info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb)
    else:
        info = extract_fields_from_word(fpath, xa_mapping, thon_mapping, log_cb=log_cb, word_session=word_session,
                                        doc_docx_cache=doc_docx_cache)

    if "error" in info:
        return {"error": info["error"]}
    return {"maxa": info.get("maxa"), "to": info.get("to"), "thua": info.get("thua"), "nguon": "nội dung/OCR"}


try:
    from app.services.file_compare_service import move_or_copy_file as _move_or_copy_file
    _FILE_COMPARE_SERVICE_MODULE_OK_C = True
except Exception:
    _FILE_COMPARE_SERVICE_MODULE_OK_C = False

    def _move_or_copy_file(src, dest_dir, copy_mode, dry_run):
        """Copy/di chuyển file KHÔNG ghi đè (tự thêm _1, _2...). dry_run=True: chỉ tính tên đích, không làm gì thật."""
        base = os.path.basename(src)
        dest = os.path.join(dest_dir, base)
        if dry_run:
            return ("CHAY_THU", dest)
        os.makedirs(dest_dir, exist_ok=True)
        b, ext = os.path.splitext(dest)
        i = 1
        while os.path.isfile(dest):
            dest = f"{b}_{i}{ext}"
            i += 1
        if copy_mode:
            shutil.copy2(src, dest)
            return ("DA_COPY", dest)
        else:
            shutil.move(src, dest)
            return ("DA_DI_CHUYEN", dest)


def compare_folders_word_pdf(folder_a, folder_b, output_folder, include_subfolders,
                              use_maxa, fallback_to_thua_only, move_a, move_b, copy_instead_of_move,
                              dry_run, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi,
                              log_cb, control=None, word_session=None):
    """
    So sánh 2 chiều file PDF/Word giữa folder_a và folder_b theo khóa (Mã xã+Tờ+Thửa, hoặc chỉ
    Tờ+Thửa nếu không có mã xã). File không khớp được copy/di chuyển ra thư mục kết quả riêng.
    KHÔNG xóa file, KHÔNG ghi đè, KHÔNG sửa/đổi tên file gốc.
    Trả về list dict log đầy đủ (1 dòng / 1 file), dùng để hiển thị bảng + xuất CSV.
    """

    def list_files(folder):
        if include_subfolders:
            found = []
            for root, dirs, fnames in os.walk(folder):
                for f in fnames:
                    if f.lower().endswith(WORD_PDF_COMPARE_EXTS) and not f.startswith("~$"):
                        found.append(os.path.join(root, f))
            return sorted(found)
        return sorted(
            os.path.join(folder, f) for f in os.listdir(folder)
            if f.lower().endswith(WORD_PDF_COMPARE_EXTS) and not f.startswith("~$")
        )

    files_a = list_files(folder_a)
    files_b = list_files(folder_b)

    # TỐI ƯU TỐC ĐỘ: quét trước các file .doc CŨ mà tên file không đủ khóa (sẽ phải đọc nội dung
    # qua LibreOffice) - gộp chuyển đổi TẤT CẢ cùng lúc bằng 1 lệnh soffice DUY NHẤT thay vì mỗi
    # file 1 lệnh riêng (cùng lớp tối ưu đã áp dụng cho mục 7e - LibreOffice tốn chi phí khởi động
    # rất lớn, gộp nhiều file nhanh hơn 3.5-4.5 lần). CHỈ áp dụng khi KHÔNG có Word COM (đường Word
    # COM đã đủ nhanh với session dùng chung, không cần gộp).
    doc_docx_cache = {}
    if word_session is None and not HAS_WIN32COM:
        doc_files_can_ban_doc = []
        for fpath in files_a + files_b:
            if os.path.splitext(fpath)[1].lower() != ".doc":
                continue
            key = extract_key_from_filename_v2(os.path.basename(fpath))
            if not (key and key["to"] and key["thua"]):
                doc_files_can_ban_doc.append(fpath)
        if doc_files_can_ban_doc:
            try:
                td = os.path.join(get_app_data_dir(), "Temp", f"doc_to_docx_{int(time.time())}")
                log_cb(f"→ Gộp chuyển đổi trước {len(doc_files_can_ban_doc)} file .doc cũ sang .docx "
                      f"để đọc nhanh hơn (LibreOffice, 1 lần gọi thay vì {len(doc_files_can_ban_doc)} lần)...")
                doc_docx_cache = convert_word_to_pdf_libreoffice_batch(
                    doc_files_can_ban_doc, td, target_format="docx")
            except Exception as e:
                log_cb(f"⚠ Không gộp chuyển đổi trước được ({e}) - sẽ chuyển đổi từng file như bình thường.")

    def build_index(files, label):
        index = {}
        rows_local = []
        for i, fpath in enumerate(files):
            if control:
                try:
                    control.checkpoint()
                except TaskCancelled:
                    log_cb(f"⏹ Đã hủy khi đang xử lý thư mục {label} (còn {len(files) - i} file).")
                    break

            fname = os.path.basename(fpath)
            loai = os.path.splitext(fname)[1].lstrip(".").upper()
            try:
                result = identify_file_key_word_or_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback,
                                                        ocr_dpi, log_cb, word_session=word_session,
                                                        doc_docx_cache=doc_docx_cache)
            except Exception as e:
                result = {"error": f"Lỗi đọc file: {e}"}

            row = {"stt": None, "thumuc": label, "file": fname, "path": fpath, "loai": loai,
                   "maxa": "", "to": "", "thua": "", "khoa": "", "co_maxa": "",
                   "status": None, "hanhdong": "", "note": ""}

            if "error" in result:
                row["status"] = "LOI_DOC_FILE" if "Lỗi đọc file" in result["error"] else "CAN_KIEM_TRA_THIEU_THONG_TIN"
                row["note"] = result["error"]
                log_cb(f"[{label}] ✗ {fname}: {result['error']}")
                rows_local.append(row)
                continue

            maxa, to, thua = result.get("maxa"), result.get("to"), result.get("thua")
            if not to or not thua:
                row["status"] = "CAN_KIEM_TRA_THIEU_THONG_TIN"
                row["note"] = "Không đủ Số tờ/Số thửa"
                rows_local.append(row)
                continue

            if maxa and use_maxa:
                khoa = f"{maxa}_{to}_{thua}"
                co_maxa = "Có"
            elif fallback_to_thua_only:
                khoa = f"{to}_{thua}"
                co_maxa = "Có" if maxa else "Không"
                if not maxa:
                    row["note"] = ("File không có mã xã, chỉ so sánh theo số tờ + số thửa; cần kiểm tra "
                                   "nếu có khả năng trùng số tờ, số thửa giữa các xã cũ.")
            elif maxa:
                khoa = f"{maxa}_{to}_{thua}"
                co_maxa = "Có"
            else:
                row["status"] = "CAN_KIEM_TRA_THIEU_THONG_TIN"
                row["note"] = "Không có Mã xã và chưa bật tùy chọn so sánh theo Tờ+Thửa"
                rows_local.append(row)
                continue

            row.update(maxa=maxa or "", to=to, thua=thua, khoa=khoa, co_maxa=co_maxa)
            index.setdefault(khoa, []).append(row)
            rows_local.append(row)
        return index, rows_local

    log_cb(f"Đang lập chỉ mục thư mục A ({len(files_a)} file)...")
    index_a, rows_a = build_index(files_a, "A")
    log_cb(f"Đang lập chỉ mục thư mục B ({len(files_b)} file)...")
    index_b, rows_b = build_index(files_b, "B")

    for khoa, entries in index_a.items():
        if len(entries) > 1:
            for r in entries:
                r["note"] = (r["note"] + " | " if r["note"] else "") + \
                            f"TRÙNG_KHÓA_NHIỀU_FILE (cùng khóa với {len(entries) - 1} file khác trong A)"
    for khoa, entries in index_b.items():
        if len(entries) > 1:
            for r in entries:
                r["note"] = (r["note"] + " | " if r["note"] else "") + \
                            f"TRÙNG_KHÓA_NHIỀU_FILE (cùng khóa với {len(entries) - 1} file khác trong B)"

    keys_a = set(index_a.keys())
    keys_b = set(index_b.keys())

    for r in rows_a:
        if r["status"] is not None:
            continue
        r["status"] = "TRUNG_KHOP" if r["khoa"] in keys_b else "KHONG_KHOP_THU_MUC_A"
    for r in rows_b:
        if r["status"] is not None:
            continue
        r["status"] = "TRUNG_KHOP" if r["khoa"] in keys_a else "KHONG_KHOP_THU_MUC_B"

    out_a = os.path.join(output_folder, "Khong_khop_Thu_muc_A")
    out_b = os.path.join(output_folder, "Khong_khop_Thu_muc_B")
    out_check = os.path.join(output_folder, "Can_kiem_tra")

    all_rows = rows_a + rows_b
    for r in all_rows:
        do_move = False
        dest_dir = None
        if r["status"] == "KHONG_KHOP_THU_MUC_A" and move_a:
            do_move, dest_dir = True, out_a
        elif r["status"] == "KHONG_KHOP_THU_MUC_B" and move_b:
            do_move, dest_dir = True, out_b
        elif r["status"] in ("CAN_KIEM_TRA_THIEU_THONG_TIN", "LOI_DOC_FILE"):
            if (r["thumuc"] == "A" and move_a) or (r["thumuc"] == "B" and move_b):
                do_move, dest_dir = True, out_check

        if do_move:
            action, dest = _move_or_copy_file(r["path"], dest_dir, copy_instead_of_move, dry_run)
            r["hanhdong"] = action
            if not dry_run:
                r["note"] = (r["note"] + " | " if r["note"] else "") + f"-> {dest}"
                log_cb(f"[{r['thumuc']}] {action}: {r['file']} -> {dest_dir}")
        else:
            r["hanhdong"] = "CHAY_THU" if dry_run else ""

    for idx, r in enumerate(all_rows, start=1):
        r["stt"] = idx

    return all_rows


def write_compare_folders_logs(output_folder, stamp, all_rows):
    """Xuất LOG_SO_SANH_THU_MUC.csv + FILE_KHONG_KHOP_A.csv + FILE_KHONG_KHOP_B.csv + FILE_CAN_KIEM_TRA.csv."""
    import csv
    os.makedirs(output_folder, exist_ok=True)
    paths = {}

    p_full = os.path.join(output_folder, f"LOG_SO_SANH_THU_MUC_{stamp}.csv")
    with open(p_full, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "ThuMucNguon", "TenFile", "DuongDanFile", "LoaiFile", "MaXa", "SoTo", "SoThua",
                    "KhoaSoSanh", "CoMaXa", "TrangThai", "HanhDong", "GhiChu"])
        for r in all_rows:
            w.writerow([r["stt"], r["thumuc"], r["file"], r["path"], r["loai"], r["maxa"], r["to"], r["thua"],
                        r["khoa"], r["co_maxa"], r["status"], r["hanhdong"], r["note"]])
    paths["full"] = p_full

    def write_filtered(fname, predicate):
        p = os.path.join(output_folder, fname)
        with open(p, "w", encoding="utf-8-sig", newline="") as f:
            w = csv.writer(f)
            w.writerow(["STT", "ThuMucNguon", "TenFile", "MaXa", "SoTo", "SoThua", "TrangThai", "GhiChu"])
            for r in all_rows:
                if predicate(r):
                    w.writerow([r["stt"], r["thumuc"], r["file"], r["maxa"], r["to"], r["thua"],
                               r["status"], r["note"]])
        return p

    paths["a"] = write_filtered(f"FILE_KHONG_KHOP_A_{stamp}.csv", lambda r: r["status"] == "KHONG_KHOP_THU_MUC_A")
    paths["b"] = write_filtered(f"FILE_KHONG_KHOP_B_{stamp}.csv", lambda r: r["status"] == "KHONG_KHOP_THU_MUC_B")
    paths["check"] = write_filtered(
        f"FILE_CAN_KIEM_TRA_{stamp}.csv",
        lambda r: r["status"] in ("CAN_KIEM_TRA_THIEU_THONG_TIN", "LOI_DOC_FILE"))

    return paths


# ============================================================================
# MỤC 6 (MỚI): XUẤT WORD HÀNG LOẠT
# ============================================================================

def pdf_to_docx_single(pdf_path, docx_path, log_cb=None):
    if not HAS_PDF2DOCX:
        raise RuntimeError("Thiếu thư viện pdf2docx. Cài: pip install pdf2docx")
    cv = PDF2DocxConverter(pdf_path)
    try:
        cv.convert(docx_path)
    finally:
        cv.close()


# ============================================================================
# CHUYỂN ĐỔI PDF SANG WORD NÂNG CAO - 3 CHẾ ĐỘ (mục IV-X tài liệu
# "CHỈNH SỬA TAB 7+9, BỔ SUNG CHUYỂN ĐỔI PDF SANG WORD")
# ============================================================================

def detect_pdf_text_layer_info(pdf_path):
    """
    Kiểm tra PDF có lớp chữ (text layer) đọc trực tiếp được không, hay là bản scan cần OCR -
    đúng bước 1 mục VI tài liệu ("Kiểm tra PDF có text layer hay không").
    Trả về dict {num_pages, has_text_layer, avg_chars_per_page, is_likely_scan}.
    """
    if not HAS_FITZ:
        raise RuntimeError("Thiếu thư viện PyMuPDF (fitz) để kiểm tra PDF.")
    doc = fitz.open(pdf_path)
    try:
        num_pages = len(doc)
        total_chars = sum(len(doc[i].get_text().strip()) for i in range(num_pages))
        avg_chars = total_chars / num_pages if num_pages else 0
        # Ngưỡng thực nghiệm: dưới ~20 ký tự/trang coi như KHÔNG có text layer đủ dùng (có thể chỉ
        # là watermark/số trang lẻ tẻ, không phải nội dung chữ thật).
        has_text = avg_chars >= 20
        return {"num_pages": num_pages, "has_text_layer": has_text,
                "avg_chars_per_page": avg_chars, "is_likely_scan": not has_text}
    finally:
        doc.close()


def convert_pdf_to_docx_giu_bo_cuc(pdf_path, docx_path, log_cb=None):
    """
    Chế độ "Giữ bố cục" (mục V.2 tài liệu) - dùng cho công văn, biểu mẫu, bảng biểu. Dùng
    pdf2docx với cấu hình ưu tiên nhận diện bảng chính xác hơn (mục VI: "Có thể dùng... thư viện
    pdf2docx"). KHÔNG chỉ chèn ảnh trang - pdf2docx phục dựng đoạn văn/bảng thành đối tượng Word
    thật khi PDF có text layer.
    """
    if not HAS_PDF2DOCX:
        raise RuntimeError("Thiếu thư viện pdf2docx. Cài: pip install pdf2docx")
    cv = PDF2DocxConverter(pdf_path)
    try:
        # connected_border_tolerance/min_border_clearance: cấu hình pdf2docx giúp nhận diện đúng
        # hơn các bảng có đường viền mảnh/không đều - hay gặp ở công văn hành chính scan-in-PDF.
        cv.convert(docx_path, table_settings={"connected_border_tolerance": 3.0})
    finally:
        cv.close()
    log_cb and log_cb(f"   Đã chuyển (Giữ bố cục): {os.path.basename(docx_path)}")


def convert_pdf_to_docx_nhanh(pdf_path, docx_path, log_cb=None):
    """Chế độ "Nhanh" (mục V.1) - PDF có text layer rõ ràng, chuyển nhanh bằng pdf2docx mặc định."""
    pdf_to_docx_single(pdf_path, docx_path, log_cb)
    log_cb and log_cb(f"   Đã chuyển (Nhanh): {os.path.basename(docx_path)}")


def convert_pdf_to_docx_ocr(pdf_path, docx_path, ocr_dpi=300, log_cb=None):
    """
    Chế độ "OCR bản quét" (mục V.3) - dùng cho PDF scan không có text layer. Render từng trang,
    OCR tiếng Việt+Anh, chèn kết quả dưới dạng ĐOẠN VĂN CÓ THỂ CHỈNH SỬA (không chỉ ảnh trang) -
    đúng yêu cầu cốt lõi mục IV ("Không chỉ chèn toàn bộ trang PDF thành ảnh vào Word").

    LƯU Ý PHẠM VI: đây là bản OCR CƠ BẢN - phục dựng đoạn văn theo khối văn bản OCR nhận diện
    được (dùng pytesseract image_to_data để nhóm theo dòng/khối), CHƯA phục dựng được bảng biểu
    thật từ ảnh scan (việc nhận diện lưới bảng từ OCR/ảnh là bài toán lớn hơn nhiều, để triển khai
    riêng khi có nhu cầu cụ thể + đủ dữ liệu mẫu kiểm thử). Trang có khả năng chứa bảng phức tạp
    sẽ được ghi chú CẦN KIỂM TRA BẢNG trong log thay vì âm thầm bỏ qua.
    """
    if not (HAS_FITZ and HAS_OCR):
        raise RuntimeError("Thiếu thư viện PyMuPDF/pytesseract để OCR PDF scan.")
    if not HAS_DOCXCOMPOSE:
        raise RuntimeError("Thiếu thư viện python-docx.")

    doc = fitz.open(pdf_path)
    docx_doc = DocxDocument()
    warnings_list = []
    try:
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img = ImageOps.autocontrast(img)

            text = pytesseract.image_to_string(img, lang="vie+eng")
            text = unicodedata.normalize("NFC", text)

            if page_idx > 0:
                docx_doc.add_page_break()
            for para_text in text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    docx_doc.add_paragraph(para_text)

            # Phát hiện khả năng có bảng (nhiều dòng ngắn liên tiếp thẳng hàng - dấu hiệu gián
            # tiếp, không chắc chắn tuyệt đối) để cảnh báo CẦN KIỂM TRA thay vì im lặng bỏ qua.
            lines = [l for l in text.split("\n") if l.strip()]
            short_lines = sum(1 for l in lines if len(l.strip()) < 25)
            if lines and short_lines / len(lines) > 0.5 and len(lines) > 6:
                warnings_list.append(f"Trang {page_idx + 1}: có thể chứa bảng biểu - "
                                     f"OCR CHƯA phục dựng được bảng thật, cần kiểm tra tay")
        docx_doc.save(docx_path)
    finally:
        doc.close()

    if warnings_list and log_cb:
        for w in warnings_list:
            log_cb(f"   ⚠ {w}")
    log_cb and log_cb(f"   Đã chuyển (OCR bản quét): {os.path.basename(docx_path)}")
    return warnings_list


def convert_pdf_to_docx_smart(pdf_path, docx_path, mode="tu_nhan_dien", ocr_dpi=300, log_cb=None):
    """
    Hàm điều phối 3 chế độ chuyển đổi PDF sang Word (mục V + VI tài liệu). mode:
    "nhanh" | "giu_bo_cuc" | "ocr" | "tu_nhan_dien" (tự chọn theo mục VI: có text layer thì
    dùng pdf2docx, không thì OCR).
    Trả về dict {mode_da_dung, canh_bao: [...], trang_thai}.
    """
    log_cb = log_cb or (lambda x: None)
    info = detect_pdf_text_layer_info(pdf_path)

    actual_mode = mode
    if mode == "tu_nhan_dien":
        actual_mode = "ocr" if info["is_likely_scan"] else "giu_bo_cuc"
        log_cb(f"   Tự nhận diện: {'PDF scan, cần OCR' if info['is_likely_scan'] else 'Có text layer'} "
              f"→ dùng chế độ '{actual_mode}'")

    canh_bao = []
    if actual_mode == "nhanh":
        if info["is_likely_scan"]:
            canh_bao.append("PDF có vẻ là bản scan (không có text layer) nhưng chọn chế độ 'Nhanh' - "
                            "kết quả có thể rỗng/thiếu nội dung, khuyến nghị dùng chế độ 'OCR bản quét'")
        convert_pdf_to_docx_nhanh(pdf_path, docx_path, log_cb)
    elif actual_mode == "giu_bo_cuc":
        if info["is_likely_scan"]:
            canh_bao.append("PDF có vẻ là bản scan (không có text layer) nhưng chọn chế độ 'Giữ bố cục' - "
                            "kết quả có thể rỗng/thiếu nội dung, khuyến nghị dùng chế độ 'OCR bản quét'")
        convert_pdf_to_docx_giu_bo_cuc(pdf_path, docx_path, log_cb)
    elif actual_mode == "ocr":
        warnings_from_ocr = convert_pdf_to_docx_ocr(pdf_path, docx_path, ocr_dpi, log_cb)
        canh_bao.extend(warnings_from_ocr)
    else:
        raise ValueError(f"Chế độ chuyển đổi không hợp lệ: {mode}")

    trang_thai = "CAN_KIEM_TRA_BANG" if canh_bao else "THANH_CONG"
    return {"mode_da_dung": actual_mode, "canh_bao": canh_bao, "trang_thai": trang_thai,
            "num_pages": info["num_pages"]}


def check_docx_conversion_quality(pdf_path, docx_path):
    """
    Kiểm tra chất lượng sau chuyển đổi PDF→Word (mục XVI tài liệu) - phát hiện bất thường thay vì
    coi mọi lần chuyển đổi "chạy xong không lỗi" là hoàn thành đúng.
    Trả về dict {ok, canh_bao: [...]}.
    """
    canh_bao = []
    try:
        pdf_info = detect_pdf_text_layer_info(pdf_path)
        docx_doc = DocxDocument(docx_path)
        text_len = sum(len(p.text) for p in docx_doc.paragraphs)
        n_tables = len(docx_doc.tables)

        if pdf_info["avg_chars_per_page"] > 50 and text_len < pdf_info["avg_chars_per_page"] * 0.3:
            canh_bao.append("Lượng văn bản trong Word ít hơn đáng kể so với PDF gốc - có thể bị thiếu nội dung")

        # Lỗi font/mã hóa hay gặp: dấu "?" thay cho ký tự tiếng Việt (mục VII.2 tài liệu)
        full_text = "\n".join(p.text for p in docx_doc.paragraphs)
        if re.search(r"[A-Za-z]\?[A-Za-z]", full_text):
            canh_bao.append("Phát hiện dấu '?' xen giữa chữ cái - có thể lỗi font/mã hóa tiếng Việt")

        if pdf_info["num_pages"] > 1 and text_len == 0 and n_tables == 0:
            canh_bao.append("File Word không có nội dung văn bản lẫn bảng - có khả năng chuyển đổi thất bại")
    except Exception as e:
        canh_bao.append(f"Không kiểm tra được chất lượng: {e}")

    return {"ok": len(canh_bao) == 0, "canh_bao": canh_bao}


def batch_export_word(folder, output_folder, merge_into_one, merged_name, log_cb, control=None):
    if not HAS_PDF2DOCX:
        raise RuntimeError("Thiếu thư viện pdf2docx. Cài: pip install pdf2docx")
    os.makedirs(output_folder, exist_ok=True)
    files = [f for f in sorted(os.listdir(folder)) if f.lower().endswith(".pdf")]
    if not files:
        raise RuntimeError("Không có file PDF nào trong thư mục đã chọn.")

    docx_paths = []
    for i, fname in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        base = os.path.splitext(fname)[0]
        out_docx = os.path.join(output_folder, base + ".docx")
        try:
            pdf_to_docx_single(os.path.join(folder, fname), out_docx, log_cb)
            log_cb(f"✓ Đã xuất: {fname}  →  {os.path.basename(out_docx)}")
            docx_paths.append(out_docx)
        except Exception as e:
            log_cb(f"✗ Lỗi xuất {fname}: {e}")

    if merge_into_one and docx_paths:
        if not HAS_DOCXCOMPOSE:
            raise RuntimeError("Thiếu thư viện docxcompose. Cài: pip install docxcompose python-docx")
        log_cb("Đang gộp tất cả vào 1 file .docx...")
        master = DocxDocument(docx_paths[0])
        composer = DocxComposer(master)
        for p in docx_paths[1:]:
            composer.doc.add_page_break()
            composer.append(DocxDocument(p))
        merged_path = os.path.join(output_folder, merged_name)
        composer.save(merged_path)
        log_cb(f"✓ Đã gộp thành: {merged_path}")
        return docx_paths, merged_path


def batch_export_word_smart(source, output_folder, mode, name_suffix, no_overwrite, log_cb,
                             ocr_dpi=300, control=None, progress_cb=None):
    """
    Chuyển đổi PDF sang Word hàng loạt - PHIÊN BẢN NÂNG CAO hỗ trợ 3 chế độ (mục V tài liệu
    "CHỈNH SỬA TAB 7+9... BỔ SUNG CHUYỂN ĐỔI PDF SANG WORD"). Khác `batch_export_word()` (giữ
    nguyên không đổi, dùng cho luồng đơn giản cũ) - hàm này có kiểm tra chất lượng sau chuyển đổi,
    hậu tố tên file tùy chọn, không ghi đè, log đầy đủ theo mục XIV.

    mode: "nhanh" | "giu_bo_cuc" | "ocr" | "tu_nhan_dien"
    Trả về list log_rows (dict) - 1 dòng/file, đủ cột theo mục XIV tài liệu.
    """
    os.makedirs(output_folder, exist_ok=True)
    files = list_files_from_source(source, ".pdf")
    if not files:
        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")

    log_rows = []
    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        fname = os.path.basename(fpath)
        if progress_cb:
            progress_cb(i + 1, len(files), fname)
        t0 = time.time()

        base = os.path.splitext(fname)[0]
        suffix_part = f"_{name_suffix}" if name_suffix else ""
        out_docx = os.path.join(output_folder, base + suffix_part + ".docx")
        if no_overwrite and os.path.isfile(out_docx):
            k = 1
            while os.path.isfile(os.path.join(output_folder, f"{base}{suffix_part}_{k}.docx")):
                k += 1
            out_docx = os.path.join(output_folder, f"{base}{suffix_part}_{k}.docx")

        row = {"stt": i + 1, "file_in": fname, "path_in": fpath, "loai_in": "PDF",
              "loai_chuyen_doi": "PDF_SANG_WORD", "so_trang": "", "co_text_layer": "",
              "co_ocr": "", "co_bang": "", "che_do": mode, "file_out": "", "path_out": "",
              "thoi_gian_xu_ly": "", "status": "", "note": ""}
        try:
            info = detect_pdf_text_layer_info(fpath)
            row["so_trang"] = info["num_pages"]
            row["co_text_layer"] = "CÓ" if info["has_text_layer"] else "KHÔNG"

            result = convert_pdf_to_docx_smart(fpath, out_docx, mode=mode, ocr_dpi=ocr_dpi, log_cb=log_cb)
            row["co_ocr"] = "CÓ" if result["mode_da_dung"] == "ocr" else "KHÔNG"
            row["che_do"] = result["mode_da_dung"]

            quality = check_docx_conversion_quality(fpath, out_docx)
            row["co_bang"] = "CẦN KIỂM TRA" if result["canh_bao"] else "OK"
            row["file_out"] = os.path.basename(out_docx)
            row["path_out"] = out_docx
            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)

            all_warnings = result["canh_bao"] + quality["canh_bao"]
            if all_warnings:
                row["status"] = "CAN_KIEM_TRA_BANG" if result["canh_bao"] else "CAN_KIEM_TRA_BO_CUC"
                row["note"] = "; ".join(all_warnings)
                log_cb(f"⚠ {fname}: {row['note']}")
            else:
                row["status"] = "THANH_CONG"
                log_cb(f"✓ Đã chuyển: {fname} → {row['file_out']}")
        except Exception as e:
            row["status"] = "LOI_CHUYEN_PDF"
            row["note"] = str(e)
            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)
            log_cb(f"✗ Lỗi chuyển {fname}: {e}")

        log_rows.append(row)

    return log_rows


def write_log_chuyen_doi_word_pdf(csv_path, log_rows):
    """Ghi LOG_CHUYEN_DOI_WORD_PDF.csv - đúng mục XIV tài liệu."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file đầu vào", "Đường dẫn đầu vào", "Loại file đầu vào",
                    "Loại chuyển đổi", "Tổng số trang", "Có text layer", "Có OCR", "Có bảng",
                    "Chế độ chuyển đổi", "Tên file đầu ra", "Đường dẫn đầu ra", "Thời gian xử lý",
                    "Trạng thái", "Ghi chú"])
        for r in log_rows:
            w.writerow([r["stt"], r["file_in"], r["path_in"], r["loai_in"], r["loai_chuyen_doi"],
                       r["so_trang"], r["co_text_layer"], r["co_ocr"], r["co_bang"], r["che_do"],
                       r["file_out"], r["path_out"], r["thoi_gian_xu_ly"], r["status"], r["note"]])


# ============================================================================
# MỤC 7b (MỚI): XUẤT WORD HÀNG LOẠT TỪ FILE EXCEL BẤT KỲ
# ============================================================================

try:
    from app.services.excel_service import load_excel_rows_generic
    _EXCEL_SERVICE_MODULE_OK_A = True
except Exception:
    _EXCEL_SERVICE_MODULE_OK_A = False

    def load_excel_rows_generic(excel_path, sheet_name, header_row, col_map):
        """
        col_map: dict {ten_truong: chu_cot} VD {'maxa': 'B', 'to': 'W', 'thua': 'X', 'ten': 'H', ...}
        Trả về list dict {ten_truong: gia_tri}, bỏ qua dòng trống hoàn toàn.
        """
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

        rows = []
        for r in range(header_row, ws.max_row + 1):
            data = {}
            any_val = False
            for field, col in col_map.items():
                col = (col or "").strip()
                if not col:
                    data[field] = ""
                    continue
                v = ws[f"{col}{r}"].value
                if v is not None and str(v).strip() != "":
                    any_val = True
                data[field] = v if v is not None else ""
            if not any_val:
                continue
            if not data.get("stt"):
                data["stt"] = len(rows) + 1
            rows.append(data)
        return rows

class _SafeFormatDict(dict):
    def __missing__(self, key):
        return ""


def render_word_template(template_text, row_data):
    """Thay {ten_truong} trong template_text bằng giá trị tương ứng trong row_data (không phân biệt hoa/thường)."""
    clean = {}
    for k, v in row_data.items():
        if isinstance(v, float):
            v = round(v, 2)
            if v == int(v):
                v = int(v)
        clean[str(k).lower()] = "" if v is None else v
    mapping = _SafeFormatDict(clean)
    try:
        return template_text.format_map(mapping)
    except Exception:
        return template_text


def export_word_from_excel(excel_path, sheet_name, header_row, col_map, template_text,
                            output_folder, combine_into_one, filename_template, log_cb, control=None):
    if not HAS_DOCXCOMPOSE:
        raise RuntimeError("Thiếu thư viện python-docx/docxcompose. Cài: pip install python-docx docxcompose")

    rows = load_excel_rows_generic(excel_path, sheet_name, header_row, col_map)
    if not rows:
        raise RuntimeError("Không có dòng dữ liệu nào trong Excel theo cấu hình cột đã chọn.")

    os.makedirs(output_folder, exist_ok=True)
    docx_paths = []

    for idx, row_data in enumerate(rows):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(rows) - idx} dòng chưa xử lý).")
                break
        content = render_word_template(template_text, row_data)
        doc = DocxDocument()
        for line in content.split("\n"):
            doc.add_paragraph(line)

        fname_base = render_word_template(filename_template, row_data).strip() or f"ho_so_{row_data.get('stt')}"
        fname_base = sanitize_filename_part(fname_base)
        out_path = os.path.join(output_folder, fname_base + ".docx")
        base, ext = os.path.splitext(out_path)
        i = 2
        while os.path.isfile(out_path):
            out_path = f"{base}_{i}{ext}"
            i += 1

        from app.services import file_safety_service as _fsafe
        _fsafe.ghi_nguyen_tu(doc.save, out_path)
        docx_paths.append(out_path)
        log_cb(f"✓ Đã tạo: {os.path.basename(out_path)}")

    merged_path = None
    if combine_into_one and docx_paths:
        log_cb("Đang gộp tất cả vào 1 file .docx...")
        master = DocxDocument(docx_paths[0])
        composer = DocxComposer(master)
        for p in docx_paths[1:]:
            composer.doc.add_page_break()
            composer.append(DocxDocument(p))
        merged_path = os.path.join(output_folder, "TONG_HOP_TU_EXCEL.docx")
        composer.save(merged_path)
        log_cb(f"✓ Đã gộp thành: {merged_path}")

    return docx_paths, merged_path


# ============================================================================
# MỤC 7c (MỚI): XUẤT WORD THEO MẪU CÓ SẴN (MAIL-MERGE, GIỮ NGUYÊN BỐ CỤC GỐC)
# ============================================================================

# Các token dùng trong file mẫu .docx thật (không có dấu ngoặc, chèn thẳng vào văn bản
# lúc soạn mẫu trên Word), VD: "a) Tên: name" -> "name" sẽ được thay bằng giá trị thực.
MAILMERGE_TOKENS = ["name", "id", "location", "numberthua", "numberto", "Sdientichm2", "Sdientich",
                    "mucdich", "core", "year", "numbertb", "numberbb", "numbercv", "numberkh",
                    "sovanban", "kyhieu", "ngayvanban",
                    # Token địa bàn/thôn sáp nhập (dùng khi mẫu văn bản cần ghi rõ xã/thôn cũ-mới)
                    "maxa", "maxacu", "maxahientai", "xacu", "xahientai", "thon", "thoncu",
                    "thonmoi", "tinh", "nghiquyet", "ngaynghiquyet"]


def infer_dob_gender_from_cccd(cccd):
    """
    Suy NĂM SINH và GIỚI TÍNH từ số CCCD 12 số theo chuẩn mã hóa hiện hành.
    Lưu ý quan trọng: CCCD chỉ mã hóa NĂM sinh, KHÔNG mã hóa ngày/tháng sinh đầy đủ
    - vì vậy hàm này KHÔNG trả về ngày sinh đầy đủ, chỉ trả về (năm_sinh, giới_tính).
    Trả về (None, None) nếu số không hợp lệ.
    """
    if not cccd:
        return None, None
    cccd = re.sub(r"\D", "", str(cccd))
    if len(cccd) != 12:
        return None, None
    try:
        century_gender_digit = int(cccd[3])
        yy = int(cccd[4:6])
    except ValueError:
        return None, None
    gioi_tinh = "Nam" if century_gender_digit % 2 == 0 else "Nữ"
    the_ky_index = century_gender_digit // 2
    base_year = 1900 + the_ky_index * 100
    nam_sinh = base_year + yy
    return nam_sinh, gioi_tinh


try:
    from app.services.word_service import (
        WordCOMSession,
        mail_merge_via_word_com,
        mail_merge_smart,
    )
    _WORD_SERVICE_MODULE_OK_A = True
except Exception:
    _WORD_SERVICE_MODULE_OK_A = False

    class WordCOMSession:
        """
        Giữ 1 PHIÊN Microsoft Word MỞ DUY NHẤT, RIÊNG BIỆT cho toàn bộ đợt xuất hàng loạt, thay vì
        mở/tắt Word cho từng file (rất chậm - mỗi lần mở/tắt Word mất 1-3 giây). Dùng:

            with WordCOMSession() as session:
                for ...:
                    session.merge_and_save(template, replacements, out_path)
                    # hoặc
                    session.convert_to_pdf(input_path, out_path)
                    # hoặc
                    session.read_text(input_path)

        Word chỉ thật sự khởi động 1 lần khi vào `with`, và chỉ Quit() khi ra khỏi `with`.

        TUÂN THỦ CÁC NGUYÊN TẮC AN TOÀN COM BẮT BUỘC (tránh treo WINWORD.EXE ngầm, tránh ảnh hưởng
        phiên Word người dùng đang mở thủ công):
        - Dùng DispatchEx (KHÔNG dùng Dispatch/gencache.EnsureDispatch) -> LUÔN tạo 1 tiến trình
          WINWORD.EXE MỚI, RIÊNG BIỆT, không bao giờ kết nối/dùng chung với phiên Word người dùng
          đang mở sẵn trên máy.
        - Mở file ReadOnly=True + AddToRecentFiles=False cho các thao tác chỉ ĐỌC (xuất PDF, đọc nội
          dung) - không sửa file gốc, không thêm vào danh sách "Recent Files" của người dùng.
        - ConfirmConversions=False để không hiện hộp thoại xác nhận định dạng khi mở .doc cũ.
        - MỌI document đều được đóng bằng doc.Close(False) trong khối finally - kể cả khi lỗi.
        - Khi thoát phiên: Quit() + giải phóng tham chiếu COM + gc.collect() để đảm bảo tiến trình
          WINWORD.EXE do phần mềm tạo ra kết thúc hẳn, không treo ngầm.
        - Gọi pythoncom.CoInitialize()/CoUninitialize() đúng cặp cho thread đang chạy COM (vì phần
          mềm chạy tác vụ này trong thread nền riêng).
        - KHÔNG BAO GIỜ dùng taskkill/kill toàn bộ WINWORD.EXE - chỉ Quit() đúng instance do chính
          phiên này tạo ra, không đụng đến Word người dùng đang mở thủ công.
        """

        def __init__(self):
            self.word = None
            self._com_initialized = False

        def __enter__(self):
            if not HAS_WIN32COM:
                raise RuntimeError("Cần Windows + Microsoft Word + pywin32 để dùng chế độ này.")

            try:
                import pythoncom
                pythoncom.CoInitialize()
                self._com_initialized = True
            except Exception:
                self._com_initialized = False

            self.word = win32com_client.DispatchEx("Word.Application")
            self.word.Visible = False
            try:
                self.word.DisplayAlerts = 0
            except Exception:
                pass
            try:
                self.word.ScreenUpdating = False
            except Exception:
                pass
            try:
                self.word.ConfirmConversions = False
            except Exception:
                pass
            return self

        def merge_and_save(self, template_path, replacements, output_path):
            tokens_sorted = sorted(replacements.keys(), key=len, reverse=True)
            doc = None
            try:
                doc = self.word.Documents.Open(os.path.abspath(template_path), AddToRecentFiles=False)
                for token in tokens_sorted:
                    value = replacements.get(token)
                    if value is None:
                        continue
                    find = doc.Content.Find
                    find.ClearFormatting()
                    find.Replacement.ClearFormatting()
                    find.Text = token
                    find.Replacement.Text = str(value)
                    find.Execute(Replace=2, Forward=True, Wrap=1)
                doc.SaveAs2(os.path.abspath(output_path), FileFormat=16)
            finally:
                if doc is not None:
                    doc.Close(False)
                    doc = None

        def convert_to_pdf(self, input_path, output_path):
            doc = None
            try:
                doc = self.word.Documents.Open(
                    os.path.abspath(input_path), ReadOnly=True, AddToRecentFiles=False)
                doc.ExportAsFixedFormat(OutputFileName=os.path.abspath(output_path), ExportFormat=17)
            finally:
                if doc is not None:
                    doc.Close(False)
                    doc = None

        def read_text(self, input_path):
            doc = None
            try:
                doc = self.word.Documents.Open(
                    os.path.abspath(input_path), ReadOnly=True, AddToRecentFiles=False)
                return doc.Content.Text
            finally:
                if doc is not None:
                    doc.Close(False)
                    doc = None

        def __exit__(self, exc_type, exc_val, exc_tb):
            if self.word is not None:
                try:
                    self.word.Quit()
                except Exception:
                    pass
                self.word = None

            import gc
            gc.collect()

            if self._com_initialized:
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
                self._com_initialized = False
            return False

    def mail_merge_via_word_com(template_path, replacements, output_path):
        if not HAS_WIN32COM:
            raise RuntimeError("Cần Windows + Microsoft Word + pywin32 để dùng chế độ này.")
        with WordCOMSession() as session:
            session.merge_and_save(template_path, replacements, output_path)

    def mail_merge_smart(template_path, replacements, output_path, log_cb=None, word_session=None):
        if word_session is not None:
            try:
                word_session.merge_and_save(template_path, replacements, output_path)
                return "com (phiên chung)"
            except Exception as e:
                if log_cb:
                    log_cb(f"   (COM lỗi ({e}), chuyển sang phương án dự phòng python-docx)")
        elif HAS_WIN32COM:
            try:
                mail_merge_via_word_com(template_path, replacements, output_path)
                return "com"
            except Exception as e:
                if log_cb:
                    log_cb(f"   (COM lỗi ({e}), chuyển sang phương án dự phòng python-docx)")

        if template_path.lower().endswith(".doc"):
            raise RuntimeError(
                "File mẫu định dạng .doc cũ cần Microsoft Word (qua pywin32) để mở trên máy này. "
                "Hãy mở file bằng Word và 'Save As' sang .docx rồi thử lại, hoặc cài Microsoft Word."
            )
        mail_merge_docx_template(template_path, replacements, output_path)
        return "python-docx"


# ============================================================================
# MỤC 7e (MỚI): CHUYỂN WORD SANG PDF HÀNG LOẠT
# ============================================================================

class TaskCancelled(Exception):
    """Ném ra khi người dùng bấm Hủy giữa lúc đang chạy 1 tác vụ hàng loạt."""
    pass


class TaskControl:
    """
    Điều khiển Tạm dừng / Tiếp tục / Hủy cho 1 tác vụ chạy trong thread nền.
    Dùng: gọi control.checkpoint() ở đầu mỗi vòng lặp xử lý từng file/dòng.
    """
    def __init__(self):
        self.pause_event = threading.Event()
        self.cancel_event = threading.Event()

    def pause(self):
        self.pause_event.set()

    def resume(self):
        self.pause_event.clear()

    def cancel(self):
        self.cancel_event.set()
        self.pause_event.clear()  # tránh treo mãi ở trạng thái tạm dừng khi đã hủy

    def reset(self):
        self.pause_event.clear()
        self.cancel_event.clear()

    def is_paused(self):
        return self.pause_event.is_set()

    def checkpoint(self):
        """Gọi ở đầu mỗi vòng lặp: block khi đang tạm dừng, ném TaskCancelled khi bị hủy."""
        while self.pause_event.is_set():
            if self.cancel_event.is_set():
                raise TaskCancelled()
            time.sleep(0.15)
        if self.cancel_event.is_set():
            raise TaskCancelled()


try:
    from app.services.word_service import (
        convert_word_to_pdf_com,
        convert_word_to_pdf_libreoffice,
        convert_word_to_pdf_libreoffice_batch,
        convert_word_to_pdf_smart,
    )
    _WORD_SERVICE_MODULE_OK_B = True
except Exception:
    _WORD_SERVICE_MODULE_OK_B = False

    def convert_word_to_pdf_com(input_path, output_path):
        if not HAS_WIN32COM:
            raise RuntimeError("Cần Windows + Microsoft Word + pywin32 để dùng chế độ này.")
        with WordCOMSession() as session:
            session.convert_to_pdf(input_path, output_path)

    def convert_word_to_pdf_libreoffice_batch(input_paths, out_dir, timeout=None):
        """
        Chuyển ĐỒNG THỜI nhiều file Word sang PDF bằng 1 LẦN GỌI soffice DUY NHẤT (thay vì 1 lần
        gọi/file) - LibreOffice tốn chi phí KHỞI ĐỘNG ứng dụng rất lớn (đo thực tế: ~2s/file khi
        gọi riêng lẻ, chỉ ~0.4-0.6s/file khi gộp - nhanh hơn 3.5-4.5 lần, càng nhiều file càng lợi
        vì chi phí khởi động chỉ trả 1 LẦN DUY NHẤT cho cả lô thay vì trả lại mỗi file). Tất cả
        input_paths PHẢI xuất ra CÙNG 1 out_dir (giới hạn của tham số --outdir của soffice).
        timeout: giây, mặc định 60 + 15s/file (đủ dư cho lô lớn, không set cứng 1 con số cho mọi
        kích thước lô).
        Trả về set các file .pdf ĐÃ TẠO THÀNH CÔNG (không raise lỗi cho từng file lẻ - người gọi tự
        đối chiếu file nào có trong set để biết file nào chuyển đổi thành công/thất bại).
        """
        soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice_bin:
            raise RuntimeError(
                "Không tìm thấy Microsoft Word (pywin32) lẫn LibreOffice trên máy này để chuyển đổi."
            )
        os.makedirs(out_dir, exist_ok=True)
        if timeout is None:
            timeout = 60 + 15 * len(input_paths)
        abs_paths = [os.path.abspath(p) for p in input_paths]
        subprocess.run(
            [soffice_bin, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", out_dir] + abs_paths,
            capture_output=True, text=True, timeout=timeout,
        )
        created = set()
        for p in input_paths:
            expected = os.path.join(out_dir, os.path.splitext(os.path.basename(p))[0] + ".pdf")
            if os.path.isfile(expected):
                created.add(p)
        return created

    def convert_word_to_pdf_libreoffice(input_path, output_path, timeout=120):
        soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice_bin:
            raise RuntimeError(
                "Không tìm thấy Microsoft Word (pywin32) lẫn LibreOffice trên máy này để chuyển đổi."
            )
        out_dir = os.path.dirname(os.path.abspath(output_path))
        os.makedirs(out_dir, exist_ok=True)
        result = subprocess.run(
            [soffice_bin, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", out_dir,
             os.path.abspath(input_path)],
            capture_output=True, text=True, timeout=timeout,
        )
        produced = os.path.join(out_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        if not os.path.isfile(produced):
            raise RuntimeError(f"Chuyển đổi thất bại: {(result.stderr or result.stdout or '').strip()[:300]}")
        if os.path.abspath(produced) != os.path.abspath(output_path):
            shutil.move(produced, output_path)

    def convert_word_to_pdf_smart(input_path, output_path, log_cb=None, word_session=None):
        if word_session is not None:
            try:
                word_session.convert_to_pdf(input_path, output_path)
                return "Microsoft Word (COM, phiên chung)"
            except Exception as e:
                if log_cb:
                    log_cb(f"   (Word COM lỗi: {e} - thử LibreOffice...)")
        elif HAS_WIN32COM:
            try:
                convert_word_to_pdf_com(input_path, output_path)
                return "Microsoft Word (COM)"
            except Exception as e:
                if log_cb:
                    log_cb(f"   (Word COM lỗi: {e} - thử LibreOffice...)")
        convert_word_to_pdf_libreoffice(input_path, output_path)
        return "LibreOffice"

WORD_EXTS = (".doc", ".docx", ".docm")


def batch_word_to_pdf(source_folder, output_folder, include_subfolders, overwrite,
                       keep_folder_structure, dry_run, log_cb, control=None, word_session=None,
                       progress_cb=None):
    """
    Chuyển hàng loạt file Word (.doc/.docx/.docm) trong source_folder sang PDF cùng tên.
    word_session: 1 WordCOMSession đã mở sẵn (nếu có) để dùng CHUNG cho toàn bộ đợt xuất,
    tránh mở/tắt Word cho từng file - nhanh hơn nhiều khi có nhiều file.
    progress_cb(idx, total, filename): gọi trước khi xử lý mỗi file - dùng cho thanh tiến độ.

    TỐI ƯU TỐC ĐỘ: khi KHÔNG có Word COM (word_session=None và không có pywin32/Windows), dùng
    LibreOffice làm phương án duy nhất - các file được GỘP LẠI theo cùng 1 thư mục đích rồi
    chuyển đổi bằng 1 LẦN GỌI soffice DUY NHẤT cho cả nhóm thay vì mỗi file 1 lần gọi riêng
    (đo thực tế: nhanh hơn 3.5-4.5 lần vì chi phí khởi động LibreOffice tốn nhất, chỉ trả 1 lần
    cho cả lô thay vì trả lại mỗi file). Khi CÓ Word COM, vẫn giữ nguyên vòng lặp tuần tự với
    session dùng chung như trước (đã đủ hiệu quả, không cần gộp).

    Trả về list dict {stt, word, pdf, status, note} dùng cho bảng trạng thái + log CSV.
    """
    files = []
    if include_subfolders:
        for root, dirs, fnames in os.walk(source_folder):
            for f in fnames:
                if f.lower().endswith(WORD_EXTS) and not f.startswith("~$"):
                    files.append(os.path.join(root, f))
    else:
        for f in sorted(os.listdir(source_folder)):
            fpath = os.path.join(source_folder, f)
            if os.path.isfile(fpath) and f.lower().endswith(WORD_EXTS) and not f.startswith("~$"):
                files.append(fpath)
    files.sort()

    if not files:
        raise RuntimeError("Không tìm thấy file Word (.doc/.docx/.docm) nào trong nguồn đã chọn.")

    results = [None] * len(files)
    to_convert = []  # (idx0, fpath, fname, out_dir, out_path, pdf_name)

    for idx0, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - idx0} file chưa xử lý).")
                results = results[:idx0]
                break

        fname = os.path.basename(fpath)
        progress_cb and progress_cb(idx0 + 1, len(files), fname)
        if keep_folder_structure:
            rel_dir = os.path.dirname(os.path.relpath(fpath, source_folder))
            out_dir = os.path.join(output_folder, rel_dir) if rel_dir else output_folder
        else:
            out_dir = output_folder
        pdf_name = os.path.splitext(fname)[0] + ".pdf"
        out_path = os.path.join(out_dir, pdf_name)

        if not overwrite and os.path.isfile(out_path):
            results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": pdf_name,
                             "status": "BỎ QUA", "note": "File PDF đã tồn tại (chưa tick ghi đè)"}
            log_cb(f"⏭ Bỏ qua (đã tồn tại): {fname}")
            continue

        if dry_run:
            results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": pdf_name,
                             "status": "OK (chạy thử - chưa xuất)", "note": ""}
            log_cb(f"✓ (chạy thử) {fname} → {pdf_name}")
            continue

        to_convert.append((idx0, fpath, fname, out_dir, out_path, pdf_name))

    use_batch_libreoffice = (word_session is None and not HAS_WIN32COM)

    if use_batch_libreoffice and to_convert:
        groups = {}
        for item in to_convert:
            groups.setdefault(item[3], []).append(item)  # nhóm theo out_dir

        for out_dir, items in groups.items():
            if control:
                try:
                    control.checkpoint()
                except TaskCancelled:
                    log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(items)} file trong nhóm chưa xử lý).")
                    break
            try:
                os.makedirs(out_dir, exist_ok=True)
                input_paths = [it[1] for it in items]
                log_cb(f"→ Đang chuyển gộp {len(items)} file cùng lúc bằng LibreOffice (nhanh hơn "
                      f"chuyển từng file riêng lẻ) vào: {out_dir}")
                created = convert_word_to_pdf_libreoffice_batch(input_paths, out_dir)
                for idx0, fpath, fname, _od, out_path, pdf_name in items:
                    if fpath in created:
                        results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": pdf_name,
                                         "status": "OK", "note": "LibreOffice (gộp nhiều file)"}
                        log_cb(f"✓ {fname} → {pdf_name}  (LibreOffice, gộp)")
                    else:
                        results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": "", "status": "LỖI",
                                         "note": "Không tạo được PDF khi chuyển gộp (xem log LibreOffice)"}
                        log_cb(f"✗ Lỗi (gộp): {fname}: không tạo được PDF")
            except Exception as e:
                for idx0, fpath, fname, _od, out_path, pdf_name in items:
                    results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": "", "status": "LỖI", "note": str(e)}
                log_cb(f"✗ Lỗi chuyển gộp nhóm {out_dir}: {e}")
    else:
        for idx0, fpath, fname, out_dir, out_path, pdf_name in to_convert:
            if control:
                try:
                    control.checkpoint()
                except TaskCancelled:
                    log_cb("⏹ Đã hủy theo yêu cầu.")
                    break
            try:
                os.makedirs(out_dir, exist_ok=True)
                method = convert_word_to_pdf_smart(fpath, out_path, log_cb, word_session=word_session)
                results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": pdf_name, "status": "OK", "note": method}
                log_cb(f"✓ {fname} → {pdf_name}  ({method})")
            except Exception as e:
                results[idx0] = {"stt": idx0 + 1, "word": fname, "pdf": "", "status": "LỖI", "note": str(e)}
                log_cb(f"✗ Lỗi: {fname}: {e}")

    return [r for r in results if r is not None]


def write_word_to_pdf_log_csv(csv_path, results):
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file Word", "Tên file PDF xuất ra", "Trạng thái", "Ghi chú lỗi"])
        for r in results:
            w.writerow([r["stt"], r["word"], r["pdf"], r["status"], r["note"]])


try:
    from app.services.excel_service import read_excel_rows_smart
    _EXCEL_SERVICE_MODULE_OK_B = True
except Exception:
    _EXCEL_SERVICE_MODULE_OK_B = False

    def read_excel_rows_smart(excel_path, sheet_name, header_row, col_map, log_cb=None):
        """
        Đọc dữ liệu Excel thành list dict theo col_map (giống load_excel_rows_generic).
        - File .xlsx/.xlsm: dùng openpyxl (nhanh, ổn định, không cần cài Excel).
        - File .xls cũ: dùng Microsoft Excel qua COM (pywin32) nếu có (Windows + đã cài Excel),
          vì openpyxl không đọc được định dạng .xls cũ.
        """
        ext = os.path.splitext(excel_path)[1].lower()
        if ext in (".xlsx", ".xlsm"):
            return load_excel_rows_generic(excel_path, sheet_name, header_row, col_map)

        if ext == ".xls":
            if not HAS_WIN32COM:
                raise RuntimeError(
                    "File .xls (Excel cũ) cần Microsoft Excel + pywin32 để đọc trên máy này. "
                    "Hãy mở file bằng Excel và 'Save As' sang .xlsx rồi thử lại."
                )
            com_initialized = False
            try:
                import pythoncom
                pythoncom.CoInitialize()
                com_initialized = True
            except Exception:
                pass

            excel = win32com_client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            wb = None
            try:
                wb = excel.Workbooks.Open(os.path.abspath(excel_path), ReadOnly=True,
                                           AddToMru=False, UpdateLinks=0)
                try:
                    ws = wb.Sheets(sheet_name) if sheet_name else wb.Sheets(1)
                    used = ws.UsedRange
                    max_row = used.Row + used.Rows.Count - 1
                    rows = []
                    for r in range(header_row, max_row + 1):
                        data = {}
                        any_val = False
                        for field, col in col_map.items():
                            col = (col or "").strip()
                            if not col:
                                data[field] = ""
                                continue
                            v = ws.Range(f"{col}{r}").Value
                            if v is not None and str(v).strip() != "":
                                any_val = True
                            data[field] = v if v is not None else ""
                        if not any_val:
                            continue
                        if not data.get("stt"):
                            data["stt"] = len(rows) + 1
                        rows.append(data)
                    return rows
                finally:
                    if wb is not None:
                        wb.Close(False)
                        wb = None
            finally:
                try:
                    excel.Quit()
                except Exception:
                    pass
                excel = None
                import gc
                gc.collect()
                if com_initialized:
                    try:
                        import pythoncom
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        raise RuntimeError(f"Định dạng file Excel không được hỗ trợ: {ext}")

try:
    from app.services.word_service import mail_merge_docx_template
    _WORD_SERVICE_MODULE_OK_D = True
except Exception:
    _WORD_SERVICE_MODULE_OK_D = False

    def mail_merge_docx_template(template_path, replacements, output_path):
        """
        Mở file .docx mẫu (có chứa các token thường như 'name', 'numberthua'... chèn thẳng
        trong văn bản), thay bằng giá trị thực trong `replacements`, GIỮ NGUYÊN toàn bộ bố cục/
        định dạng gốc của file mẫu. Lưu kết quả ra output_path.
        """
        if not HAS_DOCXCOMPOSE:
            raise RuntimeError("Thiếu thư viện python-docx. Cài: pip install python-docx")

        doc = DocxDocument(template_path)
        tokens_sorted = sorted(replacements.keys(), key=len, reverse=True)

        def replace_in_paragraph(p):
            full_text = "".join(run.text for run in p.runs)
            if not full_text:
                return
            new_text = full_text
            for token in tokens_sorted:
                value = replacements.get(token)
                if value is None:
                    continue
                new_text = re.sub(re.escape(token), lambda m, v=str(value): v, new_text)
            if new_text != full_text:
                if p.runs:
                    p.runs[0].text = new_text
                    for r in p.runs[1:]:
                        r.text = ""

        for p in doc.paragraphs:
            replace_in_paragraph(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

        from app.services import file_safety_service as _fsafe
        _fsafe.ghi_nguyen_tu(doc.save, output_path)

def batch_mail_merge_from_pdf_or_excel(rows, template_path, output_folder, filename_template, log_cb,
                                        control=None, progress_cb=None):
    """
    rows: list dict (từ compile_report_from_pdfs hoặc load_excel_rows_generic), mỗi dict cần có
    các khóa tương ứng token: name/ten, id, location/dia_chi, numberthua/thua, numberto/to,
    Sdientich/dt, mucdich, core, year.
    progress_cb(idx, total, ten_dong): gọi trước khi xử lý mỗi dòng - dùng cho thanh tiến độ.
    Trả về (danh_sach_file_da_tao, danh_sach_loi).
    """
    if not os.path.isfile(template_path):
        raise RuntimeError(f"Không tìm thấy file mẫu: {template_path}")

    os.makedirs(output_folder, exist_ok=True)
    created = []
    errors = []

    for idx, row in enumerate(rows):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(rows) - idx} dòng chưa xử lý).")
                break
        progress_cb and progress_cb(idx + 1, len(rows), row.get("ten") or row.get("name") or f"dòng {idx+1}")
        replacements = {
            "name": row.get("ten") or row.get("name") or "",
            "id": row.get("id") or "",
            "location": row.get("dia_chi") or row.get("location") or "",
            "numberthua": row.get("thua") or row.get("numberthua") or "",
            "numberto": row.get("to") or row.get("numberto") or "",
            "Sdientich": row.get("dt") if row.get("dt") is not None else (row.get("Sdientich") or ""),
            "mucdich": row.get("mucdich") or "",
            "core": row.get("core") or "",
            "year": row.get("year") or "",
        }
        if isinstance(replacements["Sdientich"], float):
            v = round(replacements["Sdientich"], 2)
            replacements["Sdientich"] = int(v) if v == int(v) else v

        try:
            fname_base = render_word_template(filename_template, row).strip() or f"HoSo_{row.get('stt', '')}"
            fname_base = sanitize_filename_part(fname_base)
            out_path = os.path.join(output_folder, fname_base + ".docx")
            base, ext = os.path.splitext(out_path)
            i = 2
            while os.path.isfile(out_path):
                out_path = f"{base}_{i}{ext}"
                i += 1

            mail_merge_docx_template(template_path, replacements, out_path)
            created.append(out_path)
            log_cb(f"✓ Đã tạo: {os.path.basename(out_path)}")
        except Exception as e:
            errors.append((row.get("stt", "?"), str(e)))
            log_cb(f"✗ Lỗi tạo file cho dòng {row.get('stt', '?')}: {e}")

    return created, errors


# ============================================================================
# MỤC 7d (MỚI): XUẤT WORD HÀNG LOẠT - Excel tổng + mẫu Đơn đăng ký/Thông báo xác nhận
# ============================================================================

def resolve_maxa_for_row(row, thon_mapping, xa_mapping):
    """
    Xác định Mã xã cho 1 dòng dữ liệu:
    - Nếu Excel có sẵn cột Mã xã (row['maxa'] có giá trị) -> dùng trực tiếp.
    - Nếu không -> tra theo tên thôn/xã trong 'location' (ưu tiên thôn trước).
    Trả về (maxa, nguồn) hoặc (None, None).
    """
    maxa = row.get("maxa")
    if maxa and str(maxa).strip():
        return str(maxa).strip(), "cột Mã xã trong Excel"

    location_text = str(row.get("location") or row.get("dia_chi") or "")
    if not location_text.strip():
        return None, None
    norm = strip_diacritics(location_text).upper()

    for thon_key in sorted(thon_mapping.keys(), key=len, reverse=True):
        if thon_key and thon_key in norm:
            return thon_mapping[thon_key], f"thôn {thon_key}"
    for xa_key in sorted(xa_mapping.keys(), key=len, reverse=True):
        if xa_key and xa_key in norm:
            return xa_mapping[xa_key], f"xã {xa_key}"
    return None, None


def enrich_row_with_commune_config(row, commune_rows):
    """
    Bổ sung các trường thôn sáp nhập (thoncu/thonmoi/xacu/xahientai/maxahientai/nghiquyet...) vào
    1 dòng dữ liệu, dựa trên cấu hình địa bàn đã nạp. Dùng địa chỉ thửa đất/thường trú trong dòng
    để tra cứu. Nếu không nhận diện được thì GIỮ NGUYÊN dòng (không thêm gì, không tự bịa).
    An toàn: nếu module cấu hình chưa sẵn sàng hoặc chưa nạp cấu hình -> trả về dòng gốc.
    """
    if not commune_rows or not _COMMUNE_CONFIG_MODULE_OK:
        return row
    dia_chi_thua = row.get("dia_chi") or row.get("location") or row.get("land_address") or ""
    dia_chi_tt = row.get("dia_chi_thuong_tru") or row.get("permanent_address") or ""
    try:
        info = detect_ma_xa_by_config(dia_chi_thua, dia_chi_tt, commune_rows)
    except Exception:
        return row
    if info.get("status") != "NHAN_DIEN_THANH_CONG":
        return row
    # Chỉ điền các trường CÒN TRỐNG, không ghi đè dữ liệu đã có sẵn trong dòng
    for src_key, dst_key in [("xa_cu", "xacu"), ("xa_hien_tai", "xahientai"),
                              ("ma_xa_hien_tai", "maxahientai"), ("thon_cu", "thoncu"),
                              ("thon_moi", "thonmoi"), ("thon", "thon"),
                              ("nghi_quyet_so", "nghiquyet"), ("ngay_ban_hanh", "ngaynghiquyet")]:
        val = info.get(src_key)
        if val and not row.get(dst_key):
            row[dst_key] = val
    if info.get("ma_xa") and not row.get("maxacu"):
        row["maxacu"] = info["ma_xa"]
    return row


def build_mailmerge_replacements(row):
    """Chuẩn bị dict token -> giá trị: name/id/location/diachithua/numberthua/numberto/Sdientich/
    Sdientichm2/mucdich/year/core/numbertb/numberbb/numbercv/numberkh/sovanban/kyhieu/ngayvanban."""
    dt_val = row.get("Sdientich")
    if dt_val is None:
        dt_val = row.get("dt")
    replacements = {
        "name": row.get("name") or row.get("ten") or "",
        "id": row.get("id") or "",
        "location": row.get("location") or row.get("dia_chi") or "",
        "diachithua": row.get("diachithua") or "",
        "numberthua": row.get("numberthua") or row.get("thua") or "",
        "numberto": row.get("numberto") or row.get("to") or "",
        "Sdientich": dt_val if dt_val is not None else "",
        "Sdientichm2": dt_val if dt_val is not None else "",
        "mucdich": row.get("mucdich") or "",
        "year": row.get("year") or "",
        "core": row.get("core") or "",
        "numbertb": row.get("numbertb") or "",
        "numberbb": row.get("numberbb") or "",
        "numbercv": row.get("numbercv") or "",
        "numberkh": row.get("numberkh") or "",
        "sovanban": row.get("sovanban") or "",
        "kyhieu": row.get("kyhieu") or "",
        "ngayvanban": row.get("ngayvanban") or "",
        # Token địa bàn/thôn sáp nhập - lấy từ dữ liệu dòng nếu có (rỗng nếu không có, không tự bịa)
        "maxa": row.get("maxa") or row.get("ma_xa") or "",
        "maxacu": row.get("maxacu") or row.get("ma_xa_cu") or row.get("maxa") or "",
        "maxahientai": row.get("maxahientai") or row.get("ma_xa_hien_tai") or "",
        "xacu": row.get("xacu") or row.get("xa_cu") or "",
        "xahientai": row.get("xahientai") or row.get("xa_hien_tai") or "",
        "thon": row.get("thon") or "",
        "thoncu": row.get("thoncu") or row.get("thon_cu") or "",
        "thonmoi": row.get("thonmoi") or row.get("thon_moi") or "",
        "tinh": row.get("tinh") or row.get("tinh_thanh") or "",
        "nghiquyet": row.get("nghiquyet") or row.get("nghi_quyet_so") or "",
        "ngaynghiquyet": row.get("ngaynghiquyet") or row.get("ngay_ban_hanh") or "",
    }
    for k, v in list(replacements.items()):
        if isinstance(v, float):
            vv = round(v, 2)
            replacements[k] = int(vv) if vv == int(vv) else vv
    return replacements


def run_word_batch_export(rows, thon_mapping, xa_mapping, template_gt, template_tbxn,
                           export_gt, export_tbxn, combine, output_folder, dry_run, log_cb, control=None,
                           word_session=None, commune_rows=None, suffix_config=None, progress_cb=None,
                           toc_do_xu_ly="can_bang", so_file_moi_dot=100, nghi_giua_dot_giay=2.0,
                           perf_log_path=None, perf_stats_cb=None, gioi_han_ram_mb=0,
                           checkpoint_task_id=None, stt_da_xong=None):
    """
    Xử lý hàng loạt: mỗi dòng Excel -> xuất 1 (hoặc nhiều) file Word theo mẫu đã chọn.
    dry_run=True: KHÔNG xuất file thật, chỉ kiểm tra dữ liệu và trả về bảng trạng thái xem trước.
    word_session: 1 WordCOMSession đã mở sẵn (nếu có) để dùng CHUNG cho toàn bộ đợt xuất,
    tránh mở/tắt Word cho từng dòng - nhanh hơn nhiều khi có nhiều hồ sơ.
    commune_rows: cấu hình địa bàn đã nạp (nếu có) - dùng làm giàu token thôn sáp nhập cho mẫu Word.
    suffix_config: dict cấu hình hậu tố tên file tùy chọn (None hoặc {"enabled": False} = GIỮ NGUYÊN
    hành vi cũ - hậu tố cố định GT/TBXN/GOP theo mẫu đang xuất). Khi enabled=True, hậu tố được xác
    định theo resolve_file_suffix() (ưu tiên cột Excel > nhập tay > combobox > mặc định) và tên file
    được render theo suffix_config["filename_template"].
    progress_cb(idx, total, ten_dong): gọi trước khi xử lý mỗi dòng - dùng cho thanh tiến độ.

    TỐI ƯU HIỆU NĂNG (đúng tài liệu "TỐI ƯU HIỆU NĂNG KHI XUẤT WORD HÀNG LOẠT, KHÔNG LÀM TREO
    WINDOWS") - chỉ áp dụng khi word_session KHÔNG None (đang dùng Word COM thật, nơi phát sinh
    tải CPU/Disk I/O đáng kể) và dry_run=False (chạy thử không cần nghỉ, không ghi file thật):
    toc_do_xu_ly: "nhanh"|"can_bang"|"tiet_kiem" - mức nghỉ SAU MỖI FILE, nhường tài nguyên cho
    Windows Explorer/Excel/trình duyệt hoạt động bình thường trong lúc xuất.
    so_file_moi_dot / nghi_giua_dot_giay: sau mỗi bấy nhiêu file, nghỉ dài hơn 1 lần (mục VI).
    perf_log_path: nếu có, ghi LOG_HIEU_NANG_WORD.csv (mục XVI) - CPU/RAM/thời gian mỗi file.
    perf_stats_cb(dict): gọi định kỳ với số liệu CPU/RAM/số tiến trình Word hiện tại (mục XII/XV) -
    dùng để hiển thị trên giao diện, không bắt buộc.
    gioi_han_ram_mb: nếu > 0, khi RAM tiến trình vượt ngưỡng này thì NGHỈ THÊM 1 lần (gấp đôi mức
    nghỉ thông thường) trước khi tiếp tục - "van an toàn" đơn giản, không dừng hẳn (mục XVIII).

    CHECKPOINT/RESUME (mục XVII tài liệu) - CHỈ áp dụng khi checkpoint_task_id KHÔNG None:
    checkpoint_task_id: ID ổn định cho tác vụ này (do caller tính qua checkpoint_service.tinh_task_id) -
    lưu tiến độ định kỳ ra đĩa (Data/checkpoints/{task_id}.json), TỰ XÓA khi hoàn thành toàn bộ.
    stt_da_xong: set các STT (từ rows) đã xử lý XONG ở lần chạy TRƯỚC (nếu đang RESUME từ checkpoint
    dở dang) - các dòng này sẽ được BỎ QUA (không xử lý lại, không ghi đè file đã có). None/rỗng =
    xử lý từ đầu như bình thường (AN TOÀN MẶC ĐỊNH - không đổi hành vi cũ nếu không dùng checkpoint).

    Trả về list dict: {stt, ten, maxa, to, thua, filename, status, note, hauto, nguon_hauto,
    mau_ten_file} - dùng để hiển thị bảng + xuất CSV.
    """
    results = []
    if not dry_run:
        os.makedirs(output_folder, exist_ok=True)

    suffix_enabled = bool(suffix_config and suffix_config.get("enabled"))
    filename_template = (suffix_config or {}).get("filename_template") or DEFAULT_FILENAME_TEMPLATE

    ap_dung_toi_uu = (word_session is not None and not dry_run)
    if ap_dung_toi_uu:
        from app.services import word_perf_service as _wperf
        _wperf.apply_low_priority()
        sleep_per_file = _wperf.get_sleep_per_file(toc_do_xu_ly)
    else:
        sleep_per_file = 0.0

    bi_huy_giua_chung = False
    for idx, row in enumerate(rows):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(rows) - idx} dòng chưa xử lý).")
                bi_huy_giua_chung = True
                break

        stt_hien_tai = row.get("stt", idx)
        if stt_da_xong and stt_hien_tai in stt_da_xong:
            continue  # RESUME: dong nay da xu ly xong o lan chay TRUOC, khong xu ly lai
        stt = row.get("stt", "")
        ten = row.get("name") or row.get("ten") or ""
        progress_cb and progress_cb(idx + 1, len(rows), ten or f"dòng {stt}")
        maxa, source_maxa = resolve_maxa_for_row(row, thon_mapping, xa_mapping)
        to = row.get("numberto") or row.get("to")
        thua = row.get("numberthua") or row.get("thua")

        missing = []
        if not maxa:
            missing.append("Mã xã")
        if not to:
            missing.append("Số tờ")
        if not thua:
            missing.append("Số thửa")

        # BẢO VỆ: "Số tờ"/"Số thửa" phải là số (hoặc số+chữ ngắn như "70a") - nếu trông giống ĐỊA
        # CHỈ (có dấu phẩy hoặc quá dài) thì gần như chắc chắn NGƯỜI DÙNG ĐÃ CẤU HÌNH NHẦM CỘT
        # (VD nhập nhầm cột "Địa chỉ" vào ô "Thửa") - CHẶN LẠI ở đây thay vì lặng lẽ tạo ra tên file
        # sai be bét (đã từng xảy ra thật: Thửa hiện ra "Đồng Lũng Cậu, Thôn Kim Vân, xã Văn Lang,
        # tỉnh Thái Nguyên" thay vì số 165 - do cột "Thửa" bị cấu hình trỏ nhầm sang cột Địa chỉ).
        def _trong_giong_dia_chi(gia_tri):
            s = str(gia_tri)
            return "," in s or len(s) > 20

        canh_bao_cau_hinh_cot = []
        if not missing:
            if _trong_giong_dia_chi(to):
                canh_bao_cau_hinh_cot.append(f"Số tờ trông giống ĐỊA CHỈ ('{to}') - kiểm tra lại cấu hình cột 'Tờ'")
            if _trong_giong_dia_chi(thua):
                canh_bao_cau_hinh_cot.append(f"Số thửa trông giống ĐỊA CHỈ ('{thua}') - kiểm tra lại cấu hình cột 'Thửa'")

        if missing:
            results.append({"stt": stt, "ten": ten, "maxa": maxa or "", "to": to or "", "thua": thua or "",
                            "filename": "", "status": "CẦN KIỂM TRA", "note": "Thiếu: " + ", ".join(missing),
                            "hauto": "", "nguon_hauto": "", "mau_ten_file": ""})
            log_cb(f"✗ Dòng {stt}: thiếu {', '.join(missing)}")
            continue

        if canh_bao_cau_hinh_cot:
            results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                            "filename": "", "status": "CẦN KIỂM TRA",
                            "note": "⚠ Có thể SAI CẤU HÌNH CỘT: " + " | ".join(canh_bao_cau_hinh_cot),
                            "hauto": "", "nguon_hauto": "", "mau_ten_file": ""})
            log_cb(f"⚠ Dòng {stt}: " + " | ".join(canh_bao_cau_hinh_cot))
            continue

        if commune_rows:
            row = enrich_row_with_commune_config(row, commune_rows)
        replacements = build_mailmerge_replacements(row)
        templates_to_run = []
        if combine and export_gt and export_tbxn:
            templates_to_run = [("GOP", None)]
        else:
            if export_gt:
                templates_to_run.append(("GT", template_gt))
            if export_tbxn:
                templates_to_run.append(("TBXN", template_tbxn))

        if not templates_to_run:
            results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                            "filename": "", "status": "BỎ QUA", "note": "Chưa chọn mẫu nào để xuất",
                            "hauto": "", "nguon_hauto": "", "mau_ten_file": ""})
            continue

        for default_suffix, tpl in templates_to_run:
            if suffix_enabled:
                hauto, nguon_hauto = resolve_file_suffix(
                    default_suffix,
                    manual_suffix=suffix_config.get("manual", ""),
                    combobox_suffix=suffix_config.get("combobox", ""),
                    excel_value=row.get("hauto", ""),
                    use_excel_col=suffix_config.get("use_excel_col", False),
                    prioritize_manual=suffix_config.get("prioritize_manual", False),
                    uppercase=suffix_config.get("uppercase", False),
                    remove_diacritics_opt=suffix_config.get("remove_diacritics", False),
                )
                try:
                    fname_base = render_output_filename(
                        filename_template, maxa=maxa, to=to, thua=thua, ten=ten,
                        cccd=row.get("id") or "", thon=row.get("thon") or "", xacu=row.get("xacu") or "",
                        hauto=hauto, stt=stt)
                except ValueError as e:
                    results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                                    "filename": "", "status": "LỖI", "note": str(e),
                                    "hauto": hauto, "nguon_hauto": nguon_hauto, "mau_ten_file": filename_template})
                    log_cb(f"✗ Dòng {stt}: {e}")
                    continue
            else:
                # Hành vi CŨ - giữ nguyên 100% để không phá vỡ luồng đang dùng
                hauto, nguon_hauto = default_suffix, "MAC_DINH"
                fname_base = sanitize_filename_part(f"CHUACOGIAY_{maxa}_{to}_{thua}_{default_suffix}")

            out_path = os.path.join(output_folder, fname_base + ".docx") if output_folder else fname_base + ".docx"
            ten_file_du_kien = os.path.basename(out_path)
            base, ext = os.path.splitext(out_path)
            i = 2
            while not dry_run and os.path.isfile(out_path):
                out_path = f"{base}_{i}{ext}"
                i += 1
            fname = os.path.basename(out_path)

            if dry_run:
                results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                                "filename": fname, "status": "OK (chạy thử - chưa xuất)", "note": source_maxa or "",
                                "hauto": hauto, "nguon_hauto": nguon_hauto, "mau_ten_file": filename_template,
                                "ten_file_du_kien": ten_file_du_kien})
                log_cb(f"✓ (chạy thử) Dòng {stt} → {fname}  [{source_maxa}]")
                continue

            t0_export = time.time()
            try:
                if default_suffix == "GOP":
                    if not (template_gt and template_tbxn):
                        raise RuntimeError("Cần chọn cả 2 mẫu (Đơn đăng ký + Thông báo xác nhận) để gộp chung.")
                    tmp_gt = out_path + ".__gt_tmp.docx"
                    tmp_tbxn = out_path + ".__tbxn_tmp.docx"
                    m1 = mail_merge_smart(template_gt, replacements, tmp_gt, log_cb, word_session=word_session)
                    m2 = mail_merge_smart(template_tbxn, replacements, tmp_tbxn, log_cb, word_session=word_session)
                    master = DocxDocument(tmp_gt)
                    composer = DocxComposer(master)
                    composer.doc.add_page_break()
                    composer.append(DocxDocument(tmp_tbxn))
                    from app.services import file_safety_service as _fsafe
                    _fsafe.ghi_nguyen_tu(composer.save, out_path)
                    os.remove(tmp_gt)
                    os.remove(tmp_tbxn)
                    method = f"{m1}+{m2}"
                else:
                    method = mail_merge_smart(tpl, replacements, out_path, log_cb, word_session=word_session)

                results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                                "filename": fname, "status": "OK", "note": f"{source_maxa} ({method})",
                                "hauto": hauto, "nguon_hauto": nguon_hauto, "mau_ten_file": filename_template,
                                "ten_file_du_kien": ten_file_du_kien})
                log_cb(f"✓ Dòng {stt} → {fname}  ({method})")
                loi_xuat = ""
            except Exception as e:
                results.append({"stt": stt, "ten": ten, "maxa": maxa, "to": to, "thua": thua,
                                "filename": fname, "status": "LỖI", "note": str(e),
                                "hauto": hauto, "nguon_hauto": nguon_hauto, "mau_ten_file": filename_template,
                                "ten_file_du_kien": ten_file_du_kien})
                log_cb(f"✗ Dòng {stt}: lỗi xuất - {e}")
                loi_xuat = str(e)

            # TỐI ƯU HIỆU NĂNG (mục III/V/VI/VII/XVI tài liệu "KHÔNG LÀM TREO WINDOWS") - chỉ áp
            # dụng khi đang dùng Word COM thật (không áp dụng cho dry_run/mail_merge_docx_template
            # thuần Python vốn đã nhẹ, không cần nghỉ).
            if ap_dung_toi_uu:
                thoi_gian_xuat = time.time() - t0_export
                if perf_log_path:
                    kich_thuoc_kb = round(os.path.getsize(out_path) / 1024, 1) if os.path.isfile(out_path) else 0
                    stats_now = _wperf.sample_perf_stats() or {}
                    _wperf.write_perf_log_row(
                        perf_log_path, time.strftime("%Y-%m-%d %H:%M:%S"),
                        stats_now.get("cpu_percent", ""), stats_now.get("ram_mb", ""),
                        fname, round(thoi_gian_xuat, 2), kich_thuoc_kb, not loi_xuat, loi_xuat)
                if perf_stats_cb:
                    stats_now2 = _wperf.sample_perf_stats()
                    if stats_now2:
                        perf_stats_cb(stats_now2)
                _wperf.collect_garbage()
                time.sleep(sleep_per_file)
                if gioi_han_ram_mb and gioi_han_ram_mb > 0:
                    stats_ram = _wperf.sample_perf_stats()
                    if stats_ram and stats_ram.get("ram_mb", 0) > gioi_han_ram_mb:
                        log_cb(f"⚠ RAM hiện tại ({stats_ram['ram_mb']:.0f}MB) vượt giới hạn đã đặt "
                              f"({gioi_han_ram_mb}MB) - nghỉ thêm để hệ thống ổn định lại...")
                        time.sleep(sleep_per_file * 3)
                if so_file_moi_dot and (idx + 1) % so_file_moi_dot == 0 and (idx + 1) < len(rows):
                    log_cb(f"⏸ Đã xử lý {idx + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài "
                          f"nguyên cho Windows (Explorer/Excel/trình duyệt...) trước khi tiếp tục...")
                    time.sleep(nghi_giua_dot_giay)

            # CHECKPOINT (mục XVII tài liệu) - độc lập với throttle, chỉ lưu khi có checkpoint_task_id.
            # Lưu SAU MỖI FILE để nếu phần mềm bị tắt đột ngột, lần sau biết CHÍNH XÁC đã xong đến đâu.
            if checkpoint_task_id and not dry_run:
                from app.services import checkpoint_service as _ckpt
                da_xong_stt = [r["stt"] for r in results if r.get("status") == "OK"]
                loi_stt = [r["stt"] for r in results if r.get("status") == "LỖI"]
                can_ktra_stt = [r["stt"] for r in results if r.get("status") not in ("OK", "LỖI")]
                _ckpt.luu_checkpoint(
                    checkpoint_task_id, "xuat_word_hang_loat",
                    [r.get("stt", i) for i, r in enumerate(rows)],
                    da_xong_stt, loi_stt, can_ktra_stt, idx + 1,
                    {"toc_do_xu_ly": toc_do_xu_ly}, output_folder)

    if ap_dung_toi_uu:
        _wperf.restore_normal_priority()
        # Mục XVII tài liệu: nếu Word bị treo/lỗi giữa chừng, KHÔNG hủy toàn bộ - đã tự động
        # chuyển sang phương án dự phòng python-docx cho từng file gặp lỗi (xem mail_merge_smart).
        # Ở đây chỉ TỔNG KẾT lại cho người dùng biết rõ nếu việc này xảy ra NHIỀU LẦN (dấu hiệu
        # Word có thể đã treo/crash thật sự, không chỉ lỗi riêng lẻ 1-2 file) - tránh người dùng
        # không để ý các dòng log riêng lẻ mà không biết Word đã gặp sự cố.
        so_dung_du_phong = sum(1 for r in results if "python-docx" in (r.get("note") or ""))
        if so_dung_du_phong >= 3:
            log_cb(f"\n⚠ LƯU Ý: {so_dung_du_phong}/{len(results)} file phải dùng PHƯƠNG ÁN DỰ PHÒNG "
                  f"(python-docx) thay vì Word thật - có thể Word đã gặp sự cố/bị treo giữa chừng. "
                  f"Toàn bộ file VẪN ĐƯỢC TẠO ĐẦY ĐỦ (không mất dữ liệu), nhưng nếu mẫu Word có định "
                  f"dạng phức tạp (bảng biểu, ảnh...) nên kiểm tra lại vài file bằng mắt để chắc "
                  f"chắn định dạng đúng như mong muốn.")

    if checkpoint_task_id and not dry_run and not bi_huy_giua_chung:
        from app.services import checkpoint_service as _ckpt
        _ckpt.xoa_checkpoint(checkpoint_task_id)  # hoàn thành TOÀN BỘ - không cần resume nữa

    return results


def write_word_batch_log_csv(csv_path, results):
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên chủ", "Mã xã", "Số tờ", "Số thửa", "HauToFile", "NguonHauTo",
                    "MauTenFile", "TenFileDuKien", "TenFileDaXuat", "Trạng thái", "Ghi chú"])
        for r in results:
            w.writerow([r["stt"], r["ten"], r["maxa"], r["to"], r["thua"],
                        r.get("hauto", ""), r.get("nguon_hauto", ""), r.get("mau_ten_file", ""),
                        r.get("ten_file_du_kien", "") or r["filename"], r["filename"],
                        r["status"], r["note"]])


# ============================================================================
# MỤC 9 (MỚI): GỘP FILE WORD HÀNG LOẠT
# ============================================================================

MERGED_WORD_PREFIX = "GOP_WORD_"


def generate_random_docx_name(prefix=MERGED_WORD_PREFIX):
    rand_part = "".join(random.choices(string.ascii_uppercase + string.digits, k=10))
    return f"{prefix}{rand_part}.docx"


def batch_merge_word_folder(source, log_cb, add_page_break=True, control=None, progress_cb=None):
    """
    Gộp file .docx từ `source` (1 thư mục HOẶC danh sách file cụ thể - có thể ở nhiều
    thư mục khác nhau) thành 1 file .docx duy nhất, tên ngẫu nhiên.
    Lưu vào: chính thư mục nguồn (nếu source là thư mục), hoặc thư mục chứa FILE ĐẦU TIÊN
    trong danh sách (nếu source là danh sách file cụ thể).
    progress_cb(idx, total, filename): gọi trước khi gộp mỗi file - dùng cho thanh tiến độ.
    Trả về (danh_sach_file_da_gop, duong_dan_file_ket_qua).
    """
    if not HAS_DOCXCOMPOSE:
        raise RuntimeError("Thiếu thư viện docxcompose/python-docx. Cài: pip install docxcompose python-docx")

    if isinstance(source, (list, tuple, set)):
        candidates = sorted(os.path.abspath(f) for f in source)
        doc_files_paths = [f for f in candidates if f.lower().endswith(".docx")
                           and not os.path.basename(f).startswith(MERGED_WORD_PREFIX)]
        skipped_doc = [f for f in candidates if f.lower().endswith(".doc") and not f.lower().endswith(".docx")]
        save_dir = os.path.dirname(doc_files_paths[0]) if doc_files_paths else None
    else:
        all_entries = sorted(os.listdir(source))
        doc_files_paths = [os.path.join(source, f) for f in all_entries
                           if f.lower().endswith(".docx") and not f.startswith(MERGED_WORD_PREFIX)]
        skipped_doc = [f for f in all_entries if f.lower().endswith(".doc") and not f.lower().endswith(".docx")]
        save_dir = source

    if len(doc_files_paths) < 2:
        raise RuntimeError(
            f"Cần ít nhất 2 file .docx để gộp (hiện có {len(doc_files_paths)}).")

    if skipped_doc:
        log_cb(f"⚠ Bỏ qua {len(skipped_doc)} file .doc (định dạng Word cũ, chưa hỗ trợ): "
               f"{', '.join(os.path.basename(f) for f in skipped_doc)}")

    log_cb(f"Sẽ gộp {len(doc_files_paths)} file theo thứ tự tên (A→Z):")
    for f in doc_files_paths:
        log_cb(f"  - {os.path.basename(f)}")

    master = DocxDocument(doc_files_paths[0])
    composer = DocxComposer(master)
    progress_cb and progress_cb(1, len(doc_files_paths), os.path.basename(doc_files_paths[0]))

    for i, fpath in enumerate(doc_files_paths[1:], start=2):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb("⏹ Đã hủy theo yêu cầu giữa lúc gộp - lưu tạm phần đã gộp được đến lúc này.")
                break
        progress_cb and progress_cb(i, len(doc_files_paths), os.path.basename(fpath))
        if add_page_break:
            composer.doc.add_page_break()
        composer.append(DocxDocument(fpath))
        log_cb(f"✓ Đã gộp: {os.path.basename(fpath)}")

    out_name = generate_random_docx_name()
    out_path = os.path.join(save_dir, out_name)
    # Đảm bảo không trùng tên (dù xác suất trùng gần như bằng 0)
    while os.path.isfile(out_path):
        out_name = generate_random_docx_name()
        out_path = os.path.join(save_dir, out_name)

    from app.services import file_safety_service as _fsafe
    _fsafe.ghi_nguyen_tu(composer.save, out_path)
    log_cb(f"✓ Đã lưu file gộp: {out_path}")
    return [os.path.basename(f) for f in doc_files_paths], out_path


# ============================================================================
# MỤC 8 (MỚI): TỔNG HỢP BÁO CÁO TỪ CÁC FILE PDF RA EXCEL (THEO MẪU)
# ============================================================================

# Cấu trúc cột xuất ra khớp với mẫu Excel THẬT đang dùng (VD file "CHUA_CO_GIAY_xaVanLangcu"),
# để có thể dùng NGAY file báo cáo này làm "File Excel tổng báo cáo" cho các tính năng khác mà
# không cần đổi cấu hình cột (mặc định: A=STT, B=Mã xã, S=Mã định danh, H=Tên chủ, I=CCCD,
# V=Số tờ, W=Số thửa, X=Địa chỉ, Y=Diện tích, Z=Mục đích SD, AB=Nguồn gốc SD, AD=Thời hạn SD,
# AX=Tên file quét, dữ liệu từ dòng 5).
REPORT_COL_STT = "A"
REPORT_COL_MAXA = "B"
REPORT_COL_MADINHDANH = "S"
REPORT_COL_TEN = "H"
REPORT_COL_ID = "I"
REPORT_COL_TO = "V"
REPORT_COL_THUA = "W"
REPORT_COL_DIACHI = "X"
REPORT_COL_DT = "Y"
REPORT_COL_MUCDICH = "Z"
REPORT_COL_CORE = "AB"
REPORT_COL_YEAR = "AD"
REPORT_COL_FILES = "AX"
REPORT_HEADER_ROW = 4
REPORT_DATA_START_ROW = 5


def compile_report_from_pdfs(source, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, debug=False,
                              control=None, debug_text_dir=None, threshold_cao=90, threshold_thap=70,
                              progress_cb=None, toc_do_xu_ly="can_bang", so_file_moi_dot=100,
                              nghi_giua_dot_giay=2.0, perf_log_path=None,
                              checkpoint_task_id=None, files_da_xong=None):
    """
    Đọc nội dung PDF từ `source` (1 thư mục HOẶC danh sách file cụ thể), GỘP các file
    thuộc cùng 1 thửa đất (cùng Mã xã + Tờ + Thửa) thành 1 dòng báo cáo.
    progress_cb(idx, total, filename): gọi trước khi xử lý mỗi file - dùng cho thanh tiến độ.
    Trả về (rows, error_files) với rows là list dict đã gộp theo thửa.

    TỐI ƯU HIỆU NĂNG (đúng tài liệu RÀ SOÁT/TỐI ƯU RELEASE CANDIDATE, mục V/VI/VII): nghỉ ngắn
    sau mỗi file (chỉ khi có OCR thật xảy ra - không nghỉ nếu file đã có sẵn lớp chữ, tránh làm
    chậm không cần thiết), nghỉ dài hơn sau mỗi đợt, hạ độ ưu tiên tiến trình khi xử lý từ 2 file
    OCR trở lên.

    CHECKPOINT/RESUME (mục XVII) - CHỈ áp dụng khi checkpoint_task_id KHÔNG None: lưu tiến độ định
    kỳ ra đĩa, TỰ XÓA khi hoàn thành TOÀN BỘ (KHÔNG xóa nếu bị hủy giữa chừng - CẢNH BÁO: đây từng
    là 1 lỗi thật đã gặp ở mục 7d, phải luôn dùng cờ theo dõi việc bị hủy). files_da_xong: set
    đường dẫn file đã xử lý xong ở lần chạy TRƯỚC (nếu đang resume) - các file này sẽ được BỎ QUA.
    """
    files = list_files_from_source(source, ".pdf")
    if not files:
        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")

    from app.services import word_perf_service as _wperf
    ap_dung_toi_uu = len(files) >= 2
    if ap_dung_toi_uu:
        _wperf.apply_low_priority()
        sleep_per_file = _wperf.get_sleep_per_file(toc_do_xu_ly)
    else:
        sleep_per_file = 0.0

    groups = {}
    order = []
    error_files = []

    FIELDS_TO_MERGE = ("ten", "dia_chi", "dt", "id", "mucdich", "core", "year")

    bi_huy_giua_chung = False
    for idx, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - idx} file chưa xử lý).")
                bi_huy_giua_chung = True
                break
        if files_da_xong and fpath in files_da_xong:
            continue  # RESUME: file nay da xu ly xong o lan chay TRUOC, khong xu ly lai
        fname = os.path.basename(fpath)
        progress_cb and progress_cb(idx + 1, len(files), fname)
        log_cb(f"→ Đang đọc: {fname}")
        t0_file = time.time()
        info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, debug=debug, debug_text_dir=debug_text_dir, threshold_cao=threshold_cao, threshold_thap=threshold_thap, ocr_cache_dir=get_ocr_cache_dir())

        if ap_dung_toi_uu:
            _wperf.collect_garbage()
            time.sleep(sleep_per_file)
            if perf_log_path:
                kich_thuoc_kb = round(os.path.getsize(fpath) / 1024, 1) if os.path.isfile(fpath) else 0
                stats_now = _wperf.sample_perf_stats() or {}
                _wperf.write_perf_log_row(
                    perf_log_path, time.strftime("%Y-%m-%d %H:%M:%S"),
                    stats_now.get("cpu_percent", ""), stats_now.get("ram_mb", ""),
                    fname, round(time.time() - t0_file, 2), kich_thuoc_kb,
                    "error" not in info, info.get("error", ""))
            if so_file_moi_dot and (idx + 1) % so_file_moi_dot == 0 and (idx + 1) < len(files):
                log_cb(f"⏸ Đã xử lý {idx + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                time.sleep(nghi_giua_dot_giay)

        if checkpoint_task_id:
            from app.services import checkpoint_service as _ckpt
            da_xu_ly_toi_idx = set(files[:idx + 1]) | (files_da_xong or set())
            _ckpt.luu_checkpoint(
                checkpoint_task_id, "tong_hop_bao_cao_pdf", files,
                list(da_xu_ly_toi_idx), [], [], idx + 1,
                {"toc_do_xu_ly": toc_do_xu_ly}, debug_text_dir or "")

        if "error" in info:
            error_files.append((fname, info["error"]))
            log_cb(f"   ✗ {info['error']}")
            if debug and info.get("raw_text"):
                log_cb("   ---- Văn bản đã đọc được ----")
                log_cb("   " + info["raw_text"].replace("\n", "\n   "))
                log_cb("   ---- Hết văn bản ----")
            continue

        if info.get("muc_do_tin_cay") == "KHONG_TU_XU_LY":
            reason = (f"Độ tin cậy OCR quá thấp ({info.get('diem_tin_cay_ocr')}%) để tự xử lý - "
                     f"{info.get('ly_do_can_kiem_tra', '')}")
            error_files.append((fname, reason))
            log_cb(f"   ⚠ KHONG_TU_XU_LY ({info.get('diem_tin_cay_ocr')}%): {reason}")
            continue

        key = (info["maxa"], info["to"], info["thua"])
        if key not in groups:
            groups[key] = {"maxa": info["maxa"], "to": info["to"], "thua": info["thua"],
                           **{f: info.get(f) for f in FIELDS_TO_MERGE}, "files": [],
                           "diem_tin_cay_ocr": info.get("diem_tin_cay_ocr"),
                           "muc_do_tin_cay": info.get("muc_do_tin_cay"),
                           "ly_do_can_kiem_tra": info.get("ly_do_can_kiem_tra")}
            order.append(key)
        g = groups[key]
        for f in FIELDS_TO_MERGE:
            if not g.get(f) and info.get(f):
                g[f] = info[f]
        g["files"].append(fname)

        nguon_txt = "OCR (bản scan)" if info.get("nguon") == "ocr" else "chữ có sẵn"
        via_txt = f"Thôn {info.get('thon_name')}" if info.get("nguon_maxa") == "thon" else f"Xã {info.get('xa_name')}"
        log_cb(f"   ✓ {via_txt} (Mã {info['maxa']}) - Tờ {info['to']} - Thửa {info['thua']} "
               f"- đọc bằng {nguon_txt} - Độ tin cậy: {info.get('diem_tin_cay_ocr')}% ({info.get('muc_do_tin_cay')})")

    rows = []
    for i, key in enumerate(order, start=1):
        g = groups[key]
        g["stt"] = i
        rows.append(g)

    if ap_dung_toi_uu:
        _wperf.restore_normal_priority()

    if checkpoint_task_id and not bi_huy_giua_chung:
        from app.services import checkpoint_service as _ckpt
        _ckpt.xoa_checkpoint(checkpoint_task_id)

    return rows, error_files


def export_report_from_pdfs(report_path, rows, error_files):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tổng hợp từ PDF"

    ws["A1"] = "BÁO CÁO TỔNG HỢP TỪ CÁC FILE PDF"
    headers = [
        (REPORT_COL_STT, "STT"), (REPORT_COL_MAXA, "Mã xã"), (REPORT_COL_MADINHDANH, "Mã định danh thửa đất"),
        (REPORT_COL_TEN, "Tên chủ SDĐ"), (REPORT_COL_ID, "Số CCCD"),
        (REPORT_COL_TO, "Số tờ"), (REPORT_COL_THUA, "Số thửa"), (REPORT_COL_DIACHI, "Địa chỉ"),
        (REPORT_COL_DT, "Diện tích"), (REPORT_COL_MUCDICH, "Mục đích SD"),
        (REPORT_COL_CORE, "Nguồn gốc SD"), (REPORT_COL_YEAR, "Thời hạn SD"),
        (REPORT_COL_FILES, "Tên file quét"),
    ]
    for col, label in headers:
        ws[f"{col}{REPORT_HEADER_ROW}"] = label
    # Cột chấm điểm tin cậy OCR - đặt ngay sau cột cuối cùng đã dùng (AY)
    col_diem, col_mucdo = "AZ", "BA"
    ws[f"{col_diem}{REPORT_HEADER_ROW}"] = "Điểm tin cậy OCR (%)"
    ws[f"{col_mucdo}{REPORT_HEADER_ROW}"] = "Mức độ tin cậy"

    # Tô đậm dòng tiêu đề (các cột không liền mạch B,H,I,...,AZ,BA)
    from openpyxl.styles import Font as _Font, PatternFill as _PatternFill, Alignment as _Alignment
    _header_font = _Font(bold=True, color="FFFFFF")
    _header_fill = _PatternFill("solid", fgColor="1F4E78")
    _center = _Alignment(horizontal="center", vertical="center")
    for col, _ in headers + [(col_diem, ""), (col_mucdo, "")]:
        cell = ws[f"{col}{REPORT_HEADER_ROW}"]
        cell.font = _header_font
        cell.fill = _header_fill
        cell.alignment = _center
    ws["A1"].font = _Font(bold=True, size=13)

    for idx, g in enumerate(rows):
        r = REPORT_DATA_START_ROW + idx
        ma_dinh_danh = f"CHUACOGIAY_{g['maxa']}_{g['to']}_{g['thua']}"
        ws[f"{REPORT_COL_STT}{r}"] = g["stt"]
        ws[f"{REPORT_COL_MAXA}{r}"] = g["maxa"]
        ws[f"{REPORT_COL_MADINHDANH}{r}"] = ma_dinh_danh
        ws[f"{REPORT_COL_TEN}{r}"] = g.get("ten") or ""
        ws[f"{REPORT_COL_ID}{r}"] = g.get("id") or ""
        ws[f"{REPORT_COL_TO}{r}"] = g["to"]
        ws[f"{REPORT_COL_THUA}{r}"] = g["thua"]
        ws[f"{REPORT_COL_DIACHI}{r}"] = g.get("dia_chi") or ""
        ws[f"{REPORT_COL_DT}{r}"] = g.get("dt") or ""
        ws[f"{REPORT_COL_MUCDICH}{r}"] = g.get("mucdich") or ""
        ws[f"{REPORT_COL_CORE}{r}"] = g.get("core") or ""
        ws[f"{REPORT_COL_YEAR}{r}"] = g.get("year") or ""
        ws[f"{REPORT_COL_FILES}{r}"] = ",".join(g["files"])
        ws[f"{col_diem}{r}"] = g.get("diem_tin_cay_ocr") or ""
        ws[f"{col_mucdo}{r}"] = g.get("muc_do_tin_cay") or ""

    ws2 = wb.create_sheet("File lỗi - cần kiểm tra tay")
    ws2.append(["Tên file PDF", "Lý do không đọc được"])
    for fname, reason in error_files:
        ws2.append([fname, reason])
    style_excel_header_row(ws2, row_num=1, auto_width=False)

    for ws_ in (ws, ws2):
        for col in ws_.columns:
            length = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            ws_.column_dimensions[col[0].column_letter].width = min(max(length + 2, 10), 60)

    wb.save(report_path)


# ============================================================================
# MỤC 8b (MỚI): THỐNG KÊ TIẾN ĐỘ TỪ FILE EXCEL BẤT KỲ
# ============================================================================

def analyze_excel_progress(excel_path, sheet_name, header_row, col_count,
                            col_status, done_values,
                            col_gcn, gcn_done_values,
                            target_total, log_cb=None, progress_cb=None):
    """
    Thống kê tiến độ từ 1 file Excel BẤT KỲ, cột do người dùng tự chọn:
    - col_count: cột dùng để xác định 1 dòng có dữ liệu hay không (VD cột STT/Số thửa).
      Để trống thì tính tất cả các dòng có bất kỳ ô nào trong col_status/col_gcn.
    - col_status + done_values: cột trạng thái "đã/chưa thực hiện" và các giá trị coi là ĐÃ thực hiện.
    - col_gcn + gcn_done_values: cột trạng thái GCN và các giá trị coi là ĐÃ nhập.
    - target_total: tổng khối lượng được giao (mẫu số tính %). Để trống/0 thì dùng luôn tổng số dòng đếm được.
    progress_cb(idx, total, mo_ta): gọi định kỳ trong lúc quét dòng - dùng cho thanh tiến độ.
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    done_norm = {strip_diacritics(v).upper().strip() for v in done_values if v.strip()}
    gcn_done_norm = {strip_diacritics(v).upper().strip() for v in gcn_done_values if v.strip()}

    total_rows = 0
    done_count = 0
    status_blank_count = 0
    gcn_done_count = 0
    gcn_notdone_count = 0
    gcn_blank_count = 0

    tong_so_dong_quet = ws.max_row - header_row + 1
    for r in range(header_row, ws.max_row + 1):
        if progress_cb and (r - header_row) % 200 == 0:
            progress_cb(r - header_row + 1, tong_so_dong_quet, f"dòng {r}")
        if col_count:
            count_val = ws[f"{col_count}{r}"].value
            if count_val is None or str(count_val).strip() == "":
                continue
        elif col_status or col_gcn:
            # Không có cột đếm riêng -> coi dòng có dữ liệu nếu 1 trong 2 cột status/gcn có giá trị
            v1 = ws[f"{col_status}{r}"].value if col_status else None
            v2 = ws[f"{col_gcn}{r}"].value if col_gcn else None
            if (v1 is None or str(v1).strip() == "") and (v2 is None or str(v2).strip() == ""):
                continue
        else:
            continue
        total_rows += 1

        if col_status:
            status_val = ws[f"{col_status}{r}"].value
            if status_val is None or str(status_val).strip() == "":
                status_blank_count += 1
            else:
                norm = strip_diacritics(str(status_val)).upper().strip()
                if norm in done_norm:
                    done_count += 1

        if col_gcn:
            gcn_val = ws[f"{col_gcn}{r}"].value
            if gcn_val is None or str(gcn_val).strip() == "":
                gcn_blank_count += 1
            else:
                normg = strip_diacritics(str(gcn_val)).upper().strip()
                if normg in gcn_done_norm:
                    gcn_done_count += 1
                else:
                    gcn_notdone_count += 1

    denom = target_total if target_total and target_total > 0 else total_rows

    result = {
        "total_rows": total_rows,
        "target_total": denom,
        "col_status_used": bool(col_status),
        "done_count": done_count,
        "not_done_count": (total_rows - done_count) if col_status else None,
        "status_blank_count": status_blank_count if col_status else None,
        "percent_done": round(done_count / denom * 100, 2) if (col_status and denom) else None,
        "col_gcn_used": bool(col_gcn),
        "gcn_done_count": gcn_done_count if col_gcn else None,
        "gcn_notdone_count": gcn_notdone_count if col_gcn else None,
        "gcn_blank_count": gcn_blank_count if col_gcn else None,
        "percent_gcn_done": round(gcn_done_count / denom * 100, 2) if (col_gcn and denom) else None,
    }
    return result


def export_progress_report(report_path, result, meta):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Thống kê tiến độ"

    rows_out = [
        ("Nguồn dữ liệu (file Excel)", meta.get("excel_path", "")),
        ("Tổng số dòng có dữ liệu đếm được", result["total_rows"]),
        ("Tổng khối lượng được giao (mẫu số tính %)", result["target_total"]),
        ("", ""),
    ]
    if result["col_status_used"]:
        rows_out += [
            ("--- Tiến độ thực hiện ---", ""),
            ("Số đã thực hiện", result["done_count"]),
            ("Số chưa thực hiện", result["not_done_count"]),
            ("Số dòng bỏ trống trạng thái", result["status_blank_count"]),
            ("Tỷ lệ đã thực hiện (%)", result["percent_done"]),
            ("", ""),
        ]
    if result["col_gcn_used"]:
        rows_out += [
            ("--- Giấy chứng nhận (GCN) ---", ""),
            ("Số đã nhập GCN", result["gcn_done_count"]),
            ("Số chưa nhập GCN", result["gcn_notdone_count"]),
            ("Số dòng bỏ trống GCN", result["gcn_blank_count"]),
            ("Tỷ lệ đã nhập GCN (%)", result["percent_gcn_done"]),
        ]

    for label, val in rows_out:
        ws.append([label, val])

    # Tô đậm cột nhãn (A) và các dòng phân nhóm "--- ... ---"
    from openpyxl.styles import Font as _Font, PatternFill as _PatternFill
    _bold = _Font(bold=True)
    _section_fill = _PatternFill("solid", fgColor="D9E1F2")
    for r in range(1, ws.max_row + 1):
        cell_a = ws.cell(row=r, column=1)
        cell_a.font = _bold
        label_text = str(cell_a.value or "")
        if label_text.startswith("---"):
            cell_a.fill = _section_fill
            ws.cell(row=r, column=2).fill = _section_fill

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 30
    wb.save(report_path)


# ============================================================================
# MỤC 8c (MỚI): ĐIỀN DỮ LIỆU TỪ PDF VÀO FILE EXCEL MẪU CÓ SẴN
# (giữ nguyên định dạng, tiêu đề, merge cell, công thức, bố cục của file mẫu)
# ============================================================================

WB8C_COLS = ["B", "H", "I", "J", "K", "L", "V", "W", "X", "Y", "Z", "AA"]


def fill_excel_template_from_pdfs(template_excel_path, source_pdfs, sheet_name, header_row,
                                   xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi,
                                   clear_old_data, infer_from_cccd, output_path, dry_run, log_cb, control=None,
                                   col_map=None, no_overwrite=False, warn_duplicate=True, debug_text_dir=None,
                                   threshold_cao=90, threshold_thap=70, progress_cb=None,
                                   toc_do_xu_ly="can_bang", so_file_moi_dot=100, nghi_giua_dot_giay=2.0,
                                   perf_log_path=None, checkpoint_task_id=None, files_da_xong=None):
    """
    Đọc từng PDF (kể cả PDF scan/ảnh qua OCR), trích thông tin, điền vào file Excel MẪU có sẵn
    (giữ nguyên định dạng/merge cell/công thức) tại các cột theo col_map (mặc định: B, H, I, J, K,
    L, V, W, X, Y, Z, AA).
    col_map: dict {ten_truong: chu_cot} - cho phép người dùng tùy chọn cột, không cố định cứng.
    Các khóa: maxa, ten, cccd, ngaysinh, gioitinh, diachitt, soto, sothua, diachithua, dientich,
    mucdich, dientichloai1. Bỏ trống 1 cột -> không ghi trường đó.
    no_overwrite: nếu True, chỉ ghi vào ô đang TRỐNG (không ghi đè ô đã có dữ liệu).
    warn_duplicate: nếu True, cảnh báo khi trùng khóa Mã xã+Tờ+Thửa.
    Nếu dry_run=True: KHÔNG ghi/lưu file Excel, chỉ trả về bảng log để xem trước.
    progress_cb(idx, total, filename): gọi trước khi xử lý mỗi file - dùng cho thanh tiến độ.
    Trả về list dict log: {file, maxa, ten, cccd, to, thua, dt, mucdich, status, note}.
    """
    # col_map mặc định (khớp sheet "Chưa GCN" - đã xác minh với file mẫu thực tế)
    if col_map is None:
        col_map = {"maxa": "B", "ten": "H", "cccd": "I", "ngaysinh": "J", "gioitinh": "K",
                   "diachitt": "L", "soto": "V", "sothua": "W", "diachithua": "X",
                   "dientich": "Y", "mucdich": "Z", "dientichloai1": "AA"}
    # Danh sách cột thực sự dùng để xóa dữ liệu cũ (bỏ cột trống)
    used_cols = [c.strip().upper() for c in col_map.values() if c and c.strip()]
    wb = None
    ws = None
    maxa_col = col_map.get("maxa", "B").strip().upper() or "B"
    if not dry_run:
        wb = openpyxl.load_workbook(template_excel_path)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

        if clear_old_data:
            for r in range(header_row, ws.max_row + 1):
                for c in used_cols:
                    ws[f"{c}{r}"] = None

    write_row = header_row
    if not dry_run and not clear_old_data:
        r = header_row
        while ws[f"{maxa_col}{r}"].value not in (None, ""):
            r += 1
        write_row = r

    files = list_files_from_source(source_pdfs, ".pdf")
    log_rows = []
    seen_keys = {}  # theo dõi khóa Mã xã+Tờ+Thửa để cảnh báo trùng
    stt = 0

    from app.services import word_perf_service as _wperf
    ap_dung_toi_uu = len(files) >= 2 and not dry_run
    if ap_dung_toi_uu:
        _wperf.apply_low_priority()
        sleep_per_file = _wperf.get_sleep_per_file(toc_do_xu_ly)
    else:
        sleep_per_file = 0.0

    def _log_row(fpath, info, status, note, dong_excel=""):
        nonlocal stt
        stt += 1
        info = info or {}
        return {
            "stt": stt, "file": os.path.basename(fpath), "path": fpath,
            "maxa": info.get("maxa") or "",
            "xa_cu": info.get("xa_name") or info.get("xa_name_1c") or "",
            "thon": info.get("thon_name") or info.get("thon_name_1c") or "",
            "ten": info.get("ten") or "", "cccd": info.get("id") or "",
            "ngaysinh": info.get("_ngaysinh_log") or "",
            "gioitinh": info.get("_gioitinh_log") or "",
            "diachitt": info.get("dia_chi_thuong_tru") or "",
            "so_to_goc": info.get("so_to_goc") or info.get("to") or "",
            "to": info.get("to") or "", "thua": info.get("thua") or "",
            "diachithua": info.get("dia_chi") or "",
            "dt": info.get("dt") if info.get("dt") is not None else "",
            "mucdich": info.get("mucdich") or "",
            "loai_dat": info.get("loai_dat") or "",
            "co_chuyen_so_to_lam_nghiep": "CÓ" if info.get("forest_map_status") == "DA_CHUYEN_SO_TO_LAM_NGHIEP" else "KHÔNG",
            "nguon_maxa": info.get("nguon_maxa") or "",
            "nguon_so_to": info.get("nguon") or "", "nguon_so_thua": info.get("nguon") or "",
            "nguon_loai_dat": info.get("nguon") or "",
            "diem_tin_cay": info.get("diem_tin_cay_ocr") if info.get("diem_tin_cay_ocr") is not None else "",
            "dong_excel": dong_excel, "status": status, "note": note,
        }

    bi_huy_giua_chung = False
    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                bi_huy_giua_chung = True
                break
        if files_da_xong and fpath in files_da_xong:
            continue  # RESUME: file nay da xu ly xong o lan chay TRUOC, khong xu ly lai
        fname = os.path.basename(fpath)
        progress_cb and progress_cb(i + 1, len(files), fname)
        log_cb(f"→ Đang đọc: {fname}")
        t0_file = time.time()
        info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, debug_text_dir=debug_text_dir, threshold_cao=threshold_cao, threshold_thap=threshold_thap, ocr_cache_dir=get_ocr_cache_dir())

        if ap_dung_toi_uu:
            _wperf.collect_garbage()
            time.sleep(sleep_per_file)
            if perf_log_path:
                kich_thuoc_kb = round(os.path.getsize(fpath) / 1024, 1) if os.path.isfile(fpath) else 0
                stats_now = _wperf.sample_perf_stats() or {}
                _wperf.write_perf_log_row(
                    perf_log_path, time.strftime("%Y-%m-%d %H:%M:%S"),
                    stats_now.get("cpu_percent", ""), stats_now.get("ram_mb", ""),
                    fname, round(time.time() - t0_file, 2), kich_thuoc_kb,
                    "error" not in info, info.get("error", ""))
            if so_file_moi_dot and (i + 1) % so_file_moi_dot == 0 and (i + 1) < len(files):
                log_cb(f"⏸ Đã xử lý {i + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                time.sleep(nghi_giua_dot_giay)

        if checkpoint_task_id and not dry_run:
            from app.services import checkpoint_service as _ckpt
            da_xu_ly_toi_idx = set(files[:i + 1]) | (files_da_xong or set())
            _ckpt.luu_checkpoint(
                checkpoint_task_id, "dien_excel_mau_tu_pdf", files,
                list(da_xu_ly_toi_idx), [], [], i + 1,
                {"toc_do_xu_ly": toc_do_xu_ly}, output_path or "")

        if "error" in info:
            log_rows.append(_log_row(fpath, None, "LOI_DOC_PDF", info["error"]))
            log_cb(f"   ✗ {info['error']}")
            continue

        if info.get("muc_do_tin_cay") == "KHONG_TU_XU_LY":
            log_rows.append(_log_row(fpath, info, "OCR_YEU",
                                     f"Độ tin cậy OCR quá thấp ({info.get('diem_tin_cay_ocr')}%)."))
            log_cb(f"   ✗ OCR yếu ({info.get('diem_tin_cay_ocr')}%)")
            continue

        if not info.get("maxa"):
            log_rows.append(_log_row(fpath, info, "THIEU_MA_XA", "Không nhận diện được Mã xã."))
            log_cb(f"   ✗ Thiếu Mã xã")
            continue
        if not info.get("to"):
            log_rows.append(_log_row(fpath, info, "THIEU_SO_TO", "Không nhận diện được Số tờ."))
            log_cb(f"   ✗ Thiếu Số tờ")
            continue
        if not info.get("thua"):
            log_rows.append(_log_row(fpath, info, "THIEU_SO_THUA", "Không nhận diện được Số thửa."))
            log_cb(f"   ✗ Thiếu Số thửa")
            continue

        missing_other = []
        if not info.get("ten"):
            missing_other.append("Họ tên")
        if not info.get("dia_chi"):
            missing_other.append("Địa chỉ thửa đất")
        if info.get("dt") is None:
            missing_other.append("Diện tích")
        if not info.get("mucdich"):
            missing_other.append("Mục đích sử dụng")
        if missing_other:
            log_rows.append(_log_row(fpath, info, "CAN_KIEM_TRA", "Thiếu: " + ", ".join(missing_other)))
            log_cb(f"   ✗ Thiếu: {', '.join(missing_other)}")
            continue

        # Kiểm tra chéo: Mã xã/Số tờ/Số thửa theo TÊN FILE so với dữ liệu đọc được từ NỘI DUNG PDF
        # (mục VII.5 tài liệu "LÀM SÂU CHỨC NĂNG TỔNG HỢP PDF RA EXCEL") - nếu mâu thuẫn THẬT (không
        # giải thích được bằng quy tắc chuẩn hóa số tờ lâm nghiệp), không ghi tự động, đưa vào CẦN
        # KIỂM TRA để người dùng tự xác nhận.
        cross_check = cross_check_filename_vs_content(
            fname, info["maxa"], info["to"], info["thua"], loai_dat=info.get("loai_dat"))
        if cross_check["co_mau_thuan"]:
            note = "Dữ liệu mâu thuẫn tên file/nội dung PDF: " + "; ".join(cross_check["chi_tiet"])
            log_rows.append(_log_row(fpath, info, "DU_LIEU_MAU_THUAN", note))
            log_cb(f"   ⚠ {note}")
            continue

        nam_sinh, gioi_tinh = None, None
        if infer_from_cccd and info.get("id"):
            nam_sinh, gioi_tinh = infer_dob_gender_from_cccd(info["id"])
        info["_ngaysinh_log"] = nam_sinh or ""
        info["_gioitinh_log"] = gioi_tinh or ""

        # Cảnh báo trùng khóa Mã xã + Tờ + Thửa
        key = f"{info['maxa']}_{info['to']}_{info['thua']}"
        if warn_duplicate and key in seen_keys:
            log_rows.append(_log_row(fpath, info, "TRUNG_KHOA", f"Trùng khóa với {seen_keys[key]}"))
            log_cb(f"   ⚠ TRÙNG KHÓA với {seen_keys[key]}")
            continue
        seen_keys.setdefault(key, fname)

        dong_excel_ghi = ""
        if not dry_run:
            r = write_row
            # Ánh xạ trường -> giá trị
            values = {
                "maxa": info["maxa"], "ten": info.get("ten") or "", "cccd": info.get("id") or "",
                "ngaysinh": nam_sinh or "", "gioitinh": gioi_tinh or "",
                "diachitt": info.get("dia_chi_thuong_tru") or "", "soto": info["to"],
                "sothua": info["thua"], "diachithua": info.get("dia_chi") or "",
                "dientich": info.get("dt"), "mucdich": info.get("mucdich") or "",
                "dientichloai1": info.get("dt"),
            }
            try:
                for field, col in col_map.items():
                    col = (col or "").strip().upper()
                    if not col or field not in values:
                        continue
                    cell = ws[f"{col}{r}"]
                    if no_overwrite and cell.value not in (None, ""):
                        continue  # không ghi đè ô đã có dữ liệu
                    cell.value = values[field]
                write_row += 1
                dong_excel_ghi = r
            except Exception as e:
                log_rows.append(_log_row(fpath, info, "LOI_GHI_EXCEL", f"Lỗi ghi Excel: {e}"))
                log_cb(f"   ✗ Lỗi ghi Excel: {e}")
                continue

        status = "CHAY_THU" if dry_run else "DA_GHI_EXCEL"
        log_rows.append(_log_row(fpath, info, status, "", dong_excel=dong_excel_ghi))
        log_cb(f"   ✓ {'(chạy thử) ' if dry_run else ''}{info.get('ten')} - Mã {info['maxa']} - "
               f"Tờ {info['to']} - Thửa {info['thua']}")

    if not dry_run:
        from app.services import file_safety_service as _fsafe
        _fsafe.ghi_nguyen_tu(wb.save, output_path)

    if ap_dung_toi_uu:
        _wperf.restore_normal_priority()

    if checkpoint_task_id and not dry_run and not bi_huy_giua_chung:
        from app.services import checkpoint_service as _ckpt
        _ckpt.xoa_checkpoint(checkpoint_task_id)

    return log_rows


_read_excel_tong_rows_cache = {}


def read_excel_tong_rows(excel_path, sheet_name, header_row, col_map_tong, dientich_dung_sai=1.0,
                          use_memory_cache=False):
    """
    Đọc file Excel TỔNG đã chỉnh sửa theo cấu hình cột - đúng mục VII tài liệu "ĐÁNH LẠI THỨ TỰ
    TAB VÀ HOÀN THIỆN TAB 8 MỤC 8D": tạo NHIỀU chỉ mục (dictionary) để đối chiếu nhanh và có dự
    phòng, KHÔNG duyệt lồng nhau toàn bộ dữ liệu.

    use_memory_cache: nếu True, GHI NHỚ TRONG BỘ NHỚ (chỉ trong phiên làm việc hiện tại, KHÔNG ghi
    ra đĩa như cache OCR) kết quả đọc theo file+kích thước+thời gian sửa+cấu hình cột - tránh đọc
    lại Excel tổng (~4s cho 6000 dòng thực tế) khi người dùng bấm "Chạy thử" rồi "Tổng hợp báo
    cáo" liền sau (2 lần đọc CÙNG dữ liệu trong CÙNG 1 phiên chạy phần mềm). Mặc định False - AN
    TOÀN TUYỆT ĐỐI, giữ nguyên hành vi cũ (luôn đọc lại từ đĩa) nếu không bật.

    dientich_dung_sai: sai số diện tích (m²) coi là "gần bằng" khi so khớp theo chỉ mục diện tích.

    Trả về (indexes, rows_list):
    - indexes: dict gồm 3 chỉ mục, mỗi chỉ mục là {khóa: [row_data, ...]} (LIST vì có thể trùng
      khóa - KHÔNG tự lấy dòng đầu, phải phát hiện trùng khóa đúng mục VII):
      - "by_key": Mã xã+Số tờ+Số thửa (chỉ mục CHÍNH, ưu tiên cao nhất)
      - "by_parcel": Số tờ+Số thửa (không có Mã xã - dự phòng khi Mã xã tên file sai/thiếu)
      - "by_parcel_area": Số tờ+Số thửa+Diện tích làm tròn theo dung_sai (dự phòng thêm 1 lớp)
    - rows_list: toàn bộ dòng đọc được (kể cả dòng thiếu khóa) - dùng cho log/preview.
    Số tờ/Số thửa được CHUẨN HÓA (kể cả áp dụng quy tắc lâm nghiệp) trước khi tạo khóa.
    """
    cache_key = None
    if use_memory_cache:
        try:
            stat = os.stat(excel_path)
            cache_key = (excel_path, stat.st_size, int(stat.st_mtime), sheet_name, header_row,
                        tuple(sorted(col_map_tong.items())), dientich_dung_sai)
            if cache_key in _read_excel_tong_rows_cache:
                return _read_excel_tong_rows_cache[cache_key]
        except OSError:
            cache_key = None

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    if sheet_name not in wb.sheetnames:
        raise ValueError(f"Không tìm thấy sheet '{sheet_name}' trong file Excel tổng.")
    ws = wb[sheet_name]

    by_key, by_parcel, by_parcel_area = {}, {}, {}
    rows_list = []
    for r in range(header_row + 1, ws.max_row + 1):
        def g(field):
            col = (col_map_tong.get(field) or "").strip().upper()
            if not col:
                return None
            v = ws[f"{col}{r}"].value
            return None if v is None else v

        maxa_raw = g("maxa")
        soto_raw = g("soto")
        sothua_raw = g("sothua")
        if maxa_raw is None and soto_raw is None and sothua_raw is None:
            continue  # dòng trống hoàn toàn ở các cột khóa - bỏ qua, không tính là dữ liệu

        maxa, _ = normalize_ma_xa(maxa_raw) if maxa_raw is not None else (None, None)
        loai_dat_tong = g("loai_dat")
        soto = norm_num(soto_raw) if soto_raw is not None else None
        if soto is not None and loai_dat_tong:
            soto = normalize_forest_map_sheet_number(soto, loai_dat_tong)["so_to_chuan"]
        sothua = norm_num(sothua_raw) if sothua_raw is not None else None
        dientich_raw = g("dientich")
        dientich_chuan = _normalize_area_for_compare(dientich_raw) if dientich_raw is not None else None

        row_data = {
            "row_index": r, "maxa": maxa, "soto": soto, "sothua": sothua,
            "ten": g("ten"), "cccd": g("cccd"), "ngaysinh": g("ngaysinh"), "gioitinh": g("gioitinh"),
            "diachitt": g("diachitt"), "diachithua": g("diachithua"), "dientich": dientich_raw,
            "dientich_chuan": dientich_chuan,
            "mucdich": g("mucdich"), "loai_dat": loai_dat_tong, "ghichu": g("ghichu"),
        }
        rows_list.append(row_data)

        if soto and sothua:
            by_parcel.setdefault(f"{soto}_{sothua}", []).append(row_data)
            if maxa:
                by_key.setdefault(f"{maxa}_{soto}_{sothua}", []).append(row_data)
            if dientich_chuan is not None:
                dt_lam_tron = round(dientich_chuan / dientich_dung_sai) * dientich_dung_sai
                by_parcel_area.setdefault(f"{soto}_{sothua}_{dt_lam_tron}", []).append(row_data)

    result = {"by_key": by_key, "by_parcel": by_parcel, "by_parcel_area": by_parcel_area}, rows_list
    if cache_key is not None:
        _read_excel_tong_rows_cache[cache_key] = result
    return result


def reconcile_excel_tong_with_pdfs(excel_tong_path, sheet_tong, header_row_tong, col_map_tong,
                                    source_pdfs, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi,
                                    log_cb, control=None, doc_so_sanh=True, dientich_dung_sai=1.0,
                                    progress_cb=None):
    """
    Đối chiếu Excel TỔNG đã chỉnh sửa với PDF theo NHIỀU BƯỚC dự phòng - đúng mục IV-VII tài liệu
    "ĐÁNH LẠI THỨ TỰ TAB VÀ HOÀN THIỆN TAB 8 MỤC 8D". Excel tổng là NGUỒN CHÍNH để ghi báo cáo.

    TUYỆT ĐỐI KHÔNG kết luận "không có trong Excel tổng" chỉ vì khóa từ TÊN FILE chưa khớp ngay -
    phải thử đủ các bước: tên file → chuẩn hóa lâm nghiệp → đọc nội dung PDF → đối chiếu theo
    Tờ+Thửa (không cần Mã xã) → đối chiếu thêm Diện tích để xác nhận. Nếu 1 khóa khớp NHIỀU dòng
    Excel tổng, KHÔNG tự chọn - đưa vào TRUNG_KHOA_EXCEL_TONG.

    doc_so_sanh=True: đọc nội dung PDF để tạo dữ liệu so sánh (sheet So_sanh) cho MỌI hồ sơ đã
    khớp - chính xác hơn nhưng CHẬM (OCR từng file). doc_so_sanh=False: bỏ qua bước này khi đã
    khớp được từ tên file - nhanh hơn NHIỀU cho lô lớn, phù hợp khi chỉ cần điền báo cáo, không
    cần xem trước sự khác biệt so với PDF gốc.

    Trả về dict:
    - "matched": list {pdf_file, pdf_path, info, excel_row, key, phuong_thuc_khop}
    - "pdf_khong_co_trong_excel": list (pdf_file, key, ghi_chu)
    - "trung_khoa": list (pdf_file, key, list_dong_excel_trung) - cần người dùng tự chọn
    - "excel_khong_co_pdf": list row_data (Excel tổng không có PDF tương ứng)
    - "so_sanh_rows": list đầy đủ các dòng cho sheet "So_sanh"
    """
    log_cb("Đang đọc Excel tổng đã chỉnh sửa...")
    indexes, rows_list_tong = read_excel_tong_rows(excel_tong_path, sheet_tong, header_row_tong,
                                                    col_map_tong, dientich_dung_sai=dientich_dung_sai,
                                                    use_memory_cache=True)
    log_cb(f"  -> Đọc được {len(rows_list_tong)} dòng, {len(indexes['by_key'])} khóa Mã xã+Tờ+Thửa duy nhất.")

    files = list_files_from_source(source_pdfs, ".pdf")
    if not files:
        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")

    matched = []
    pdf_khong_co_trong_excel = []
    trung_khoa = []
    so_sanh_rows = []
    used_row_indexes = set()
    stt_so_sanh = 0

    FIELDS_SO_SANH = [
        ("maxa", "Mã xã"), ("to", "Số tờ"), ("thua", "Số thửa"), ("ten", "Họ tên chủ sử dụng"),
        ("cccd", "CCCD/số định danh"), ("dia_chi_thuong_tru", "Địa chỉ thường trú"),
        ("dia_chi", "Địa chỉ thửa đất"), ("dt", "Diện tích"), ("mucdich", "Mục đích sử dụng đất"),
    ]
    FIELD_KEY_IN_TONG = {"maxa": "maxa", "to": "soto", "thua": "sothua", "ten": "ten",
                        "cccd": "cccd", "dia_chi_thuong_tru": "diachitt", "dia_chi": "diachithua",
                        "dt": "dientich", "mucdich": "mucdich"}
    FIELD_TYPE_FOR_COMPARE = {"maxa": "maxa", "to": "so_to", "thua": "so_thua", "ten": "ten",
                             "cccd": "cccd", "dia_chi_thuong_tru": "diachitt", "dia_chi": "diachithua",
                             "dt": "dientich", "mucdich": "mucdich"}

    def _try_by_key(maxa, soto, sothua):
        if not (maxa and soto and sothua):
            return None
        rows = indexes["by_key"].get(f"{maxa}_{soto}_{sothua}")
        return rows

    def _try_forest_variants(maxa, so_to_goc, sothua, loai_dat=None):
        """Mục VI: cả 2 phía phải chuẩn hóa lâm nghiệp thống nhất. Nếu tên file ghi tờ gốc 1/2/3."""
        if so_to_goc not in ("1", "2", "3"):
            return None, None
        for so_to_thu in ("110000", "210000", "310000"):
            rows = indexes["by_key"].get(f"{maxa}_{so_to_thu}_{sothua}")
            if rows:
                return rows, so_to_thu
        return None, None

    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        fname = os.path.basename(fpath)
        progress_cb and progress_cb(i + 1, len(files), fname)
        log_cb(f"→ Đang đối chiếu: {fname}")

        excel_rows = None
        phuong_thuc_khop = None
        key_used = None
        info = None  # nội dung PDF - chỉ đọc khi thực sự cần (tốc độ - mục XI/XVII)

        # BƯỚC 1: tách khóa từ TÊN FILE trước (nhanh, không cần đọc nội dung/OCR).
        key_from_name = parse_land_key_from_filename(fname)
        if key_from_name["status"] == "DA_TACH_TU_TEN_FILE" and key_from_name["ma_xa"]:
            maxa1, soto1, sothua1 = key_from_name["ma_xa"], key_from_name["so_to"], key_from_name["so_thua"]
            rows = _try_by_key(maxa1, soto1, sothua1)
            if rows:
                excel_rows, phuong_thuc_khop, key_used = rows, "KHOP_TU_TEN_FILE", f"{maxa1}_{soto1}_{sothua1}"
            else:
                # BƯỚC tiếp: thử số tờ chuẩn lâm nghiệp (đã chuẩn hóa 2 phía thống nhất - mục VI)
                rows, so_to_thu = _try_forest_variants(maxa1, soto1, sothua1)
                if rows:
                    excel_rows = rows
                    phuong_thuc_khop = "KHOP_SAU_CHUAN_HOA"
                    key_used = f"{maxa1}_{so_to_thu}_{sothua1}"
                    log_cb(f"   🌲 Số tờ gốc {soto1} khớp Số tờ chuẩn lâm nghiệp {so_to_thu} theo Excel tổng.")

        # BƯỚC 2: tên file không tách được HOẶC khớp tên file thất bại - đọc NỘI DUNG PDF (chậm
        # hơn, có thể cần OCR) rồi thử lại theo mục V.3 - KHÔNG được kết luận sớm là không có.
        if not excel_rows:
            info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, ocr_cache_dir=get_ocr_cache_dir())
            if "error" not in info:
                rows = _try_by_key(info.get("maxa"), info.get("to"), info.get("thua"))
                if rows:
                    excel_rows = rows
                    phuong_thuc_khop = "KHOP_TU_NOI_DUNG_PDF"
                    key_used = f"{info.get('maxa')}_{info.get('to')}_{info.get('thua')}"
                    log_cb(f"   📄 Không khớp theo tên file - đã đọc nội dung PDF và khớp được: {key_used}")

        # BƯỚC 3: vẫn chưa khớp - thử theo Tờ+Thửa KHÔNG cần Mã xã (mục V.5, VII.2) - dự phòng khi
        # Mã xã tên file sai/thiếu. Ưu tiên dùng Tờ/Thửa từ tên file nếu có, không thì từ nội dung.
        if not excel_rows:
            soto_fallback = (key_from_name.get("so_to") if key_from_name["status"] == "DA_TACH_TU_TEN_FILE"
                            else None) or (info.get("to") if info else None)
            sothua_fallback = (key_from_name.get("so_thua") if key_from_name["status"] == "DA_TACH_TU_TEN_FILE"
                              else None) or (info.get("thua") if info else None)
            if soto_fallback and sothua_fallback:
                rows = indexes["by_parcel"].get(f"{soto_fallback}_{sothua_fallback}")
                if rows and len(rows) == 1:
                    excel_rows = rows
                    phuong_thuc_khop = "KHOP_CAN_XAC_NHAN"
                    key_used = f"(?)_{soto_fallback}_{sothua_fallback}"
                    log_cb(f"   ⚠ Khớp theo Tờ+Thửa (không xác nhận được Mã xã) - CẦN NGƯỜI DÙNG XÁC NHẬN: {key_used}")
                elif rows and len(rows) > 1:
                    # BƯỚC 4: nhiều dòng cùng Tờ+Thửa - thử thu hẹp bằng Diện tích (mục V.5.2, VII.3)
                    dt_content = _normalize_area_for_compare(info.get("dt")) if info else None
                    if dt_content is not None:
                        dt_lam_tron = round(dt_content / dientich_dung_sai) * dientich_dung_sai
                        rows_dt = indexes["by_parcel_area"].get(f"{soto_fallback}_{sothua_fallback}_{dt_lam_tron}")
                        if rows_dt and len(rows_dt) == 1:
                            excel_rows = rows_dt
                            phuong_thuc_khop = "KHOP_THEO_DIEN_TICH"
                            key_used = f"(?)_{soto_fallback}_{sothua_fallback}_dt={dt_lam_tron}"
                            log_cb(f"   📐 Nhiều dòng cùng Tờ+Thửa - đã thu hẹp đúng 1 dòng nhờ Diện tích: {key_used}")
                    if not excel_rows:
                        trung_khoa.append((fname, f"{soto_fallback}_{sothua_fallback}", rows))
                        log_cb(f"   ⚠ TRÙNG KHÓA: {len(rows)} dòng Excel tổng cùng Tờ {soto_fallback}+Thửa "
                              f"{sothua_fallback} - cần người dùng tự chọn, KHÔNG tự lấy dòng đầu.")
                        continue

        if not excel_rows:
            pdf_khong_co_trong_excel.append((fname, key_used or "", "Không tìm thấy trong Excel tổng "
                                             "(đã thử tên file, chuẩn hóa, nội dung PDF, Tờ+Thửa)"))
            log_cb(f"   ✗ KHONG_TIM_THAY_TRONG_EXCEL_TONG: {fname}")
            continue

        excel_row = excel_rows[0]  # tại đây excel_rows luôn có đúng 1 dòng (đã lọc trùng khóa ở trên)
        used_row_indexes.add(excel_row["row_index"])

        # BƯỚC so sánh: đọc nội dung PDF nếu CHƯA đọc và người dùng CÓ bật so sánh (mặc định có)
        if info is None and doc_so_sanh:
            info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb, ocr_cache_dir=get_ocr_cache_dir())
        if info is None or "error" in info:
            info = {"maxa": None, "to": None, "thua": None}

        matched.append({"pdf_file": fname, "pdf_path": fpath, "info": info, "excel_row": excel_row,
                        "key": key_used, "phuong_thuc_khop": phuong_thuc_khop})

        if doc_so_sanh:
            for field_key, field_label in FIELDS_SO_SANH:
                pdf_val = info.get(field_key)
                excel_val = excel_row.get(FIELD_KEY_IN_TONG[field_key])
                cmp = compare_field_values(FIELD_TYPE_FOR_COMPARE[field_key], pdf_val, excel_val)
                stt_so_sanh += 1
                so_sanh_rows.append({
                    "stt": stt_so_sanh, "file": fname, "path": fpath,
                    "dong_excel_tong": excel_row["row_index"], "dong_excel_baocao": "",
                    "khoa": key_used, "truong": field_label,
                    "gia_tri_pdf": pdf_val if pdf_val is not None else "",
                    "gia_tri_excel_tong": excel_val if excel_val is not None else "",
                    "ket_qua": cmp["ket_qua"], "co_chinh_sua": cmp["co_chinh_sua"], "ghi_chu": cmp["ghi_chu"],
                })

    excel_khong_co_pdf = [row for row in rows_list_tong if row["row_index"] not in used_row_indexes
                          and row["soto"] and row["sothua"]]

    log_cb(f"\nKết quả: {len(matched)} khớp, {len(pdf_khong_co_trong_excel)} không tìm thấy, "
          f"{len(trung_khoa)} trùng khóa cần tự chọn, {len(excel_khong_co_pdf)} dòng Excel tổng không có PDF.")

    return {"matched": matched, "pdf_khong_co_trong_excel": pdf_khong_co_trong_excel,
            "trung_khoa": trung_khoa, "excel_khong_co_pdf": excel_khong_co_pdf, "so_sanh_rows": so_sanh_rows}


def clear_old_report_data_safe(ws, columns, start_row, end_row, log_cb=None):
    """
    Xóa dữ liệu CŨ trong vùng báo cáo trước khi ghi mới - đúng mục IV tài liệu "CHỈNH SỬA TAB 7,
    TAB 8 VÀ HOÀN THIỆN MỤC 8D": CHỈ xóa GIÁ TRỊ trong đúng các cột được ánh xạ, TUYỆT ĐỐI KHÔNG
    đụng đến công thức, merge cell, định dạng, macro hay bất kỳ ô nào ngoài phạm vi.

    columns: list các chữ cái cột cần xóa (VD ["B","H","I",...]) - lấy từ giá trị col_map_report.
    start_row, end_row: phạm vi dòng cần xóa (bao gồm cả 2 đầu).

    Trả về dict {"da_xoa": số_ô_đã_xóa, "bo_qua_cong_thuc": số_ô_có_công_thức_được_giữ_nguyên}.
    """
    log_cb = log_cb or (lambda x: None)
    da_xoa = 0
    bo_qua_cong_thuc = 0
    for col in columns:
        col = col.strip().upper()
        if not col:
            continue
        for r in range(start_row, end_row + 1):
            cell = ws[f"{col}{r}"]
            # Ô có công thức (bắt đầu bằng "=") - TUYỆT ĐỐI KHÔNG xóa, chỉ xóa giá trị THƯỜNG.
            if isinstance(cell.value, str) and cell.value.startswith("="):
                bo_qua_cong_thuc += 1
                continue
            if cell.value is not None:
                cell.value = None
                da_xoa += 1
    log_cb(f"Đã xóa {da_xoa} ô dữ liệu cũ trong vùng báo cáo"
          + (f" (giữ nguyên {bo_qua_cong_thuc} ô có công thức)" if bo_qua_cong_thuc else "") + ".")
    return {"da_xoa": da_xoa, "bo_qua_cong_thuc": bo_qua_cong_thuc}


def write_report_with_so_sanh(template_excel_path, sheet_report, header_row_report, col_map_report,
                               reconcile_result, output_path, log_cb, clear_old_data=False,
                               clear_start_row=None, clear_end_row=None):
    """
    Ghi báo cáo Excel từ kết quả reconcile_excel_tong_with_pdfs() - dữ liệu lấy từ EXCEL TỔNG
    (không phải PDF), giữ nguyên định dạng file mẫu, thêm sheet "So_sanh" (mục VII/VIII tài liệu).
    KHÔNG sửa file mẫu gốc - copy ra file MỚI.

    Hỗ trợ ĐÚNG file mẫu .xlsm có macro (mục V tài liệu "CHỈNH SỬA TAB 7, TAB 8 VÀ HOÀN THIỆN MỤC
    8D") - tự phát hiện qua đuôi file, dùng keep_vba=True khi mở để KHÔNG làm mất VBA project.
    output_path SẼ TỰ ĐỘNG đổi đuôi cho khớp với file mẫu gốc (.xlsm giữ .xlsm, .xlsx giữ .xlsx)
    nếu người gọi lỡ truyền sai đuôi - tuyệt đối không được âm thầm biến .xlsm thành .xlsx.
    """
    from openpyxl.styles import PatternFill
    is_xlsm = template_excel_path.lower().endswith((".xlsm", ".xltm"))
    if is_xlsm and not output_path.lower().endswith(".xlsm"):
        output_path = os.path.splitext(output_path)[0] + ".xlsm"
        log_cb(f"⚠ File mẫu có macro (.xlsm) - đã tự đổi tên file đầu ra thành: {os.path.basename(output_path)}")

    shutil.copy(template_excel_path, output_path)
    wb = openpyxl.load_workbook(output_path, keep_vba=is_xlsm)
    if sheet_report not in wb.sheetnames:
        raise ValueError(f"Không tìm thấy sheet '{sheet_report}' trong file Excel báo cáo mẫu.")
    ws = wb[sheet_report]

    if clear_old_data:
        cols_to_clear = [c for c in col_map_report.values() if c and c.strip()]
        clear_start = clear_start_row if clear_start_row is not None else header_row_report + 1
        clear_end = clear_end_row if clear_end_row is not None else ws.max_row
        clear_old_report_data_safe(ws, cols_to_clear, clear_start, clear_end, log_cb=log_cb)

    write_row = header_row_report + 1
    log_rows = []
    for i, m in enumerate(reconcile_result["matched"], start=1):
        excel_row = m["excel_row"]
        values = {
            "maxa": excel_row.get("maxa") or "", "ten": excel_row.get("ten") or "",
            "cccd": excel_row.get("cccd") or "", "ngaysinh": excel_row.get("ngaysinh") or "",
            "gioitinh": excel_row.get("gioitinh") or "", "diachitt": excel_row.get("diachitt") or "",
            "soto": excel_row.get("soto") or "", "sothua": excel_row.get("sothua") or "",
            "diachithua": excel_row.get("diachithua") or "", "dientich": excel_row.get("dientich"),
            "mucdich": excel_row.get("mucdich") or "", "dientichloai1": excel_row.get("dientich"),
        }
        r = write_row
        for field, col in col_map_report.items():
            col = (col or "").strip().upper()
            if not col or field not in values:
                continue
            ws[f"{col}{r}"] = values[field]
        write_row += 1
        log_rows.append({"stt": i, "dong_excel_baocao": r, **{k: v for k, v in excel_row.items() if k != "row_index"},
                         "nguon": "excel_tong", "status": "DA_GHI_EXCEL", "note": ""})
        # Cập nhật lại "dòng Excel báo cáo" tương ứng trong so_sanh_rows
        for sr in reconcile_result["so_sanh_rows"]:
            if sr["file"] == m["pdf_file"] and sr["khoa"] == m["key"]:
                sr["dong_excel_baocao"] = r

    ws_ss = wb.create_sheet("So_sanh")
    headers_ss = ["STT", "Tên file PDF", "Đường dẫn PDF", "Dòng Excel tổng", "Dòng Excel báo cáo",
                 "Khóa đối chiếu", "Trường so sánh", "Giá trị trong PDF", "Giá trị trong Excel tổng",
                 "Kết quả so sánh", "Có chỉnh sửa không", "Ghi chú"]
    ws_ss.append(headers_ss)
    for cell in ws_ss[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
    for sr in reconcile_result["so_sanh_rows"]:
        ws_ss.append([sr["stt"], sr["file"], sr["path"], sr["dong_excel_tong"], sr["dong_excel_baocao"],
                     sr["khoa"], sr["truong"], sr["gia_tri_pdf"], sr["gia_tri_excel_tong"],
                     sr["ket_qua"], sr["co_chinh_sua"], sr["ghi_chu"]])
    widths_ss = [6, 35, 45, 14, 14, 20, 22, 30, 30, 20, 16, 45]
    for c, w in enumerate(widths_ss, start=1):
        ws_ss.column_dimensions[ws_ss.cell(row=1, column=c).column_letter].width = w

    from app.services import file_safety_service as _fsafe
    _fsafe.ghi_nguyen_tu(wb.save, output_path)
    log_cb(f"Đã ghi {len(log_rows)} dòng vào báo cáo + sheet 'So_sanh' ({len(reconcile_result['so_sanh_rows'])} dòng so sánh).")
    return log_rows, output_path


def write_log_doi_chieu_excel_tong_va_pdf(csv_path, reconcile_result):
    """Ghi LOG_DOI_CHIEU_EXCEL_TONG_VA_PDF.csv - mục XV.1 tài liệu "BỔ SUNG TAB 8 MỤC 8C"."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file PDF", "Đường dẫn PDF", "Khóa PDF", "Dòng Excel tổng khớp",
                    "Mã xã Excel tổng", "Số tờ Excel tổng", "Số thửa Excel tổng", "Khóa Excel tổng",
                    "Trạng thái", "Ghi chú"])
        stt = 0
        for m in reconcile_result["matched"]:
            stt += 1
            er = m["excel_row"]
            w.writerow([stt, m["pdf_file"], m["pdf_path"], m["key"], er["row_index"],
                       er["maxa"], er["soto"], er["sothua"], m["key"], "KHOP_PDF_VA_EXCEL_TONG", ""])
        for fname, key, note in reconcile_result["pdf_khong_co_trong_excel"]:
            stt += 1
            w.writerow([stt, fname, "", key, "", "", "", "", "", "PDF_KHONG_CO_TRONG_EXCEL_TONG", note])
        for er in reconcile_result["excel_khong_co_pdf"]:
            stt += 1
            key = f"{er['maxa']}_{er['soto']}_{er['sothua']}"
            w.writerow([stt, "", "", "", er["row_index"], er["maxa"], er["soto"], er["sothua"],
                       key, "EXCEL_TONG_KHONG_CO_PDF", ""])


def write_log_so_sanh_pdf_va_excel_tong(csv_path, so_sanh_rows):
    """Ghi LOG_SO_SANH_PDF_VA_EXCEL_TONG.csv - mục XV.3 tài liệu."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file PDF", "Khóa hồ sơ", "Trường so sánh", "Giá trị PDF",
                    "Giá trị Excel tổng", "Kết quả so sánh", "Có chỉnh sửa không", "Ghi chú"])
        for sr in so_sanh_rows:
            w.writerow([sr["stt"], sr["file"], sr["khoa"], sr["truong"], sr["gia_tri_pdf"],
                       sr["gia_tri_excel_tong"], sr["ket_qua"], sr["co_chinh_sua"], sr["ghi_chu"]])


def write_log_ghi_bao_cao_tu_excel_tong(csv_path, log_rows):
    """Ghi LOG_GHI_BAO_CAO_TU_EXCEL_TONG.csv - mục XV.2 tài liệu (log_rows từ write_report_with_so_sanh())."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Dòng Excel báo cáo", "Mã xã", "Số tờ", "Số thửa", "Họ tên", "CCCD",
                    "Ngày sinh", "Giới tính", "Địa chỉ thường trú", "Địa chỉ thửa đất", "Diện tích",
                    "Mục đích sử dụng đất", "Loại đất", "Nguồn dữ liệu", "Trạng thái", "Ghi chú"])
        for r in log_rows:
            w.writerow([r["stt"], r["dong_excel_baocao"], r.get("maxa", ""), r.get("soto", ""),
                       r.get("sothua", ""), r.get("ten", ""), r.get("cccd", ""), r.get("ngaysinh", ""),
                       r.get("gioitinh", ""), r.get("diachitt", ""), r.get("diachithua", ""),
                       r.get("dientich", ""), r.get("mucdich", ""), r.get("loai_dat", ""),
                       r.get("nguon", ""), r.get("status", ""), r.get("note", "")])


def write_danh_sach_can_kiem_tra_tong(csv_path, reconcile_result):
    """
    Ghi DANH_SACH_CAN_KIEM_TRA.csv - mục XV.4 tài liệu: tổng hợp PDF không có trong Excel tổng,
    Excel tổng không có PDF, và các dòng có mâu thuẫn dữ liệu nghiêm trọng (KHAC ở trường khóa).
    """
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Loại vấn đề", "Tên file PDF", "Khóa", "Chi tiết"])
        for fname, key, note in reconcile_result["pdf_khong_co_trong_excel"]:
            w.writerow(["PDF_KHONG_CO_TRONG_EXCEL_TONG", fname, key, note])
        for er in reconcile_result["excel_khong_co_pdf"]:
            key = f"{er['maxa']}_{er['soto']}_{er['sothua']}"
            w.writerow(["EXCEL_TONG_KHONG_CO_PDF", "", key,
                       f"Dòng Excel tổng {er['row_index']} không có PDF tương ứng"])
        for sr in reconcile_result["so_sanh_rows"]:
            if sr["truong"] in ("Mã xã", "Số tờ", "Số thửa") and sr["ket_qua"] == "KHAC":
                w.writerow(["DU_LIEU_MAU_THUAN_NGHIEM_TRONG", sr["file"], sr["khoa"],
                           f"{sr['truong']}: PDF={sr['gia_tri_pdf']}, Excel tổng={sr['gia_tri_excel_tong']}"])


def archive_report(report_path, metadata, log_cb=None):
    """
    Lưu 1 bản sao báo cáo đã hoàn thiện vào kho lưu trữ nội bộ - đúng mục XI tài liệu "ĐÁNH LẠI
    THỨ TỰ TAB VÀ HOÀN THIỆN TAB 8 MỤC 8D". Cấu trúc: Output/Bao_cao_da_hoan_thien/YYYY/MM/.
    Đồng thời ghi 1 dòng vào Logs/LICH_SU_BAO_CAO.csv để tra cứu/phục hồi sau này.

    report_path: đường dẫn file báo cáo đã xuất (đã có ở thư mục người dùng chọn - hàm này chỉ
    COPY THÊM 1 bản vào kho lưu trữ, KHÔNG di chuyển/xóa bản gốc).
    metadata: dict các thông tin mô tả lần chạy - xem đủ cột trong LICH_SU_BAO_CAO.csv bên dưới.

    Trả về đường dẫn file đã lưu trong kho lưu trữ.
    """
    log_cb = log_cb or (lambda x: None)
    now = datetime.datetime.now()
    archive_dir = os.path.join(get_app_data_dir(), "Output", "Bao_cao_da_hoan_thien",
                               now.strftime("%Y"), now.strftime("%m"))
    os.makedirs(archive_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(report_path))[0]
    archive_path = os.path.join(archive_dir, f"{base_name}_{now.strftime('%Y%m%d_%H%M%S')}.xlsx")
    shutil.copy(report_path, archive_path)
    log_cb(f"Đã lưu vào kho lưu trữ: {archive_path}")

    history_csv = os.path.join(get_app_data_dir(), "Logs", "LICH_SU_BAO_CAO.csv")
    os.makedirs(os.path.dirname(history_csv), exist_ok=True)
    is_new = not os.path.isfile(history_csv)
    import csv
    with open(history_csv, "a", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        if is_new:
            w.writerow(["MaBaoCao", "TenBaoCao", "DuongDanFile", "FileExcelTongNguon", "FileBaoCaoMau",
                       "ThuMucHoSoNguon", "SheetExcelTong", "SheetBaoCao", "SoHoSoDaQuet", "SoHoSoKhop",
                       "SoHoSoDaCapNhat", "SoHoSoCanKiemTra", "NguoiThucHien", "ThoiGianBatDau",
                       "ThoiGianHoanThanh", "TrangThai", "GhiChu"])
        ma_bao_cao = now.strftime("%Y%m%d%H%M%S")
        w.writerow([
            ma_bao_cao, os.path.basename(archive_path), archive_path,
            metadata.get("excel_tong_nguon", ""), metadata.get("file_bao_cao_mau", ""),
            metadata.get("thu_muc_ho_so_nguon", ""), metadata.get("sheet_excel_tong", ""),
            metadata.get("sheet_bao_cao", ""), metadata.get("so_ho_so_da_quet", ""),
            metadata.get("so_ho_so_khop", ""), metadata.get("so_ho_so_da_cap_nhat", ""),
            metadata.get("so_ho_so_can_kiem_tra", ""), metadata.get("nguoi_thuc_hien", ""),
            metadata.get("thoi_gian_bat_dau", ""), now.strftime("%Y-%m-%d %H:%M:%S"),
            metadata.get("trang_thai", "DA_LUU_TRU"), metadata.get("ghi_chu", ""),
        ])
    log_cb(f"Đã ghi lịch sử báo cáo: {history_csv}")
    return archive_path


def read_report_history():
    """
    Đọc lịch sử báo cáo từ Logs/LICH_SU_BAO_CAO.csv - dùng cho giao diện kho lưu trữ (mục XI).
    Trả về list dict (mỗi dòng CSV thành 1 dict), rỗng nếu chưa có lịch sử nào.
    """
    history_csv = os.path.join(get_app_data_dir(), "Logs", "LICH_SU_BAO_CAO.csv")
    if not os.path.isfile(history_csv):
        return []
    import csv
    with open(history_csv, encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_wb8c_log_csv(csv_path, log_rows):
    """Ghi log CSV đầy đủ các cột theo mục XV tài liệu "LÀM SÂU CHỨC NĂNG TỔNG HỢP PDF RA EXCEL"."""
    import csv
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file PDF", "Đường dẫn PDF", "Mã xã", "Xã cũ", "Thôn", "Họ tên", "CCCD",
                    "Ngày sinh", "Giới tính", "Địa chỉ thường trú", "Số tờ gốc đọc được",
                    "Số tờ chuẩn sau xử lý", "Số thửa", "Địa chỉ thửa đất", "Diện tích",
                    "Mục đích sử dụng đất", "Loại đất", "Có chuyển số tờ lâm nghiệp",
                    "Nguồn mã xã", "Nguồn số tờ", "Nguồn số thửa", "Nguồn loại đất",
                    "Điểm tin cậy", "Dòng Excel đã ghi", "Trạng thái", "Ghi chú"])
        for r in log_rows:
            w.writerow([r.get("stt", ""), r["file"], r.get("path", ""), r["maxa"],
                        r.get("xa_cu", ""), r.get("thon", ""), r["ten"], r["cccd"],
                        r.get("ngaysinh", ""), r.get("gioitinh", ""), r.get("diachitt", ""),
                        r.get("so_to_goc", r.get("to", "")), r["to"], r["thua"],
                        r.get("diachithua", ""), r["dt"], r["mucdich"],
                        r.get("loai_dat", ""), r.get("co_chuyen_so_to_lam_nghiep", ""),
                        r.get("nguon_maxa", ""), r.get("nguon_so_to", ""), r.get("nguon_so_thua", ""),
                        r.get("nguon_loai_dat", ""), r.get("diem_tin_cay", ""),
                        r.get("dong_excel", ""), r["status"], r["note"]])


def write_wb8c_check_csv(csv_path, log_rows):
    """Xuất riêng danh sách các file CẦN KIỂM TRA (mọi trạng thái khác DA_GHI_EXCEL/CHAY_THU)."""
    import csv
    check_statuses = {"CAN_KIEM_TRA", "THIEU_MA_XA", "THIEU_SO_TO", "THIEU_SO_THUA",
                      "OCR_YEU", "TRUNG_KHOA", "LOI_DOC_PDF", "LOI_GHI_EXCEL"}
    rows_to_check = [r for r in log_rows if r["status"] in check_statuses]
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["STT", "Tên file PDF", "Đường dẫn PDF", "Mã xã", "Số tờ", "Số thửa",
                    "Trạng thái", "Ghi chú"])
        for r in rows_to_check:
            w.writerow([r.get("stt", ""), r["file"], r.get("path", ""), r["maxa"], r["to"], r["thua"],
                        r["status"], r["note"]])
    return len(rows_to_check)


# ============================================================================
# MỤC 8b (MỚI): KHỐI LƯỢNG VÀ TIẾN ĐỘ THỰC HIỆN
# ============================================================================

def scan_completed_pdfs(source_pdfs, xa_mapping, thon_mapping, detect_from_filename, detect_from_content,
                         use_ocr_fallback, ocr_dpi, log_cb, control=None):
    """
    Quét danh sách PDF "đã thực hiện", nhận diện Mã xã/Số tờ/Số thửa:
    1) Từ TÊN FILE trước (nhanh) nếu detect_from_filename=True.
    2) Nếu vẫn thiếu Số tờ/Số thửa và detect_from_content=True thì đọc NỘI DUNG (kể cả OCR).
    Trả về list dict: {stt, maxa, to, thua, file, path, status, note}.
    status: DA_NHAN_DIEN (đủ tờ+thửa) / CAN_KIEM_TRA (thiếu).
    """
    files = list_files_from_source(source_pdfs, ".pdf")
    results = []
    for i, fpath in enumerate(files):
        if control:
            try:
                control.checkpoint()
            except TaskCancelled:
                log_cb(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                break
        fname = os.path.basename(fpath)
        maxa = to = thua = None
        note = ""

        if detect_from_filename:
            key = extract_key_from_filename_v2(fname)
            if key:
                maxa, to, thua = key.get("maxa"), key.get("to"), key.get("thua")

        if (not to or not thua) and detect_from_content:
            log_cb(f"→ Đọc nội dung: {fname}")
            info = extract_fields_from_pdf(fpath, xa_mapping, thon_mapping, use_ocr_fallback, ocr_dpi, log_cb)
            if "error" in info:
                note = info["error"]
            else:
                maxa = maxa or info.get("maxa")
                to = to or info.get("to")
                thua = thua or info.get("thua")

        status = "DA_NHAN_DIEN" if (to and thua) else "CAN_KIEM_TRA"
        if status == "CAN_KIEM_TRA" and not note:
            note = "Không nhận diện được Số tờ/Số thửa (từ tên file" + \
                   (" lẫn nội dung)." if detect_from_content else "; chưa bật đọc nội dung).")

        results.append({"stt": i + 1, "maxa": maxa or "", "to": to or "", "thua": thua or "",
                        "file": fname, "path": fpath, "status": status, "note": note})
    return results


def compute_progress(total, done, in_progress):
    """Tính Chưa thực hiện + Tỷ lệ hoàn thành (%) từ Tổng khối lượng/Đã thực hiện/Đang thực hiện."""
    total = total or 0
    done = done or 0
    in_progress = in_progress or 0
    chua = max(total - done - in_progress, 0)
    ty_le = round(done / total * 100, 1) if total > 0 else 0.0
    return chua, ty_le


def export_execution_progress_report(path, total, done, in_progress, chua, ty_le, completed_list, check_list):
    """
    Xuất BAO_CAO_TIEN_DO_THUC_HIEN.xlsx gồm 3 sheet:
    - Tong_hop: Tổng/Đã/Đang/Chưa/Tỷ lệ/Ngày xuất (tô đậm).
    - Danh_sach_da_thuc_hien: STT/Mã xã/Số tờ/Số thửa/Tên file/Đường dẫn/Trạng thái/Ghi chú.
    - Can_kiem_tra: các file chưa nhận diện được số tờ/số thửa.
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    bold = Font(bold=True)
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F4E78")
    center = Alignment(horizontal="center", vertical="center")
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws1 = wb.active
    ws1.title = "Tong_hop"
    ws1.append(["Chỉ tiêu", "Giá trị"])
    for c in range(1, 3):
        cell = ws1.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    rows = [
        ("Tổng khối lượng", total), ("Đã thực hiện", done), ("Đang thực hiện", in_progress),
        ("Chưa thực hiện", chua), ("Tỷ lệ hoàn thành (%)", ty_le),
        ("Ngày xuất báo cáo", datetime.datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]
    for label, val in rows:
        r = ws1.max_row + 1
        c1 = ws1.cell(row=r, column=1, value=label)
        c2 = ws1.cell(row=r, column=2, value=val)
        c1.font = bold
        c1.border = border
        c2.border = border
    ws1.column_dimensions["A"].width = 26
    ws1.column_dimensions["B"].width = 20

    ws2 = wb.create_sheet("Danh_sach_da_thuc_hien")
    headers2 = ["STT", "Mã xã", "Số tờ", "Số thửa", "Tên file PDF", "Đường dẫn", "Trạng thái", "Ghi chú"]
    ws2.append(headers2)
    for c in range(1, len(headers2) + 1):
        cell = ws2.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    for item in completed_list:
        ws2.append([item.get("stt"), item.get("maxa"), item.get("to"), item.get("thua"),
                   item.get("file"), item.get("path"), item.get("status"), item.get("note")])
    widths2 = [6, 10, 8, 8, 34, 46, 14, 30]
    for c, w in enumerate(widths2, start=1):
        ws2.column_dimensions[ws2.cell(row=1, column=c).column_letter].width = w

    ws3 = wb.create_sheet("Can_kiem_tra")
    headers3 = ["STT", "Tên file PDF", "Đường dẫn", "Ghi chú"]
    ws3.append(headers3)
    for c in range(1, len(headers3) + 1):
        cell = ws3.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    for item in check_list:
        ws3.append([item.get("stt"), item.get("file"), item.get("path"), item.get("note")])
    widths3 = [6, 34, 46, 40]
    for c, w in enumerate(widths3, start=1):
        ws3.column_dimensions[ws3.cell(row=1, column=c).column_letter].width = w

    dirpath = os.path.dirname(path)
    if dirpath:
        os.makedirs(dirpath, exist_ok=True)
    wb.save(path)


# ------------------------------ GIAO DIỆN ------------------------------

# ------------------------------ WIDGET DÙNG CHUNG ------------------------------

def styled_labelframe(parent, text="", **kwargs):
    """
    LabelFrame CHUẨN dùng cho MỌI nhóm chức năng trong toàn phần mềm — tiêu đề TÔ ĐẬM, màu chữ
    nổi bật (xanh đậm), để các nhóm (VD "Đối chiếu hồ sơ", "Xuất hồ sơ Word/PDF", "Tổng hợp báo
    cáo từ PDF"...) hiển thị rõ ràng, dễ phân biệt trên giao diện.
    Dùng thay cho tk.LabelFrame(...) trực tiếp ở mọi nơi trong phần mềm.
    Cho phép ghi đè font/fg riêng qua kwargs nếu 1 nhóm cụ thể cần nổi bật hơn nữa.
    """
    kwargs.setdefault("font", ("Segoe UI", 9, "bold"))
    kwargs.setdefault("fg", "#1a237e")
    return tk.LabelFrame(parent, text=text, **kwargs)


class SourcePicker(tk.LabelFrame):
    """
    Widget cho phép chọn NGUỒN xử lý theo 1 trong 2 cách:
    - Chọn 1 THƯ MỤC (xử lý toàn bộ file đúng loại trong đó)
    - Chọn NHIỀU FILE cụ thể (có thể ở nhiều thư mục khác nhau, chọn 1 hoặc nhiều lần)
    get_source() trả về: str (đường dẫn thư mục) HOẶC list[str] (danh sách file) HOẶC None.
    """

    def __init__(self, parent, title, filetypes, file_label="file", **kwargs):
        kwargs.setdefault("font", ("Segoe UI", 9, "bold"))
        kwargs.setdefault("fg", "#1a237e")
        super().__init__(parent, text=title, **kwargs)
        self.filetypes = filetypes
        self.file_label = file_label
        self.folder_var = tk.StringVar()
        self.selected_files = []

        row1 = tk.Frame(self)
        row1.pack(fill="x", padx=6, pady=3)
        tk.Label(row1, text="Thư mục:", width=12, anchor="w").pack(side="left")
        tk.Entry(row1, textvariable=self.folder_var).pack(side="left", fill="x", expand=True, padx=4)
        tk.Button(row1, text="Chọn thư mục...", command=self._pick_folder).pack(side="left")

        row2 = tk.Frame(self)
        row2.pack(fill="x", padx=6, pady=(0, 3))
        tk.Button(row2, text=f"Hoặc chọn {file_label} (chọn được nhiều)...", command=self._pick_files).pack(side="left")
        tk.Button(row2, text="Xóa danh sách đã chọn", command=self._clear_files).pack(side="left", padx=6)

        self.lbl_status = tk.Label(self, text="Chưa chọn nguồn nào.", fg="#555", anchor="w",
                                    justify="left", wraplength=860)
        self.lbl_status.pack(fill="x", padx=6, pady=(0, 6))

    def _pick_folder(self):
        p = filedialog.askdirectory(title="Chọn thư mục")
        if p:
            self.folder_var.set(p)
            self.selected_files = []
            self._update_status()

    def _pick_files(self):
        paths = filedialog.askopenfilenames(title=f"Chọn {self.file_label} (có thể chọn nhiều)",
                                             filetypes=self.filetypes)
        if paths:
            self.selected_files = list(paths)
            self.folder_var.set("")
            self._update_status()

    def _clear_files(self):
        self.selected_files = []
        self.folder_var.set("")
        self._update_status()

    def _update_status(self):
        if self.selected_files:
            names = ", ".join(os.path.basename(f) for f in self.selected_files[:5])
            more = f" (+{len(self.selected_files) - 5} file khác)" if len(self.selected_files) > 5 else ""
            self.lbl_status.config(text=f"Đã chọn {len(self.selected_files)} file cụ thể: {names}{more}")
        elif self.folder_var.get().strip():
            self.lbl_status.config(text=f"Sẽ xử lý TOÀN BỘ file trong thư mục: {self.folder_var.get().strip()}")
        else:
            self.lbl_status.config(text="Chưa chọn nguồn nào.")

    def get_source(self):
        if self.selected_files:
            return list(self.selected_files)
        f = self.folder_var.get().strip()
        return f if f else None


try:
    from app.config.config_loader import load_app_config, save_app_config, create_default_settings_file
    from app.core.models import AppConfig as _AppConfig
    _CONFIG_LOADER_MODULE_OK = True
except Exception:
    _CONFIG_LOADER_MODULE_OK = False

try:
    from app.core.commune_config import (
        load_commune_config, create_sample_config_file as create_commune_sample,
        export_config_to_excel, detect_ma_xa_by_config, build_lookup_from_config,
        load_commune_config_auto, create_sample_config_file_multi_sheet,
    )
    _COMMUNE_CONFIG_MODULE_OK = True
except Exception:
    _COMMUNE_CONFIG_MODULE_OK = False


def load_app_config_safe():
    """Nạp AppConfig từ data/app_settings.json (dùng config_loader nếu có). Luôn trả về 1 object
    cấu hình có đủ thuộc tính - kể cả khi thiếu module/file (dùng đối tượng mặc định đơn giản)."""
    if _CONFIG_LOADER_MODULE_OK:
        return load_app_config(get_app_dir())

    # Fallback tối giản: 1 object có sẵn các thuộc tính cấu hình mặc định
    class _FallbackConfig:
        output_folder = "Output"
        logs_folder = "Logs"
        temp_folder = "Temp"
        backup_folder = "Backup"
        undo_folder = "Undo"
        commune_config_file = "Cau_hinh_ma_xa_thon.xlsx"
        default_file_name_template = "CHUACOGIAY_{maxa}_{soto}_{sothua}_{loaihoso}"
        ocr_language = "vie+eng"
        ocr_mode = "Chuẩn"
        require_dry_run = False
        auto_backup = False
        word_pdf_safe_mode = True
        update_check_enabled = True
        nguong_tin_cay_cao = 90.0
        nguong_tin_cay_thap = 70.0
        che_do_nang_cao = False
        word_toc_do_xu_ly = "can_bang"
        word_so_file_moi_dot = 100
        word_nghi_giua_dot_giay = 2.0
        word_hien_thi_thong_ke_hieu_nang = True
        word_tu_khoi_dong_lai_khi_loi = True
        word_gioi_han_ram_mb = 0
    return _FallbackConfig()


def save_app_config_safe(config):
    """Lưu cấu hình ra data/app_settings.json nếu module config_loader có sẵn. Trả về True/False."""
    if _CONFIG_LOADER_MODULE_OK:
        return save_app_config(get_app_dir(), config)
    return False


try:
    from app.ui.tab_settings import SettingsTabMixin
    _UI_SETTINGS_MODULE_OK = True
except Exception:
    _UI_SETTINGS_MODULE_OK = False

    class SettingsTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_settings.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_settings(self, parent, pad):
            frm = styled_labelframe(parent, text="⚙ CÀI ĐẶT PHẦN MỀM")
            frm.pack(fill="both", expand=True, **pad)

            self._build_license_section(frm, pad)

            # ===== Hướng dẫn sử dụng nhanh =====
            frm_guide = styled_labelframe(frm, text="📖 Hướng dẫn sử dụng nhanh")
            frm_guide.pack(fill="x", padx=10, pady=(10, 6))
            guide_text = (
                "• Lần đầu dùng: vào mục '🗺 Cấu hình địa bàn nâng cao' bên dưới → 'Tạo mẫu' → mở "
                "Excel bổ sung xã/thôn của bạn → 'Nạp & kiểm tra cấu hình'. Nếu chỉ dùng cho Văn Lang, "
                "phần mềm đã có sẵn bảng tra mặc định, có thể bỏ qua bước này.\n"
                "• Chế độ OCR: 'Nhanh' cho hồ sơ rõ nét cần xử lý số lượng lớn; 'Kỹ' cho hồ sơ scan mờ, "
                "chữ nhỏ (chậm hơn). 'Chuẩn' phù hợp đa số trường hợp.\n"
                "• Bật 'Tự động sao lưu' nếu muốn phần mềm tự giữ 1 bản sao file gốc trước khi đổi tên/"
                "di chuyển/chỉnh sửa — an toàn hơn nhưng chiếm thêm dung lượng đĩa theo thời gian.\n"
                "• Bật 'Bắt buộc chạy thử' nếu muốn phần mềm nhắc lại trước khi cho chạy thật mà chưa "
                "từng chạy thử trong phiên làm việc — giảm rủi ro thao tác nhầm.\n"
                "• Mục 'Cấu hình cột Excel báo cáo' (Tab 8) và 'Cấu hình hậu tố tên file' (Tab Word) "
                "nằm ở đúng Tab tương ứng, không nằm ở đây - vì cần xem trực tiếp cùng dữ liệu đang xử lý.\n"
                "• Mọi thao tác đổi tên/di chuyển/ghi đè đều có 'Chạy thử' trước — luôn kiểm tra kỹ bảng "
                "kết quả xem trước khi bấm nút chạy thật."
            )
            tk.Label(frm_guide, text=guide_text, fg="#333", wraplength=980, justify="left",
                    anchor="w").pack(fill="x", padx=8, pady=8)

            # ===== Có gì mới (lịch sử cập nhật) =====
            frm_changelog = styled_labelframe(frm, text=f"🆕 Có gì mới — SỸ LAND {RELEASE_LABEL} (build {APP_VERSION})")
            frm_changelog.pack(fill="x", padx=10, pady=(0, 10))
            changelog_box = tk.Text(frm_changelog, height=9, wrap="word", fg="#333",
                                    font=("Segoe UI", 9), relief="flat", bg="#fafafa", padx=8, pady=6)
            changelog_scroll = tk.Scrollbar(frm_changelog, orient="vertical", command=changelog_box.yview)
            changelog_box.configure(yscrollcommand=changelog_scroll.set)
            changelog_box.tag_configure("ver", font=("Segoe UI", 9, "bold"), foreground="#1a237e")
            for ver, desc in CHANGELOG:
                changelog_box.insert("end", f"v{ver}", "ver")
                changelog_box.insert("end", f"  —  {desc}\n\n")
            changelog_box.configure(state="disabled")
            changelog_box.pack(side="left", fill="both", expand=True, padx=(8, 0), pady=8)
            changelog_scroll.pack(side="left", fill="y", padx=(0, 8), pady=8)

            if not _CONFIG_LOADER_MODULE_OK:
                tk.Label(frm, text="(Module cấu hình chưa sẵn sàng - các thay đổi sẽ không được lưu lại "
                                   "vào file. Phần mềm vẫn dùng cấu hình mặc định.)",
                         fg="#c62828", wraplength=900, justify="left").pack(anchor="w", padx=10, pady=8)

            tk.Label(frm, text="Các cài đặt dưới đây được lưu vào data/app_settings.json và áp dụng cho "
                               "lần chạy sau. Một số tùy chọn (VD chế độ OCR mặc định) sẽ dùng làm giá trị "
                               "khởi tạo cho các tab tương ứng.",
                     fg="#555", wraplength=900, justify="left").pack(anchor="w", padx=10, pady=(8, 10))

            cfg = self.app_config

            # --- Chế độ OCR mặc định ---
            row_ocr = tk.Frame(frm); row_ocr.pack(fill="x", padx=10, pady=4)
            tk.Label(row_ocr, text="Chế độ OCR mặc định:", width=28, anchor="w").pack(side="left")
            self.set_var_ocr_mode = tk.StringVar(value=getattr(cfg, "ocr_mode", "Chuẩn"))
            for mode in ("Nhanh", "Chuẩn", "Kỹ"):
                tk.Radiobutton(row_ocr, text=mode, variable=self.set_var_ocr_mode, value=mode).pack(side="left", padx=4)

            # --- Ngôn ngữ OCR ---
            row_lang = tk.Frame(frm); row_lang.pack(fill="x", padx=10, pady=4)
            tk.Label(row_lang, text="Ngôn ngữ OCR:", width=28, anchor="w").pack(side="left")
            self.set_var_ocr_lang = tk.StringVar(value=getattr(cfg, "ocr_language", "vie+eng"))
            tk.Entry(row_lang, textvariable=self.set_var_ocr_lang, width=20).pack(side="left")
            tk.Label(row_lang, text="(VD: vie+eng)", fg="#888").pack(side="left", padx=6)

            # --- Mẫu tên file mặc định ---
            row_tmpl = tk.Frame(frm); row_tmpl.pack(fill="x", padx=10, pady=4)
            tk.Label(row_tmpl, text="Mẫu tên file mặc định:", width=28, anchor="w").pack(side="left")
            self.set_var_name_tmpl = tk.StringVar(value=getattr(cfg, "default_file_name_template",
                                                                "CHUACOGIAY_{maxa}_{soto}_{sothua}_{loaihoso}"))
            tk.Entry(row_tmpl, textvariable=self.set_var_name_tmpl, width=45).pack(side="left")

            # --- Ngưỡng điểm tin cậy (dùng cho Tổng hợp báo cáo Tab 8) ---
            row_nguong = tk.Frame(frm); row_nguong.pack(fill="x", padx=10, pady=4)
            tk.Label(row_nguong, text="Ngưỡng tin cậy TỰ ĐỘNG xử lý (%):", width=28, anchor="w").pack(side="left")
            self.set_var_nguong_cao = tk.StringVar(value=str(getattr(cfg, "nguong_tin_cay_cao", 90.0)))
            tk.Entry(row_nguong, textvariable=self.set_var_nguong_cao, width=8).pack(side="left")
            tk.Label(row_nguong, text="   Ngưỡng KHÔNG TỰ xử lý (%):").pack(side="left", padx=(16, 4))
            self.set_var_nguong_thap = tk.StringVar(value=str(getattr(cfg, "nguong_tin_cay_thap", 70.0)))
            tk.Entry(row_nguong, textvariable=self.set_var_nguong_thap, width=8).pack(side="left")
            tk.Label(row_nguong, text="(điểm ≥ ngưỡng cao: tự xử lý; điểm < ngưỡng thấp: bắt buộc CẦN KIỂM TRA)",
                    fg="#888").pack(side="left", padx=6)

            # --- Chế độ Cơ bản / Nâng cao (ẩn/hiện tùy chọn kỹ thuật cho người mới) ---
            row_nangcao = tk.Frame(frm); row_nangcao.pack(fill="x", padx=10, pady=4)
            self.set_var_nang_cao = tk.BooleanVar(value=getattr(cfg, "che_do_nang_cao", False))
            tk.Checkbutton(row_nangcao, text="Chế độ NÂNG CAO (hiện thêm các tùy chọn kỹ thuật cho người dùng có kinh nghiệm)",
                           variable=self.set_var_nang_cao,
                           command=lambda: self._apply_advanced_mode_visibility(self.set_var_nang_cao.get())
                           ).pack(anchor="w")
            tk.Label(row_nangcao, text="(Tắt = Cơ bản: ẩn bớt tùy chọn nâng cao/gỡ lỗi ít dùng, giao diện gọn hơn cho người mới)",
                    fg="#888").pack(anchor="w", padx=24)

            # --- Các tùy chọn bật/tắt ---
            self.set_var_auto_backup = tk.BooleanVar(value=getattr(cfg, "auto_backup", False))
            tk.Checkbutton(frm, text="Tự động sao lưu trước các thao tác rủi ro (đổi tên, di chuyển file)",
                           variable=self.set_var_auto_backup).pack(anchor="w", padx=12, pady=3)

            self.set_var_require_dry_run = tk.BooleanVar(value=getattr(cfg, "require_dry_run", False))
            tk.Checkbutton(frm, text="Bắt buộc chạy thử trước khi chạy thật (với thao tác rủi ro)",
                           variable=self.set_var_require_dry_run).pack(anchor="w", padx=12, pady=3)

            self.set_var_word_safe = tk.BooleanVar(value=getattr(cfg, "word_pdf_safe_mode", True))
            tk.Checkbutton(frm, text="Chế độ an toàn khi thao tác Word/Excel (DispatchEx - không ảnh hưởng "
                                     "file đang mở) — khuyến nghị BẬT",
                           variable=self.set_var_word_safe).pack(anchor="w", padx=12, pady=3)

            self.set_var_update_check = tk.BooleanVar(value=getattr(cfg, "update_check_enabled", True))
            tk.Checkbutton(frm, text="Tự động kiểm tra cập nhật khi khởi động (tối đa 1 lần/ngày)",
                           variable=self.set_var_update_check).pack(anchor="w", padx=12, pady=3)

            # --- Nút lưu / khôi phục mặc định ---
            row_btn = tk.Frame(frm); row_btn.pack(fill="x", padx=10, pady=(14, 8))
            tk.Button(row_btn, text="💾 Lưu cài đặt", font=("Segoe UI", 10, "bold"),
                      bg="#2e7d32", fg="white", command=self.save_settings_clicked).pack(side="left", padx=(0, 6))
            tk.Button(row_btn, text="↺ Khôi phục mặc định", command=self.reset_settings_clicked).pack(side="left", padx=6)
            tk.Button(row_btn, text="📂 Mở thư mục cấu hình", command=self.open_settings_folder).pack(side="left", padx=6)
            self.lbl_settings_status = tk.Label(row_btn, text="", fg="#2e7d32")
            self.lbl_settings_status.pack(side="left", padx=12)

            # ================= Cấu hình địa bàn nâng cao (dùng cho bất kỳ xã/thôn nào) =================
            frm_dia_ban = styled_labelframe(parent, text="🗺 Cấu hình địa bàn nâng cao (dùng phần mềm cho BẤT KỲ xã/phường/thôn nào)")
            frm_dia_ban.pack(fill="x", padx=10, pady=(10, 6))

            tk.Label(frm_dia_ban,
                     text="File cấu hình mở rộng (Excel) cho phép khai báo Tỉnh, Xã/phường hiện tại, Xã cũ, "
                          "Mã xã cũ, Thôn/tổ dân phố, Thôn cũ/mới sau sáp nhập, số Nghị quyết HĐND, tên gọi khác, "
                          "từ khóa nhận diện... Dùng để nhận diện mã xã cho địa bàn bất kỳ, không giới hạn Văn Lang.",
                     fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 4))

            if not _COMMUNE_CONFIG_MODULE_OK:
                tk.Label(frm_dia_ban, text="(Module cấu hình địa bàn chưa sẵn sàng.)",
                         fg="#c62828").pack(anchor="w", padx=8, pady=4)
            else:
                self.var_commune_config_path = tk.StringVar(value=self._get_commune_config_default_path())
                row_path = tk.Frame(frm_dia_ban); row_path.pack(fill="x", padx=8, pady=4)
                tk.Label(row_path, text="File cấu hình:", width=14, anchor="w").pack(side="left")
                tk.Entry(row_path, textvariable=self.var_commune_config_path, width=70).pack(side="left", padx=(0, 6))
                tk.Button(row_path, text="Chọn...", command=self.choose_commune_config_file).pack(side="left")

                row_cbtn = tk.Frame(frm_dia_ban); row_cbtn.pack(fill="x", padx=8, pady=(4, 6))
                tk.Button(row_cbtn, text="+ Tạo mẫu đơn giản (1 sheet)", command=self.create_commune_config_sample).pack(side="left", padx=(0, 6))
                tk.Button(row_cbtn, text="+ Tạo mẫu linh động (Xã+Thôn, hỗ trợ sáp nhập)",
                         command=self.create_commune_config_sample_multi).pack(side="left", padx=(0, 6))
                tk.Button(row_cbtn, text="↻ Nạp & kiểm tra cấu hình", command=self.load_commune_config_clicked).pack(side="left", padx=(0, 6))
                tk.Button(row_cbtn, text="📂 Mở file cấu hình", command=self.open_commune_config_file).pack(side="left")
                tk.Label(frm_dia_ban, text="Mẫu 'linh động' dùng khi 1 xã hiện tại được hình thành từ nhiều xã cũ, hoặc "
                                           "nhiều thôn cũ (có thể khác xã cũ) gộp thành 1 thôn mới - phần mềm tự phát hiện "
                                           "đúng định dạng file khi nạp, không cần chọn thủ công.",
                        fg="#777", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(0, 4))
                self.lbl_commune_config_status = tk.Label(frm_dia_ban, text="(chưa nạp)", fg="#777", wraplength=980, justify="left")
                self.lbl_commune_config_status.pack(anchor="w", padx=8, pady=(0, 8))

            # ================= Cache OCR (tăng tốc khi Chạy thử rồi Chạy thật cùng lô hồ sơ) =================
            # LƯU Ý: đặt NGOÀI khối if/else phía trên (không phụ thuộc _COMMUNE_CONFIG_MODULE_OK) -
            # đây là tính năng ĐỘC LẬP với cấu hình địa bàn. (Đã từng có lỗi thật: 2 nhóm này bị
            # lồng nhầm trong nhánh else, khiến chúng biến mất hoàn toàn nếu app/core/commune_config.py
            # không tải được - đã sửa.)
            frm_cache = styled_labelframe(parent, text="⚡ Cache OCR (tăng tốc xử lý)")
            frm_cache.pack(fill="x", padx=10, pady=(0, 10))
            tk.Label(frm_cache,
                    text="Khi OCR 1 file PDF, phần mềm tự lưu lại kết quả. Lần sau OCR CÙNG file (chưa "
                         "bị sửa) với CÙNG chế độ, dùng lại kết quả cũ thay vì OCR lại - nhanh hơn rất "
                         "nhiều khi 'Chạy thử' rồi 'Chạy thật' trên cùng 1 lô hồ sơ.",
                    fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 6))
            row_cache = tk.Frame(frm_cache)
            row_cache.pack(fill="x", padx=8, pady=(0, 8))
            tk.Button(row_cache, text="🗑 Xóa cache OCR", command=self.clear_ocr_cache_clicked).pack(side="left")
            self.lbl_cache_status = tk.Label(row_cache, text="", fg="#777")
            self.lbl_cache_status.pack(side="left", padx=10)

            # ================= Kiểm thử hệ thống =================
            frm_test = styled_labelframe(parent, text="🧪 Kiểm thử hệ thống")
            frm_test.pack(fill="x", padx=10, pady=(0, 10))
            tk.Label(frm_test,
                    text="Tự động kiểm tra 15 luồng chính (OCR, đối chiếu, lọc trùng, xuất Word, tổng "
                         "hợp báo cáo...) bằng dữ liệu mẫu có sẵn (test_data/) - dùng khi nghi ngờ có gì "
                         "không hoạt động đúng, hoặc muốn kiểm tra nhanh trước khi dùng cho việc quan trọng. "
                         "Không ảnh hưởng dữ liệu thật của bạn, chạy trong khoảng 30-60 giây.",
                    fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 6))
            row_test = tk.Frame(frm_test)
            row_test.pack(fill="x", padx=8, pady=(0, 4))
            self.btn_run_system_test = tk.Button(row_test, text="🧪 Chạy kiểm thử hệ thống",
                                                 font=("Arial", 10, "bold"), bg="#00838f", fg="white",
                                                 command=self.run_system_test_clicked)
            self.btn_run_system_test.pack(side="left")
            self.lbl_test_status = tk.Label(row_test, text="", fg="#777")
            self.lbl_test_status.pack(side="left", padx=10)
            self.var_export_acceptance_report = tk.BooleanVar(value=False)
            tk.Checkbutton(frm_test, text="Xuất kèm Báo cáo nghiệm thu chính thức "
                                          "(BAO_CAO_NGHIEM_THU_SY_LAND - dùng khi chốt bản phát hành)",
                          variable=self.var_export_acceptance_report).pack(anchor="w", padx=8, pady=(0, 8))

            self._build_env_setup_section(parent, pad)
            self._build_word_perf_section(parent, pad)
            self._build_diagnostic_section(parent, pad)

        def _import_run_tests_module(self):
            import importlib.util
            base_dir = os.path.dirname(sys.executable) if getattr(sys, "frozen", False) else get_base_dir()
            run_tests_path = os.path.join(base_dir, "run_tests.py")
            if not os.path.isfile(run_tests_path):
                return None
            spec = importlib.util.spec_from_file_location("run_tests", run_tests_path)
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            return module

        def run_system_test_clicked(self):
            run_tests_module = self._import_run_tests_module()
            if run_tests_module is None:
                messagebox.showerror("Không tìm thấy",
                                     "Không tìm thấy file run_tests.py cùng thư mục phần mềm. "
                                     "Đây là file đi kèm bộ cài, hãy kiểm tra lại thư mục cài đặt.")
                return

            self.txt_log.delete("1.0", "end")
            self.btn_run_system_test.config(state="disabled")
            self.lbl_test_status.config(text="Đang chạy kiểm thử...", fg="#f57c00")

            def worker():
                try:
                    results, n_pass, n_fail = run_tests_module.run_all_scenarios(output_fn=self.log)

                    save_dir = get_app_data_dir()
                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    report_path = os.path.join(save_dir, f"BAO_CAO_KIEM_THU_HE_THONG_{stamp}.xlsx")
                    run_tests_module.export_test_report_xlsx(results, report_path)
                    self.log(f"\nĐã xuất báo cáo kiểm thử: {report_path}")

                    acceptance_path = None
                    if self.var_export_acceptance_report.get():
                        acceptance_path = os.path.join(save_dir, f"BAO_CAO_NGHIEM_THU_SY_LAND_v1.0_{stamp}.xlsx")
                        run_tests_module.export_acceptance_report_xlsx(
                            results, acceptance_path, release_label=RELEASE_LABEL, build_version=APP_VERSION)
                        self.log(f"Đã xuất báo cáo nghiệm thu: {acceptance_path}")

                    status_text = f"✓ {n_pass}/{len(results)} PASS" if n_fail == 0 else f"⚠ {n_pass}/{len(results)} PASS, {n_fail} lỗi"
                    status_color = "#2e7d32" if n_fail == 0 else "#c62828"
                    self.after(0, lambda: self.lbl_test_status.config(text=status_text, fg=status_color))

                    msg = f"Kết quả: {n_pass}/{len(results)} PASS\n\nBáo cáo chi tiết:\n{report_path}"
                    if acceptance_path:
                        msg += f"\n\nBáo cáo nghiệm thu:\n{acceptance_path}"
                    messagebox.showinfo("Hoàn tất kiểm thử" if n_fail == 0 else "Kiểm thử có lỗi", msg)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi kiểm thử", str(e))
                finally:
                    self.btn_run_system_test.config(state="normal")

            threading.Thread(target=worker, daemon=True).start()

        def clear_ocr_cache_clicked(self):
            cache_dir = get_ocr_cache_dir()
            if not messagebox.askyesno("Xác nhận",
                                        "Xóa cache OCR sẽ khiến lần OCR tiếp theo cho các file (kể cả đã OCR "
                                        "trước đó) chạy chậm hơn (phải OCR lại từ đầu). Bạn có chắc muốn xóa?"):
                return
            n = clear_ocr_cache(cache_dir)
            self.lbl_cache_status.config(text=f"✓ Đã xóa {n} file cache.", fg="#2e7d32")
            messagebox.showinfo("Hoàn tất", f"Đã xóa {n} file cache OCR.")

        def _get_commune_config_default_path(self):
            return os.path.join(get_app_dir(), "data", "Cau_hinh_ma_xa_thon.xlsx")

        def choose_commune_config_file(self):
            path = filedialog.askopenfilename(title="Chọn file cấu hình địa bàn",
                                              filetypes=[("Excel", "*.xlsx"), ("Tất cả", "*.*")])
            if path:
                self.var_commune_config_path.set(path)

        def create_commune_config_sample(self):
            path = self.var_commune_config_path.get().strip() or self._get_commune_config_default_path()
            if os.path.isfile(path):
                if not messagebox.askyesno("File đã tồn tại", f"File {os.path.basename(path)} đã tồn tại. Ghi đè bằng mẫu mới?"):
                    return
            try:
                create_commune_sample(path)
                self.var_commune_config_path.set(path)
                messagebox.showinfo("Đã tạo", f"Đã tạo file cấu hình mẫu:\n{path}\n\nMở bằng Excel để bổ sung địa bàn, "
                                              f"sau đó bấm 'Nạp & kiểm tra cấu hình'.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không tạo được file: {e}")

        def create_commune_config_sample_multi(self):
            path = self.var_commune_config_path.get().strip() or self._get_commune_config_default_path()
            if os.path.isfile(path):
                if not messagebox.askyesno("File đã tồn tại", f"File {os.path.basename(path)} đã tồn tại. Ghi đè bằng mẫu mới?"):
                    return
            try:
                create_sample_config_file_multi_sheet(path)
                self.var_commune_config_path.set(path)
                messagebox.showinfo("Đã tạo", f"Đã tạo file cấu hình mẫu LINH ĐỘNG (2 sheet Xã+Thôn):\n{path}\n\n"
                                              f"Mở bằng Excel, xem sheet 'Huong_dan' để biết cách nhập 1 xã hiện tại "
                                              f"gồm nhiều xã cũ, hoặc nhiều thôn cũ gộp thành 1 thôn mới, "
                                              f"sau đó bấm 'Nạp & kiểm tra cấu hình'.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không tạo được file: {e}")

        def load_commune_config_clicked(self):
            path = self.var_commune_config_path.get().strip()
            if not path or not os.path.isfile(path):
                messagebox.showerror("Chưa có file", "Vui lòng chọn hoặc tạo file cấu hình địa bàn trước.")
                return
            try:
                rows, errors = load_commune_config_auto(path)
            except ValueError as e:
                self.lbl_commune_config_status.config(text=f"✗ Lỗi nghiêm trọng: {e}", fg="#c62828")
                messagebox.showerror("Lỗi cấu hình", f"File cấu hình thiếu cột bắt buộc:\n{e}")
                return
            except Exception as e:
                self.lbl_commune_config_status.config(text=f"✗ Không đọc được file: {e}", fg="#c62828")
                return

            thon_mapping, xa_mapping = build_lookup_from_config(rows)
            self.commune_config_rows = rows  # lưu lại để làm giàu dữ liệu khi xuất Word (token thôn sáp nhập)
            # Áp dụng ngay vào bảng tra dùng chung ở Tab 5 (nếu đã dựng)
            if hasattr(self, "txt_thon_mapping"):
                self.txt_thon_mapping.delete("1.0", "end")
                self.txt_thon_mapping.insert("1.0", "\n".join(f"{k}={v}" for k, v in thon_mapping.items()))
            if hasattr(self, "txt_xa_mapping"):
                self.txt_xa_mapping.delete("1.0", "end")
                self.txt_xa_mapping.insert("1.0", "\n".join(f"{k}={v}" for k, v in xa_mapping.items()))

            msg = f"✓ Đã nạp {len(rows)} dòng địa bàn ({len(thon_mapping)} tên thôn/tổ dân phố)."
            if errors:
                msg += f"  ⚠ {len(errors)} cảnh báo cấu hình (xem Nhật ký)."
                self.log(f"Cảnh báo cấu hình địa bàn ({len(errors)}):")
                for e in errors[:30]:
                    self.log(f"  - Dòng {e['dong']}, {e['truong']}: {e['loi']} → {e['goi_y']}")
                self.lbl_commune_config_status.config(text=msg, fg="#f57c00")
            else:
                self.lbl_commune_config_status.config(text=msg, fg="#2e7d32")
            self.log(msg + f"  Áp dụng vào bảng tra Thôn/Xã (Tab 5).")

        def open_commune_config_file(self):
            path = self.var_commune_config_path.get().strip()
            if not path or not os.path.isfile(path):
                messagebox.showinfo("Chưa có file", "Chưa có file cấu hình để mở. Hãy tạo file mẫu trước.")
                return
            try:
                if sys.platform == "win32":
                    os.startfile(path)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", path])
                else:
                    subprocess.Popen(["xdg-open", path])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được file: {e}")

        def _collect_settings_from_ui(self):
            """Đọc các giá trị đang chọn trong Tab Cài đặt, cập nhật vào self.app_config."""
            cfg = self.app_config
            cfg.ocr_mode = self.set_var_ocr_mode.get()
            cfg.ocr_language = self.set_var_ocr_lang.get().strip() or "vie+eng"
            cfg.default_file_name_template = self.set_var_name_tmpl.get().strip()
            cfg.auto_backup = self.set_var_auto_backup.get()
            cfg.require_dry_run = self.set_var_require_dry_run.get()
            cfg.word_pdf_safe_mode = self.set_var_word_safe.get()
            cfg.update_check_enabled = self.set_var_update_check.get()
            try:
                cfg.nguong_tin_cay_cao = float(self.set_var_nguong_cao.get().strip())
            except ValueError:
                cfg.nguong_tin_cay_cao = 90.0
            try:
                cfg.nguong_tin_cay_thap = float(self.set_var_nguong_thap.get().strip())
            except ValueError:
                cfg.nguong_tin_cay_thap = 70.0
            cfg.che_do_nang_cao = self.set_var_nang_cao.get()
            return cfg

        def save_settings_clicked(self):
            cfg = self._collect_settings_from_ui()
            if save_app_config_safe(cfg):
                self.lbl_settings_status.config(text="✓ Đã lưu cài đặt.", fg="#2e7d32")
                self.log("✓ Đã lưu cài đặt vào data/app_settings.json")
            else:
                self.lbl_settings_status.config(text="⚠ Không lưu được (module cấu hình chưa sẵn sàng).", fg="#c62828")

        def reset_settings_clicked(self):
            if not messagebox.askyesno("Khôi phục mặc định",
                                        "Đưa toàn bộ cài đặt về giá trị mặc định ban đầu?"):
                return
            if _CONFIG_LOADER_MODULE_OK:
                self.app_config = _AppConfig()
            else:
                self.app_config = load_app_config_safe()
            # Cập nhật lại các ô trên giao diện
            cfg = self.app_config
            self.set_var_ocr_mode.set(cfg.ocr_mode)
            self.set_var_ocr_lang.set(cfg.ocr_language)
            self.set_var_name_tmpl.set(cfg.default_file_name_template)
            self.set_var_auto_backup.set(cfg.auto_backup)
            self.set_var_require_dry_run.set(cfg.require_dry_run)
            self.set_var_word_safe.set(cfg.word_pdf_safe_mode)
            self.set_var_update_check.set(cfg.update_check_enabled)
            self.set_var_nguong_cao.set(str(getattr(cfg, "nguong_tin_cay_cao", 90.0)))
            self.set_var_nguong_thap.set(str(getattr(cfg, "nguong_tin_cay_thap", 70.0)))
            self.set_var_nang_cao.set(getattr(cfg, "che_do_nang_cao", False))
            self._apply_advanced_mode_visibility(self.set_var_nang_cao.get())
            save_app_config_safe(cfg)
            self.lbl_settings_status.config(text="✓ Đã khôi phục mặc định.", fg="#2e7d32")

        def open_settings_folder(self):
            data_dir = os.path.join(get_app_dir(), "data")
            os.makedirs(data_dir, exist_ok=True)
            try:
                if sys.platform == "win32":
                    os.startfile(data_dir)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", data_dir])
                else:
                    subprocess.Popen(["xdg-open", data_dir])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được thư mục: {e}")


        # ========================================================================
        # CÀI ĐẶT PYTHON VÀ OCR/TESSERACT - đúng tài liệu "LỆNH BỔ SUNG PHẦN CÀI ĐẶT
        # VÀ QUẢN LÝ BẢN QUYỀN PHẦN MỀM SỸ LAND" (Phần I)
        # ========================================================================
        def _build_license_section(self, parent, pad):
            """
            Mục "🔑 Đăng ký & Kích hoạt bản quyền" - đúng mục X-XVII tài liệu "LỆNH BỔ SUNG PHẦN CÀI
            ĐẶT VÀ QUẢN LÝ BẢN QUYỀN PHẦN MỀM SỸ LAND". Kích hoạt NGOẠI TUYẾN bằng chữ ký số Ed25519 -
            không cần máy chủ Internet, người dùng trao đổi mã qua Zalo/điện thoại với Nguyễn Sỹ.
            """
            from app.services import license_service as _ls
            frm = styled_labelframe(parent, text="🔑 Đăng ký & Kích hoạt bản quyền")
            frm.pack(fill="x", padx=10, pady=(0, 6))

            # --- Trạng thái hiện tại ---
            self.lbl_license_status = tk.Label(frm, text="", justify="left", anchor="w", font=("Arial", 9))
            self.lbl_license_status.pack(fill="x", padx=8, pady=(8, 8))

            # --- Mã thiết bị (luôn hiển thị - cần để gửi kèm khi xin cấp key) ---
            row_device = tk.Frame(frm); row_device.pack(fill="x", padx=8, pady=(0, 8))
            tk.Label(row_device, text="Mã thiết bị của máy này:", font=("Arial", 9, "bold")).pack(side="left")
            self.var_license_device_id = tk.StringVar(value=_ls.get_device_id())
            tk.Entry(row_device, textvariable=self.var_license_device_id, width=22, state="readonly",
                    font=("Consolas", 9)).pack(side="left", padx=(6, 4))
            tk.Button(row_device, text="📋 Sao chép", command=lambda: self._sao_chep_clipboard(
                self.var_license_device_id.get())).pack(side="left")

            # --- Thông tin đăng ký (lưu cục bộ trên máy này) ---
            frm_dangky = tk.LabelFrame(frm, text="Thông tin đăng ký (lưu trên máy này)", font=("Arial", 9, "bold"))
            frm_dangky.pack(fill="x", padx=8, pady=(0, 8))
            cfg_lic = _ls.load_local_token(self._duong_dan_ho_so_dang_ky()) or {}
            self.var_license_hoten = tk.StringVar(value=cfg_lic.get("ho_ten", ""))
            self.var_license_dienthoai = tk.StringVar(value=cfg_lic.get("dien_thoai", ""))
            self.var_license_donvi = tk.StringVar(value=cfg_lic.get("don_vi", ""))
            row_dk1 = tk.Frame(frm_dangky); row_dk1.pack(fill="x", padx=6, pady=(6, 3))
            tk.Label(row_dk1, text="Họ và tên:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_dk1, textvariable=self.var_license_hoten, width=30).pack(side="left", padx=(0, 16))
            tk.Label(row_dk1, text="Số điện thoại:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_dk1, textvariable=self.var_license_dienthoai, width=16).pack(side="left")
            row_dk2 = tk.Frame(frm_dangky); row_dk2.pack(fill="x", padx=6, pady=(0, 6))
            tk.Label(row_dk2, text="Đơn vị công tác:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_dk2, textvariable=self.var_license_donvi, width=46).pack(side="left")

            # --- Bước 1: tạo mã yêu cầu kích hoạt ---
            frm_b1 = tk.LabelFrame(frm, text="Bước 1: Tạo mã yêu cầu kích hoạt (gửi cho Nguyễn Sỹ)",
                                  font=("Arial", 9, "bold"))
            frm_b1.pack(fill="x", padx=8, pady=(0, 8))
            row_key = tk.Frame(frm_b1); row_key.pack(fill="x", padx=6, pady=(6, 4))
            tk.Label(row_key, text="Mã key đã nhận (nếu có):").pack(side="left")
            self.var_license_key_nhap = tk.StringVar()
            tk.Entry(row_key, textvariable=self.var_license_key_nhap, width=24).pack(side="left", padx=(6, 10))
            tk.Button(row_key, text="① Tạo mã yêu cầu", command=self.tao_ma_yeu_cau_kich_hoat_clicked
                      ).pack(side="left")
            self.txt_license_request = tk.Text(frm_b1, height=3, font=("Consolas", 8), wrap="char")
            self.txt_license_request.pack(fill="x", padx=6, pady=(0, 4))
            tk.Button(frm_b1, text="📋 Sao chép mã yêu cầu", command=self.sao_chep_ma_yeu_cau_clicked
                      ).pack(anchor="w", padx=6, pady=(0, 6))

            # --- Bước 2: nhập mã phản hồi ---
            frm_b2 = tk.LabelFrame(frm, text="Bước 2: Dán mã phản hồi Nguyễn Sỹ gửi lại, rồi Kích hoạt",
                                  font=("Arial", 9, "bold"))
            frm_b2.pack(fill="x", padx=8, pady=(0, 8))
            self.txt_license_response = tk.Text(frm_b2, height=3, font=("Consolas", 8), wrap="char")
            self.txt_license_response.pack(fill="x", padx=6, pady=(6, 4))
            row_b2_btn = tk.Frame(frm_b2); row_b2_btn.pack(fill="x", padx=6, pady=(0, 6))
            tk.Button(row_b2_btn, text="② Kích hoạt", bg="#2e7d32", fg="white",
                      command=self.kich_hoat_ban_quyen_clicked).pack(side="left")
            tk.Button(row_b2_btn, text="🗑 Hủy kích hoạt thiết bị này", command=self.huy_kich_hoat_clicked
                      ).pack(side="left", padx=(10, 0))

            tk.Label(frm, text="Liên hệ cấp key: Nguyễn Sỹ - SĐT/Zalo: 0972560335 - Email: minhsybk@gmail.com",
                    fg="#555").pack(anchor="w", padx=8, pady=(0, 8))

            self._cap_nhat_hien_thi_trang_thai_ban_quyen()

        def _duong_dan_ho_so_dang_ky(self):
            return os.path.join(get_app_data_dir(), "license_ho_so_dang_ky.dat")

        def _sao_chep_clipboard(self, text):
            self.clipboard_clear()
            self.clipboard_append(text)

        def _cap_nhat_hien_thi_trang_thai_ban_quyen(self):
            trang_thai, data = self._lay_trang_thai_ban_quyen()
            from app.services import license_service as _ls
            mau = {"DA_KICH_HOAT": "#2e7d32", "DANG_DUNG_THU": "#f57f17", "SAP_HET_HAN": "#f57f17",
                  "CHUA_KICH_HOAT": "#c62828", "HET_HAN": "#c62828", "BI_KHOA": "#c62828"}.get(trang_thai, "#777")
            if data:
                so_ngay = _ls.so_ngay_con_lai(data)
                dong_ngay = f" - còn {so_ngay} ngày" if so_ngay is not None else ""
                text = (f"Trạng thái: {trang_thai}{dong_ngay}\n"
                       f"Key: {_ls.mask_license_key(data.get('license_code', ''))}   "
                       f"Gói: {data.get('plan', '')}   Hết hạn: {data.get('expiry_date', 'Không giới hạn')}")
            else:
                text = "Trạng thái: CHƯA KÍCH HOẠT - vui lòng thực hiện Bước 1 và Bước 2 bên dưới."
            self.lbl_license_status.config(text=text, fg=mau)

        def tao_ma_yeu_cau_kich_hoat_clicked(self):
            from app.services import license_service as _ls
            ho_ten = self.var_license_hoten.get().strip()
            if not ho_ten:
                messagebox.showwarning("Thiếu thông tin", "Vui lòng nhập Họ và tên trước khi tạo mã yêu cầu.")
                return
            _ls.save_local_token({"ho_ten": ho_ten, "dien_thoai": self.var_license_dienthoai.get().strip(),
                                  "don_vi": self.var_license_donvi.get().strip()},
                                 self._duong_dan_ho_so_dang_ky())
            ma_yc = _ls.create_activation_request(ho_ten, self.var_license_key_nhap.get().strip())
            self.txt_license_request.delete("1.0", "end")
            self.txt_license_request.insert("1.0", ma_yc)

        def sao_chep_ma_yeu_cau_clicked(self):
            noi_dung = self.txt_license_request.get("1.0", "end").strip()
            if noi_dung:
                self._sao_chep_clipboard(noi_dung)
                messagebox.showinfo("Đã sao chép", "Đã sao chép mã yêu cầu - gửi cho Nguyễn Sỹ qua Zalo/tin nhắn.")

        def kich_hoat_ban_quyen_clicked(self):
            from app.services import license_service as _ls
            ma_phan_hoi = self.txt_license_response.get("1.0", "end").strip()
            if not ma_phan_hoi:
                messagebox.showwarning("Thiếu mã phản hồi", "Vui lòng dán mã phản hồi Nguyễn Sỹ đã gửi.")
                return
            ok, ket_qua = _ls.verify_activation_response(ma_phan_hoi)
            if not ok:
                messagebox.showerror("Kích hoạt thất bại", str(ket_qua))
                return
            _ls.save_local_token(ket_qua, self._duong_dan_token_ban_quyen())
            self._cap_nhat_hien_thi_trang_thai_ban_quyen()
            self._ap_dung_gioi_han_ban_quyen()
            messagebox.showinfo("Kích hoạt thành công", "Đã kích hoạt bản quyền thành công! Toàn bộ chức năng đã được mở.")

        def huy_kich_hoat_clicked(self):
            if not messagebox.askyesno("Xác nhận", "Hủy kích hoạt bản quyền trên MÁY NÀY? Bạn sẽ cần "
                                       "kích hoạt lại (xin mã phản hồi mới) để dùng tiếp."):
                return
            from app.services import license_service as _ls
            _ls.xoa_local_token(self._duong_dan_token_ban_quyen())
            self._cap_nhat_hien_thi_trang_thai_ban_quyen()
            self._ap_dung_gioi_han_ban_quyen()

        def _build_env_setup_section(self, parent, pad):
            from app.services import env_setup_service as _env

            frm = styled_labelframe(parent, text="🐍 Cài đặt Python và OCR/Tesseract")
            frm.pack(fill="x", **pad)
            tk.Label(frm, text="Kiểm tra và khắc phục khi phần mềm báo thiếu Python hoặc Tesseract OCR - "
                              "không cần biết dòng lệnh. Dùng khi cài tự động (đi kèm bộ cài .exe) không "
                              "hoạt động, hoặc muốn kiểm tra/đổi sang bản Python/Tesseract khác.",
                    fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 8))

            env_cfg = _env.load_env_config(os.path.join(get_app_dir(), "data", "app_settings.json"))
            self.var_env_python_path = tk.StringVar(value=env_cfg.get("python_path", ""))
            self.var_env_tesseract_path = tk.StringVar(value=env_cfg.get("tesseract_path", ""))
            self.var_env_tessdata_path = tk.StringVar(value=env_cfg.get("tessdata_path", ""))

            # --------- Khối trạng thái Python ---------
            frm_py = tk.LabelFrame(frm, text="Python", font=("Arial", 9, "bold"))
            frm_py.pack(fill="x", padx=8, pady=(0, 8))
            self.lbl_env_python_status = tk.Label(frm_py, text="Chưa kiểm tra", fg="#777", justify="left",
                                                  anchor="w", wraplength=960)
            self.lbl_env_python_status.pack(fill="x", padx=8, pady=(6, 4))

            row_py_path = tk.Frame(frm_py); row_py_path.pack(fill="x", padx=8, pady=(0, 4))
            tk.Label(row_py_path, text="Đường dẫn python.exe:").pack(side="left")
            tk.Entry(row_py_path, textvariable=self.var_env_python_path, width=60).pack(side="left", padx=(4, 4))
            tk.Button(row_py_path, text="📂 Chọn thủ công...",
                      command=self.browse_python_exe_clicked).pack(side="left", padx=(0, 4))

            row_py_btn = tk.Frame(frm_py); row_py_btn.pack(fill="x", padx=8, pady=(0, 8))
            tk.Button(row_py_btn, text="🔍 Kiểm tra Python", command=self.refresh_python_status_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_py_btn, text="🔎 Tự động tìm Python", command=self.auto_find_python_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_py_btn, text="➕ Thêm vào User PATH", command=self.add_python_to_path_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_py_btn, text="🌐 Mở trang tải Python",
                      command=lambda: self._mo_trang_web("https://www.python.org/downloads/")).pack(side="left", padx=(0, 4))

            # --------- Khối trạng thái Tesseract ---------
            frm_tess = tk.LabelFrame(frm, text="Tesseract OCR", font=("Arial", 9, "bold"))
            frm_tess.pack(fill="x", padx=8, pady=(0, 8))
            self.lbl_env_tesseract_status = tk.Label(frm_tess, text="Chưa kiểm tra", fg="#777", justify="left",
                                                     anchor="w", wraplength=960)
            self.lbl_env_tesseract_status.pack(fill="x", padx=8, pady=(6, 4))

            row_tess_path = tk.Frame(frm_tess); row_tess_path.pack(fill="x", padx=8, pady=(0, 4))
            tk.Label(row_tess_path, text="Đường dẫn tesseract.exe:").pack(side="left")
            tk.Entry(row_tess_path, textvariable=self.var_env_tesseract_path, width=50).pack(side="left", padx=(4, 4))
            tk.Button(row_tess_path, text="📂 Chọn thủ công...",
                      command=self.browse_tesseract_exe_clicked).pack(side="left")

            row_tessdata_path = tk.Frame(frm_tess); row_tessdata_path.pack(fill="x", padx=8, pady=(0, 4))
            tk.Label(row_tessdata_path, text="Thư mục tessdata:      ").pack(side="left")
            tk.Entry(row_tessdata_path, textvariable=self.var_env_tessdata_path, width=50).pack(side="left", padx=(4, 4))
            tk.Button(row_tessdata_path, text="📂 Chọn thư mục...",
                      command=self.browse_tessdata_dir_clicked).pack(side="left")

            row_tess_btn = tk.Frame(frm_tess); row_tess_btn.pack(fill="x", padx=8, pady=(0, 4))
            tk.Button(row_tess_btn, text="🔍 Kiểm tra Tesseract", command=self.refresh_tesseract_status_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_tess_btn, text="🔎 Tự động tìm Tesseract", command=self.auto_find_tesseract_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_tess_btn, text="📦 Dùng bản đi kèm phần mềm",
                      command=self.use_bundled_tesseract_clicked).pack(side="left", padx=(0, 4))

            row_tess_btn2 = tk.Frame(frm_tess); row_tess_btn2.pack(fill="x", padx=8, pady=(0, 8))
            tk.Button(row_tess_btn2, text="➕ Thêm vào User PATH", command=self.add_tesseract_to_path_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_tess_btn2, text="🌐 Mở trang tải Tesseract",
                      command=lambda: self._mo_trang_web(
                          "https://github.com/UB-Mannheim/tesseract/wiki")).pack(side="left", padx=(0, 4))
            tk.Button(row_tess_btn2, text="📁 Mở thư mục tessdata", command=self.open_tessdata_folder_clicked).pack(side="left")

            # --------- Khôi phục PATH (mục VII tài liệu) ---------
            row_path_restore = tk.Frame(frm); row_path_restore.pack(fill="x", padx=8, pady=(0, 8))
            tk.Button(row_path_restore, text="🔄 Khôi phục PATH từ bản sao lưu...",
                      command=self.restore_path_clicked).pack(side="left")
            tk.Label(row_path_restore, text="(dùng khi việc \"Thêm vào PATH\" ở trên gây ra sự cố - "
                    "mỗi lần thêm PATH đều tự động sao lưu lại trước)",
                    fg="#888").pack(side="left", padx=6)

            # --------- Kiểm tra thư viện Python ---------
            frm_lib = tk.LabelFrame(frm, text="Thư viện Python cần thiết", font=("Arial", 9, "bold"))
            frm_lib.pack(fill="x", padx=8, pady=(0, 8))
            row_lib_btn = tk.Frame(frm_lib); row_lib_btn.pack(fill="x", padx=8, pady=(6, 4))
            tk.Button(row_lib_btn, text="🔍 Kiểm tra thư viện", command=self.check_libraries_clicked).pack(side="left", padx=(0, 4))
            tk.Button(row_lib_btn, text="⬇ Cài thư viện còn thiếu", command=self.install_missing_libraries_clicked).pack(side="left")

            cols_lib = ("ten", "trang_thai", "phien_ban")
            self.tree_env_libs = ttk.Treeview(frm_lib, columns=cols_lib, show="headings", height=6)
            for c, h, w in zip(cols_lib, ["Thư viện", "Trạng thái", "Phiên bản"], [180, 120, 140]):
                self.tree_env_libs.heading(c, text=h)
                self.tree_env_libs.column(c, width=w, anchor="w")
            self.tree_env_libs.pack(fill="x", padx=8, pady=(0, 8))

            # --------- Chạy thử OCR ---------
            frm_ocr_test = tk.LabelFrame(frm, text="Chạy thử OCR", font=("Arial", 9, "bold"))
            frm_ocr_test.pack(fill="x", padx=8, pady=(0, 8))
            self._register_advanced_widget(frm_ocr_test, fill="x", padx=8, pady=(0, 8))
            row_ocr_test = tk.Frame(frm_ocr_test); row_ocr_test.pack(fill="x", padx=8, pady=6)
            self.var_env_ocr_test_lang = tk.StringVar(value="vie")
            tk.Label(row_ocr_test, text="Ngôn ngữ:").pack(side="left")
            ttk.Combobox(row_ocr_test, textvariable=self.var_env_ocr_test_lang, values=["vie", "eng", "vie+eng"],
                        width=8, state="readonly").pack(side="left", padx=(4, 10))
            tk.Button(row_ocr_test, text="🖼 Chọn ảnh & chạy thử OCR",
                      command=self.run_ocr_test_clicked).pack(side="left")
            self.txt_env_ocr_result = tk.Text(frm_ocr_test, height=5, wrap="word")
            self.txt_env_ocr_result.pack(fill="x", padx=8, pady=(0, 8))

            # Kiem tra trang thai ngay khi mo Tab (khong cho nguoi dung phai bam nut truoc)
            self.refresh_python_status_clicked()
            self.refresh_tesseract_status_clicked()

        def _build_word_perf_section(self, parent, pad):
            cfg = load_app_config_safe()
            frm = styled_labelframe(parent, text="⚡ Hiệu năng xuất Word hàng loạt (mục 7d)")
            frm.pack(fill="x", **pad)
            self._register_advanced_widget(frm, fill="x", **pad)
            tk.Label(frm, text="Khi xuất HÀNG LOẠT nhiều file Word (mục 7d, dùng Microsoft Word thật), "
                              "phần mềm sẽ chủ động NGHỈ NGẮN giữa các file để bạn vẫn mở được thư mục/"
                              "Excel/PDF và làm việc khác bình thường trong lúc đang xuất - không chiếm "
                              "hết CPU/ổ đĩa khiến Windows bị lag.",
                    fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 8))

            row_tocdo = tk.Frame(frm); row_tocdo.pack(fill="x", padx=8, pady=(0, 6))
            tk.Label(row_tocdo, text="Tốc độ xử lý:", font=("Arial", 9, "bold")).pack(side="left")
            self.var_word_toc_do = tk.StringVar(value=getattr(cfg, "word_toc_do_xu_ly", "can_bang"))
            for gia_tri, nhan in [("nhanh", "Nhanh (CPU ~70-80%)"), ("can_bang", "Cân bằng (mặc định, CPU ~40-60%)"),
                                  ("tiet_kiem", "Tiết kiệm tài nguyên (CPU ~20-30%)")]:
                tk.Radiobutton(row_tocdo, text=nhan, variable=self.var_word_toc_do,
                              value=gia_tri).pack(side="left", padx=(10, 0))

            row_dot = tk.Frame(frm); row_dot.pack(fill="x", padx=8, pady=(0, 6))
            tk.Label(row_dot, text="Số file mỗi đợt:").pack(side="left")
            self.var_word_so_file_dot = tk.StringVar(value=str(getattr(cfg, "word_so_file_moi_dot", 100)))
            ttk.Combobox(row_dot, textvariable=self.var_word_so_file_dot, width=8, state="readonly",
                        values=["50", "100", "200", "500", "1000"]).pack(side="left", padx=(4, 16))
            tk.Label(row_dot, text="Nghỉ giữa các đợt (giây):").pack(side="left")
            self.var_word_nghi_dot = tk.StringVar(value=str(getattr(cfg, "word_nghi_giua_dot_giay", 2.0)))
            tk.Entry(row_dot, textvariable=self.var_word_nghi_dot, width=6).pack(side="left", padx=(4, 16))
            tk.Label(row_dot, text="Giới hạn RAM tối đa (MB, 0=không giới hạn):").pack(side="left")
            self.var_word_gioi_han_ram = tk.StringVar(value=str(getattr(cfg, "word_gioi_han_ram_mb", 0)))
            tk.Entry(row_dot, textvariable=self.var_word_gioi_han_ram, width=8).pack(side="left", padx=(4, 0))

            row_check = tk.Frame(frm); row_check.pack(fill="x", padx=8, pady=(0, 8))
            self.var_word_hien_thong_ke = tk.BooleanVar(value=getattr(cfg, "word_hien_thi_thong_ke_hieu_nang", True))
            tk.Checkbutton(row_check, text="Ghi log hiệu năng (Logs/LOG_HIEU_NANG_WORD.csv - CPU/RAM/thời "
                          "gian mỗi file)", variable=self.var_word_hien_thong_ke).pack(anchor="w")
            self.var_word_tu_khoi_dong = tk.BooleanVar(value=getattr(cfg, "word_tu_khoi_dong_lai_khi_loi", True))
            tk.Checkbutton(row_check, text="Tự động chuyển sang phương án dự phòng nếu Word gặp lỗi giữa "
                          "chừng (không dừng cả đợt xuất)",
                          variable=self.var_word_tu_khoi_dong).pack(anchor="w")

            tk.Button(frm, text="💾 Lưu cấu hình hiệu năng", command=self.save_word_perf_settings_clicked
                      ).pack(anchor="w", padx=8, pady=(0, 8))

        def save_word_perf_settings_clicked(self):
            cfg = load_app_config_safe()
            cfg.word_toc_do_xu_ly = self.var_word_toc_do.get()
            try:
                cfg.word_so_file_moi_dot = int(self.var_word_so_file_dot.get())
            except ValueError:
                cfg.word_so_file_moi_dot = 100
            try:
                cfg.word_nghi_giua_dot_giay = float(self.var_word_nghi_dot.get())
            except ValueError:
                cfg.word_nghi_giua_dot_giay = 2.0
            try:
                cfg.word_gioi_han_ram_mb = int(self.var_word_gioi_han_ram.get())
            except ValueError:
                cfg.word_gioi_han_ram_mb = 0
            cfg.word_hien_thi_thong_ke_hieu_nang = self.var_word_hien_thong_ke.get()
            cfg.word_tu_khoi_dong_lai_khi_loi = self.var_word_tu_khoi_dong.get()
            save_app_config_safe(cfg)
            self.app_config = cfg
            messagebox.showinfo("Đã lưu", "Đã lưu cấu hình hiệu năng xuất Word.")

        def _build_diagnostic_section(self, parent, pad):
            frm = styled_labelframe(parent, text="🩺 Xuất gói chẩn đoán (hỗ trợ kỹ thuật từ xa)")
            frm.pack(fill="x", padx=10, pady=(0, 8))
            tk.Label(frm, text="Khi gặp sự cố, xuất 1 file duy nhất gửi cho Nguyễn Sỹ để được hỗ trợ "
                              "nhanh hơn - KHÔNG chứa hồ sơ đất đai, mật khẩu, hay key bản quyền đầy đủ, "
                              "chỉ chứa thông tin môi trường/cấu hình cần thiết để chẩn đoán.",
                    fg="#555", wraplength=980, justify="left").pack(anchor="w", padx=8, pady=(8, 8))
            tk.Button(frm, text="🩺 Xuất gói chẩn đoán...", command=self.xuat_goi_chan_doan_clicked,
                     bg="#00838f", fg="white").pack(anchor="w", padx=8, pady=(0, 8))

        def xuat_goi_chan_doan_clicked(self):
            from app.services import diagnostic_service as _diag
            duong_dan = filedialog.asksaveasfilename(
                defaultextension=".zip", filetypes=[("File ZIP", "*.zip")],
                initialfile=f"ChanDoan_SyLand_{time.strftime('%Y%m%d_%H%M%S')}.zip")
            if not duong_dan:
                return

            module_flags = {
                "tab_settings": _UI_SETTINGS_MODULE_OK, "tab_dup": _UI_DUP_MODULE_OK,
                "tab_merge_word": _UI_MERGE_WORD_MODULE_OK, "tab_split": _UI_SPLIT_MODULE_OK,
                "tab_excel": _UI_EXCEL_MODULE_OK, "tab_reconcile": _UI_RECONCILE_MODULE_OK,
                "tab_workflow": _UI_WORKFLOW_MODULE_OK, "tab_content": _UI_CONTENT_MODULE_OK,
                "tab_pdfedit": _UI_PDFEDIT_MODULE_OK, "tab_report": _UI_REPORT_MODULE_OK,
                "tab_word": _UI_WORD_MODULE_OK,
            }
            trang_thai, license_data = self._lay_trang_thai_ban_quyen() if hasattr(self, "_lay_trang_thai_ban_quyen") else (None, None)

            thanh_cong, thong_bao = _diag.tao_goi_chan_doan(
                duong_dan, APP_VERSION, self.app_config, license_data, module_flags,
                get_standard_logs_dir())
            if thanh_cong:
                messagebox.showinfo("Đã xuất gói chẩn đoán", thong_bao)
            else:
                messagebox.showerror("Lỗi", thong_bao)

        def _mo_trang_web(self, url):
            import webbrowser
            webbrowser.open(url)

        def _cap_nhat_mau_trang_thai(self, label_widget, info, ten_phan_mem):
            mau_nen = {"xanh": "#c8e6c9", "vang": "#fff9c4", "do": "#ffcdd2"}.get(info["trang_thai_mau"], "#eeeeee")
            if info["installed"]:
                dong1 = f"✓ Đã tìm thấy {ten_phan_mem} {info.get('version', '')} tại: {info.get('path', '')}"
            else:
                dong1 = f"✗ Chưa tìm thấy {ten_phan_mem} trên máy."
            chi_tiet = []
            if "scripts_path" in info:
                chi_tiet.append(f"Scripts: {info.get('scripts_path') or '(không có)'}")
                chi_tiet.append(f"pip: {'hoạt động' if info.get('pip_ok') else 'CHƯA hoạt động'}")
            if "tessdata_path" in info:
                chi_tiet.append(f"tessdata: {info.get('tessdata_path') or '(không có)'}")
                chi_tiet.append(f"vie.traineddata: {'có' if info.get('has_vie') else 'THIẾU'}")
                chi_tiet.append(f"eng.traineddata: {'có' if info.get('has_eng') else 'THIẾU'}")
            chi_tiet.append(f"Trong PATH: {'có' if info.get('in_path') else 'CHƯA'}")
            dong2 = " | ".join(chi_tiet)
            text_full = dong1 + ("\n" + info["ghi_chu"] if info.get("ghi_chu") else "") + "\n" + dong2
            label_widget.config(text=text_full, bg=mau_nen, fg="#212121")

        def refresh_python_status_clicked(self):
            from app.services import env_setup_service as _env
            info = _env.detect_python_info(configured_path=self.var_env_python_path.get().strip() or None)
            self.env_python_info = info
            if info["installed"]:
                self.var_env_python_path.set(info["path"])
            self._cap_nhat_mau_trang_thai(self.lbl_env_python_status, info, "Python")

        def refresh_tesseract_status_clicked(self):
            from app.services import env_setup_service as _env
            info = _env.detect_tesseract_info(
                configured_path=self.var_env_tesseract_path.get().strip() or None,
                configured_tessdata=self.var_env_tessdata_path.get().strip() or None)
            self.env_tesseract_info = info
            if info["installed"]:
                self.var_env_tesseract_path.set(info["path"])
            if info.get("tessdata_path"):
                self.var_env_tessdata_path.set(info["tessdata_path"])
            self._cap_nhat_mau_trang_thai(self.lbl_env_tesseract_status, info, "Tesseract OCR")

        def auto_find_python_clicked(self):
            from app.services import env_setup_service as _env
            candidates = _env.find_python_candidates()
            if not candidates:
                messagebox.showinfo("Không tìm thấy", "Không tự động tìm thấy bản Python nào trên máy.\n"
                                    "Hãy dùng nút 'Chọn thủ công...' hoặc cài Python mới.")
                return
            if len(candidates) == 1:
                self.var_env_python_path.set(candidates[0]["path"])
                self.refresh_python_status_clicked()
                return
            win = tk.Toplevel(self)
            win.title("Chọn bản Python")
            tk.Label(win, text="Tìm thấy nhiều bản Python - chọn 1 bản để dùng:",
                    font=("Arial", 9, "bold")).pack(anchor="w", padx=10, pady=(10, 6))
            for c in candidates:
                row = tk.Frame(win); row.pack(fill="x", padx=10, pady=3)
                tk.Label(row, text=f"Python {c['version']}  —  {c['path']}  ({c['source']})",
                        anchor="w").pack(side="left", fill="x", expand=True)
                def _chon(path=c["path"]):
                    self.var_env_python_path.set(path)
                    self.refresh_python_status_clicked()
                    win.destroy()
                tk.Button(row, text="Dùng bản này", command=_chon).pack(side="right")

        def auto_find_tesseract_clicked(self):
            from app.services import env_setup_service as _env
            candidates = _env.find_tesseract_candidates(saved_path=self.var_env_tesseract_path.get().strip() or None)
            if not candidates:
                messagebox.showinfo("Không tìm thấy", "Không tự động tìm thấy Tesseract OCR nào trên máy.\n"
                                    "Hãy dùng nút 'Dùng bản đi kèm phần mềm' hoặc 'Chọn thủ công...'.")
                return
            self.var_env_tesseract_path.set(candidates[0]["path"])
            self.refresh_tesseract_status_clicked()
            messagebox.showinfo("Đã tìm thấy", f"Đã chọn: {candidates[0]['path']}\n(nguồn: {candidates[0]['source']})")

        def use_bundled_tesseract_clicked(self):
            exe_name = "tesseract.exe" if sys.platform == "win32" else "tesseract"
            bundled = os.path.join(get_app_dir(), "tools", "tesseract", exe_name)
            if not os.path.isfile(bundled):
                messagebox.showerror("Không tìm thấy", f"Không tìm thấy bản Tesseract đi kèm phần mềm tại:\n{bundled}\n\n"
                                     "Hãy dùng 'Tự động tìm' hoặc cài đặt riêng.")
                return
            self.var_env_tesseract_path.set(bundled)
            self.refresh_tesseract_status_clicked()

        def browse_python_exe_clicked(self):
            pattern = [("python.exe", "python.exe")] if sys.platform == "win32" else [("Python", "python3;python")]
            path = filedialog.askopenfilename(title="Chọn file python.exe")
            if path:
                self.var_env_python_path.set(path)
                self.refresh_python_status_clicked()

        def browse_tesseract_exe_clicked(self):
            path = filedialog.askopenfilename(title="Chọn file tesseract.exe")
            if path:
                self.var_env_tesseract_path.set(path)
                self.refresh_tesseract_status_clicked()

        def browse_tessdata_dir_clicked(self):
            path = filedialog.askdirectory(title="Chọn thư mục tessdata")
            if path:
                self.var_env_tessdata_path.set(path)
                self.refresh_tesseract_status_clicked()

        def add_python_to_path_clicked(self):
            from app.services import env_setup_service as _env
            python_path = self.var_env_python_path.get().strip()
            if not python_path or not os.path.isfile(python_path):
                messagebox.showwarning("Chưa có Python", "Hãy kiểm tra/chọn đường dẫn Python hợp lệ trước.")
                return
            python_dir = os.path.dirname(python_path)
            scripts_dir = _env.find_python_scripts_dir(python_path)
            ket_qua_1 = _env.add_to_path_safe(python_dir, scope="user")
            thong_bao = ket_qua_1["thong_bao"]
            if scripts_dir:
                ket_qua_2 = _env.add_to_path_safe(scripts_dir, scope="user")
                thong_bao += "\n" + ket_qua_2["thong_bao"]
            messagebox.showinfo("Kết quả", thong_bao)
            self.refresh_python_status_clicked()

        def add_tesseract_to_path_clicked(self):
            from app.services import env_setup_service as _env
            tess_path = self.var_env_tesseract_path.get().strip()
            if not tess_path or not os.path.isfile(tess_path):
                messagebox.showwarning("Chưa có Tesseract", "Hãy kiểm tra/chọn đường dẫn Tesseract hợp lệ trước.")
                return
            ket_qua = _env.add_to_path_safe(os.path.dirname(tess_path), scope="user")
            messagebox.showinfo("Kết quả", ket_qua["thong_bao"])
            self.refresh_tesseract_status_clicked()

        def open_tessdata_folder_clicked(self):
            path = self.var_env_tessdata_path.get().strip()
            if not path or not os.path.isdir(path):
                messagebox.showwarning("Chưa có thư mục", "Chưa xác định được thư mục tessdata hợp lệ.")
                return
            try:
                if sys.platform == "win32":
                    os.startfile(path)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", path])
                else:
                    subprocess.Popen(["xdg-open", path])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được thư mục: {e}")

        def restore_path_clicked(self):
            from app.services import env_setup_service as _env
            backups = _env.list_path_backups()
            if not backups:
                messagebox.showinfo("Chưa có bản sao lưu",
                                    "Chưa có bản sao lưu PATH nào (chỉ tạo ra khi bạn dùng nút "
                                    "\"Thêm vào User PATH\" ở trên ít nhất 1 lần).")
                return

            win = tk.Toplevel(self)
            win.title("Khôi phục PATH từ bản sao lưu")
            win.geometry("640x360")
            tk.Label(win, text="Chọn 1 bản sao lưu PATH để khôi phục lại (mới nhất ở trên cùng):",
                    fg="#555").pack(anchor="w", padx=10, pady=(10, 4))

            listbox = tk.Listbox(win, font=("Consolas", 9))
            for b in backups:
                listbox.insert("end", os.path.basename(b))
            listbox.pack(fill="both", expand=True, padx=10, pady=(0, 8))
            listbox.selection_set(0)

            def _thuc_hien_khoi_phuc():
                sel = listbox.curselection()
                if not sel:
                    return
                backup_path = backups[sel[0]]
                if not messagebox.askyesno("Xác nhận khôi phục",
                                           f"Khôi phục PATH từ:\n{os.path.basename(backup_path)}\n\n"
                                           "PATH hiện tại sẽ bị THAY THẾ bằng nội dung trong bản sao lưu này. "
                                           "Bạn có chắc chắn không?"):
                    return
                ket_qua = _env.restore_path_backup(backup_path)
                messagebox.showinfo("Kết quả", ket_qua["thong_bao"])
                if ket_qua["thanh_cong"]:
                    win.destroy()
                    self.refresh_python_status_clicked()
                    self.refresh_tesseract_status_clicked()

            row_btn = tk.Frame(win); row_btn.pack(fill="x", padx=10, pady=(0, 10))
            tk.Button(row_btn, text="✓ Khôi phục bản đã chọn", bg="#c62828", fg="white",
                      command=_thuc_hien_khoi_phuc).pack(side="left")
            tk.Button(row_btn, text="Hủy", command=win.destroy).pack(side="left", padx=(8, 0))

        def check_libraries_clicked(self):
            from app.services import env_setup_service as _env
            python_path = self.var_env_python_path.get().strip()
            if not python_path or not os.path.isfile(python_path):
                messagebox.showwarning("Chưa có Python", "Hãy kiểm tra/chọn đường dẫn Python hợp lệ trước.")
                return
            results = _env.check_required_libraries(python_path)
            self._env_lib_results = results
            for item in self.tree_env_libs.get_children():
                self.tree_env_libs.delete(item)
            for r in results:
                trang_thai = "✓ Đã cài" if r["da_cai"] else "✗ Chưa cài"
                self._insert_row_colored(self.tree_env_libs, (r["ten"], trang_thai, r["phien_ban"]),
                                         "THANH_CONG" if r["da_cai"] else "LOI")

        def install_missing_libraries_clicked(self):
            from app.services import env_setup_service as _env
            python_path = self.var_env_python_path.get().strip()
            if not python_path or not os.path.isfile(python_path):
                messagebox.showwarning("Chưa có Python", "Hãy kiểm tra/chọn đường dẫn Python hợp lệ trước.")
                return
            results = getattr(self, "_env_lib_results", None) or _env.check_required_libraries(python_path)
            thieu = [r["ten"] for r in results if not r["da_cai"]]
            if not thieu:
                messagebox.showinfo("Đã đủ", "Tất cả thư viện cần thiết đã được cài đặt.")
                return
            if not messagebox.askyesno("Xác nhận cài đặt",
                                       f"Sẽ cài {len(thieu)} thư viện còn thiếu: {', '.join(thieu)}\n\n"
                                       "Cần kết nối Internet. Tiếp tục?"):
                return
            thanh_cong, thong_bao = _env.install_missing_libraries(python_path, thieu, log_cb=None)
            messagebox.showinfo("Kết quả cài đặt", thong_bao)
            self.check_libraries_clicked()

        def run_ocr_test_clicked(self):
            from app.services import env_setup_service as _env
            tess_path = self.var_env_tesseract_path.get().strip()
            if not tess_path or not os.path.isfile(tess_path):
                messagebox.showwarning("Chưa có Tesseract", "Hãy kiểm tra/chọn đường dẫn Tesseract hợp lệ trước.")
                return
            image_path = filedialog.askopenfilename(
                title="Chọn ảnh để chạy thử OCR",
                filetypes=[("Ảnh", "*.png;*.jpg;*.jpeg;*.bmp;*.tiff"), ("Tất cả file", "*.*")])
            if not image_path:
                return
            self.txt_env_ocr_result.delete("1.0", "end")
            self.txt_env_ocr_result.insert("end", "Đang xử lý...")
            self.update()
            ket_qua = _env.run_ocr_test(image_path, tess_path, lang=self.var_env_ocr_test_lang.get(),
                                        tessdata_dir=self.var_env_tessdata_path.get().strip() or None)
            self.txt_env_ocr_result.delete("1.0", "end")
            if ket_qua["thanh_cong"]:
                self.txt_env_ocr_result.insert("end",
                    f"✓ Thành công ({ket_qua['thoi_gian_giay']}s) - Tesseract: {ket_qua['tesseract_dang_dung']}\n"
                    f"{'-'*60}\n{ket_qua['text']}")
            else:
                self.txt_env_ocr_result.insert("end", f"✗ Lỗi: {ket_qua['loi']}")


try:
    from app.ui.tab_dup import DupTabMixin
    _UI_DUP_MODULE_OK = True
except Exception:
    _UI_DUP_MODULE_OK = False

    class DupTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_dup.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_dup(self, parent, pad):
            frm = styled_labelframe(parent, text="NHÓM A — Lọc file PDF trùng giữa 2 thư mục (theo Mã xã-Tờ-Thửa)")
            frm.pack(fill="x", **pad)
            tk.Label(frm,
                     text="So sánh tên file PDF giữa 2 thư mục theo Mã xã + Số tờ + Số thửa tách được từ TÊN FILE "
                          "(bỏ qua phần đuôi/hậu tố khác nhau ở cuối, VD _GT, _TBXN, hay bất kỳ chữ gì khác).\n"
                          "File nào trong 'Thư mục cần so sánh' có Mã xã+Tờ+Thửa TRÙNG với 1 file đã có sẵn trong "
                          "'Thư mục tham chiếu' sẽ được TỰ ĐỘNG DI CHUYỂN sang thư mục thứ 3 bạn chọn.\n"
                          "Yêu cầu: tên file phải theo đúng chuẩn có chứa Mã xã (5 số) + Tờ + Thửa, VD "
                          "CHUACOGIAY_02140_111_13_GT.pdf (dùng Tab 5 để đổi tên trước nếu file chưa đúng chuẩn).",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 8))

            row1 = tk.Frame(frm)
            row1.pack(fill="x", padx=6, pady=3)
            tk.Label(row1, text="Thư mục THAM CHIẾU (đã có sẵn, đặt tên đúng chuẩn):", width=42, anchor="w").pack(side="left")
            tk.Entry(row1, textvariable=self.var_dup_reference).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row1, text="Chọn...", command=self.pick_dup_reference).pack(side="left")

            row2 = tk.Frame(frm)
            row2.pack(fill="x", padx=6, pady=3)
            tk.Label(row2, text="Thư mục CẦN SO SÁNH (di chuyển file trùng ra khỏi đây):", width=42, anchor="w").pack(side="left")
            tk.Entry(row2, textvariable=self.var_dup_compare).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row2, text="Chọn...", command=self.pick_dup_compare).pack(side="left")

            row3 = tk.Frame(frm)
            row3.pack(fill="x", padx=6, pady=3)
            tk.Label(row3, text="Thư mục LƯU file trùng đã di chuyển (mới):", width=42, anchor="w").pack(side="left")
            tk.Entry(row3, textvariable=self.var_dup_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row3, text="Chọn...", command=self.pick_dup_output).pack(side="left")

            row_run = tk.Frame(frm)
            row_run.pack(fill="x", padx=6, pady=(6, 8))
            self.btn_dup = tk.Button(row_run, text="🔍 LỌC & DI CHUYỂN FILE TRÙNG", font=("Arial", 11, "bold"),
                                      bg="#bf360c", fg="white", command=self.run_dup_clicked)
            self.btn_dup.pack(side="left")
            self.btn_dup_pause, self.btn_dup_cancel = self._add_pause_cancel(row_run, "control_dup")
            self.btn_dup_pause.pack(side="left", padx=(8, 3))
            self.btn_dup_cancel.pack(side="left", padx=3)
            tk.Label(row_run,
                     text="⚠️ Đây là DI CHUYỂN (move) thật sự khỏi thư mục so sánh, không phải copy — hãy sao lưu trước nếu cần.",
                     fg="#b71c1c", wraplength=600, justify="left").pack(side="left", padx=10)

        def pick_dup_reference(self):
            p = filedialog.askdirectory(title="Chọn thư mục THAM CHIẾU (đã có sẵn, đặt tên đúng chuẩn)")
            if p:
                self.var_dup_reference.set(p)

        def pick_dup_compare(self):
            p = filedialog.askdirectory(title="Chọn thư mục CẦN SO SÁNH")
            if p:
                self.var_dup_compare.set(p)

        def pick_dup_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục LƯU file trùng đã di chuyển")
            if p:
                self.var_dup_output.set(p)

        def run_dup_clicked(self):
            folder_reference = self.var_dup_reference.get().strip()
            folder_compare = self.var_dup_compare.get().strip()
            output_folder = self.var_dup_output.get().strip()

            if not folder_reference or not os.path.isdir(folder_reference):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục THAM CHIẾU hợp lệ.")
                return
            if not folder_compare or not os.path.isdir(folder_compare):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục CẦN SO SÁNH hợp lệ.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu file trùng đã di chuyển.")
                return
            if os.path.abspath(folder_reference) == os.path.abspath(folder_compare):
                messagebox.showerror("Lỗi", "Thư mục tham chiếu và thư mục so sánh phải khác nhau.")
                return

            if not messagebox.askyesno(
                    "Xác nhận di chuyển file",
                    f"Thao tác này sẽ DI CHUYỂN (không phải copy) các file PDF trùng khỏi:\n{folder_compare}\n\n"
                    f"sang:\n{output_folder}\n\nBạn có chắc chắn muốn tiếp tục?"):
                return

            self.txt_log.delete("1.0", "end")
            self.btn_dup.config(state="disabled")
            self._enable_pause_cancel("control_dup", self.btn_dup_pause, self.btn_dup_cancel)

            def worker():
                try:
                    files_to_backup = [os.path.join(folder_compare, f) for f in os.listdir(folder_compare)
                                       if f.lower().endswith(".pdf")]
                    self._maybe_auto_backup(files_to_backup, "Di chuyển file PDF trùng (Tab 9)")
                    result = find_and_move_duplicate_pdfs(folder_reference, folder_compare, output_folder, self.log,
                                                            control=self.control_dup)
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {len(result['moved'])} file trùng đã di chuyển | "
                             f"{len(result['unmatched'])} file không trùng (giữ nguyên) | "
                             f"{len(result['unparsed_compare'])} file không tách được tên (bỏ qua)")
                    messagebox.showinfo(
                        "Hoàn tất",
                        f"Đã di chuyển: {len(result['moved'])} file\n"
                        f"Không trùng: {len(result['unmatched'])} file\n"
                        f"Không tách được tên: {len(result['unparsed_compare'])} file\n\n"
                        f"File trùng đã lưu tại:\n{output_folder}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_dup.config(state="normal")
                    self._disable_pause_cancel(self.btn_dup_pause, self.btn_dup_cancel)

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_merge_word import MergeWordTabMixin
    _UI_MERGE_WORD_MODULE_OK = True
except Exception:
    _UI_MERGE_WORD_MODULE_OK = False

    class MergeWordTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_merge_word.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_merge_word(self, parent, pad):
            frm = styled_labelframe(parent, text="7b. Gộp file Word (.docx) hàng loạt")
            frm.pack(fill="x", **pad)
            tk.Label(frm,
                     text="Gộp nhiều file .docx thành 1 file .docx tổng duy nhất (theo thứ tự tên file A→Z, "
                          "mỗi file cách nhau 1 ngắt trang). File kết quả đặt tên NGẪU NHIÊN, lưu vào chính thư mục "
                          "nguồn (nếu chọn thư mục) hoặc thư mục chứa file ĐẦU TIÊN trong danh sách (nếu chọn file cụ thể).",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 8))

            self.mergeword_picker = SourcePicker(
                parent, "Nguồn Word cần gộp (chọn 1 thư mục, HOẶC chọn 1/nhiều file .docx bất kỳ)",
                filetypes=[("Word files", "*.docx")], file_label="file Word")
            self.mergeword_picker.pack(fill="x", **pad)

            row_opt = tk.Frame(frm)
            row_opt.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_opt, text="Ngắt trang giữa các file khi gộp (khuyến nghị giữ bật)",
                           variable=self.var_merge_word_pagebreak).pack(side="left")

            tk.Label(frm,
                     text="Lưu ý: chỉ gộp file .docx (Word mới). File .doc (Word cũ) sẽ bị bỏ qua và báo rõ trong Nhật ký. "
                          "File gộp lần trước (tên bắt đầu bằng 'GOP_WORD_') sẽ tự động không bị gộp lại vào lần chạy sau.",
                     fg="#777", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(0, 6))

            row_run = tk.Frame(frm)
            row_run.pack(fill="x", padx=6, pady=(2, 3))
            self.btn_merge_word = tk.Button(row_run, text="📎 GỘP FILE WORD HÀNG LOẠT", font=("Arial", 11, "bold"),
                                             bg="#4527a0", fg="white", command=self.run_merge_word_clicked)
            self.btn_merge_word.pack(side="left")
            self.btn_merge_word_pause, self.btn_merge_word_cancel = self._add_pause_cancel(row_run, "control_merge_word")
            self.btn_merge_word_pause.pack(side="left", padx=(8, 3))
            self.btn_merge_word_cancel.pack(side="left", padx=3)

            row_merge_progress = tk.Frame(frm); row_merge_progress.pack(fill="x", padx=6, pady=(0, 8))
            self.pb_merge_word = ttk.Progressbar(row_merge_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_merge_word.pack(side="left", padx=(0, 10))
            self.lbl_merge_word_progress = tk.Label(row_merge_progress, text="", fg="#555", anchor="w")
            self.lbl_merge_word_progress.pack(side="left", fill="x", expand=True)

        def run_merge_word_clicked(self):
            source = self.mergeword_picker.get_source()

            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file Word cần gộp.")
                return
            if not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install docxcompose python-docx")
                return

            self.txt_log.delete("1.0", "end")
            self.btn_merge_word.config(state="disabled")
            self._enable_pause_cancel("control_merge_word", self.btn_merge_word_pause, self.btn_merge_word_cancel)
            self.pb_merge_word.config(value=0, maximum=100)
            self._merge_word_start_time = time.time()

            def update_merge_progress(idx, total, fname):
                elapsed = time.time() - self._merge_word_start_time
                self.pb_merge_word.config(value=idx, maximum=total)
                self.lbl_merge_word_progress.config(text=f"Đang gộp {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_merge_progress(idx, total, fname))
                append_tab_log_row(7, "7b. Gộp file Word hàng loạt", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XU_LY")

            def worker():
                try:
                    files, out_path = batch_merge_word_folder(
                        source, self.log, add_page_break=self.var_merge_word_pagebreak.get(),
                        control=self.control_merge_word, progress_cb=progress_cb)
                    self.after(0, lambda: self.lbl_merge_word_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._merge_word_start_time)}s)"))

                    try:
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "GOP_FILE_WORD", run_id=stamp)
                        for f in files:
                            std_log.add(action="MERGE", source_file=f, output_file=out_path, status="DA_GOP")
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng kết quả chính): {e}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã gộp {len(files)} file Word thành 1 file duy nhất.")
                    self.log(f"File kết quả: {out_path}")
                    messagebox.showinfo("Hoàn tất gộp Word",
                                         f"Đã gộp {len(files)} file Word.\n\nFile kết quả:\n{out_path}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_merge_word.config(state="normal")
                    self._disable_pause_cancel(self.btn_merge_word_pause, self.btn_merge_word_cancel)

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_split import SplitTabMixin
    _UI_SPLIT_MODULE_OK = True
except Exception:
    _UI_SPLIT_MODULE_OK = False

    class SplitTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_split.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_split(self, parent, pad):
            frm_split = styled_labelframe(parent, text="NHÓM: Tách / Gộp PDF — 6a. Tách trang PDF theo mẫu tùy chọn (trang 1 / trang 2 / trang 1+2 / tùy ý)")
            frm_split.pack(fill="x", **pad)

            self.split_picker = SourcePicker(
                frm_split, "Nguồn PDF cần tách (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF bất kỳ)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.split_picker.pack(fill="x", padx=6, pady=3)

            row_out = tk.Frame(frm_split)
            row_out.pack(fill="x", padx=6, pady=3)
            tk.Label(row_out, text="Thư mục lưu kết quả:", width=28, anchor="w").pack(side="left")
            tk.Entry(row_out, textvariable=self.var_split_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_out, text="Chọn...", command=self.pick_split_output).pack(side="left")

            row_preset = tk.Frame(frm_split)
            row_preset.pack(fill="x", padx=6, pady=3)
            tk.Label(row_preset, text="Chọn nhanh:", width=12, anchor="w").pack(side="left")
            tk.Button(row_preset, text="Trang 1",
                      command=lambda: self._set_split_preset("1", "_trang1")).pack(side="left", padx=3)
            tk.Button(row_preset, text="Trang 2",
                      command=lambda: self._set_split_preset("2", "_trang2")).pack(side="left", padx=3)
            tk.Button(row_preset, text="Trang 1+2",
                      command=lambda: self._set_split_preset("1,2", "_trang1_2")).pack(side="left", padx=3)

            row_custom = tk.Frame(frm_split)
            row_custom.pack(fill="x", padx=6, pady=3)
            tk.Label(row_custom, text="Trang cần lấy (VD: 1  hoặc  1,3  hoặc  2-4):", width=32, anchor="w").pack(side="left")
            tk.Entry(row_custom, textvariable=self.var_split_page_spec, width=15).pack(side="left", padx=4)
            tk.Label(row_custom, text="Hậu tố tên file:", width=14, anchor="w").pack(side="left", padx=(12, 0))
            tk.Entry(row_custom, textvariable=self.var_split_suffix, width=15).pack(side="left")

            row_run_split = tk.Frame(frm_split)
            row_run_split.pack(fill="x", padx=6, pady=(4, 6))
            tk.Button(row_run_split, text="✂ TÁCH TRANG HÀNG LOẠT", font=("Arial", 11, "bold"),
                      bg="#6a1b9a", fg="white", command=self.run_split_clicked).pack(side="left")
            self.btn_split_pause, self.btn_split_cancel = self._add_pause_cancel(row_run_split, "control_split")
            self.btn_split_pause.pack(side="left", padx=(8, 3))
            self.btn_split_cancel.pack(side="left", padx=3)

            frm_merge = styled_labelframe(parent, text="6b. Gộp nhiều file PDF thành 1 file tổng (nhiều trang)")
            frm_merge.pack(fill="both", expand=True, **pad)

            row_mbtn = tk.Frame(frm_merge)
            row_mbtn.pack(fill="x", padx=6, pady=3)
            tk.Button(row_mbtn, text="+ Thêm file PDF...", command=self.merge_add_files).pack(side="left", padx=2)
            tk.Button(row_mbtn, text="+ Thêm CẢ THƯ MỤC...", command=self.merge_add_folder).pack(side="left", padx=2)
            tk.Button(row_mbtn, text="Xóa mục đã chọn", command=self.merge_remove_selected).pack(side="left", padx=2)
            tk.Button(row_mbtn, text="↑ Lên", command=self.merge_move_up).pack(side="left", padx=2)
            tk.Button(row_mbtn, text="↓ Xuống", command=self.merge_move_down).pack(side="left", padx=2)
            tk.Button(row_mbtn, text="Xóa toàn bộ danh sách", command=self.merge_clear_all).pack(side="left", padx=2)

            self.lst_merge_files = tk.Listbox(frm_merge, height=7, font=("Consolas", 9))
            self.lst_merge_files.pack(fill="both", expand=True, padx=6, pady=3)

            row_mout = tk.Frame(frm_merge)
            row_mout.pack(fill="x", padx=6, pady=3)
            tk.Label(row_mout, text="Lưu file gộp tại:", width=16, anchor="w").pack(side="left")
            tk.Entry(row_mout, textvariable=self.var_merge_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_mout, text="Chọn nơi lưu...", command=self.pick_merge_output).pack(side="left")

            row_run_merge = tk.Frame(frm_merge)
            row_run_merge.pack(fill="x", padx=6, pady=(4, 6))
            tk.Button(row_run_merge, text="📎 GỘP FILE PDF", font=("Arial", 11, "bold"),
                      bg="#00695c", fg="white", command=self.run_merge_clicked).pack(side="left")
            tk.Label(row_run_merge, text="Các file sẽ được gộp theo đúng thứ tự hiện trong danh sách ở trên.",
                     fg="#555").pack(side="left", padx=10)

            frm_group = styled_labelframe(parent, text="6c. Tách 1 file PDF nhiều trang thành NHIỀU FILE NHỎ, mỗi file N trang liên tiếp")
            frm_group.pack(fill="x", **pad)
            tk.Label(frm_group,
                     text="VD: file tổng 6 trang, đặt 'Số trang mỗi file' = 2 → tự tách thành 3 file: "
                          "trang 1-2, trang 3-4, trang 5-6, tên tự động đặt lần lượt, lưu vào thư mục bạn chọn.\n"
                          "Có thể dùng ngay file vừa GỘP ở mục 6b làm nguồn để tách tiếp.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_g1 = tk.Frame(frm_group)
            row_g1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_g1, text="File PDF nguồn (1 file):", width=28, anchor="w").pack(side="left")
            tk.Entry(row_g1, textvariable=self.var_group_input).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_g1, text="Chọn file...", command=self.pick_group_input).pack(side="left")
            tk.Button(row_g1, text="Dùng file vừa gộp ở 6b", command=self.use_merged_as_group_input).pack(side="left", padx=4)

            row_g2 = tk.Frame(frm_group)
            row_g2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_g2, text="Thư mục lưu kết quả:", width=28, anchor="w").pack(side="left")
            tk.Entry(row_g2, textvariable=self.var_group_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_g2, text="Chọn...", command=self.pick_group_output).pack(side="left")

            row_g3 = tk.Frame(frm_group)
            row_g3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_g3, text="Số trang mỗi file:", width=28, anchor="w").pack(side="left")
            tk.Entry(row_g3, textvariable=self.var_group_size, width=8).pack(side="left")

            row_g_run = tk.Frame(frm_group)
            row_g_run.pack(fill="x", padx=6, pady=(4, 6))
            tk.Button(row_g_run, text="✂ TÁCH THÀNH NHIỀU FILE NHỎ", font=("Arial", 11, "bold"),
                      bg="#37474f", fg="white", command=self.run_group_split_clicked).pack(side="left")

        def pick_split_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_split_output.set(p)

        def _set_split_preset(self, spec, suffix):
            self.var_split_page_spec.set(spec)
            self.var_split_suffix.set(suffix)

        def run_split_clicked(self):
            source = self.split_picker.get_source()
            output = self.var_split_output.get().strip()
            spec = self.var_split_page_spec.get().strip()
            suffix = self.var_split_suffix.get().strip() or "_tach"

            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF nguồn.")
                return
            if not output:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return

            self.txt_log.delete("1.0", "end")
            self._enable_pause_cancel("control_split", self.btn_split_pause, self.btn_split_cancel)

            def worker():
                try:
                    count = batch_split_folder(source, spec, suffix, output, self.log, control=self.control_split)

                    try:
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "TACH_PDF", run_id=stamp)
                        std_log.add(action="SPLIT", source_file=source, output_file=output,
                                   status="DA_TACH", message=f"{count} file")
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng kết quả chính): {e}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã tách {count} file.")
                    messagebox.showinfo("Hoàn tất", f"Đã tách {count} file PDF.\nLưu tại: {output}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self._disable_pause_cancel(self.btn_split_pause, self.btn_split_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def merge_add_files(self):
            paths = filedialog.askopenfilenames(title="Chọn các file PDF cần gộp", filetypes=[("PDF files", "*.pdf")])
            for p in paths:
                self.lst_merge_files.insert("end", p)

        def merge_add_folder(self):
            folder = filedialog.askdirectory(title="Chọn thư mục chứa PDF cần thêm vào danh sách gộp")
            if not folder:
                return
            added = 0
            for fname in sorted(os.listdir(folder)):
                if fname.lower().endswith(".pdf"):
                    self.lst_merge_files.insert("end", os.path.join(folder, fname))
                    added += 1
            if added == 0:
                messagebox.showinfo("Không có file", "Thư mục này không có file PDF nào.")

        def merge_remove_selected(self):
            for i in reversed(self.lst_merge_files.curselection()):
                self.lst_merge_files.delete(i)

        def merge_clear_all(self):
            self.lst_merge_files.delete(0, "end")

        def merge_move_up(self):
            sel = self.lst_merge_files.curselection()
            for i in sel:
                if i == 0:
                    continue
                text = self.lst_merge_files.get(i)
                self.lst_merge_files.delete(i)
                self.lst_merge_files.insert(i - 1, text)
                self.lst_merge_files.selection_set(i - 1)

        def merge_move_down(self):
            sel = list(self.lst_merge_files.curselection())
            for i in reversed(sel):
                if i >= self.lst_merge_files.size() - 1:
                    continue
                text = self.lst_merge_files.get(i)
                self.lst_merge_files.delete(i)
                self.lst_merge_files.insert(i + 1, text)
                self.lst_merge_files.selection_set(i + 1)

        def pick_merge_output(self):
            p = filedialog.asksaveasfilename(title="Lưu file PDF gộp tại", defaultextension=".pdf",
                                              filetypes=[("PDF files", "*.pdf")])
            if p:
                self.var_merge_output.set(p)

        def run_merge_clicked(self):
            files = list(self.lst_merge_files.get(0, "end"))
            output = self.var_merge_output.get().strip()

            if len(files) < 2:
                messagebox.showerror("Thiếu thông tin", "Cần ít nhất 2 file PDF để gộp.")
                return
            if not output:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn nơi lưu file gộp.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return

            self.txt_log.delete("1.0", "end")

            def worker():
                try:
                    self.log("Đang gộp các file theo thứ tự:")
                    for f in files:
                        self.log(f"  - {f}")
                    merge_pdfs(files, output)

                    try:
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "GOP_PDF", run_id=stamp)
                        for f in files:
                            std_log.add(action="MERGE", source_file=f, output_file=output, status="DA_GOP")
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng kết quả chính): {e}")

                    self.log(f"\n✓ Đã gộp xong: {output}")
                    messagebox.showinfo("Hoàn tất", f"Đã gộp {len(files)} file thành:\n{output}")
                    self.var_group_input.set(output)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 6c: Tách 1 file thành nhiều file nhỏ (nhóm trang cố định) -------------------

        def pick_group_input(self):
            p = filedialog.askopenfilename(title="Chọn file PDF nguồn", filetypes=[("PDF files", "*.pdf")])
            if p:
                self.var_group_input.set(p)

        def pick_group_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_group_output.set(p)

        def use_merged_as_group_input(self):
            if self.var_merge_output.get().strip():
                self.var_group_input.set(self.var_merge_output.get().strip())
            else:
                messagebox.showinfo("Chưa có file", "Bạn chưa gộp file nào ở mục 6b.")

        def run_group_split_clicked(self):
            input_path = self.var_group_input.get().strip()
            output_folder = self.var_group_output.get().strip()

            if not input_path or not os.path.isfile(input_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn 1 file PDF nguồn hợp lệ.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return
            try:
                group_size = int(self.var_group_size.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Số trang mỗi file phải là số nguyên.")
                return

            self.txt_log.delete("1.0", "end")

            def worker():
                try:
                    count = split_pdf_into_fixed_groups(input_path, group_size, output_folder, log_cb=self.log)

                    try:
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "TACH_NHOM_PDF", run_id=stamp)
                        std_log.add(action="SPLIT_GROUP", source_file=input_path, output_file=output_folder,
                                   status="DA_TACH", message=f"{count} file, {group_size} trang/file")
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng kết quả chính): {e}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã tách thành {count} file nhỏ.")
                    messagebox.showinfo("Hoàn tất", f"Đã tách thành {count} file.\nLưu tại: {output_folder}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_excel import ExcelTabMixin
    _UI_EXCEL_MODULE_OK = True
except Exception:
    _UI_EXCEL_MODULE_OK = False

    class ExcelTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_excel.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_excel(self, parent, pad):
            frm_paths = styled_labelframe(parent, text="1. Chọn thư mục và file")
            frm_paths.pack(fill="x", **pad)
            self._row_picker(frm_paths, "Thư mục PDF 1:", self.var_folder1, self.pick_folder1)
            self._row_picker(frm_paths, "Thư mục PDF 2 (không bắt buộc):", self.var_folder2, self.pick_folder2)
            self._row_picker(frm_paths, "File Excel tổng báo cáo:", self.var_excel, self.pick_excel)
            self._row_picker(frm_paths, "Thư mục lưu file KHỚP:", self.var_output_matched, self.pick_output_matched)
            self._row_picker(frm_paths, "Thư mục lưu file KHÔNG khớp:", self.var_output_unmatched, self.pick_output_unmatched)

            frm_cfg = styled_labelframe(parent, text="2. Cấu hình cột Excel (chữ cái cột, VD: B, W, X...)")
            frm_cfg.pack(fill="x", **pad)
            self._row_entry(frm_cfg, "Tên Sheet (để trống = sheet đầu tiên):", self.var_sheet)
            self._row_entry(frm_cfg, "Dòng bắt đầu dữ liệu:", self.var_header_row)
            self._row_entry(frm_cfg, "Cột Số TT:", self.var_col_stt)
            self._row_entry(frm_cfg, "Cột Mã xã:", self.var_col_maxa)
            self._row_entry(frm_cfg, "Cột Số tờ:", self.var_col_to)
            self._row_entry(frm_cfg, "Cột Số thửa:", self.var_col_thua)
            self._row_entry(frm_cfg, "Cột Diện tích:", self.var_col_dt)
            self._row_entry(frm_cfg, "Cột Tên chủ SD:", self.var_col_ten)
            self._row_entry(frm_cfg, "Cột Tên file quét (để trống nếu không có):", self.var_col_files)

            frm_crit = styled_labelframe(parent, text="3. Tiêu chí BẮT BUỘC phải khớp")
            frm_crit.pack(fill="x", **pad)

            row_preset = tk.Frame(frm_crit)
            row_preset.pack(fill="x", padx=6, pady=(4, 2))
            tk.Label(row_preset, text="Chọn nhanh:", width=12, anchor="w").pack(side="left")
            tk.Button(row_preset, text="Mã xã + Tờ + Thửa",
                      command=lambda: self.set_preset(maxa=True, to=True, thua=True, dt=False, ten=False)).pack(side="left", padx=3)
            tk.Button(row_preset, text="Tờ + Thửa",
                      command=lambda: self.set_preset(maxa=False, to=True, thua=True, dt=False, ten=False)).pack(side="left", padx=3)
            tk.Button(row_preset, text="Tên chủ",
                      command=lambda: self.set_preset(maxa=False, to=False, thua=False, dt=False, ten=True)).pack(side="left", padx=3)
            tk.Button(row_preset, text="Mã xã + Tên chủ",
                      command=lambda: self.set_preset(maxa=True, to=False, thua=False, dt=False, ten=True)).pack(side="left", padx=3)

            row_c = tk.Frame(frm_crit)
            row_c.pack(fill="x", padx=6, pady=4)
            tk.Checkbutton(row_c, text="Mã xã", variable=self.var_crit_maxa).pack(side="left", padx=8)
            tk.Checkbutton(row_c, text="Số tờ", variable=self.var_crit_to).pack(side="left", padx=8)
            tk.Checkbutton(row_c, text="Số thửa", variable=self.var_crit_thua).pack(side="left", padx=8)
            tk.Checkbutton(row_c, text="Tên chủ", variable=self.var_crit_ten).pack(side="left", padx=8)
            tk.Checkbutton(row_c, text="Diện tích, sai số ±", variable=self.var_crit_dt).pack(side="left", padx=8)
            tk.Entry(row_c, textvariable=self.var_dt_tolerance, width=6).pack(side="left")
            tk.Label(row_c, text="m²").pack(side="left")

            tk.Label(frm_crit,
                     text="Tên chủ: so khớp nếu TÊN CHỦ (bỏ dấu) xuất hiện trong tên file. Chỉ dùng khi tên file có ghi tên chủ.",
                     fg="#555").pack(anchor="w", padx=6, pady=(0, 4))

            frm_run = tk.Frame(parent)
            frm_run.pack(fill="x", **pad)
            self.btn_run = tk.Button(frm_run, text="▶ CHẠY ĐỐI CHIẾU", font=("Arial", 11, "bold"),
                                      bg="#2e7d32", fg="white", command=self.run_clicked)
            self.btn_run.pack(side="left", padx=4)
            self.lbl_status = tk.Label(frm_run, text="", font=("Arial", 10))
            self.lbl_status.pack(side="left", padx=10)

            frm_rename = styled_labelframe(parent, text="4. Đổi tên hàng loạt file PDF (mẫu tên tự chọn, 1 thao tác)")
            frm_rename.pack(fill="x", **pad)

            row_r = tk.Frame(frm_rename)
            row_r.pack(fill="x", padx=6, pady=3)
            tk.Label(row_r, text="Thư mục cần đổi tên:", width=32, anchor="w").pack(side="left")
            tk.Entry(row_r, textvariable=self.var_rename_folder).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_r, text="Chọn...", command=self.pick_rename_folder).pack(side="left")
            tk.Button(row_r, text="Dùng thư mục KHỚP ở trên", command=self.use_matched_as_rename).pack(side="left", padx=4)

            row_tpl = tk.Frame(frm_rename)
            row_tpl.pack(fill="x", padx=6, pady=3)
            tk.Label(row_tpl, text="Mẫu tên mới:", width=32, anchor="w").pack(side="left")
            tk.Entry(row_tpl, textvariable=self.var_name_template).pack(side="left", fill="x", expand=True, padx=4)

            row_tpl_hint = tk.Frame(frm_rename)
            row_tpl_hint.pack(fill="x", padx=6, pady=(0, 3))
            tk.Label(row_tpl_hint,
                     text="Placeholder dùng được: {maxa} {to} {thua} {ten} {loai} {stt} {dt} {original} — "
                          "VD: TBA_{maxa}_{to}_{thua}   hoặc   {ten}_To{to}_Thua{thua}_{loai}",
                     fg="#555", wraplength=880, justify="left").pack(anchor="w")

            row_tpl_presets = tk.Frame(frm_rename)
            row_tpl_presets.pack(fill="x", padx=6, pady=(0, 4))
            tk.Label(row_tpl_presets, text="Mẫu có sẵn:", width=12, anchor="w").pack(side="left")
            tk.Button(row_tpl_presets, text="CHUACOGIAY_xã_tờ_thửa_loại",
                      command=lambda: self.var_name_template.set("CHUACOGIAY_{maxa}_{to}_{thua}_{loai}")).pack(side="left", padx=3)
            tk.Button(row_tpl_presets, text="xã_tờ_thửa",
                      command=lambda: self.var_name_template.set("{maxa}_{to}_{thua}")).pack(side="left", padx=3)
            tk.Button(row_tpl_presets, text="Tên chủ_tờ_thửa_loại",
                      command=lambda: self.var_name_template.set("{ten}_To{to}_Thua{thua}_{loai}")).pack(side="left", padx=3)

            row_fallback = tk.Frame(frm_rename)
            row_fallback.pack(fill="x", padx=6, pady=(0, 3))
            tk.Checkbutton(row_fallback,
                           text="Nếu 1 file không khớp Excel theo TÊN FILE, tự đọc trực tiếp NỘI DUNG PDF để lấy "
                                "Mã xã/Tờ/Thửa thay vì bỏ qua (dùng chung cấu hình OCR + bảng tra Thôn/Xã ở Tab 5)",
                           variable=self.var_rename_content_fallback).pack(anchor="w")

            row_r2 = tk.Frame(frm_rename)
            row_r2.pack(fill="x", padx=6, pady=(2, 6))
            self.btn_rename = tk.Button(row_r2, text="✎ ĐỔI TÊN HÀNG LOẠT", font=("Arial", 11, "bold"),
                                         bg="#1565c0", fg="white", command=self.rename_clicked)
            self.btn_rename.pack(side="left", padx=4)
            self.btn_rename_pause, self.btn_rename_cancel = self._add_pause_cancel(row_r2, "control_rename")
            self.btn_rename_pause.pack(side="left", padx=(8, 3))
            self.btn_rename_cancel.pack(side="left", padx=3)
            tk.Label(row_r2, text="Việc đổi tên sẽ tự động dùng tiêu chí đã tick ở mục 3 để xác định Mã xã/Tờ/Thửa/Tên chủ cho từng file.",
                     fg="#555", wraplength=700, justify="left").pack(side="left", padx=10)

        def _row_picker(self, parent, label, var, cmd):
            row = tk.Frame(parent)
            row.pack(fill="x", padx=6, pady=3)
            tk.Label(row, text=label, width=32, anchor="w").pack(side="left")
            tk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row, text="Chọn...", command=cmd).pack(side="left")

        def _row_entry(self, parent, label, var):
            row = tk.Frame(parent)
            row.pack(fill="x", padx=6, pady=2)
            tk.Label(row, text=label, width=45, anchor="w").pack(side="left")
            tk.Entry(row, textvariable=var, width=15).pack(side="left")

        def set_preset(self, maxa, to, thua, dt, ten):
            self.var_crit_maxa.set(maxa)
            self.var_crit_to.set(to)
            self.var_crit_thua.set(thua)
            self.var_crit_dt.set(dt)
            self.var_crit_ten.set(ten)

        def pick_folder1(self):
            p = filedialog.askdirectory(title="Chọn Thư mục PDF 1")
            if p:
                self.var_folder1.set(p)

        def pick_folder2(self):
            p = filedialog.askdirectory(title="Chọn Thư mục PDF 2")
            if p:
                self.var_folder2.set(p)

        def pick_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel", filetypes=[("Excel files", "*.xlsx *.xlsm")])
            if p:
                self.var_excel.set(p)

        def pick_output_matched(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu file KHỚP")
            if p:
                self.var_output_matched.set(p)

        def pick_output_unmatched(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu file KHÔNG khớp")
            if p:
                self.var_output_unmatched.set(p)

        def pick_rename_folder(self):
            p = filedialog.askdirectory(title="Chọn thư mục cần đổi tên")
            if p:
                self.var_rename_folder.set(p)

        def use_matched_as_rename(self):
            if self.var_output_matched.get().strip():
                self.var_rename_folder.set(self.var_output_matched.get().strip())
            else:
                messagebox.showinfo("Chưa có thư mục", "Bạn chưa chọn Thư mục lưu file KHỚP ở mục 1.")

        def get_criteria(self):
            c = set()
            if self.var_crit_maxa.get():
                c.add("maxa")
            if self.var_crit_to.get():
                c.add("to")
            if self.var_crit_thua.get():
                c.add("thua")
            if self.var_crit_dt.get():
                c.add("dt")
            if self.var_crit_ten.get():
                c.add("ten")
            return c

        def get_dt_tolerance(self):
            try:
                return float(self.var_dt_tolerance.get().strip())
            except ValueError:
                return 5.0

        def run_clicked(self):
            excel_path = self.var_excel.get().strip()
            output_matched = self.var_output_matched.get().strip()
            output_unmatched = self.var_output_unmatched.get().strip()
            folder1 = self.var_folder1.get().strip()
            folder2 = self.var_folder2.get().strip()

            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel hợp lệ.")
                return
            if not output_matched or not output_unmatched:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn cả 2 thư mục đích (KHỚP và KHÔNG khớp).")
                return
            if not folder1 and not folder2:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn ít nhất 1 thư mục PDF.")
                return

            criteria = self.get_criteria()
            if not criteria:
                messagebox.showerror("Thiếu tiêu chí", "Vui lòng tick ít nhất 1 tiêu chí đối chiếu.")
                return

            try:
                header_row = int(self.var_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            self.txt_log.delete("1.0", "end")
            self.btn_run.config(state="disabled")
            self.lbl_status.config(text="Đang xử lý...")

            def worker():
                try:
                    matched, unmatched, ambiguous, missing_rows, report_path = process(
                        folders=[folder1, folder2],
                        excel_path=excel_path,
                        output_matched=output_matched,
                        output_unmatched=output_unmatched,
                        sheet_name=self.var_sheet.get().strip() or None,
                        header_row=header_row,
                        col_stt=self.var_col_stt.get().strip() or None,
                        col_maxa=self.var_col_maxa.get().strip(),
                        col_to=self.var_col_to.get().strip(),
                        col_thua=self.var_col_thua.get().strip(),
                        col_ten=self.var_col_ten.get().strip() or None,
                        col_files=self.var_col_files.get().strip() or None,
                        col_dt=self.var_col_dt.get().strip() or None,
                        criteria=criteria,
                        dt_tolerance=self.get_dt_tolerance(),
                        log_cb=self.log,
                    )
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {matched} file khớp | {unmatched} file KHÔNG khớp (trong đó {ambiguous} khớp nhiều dòng)")
                    self.log(f"          {len(missing_rows)} dòng Excel CHƯA có file PDF tương ứng.")
                    self.log(f"Báo cáo chi tiết: {report_path}")
                    self.lbl_status.config(text=f"Hoàn tất: {matched} khớp / {unmatched} không khớp")
                    messagebox.showinfo("Hoàn tất",
                                         f"Khớp: {matched} file\nKhông khớp: {unmatched} file\n\nBáo cáo tại:\n{report_path}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_run.config(state="normal")

            threading.Thread(target=worker, daemon=True).start()

        def rename_clicked(self):
            folder = self.var_rename_folder.get().strip()
            excel_path = self.var_excel.get().strip()
            template = self.var_name_template.get().strip()

            if not folder or not os.path.isdir(folder):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục cần đổi tên.")
                return
            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel ở mục 1 (dùng để tra cứu dữ liệu).")
                return
            if not template:
                messagebox.showerror("Thiếu thông tin", "Vui lòng nhập mẫu tên mới ở mục 4.")
                return

            criteria = self.get_criteria()
            if not criteria:
                messagebox.showerror("Thiếu tiêu chí", "Vui lòng tick ít nhất 1 tiêu chí ở mục 3.")
                return

            try:
                header_row = int(self.var_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            if not messagebox.askyesno(
                    "Xác nhận đổi tên",
                    f"Sẽ đổi tên các file PDF trong:\n{folder}\n\n"
                    f"Theo mẫu: {template}\n\n"
                    "Hành động này SỬA TRỰC TIẾP tên file gốc (không phải bản copy). Tiếp tục?"):
                return

            self.txt_log.delete("1.0", "end")
            self.btn_rename.config(state="disabled")
            self._enable_pause_cancel("control_rename", self.btn_rename_pause, self.btn_rename_cancel)

            def worker():
                try:
                    use_fallback = self.var_rename_content_fallback.get()
                    xa_mapping = None
                    thon_mapping = None
                    use_ocr = False
                    ocr_dpi = 300
                    debug_flag = False
                    if use_fallback:
                        xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
                        thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
                        use_ocr = self.var_use_ocr.get()
                        try:
                            ocr_dpi = int(self.var_ocr_dpi.get().strip())
                        except ValueError:
                            ocr_dpi = 300
                        debug_flag = self.var_debug_content.get()

                    renamed, skipped, already_ok, report_path, fallback_count = rename_bulk(
                        folder=folder,
                        excel_path=excel_path,
                        sheet_name=self.var_sheet.get().strip() or None,
                        header_row=header_row,
                        col_stt=self.var_col_stt.get().strip() or None,
                        col_maxa=self.var_col_maxa.get().strip(),
                        col_to=self.var_col_to.get().strip(),
                        col_thua=self.var_col_thua.get().strip(),
                        col_ten=self.var_col_ten.get().strip() or None,
                        col_files=self.var_col_files.get().strip() or None,
                        col_dt=self.var_col_dt.get().strip() or None,
                        criteria=criteria,
                        dt_tolerance=self.get_dt_tolerance(),
                        name_template=template,
                        log_cb=self.log,
                        xa_mapping=xa_mapping,
                        thon_mapping=thon_mapping,
                        use_ocr_fallback=use_ocr,
                        ocr_dpi=ocr_dpi,
                        debug=debug_flag,
                        control=self.control_rename,
                    )
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT ĐỔI TÊN: {renamed} file đã đổi tên | {already_ok} file đã đúng chuẩn | {skipped} file bỏ qua (không khớp)")
                    if fallback_count:
                        self.log(f"  (Trong đó {fallback_count} file được xác định nhờ ĐỌC TRỰC TIẾP nội dung PDF, không khớp Excel theo tên file)")
                    if report_path:
                        self.log(f"Báo cáo đổi tên: {report_path}")
                    messagebox.showinfo("Hoàn tất đổi tên",
                                         f"Đã đổi tên: {renamed} file\nĐã đúng chuẩn: {already_ok} file\nBỏ qua: {skipped} file")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi đổi tên", str(e))
                finally:
                    self.btn_rename.config(state="normal")
                    self._disable_pause_cancel(self.btn_rename_pause, self.btn_rename_cancel)

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_reconcile import ReconcileTabMixin
    _UI_RECONCILE_MODULE_OK = True
except Exception:
    _UI_RECONCILE_MODULE_OK = False

    class ReconcileTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_reconcile.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_reconcile(self, parent, pad):
            frm = styled_labelframe(parent, text="NHÓM B — Lọc Excel tổng theo thư mục PDF & hoàn thiện báo cáo")
            frm.pack(fill="both", expand=True, **pad)
            tk.Label(frm,
                     text="Đối chiếu từng dòng Excel TỔNG với các file PDF trong thư mục theo Mã xã+Tờ+Thửa "
                          "(hoặc chỉ Tờ+Thửa nếu PDF không xác định được mã xã). Dòng KHÔNG có PDF tương ứng sẽ "
                          "bị XÓA khỏi báo cáo; dòng CÓ PDF thì giữ lại, có thể hoàn thiện thêm dữ liệu từ PDF.\n"
                          "Nhận diện khóa PDF ưu tiên từ TÊN FILE trước (CHUACOGIAY_maxa_to_thua..., hoặc chỉ "
                          "to_thua.pdf), nếu tên file không đủ mới đọc nội dung/OCR. File Excel gốc KHÔNG bị sửa "
                          "— luôn lưu ra file báo cáo mới.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=6, pady=(6, 6))

            row_1 = tk.Frame(frm)
            row_1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_1, text="① File Excel tổng:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_1, textvariable=self.var_rec_excel).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_1, text="Chọn...", command=self.pick_rec_excel).pack(side="left")

            row_2 = tk.Frame(frm)
            row_2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_2, text="② Thư mục chứa PDF:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_2, textvariable=self.var_rec_pdf_folder).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_2, text="Chọn...", command=self.pick_rec_pdf_folder).pack(side="left")

            row_3 = tk.Frame(frm)
            row_3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_3, text="③ Thư mục xuất kết quả:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_3, textvariable=self.var_rec_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_3, text="Chọn...", command=self.pick_rec_output).pack(side="left")

            row_4 = tk.Frame(frm)
            row_4.pack(fill="x", padx=6, pady=3)
            tk.Label(row_4, text="Sheet:").pack(side="left")
            tk.Entry(row_4, textvariable=self.var_rec_sheet, width=15).pack(side="left", padx=(4, 20))
            tk.Label(row_4, text="Dòng bắt đầu dữ liệu:").pack(side="left")
            tk.Entry(row_4, textvariable=self.var_rec_header_row, width=6).pack(side="left", padx=4)

            row_5 = tk.Frame(frm)
            row_5.pack(fill="x", padx=6, pady=3)
            tk.Label(row_5, text="Cột Mã xã:").pack(side="left")
            tk.Entry(row_5, textvariable=self.var_rec_col_maxa, width=5).pack(side="left", padx=(4, 16))
            tk.Label(row_5, text="Cột Số tờ:").pack(side="left")
            tk.Entry(row_5, textvariable=self.var_rec_col_to, width=5).pack(side="left", padx=(4, 16))
            tk.Label(row_5, text="Cột Số thửa:").pack(side="left")
            tk.Entry(row_5, textvariable=self.var_rec_col_thua, width=5).pack(side="left", padx=(4, 16))
            tk.Label(row_5, text="Cột STT:").pack(side="left")
            tk.Entry(row_5, textvariable=self.var_rec_col_stt, width=5).pack(side="left", padx=4)

            row_6 = tk.Frame(frm)
            row_6.pack(fill="x", padx=6, pady=(8, 3))
            tk.Checkbutton(row_6, text="Quét cả thư mục con", variable=self.var_rec_subfolders).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_6, text="Sau khi lọc, hoàn thiện dữ liệu từ PDF vào Excel",
                           variable=self.var_rec_complete).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_6, text="Cho phép ghi đè ô đã có dữ liệu bằng dữ liệu từ PDF",
                           variable=self.var_rec_overwrite).pack(side="left")

            row_7 = tk.Frame(frm)
            row_7.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_7, text="Tự đánh lại số thứ tự (STT) sau khi xóa dòng",
                           variable=self.var_rec_renumber).pack(side="left")

            tk.Label(frm,
                     text="Lưu ý: dùng chung cấu hình OCR + bảng tra Thôn/Xã ở Tab 5 khi cần đọc nội dung PDF.",
                     fg="#777", justify="left", wraplength=1000).pack(anchor="w", padx=6, pady=(4, 4))

            row_run = tk.Frame(frm)
            row_run.pack(fill="x", padx=6, pady=(6, 6))
            self.btn_rec_dryrun = tk.Button(row_run, text="CHẠY THỬ ĐỐI CHIẾU (không xóa/không lưu)",
                                             font=("Arial", 10, "bold"), bg="#455a64", fg="white",
                                             command=lambda: self.run_reconcile_clicked(dry_run=True))
            self.btn_rec_dryrun.pack(side="left", padx=(0, 6))
            self.btn_rec_run = tk.Button(row_run, text="LỌC & HOÀN THIỆN BÁO CÁO",
                                          font=("Arial", 10, "bold"), bg="#bf360c", fg="white",
                                          command=lambda: self.run_reconcile_clicked(dry_run=False))
            self.btn_rec_run.pack(side="left")
            self.btn_rec_pause, self.btn_rec_cancel = self._add_pause_cancel(row_run, "control_rec")
            self.btn_rec_pause.pack(side="left", padx=(12, 3))
            self.btn_rec_cancel.pack(side="left", padx=3)

            tk.Label(frm, text="Bảng log kết quả:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 2))
            cols = ("dong", "maxa", "to", "thua", "status", "note")
            headers = ["Dòng Excel", "Mã xã", "Số tờ", "Số thửa", "Trạng thái", "Ghi chú"]
            widths = [70, 60, 60, 60, 150, 340]
            self.tree_rec = ttk.Treeview(frm, columns=cols, show="headings", height=10)
            for c, h, w in zip(cols, headers, widths):
                self.tree_rec.heading(c, text=h)
                self.tree_rec.column(c, width=w, anchor="w")
            vsb_rec = tk.Scrollbar(frm, orient="vertical", command=self.tree_rec.yview)
            self.tree_rec.configure(yscrollcommand=vsb_rec.set)
            self.tree_rec.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_rec.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            # ================= NHÓM C: So sánh thư mục Word/PDF theo Mã xã-Tờ-Thửa =================
            frm_c = styled_labelframe(parent, text="NHÓM C — So sánh thư mục Word/PDF theo Mã xã-Tờ-Thửa")
            frm_c.pack(fill="both", expand=True, padx=6, pady=6)
            tk.Label(frm_c,
                     text="So sánh 2 CHIỀU giữa 2 thư mục chứa file .pdf/.doc/.docx/.docm theo khóa "
                          "Mã xã+Tờ+Thửa (hoặc chỉ Tờ+Thửa nếu không có mã xã). File không khớp được "
                          "COPY hoặc DI CHUYỂN ra thư mục kết quả riêng — KHÔNG xóa file, KHÔNG ghi đè, "
                          "KHÔNG đổi tên/sửa nội dung file gốc.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=6, pady=(6, 6))

            row_c1 = tk.Frame(frm_c)
            row_c1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c1, text="Thư mục A:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_c1, textvariable=self.var_cmp_folder_a).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_c1, text="Chọn...", command=self.pick_cmp_folder_a).pack(side="left")

            row_c2 = tk.Frame(frm_c)
            row_c2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c2, text="Thư mục B:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_c2, textvariable=self.var_cmp_folder_b).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_c2, text="Chọn...", command=self.pick_cmp_folder_b).pack(side="left")

            row_c3 = tk.Frame(frm_c)
            row_c3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c3, text="Thư mục lưu kết quả:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_c3, textvariable=self.var_cmp_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_c3, text="Chọn...", command=self.pick_cmp_output).pack(side="left")

            row_c4 = tk.Frame(frm_c)
            row_c4.pack(fill="x", padx=6, pady=(6, 3))
            tk.Checkbutton(row_c4, text="Quét cả thư mục con", variable=self.var_cmp_subfolders).pack(side="left", padx=(0, 12))
            tk.Checkbutton(row_c4, text="So sánh kèm mã xã nếu có", variable=self.var_cmp_use_maxa).pack(side="left", padx=(0, 12))
            tk.Checkbutton(row_c4, text="Nếu không có mã xã thì so sánh theo Số tờ + Số thửa",
                           variable=self.var_cmp_fallback_thua).pack(side="left")

            row_c5 = tk.Frame(frm_c)
            row_c5.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_c5, text="Di chuyển/copy file không khớp của thư mục A",
                           variable=self.var_cmp_move_a).pack(side="left", padx=(0, 12))
            tk.Checkbutton(row_c5, text="Di chuyển/copy file không khớp của thư mục B",
                           variable=self.var_cmp_move_b).pack(side="left", padx=(0, 12))
            tk.Checkbutton(row_c5, text="Copy file thay vì di chuyển (giữ nguyên file gốc)",
                           variable=self.var_cmp_copy_mode).pack(side="left")

            row_c_run = tk.Frame(frm_c)
            row_c_run.pack(fill="x", padx=6, pady=(6, 6))
            self.btn_cmp_dryrun = tk.Button(row_c_run, text="Chạy thử so sánh", font=("Arial", 10, "bold"),
                                             bg="#455a64", fg="white", command=lambda: self.run_compare_folders_clicked(dry_run=True))
            self.btn_cmp_dryrun.pack(side="left", padx=(0, 6))
            self.btn_cmp_run = tk.Button(row_c_run, text="So sánh và di chuyển file không khớp",
                                          font=("Arial", 10, "bold"), bg="#bf360c", fg="white",
                                          command=lambda: self.run_compare_folders_clicked(dry_run=False))
            self.btn_cmp_run.pack(side="left")
            self.btn_cmp_pause, self.btn_cmp_cancel = self._add_pause_cancel(row_c_run, "control_cmp")
            self.btn_cmp_pause.pack(side="left", padx=(12, 3))
            self.btn_cmp_cancel.pack(side="left", padx=3)

            tk.Label(frm_c, text="Bảng log:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(4, 2))
            cols_c = ("stt", "thumuc", "file", "maxa", "to", "thua", "khoa", "status", "note")
            headers_c = ["STT", "Thư mục", "Tên file", "Mã xã", "Tờ", "Thửa", "Khóa so sánh", "Trạng thái", "Ghi chú"]
            widths_c = [40, 55, 180, 55, 40, 45, 100, 170, 260]
            self.tree_cmp = ttk.Treeview(frm_c, columns=cols_c, show="headings", height=10)
            for c, h, w in zip(cols_c, headers_c, widths_c):
                self.tree_cmp.heading(c, text=h)
                self.tree_cmp.column(c, width=w, anchor="w")
            vsb_cmp = tk.Scrollbar(frm_c, orient="vertical", command=self.tree_cmp.yview)
            self.tree_cmp.configure(yscrollcommand=vsb_cmp.set)
            self.tree_cmp.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_cmp.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

        def pick_rec_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel tổng", filetypes=[("Excel files", "*.xlsx *.xlsm")])
            if p:
                self.var_rec_excel.set(p)

        def pick_rec_pdf_folder(self):
            p = filedialog.askdirectory(title="Chọn thư mục chứa PDF")
            if p:
                self.var_rec_pdf_folder.set(p)

        def pick_rec_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục xuất kết quả")
            if p:
                self.var_rec_output.set(p)

        def _refresh_rec_tree(self, log_rows):
            for item in self.tree_rec.get_children():
                self.tree_rec.delete(item)
            for r in log_rows:
                self.tree_rec.insert("", "end", values=(r["dong"], r["maxa"], r["to"], r["thua"], r["status"], r["note"]))

        def run_reconcile_clicked(self, dry_run=True):
            excel_path = self.var_rec_excel.get().strip()
            pdf_folder = self.var_rec_pdf_folder.get().strip()
            output_dir = self.var_rec_output.get().strip()

            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel tổng hợp lệ.")
                return
            if not pdf_folder or not os.path.isdir(pdf_folder):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục chứa PDF hợp lệ.")
                return
            if not dry_run and not output_dir:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục xuất kết quả.")
                return

            try:
                header_row = int(self.var_rec_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            col_maxa = self.var_rec_col_maxa.get().strip()
            col_to = self.var_rec_col_to.get().strip()
            col_thua = self.var_rec_col_thua.get().strip()
            col_stt = self.var_rec_col_stt.get().strip()
            if not col_to or not col_thua:
                messagebox.showerror("Thiếu thông tin", "Vui lòng nhập ít nhất cột Số tờ và Số thửa.")
                return

            include_sub = self.var_rec_subfolders.get()
            complete_from_pdf = self.var_rec_complete.get()
            allow_overwrite = self.var_rec_overwrite.get()
            renumber = self.var_rec_renumber.get()

            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            if not dry_run:
                if not self._check_require_dry_run("_dry_run_done_reconcile", True, "Đối chiếu và hoàn thiện báo cáo (Tab 9)"):
                    return
                if not messagebox.askyesno(
                        "Xác nhận lọc & hoàn thiện báo cáo",
                        "Thao tác này sẽ XÓA khỏi báo cáo các dòng Excel không có file PDF tương ứng, và có thể "
                        "ghi đè dữ liệu vào các dòng còn lại (nếu bật tùy chọn). File Excel GỐC không bị ảnh "
                        "hưởng — kết quả lưu ra file MỚI. Tiếp tục?"):
                    return

            self.txt_log.delete("1.0", "end")
            self._refresh_rec_tree([])
            self.btn_rec_dryrun.config(state="disabled")
            self.btn_rec_run.config(state="disabled")
            self._enable_pause_cancel("control_rec", self.btn_rec_pause, self.btn_rec_cancel)

            def worker():
                try:
                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    output_path = None
                    if not dry_run:
                        output_path = os.path.join(output_dir, f"BAO_CAO_DA_DOI_KHOP_HOAN_THIEN_{stamp}.xlsx")
                        self._maybe_auto_backup([excel_path], "Đối chiếu và hoàn thiện báo cáo (Tab 9)")

                    result = reconcile_excel_with_pdfs(
                        excel_path, pdf_folder, self.var_rec_sheet.get().strip() or None, header_row,
                        col_maxa or None, col_to, col_thua, col_stt or None,
                        include_sub, complete_from_pdf, allow_overwrite, renumber,
                        xa_mapping, thon_mapping, use_ocr, ocr_dpi,
                        dry_run, output_path, self.log, control=self.control_rec)
                    if dry_run:
                        self._dry_run_done_reconcile = True

                    self.after(0, lambda: self._refresh_rec_tree(result["log_rows"]))

                    log_dir = output_dir or os.path.dirname(excel_path)
                    log_paths = write_reconcile_logs(log_dir, stamp, result)

                    try:
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "DOI_CHIEU_EXCEL", run_id=stamp)
                        for r in result["log_rows"]:
                            std_log.add(action="RECONCILE", source_file=excel_path,
                                       ma_xa=r.get("maxa"), so_to=r.get("to"), so_thua=r.get("thua"),
                                       key=f"{r.get('maxa')}_{r.get('to')}_{r.get('thua')}",
                                       status=r["status"], message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: giữ lại {result['kept_count']} dòng | xóa {result['deleted_count']} dòng | "
                             f"{len(result['pdf_need_check'])} PDF cần kiểm tra")
                    self.log(f"Log dòng bị xóa: {log_paths['deleted_csv']}")
                    self.log(f"Log PDF cần kiểm tra: {log_paths['pdf_check_csv']}")
                    self.log(f"Log đầy đủ: {log_paths['full_log_csv']}")
                    if output_path:
                        self.log(f"Báo cáo kết quả: {output_path}")

                    msg = (f"Giữ lại: {result['kept_count']} dòng\nXóa: {result['deleted_count']} dòng\n"
                           f"PDF cần kiểm tra: {len(result['pdf_need_check'])}\n\n")
                    msg += f"Báo cáo: {output_path}" if output_path else "(Chạy thử - chưa lưu file)"
                    self.after(0, lambda: self.mark_workflow_step("doi_chieu", "Đã chạy thử" if dry_run else "Đã hoàn thành"))
                    messagebox.showinfo("Chạy thử hoàn tất" if dry_run else "Lọc & hoàn thiện hoàn tất", msg)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_rec_dryrun.config(state="normal")
                    self.btn_rec_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_rec_pause, self.btn_rec_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- NHÓM C: So sánh thư mục Word/PDF -------------------

        def pick_cmp_folder_a(self):
            p = filedialog.askdirectory(title="Chọn thư mục A")
            if p:
                self.var_cmp_folder_a.set(p)

        def pick_cmp_folder_b(self):
            p = filedialog.askdirectory(title="Chọn thư mục B")
            if p:
                self.var_cmp_folder_b.set(p)

        def pick_cmp_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_cmp_output.set(p)

        def _refresh_cmp_tree(self, rows):
            for item in self.tree_cmp.get_children():
                self.tree_cmp.delete(item)
            for r in rows:
                self.tree_cmp.insert("", "end", values=(r["stt"], r["thumuc"], r["file"], r["maxa"], r["to"],
                                                          r["thua"], r["khoa"], r["status"], r["note"]))

        def run_compare_folders_clicked(self, dry_run=True):
            folder_a = self.var_cmp_folder_a.get().strip()
            folder_b = self.var_cmp_folder_b.get().strip()
            output_folder = self.var_cmp_output.get().strip()

            if not folder_a or not os.path.isdir(folder_a):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn Thư mục A hợp lệ.")
                return
            if not folder_b or not os.path.isdir(folder_b):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn Thư mục B hợp lệ.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return

            use_maxa = self.var_cmp_use_maxa.get()
            fallback_thua = self.var_cmp_fallback_thua.get()
            move_a = self.var_cmp_move_a.get()
            move_b = self.var_cmp_move_b.get()
            copy_mode = self.var_cmp_copy_mode.get()
            include_sub = self.var_cmp_subfolders.get()

            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            if not dry_run and (move_a or move_b):
                action_word = "COPY" if copy_mode else "DI CHUYỂN"
                if not messagebox.askyesno(
                        "Xác nhận",
                        f"Sẽ {action_word} các file không khớp ra thư mục kết quả. File gốc KHÔNG bị xóa"
                        + ("" if copy_mode else ", chỉ bị chuyển sang vị trí khác") + ". Tiếp tục?"):
                    return

            self.txt_log.delete("1.0", "end")
            self._refresh_cmp_tree([])
            self.btn_cmp_dryrun.config(state="disabled")
            self.btn_cmp_run.config(state="disabled")
            self._enable_pause_cancel("control_cmp", self.btn_cmp_pause, self.btn_cmp_cancel)

            def worker():
                try:
                    use_com_session = HAS_WIN32COM and not dry_run
                    if use_com_session:
                        self.log("Đang mở 1 phiên Microsoft Word dùng chung để đọc nhanh các file .doc/.docx...")
                        with WordCOMSession() as word_session:
                            all_rows = compare_folders_word_pdf(
                                folder_a, folder_b, output_folder, include_sub, use_maxa, fallback_thua,
                                move_a, move_b, copy_mode, dry_run, xa_mapping, thon_mapping, use_ocr, ocr_dpi,
                                self.log, control=self.control_cmp, word_session=word_session)
                    else:
                        all_rows = compare_folders_word_pdf(
                            folder_a, folder_b, output_folder, include_sub, use_maxa, fallback_thua,
                            move_a, move_b, copy_mode, dry_run, xa_mapping, thon_mapping, use_ocr, ocr_dpi,
                            self.log, control=self.control_cmp)

                    self.after(0, lambda: self._refresh_cmp_tree(all_rows))

                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    log_paths = write_compare_folders_logs(output_folder, stamp, all_rows)

                    trung_khop = sum(1 for r in all_rows if r["status"] == "TRUNG_KHOP")
                    khong_khop_a = sum(1 for r in all_rows if r["status"] == "KHONG_KHOP_THU_MUC_A")
                    khong_khop_b = sum(1 for r in all_rows if r["status"] == "KHONG_KHOP_THU_MUC_B")
                    can_kiem_tra = sum(1 for r in all_rows if r["status"] in
                                       ("CAN_KIEM_TRA_THIEU_THONG_TIN", "LOI_DOC_FILE"))

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {trung_khop} trùng khớp | {khong_khop_a} không khớp (A) | "
                             f"{khong_khop_b} không khớp (B) | {can_kiem_tra} cần kiểm tra")
                    self.log(f"Log đầy đủ: {log_paths['full']}")

                    messagebox.showinfo(
                        "Chạy thử hoàn tất" if dry_run else "So sánh hoàn tất",
                        f"Trùng khớp: {trung_khop}\nKhông khớp (A): {khong_khop_a}\n"
                        f"Không khớp (B): {khong_khop_b}\nCần kiểm tra: {can_kiem_tra}\n\n"
                        f"Log: {log_paths['full']}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_cmp_dryrun.config(state="normal")
                    self.btn_cmp_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_cmp_pause, self.btn_cmp_cancel)

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_workflow import WorkflowTabMixin
    _UI_WORKFLOW_MODULE_OK = True
except Exception:
    _UI_WORKFLOW_MODULE_OK = False

    class WorkflowTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_workflow.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_workflow(self, parent, pad, nav):
            frm = styled_labelframe(parent, text="★ QUY TRÌNH XỬ LÝ HỒ SƠ — làm theo từng bước cho hồ sơ mới")
            frm.pack(fill="both", expand=True, **pad)
            tk.Label(frm,
                     text="Gợi ý trình tự xử lý 1 đợt hồ sơ đất đai từ đầu đến cuối. Mỗi bước có nút "
                          "'Mở chức năng' để nhảy thẳng tới đúng tab xử lý. Bạn có thể bỏ qua bước không cần dùng.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=6, pady=(6, 10))

            self.workflow_status_labels = {}
            steps = [
                ("chuan_hoa", "Bước 1: Chuẩn hóa PDF",
                 "Xoay PDF về đúng hướng, sửa nghiêng, tăng nét/tương phản, kiểm tra file scan/ảnh.",
                 "pdftools"),
                ("ocr", "Bước 2: OCR và nhận diện dữ liệu",
                 "Đọc PDF scan bằng OCR, nhận diện Mã xã/Số tờ/Số thửa/Họ tên/CCCD/Địa chỉ/Diện tích/Mục đích SD.",
                 "content"),
                ("doi_ten", "Bước 3: Đổi tên file PDF",
                 "Đổi tên theo cú pháp chuẩn CHUACOGIAY_[Mã xã]_[Tờ]_[Thửa]_GT.pdf. Luôn Chạy thử trước, "
                 "xuất log kiểm tra.", "content"),
                ("doi_chieu", "Bước 4: Đối chiếu Excel tổng với thư mục PDF",
                 "So sánh Mã xã/Số tờ/Số thửa, lọc dòng không khớp, hoàn thiện dữ liệu dòng khớp.", "reconcile"),
                ("xuat_word", "Bước 5: Xuất Word hàng loạt",
                 "Xuất Đơn đăng ký, Thông báo xác nhận, Biên bản/Công văn/Kế hoạch nếu có mẫu.", "word"),
                ("chuyen_pdf", "Bước 6: Chuyển Word sang PDF",
                 "Chuyển hàng loạt, giữ nguyên định dạng, không ảnh hưởng Word đang mở.", "word"),
                ("bao_cao", "Bước 7: Tổng hợp báo cáo hoàn chỉnh",
                 "Xuất Excel báo cáo, xuất log, xuất danh sách cần kiểm tra.", "report"),
            ]

            for key, title, desc, nav_key in steps:
                row = tk.Frame(frm, relief="groove", borderwidth=1)
                row.pack(fill="x", padx=6, pady=4)
                left = tk.Frame(row)
                left.pack(side="left", fill="both", expand=True, padx=8, pady=6)
                tk.Label(left, text=title, font=("Segoe UI", 10, "bold"), anchor="w").pack(anchor="w")
                tk.Label(left, text=desc, fg="#555", anchor="w", justify="left", wraplength=680).pack(anchor="w")

                right = tk.Frame(row)
                right.pack(side="right", padx=8, pady=6)
                status_lbl = tk.Label(right, text="Chưa chạy", fg="#888", width=16, anchor="e")
                status_lbl.pack(side="left", padx=(0, 10))
                self.workflow_status_labels[key] = status_lbl
                tk.Button(right, text="Mở chức năng →",
                          command=lambda w=nav[nav_key]: self.notebook.select(w)).pack(side="left")

            tk.Label(frm, text="Tiến độ tổng thể quy trình:", font=("Segoe UI", 9, "bold")).pack(
                anchor="w", padx=6, pady=(12, 2))
            self.workflow_progress = ttk.Progressbar(frm, orient="horizontal", mode="determinate", maximum=len(steps))
            self.workflow_progress.pack(fill="x", padx=6, pady=(0, 10))

            # ================= Giới thiệu =================
            frm_about = styled_labelframe(parent, text="ℹ Giới thiệu")
            frm_about.pack(fill="x", padx=6, pady=(0, 6))
            tk.Label(frm_about,
                     text="Phần mềm hỗ trợ xử lý hồ sơ đăng ký đất đai, đổi tên PDF, xuất Word/PDF hàng loạt "
                          "và tổng hợp báo cáo.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=8, pady=(8, 2))
            tk.Label(frm_about, text=f"Phiên bản hiện tại: SỸ LAND {RELEASE_LABEL} (build {APP_VERSION})",
                     font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=8, pady=(4, 2))

            row_about_btn = tk.Frame(frm_about)
            row_about_btn.pack(anchor="w", padx=8, pady=(0, 4))
            tk.Button(row_about_btn, text="🔄 Kiểm tra cập nhật",
                      command=lambda: self._check_for_update_async(force=True)).pack(side="left", padx=(0, 6))
            tk.Button(row_about_btn, text="⬇ Mở trang tải bản mới",
                      command=self.open_download_page).pack(side="left")

            # --- Kiểm tra môi trường ---
            tk.Label(frm_about, text="Kiểm tra môi trường:", font=("Segoe UI", 9, "bold")).pack(
                anchor="w", padx=8, pady=(10, 2))
            env_row = tk.Frame(frm_about)
            env_row.pack(anchor="w", padx=8, pady=(0, 4), fill="x")
            self.lbl_env_word = tk.Label(env_row, text="Microsoft Word: đang kiểm tra...", fg="#888")
            self.lbl_env_word.pack(anchor="w")
            self.lbl_env_excel = tk.Label(env_row, text="Microsoft Excel: đang kiểm tra...", fg="#888")
            self.lbl_env_excel.pack(anchor="w")
            self.lbl_env_tesseract = tk.Label(env_row, text="Tesseract OCR: đang kiểm tra...", fg="#888")
            self.lbl_env_tesseract.pack(anchor="w")
            self.lbl_env_tessdata = tk.Label(env_row, text="Tessdata tiếng Việt: đang kiểm tra...", fg="#888")
            self.lbl_env_tessdata.pack(anchor="w")
            self.lbl_env_libs = tk.Label(env_row, text="Thư viện Excel/PDF/nén file: đang kiểm tra...", fg="#888")
            self.lbl_env_libs.pack(anchor="w")
            self.lbl_env_config = tk.Label(env_row, text="Cấu hình mã xã/thôn: đang kiểm tra...", fg="#888")
            self.lbl_env_config.pack(anchor="w")
            self.lbl_env_folders = tk.Label(env_row, text="Thư mục làm việc (Output/Logs/Temp/Backup/Undo): đang kiểm tra...",
                                             fg="#888")
            self.lbl_env_folders.pack(anchor="w")
            tk.Button(frm_about, text="🔍 Kiểm tra lại môi trường",
                      command=self.run_environment_check).pack(anchor="w", padx=8, pady=(4, 4))

            tk.Label(frm_about, text="Phần mềm thuộc Nguyễn Sỹ  |  SĐT/Zalo: 0972560335  |  Email: minhsybk@gmail.com",
                     fg="#777", justify="left", wraplength=1000).pack(anchor="w", padx=8, pady=(8, 8))

            self.after(300, self.run_environment_check)

        def run_environment_check(self):
            """Chạy kiểm tra môi trường trong luồng nền (không làm chậm khởi động), cập nhật nhãn trạng thái."""
            def worker():
                result = check_environment()

                def update_ui():
                    self.lbl_env_word.config(
                        text=f"Microsoft Word: {'Đã có' if result['word'] else 'Chưa có — chức năng chuyển Word sang PDF không khả dụng'}",
                        fg="#2e7d32" if result["word"] else "#c62828")
                    self.lbl_env_excel.config(
                        text=f"Microsoft Excel: {'Đã có' if result['excel'] else 'Chưa có — một số chức năng xử lý Excel .xls cũ có thể không hoạt động'}",
                        fg="#2e7d32" if result["excel"] else "#c62828")
                    if result["tesseract"] and result.get("tesseract_bundled"):
                        tess_text = "Tesseract OCR: Đã có (bản portable kèm theo phần mềm, không cần cài thêm)"
                    elif result["tesseract"]:
                        tess_text = "Tesseract OCR: Đã có (đã cài trên máy)"
                    else:
                        tess_text = "Tesseract OCR: Chưa có — chức năng OCR PDF scan không hoạt động"
                    self.lbl_env_tesseract.config(text=tess_text, fg="#2e7d32" if result["tesseract"] else "#c62828")
                    self.lbl_env_tessdata.config(
                        text=f"Tessdata tiếng Việt: {'Đã có' if result.get('tessdata_vie') else 'Chưa có/thiếu — OCR tiếng Việt có thể không chính xác'}",
                        fg="#2e7d32" if result.get("tessdata_vie") else "#c62828")
                    lib_parts = []
                    lib_parts.append("Excel: " + ("✓" if result.get("openpyxl") else "✗ THIẾU"))
                    lib_parts.append("PDF: " + ("✓" if result.get("pdf_library") else "✗ THIẾU"))
                    lib_parts.append(".zip: " + ("✓" if result.get("zip_support") else "✗"))
                    lib_parts.append(".rar: " + ("✓" if result.get("rar_support") else "chưa hỗ trợ"))
                    libs_ok = result.get("openpyxl") and result.get("pdf_library")
                    self.lbl_env_libs.config(text="Thư viện Excel/PDF/nén file: " + " | ".join(lib_parts),
                                             fg="#2e7d32" if libs_ok else "#c62828")
                    self.lbl_env_config.config(
                        text=f"Cấu hình mã xã/thôn: {'Đã có file riêng' if result['config_xa_thon'] else 'Chưa có (đang dùng bảng mặc định có sẵn)'}",
                        fg="#2e7d32" if result["config_xa_thon"] else "#f57c00")
                    n_folders = sum(1 for v in result["folders"].values() if v)
                    self.lbl_env_folders.config(
                        text=f"Thư mục làm việc (Output/Logs/Temp/Backup/Undo/Can_kiem_tra): "
                             f"đã tạo {n_folders}/{len(result['folders'])} tại {get_app_data_dir()}",
                        fg="#2e7d32" if n_folders == len(result["folders"]) else "#c62828")

                self.after(0, update_ui)

            threading.Thread(target=worker, daemon=True).start()

        def mark_workflow_step(self, key, status):
            """Cập nhật trạng thái 1 bước trong Tab Quy trình (Chưa chạy/Đã chạy thử/Đã hoàn thành/Có lỗi)."""
            colors = {"Chưa chạy": "#888", "Đã chạy thử": "#f57c00", "Đã hoàn thành": "#2e7d32", "Có lỗi": "#c62828"}
            lbl = self.workflow_status_labels.get(key)
            if lbl:
                lbl.config(text=status, fg=colors.get(status, "#888"))
            done_count = sum(1 for l in self.workflow_status_labels.values() if l.cget("text") == "Đã hoàn thành")
            self.workflow_progress["value"] = done_count


try:
    from app.ui.tab_content import ContentTabMixin
    _UI_CONTENT_MODULE_OK = True
except Exception:
    _UI_CONTENT_MODULE_OK = False

    class ContentTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_content.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_content(self, parent, pad):
            frm_all = styled_labelframe(parent, text="5. Đổi tên theo nội dung PDF (OCR cho file scan, hoặc đọc trực tiếp cho file có chữ)")
            frm_all.pack(fill="both", expand=True, **pad)

            def _section_title(text, fg="#1b3a2b", pady_top=10):
                tk.Label(frm_all, text=text, font=("Arial", 10, "bold"), fg=fg, anchor="w").pack(
                    fill="x", padx=8, pady=(pady_top, 2))
                ttk.Separator(frm_all, orient="horizontal").pack(fill="x", padx=8, pady=(0, 6))

            # --- ★ Chế độ OCR cho PDF scan/ảnh ---
            _section_title("★ CHẾ ĐỘ OCR CHO PDF SCAN/ẢNH — đổi tên hàng loạt (khuyến nghị cho file scan không có chữ)",
                            pady_top=6)
            tk.Label(frm_all,
                     text="Chuyên xử lý PDF là bản SCAN/ẢNH (không có text layer): render từng trang → ảnh độ phân giải cao "
                          "→ OCR tiếng Việt (vie+eng) → chuẩn hóa chữ → tìm mục 2.a (Thửa đất số, Tờ bản đồ số) và "
                          "mục 2.b (thôn/xã) → đặt tên CHUACOGIAY_[Mã xã]_[Số tờ]_[Số thửa]_GT.pdf.\n"
                          "Mặc định KHÔNG đổi tên ngay — xuất file CSV để bạn kiểm tra trước + file .bat chứa lệnh ren.",
                     fg="#33691e", justify="left", wraplength=1000).pack(anchor="w", padx=10, pady=(0, 6))

            self.ocr_picker = SourcePicker(
                frm_all, "Nguồn PDF scan (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF bất kỳ)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.ocr_picker.pack(fill="x", padx=10, pady=3)

            row_ocr_opt = tk.Frame(frm_all)
            row_ocr_opt.pack(fill="x", padx=10, pady=3)
            tk.Label(row_ocr_opt, text="Chế độ OCR:").pack(side="left")
            ttk.Combobox(row_ocr_opt, textvariable=self.var_ocr_mode, values=["Nhanh", "Chuẩn", "Kỹ"],
                         state="readonly", width=10).pack(side="left", padx=4)
            tk.Label(row_ocr_opt, text="   Hậu tố tên file:").pack(side="left")
            tk.Entry(row_ocr_opt, textvariable=self.var_ocr_suffix, width=8).pack(side="left")
            tk.Label(frm_all,
                     text="Nhanh: render 2.5x, không tiền xử lý ảnh, OCR psm 6 (nhanh nhất, dùng cho file scan rõ nét).\n"
                          "Chuẩn: render 3.5x, có tiền xử lý ảnh, thử psm 6+4 (cân bằng tốc độ/độ chính xác - khuyến nghị).\n"
                          "Kỹ: render 4x, thử nhiều mức threshold, psm 6+4+11 (chậm nhất, dùng cho file scan mờ/khó đọc).",
                     fg="#777", justify="left", wraplength=1000).pack(anchor="w", padx=10, pady=(0, 4))

            row_ocr_out = tk.Frame(frm_all)
            row_ocr_out.pack(fill="x", padx=10, pady=3)
            tk.Label(row_ocr_out, text="Thư mục lưu CSV + file .bat:", width=26, anchor="w").pack(side="left")
            tk.Entry(row_ocr_out, textvariable=self.var_ocr_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_ocr_out, text="Chọn...", command=self.pick_ocr_output).pack(side="left")
            tk.Label(frm_all, text="(Để trống → lưu cùng thư mục nguồn, hoặc thư mục chứa file đầu tiên nếu chọn file rải rác)",
                     fg="#777").pack(anchor="w", padx=10)

            row_ocr_debug = tk.Frame(frm_all)
            row_ocr_debug.pack(fill="x", padx=10, pady=3)
            tk.Checkbutton(row_ocr_debug,
                           text="Lưu ảnh debug (debug_page_N.png, debug_page_N_processed.png) + text OCR đầy đủ "
                                "(debug_ocr_text__*.txt) — để kiểm tra bằng mắt khi cần",
                           variable=self.var_ocr_debug).pack(side="left")
            self._register_advanced_widget(row_ocr_debug, fill="x", padx=10, pady=3)

            row_ocr_run = tk.Frame(frm_all)
            row_ocr_run.pack(fill="x", padx=10, pady=(6, 4))
            self.btn_ocr_preview = tk.Button(row_ocr_run, text="① KIỂM TRA TRƯỚC (xuất CSV + .bat, KHÔNG đổi tên)",
                                              font=("Arial", 10, "bold"), bg="#00695c", fg="white",
                                              command=lambda: self.run_ocr_rename_clicked(do_rename=False))
            self.btn_ocr_preview.pack(side="left", padx=(0, 6))
            self.btn_ocr_do = tk.Button(row_ocr_run, text="② ĐỔI TÊN NGAY (sau khi đã kiểm tra CSV)",
                                         font=("Arial", 10, "bold"), bg="#bf360c", fg="white",
                                         command=lambda: self.run_ocr_rename_clicked(do_rename=True))
            self.btn_ocr_do.pack(side="left")
            self.btn_ocr_pause, self.btn_ocr_cancel = self._add_pause_cancel(row_ocr_run, "control_ocr")
            self.btn_ocr_pause.pack(side="left", padx=(12, 3))
            self.btn_ocr_cancel.pack(side="left", padx=3)
            self.btn_ocr_undo = tk.Button(row_ocr_run, text="↩ Hoàn tác lần đổi tên gần nhất",
                                           bg="#795548", fg="white", state="disabled",
                                           command=self.undo_ocr_rename_clicked)
            self.btn_ocr_undo.pack(side="left", padx=(12, 3))

            # --- Thanh tiến độ trực quan ---
            frm_ocr_progress = tk.Frame(frm_all)
            frm_ocr_progress.pack(fill="x", padx=10, pady=(0, 6))
            self.pb_ocr = ttk.Progressbar(frm_ocr_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_ocr.pack(side="left", padx=(0, 10))
            self.lbl_ocr_progress = tk.Label(frm_ocr_progress, text="", fg="#555", anchor="w")
            self.lbl_ocr_progress.pack(side="left", fill="x", expand=True)

            # --- Đọc nội dung PDF có sẵn chữ (mẫu tên tùy chỉnh) ---
            _section_title("Đọc nội dung PDF có sẵn chữ (không phải scan) — mẫu tên tùy chỉnh tự do")

            tk.Label(frm_all,
                     text="Lấy 'Thửa đất số', 'Tờ bản đồ số', 'Địa chỉ' (suy ra Tên xã) ngay trong nội dung file PDF.\n"
                          "Nếu file là bản scan (không có chữ để đọc), phần mềm sẽ tự OCR trang đầu.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=10, pady=(0, 6))

            row_ocr = tk.Frame(frm_all)
            row_ocr.pack(fill="x", padx=10, pady=3)
            tk.Checkbutton(row_ocr, text="Tự động OCR nếu file là bản scan (cần cài Tesseract-OCR)",
                           variable=self.var_use_ocr).pack(side="left")
            tk.Label(row_ocr, text="   Độ phân giải OCR (dpi):").pack(side="left")
            tk.Entry(row_ocr, textvariable=self.var_ocr_dpi, width=6).pack(side="left")

            row_exec = tk.Frame(frm_all)
            row_exec.pack(fill="x", padx=10, pady=3)
            tk.Checkbutton(row_exec, text="Đổi tên NGAY sau khi quét (không cần tự chạy file .bat) — vẫn tạo file .bat để lưu lại lịch sử",
                           variable=self.var_ren_execute_now).pack(side="left")

            row_debug = tk.Frame(frm_all)
            row_debug.pack(fill="x", padx=10, pady=3)
            tk.Checkbutton(row_debug,
                           text="Hiện văn bản đã đọc được khi LỖI (để chẩn đoán vì sao 1 file không đổi tên được)",
                           variable=self.var_debug_content).pack(side="left")
            self._register_advanced_widget(row_debug, fill="x", padx=10, pady=3)

            row_tpl = tk.Frame(frm_all)
            row_tpl.pack(fill="x", padx=10, pady=(6, 3))
            tk.Label(row_tpl, text="Mẫu tên mới:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_tpl, textvariable=self.var_content_name_template).pack(side="left", fill="x", expand=True, padx=4)

            row_tpl_hint = tk.Frame(frm_all)
            row_tpl_hint.pack(fill="x", padx=10, pady=(0, 3))
            tk.Label(row_tpl_hint,
                     text="Placeholder: {maxa} {to} {thua} {loai} {ten} {dt} {id} {mucdich} {core} {year} {original} "
                          "— {loai} tự nhận GT/TBXN/KHAC, các phần còn lại bạn tự thêm bớt tùy ý.",
                     fg="#555", wraplength=1000, justify="left").pack(anchor="w")

            row_tpl_presets = tk.Frame(frm_all)
            row_tpl_presets.pack(fill="x", padx=10, pady=(0, 6))
            tk.Label(row_tpl_presets, text="Mẫu có sẵn:", width=14, anchor="w").pack(side="left")
            tk.Button(row_tpl_presets, text="CHUACOGIAY_xã_tờ_thửa_loại",
                      command=lambda: self.var_content_name_template.set("CHUACOGIAY_{maxa}_{to}_{thua}_{loai}")
                      ).pack(side="left", padx=3)
            tk.Button(row_tpl_presets, text="...+ CCCD",
                      command=lambda: self.var_content_name_template.set("CHUACOGIAY_{maxa}_{to}_{thua}_{loai}_{id}")
                      ).pack(side="left", padx=3)
            tk.Button(row_tpl_presets, text="Tên chủ_tờ_thửa_loại",
                      command=lambda: self.var_content_name_template.set("{ten}_To{to}_Thua{thua}_{loai}")
                      ).pack(side="left", padx=3)

            self.pdfcontent_picker = SourcePicker(
                frm_all, "Nguồn PDF cần đổi tên (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF bất kỳ)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.pdfcontent_picker.pack(fill="x", padx=10, pady=3)

            frm_run5 = tk.Frame(frm_all)
            frm_run5.pack(fill="x", padx=10, pady=(4, 4))
            self.btn_content_rename = tk.Button(frm_run5, text="🔍 QUÉT NỘI DUNG & TẠO LỆNH ĐỔI TÊN",
                                                 font=("Arial", 11, "bold"), bg="#e65100", fg="white",
                                                 command=self.run_content_rename_clicked)
            self.btn_content_rename.pack(side="left", padx=4)
            self.btn_content_pause, self.btn_content_cancel = self._add_pause_cancel(frm_run5, "control_content")
            self.btn_content_pause.pack(side="left", padx=(8, 3))
            self.btn_content_cancel.pack(side="left", padx=3)
            tk.Label(frm_run5,
                     text="Kết quả: file DOI_TEN_HANG_LOAT.bat (lệnh ren) được tạo trong chính thư mục đó.",
                     fg="#555").pack(side="left", padx=10)

            # --- Bảng tra Thôn/Xã -> Mã xã (dùng chung cho cả 2 chế độ ở trên) ---
            _section_title("Bảng tra Thôn / Xã → Mã xã (dùng chung cho cả 2 chế độ phía trên)")

            row_cfg = tk.Frame(frm_all)
            row_cfg.pack(fill="x", padx=10, pady=(0, 6))
            tk.Label(row_cfg,
                     text=f"Có thể tách bảng tra ra file cấu hình riêng ({CONFIG_XA_THON_FILENAME}, đặt cạnh phần mềm) "
                          "để bổ sung địa bàn mới mà KHÔNG cần sửa phần mềm.",
                     fg="#555", wraplength=800, justify="left").pack(anchor="w")
            row_cfg_btn = tk.Frame(frm_all)
            row_cfg_btn.pack(fill="x", padx=10, pady=(0, 8))
            tk.Button(row_cfg_btn, text="↻ Nạp lại cấu hình", command=self.reload_xa_thon_config).pack(side="left", padx=(0, 6))
            tk.Button(row_cfg_btn, text="📂 Mở file cấu hình", command=self.open_xa_thon_config_file).pack(side="left", padx=(0, 6))
            tk.Button(row_cfg_btn, text="+ Tạo file cấu hình mẫu", command=self.create_xa_thon_config_sample).pack(side="left")
            self.lbl_xa_thon_config_status = tk.Label(row_cfg_btn, text="", fg="#777")
            self.lbl_xa_thon_config_status.pack(side="left", padx=12)

            row_maps = tk.Frame(frm_all)
            row_maps.pack(fill="x", padx=10, pady=(0, 8))

            col_thon = tk.Frame(row_maps)
            col_thon.pack(side="left", fill="both", expand=True, padx=(0, 8))
            tk.Label(col_thon, text="Tên THÔN → Mã xã (ƯU TIÊN cao hơn, mỗi dòng: Tên thôn=Mã xã)",
                     anchor="w", wraplength=480).pack(anchor="w")
            tk.Label(col_thon,
                     text="Dùng khi PDF ghi tên xã cũ (trước sáp nhập) nhưng tên thôn vẫn giữ nguyên.",
                     fg="#777", anchor="w", wraplength=480, justify="left").pack(anchor="w", pady=(0, 3))
            self.txt_thon_mapping = tk.Text(col_thon, height=8, font=("Consolas", 9))
            self.txt_thon_mapping.pack(fill="both", expand=True)
            self.txt_thon_mapping.insert("1.0", DEFAULT_THON_MAPPING_TEXT)

            col_xa = tk.Frame(row_maps)
            col_xa.pack(side="left", fill="both", expand=True, padx=(8, 0))
            tk.Label(col_xa, text="Tên XÃ → Mã xã (dự phòng, mỗi dòng: Tên xã=Mã xã)",
                     anchor="w", wraplength=480).pack(anchor="w")
            tk.Label(col_xa, text="Chỉ dùng khi không khớp thôn nào ở bảng bên trái.",
                     fg="#777", anchor="w", wraplength=480, justify="left").pack(anchor="w", pady=(0, 3))
            self.txt_xa_mapping = tk.Text(col_xa, height=8, font=("Consolas", 9))
            self.txt_xa_mapping.pack(fill="both", expand=True)
            self.txt_xa_mapping.insert("1.0", DEFAULT_XA_MAPPING_TEXT)

            # Tự động nạp cấu hình từ file Excel bên ngoài nếu có sẵn (không có thì giữ mặc định vừa chèn ở trên)
            self.after(200, self.reload_xa_thon_config, True)

        def reload_xa_thon_config(self, silent=False):
            """Đọc lại Cau_hinh_ma_xa_thon.xlsx (nếu có) và nạp vào 2 bảng tra Thôn/Xã. silent=True: không hiện popup khi không có file (dùng lúc khởi động)."""
            path = get_xa_thon_config_path()
            thon_mapping, xa_mapping = load_xa_thon_config_from_excel(path, log_cb=self.log if hasattr(self, "txt_log") else None)
            if thon_mapping is None and xa_mapping is None:
                msg = f"Không tìm thấy file cấu hình ({CONFIG_XA_THON_FILENAME}) - đang dùng bảng tra mặc định có sẵn."
                self.lbl_xa_thon_config_status.config(text="Đang dùng cấu hình MẶC ĐỊNH (chưa có file ngoài)", fg="#777")
                if not silent:
                    messagebox.showinfo("Không có file cấu hình", msg + f"\n\nBấm '+ Tạo file cấu hình mẫu' để tạo file tại:\n{path}")
                return

            def dict_to_lines(d):
                return "\n".join(f"{k}={v}" for k, v in d.items())

            self.txt_thon_mapping.delete("1.0", "end")
            self.txt_thon_mapping.insert("1.0", dict_to_lines(thon_mapping))
            self.txt_xa_mapping.delete("1.0", "end")
            self.txt_xa_mapping.insert("1.0", dict_to_lines(xa_mapping))
            self.lbl_xa_thon_config_status.config(
                text=f"✓ Đã nạp từ file cấu hình ({len(thon_mapping)} thôn, {len(xa_mapping)} xã cũ)", fg="#2e7d32")
            if not silent:
                messagebox.showinfo("Đã nạp cấu hình", f"Đã nạp {len(thon_mapping)} thôn, {len(xa_mapping)} xã cũ từ:\n{path}")

        def open_xa_thon_config_file(self):
            path = get_xa_thon_config_path()
            if not os.path.isfile(path):
                if not messagebox.askyesno("Chưa có file cấu hình",
                                            f"Chưa có file {CONFIG_XA_THON_FILENAME}. Tạo file mẫu ngay bây giờ?"):
                    return
                self.create_xa_thon_config_sample()
            try:
                if sys.platform == "win32":
                    os.startfile(path)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", path])
                else:
                    subprocess.Popen(["xdg-open", path])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được file: {e}")

        def create_xa_thon_config_sample(self):
            path = get_xa_thon_config_path()
            if os.path.isfile(path):
                if not messagebox.askyesno("File đã tồn tại",
                                            f"File {CONFIG_XA_THON_FILENAME} đã tồn tại. Ghi đè bằng file mẫu mới?"):
                    return
            try:
                create_default_xa_thon_config_file(path)
                messagebox.showinfo("Đã tạo file cấu hình mẫu",
                                     f"Đã tạo tại:\n{path}\n\nBạn có thể mở file này bằng Excel để bổ sung/sửa "
                                     f"địa bàn, sau đó bấm 'Nạp lại cấu hình'.")
                self.reload_xa_thon_config()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không tạo được file cấu hình: {e}")
        def pick_ocr_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu CSV + file .bat")
            if p:
                self.var_ocr_output.set(p)

        def run_ocr_rename_clicked(self, do_rename=False):
            source = self.ocr_picker.get_source()
            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF scan cần đổi tên.")
                return
            if not (HAS_FITZ and HAS_OCR):
                messagebox.showerror(
                    "Thiếu thư viện OCR",
                    "Chế độ này BẮT BUỘC cần OCR. Vui lòng cài:\n"
                    "  pip install pymupdf pytesseract pillow\n"
                    "và cài chương trình Tesseract-OCR (kèm gói ngôn ngữ Vietnamese - vie).")
                return

            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            if not thon_mapping and not xa_mapping:
                messagebox.showerror("Thiếu bảng tra", "Bảng tra Thôn/Xã → Mã xã (mục 5b/5c) đang trống.")
                return

            mode = self.var_ocr_mode.get().strip() or "Chuẩn"
            suffix = self.var_ocr_suffix.get().strip() or "GT"

            if do_rename:
                if not self._check_require_dry_run("_dry_run_done_ocr", True, "Đổi tên theo nội dung PDF (Tab 5)"):
                    return
                if not messagebox.askyesno(
                        "Xác nhận đổi tên",
                        "Bạn đã KIỂM TRA file CSV và chắc chắn kết quả đúng?\n\n"
                        "Thao tác này sẽ ĐỔI TÊN THẬT các file PDF. Những file không nhận diện đủ "
                        "Mã xã/Số tờ/Số thửa sẽ được giữ nguyên (không đổi)."):
                    return

            # Xác định thư mục lưu CSV/.bat
            out_dir = self.var_ocr_output.get().strip()
            if not out_dir:
                if isinstance(source, str):
                    out_dir = source
                else:
                    files_all = list_files_from_source(source, ".pdf")
                    out_dir = os.path.dirname(files_all[0]) if files_all else os.getcwd()

            self.txt_log.delete("1.0", "end")

            checkpoint_task_id = None
            files_da_xong = None
            if do_rename:
                from app.services import checkpoint_service as _ckpt
                checkpoint_task_id = _ckpt.tinh_task_id("ocr_doi_ten_pdf", str(source), "")
                data_cu = _ckpt.doc_checkpoint(checkpoint_task_id)
                if data_cu and data_cu.get("trang_thai") == "DANG_XU_LY":
                    so_da_xong = len(data_cu.get("da_hoan_thanh", []))
                    so_tong = len(data_cu.get("danh_sach_file", []))
                    if messagebox.askyesno(
                        "Phát hiện tác vụ dở dang",
                        f"Phát hiện đợt đổi tên OCR TRƯỚC ĐÓ chưa hoàn thành (đã xong {so_da_xong}/{so_tong} "
                        f"file, có thể do phần mềm bị đóng đột ngột).\n\n"
                        f"Bấm CÓ để TIẾP TỤC từ chỗ dở dang.\nBấm KHÔNG để BẮT ĐẦU LẠI TỪ ĐẦU."):
                        files_da_xong = set(data_cu.get("da_hoan_thanh", []))
                    else:
                        _ckpt.xoa_checkpoint(checkpoint_task_id)

            self.btn_ocr_preview.config(state="disabled")
            self.btn_ocr_do.config(state="disabled")
            self._enable_pause_cancel("control_ocr", self.btn_ocr_pause, self.btn_ocr_cancel)
            self.pb_ocr.config(value=0, maximum=100)
            self.lbl_ocr_progress.config(text="Đang chuẩn bị...")
            self._ocr_progress_start_time = time.time()

            from app.services import task_manager_service as _tm3
            self.task_info_ocr = _tm3.TaskInfo(
                ten_chuc_nang="5. Đổi tên theo nội dung PDF (OCR)", thu_muc_dau_ra=out_dir or "",
                trang_thai=_tm3.TrangThaiTask.DANG_KHOI_TAO)

            def update_progress_ui(idx, total, fname):
                elapsed = time.time() - self._ocr_progress_start_time
                if idx > 1:
                    avg_per_file = elapsed / (idx - 1)
                    remaining = avg_per_file * (total - idx + 1)
                    eta_text = f" | Còn khoảng {int(remaining // 60)} phút {int(remaining % 60)} giây"
                else:
                    eta_text = ""
                text = f"Đang xử lý {idx}/{total}: {fname}  (đã chạy {int(elapsed)}s{eta_text})"
                self.pb_ocr.config(value=idx, maximum=total)
                self.lbl_ocr_progress.config(text=text)

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_progress_ui(idx, total, fname))
                self.task_info_ocr.cap_nhat_tien_do(tien_do_toan_bo=idx, trang_thai=_tm3.TrangThaiTask.DANG_OCR)
                self.task_info_ocr.tong_so_luong = total

            def worker():
                try:
                    self.log(f"Bắt đầu OCR chế độ '{mode}' (có thể mất thời gian tùy số lượng file)...")
                    debug_dir = out_dir if self.var_ocr_debug.get() else None
                    if debug_dir:
                        self.log(f"Chế độ debug: sẽ lưu ảnh + text OCR đầy đủ vào {debug_dir}")
                    if do_rename:
                        files_to_backup = list_files_from_source(source, ".pdf")
                        self._maybe_auto_backup(files_to_backup, "Đổi tên theo nội dung PDF (Tab 5)")
                    perf_log_path = None
                    if getattr(self.app_config, "word_hien_thi_thong_ke_hieu_nang", True):
                        perf_log_path = os.path.join(get_standard_logs_dir(), "LOG_HIEU_NANG_XU_LY.csv")
                    result = ocr_rename_scan_pdfs(
                        source, thon_mapping, xa_mapping, None, self.log,
                        name_suffix=suffix, do_rename=do_rename, debug_dir=debug_dir, control=self.control_ocr,
                        mode=mode, progress_cb=progress_cb,
                        toc_do_xu_ly=getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"),
                        so_file_moi_dot=getattr(self.app_config, "word_so_file_moi_dot", 100),
                        nghi_giua_dot_giay=getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0),
                        perf_log_path=perf_log_path,
                        checkpoint_task_id=checkpoint_task_id, files_da_xong=files_da_xong)
                    self.after(0, lambda: self.lbl_ocr_progress.config(
                        text=f"✓ Hoàn tất — đã xử lý {len(result['ok']) + len(result['need_check'])} file "
                             f"({int(time.time() - self._ocr_progress_start_time)}s)"))
                    self.task_info_ocr.so_luong_thanh_cong = len(result["ok"])
                    self.task_info_ocr.so_luong_can_kiem_tra = len(result["need_check"])
                    self.task_info_ocr.trang_thai = _tm3.TrangThaiTask.THANH_CONG
                    if not do_rename:
                        self._dry_run_done_ocr = True

                    import time
                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    csv_path = os.path.join(out_dir, f"KIEMTRA_DOITEN_{stamp}.csv")
                    cmd_path = os.path.join(out_dir, f"DOITEN_{stamp}.bat")
                    write_ocr_rename_csv(csv_path, result)
                    write_ocr_rename_cmd(cmd_path, result)

                    # Ghi thêm log CHUẨN (Logs/<run_id>_DOI_TEN_PDF.csv) song song với log đặc thù ở trên,
                    # dùng CÙNG run_id (stamp) để dễ đối chiếu 2 file log với nhau.
                    try:
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "DOI_TEN_PDF", run_id=stamp)
                        for r in result["ok"]:
                            std_log.add(action="RENAME", source_file=r["old_path"], output_file=r.get("new_name", ""),
                                       ma_xa=r.get("maxa"), so_to=r.get("to"), so_thua=r.get("thua"),
                                       key=f"{r.get('maxa')}_{r.get('to')}_{r.get('thua')}",
                                       status=r["status"], message=r.get("note", ""))
                        for r in result["need_check"]:
                            std_log.add(action="RENAME", source_file=r.get("old_name", ""), status=r["status"],
                                       message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {len(result['ok'])} file nhận diện đủ | "
                             f"{len(result['need_check'])} file CẦN KIỂM TRA (không đổi tên)")
                    self.log(f"File CSV kiểm tra: {csv_path}")
                    self.log(f"File lệnh .bat:   {cmd_path}")
                    if result["need_check"]:
                        self.log("\nDanh sách CẦN KIỂM TRA:")
                        for r in result["need_check"]:
                            self.log(f"  - {r['old_name']}: {r['reason']}")

                    if do_rename:
                        undo_path = os.path.join(out_dir, f"UNDO_{stamp}.cmd")
                        write_undo_rename_cmd(undo_path, result)
                        self.log(f"File hoàn tác (chạy tay nếu cần): {undo_path}")
                        self.last_ocr_rename_result = result
                        self.after(0, lambda: self.btn_ocr_undo.config(state="normal"))
                        step_status = "Đã hoàn thành" if not result["need_check"] else "Có lỗi"
                        self.after(0, lambda: self.mark_workflow_step("doi_ten", step_status))
                        self.after(0, lambda: self.mark_workflow_step("ocr", "Đã hoàn thành"))
                        msg = (f"Đã đổi tên: {len(result['ok'])} file\n"
                               f"Cần kiểm tra (giữ nguyên): {len(result['need_check'])} file")
                    else:
                        self.after(0, lambda: self.mark_workflow_step("doi_ten", "Đã chạy thử"))
                        self.after(0, lambda: self.mark_workflow_step("ocr", "Đã chạy thử"))
                        msg = (f"Nhận diện đủ: {len(result['ok'])} file\n"
                               f"Cần kiểm tra: {len(result['need_check'])} file\n\n"
                               f"Đã xuất CSV + .bat để bạn kiểm tra trước khi đổi tên.")
                    messagebox.showinfo("Hoàn tất OCR", msg + f"\n\nLưu tại:\n{out_dir}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_ocr_preview.config(state="normal")
                    self.btn_ocr_do.config(state="normal")
                    self._disable_pause_cancel(self.btn_ocr_pause, self.btn_ocr_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def undo_ocr_rename_clicked(self):
            if not self.last_ocr_rename_result:
                messagebox.showinfo("Không có gì để hoàn tác", "Chưa có lần đổi tên thật nào trong phiên làm việc này.")
                return
            n_renamed = sum(1 for r in self.last_ocr_rename_result["ok"] if r.get("status") == "DA_DOI_TEN")
            if not messagebox.askyesno(
                    "Xác nhận hoàn tác",
                    f"Sẽ đưa {n_renamed} file vừa đổi tên TRỞ VỀ tên gốc ban đầu. Chỉ hoàn tác được nếu file "
                    f"chưa bị di chuyển/đổi tên thêm lần nữa sau đó. Tiếp tục?"):
                return

            def worker():
                undone, errors = undo_rename_in_app(self.last_ocr_rename_result, self.log)
                self.log("\n" + "=" * 60)
                self.log(f"TỔNG KẾT HOÀN TÁC: {undone} file đã đưa về tên gốc | {len(errors)} lỗi")
                for e in errors:
                    self.log(f"  ✗ {e}")
                msg = f"Đã hoàn tác: {undone} file"
                if errors:
                    msg += f"\nLỗi: {len(errors)} file (xem chi tiết trong Nhật ký)"
                messagebox.showinfo("Hoàn tất hoàn tác", msg)
                self.last_ocr_rename_result = None
                self.after(0, lambda: self.btn_ocr_undo.config(state="disabled"))

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 5: Đổi tên theo nội dung PDF -------------------

        def run_content_rename_clicked(self):
            source = self.pdfcontent_picker.get_source()
            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF cần đổi tên.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return

            mapping_text = self.txt_xa_mapping.get("1.0", "end")
            xa_mapping = parse_xa_mapping(mapping_text)
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            if not xa_mapping and not thon_mapping:
                messagebox.showerror("Thiếu bảng tra", "Bảng tra Thôn/Xã → Mã xã đang trống. Vui lòng nhập ít nhất 1 dòng.")
                return

            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            if use_ocr and not (HAS_FITZ and HAS_OCR):
                if not messagebox.askyesno(
                        "Thiếu thư viện OCR",
                        "Chưa cài đủ thư viện OCR (pymupdf, pytesseract, pillow) hoặc chưa cài chương trình "
                        "Tesseract-OCR.\nBạn có muốn tiếp tục mà KHÔNG dùng OCR không? "
                        "(các file scan sẽ báo lỗi thiếu thông tin)"):
                    return
                use_ocr = False

            self.txt_log.delete("1.0", "end")
            self.btn_content_rename.config(state="disabled")
            self._enable_pause_cancel("control_content", self.btn_content_pause, self.btn_content_cancel)

            def worker():
                try:
                    name_template = self.var_content_name_template.get().strip() or "CHUACOGIAY_{maxa}_{to}_{thua}_{loai}"
                    ok_list, err_list = build_rename_commands_from_pdf_content(
                        source, xa_mapping, thon_mapping, use_ocr, ocr_dpi, self.log,
                        debug=self.var_debug_content.get(), name_template=name_template, control=self.control_content)

                    # Nơi lưu file lệnh .bat: thư mục nguồn (nếu chọn thư mục), hoặc thư mục chứa file ĐẦU
                    # TIÊN trong danh sách (nếu chọn nhiều file cụ thể ở nhiều nơi khác nhau)
                    if isinstance(source, str):
                        bat_dir = source
                    elif ok_list:
                        bat_dir = os.path.dirname(ok_list[0][0])
                    else:
                        files_all = list_files_from_source(source, ".pdf")
                        bat_dir = os.path.dirname(files_all[0]) if files_all else os.getcwd()
                    bat_path = os.path.join(bat_dir, "DOI_TEN_HANG_LOAT.bat")
                    write_ren_bat_file(ok_list, bat_path)

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {len(ok_list)} file đọc được đủ thông tin | {len(err_list)} file lỗi/thiếu thông tin")
                    self.log(f"Đã tạo file lệnh: {bat_path}")
                    if err_list:
                        self.log("\nDanh sách file LỖI (cần kiểm tra tay):")
                        for fname, reason in err_list:
                            self.log(f"  - {fname}: {reason}")

                    executed_msg = ""
                    if self.var_ren_execute_now.get() and ok_list:
                        self.log("\nĐang đổi tên ngay theo yêu cầu...")
                        for old_full, new_name in ok_list:
                            try:
                                new_full = os.path.join(os.path.dirname(old_full), new_name)
                                os.rename(old_full, new_full)
                                self.log(f"✎ Đã đổi: {os.path.basename(old_full)} → {new_name}")
                            except Exception as e:
                                self.log(f"✗ Lỗi đổi tên {os.path.basename(old_full)}: {e}")
                        executed_msg = "\n\nĐã đổi tên trực tiếp theo yêu cầu."

                    messagebox.showinfo(
                        "Hoàn tất quét nội dung PDF",
                        f"Đọc được: {len(ok_list)} file\nLỗi/thiếu thông tin: {len(err_list)} file\n\n"
                        f"File lệnh ren: {bat_path}{executed_msg}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_content_rename.config(state="normal")
                    self._disable_pause_cancel(self.btn_content_pause, self.btn_content_cancel)

            threading.Thread(target=worker, daemon=True).start()


try:
    from app.ui.tab_pdfedit import PdfEditTabMixin
    _UI_PDFEDIT_MODULE_OK = True
except Exception:
    _UI_PDFEDIT_MODULE_OK = False

    class PdfEditTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_pdfedit.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_pdfedit(self, parent, pad):
            frm = styled_labelframe(parent, text="NHÓM: Chỉnh sửa / Xoay / Xem trước / Cải thiện chất lượng PDF")
            frm.pack(fill="both", expand=True, **pad)

            row1 = tk.Frame(frm)
            row1.pack(fill="x", padx=6, pady=3)
            tk.Button(row1, text="+ Chọn file PDF...", command=self.pdfedit_add_files).pack(side="left", padx=2)
            tk.Button(row1, text="+ Chọn thư mục PDF...", command=self.pdfedit_add_folder).pack(side="left", padx=2)
            tk.Checkbutton(row1, text="Quét cả thư mục con", variable=self.var_pdfedit_subfolders).pack(side="left", padx=8)
            tk.Button(row1, text="Xóa danh sách", command=self.pdfedit_clear_files).pack(side="left", padx=2)

            row2 = tk.Frame(frm)
            row2.pack(fill="x", padx=6, pady=3)
            tk.Label(row2, text="Thư mục lưu kết quả:", width=20, anchor="w").pack(side="left")
            tk.Entry(row2, textvariable=self.var_pdfedit_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row2, text="Chọn...", command=self.pdfedit_pick_output).pack(side="left")

            main_area = tk.Frame(frm)
            main_area.pack(fill="both", expand=True, padx=6, pady=6)

            left_col = tk.Frame(main_area)
            left_col.pack(side="left", fill="y", padx=(0, 8))
            tk.Label(left_col, text="Danh sách file:", font=("Arial", 9, "bold")).pack(anchor="w")
            self.lst_pdfedit_files = tk.Listbox(left_col, width=34, height=16, exportselection=False)
            self.lst_pdfedit_files.pack(fill="y", expand=False)
            self.lst_pdfedit_files.bind("<<ListboxSelect>>", self.pdfedit_on_file_selected)

            right_col = tk.Frame(main_area)
            right_col.pack(side="left", fill="both", expand=True)

            nav_row = tk.Frame(right_col)
            nav_row.pack(fill="x")
            tk.Button(nav_row, text="◀ Trang trước", command=self.pdfedit_prev_page).pack(side="left")
            self.var_pdfedit_pageno = tk.StringVar(value="0")
            tk.Entry(nav_row, textvariable=self.var_pdfedit_pageno, width=5, justify="center").pack(side="left", padx=4)
            self.lbl_pdfedit_total = tk.Label(nav_row, text="/ 0")
            self.lbl_pdfedit_total.pack(side="left")
            tk.Button(nav_row, text="Đi tới", command=self.pdfedit_goto_page).pack(side="left", padx=4)
            tk.Button(nav_row, text="Trang sau ▶", command=self.pdfedit_next_page).pack(side="left", padx=(4, 16))
            tk.Button(nav_row, text="Vừa khung", command=lambda: self.pdfedit_set_zoom(1.0)).pack(side="left", padx=2)
            tk.Button(nav_row, text="Phóng to +", command=lambda: self.pdfedit_zoom_delta(0.25)).pack(side="left", padx=2)
            tk.Button(nav_row, text="Thu nhỏ -", command=lambda: self.pdfedit_zoom_delta(-0.25)).pack(side="left", padx=2)

            canvas_frame = tk.Frame(right_col, bg="#cfcfcf", height=420)
            canvas_frame.pack(fill="both", expand=True, pady=4)
            canvas_frame.pack_propagate(False)
            self.canvas_pdfedit = tk.Canvas(canvas_frame, bg="#cfcfcf", highlightthickness=0)
            self.canvas_pdfedit.pack(fill="both", expand=True)

            rot_row = tk.Frame(right_col)
            rot_row.pack(fill="x", pady=(4, 2))
            tk.Label(rot_row, text="Xoay trang đang xem:", font=("Arial", 9, "bold")).pack(side="left")
            tk.Button(rot_row, text="↺ Trái 90°", command=lambda: self.pdfedit_rotate_current(-90)).pack(side="left", padx=3)
            tk.Button(rot_row, text="↻ Phải 90°", command=lambda: self.pdfedit_rotate_current(90)).pack(side="left", padx=3)
            tk.Button(rot_row, text="⟲ 180°", command=lambda: self.pdfedit_rotate_current(180)).pack(side="left", padx=3)
            tk.Button(rot_row, text="Về góc gốc", command=self.pdfedit_reset_current).pack(side="left", padx=3)
            self.lbl_pdfedit_current_rot = tk.Label(rot_row, text="Đang xoay thêm: 0°", fg="#555")
            self.lbl_pdfedit_current_rot.pack(side="left", padx=12)

            tk.Label(right_col, text="Áp dụng góc xoay của trang đang xem cho:",
                     font=("Arial", 9, "bold")).pack(anchor="w", pady=(4, 0))
            apply_btn_row = tk.Frame(right_col)
            apply_btn_row.pack(fill="x", pady=(0, 4))
            tk.Button(apply_btn_row, text="Cả file đang xem",
                      command=lambda: self.pdfedit_apply_scope("file")).pack(side="left", padx=2)
            tk.Button(apply_btn_row, text="Toàn bộ danh sách file",
                      command=lambda: self.pdfedit_apply_scope("all_files")).pack(side="left", padx=2)
            tk.Button(apply_btn_row, text="Trang chẵn",
                      command=lambda: self.pdfedit_apply_scope("even")).pack(side="left", padx=2)
            tk.Button(apply_btn_row, text="Trang lẻ",
                      command=lambda: self.pdfedit_apply_scope("odd")).pack(side="left", padx=2)
            tk.Entry(apply_btn_row, textvariable=self.var_pdfedit_range, width=12).pack(side="left", padx=(10, 2))
            tk.Button(apply_btn_row, text="Áp dụng khoảng trang (VD 1-5,8)",
                      command=lambda: self.pdfedit_apply_scope("range")).pack(side="left", padx=2)

            auto_row = tk.Frame(frm)
            auto_row.pack(fill="x", padx=6, pady=(2, 3))
            tk.Button(auto_row, text="🔍 Tự phát hiện hướng (file đang xem)",
                      command=lambda: self.pdfedit_autodetect("file")).pack(side="left", padx=2)
            tk.Button(auto_row, text="🔍 Tự phát hiện hướng (toàn bộ danh sách)",
                      command=lambda: self.pdfedit_autodetect("all")).pack(side="left", padx=2)
            tk.Label(auto_row, text="(OCR nhanh 4 hướng, chỉ đề xuất khi đủ tin cậy, còn lại đưa vào CẦN KIỂM TRA)",
                     fg="#777").pack(side="left", padx=10)

            enh_frame = styled_labelframe(
                frm, text="Cải thiện ảnh PDF scan (tùy chọn, áp dụng khi lưu — LƯU Ý: sẽ làm mất lớp chữ gốc nếu file có)")
            enh_frame.pack(fill="x", padx=6, pady=6)
            self._register_advanced_widget(enh_frame, fill="x", padx=6, pady=6)
            row_e1 = tk.Frame(enh_frame)
            row_e1.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_e1, text="Bật cải thiện ảnh (nếu tắt, chỉ xoay trang, giữ nguyên chất lượng gốc)",
                           variable=self.var_pdfedit_enhance_on).pack(side="left")
            row_e2 = tk.Frame(enh_frame)
            row_e2.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_e2, text="Chuyển ảnh xám", variable=self.var_pdfedit_grayscale).pack(side="left", padx=6)
            tk.Checkbutton(row_e2, text="Tăng tương phản, mức:", variable=self.var_pdfedit_contrast_on).pack(side="left", padx=6)
            tk.Entry(row_e2, textvariable=self.var_pdfedit_contrast_val, width=5).pack(side="left")
            tk.Checkbutton(row_e2, text="Làm nét nhẹ", variable=self.var_pdfedit_sharpen).pack(side="left", padx=6)
            tk.Checkbutton(row_e2, text="Khử nhiễu nhẹ", variable=self.var_pdfedit_denoise).pack(side="left", padx=6)
            row_e3 = tk.Frame(enh_frame)
            row_e3.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_e3, text="Nhị phân hóa đen-trắng, ngưỡng:",
                           variable=self.var_pdfedit_threshold_on).pack(side="left", padx=6)
            tk.Entry(row_e3, textvariable=self.var_pdfedit_threshold_val, width=5).pack(side="left")
            deskew_label = "Tự sửa nghiêng nhẹ" + ("" if HAS_CV2 else " (⚠ chưa cài OpenCV trên máy này - sẽ bỏ qua)")
            tk.Checkbutton(row_e3, text=deskew_label, variable=self.var_pdfedit_deskew).pack(side="left", padx=6)
            tk.Checkbutton(row_e3, text="OCR lại sau xử lý (tạo bản có thể tìm kiếm)",
                           variable=self.var_pdfedit_ocr_after).pack(side="left", padx=6)

            row_e4 = tk.Frame(enh_frame)
            row_e4.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_e4, text="Cắt viền trắng/đen thừa quanh trang",
                           variable=self.var_pdfedit_crop_border).pack(side="left", padx=6)
            tk.Checkbutton(row_e4, text="Chuẩn hóa kích thước trang về A4",
                           variable=self.var_pdfedit_normalize_a4).pack(side="left", padx=6)

            run_row = tk.Frame(frm)
            run_row.pack(fill="x", padx=6, pady=(4, 8))
            self.btn_pdfedit_dryrun = tk.Button(run_row, text="Chạy thử", font=("Arial", 10, "bold"),
                                                 bg="#455a64", fg="white", command=lambda: self.pdfedit_run(dry_run=True))
            self.btn_pdfedit_dryrun.pack(side="left", padx=2)
            self.btn_pdfedit_run = tk.Button(run_row, text="Lưu PDF đã chỉnh sửa", font=("Arial", 10, "bold"),
                                              bg="#2e7d32", fg="white", command=lambda: self.pdfedit_run(dry_run=False))
            self.btn_pdfedit_run.pack(side="left", padx=2)
            self.btn_pdfedit_pause, self.btn_pdfedit_cancel = self._add_pause_cancel(run_row, "control_pdfedit")
            self.btn_pdfedit_pause.pack(side="left", padx=(12, 3))
            self.btn_pdfedit_cancel.pack(side="left", padx=3)
            tk.Button(run_row, text="Mở thư mục kết quả", command=self.pdfedit_open_output_folder).pack(side="left", padx=12)

            tk.Label(frm, text="Bảng log:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(4, 2))
            cols = ("stt", "file_goc", "so_trang", "trang_da_xoay", "goc_xoay", "status", "note")
            headers = ["STT", "Tên file gốc", "Số trang", "Trang đã chỉnh sửa", "Góc xoay", "Trạng thái", "Ghi chú"]
            widths = [40, 160, 60, 130, 120, 100, 260]
            self.tree_pdfedit = ttk.Treeview(frm, columns=cols, show="headings", height=8)
            for c, h, w in zip(cols, headers, widths):
                self.tree_pdfedit.heading(c, text=h)
                self.tree_pdfedit.column(c, width=w, anchor="w")
            vsb_pe = tk.Scrollbar(frm, orient="vertical", command=self.tree_pdfedit.yview)
            self.tree_pdfedit.configure(yscrollcommand=vsb_pe.set)
            self.tree_pdfedit.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_pe.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            # ================= Tự nhận diện, xoay và làm thẳng PDF (mục I-IX tài liệu) =================
            frm_auto = tk.LabelFrame(parent, labelwidget=tk.Label(
                parent, text="  Tự nhận diện, xoay và làm thẳng PDF  ", font=("Segoe UI", 10, "bold"),
                fg="#1a237e", bg="#c5cae9", relief="raised", bd=1))
            frm_auto.pack(fill="x", **pad)
            tk.Label(frm_auto,
                    text="Tự động kiểm tra từng trang PDF, xác định hướng đúng (0°/90°/180°/270°) và làm thẳng góc "
                         "nghiêng nhỏ do scan - không chỉ dựa vào rotation metadata (nhiều file scan có ảnh đã bị "
                         "xoay nhưng metadata vẫn bằng 0). Giữ nguyên khổ giấy gốc, chỉ xoay để chữ đúng chiều.",
                    fg="#555", justify="left", wraplength=900).pack(anchor="w", padx=6, pady=(6, 6))

            row_auto_in = tk.Frame(frm_auto); row_auto_in.pack(fill="x", padx=6, pady=3)
            tk.Label(row_auto_in, text="Nguồn PDF:", width=16, anchor="w").pack(side="left")
            self.auto_rotate_picker = SourcePicker(
                row_auto_in, "Nguồn PDF", [("PDF files", "*.pdf")], file_label="PDF")
            self.auto_rotate_picker.pack(side="left", fill="x", expand=True)

            row_auto_out = tk.Frame(frm_auto); row_auto_out.pack(fill="x", padx=6, pady=3)
            tk.Label(row_auto_out, text="Thư mục xuất kết quả:", width=16, anchor="w").pack(side="left")
            self.var_auto_rotate_output = tk.StringVar()
            tk.Entry(row_auto_out, textvariable=self.var_auto_rotate_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_auto_out, text="Chọn...", command=self.pick_auto_rotate_output).pack(side="left")

            row_auto_opt = tk.Frame(frm_auto); row_auto_opt.pack(fill="x", padx=6, pady=3)
            tk.Label(row_auto_opt, text="Ngưỡng tự động xử lý:").pack(side="left")
            self.var_auto_rotate_threshold_auto = tk.StringVar(value="60")
            tk.Entry(row_auto_opt, textvariable=self.var_auto_rotate_threshold_auto, width=5).pack(side="left", padx=(4, 12))
            tk.Label(row_auto_opt, text="Ngưỡng cần kiểm tra (dưới mức này):").pack(side="left")
            self.var_auto_rotate_threshold_warn = tk.StringVar(value="40")
            tk.Entry(row_auto_opt, textvariable=self.var_auto_rotate_threshold_warn, width=5).pack(side="left", padx=(4, 12))
            self._register_advanced_widget(row_auto_opt, fill="x", padx=6, pady=3)

            row_auto_suffix = tk.Frame(frm_auto); row_auto_suffix.pack(fill="x", padx=6, pady=3)
            tk.Label(row_auto_suffix, text="Hậu tố tên file:").pack(side="left")
            self.var_auto_rotate_suffix = tk.StringVar(value="DA_XOAY_DUNG")
            tk.Entry(row_auto_suffix, textvariable=self.var_auto_rotate_suffix, width=14).pack(side="left", padx=(4, 0))

            row_auto_run = tk.Frame(frm_auto); row_auto_run.pack(fill="x", padx=6, pady=(6, 3))
            tk.Button(row_auto_run, text="🔍 Quét và phân tích (xem trước)", font=("Arial", 10),
                      bg="#00838f", fg="white", command=lambda: self.run_auto_rotate_clicked(dry_run=True)).pack(side="left", padx=(0, 6))
            tk.Button(row_auto_run, text="🔄 XỬ LÝ TOÀN BỘ (xoay/làm thẳng thật)", font=("Arial", 10, "bold"),
                      bg="#1a237e", fg="white", command=lambda: self.run_auto_rotate_clicked(dry_run=False)).pack(side="left")
            self.btn_auto_rotate_pause, self.btn_auto_rotate_cancel = self._add_pause_cancel(row_auto_run, "control_auto_rotate")
            self.btn_auto_rotate_pause.pack(side="left", padx=(8, 3))
            self.btn_auto_rotate_cancel.pack(side="left", padx=3)

            row_auto_progress1 = tk.Frame(frm_auto); row_auto_progress1.pack(fill="x", padx=6, pady=(3, 0))
            tk.Label(row_auto_progress1, text="File:", width=6, anchor="w").pack(side="left")
            self.pb_auto_rotate_file = ttk.Progressbar(row_auto_progress1, orient="horizontal", mode="determinate", length=350)
            self.pb_auto_rotate_file.pack(side="left", padx=(0, 10))
            self.lbl_auto_rotate_file_progress = tk.Label(row_auto_progress1, text="", fg="#555", anchor="w")
            self.lbl_auto_rotate_file_progress.pack(side="left", fill="x", expand=True)

            row_auto_progress2 = tk.Frame(frm_auto); row_auto_progress2.pack(fill="x", padx=6, pady=(0, 6))
            tk.Label(row_auto_progress2, text="Trang:", width=6, anchor="w").pack(side="left")
            self.pb_auto_rotate_page = ttk.Progressbar(row_auto_progress2, orient="horizontal", mode="determinate", length=350)
            self.pb_auto_rotate_page.pack(side="left", padx=(0, 10))
            self.lbl_auto_rotate_page_progress = tk.Label(row_auto_progress2, text="", fg="#555", anchor="w")
            self.lbl_auto_rotate_page_progress.pack(side="left", fill="x", expand=True)

            row_auto_tools = tk.Frame(frm_auto); row_auto_tools.pack(fill="x", padx=6, pady=(0, 6))
            tk.Button(row_auto_tools, text="Mở thư mục kết quả", command=self.auto_rotate_open_output_folder).pack(side="left", padx=(0, 6))
            tk.Button(row_auto_tools, text="Mở log", command=self.auto_rotate_open_log_folder).pack(side="left")

            tk.Label(frm_auto, text="Kết quả từng file:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(4, 2))
            cols_auto = ("stt", "file", "tong_trang", "da_dung", "xoay90", "xoay180", "xoay270", "lam_thang", "can_kt", "status")
            headers_auto = ["STT", "Tên file", "Tổng trang", "Đã đúng", "Xoay 90°", "Xoay 180°", "Xoay 270°",
                           "Làm thẳng", "Cần kiểm tra", "Trạng thái"]
            widths_auto = [40, 200, 70, 65, 70, 70, 70, 75, 80, 100]
            self.tree_auto_rotate = ttk.Treeview(frm_auto, columns=cols_auto, show="headings", height=6)
            for c, h, w in zip(cols_auto, headers_auto, widths_auto):
                self.tree_auto_rotate.heading(c, text=h)
                self.tree_auto_rotate.column(c, width=w, anchor="w")
            vsb_auto = tk.Scrollbar(frm_auto, orient="vertical", command=self.tree_auto_rotate.yview)
            self.tree_auto_rotate.configure(yscrollcommand=vsb_auto.set)
            self.tree_auto_rotate.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_auto.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            tk.Label(frm_auto, text="Xem trước từng trang (Quét và phân tích):", font=("Arial", 9, "bold")).pack(
                anchor="w", padx=6, pady=(4, 2))
            cols_page = ("stt", "so_trang", "rotation_meta", "goc_de_xuat", "goc_nghieng", "tin_cay", "nguon", "trang_thai", "ghi_chu")
            headers_page = ["STT", "Số trang", "Rotation gốc", "Góc đề xuất", "Góc nghiêng", "Tin cậy",
                           "Nguồn nhận diện", "Trạng thái", "Ghi chú"]
            widths_page = [40, 60, 80, 85, 80, 60, 130, 140, 260]
            self.tree_auto_rotate_pages = ttk.Treeview(frm_auto, columns=cols_page, show="headings", height=6)
            for c, h, w in zip(cols_page, headers_page, widths_page):
                self.tree_auto_rotate_pages.heading(c, text=h)
                self.tree_auto_rotate_pages.column(c, width=w, anchor="w")
            vsb_auto_page = tk.Scrollbar(frm_auto, orient="vertical", command=self.tree_auto_rotate_pages.yview)
            self.tree_auto_rotate_pages.configure(yscrollcommand=vsb_auto_page.set)
            self.tree_auto_rotate_pages.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_auto_page.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            row_auto_page_actions = tk.Frame(frm_auto); row_auto_page_actions.pack(fill="x", padx=6, pady=(0, 8))
            tk.Label(row_auto_page_actions, text="Chọn 1 dòng ở bảng trên rồi:", fg="#555").pack(side="left", padx=(0, 8))
            tk.Button(row_auto_page_actions, text="Bỏ qua trang này",
                      command=lambda: self.auto_rotate_page_action("skip")).pack(side="left", padx=3)
            tk.Button(row_auto_page_actions, text="Xoay thêm +90°",
                      command=lambda: self.auto_rotate_page_action("rotate90")).pack(side="left", padx=3)
            tk.Button(row_auto_page_actions, text="Đặt lại về hướng gốc (0°)",
                      command=lambda: self.auto_rotate_page_action("reset0")).pack(side="left", padx=3)
            tk.Button(row_auto_page_actions, text="Xóa ghi đè (dùng lại đề xuất tự động)",
                      command=lambda: self.auto_rotate_page_action("clear")).pack(side="left", padx=3)
            self.tree_auto_rotate_pages.bind("<<TreeviewSelect>>", self.auto_rotate_page_selected)

            tk.Label(frm_auto, text="Xem trước hình ảnh (chọn 1 dòng ở bảng trên để tải):",
                    font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(4, 2))
            row_auto_images = tk.Frame(frm_auto); row_auto_images.pack(fill="both", expand=True, padx=6, pady=(0, 8))
            frame_before = tk.LabelFrame(row_auto_images, text="Trước (gốc)")
            frame_before.pack(side="left", fill="both", expand=True, padx=(0, 4))
            self.canvas_auto_rotate_before = tk.Canvas(frame_before, bg="#cfcfcf", highlightthickness=0, height=280)
            self.canvas_auto_rotate_before.pack(fill="both", expand=True)
            frame_after = tk.LabelFrame(row_auto_images, text="Sau (dự kiến)")
            frame_after.pack(side="left", fill="both", expand=True, padx=(4, 0))
            self.canvas_auto_rotate_after = tk.Canvas(frame_after, bg="#cfcfcf", highlightthickness=0, height=280)
            self.canvas_auto_rotate_after.pack(fill="both", expand=True)

        def pdfedit_add_files(self):
            paths = filedialog.askopenfilenames(title="Chọn file PDF", filetypes=[("PDF files", "*.pdf")])
            for p in paths:
                if p not in self.pdfedit_files:
                    self.pdfedit_files.append(p)
                    self.lst_pdfedit_files.insert("end", os.path.basename(p))
            if paths and self.pdfedit_current_file is None:
                self.lst_pdfedit_files.selection_set(0)
                self.pdfedit_on_file_selected(None)

        def pdfedit_add_folder(self):
            folder = filedialog.askdirectory(title="Chọn thư mục chứa PDF")
            if not folder:
                return
            if self.var_pdfedit_subfolders.get():
                found = []
                for root, dirs, fnames in os.walk(folder):
                    for f in fnames:
                        if f.lower().endswith(".pdf"):
                            found.append(os.path.join(root, f))
                found.sort()
            else:
                found = [os.path.join(folder, f) for f in sorted(os.listdir(folder)) if f.lower().endswith(".pdf")]
            added = 0
            for p in found:
                if p not in self.pdfedit_files:
                    self.pdfedit_files.append(p)
                    self.lst_pdfedit_files.insert("end", os.path.basename(p))
                    added += 1
            if added and self.pdfedit_current_file is None:
                self.lst_pdfedit_files.selection_set(0)
                self.pdfedit_on_file_selected(None)
            if added == 0:
                messagebox.showinfo("Không có file mới", "Không tìm thấy file PDF mới nào để thêm.")

        def pdfedit_clear_files(self):
            self.pdfedit_files = []
            self.pdfedit_rotations = {}
            self.pdfedit_current_file = None
            self.pdfedit_total_pages = 0
            self.lst_pdfedit_files.delete(0, "end")
            self.canvas_pdfedit.delete("all")
            self.var_pdfedit_pageno.set("0")
            self.lbl_pdfedit_total.config(text="/ 0")

        def pdfedit_pick_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_pdfedit_output.set(p)

        def pdfedit_open_output_folder(self):
            out_dir = self.var_pdfedit_output.get().strip()
            if not out_dir or not os.path.isdir(out_dir):
                messagebox.showinfo("Chưa có thư mục", "Vui lòng chọn thư mục lưu kết quả hợp lệ trước.")
                return
            try:
                if sys.platform == "win32":
                    os.startfile(out_dir)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", out_dir])
                else:
                    subprocess.Popen(["xdg-open", out_dir])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được thư mục: {e}")

        def pdfedit_on_file_selected(self, event):
            sel = self.lst_pdfedit_files.curselection()
            if not sel:
                return
            fpath = self.pdfedit_files[sel[0]]
            self.pdfedit_current_file = fpath
            try:
                self.pdfedit_total_pages = get_pdf_page_count(fpath)
            except Exception as e:
                messagebox.showerror("Lỗi mở PDF", f"Không đọc được file: {e}")
                return
            self.pdfedit_current_page = 0
            self.lbl_pdfedit_total.config(text=f"/ {self.pdfedit_total_pages}")
            self.var_pdfedit_pageno.set("1")
            self.pdfedit_refresh_preview()

        def pdfedit_refresh_preview(self):
            if not self.pdfedit_current_file or not (HAS_FITZ and HAS_OCR or HAS_FITZ):
                return
            if not HAS_FITZ:
                self.canvas_pdfedit.delete("all")
                self.canvas_pdfedit.create_text(200, 100, text="Cần cài PyMuPDF (pymupdf) để xem trước.", fill="#900")
                return
            try:
                delta = self._pdfedit_get_delta(self.pdfedit_current_file, self.pdfedit_current_page)
                img = render_pdf_page_to_image(self.pdfedit_current_file, self.pdfedit_current_page,
                                                zoom=1.3 * self.pdfedit_zoom, extra_rotation=delta)
                from PIL import ImageTk
                self.pdfedit_photo_ref = ImageTk.PhotoImage(img)
                self.canvas_pdfedit.delete("all")
                self.canvas_pdfedit.config(scrollregion=(0, 0, img.width, img.height))
                self.canvas_pdfedit.create_image(0, 0, anchor="nw", image=self.pdfedit_photo_ref)
                self.lbl_pdfedit_current_rot.config(text=f"Đang xoay thêm: {delta}°")
            except Exception as e:
                self.canvas_pdfedit.delete("all")
                self.canvas_pdfedit.create_text(200, 100, text=f"Lỗi xem trước: {e}", fill="#900")

        def _pdfedit_get_delta(self, fpath, page_index):
            return self.pdfedit_rotations.get(fpath, {}).get(page_index, 0)

        def _pdfedit_set_delta(self, fpath, page_index, delta):
            self.pdfedit_rotations.setdefault(fpath, {})[page_index] = delta % 360

        def pdfedit_prev_page(self):
            if not self.pdfedit_current_file:
                return
            if self.pdfedit_current_page > 0:
                self.pdfedit_current_page -= 1
                self.var_pdfedit_pageno.set(str(self.pdfedit_current_page + 1))
                self.pdfedit_refresh_preview()

        def pdfedit_next_page(self):
            if not self.pdfedit_current_file:
                return
            if self.pdfedit_current_page < self.pdfedit_total_pages - 1:
                self.pdfedit_current_page += 1
                self.var_pdfedit_pageno.set(str(self.pdfedit_current_page + 1))
                self.pdfedit_refresh_preview()

        def pdfedit_goto_page(self):
            if not self.pdfedit_current_file:
                return
            try:
                p = int(self.var_pdfedit_pageno.get().strip())
            except ValueError:
                return
            if 1 <= p <= self.pdfedit_total_pages:
                self.pdfedit_current_page = p - 1
                self.pdfedit_refresh_preview()
            else:
                messagebox.showerror("Lỗi", f"Số trang phải từ 1 đến {self.pdfedit_total_pages}.")

        def pdfedit_set_zoom(self, value):
            self.pdfedit_zoom = value
            self.pdfedit_refresh_preview()

        def pdfedit_zoom_delta(self, delta):
            self.pdfedit_zoom = max(0.25, min(4.0, self.pdfedit_zoom + delta))
            self.pdfedit_refresh_preview()

        def pdfedit_rotate_current(self, delta):
            if not self.pdfedit_current_file:
                messagebox.showinfo("Chưa chọn file", "Vui lòng chọn 1 file PDF trong danh sách trước.")
                return
            current = self._pdfedit_get_delta(self.pdfedit_current_file, self.pdfedit_current_page)
            self._pdfedit_set_delta(self.pdfedit_current_file, self.pdfedit_current_page, current + delta)
            self.pdfedit_refresh_preview()

        def pdfedit_reset_current(self):
            if not self.pdfedit_current_file:
                return
            self._pdfedit_set_delta(self.pdfedit_current_file, self.pdfedit_current_page, 0)
            self.pdfedit_refresh_preview()

        def pdfedit_apply_scope(self, scope):
            if not self.pdfedit_current_file:
                messagebox.showinfo("Chưa chọn file", "Vui lòng chọn 1 file PDF và xoay trang xem trước để lấy góc xoay.")
                return
            delta = self._pdfedit_get_delta(self.pdfedit_current_file, self.pdfedit_current_page)
            if delta == 0:
                if not messagebox.askyesno("Góc xoay = 0°",
                                            "Trang đang xem hiện chưa xoay (0°). Vẫn áp dụng 0° cho phạm vi đã chọn "
                                            "(coi như đặt về góc gốc)?"):
                    return

            if scope == "file":
                for i in range(self.pdfedit_total_pages):
                    self._pdfedit_set_delta(self.pdfedit_current_file, i, delta)
                self.log(f"Đã áp dụng {delta}° cho toàn bộ {self.pdfedit_total_pages} trang của file đang xem.")

            elif scope == "all_files":
                for fpath in self.pdfedit_files:
                    try:
                        n = get_pdf_page_count(fpath)
                    except Exception:
                        continue
                    for i in range(n):
                        self._pdfedit_set_delta(fpath, i, delta)
                self.log(f"Đã áp dụng {delta}° cho toàn bộ {len(self.pdfedit_files)} file trong danh sách.")

            elif scope in ("even", "odd"):
                for i in range(self.pdfedit_total_pages):
                    page_no = i + 1
                    if (scope == "even" and page_no % 2 == 0) or (scope == "odd" and page_no % 2 == 1):
                        self._pdfedit_set_delta(self.pdfedit_current_file, i, delta)
                self.log(f"Đã áp dụng {delta}° cho các trang {'chẵn' if scope == 'even' else 'lẻ'} của file đang xem.")

            elif scope == "range":
                try:
                    pages = parse_page_range_spec(self.var_pdfedit_range.get(), self.pdfedit_total_pages)
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Khoảng trang không hợp lệ: {e}")
                    return
                for i in pages:
                    self._pdfedit_set_delta(self.pdfedit_current_file, i, delta)
                self.log(f"Đã áp dụng {delta}° cho {len(pages)} trang theo khoảng '{self.var_pdfedit_range.get()}'.")

            self.pdfedit_refresh_preview()

        def pdfedit_autodetect(self, scope):
            if not self.pdfedit_files:
                messagebox.showinfo("Chưa có file", "Vui lòng thêm file PDF trước.")
                return
            if not (HAS_FITZ and HAS_OCR):
                messagebox.showerror("Thiếu thư viện OCR",
                                      "Cần cài pymupdf + pytesseract + Tesseract-OCR để dùng tính năng này.")
                return

            targets = [self.pdfedit_current_file] if scope == "file" and self.pdfedit_current_file else self.pdfedit_files
            if not targets:
                messagebox.showinfo("Chưa chọn file", "Vui lòng chọn 1 file trong danh sách trước.")
                return

            self.txt_log.delete("1.0", "end")

            def worker():
                need_check_pages = []
                for fpath in targets:
                    fname = os.path.basename(fpath)
                    try:
                        n = get_pdf_page_count(fpath)
                    except Exception as e:
                        self.log(f"✗ {fname}: {e}")
                        continue
                    self.log(f"→ {fname}: đang phân tích {n} trang...")
                    for i in range(n):
                        try:
                            result = detect_best_orientation(fpath, i)
                        except Exception as e:
                            self.log(f"   Trang {i+1}: lỗi phân tích - {e}")
                            continue
                        if result["chac_chan"] and result["angle_de_xuat"] != 0:
                            self._pdfedit_set_delta(fpath, i, result["angle_de_xuat"])
                            self.log(f"   Trang {i+1}: đề xuất xoay +{result['angle_de_xuat']}° "
                                     f"({result['ly_do']}) → ĐÃ ÁP DỤNG vào bản xem trước")
                        elif not result["chac_chan"]:
                            need_check_pages.append((fname, i + 1, result["ly_do"]))
                            self.log(f"   Trang {i+1}: CẦN KIỂM TRA - {result['ly_do']}")
                        else:
                            self.log(f"   Trang {i+1}: đã đúng hướng (0°)")
                self.log("\nHoàn tất tự phát hiện hướng. Các góc đề xuất đã được áp dụng vào bản XEM TRƯỚC "
                         "(chưa lưu file) — hãy kiểm tra lại rồi bấm 'Lưu PDF đã chỉnh sửa'.")
                if need_check_pages:
                    self.log(f"\n⚠ {len(need_check_pages)} trang KHÔNG đủ tin cậy để tự xoay, cần bạn tự kiểm tra bằng mắt:")
                    for fn, pg, reason in need_check_pages:
                        self.log(f"   - {fn} (trang {pg}): {reason}")
                self.after(0, self.pdfedit_refresh_preview)

            threading.Thread(target=worker, daemon=True).start()

        def pdfedit_run(self, dry_run=True):
            if not self.pdfedit_files:
                messagebox.showerror("Chưa có file", "Vui lòng thêm ít nhất 1 file PDF.")
                return
            output_folder = self.var_pdfedit_output.get().strip()
            if not dry_run and not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return

            enhance_opts = None
            if self.var_pdfedit_enhance_on.get():
                try:
                    contrast = float(self.var_pdfedit_contrast_val.get()) if self.var_pdfedit_contrast_on.get() else 1.0
                except ValueError:
                    contrast = 1.0
                try:
                    threshold = int(self.var_pdfedit_threshold_val.get()) if self.var_pdfedit_threshold_on.get() else None
                except ValueError:
                    threshold = None
                enhance_opts = {
                    "contrast": contrast,
                    "sharpen": self.var_pdfedit_sharpen.get(),
                    "grayscale": self.var_pdfedit_grayscale.get(),
                    "threshold": threshold,
                    "denoise": self.var_pdfedit_denoise.get(),
                    "deskew": self.var_pdfedit_deskew.get() and HAS_CV2,
                    "crop_border": self.var_pdfedit_crop_border.get(),
                    "normalize_a4": self.var_pdfedit_normalize_a4.get(),
                }

            ocr_after = self.var_pdfedit_ocr_after.get()
            rotations_snapshot = {f: dict(self.pdfedit_rotations.get(f, {})) for f in self.pdfedit_files}

            def get_rotations_for_file(fpath, total_pages):
                return rotations_snapshot.get(fpath, {})

            if not dry_run:
                if not self._check_require_dry_run("_dry_run_done_pdfedit", True, "Chỉnh sửa PDF hàng loạt (Tab 6)"):
                    return
                if not messagebox.askyesno("Xác nhận lưu",
                                            "Sẽ lưu các file PDF đã chỉnh sửa ra thư mục kết quả (file gốc KHÔNG bị "
                                            "thay đổi). Tiếp tục?"):
                    return

            self.txt_log.delete("1.0", "end")
            for item in self.tree_pdfedit.get_children():
                self.tree_pdfedit.delete(item)
            self.btn_pdfedit_dryrun.config(state="disabled")
            self.btn_pdfedit_run.config(state="disabled")
            self._enable_pause_cancel("control_pdfedit", self.btn_pdfedit_pause, self.btn_pdfedit_cancel)

            def worker():
                try:
                    if not dry_run:
                        self._maybe_auto_backup(self.pdfedit_files, "Chỉnh sửa PDF hàng loạt (Tab 6)")
                    results = batch_edit_pdfs(self.pdfedit_files, output_folder, get_rotations_for_file,
                                              enhance_opts, dry_run, ocr_after, self.log, control=self.control_pdfedit)
                    if dry_run:
                        self._dry_run_done_pdfedit = True

                    def _fill_tree():
                        for r in results:
                            self.tree_pdfedit.insert("", "end", values=(r["stt"], r["file_goc"], r["so_trang"],
                                                                          r["trang_da_xoay"], r["goc_xoay"],
                                                                          r["status"], r["note"]))
                    self.after(0, _fill_tree)

                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    log_dir = output_folder or os.path.dirname(self.pdfedit_files[0])
                    csv_path = os.path.join(log_dir, f"LOG_CHINH_SUA_PDF_{stamp}.csv")
                    write_pdf_edit_log_csv(csv_path, results)

                    try:
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "CHINH_SUA_PDF", run_id=stamp)
                        for r in results:
                            std_log.add(action="EDIT_PDF", source_file=r["file_goc"], output_file=r.get("file_xuat", ""),
                                       status=r["status"], message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    ok_count = sum(1 for r in results if r["status"] in ("DA_XU_LY", "CHAY_THU"))
                    other_count = len(results) - ok_count

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {ok_count} OK | {other_count} lỗi/cần kiểm tra")
                    self.log(f"File log CSV: {csv_path}")

                    self.after(0, lambda: self.mark_workflow_step("chuan_hoa", "Đã chạy thử" if dry_run else "Đã hoàn thành"))
                    messagebox.showinfo("Chạy thử hoàn tất" if dry_run else "Lưu hoàn tất",
                                         f"OK: {ok_count}\nLỗi/cần kiểm tra: {other_count}\n\nLog CSV: {csv_path}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_pdfedit_dryrun.config(state="normal")
                    self.btn_pdfedit_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_pdfedit_pause, self.btn_pdfedit_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def pick_auto_rotate_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục xuất kết quả")
            if p:
                self.var_auto_rotate_output.set(p)

        def auto_rotate_open_output_folder(self):
            out_dir = self.var_auto_rotate_output.get().strip()
            if not out_dir or not os.path.isdir(out_dir):
                messagebox.showinfo("Chưa có thư mục", "Vui lòng chọn thư mục lưu kết quả hợp lệ trước.")
                return
            try:
                if sys.platform == "win32":
                    os.startfile(out_dir)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", out_dir])
                else:
                    subprocess.Popen(["xdg-open", out_dir])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được thư mục: {e}")

        def auto_rotate_open_log_folder(self):
            out_dir = self.var_auto_rotate_output.get().strip()
            if not out_dir:
                messagebox.showinfo("Chưa có thư mục", "Vui lòng chọn thư mục lưu kết quả hợp lệ trước.")
                return
            log_dir = os.path.join(out_dir, "Logs")
            if not os.path.isdir(log_dir):
                messagebox.showinfo("Chưa có log", "Chưa có log nào được tạo - hãy chạy xử lý trước.")
                return
            try:
                if sys.platform == "win32":
                    os.startfile(log_dir)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", log_dir])
                else:
                    subprocess.Popen(["xdg-open", log_dir])
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không mở được thư mục log: {e}")

        def run_auto_rotate_clicked(self, dry_run=False):
            source = self.auto_rotate_picker.get_source()
            output = self.var_auto_rotate_output.get().strip()
            suffix = self.var_auto_rotate_suffix.get().strip() or "DA_XOAY_DUNG"

            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục hoặc (các) file PDF nguồn.")
                return
            if not dry_run and not output:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục xuất kết quả.")
                return
            try:
                threshold_auto = float(self.var_auto_rotate_threshold_auto.get().strip())
                threshold_warn = float(self.var_auto_rotate_threshold_warn.get().strip())
            except ValueError:
                messagebox.showerror("Sai định dạng", "Ngưỡng tin cậy phải là số.")
                return

            self.txt_log.delete("1.0", "end")

            checkpoint_task_id = None
            files_da_xong = None
            if not dry_run:
                from app.services import checkpoint_service as _ckpt
                checkpoint_task_id = _ckpt.tinh_task_id("xoay_lam_thang_pdf", str(source), output)
                data_cu = _ckpt.doc_checkpoint(checkpoint_task_id)
                if data_cu and data_cu.get("trang_thai") == "DANG_XU_LY":
                    so_da_xong = len(data_cu.get("da_hoan_thanh", []))
                    so_tong = len(data_cu.get("danh_sach_file", []))
                    if messagebox.askyesno(
                        "Phát hiện tác vụ dở dang",
                        f"Phát hiện đợt xoay/làm thẳng PDF TRƯỚC ĐÓ chưa hoàn thành (đã xong "
                        f"{so_da_xong}/{so_tong} file, có thể do phần mềm bị đóng đột ngột).\n\n"
                        f"Bấm CÓ để TIẾP TỤC từ chỗ dở dang.\nBấm KHÔNG để BẮT ĐẦU LẠI TỪ ĐẦU."):
                        files_da_xong = set(data_cu.get("da_hoan_thanh", []))
                    else:
                        _ckpt.xoa_checkpoint(checkpoint_task_id)

            from app.services import task_manager_service as _tm4
            self.task_info_rotate = _tm4.TaskInfo(
                ten_chuc_nang="6. Xử lý PDF (Xoay/Làm thẳng)", thu_muc_dau_ra=output or "",
                trang_thai=_tm4.TrangThaiTask.DANG_KHOI_TAO)

            for item in self.tree_auto_rotate.get_children():
                self.tree_auto_rotate.delete(item)
            for item in self.tree_auto_rotate_pages.get_children():
                self.tree_auto_rotate_pages.delete(item)
            self._auto_rotate_page_map = {}
            if not hasattr(self, "auto_rotate_page_overrides"):
                self.auto_rotate_page_overrides = {}
            self._enable_pause_cancel("control_auto_rotate", self.btn_auto_rotate_pause, self.btn_auto_rotate_cancel)
            self.pb_auto_rotate_file.config(value=0, maximum=100)
            self.pb_auto_rotate_page.config(value=0, maximum=100)
            self._auto_rotate_start_time = time.time()

            def update_file_progress(idx, total, fname):
                elapsed = time.time() - self._auto_rotate_start_time
                self.pb_auto_rotate_file.config(value=idx, maximum=total)
                self.lbl_auto_rotate_file_progress.config(
                    text=f"Đang xử lý file {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def update_page_progress(so_trang, tong):
                self.pb_auto_rotate_page.config(value=so_trang, maximum=tong)
                self.lbl_auto_rotate_page_progress.config(text=f"Đang phân tích trang {so_trang}/{tong}")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_file_progress(idx, total, fname))

            def worker():
                try:
                    def page_progress_wrapper(so_trang, tong):
                        self.after(0, lambda: update_page_progress(so_trang, tong))

                    files = list_files_from_source(source, ".pdf")
                    if not files:
                        raise RuntimeError("Không có file PDF nào trong nguồn đã chọn.")
                    log_rows, chi_tiet_rows = [], []
                    if not dry_run:
                        os.makedirs(output, exist_ok=True)
                    stt_page = 0

                    from app.services import word_perf_service as _wperf
                    ap_dung_toi_uu = len(files) >= 2
                    if ap_dung_toi_uu:
                        _wperf.apply_low_priority()
                        sleep_per_file = _wperf.get_sleep_per_file(
                            getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"))
                        so_file_moi_dot = getattr(self.app_config, "word_so_file_moi_dot", 100)
                        nghi_giua_dot_giay = getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0)
                    else:
                        sleep_per_file = 0.0

                    bi_huy_giua_chung = False
                    for i, fpath in enumerate(files):
                        if files_da_xong and fpath in files_da_xong:
                            continue  # RESUME: file nay da xu ly xong o lan chay TRUOC
                        fname = os.path.basename(fpath)
                        progress_cb(i + 1, len(files), fname)
                        self.task_info_rotate.cap_nhat_tien_do(tien_do_toan_bo=i + 1,
                                                               trang_thai=_tm4.TrangThaiTask.DANG_PHAN_TICH)
                        self.task_info_rotate.tong_so_luong = len(files)
                        t0 = time.time()
                        base = os.path.splitext(fname)[0]
                        out_path = os.path.join(output, f"{base}_{suffix}.pdf") if output else ""
                        if not dry_run and os.path.isfile(out_path):
                            k = 1
                            while os.path.isfile(os.path.join(output, f"{base}_{suffix}_{k}.pdf")):
                                k += 1
                            out_path = os.path.join(output, f"{base}_{suffix}_{k}.pdf")
                        row = {"stt": i + 1, "file_in": fname, "path_in": fpath, "tong_so_trang": 0,
                              "so_trang_da_dung": 0, "so_trang_xoay_90": 0, "so_trang_xoay_180": 0,
                              "so_trang_xoay_270": 0, "so_trang_lam_thang": 0, "so_trang_can_kiem_tra": 0,
                              "file_out": "", "path_out": "", "thoi_gian_xu_ly": "", "status": "", "note": ""}
                        try:
                            overrides_file = self.auto_rotate_page_overrides.get(fpath, {})
                            skip_pages = {p for p, v in overrides_file.items() if v == "skip"}
                            manual_rotations = {p: v for p, v in overrides_file.items() if isinstance(v, int)}
                            result = auto_rotate_deskew_pdf(
                                fpath, out_path, self.log, threshold_auto=threshold_auto,
                                threshold_warn=threshold_warn, control=self.control_auto_rotate,
                                page_progress_cb=page_progress_wrapper, dry_run=dry_run,
                                skip_pages=skip_pages, manual_rotations=manual_rotations)
                            row["tong_so_trang"] = len(result["trang_ket_qua"])
                            row["so_trang_da_dung"] = result["so_trang_da_dung"]
                            row["so_trang_lam_thang"] = result["so_trang_da_lam_thang"]
                            row["so_trang_can_kiem_tra"] = result["so_trang_can_kiem_tra"]
                            for tr in result["trang_ket_qua"]:
                                if tr["trang_thai"] == "DA_XOAY_90":
                                    row["so_trang_xoay_90"] += 1
                                elif tr["trang_thai"] == "DA_XOAY_180":
                                    row["so_trang_xoay_180"] += 1
                                elif tr["trang_thai"] == "DA_XOAY_270":
                                    row["so_trang_xoay_270"] += 1
                                chi_tiet_rows.append({"file": fname, **tr})
                                stt_page += 1
                                self._auto_rotate_page_map[stt_page] = (fpath, tr["so_trang"])
                                self.after(0, lambda stt=stt_page, t=tr: self.tree_auto_rotate_pages.insert(
                                    "", "end", values=(stt, t["so_trang"], t["rotation_metadata"],
                                                      t["goc_xoay_de_xuat"], t["goc_nghieng"], t["diem_tin_cay"],
                                                      t["nguon_nhan_dien"], t["trang_thai"], t["ghi_chu"])))
                            row["file_out"] = os.path.basename(out_path) if not dry_run else "(chưa ghi - chỉ xem trước)"
                            row["path_out"] = out_path
                            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)
                            row["status"] = "CAN_KIEM_TRA" if result["so_trang_can_kiem_tra"] else "THANH_CONG"
                            row["note"] = (f"{result['so_trang_can_kiem_tra']} trang cần kiểm tra thủ công"
                                          if result["so_trang_can_kiem_tra"] else "")
                            self.log(f"✓ {'Đã phân tích' if dry_run else 'Đã xử lý'}: {fname}"
                                    + (f" → {row['file_out']}" if not dry_run else ""))
                        except TaskCancelled:
                            self.log(f"⏹ Đã hủy theo yêu cầu (còn {len(files) - i} file chưa xử lý).")
                            bi_huy_giua_chung = True
                            break
                        except Exception as e:
                            row["status"] = "LOI_DOC_PDF"
                            row["note"] = str(e)
                            row["thoi_gian_xu_ly"] = round(time.time() - t0, 2)
                            self.log(f"✗ Lỗi xử lý {fname}: {e}")
                        log_rows.append(row)
                        self.after(0, lambda r=dict(row): self.tree_auto_rotate.insert("", "end", values=(
                            r["stt"], r["file_in"], r["tong_so_trang"], r["so_trang_da_dung"],
                            r["so_trang_xoay_90"], r["so_trang_xoay_180"], r["so_trang_xoay_270"],
                            r["so_trang_lam_thang"], r["so_trang_can_kiem_tra"], r["status"])))

                        if ap_dung_toi_uu:
                            _wperf.collect_garbage()
                            time.sleep(sleep_per_file)
                            if so_file_moi_dot and (i + 1) % so_file_moi_dot == 0 and (i + 1) < len(files):
                                self.log(f"⏸ Đã xử lý {i + 1} file - nghỉ {nghi_giua_dot_giay:.0f}s để nhường tài nguyên...")
                                time.sleep(nghi_giua_dot_giay)

                        if checkpoint_task_id:
                            from app.services import checkpoint_service as _ckpt2
                            da_xu_ly_toi_idx = set(files[:i + 1]) | (files_da_xong or set())
                            _ckpt2.luu_checkpoint(
                                checkpoint_task_id, "xoay_lam_thang_pdf", files,
                                list(da_xu_ly_toi_idx), [], [], i + 1,
                                {}, output or "")

                    if ap_dung_toi_uu:
                        _wperf.restore_normal_priority()

                    if checkpoint_task_id and not bi_huy_giua_chung:
                        from app.services import checkpoint_service as _ckpt3
                        _ckpt3.xoa_checkpoint(checkpoint_task_id)

                    n_ok = sum(1 for r in log_rows if r["status"] == "THANH_CONG")
                    n_check = sum(1 for r in log_rows if r["status"] == "CAN_KIEM_TRA")
                    n_loi = sum(1 for r in log_rows if r["status"] == "LOI_DOC_PDF")

                    self.task_info_rotate.so_luong_thanh_cong = n_ok
                    self.task_info_rotate.so_luong_can_kiem_tra = n_check
                    self.task_info_rotate.so_luong_loi = n_loi
                    self.task_info_rotate.trang_thai = (_tm4.TrangThaiTask.DA_DUNG_BOI_NGUOI_DUNG
                                                        if bi_huy_giua_chung else _tm4.TrangThaiTask.THANH_CONG)

                    if not dry_run:
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        log_dir = os.path.join(output, "Logs")
                        os.makedirs(log_dir, exist_ok=True)
                        write_log_tu_dong_xoay_pdf(
                            os.path.join(log_dir, f"LOG_TU_DONG_XOAY_PDF_{stamp}.csv"), log_rows)
                        write_log_chi_tiet_huong_trang(
                            os.path.join(log_dir, f"LOG_CHI_TIET_HUONG_TRANG_{stamp}.csv"), chi_tiet_rows)

                    self.after(0, lambda: self.lbl_auto_rotate_file_progress.config(
                        text=f"✓ Hoàn tất — {len(log_rows)} file ({int(time.time() - self._auto_rotate_start_time)}s)"))
                    self.log("\n" + "=" * 60)
                    if dry_run:
                        self.log(f"TỔNG KẾT PHÂN TÍCH (chưa ghi file): {n_ok} trang đã đúng dự kiến, "
                                f"{n_check} cần kiểm tra. Xem bảng bên dưới, bấm 'XỬ LÝ TOÀN BỘ' để áp dụng thật.")
                        messagebox.showinfo("Đã phân tích xong",
                                            f"Kết quả DỰ KIẾN (chưa ghi file nào):\n"
                                            f"Sẽ đúng/không đổi: {n_ok}\nCần kiểm tra: {n_check}\n\n"
                                            f"Xem chi tiết ở bảng, sau đó bấm 'XỬ LÝ TOÀN BỘ' nếu đồng ý.")
                    else:
                        self.log(f"TỔNG KẾT: {n_ok} thành công, {n_check} cần kiểm tra, {n_loi} lỗi.")
                        messagebox.showinfo("Hoàn tất",
                                            f"Thành công: {n_ok}\nCần kiểm tra: {n_check}\nLỗi: {n_loi}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self._disable_pause_cancel(self.btn_auto_rotate_pause, self.btn_auto_rotate_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def auto_rotate_page_selected(self, event):
            sel = self.tree_auto_rotate_pages.selection()
            if not sel:
                return
            values = self.tree_auto_rotate_pages.item(sel[0], "values")
            stt = int(values[0])
            if stt not in self._auto_rotate_page_map:
                return
            fpath, so_trang = self._auto_rotate_page_map[stt]

            overrides_file = self.auto_rotate_page_overrides.get(fpath, {})
            override_val = overrides_file.get(so_trang)
            if override_val == "skip":
                goc_du_kien = 0
            elif isinstance(override_val, int):
                goc_du_kien = override_val
            else:
                try:
                    goc_du_kien = int(values[3])
                except (ValueError, TypeError):
                    goc_du_kien = 0

            try:
                img_before = render_pdf_page_to_image(fpath, so_trang - 1, zoom=0.8, extra_rotation=0)
                self._auto_rotate_photo_before = self._pdfedit_tk_image(img_before)
                self.canvas_auto_rotate_before.delete("all")
                self.canvas_auto_rotate_before.config(scrollregion=(0, 0, img_before.width, img_before.height))
                self.canvas_auto_rotate_before.create_image(0, 0, anchor="nw", image=self._auto_rotate_photo_before)
            except Exception as e:
                self.canvas_auto_rotate_before.delete("all")
                self.canvas_auto_rotate_before.create_text(140, 100, text=f"Lỗi xem trước: {e}", fill="#900")

            try:
                img_after = render_pdf_page_to_image(fpath, so_trang - 1, zoom=0.8, extra_rotation=goc_du_kien)
                self._auto_rotate_photo_after = self._pdfedit_tk_image(img_after)
                self.canvas_auto_rotate_after.delete("all")
                self.canvas_auto_rotate_after.config(scrollregion=(0, 0, img_after.width, img_after.height))
                self.canvas_auto_rotate_after.create_image(0, 0, anchor="nw", image=self._auto_rotate_photo_after)
            except Exception as e:
                self.canvas_auto_rotate_after.delete("all")
                self.canvas_auto_rotate_after.create_text(140, 100, text=f"Lỗi xem trước: {e}", fill="#900")

        def _pdfedit_tk_image(self, pil_image):
            from PIL import ImageTk
            return ImageTk.PhotoImage(pil_image)

        def auto_rotate_page_action(self, action):
            sel = self.tree_auto_rotate_pages.selection()
            if not sel:
                messagebox.showinfo("Chưa chọn dòng", "Vui lòng chọn 1 dòng trong bảng xem trước trước.")
                return
            item = sel[0]
            values = self.tree_auto_rotate_pages.item(item, "values")
            stt = int(values[0])
            if stt not in self._auto_rotate_page_map:
                return
            fpath, so_trang = self._auto_rotate_page_map[stt]
            overrides_file = self.auto_rotate_page_overrides.setdefault(fpath, {})

            if action == "skip":
                overrides_file[so_trang] = "skip"
                new_status, new_note = "DA_DUNG_BOI_NGUOI_DUNG", "Người dùng chọn giữ nguyên trang này (chưa xử lý lại)"
                new_goc = 0
            elif action == "rotate90":
                current_goc = overrides_file.get(so_trang, 0)
                current_goc = current_goc if isinstance(current_goc, int) else 0
                new_goc = (current_goc + 90) % 360
                overrides_file[so_trang] = new_goc
                new_status, new_note = f"Sẽ xoay {new_goc}°", "Người dùng tự chọn góc xoay (chưa xử lý lại)"
            elif action == "reset0":
                overrides_file[so_trang] = 0
                new_status, new_note = "Sẽ giữ 0°", "Người dùng đặt lại về hướng gốc (chưa xử lý lại)"
                new_goc = 0
            elif action == "clear":
                overrides_file.pop(so_trang, None)
                new_status, new_note = "(dùng lại đề xuất tự động)", "Đã xóa ghi đè - bấm 'Quét và phân tích' lại để xem đề xuất tự động"
                new_goc = values[3]
            else:
                return

            new_values = list(values)
            new_values[3] = new_goc if action != "clear" else values[3]
            new_values[7] = new_status
            new_values[8] = new_note
            self.tree_auto_rotate_pages.item(item, values=tuple(new_values))
            self.log(f"Trang {so_trang} ({os.path.basename(fpath)}): {new_note}")

try:
    from app.ui.tab_report import ReportTabMixin
    _UI_REPORT_MODULE_OK = True
except Exception:
    _UI_REPORT_MODULE_OK = False

    class ReportTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_report.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_report(self, parent, pad):
            frm = styled_labelframe(parent, text="8a. Tổng hợp báo cáo từ các file PDF ra Excel (theo mẫu)")
            frm.pack(fill="x", **pad)
            tk.Label(frm,
                     text="Đọc PDF (chọn thư mục hoặc chọn file bất kỳ), gộp các file cùng 1 thửa đất (cùng Mã xã+Tờ+Thửa) "
                          "thành 1 dòng, xuất ra file Excel theo đúng cấu trúc cột đang dùng ở Tab 1-4 "
                          "(Mã xã=B, Tờ=W, Thửa=X, Tên chủ=H, Địa chỉ=Y, Diện tích=Z, Tên file=AY, dữ liệu từ dòng 5).\n"
                          "→ File Excel này dùng được NGAY cho Tab 1-4 (đối chiếu/đổi tên) mà không cần cấu hình lại cột.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 8))

            self.report_picker = SourcePicker(
                parent, "Nguồn PDF cần tổng hợp (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF bất kỳ)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.report_picker.pack(fill="x", **pad)

            row_out = tk.Frame(frm)
            row_out.pack(fill="x", padx=6, pady=3)
            tk.Label(row_out, text="Lưu file Excel báo cáo tại:", width=32, anchor="w").pack(side="left")
            tk.Entry(row_out, textvariable=self.var_report_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_out, text="Chọn nơi lưu...", command=self.pick_report_output).pack(side="left")

            tk.Label(frm,
                     text="Lưu ý: dùng chung cấu hình OCR và bảng tra Tên xã → Mã xã ở Tab 5 "
                          "(bật/tắt OCR, độ phân giải, xem văn bản khi lỗi).",
                     fg="#777", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(0, 6))

            row_run = tk.Frame(frm)
            row_run.pack(fill="x", padx=6, pady=(2, 3))
            self.btn_report = tk.Button(row_run, text="📊 TỔNG HỢP BÁO CÁO RA EXCEL", font=("Arial", 11, "bold"),
                                         bg="#ad1457", fg="white", command=self.run_report_clicked)
            self.btn_report.pack(side="left")
            self.btn_report_pause, self.btn_report_cancel = self._add_pause_cancel(row_run, "control_report")
            self.btn_report_pause.pack(side="left", padx=(8, 3))
            self.btn_report_cancel.pack(side="left", padx=3)

            row_report_progress = tk.Frame(frm); row_report_progress.pack(fill="x", padx=6, pady=(0, 8))
            self.pb_report = ttk.Progressbar(row_report_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_report.pack(side="left", padx=(0, 10))
            self.lbl_report_progress = tk.Label(row_report_progress, text="", fg="#555", anchor="w")
            self.lbl_report_progress.pack(side="left", fill="x", expand=True)

            frm2 = styled_labelframe(parent, text="8b. Thống kê tiến độ từ file Excel BẤT KỲ (tự chọn cột)")
            frm2.pack(fill="x", **pad)
            tk.Label(frm2,
                     text="Chọn 1 file Excel bất kỳ (VD file tổng báo cáo đang quản lý), tự chọn cột nào thể hiện "
                          "trạng thái thực hiện / giấy chứng nhận, phần mềm sẽ đếm và tính % giúp bạn.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_e1 = tk.Frame(frm2)
            row_e1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_e1, text="File Excel:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_e1, textvariable=self.var_stat_excel).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_e1, text="Chọn...", command=self.pick_stat_excel).pack(side="left")

            row_e2 = tk.Frame(frm2)
            row_e2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_e2, text="Sheet (để trống=đầu tiên):", width=20, anchor="w").pack(side="left")
            tk.Entry(row_e2, textvariable=self.var_stat_sheet, width=15).pack(side="left", padx=(0, 14))
            tk.Label(row_e2, text="Dòng bắt đầu dữ liệu:").pack(side="left")
            tk.Entry(row_e2, textvariable=self.var_stat_header_row, width=6).pack(side="left", padx=(0, 14))
            tk.Label(row_e2, text="Cột đếm số dòng (VD A):").pack(side="left")
            tk.Entry(row_e2, textvariable=self.var_stat_col_count, width=6).pack(side="left")

            row_e3 = tk.Frame(frm2)
            row_e3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_e3, text="Cột Trạng thái thực hiện:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_e3, textvariable=self.var_stat_col_status, width=6).pack(side="left", padx=(0, 8))
            tk.Label(row_e3, text="Giá trị coi là ĐÃ thực hiện (phân cách bởi dấu phẩy):").pack(side="left")
            tk.Entry(row_e3, textvariable=self.var_stat_done_values).pack(side="left", fill="x", expand=True, padx=4)

            row_e4 = tk.Frame(frm2)
            row_e4.pack(fill="x", padx=6, pady=3)
            tk.Label(row_e4, text="Cột Giấy chứng nhận (GCN):", width=20, anchor="w").pack(side="left")
            tk.Entry(row_e4, textvariable=self.var_stat_col_gcn, width=6).pack(side="left", padx=(0, 8))
            tk.Label(row_e4, text="Giá trị coi là ĐÃ nhập GCN (phân cách bởi dấu phẩy):").pack(side="left")
            tk.Entry(row_e4, textvariable=self.var_stat_gcn_values).pack(side="left", fill="x", expand=True, padx=4)

            row_e5 = tk.Frame(frm2)
            row_e5.pack(fill="x", padx=6, pady=3)
            tk.Label(row_e5, text="Tổng khối lượng được giao (để trống = tự đếm tổng số dòng):",
                     anchor="w").pack(side="left")
            tk.Entry(row_e5, textvariable=self.var_stat_target_total, width=10).pack(side="left", padx=6)

            row_e_run = tk.Frame(frm2)
            row_e_run.pack(fill="x", padx=6, pady=(4, 3))
            self.btn_stat = tk.Button(row_e_run, text="📈 THỐNG KÊ TIẾN ĐỘ", font=("Arial", 11, "bold"),
                                       bg="#00838f", fg="white", command=self.run_stat_clicked)
            self.btn_stat.pack(side="left")
            tk.Label(row_e_run, text="Kết quả hiện ở khung Nhật ký bên dưới, kèm xuất file Excel tóm tắt.",
                     fg="#555").pack(side="left", padx=10)

            row_stat_progress = tk.Frame(frm2); row_stat_progress.pack(fill="x", padx=6, pady=(0, 8))
            self.pb_stat = ttk.Progressbar(row_stat_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_stat.pack(side="left", padx=(0, 10))
            self.lbl_stat_progress = tk.Label(row_stat_progress, text="", fg="#555", anchor="w")
            self.lbl_stat_progress.pack(side="left", fill="x", expand=True)

            # ================= Khối lượng và tiến độ thực hiện (MỚI) =================
            frm_prog = styled_labelframe(parent, text="📊 Khối lượng và tiến độ thực hiện",
                                     font=("Segoe UI", 9, "bold"), fg="#1a237e")
            frm_prog.pack(fill="x", **pad)
            tk.Label(frm_prog,
                     text="Nhập thủ công tổng khối lượng/đã/đang thực hiện, HOẶC chọn thư mục/file PDF đã thực hiện "
                          "để phần mềm tự đếm và nhận diện Số tờ/Số thửa.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_p1 = tk.Frame(frm_prog)
            row_p1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_p1, text="Tổng khối lượng:", width=16, anchor="w").pack(side="left")
            tk.Entry(row_p1, textvariable=self.var_prog_total, width=10).pack(side="left", padx=(0, 16))
            tk.Label(row_p1, text="Đã thực hiện:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_p1, textvariable=self.var_prog_done, width=10).pack(side="left", padx=(0, 16))
            tk.Label(row_p1, text="Đang thực hiện:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_p1, textvariable=self.var_prog_inprogress, width=10).pack(side="left")

            row_p2 = tk.Frame(frm_prog)
            row_p2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_p2, text="Chưa thực hiện:", width=16, anchor="w").pack(side="left")
            tk.Label(row_p2, textvariable=self.var_prog_remaining, font=("Arial", 10, "bold"),
                    fg="#c62828", width=10, anchor="w").pack(side="left", padx=(0, 16))
            tk.Label(row_p2, text="Tỷ lệ hoàn thành:", width=14, anchor="w").pack(side="left")
            tk.Label(row_p2, textvariable=self.var_prog_percent, font=("Arial", 10, "bold"),
                    fg="#2e7d32", width=10, anchor="w").pack(side="left")

            self.prog_picker = SourcePicker(
                frm_prog, "Nguồn PDF đã thực hiện (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.prog_picker.pack(fill="x", padx=6, pady=3)

            row_p3 = tk.Frame(frm_prog)
            row_p3.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_p3, text="Tự đếm số file PDF đã thực hiện (ghi đè ô 'Đã thực hiện' ở trên)",
                           variable=self.var_prog_auto_count).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_p3, text="Tự nhận diện Tờ/Thửa từ TÊN FILE",
                           variable=self.var_prog_detect_filename).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_p3, text="Tự nhận diện từ NỘI DUNG PDF nếu tên file thiếu (chậm hơn)",
                           variable=self.var_prog_detect_content).pack(side="left")

            row_p_run = tk.Frame(frm_prog)
            row_p_run.pack(fill="x", padx=6, pady=(6, 4))
            tk.Button(row_p_run, text="🔄 CẬP NHẬT TIẾN ĐỘ", font=("Arial", 10, "bold"),
                      bg="#00838f", fg="white", command=self.run_progress_update_clicked).pack(side="left", padx=(0, 6))
            tk.Button(row_p_run, text="📊 HIỆN BIỂU ĐỒ", font=("Arial", 10, "bold"),
                      bg="#5e35b1", fg="white", command=self.show_progress_chart).pack(side="left", padx=6)
            tk.Button(row_p_run, text="💾 XUẤT BÁO CÁO TIẾN ĐỘ", font=("Arial", 10, "bold"),
                      bg="#2e7d32", fg="white", command=self.export_progress_report_clicked).pack(side="left", padx=6)
            self.prog_pause, self.prog_cancel = self._add_pause_cancel(row_p_run, "control_prog")
            self.prog_pause.pack(side="left", padx=(8, 3))
            self.prog_cancel.pack(side="left", padx=3)

            tk.Label(frm_prog, text="Danh sách đã thực hiện (sau khi Cập nhật tiến độ):",
                    font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 2))
            cols_p = ("stt", "maxa", "to", "thua", "file", "status", "note")
            headers_p = ["STT", "Mã xã", "Số tờ", "Số thửa", "Tên file PDF", "Trạng thái", "Ghi chú"]
            widths_p = [45, 55, 50, 55, 220, 100, 220]
            self.tree_prog = ttk.Treeview(frm_prog, columns=cols_p, show="headings", height=8)
            for c, h, w in zip(cols_p, headers_p, widths_p):
                self.tree_prog.heading(c, text=h)
                self.tree_prog.column(c, width=w, anchor="w")
            vsb_p = tk.Scrollbar(frm_prog, orient="vertical", command=self.tree_prog.yview)
            self.tree_prog.configure(yscrollcommand=vsb_p.set)
            self.tree_prog.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_p.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            # ================= 8c: Điền dữ liệu vào Excel MẪU có sẵn (giữ nguyên định dạng) =================
            frm3 = styled_labelframe(
                parent, text="8c. Điền dữ liệu từ PDF vào file Excel MẪU có sẵn (giữ nguyên định dạng/merge/công thức)")
            frm3.pack(fill="both", expand=True, **pad)
            tk.Label(frm3,
                     text="Khác với 8a (tạo file Excel MỚI hoàn toàn), mục này MỞ file Excel mẫu bạn đang dùng "
                          "(VD 'Bản sao của CHUA_CO_GIAY_xaVanLangcu.xlsx'), giữ nguyên toàn bộ tiêu đề/merge cell/"
                          "công thức/bố cục, chỉ điền dữ liệu vào các cột: B, H, I, J, K, L, V, W, X, Y, Z, AA.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_c1 = tk.Frame(frm3)
            row_c1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c1, text="① File Excel MẪU:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_c1, textvariable=self.var_wb8c_template).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_c1, text="Chọn...", command=self.pick_wb8c_template).pack(side="left")

            self.wb8c_picker = SourcePicker(
                frm3, "② Nguồn PDF (chọn 1 thư mục, HOẶC chọn 1/nhiều file PDF bất kỳ)",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.wb8c_picker.pack(fill="x", padx=6, pady=3)

            row_c3 = tk.Frame(frm3)
            row_c3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c3, text="③ Thư mục xuất Excel kết quả:", width=28, anchor="w").pack(side="left")
            tk.Entry(row_c3, textvariable=self.var_wb8c_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_c3, text="Chọn...", command=self.pick_wb8c_output).pack(side="left")

            row_c4 = tk.Frame(frm3)
            row_c4.pack(fill="x", padx=6, pady=3)
            tk.Label(row_c4, text="④ Tên sheet:").pack(side="left")
            tk.Entry(row_c4, textvariable=self.var_wb8c_sheet, width=15).pack(side="left", padx=(4, 20))
            tk.Label(row_c4, text="⑤ Dòng bắt đầu ghi dữ liệu:").pack(side="left")
            tk.Entry(row_c4, textvariable=self.var_wb8c_header_row, width=6).pack(side="left", padx=4)

            # ===== Cấu hình cột Excel báo cáo (động, không cố định cứng) =====
            frm_cols = styled_labelframe(frm3, text="⚙ Cấu hình cột Excel báo cáo (tùy chọn - để trống cột nào thì KHÔNG ghi trường đó)",
                                     font=("Segoe UI", 9, "bold"), fg="#1a237e")
            frm_cols.pack(fill="x", padx=6, pady=(6, 4))
            col_fields = [
                ("maxa", "Mã xã"), ("ten", "Họ tên"), ("cccd", "CCCD/Định danh"), ("ngaysinh", "Ngày sinh"),
                ("gioitinh", "Giới tính"), ("diachitt", "Địa chỉ thường trú"), ("soto", "Số tờ"),
                ("sothua", "Số thửa"), ("diachithua", "Địa chỉ thửa"), ("dientich", "Diện tích"),
                ("mucdich", "Mục đích SD"), ("dientichloai1", "DT loại đất 1"),
            ]
            grid = tk.Frame(frm_cols)
            grid.pack(fill="x", padx=6, pady=4)
            for idx, (key, label) in enumerate(col_fields):
                r, c = divmod(idx, 4)
                cell = tk.Frame(grid)
                cell.grid(row=r, column=c, sticky="w", padx=6, pady=2)
                tk.Label(cell, text=f"{label}:", width=15, anchor="w").pack(side="left")
                tk.Entry(cell, textvariable=self.var_col8c[key], width=5).pack(side="left")
            row_cols_btn = tk.Frame(frm_cols)
            row_cols_btn.pack(fill="x", padx=6, pady=(2, 6))
            tk.Button(row_cols_btn, text="↺ Khôi phục mặc định", command=self.reset_col8c_config).pack(side="left", padx=(0, 6))
            tk.Button(row_cols_btn, text="💾 Lưu cấu hình cột", command=self.save_col8c_config).pack(side="left", padx=6)
            tk.Button(row_cols_btn, text="📂 Tải cấu hình cột đã lưu", command=self.load_col8c_config).pack(side="left", padx=6)
            self.lbl_col8c_status = tk.Label(row_cols_btn, text="", fg="#2e7d32")
            self.lbl_col8c_status.pack(side="left", padx=10)

            row_c5 = tk.Frame(frm3)
            row_c5.pack(fill="x", padx=6, pady=(6, 3))
            tk.Checkbutton(row_c5, text="⑥ Xóa dữ liệu cũ từ dòng bắt đầu trước khi tổng hợp",
                           variable=self.var_wb8c_clear_old).pack(side="left", padx=(0, 20))
            tk.Checkbutton(row_c5, text="⑦ Tự suy năm sinh/giới tính từ CCCD nếu PDF không có "
                                         "(lưu ý: CCCD chỉ suy được NĂM sinh, không suy được ngày/tháng)",
                           variable=self.var_wb8c_infer_cccd).pack(side="left")

            row_c5b = tk.Frame(frm3)
            row_c5b.pack(fill="x", padx=6, pady=(0, 3))
            tk.Checkbutton(row_c5b, text="Không ghi đè ô đã có dữ liệu (chỉ ghi vào ô trống)",
                           variable=self.var_wb8c_no_overwrite).pack(side="left", padx=(0, 20))
            tk.Checkbutton(row_c5b, text="Cảnh báo khi trùng Mã xã+Tờ+Thửa",
                           variable=self.var_wb8c_warn_dup).pack(side="left", padx=(0, 20))
            tk.Checkbutton(row_c5b, text="Xuất danh sách file cần kiểm tra (.csv)",
                           variable=self.var_wb8c_export_check).pack(side="left", padx=(0, 20))
            chk_debug_text = tk.Checkbutton(row_c5b, text="Xuất text debug từng file (chẩn đoán khi tổng hợp sai)",
                                            variable=self.var_wb8c_debug_text)
            chk_debug_text.pack(side="left")
            self._register_advanced_widget(chk_debug_text, side="left")

            row_c_run = tk.Frame(frm3)
            row_c_run.pack(fill="x", padx=6, pady=(6, 8))
            self.btn_wb8c_dryrun = tk.Button(row_c_run, text="⑧ CHẠY THỬ (chỉ kiểm tra, KHÔNG ghi Excel)",
                                              font=("Arial", 10, "bold"), bg="#455a64", fg="white",
                                              command=lambda: self.run_wb8c_clicked(dry_run=True))
            self.btn_wb8c_dryrun.pack(side="left", padx=(0, 6))
            self.btn_wb8c_run = tk.Button(row_c_run, text="⑨ TỔNG HỢP BÁO CÁO",
                                           font=("Arial", 10, "bold"), bg="#2e7d32", fg="white",
                                           command=lambda: self.run_wb8c_clicked(dry_run=False))
            self.btn_wb8c_run.pack(side="left")
            self.btn_wb8c_pause, self.btn_wb8c_cancel = self._add_pause_cancel(row_c_run, "control_wb8c")
            self.btn_wb8c_pause.pack(side="left", padx=(12, 3))
            self.btn_wb8c_cancel.pack(side="left", padx=3)
            self.btn_wb8c_retry = tk.Button(row_c_run, text="🔁 Chạy lại file lỗi/cần kiểm tra",
                                            state="disabled", command=self.retry_wb8c_failed)
            self.btn_wb8c_retry.pack(side="left", padx=(12, 0))

            row_wb8c_progress = tk.Frame(frm3); row_wb8c_progress.pack(fill="x", padx=6, pady=(0, 8))
            self.pb_wb8c = ttk.Progressbar(row_wb8c_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_wb8c.pack(side="left", padx=(0, 10))
            self.lbl_wb8c_progress = tk.Label(row_wb8c_progress, text="", fg="#555", anchor="w")
            self.lbl_wb8c_progress.pack(side="left", fill="x", expand=True)

            tk.Label(frm3, text="⑩ Bảng log:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 2))
            cols = ("file", "maxa", "ten", "cccd", "diachitt", "so_to_goc", "to", "thua", "dt",
                   "loai_dat", "nguon", "diem_tin_cay", "status", "note")
            headers = ["Tên file PDF", "Mã xã", "Họ tên", "CCCD", "Địa chỉ thường trú", "Tờ gốc",
                      "Tờ chuẩn", "Thửa", "Diện tích", "Loại đất", "Nguồn", "Điểm tin cậy",
                      "Trạng thái", "Ghi chú lỗi"]
            widths = [140, 55, 120, 95, 150, 50, 55, 45, 65, 60, 55, 75, 110, 180]
            self.tree_wb8c = ttk.Treeview(frm3, columns=cols, show="headings", height=10)
            for c, h, w in zip(cols, headers, widths):
                self.tree_wb8c.heading(c, text=h)
                self.tree_wb8c.column(c, width=w, anchor="w")
            vsb3 = tk.Scrollbar(frm3, orient="vertical", command=self.tree_wb8c.yview)
            self.tree_wb8c.configure(yscrollcommand=vsb3.set)
            self.tree_wb8c.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb3.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            # ================= 8d. Excel tổng khớp PDF + báo cáo So sánh =================
            frm4 = styled_labelframe(
                parent, text="8d. Lấy dữ liệu từ Excel tổng đã chỉnh sửa khớp PDF + báo cáo có sheet So_sánh")
            frm4.pack(fill="x", padx=10, pady=(0, 10))
            tk.Label(frm4,
                    text="Đối chiếu PDF với Excel TỔNG đã chỉnh sửa (theo Mã xã+Số tờ+Số thửa) - dữ liệu ghi vào\n"
                         "báo cáo lấy từ Excel TỔNG (không phải PDF). File báo cáo kết quả có thêm sheet 'So_sanh'\n"
                         "để xem trường nào đã được chỉnh sửa so với nội dung PDF gốc.",
                    fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=8, pady=(8, 6))

            row_tong = tk.Frame(frm4); row_tong.pack(fill="x", padx=8, pady=3)
            tk.Label(row_tong, text="File Excel tổng đã chỉnh sửa:", width=22, anchor="w").pack(side="left")
            self.var_tong_excel = tk.StringVar()
            tk.Entry(row_tong, textvariable=self.var_tong_excel, width=45).pack(side="left", padx=(0, 6))
            tk.Button(row_tong, text="Chọn file...", command=self.pick_tong_excel).pack(side="left")
            tk.Label(row_tong, text="  Sheet:").pack(side="left", padx=(12, 4))
            self.var_tong_sheet = tk.StringVar()
            tk.Entry(row_tong, textvariable=self.var_tong_sheet, width=14).pack(side="left")
            tk.Label(row_tong, text="  Dòng tiêu đề:").pack(side="left", padx=(12, 4))
            self.var_tong_header_row = tk.StringVar(value="1")
            tk.Entry(row_tong, textvariable=self.var_tong_header_row, width=5).pack(side="left")

            row_tong_cols = tk.Frame(frm4); row_tong_cols.pack(fill="x", padx=8, pady=3)
            tk.Label(row_tong_cols, text="Cột khóa Excel tổng:", width=22, anchor="w").pack(side="left")
            _tong_col_labels = [("maxa", "Mã xã"), ("soto", "Số tờ"), ("sothua", "Số thửa")]
            self.var_col_tong = {}
            for key, label in _tong_col_labels:
                tk.Label(row_tong_cols, text=f"{label}:").pack(side="left", padx=(8, 2))
                self.var_col_tong[key] = tk.StringVar(value={"maxa": "D", "soto": "B", "sothua": "C"}[key])
                tk.Entry(row_tong_cols, textvariable=self.var_col_tong[key], width=4).pack(side="left")

            row_tong_cols2 = tk.Frame(frm4); row_tong_cols2.pack(fill="x", padx=8, pady=3)
            tk.Label(row_tong_cols2, text="Cột dữ liệu Excel tổng:", width=22, anchor="w").pack(side="left")
            _tong_data_labels = [("ten", "Tên", "K"), ("cccd", "CCCD", "L"), ("ngaysinh", "N.sinh", "M"),
                                ("gioitinh", "G.tính", "N"), ("diachitt", "ĐC t.trú", "P"),
                                ("diachithua", "ĐC thửa", "Q"), ("dientich", "D.tích", "H"),
                                ("mucdich", "M.đích", ""), ("loai_dat", "Loại đất", "G")]
            for key, label, default in _tong_data_labels:
                tk.Label(row_tong_cols2, text=f"{label}:").pack(side="left", padx=(6, 2))
                self.var_col_tong[key] = tk.StringVar(value=default)
                tk.Entry(row_tong_cols2, textvariable=self.var_col_tong[key], width=4).pack(side="left")

            row_baocao = tk.Frame(frm4); row_baocao.pack(fill="x", padx=8, pady=3)
            tk.Label(row_baocao, text="File Excel báo cáo mẫu:", width=22, anchor="w").pack(side="left")
            self.var_tong_report_template = tk.StringVar()
            tk.Entry(row_baocao, textvariable=self.var_tong_report_template, width=45).pack(side="left", padx=(0, 6))
            tk.Button(row_baocao, text="Chọn file...", command=self.pick_tong_report_template).pack(side="left")
            tk.Label(row_baocao, text="  Sheet:").pack(side="left", padx=(12, 4))
            self.var_tong_report_sheet = tk.StringVar(value="Chưa GCN")
            tk.Entry(row_baocao, textvariable=self.var_tong_report_sheet, width=14).pack(side="left")
            tk.Label(row_baocao, text="  Dòng bắt đầu ghi:").pack(side="left", padx=(12, 4))
            self.var_tong_report_header_row = tk.StringVar(value="2041")
            tk.Entry(row_baocao, textvariable=self.var_tong_report_header_row, width=6).pack(side="left")

            row_pdf4 = tk.Frame(frm4); row_pdf4.pack(fill="x", padx=8, pady=3)
            tk.Label(row_pdf4, text="Thư mục PDF:", width=22, anchor="w").pack(side="left")
            self.tong_pdf_picker = SourcePicker(row_pdf4, "Nguồn PDF", [("PDF files", "*.pdf")], file_label="PDF")
            self.tong_pdf_picker.pack(side="left", fill="x", expand=True)

            row_pdf4b = tk.Frame(frm4); row_pdf4b.pack(fill="x", padx=8, pady=(0, 3))
            tk.Label(row_pdf4b, text="  hoặc file nén PDF (.zip/.rar):", anchor="w").pack(side="left", padx=(22, 4))
            self.var_tong_pdf_zip = tk.StringVar()
            tk.Entry(row_pdf4b, textvariable=self.var_tong_pdf_zip, width=35).pack(side="left", padx=(0, 6))
            tk.Button(row_pdf4b, text="Chọn file .zip...", command=self.pick_tong_pdf_zip).pack(side="left", padx=(0, 4))
            tk.Button(row_pdf4b, text="Chọn file .rar...", command=self.pick_tong_pdf_rar).pack(side="left")
            tk.Label(row_pdf4b, text="(.rar cần bổ sung tools/7zip/7z.exe - xem README trong thư mục đó)",
                    fg="#888").pack(side="left", padx=6)

            row_tong_run = tk.Frame(frm4); row_tong_run.pack(fill="x", padx=8, pady=(6, 3))
            self.btn_tong_dryrun = tk.Button(row_tong_run, text="🔍 Chạy thử đối chiếu", bg="#00838f", fg="white",
                                             command=lambda: self.run_tong_reconcile_clicked(dry_run=True))
            self.btn_tong_dryrun.pack(side="left", padx=(0, 6))
            self.btn_tong_run = tk.Button(row_tong_run, text="✓ Tạo báo cáo + sheet So_sánh", bg="#2e7d32", fg="white",
                                          command=lambda: self.run_tong_reconcile_clicked(dry_run=False))
            self.btn_tong_run.pack(side="left")
            self.btn_tong_pause, self.btn_tong_cancel = self._add_pause_cancel(row_tong_run, "control_tong")
            self.btn_tong_pause.pack(side="left", padx=(12, 3))

            row_tong_progress = tk.Frame(frm4); row_tong_progress.pack(fill="x", padx=8, pady=(0, 3))
            self.pb_tong = ttk.Progressbar(row_tong_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_tong.pack(side="left", padx=(0, 10))
            self.lbl_tong_progress = tk.Label(row_tong_progress, text="", fg="#555", anchor="w")
            self.lbl_tong_progress.pack(side="left", fill="x", expand=True)

            row_tong_opt = tk.Frame(frm4); row_tong_opt.pack(fill="x", padx=8, pady=(0, 3))
            self.var_tong_doc_so_sanh = tk.BooleanVar(value=True)
            tk.Checkbutton(row_tong_opt,
                           text="Đọc nội dung PDF để so sánh chi tiết (sheet So_sánh) - TẮT để đối chiếu NHANH HƠN "
                                "nhiều với lô lớn khi chỉ cần điền báo cáo, không cần xem khác biệt so với PDF gốc",
                           variable=self.var_tong_doc_so_sanh).pack(anchor="w")

            row_tong_opt2 = tk.Frame(frm4); row_tong_opt2.pack(fill="x", padx=8, pady=(0, 3))
            tk.Label(row_tong_opt2, text="Sai số Diện tích coi là \"gần bằng\" khi đối chiếu (m²):").pack(side="left")
            self.var_tong_dientich_dung_sai = tk.StringVar(value="1.0")
            tk.Entry(row_tong_opt2, textvariable=self.var_tong_dientich_dung_sai, width=8).pack(side="left", padx=(6, 0))
            tk.Label(row_tong_opt2, text="(dùng khi có nhiều dòng Excel tổng cùng Tờ+Thửa, cần Diện tích để thu hẹp)",
                    fg="#888").pack(side="left", padx=6)
            self._register_advanced_widget(row_tong_opt2, fill="x", padx=8, pady=(0, 3))

            row_tong_opt3 = tk.Frame(frm4); row_tong_opt3.pack(fill="x", padx=8, pady=(0, 3))
            self.var_tong_luu_tru = tk.BooleanVar(value=True)
            tk.Checkbutton(row_tong_opt3,
                           text="Lưu thêm 1 bản vào kho lưu trữ báo cáo nội bộ (Output/Bao_cao_da_hoan_thien/) "
                                "và ghi lịch sử để tra cứu/phục hồi sau này",
                           variable=self.var_tong_luu_tru).pack(anchor="w")

            row_tong_opt4 = tk.Frame(frm4); row_tong_opt4.pack(fill="x", padx=8, pady=(0, 3))
            self.var_tong_clear_old = tk.BooleanVar(value=False)
            tk.Checkbutton(row_tong_opt4,
                           text="Xóa dữ liệu CŨ trong vùng báo cáo trước khi ghi mới (chỉ xóa giá trị ở đúng các "
                                "cột đã ánh xạ - KHÔNG đụng công thức/macro/định dạng/merge cell)",
                           variable=self.var_tong_clear_old).pack(anchor="w")

            row_tong_archive = tk.Frame(frm4); row_tong_archive.pack(fill="x", padx=8, pady=(0, 6))
            tk.Button(row_tong_archive, text="📂 Xem lịch sử báo cáo đã lưu trữ",
                      command=self.open_report_history_window).pack(side="left")

            tk.Label(frm4, text="Bảng xem trước (So_sánh):", font=("Arial", 9, "bold")).pack(anchor="w", padx=8, pady=(4, 2))
            cols4 = ("file", "truong", "pdf_val", "excel_val", "ketqua", "chinhsua")
            headers4 = ["Tên file PDF", "Trường so sánh", "Giá trị PDF", "Giá trị Excel tổng", "Kết quả", "Chỉnh sửa"]
            widths4 = [180, 130, 160, 160, 120, 80]
            self.tree_tong = ttk.Treeview(frm4, columns=cols4, show="headings", height=8)
            for c, h, w in zip(cols4, headers4, widths4):
                self.tree_tong.heading(c, text=h)
                self.tree_tong.column(c, width=w, anchor="w")
            vsb4 = tk.Scrollbar(frm4, orient="vertical", command=self.tree_tong.yview)
            self.tree_tong.configure(yscrollcommand=vsb4.set)
            self.tree_tong.pack(side="left", fill="both", expand=True, padx=(8, 0), pady=(0, 8))
            vsb4.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

        def pick_report_output(self):
            import time
            default_name = f"TONG_HOP_CHUA_CO_GIAY_TU_PDF_{time.strftime('%Y%m%d_%H%M')}.xlsx"
            p = filedialog.asksaveasfilename(title="Lưu file Excel báo cáo tại", defaultextension=".xlsx",
                                              filetypes=[("Excel files", "*.xlsx")],
                                              initialfile=default_name)
            if p:
                self.var_report_output.set(p)

        def run_report_clicked(self):
            source = self.report_picker.get_source()
            output_path = self.var_report_output.get().strip()

            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF cần tổng hợp.")
                return
            if not output_path:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn nơi lưu file Excel báo cáo.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return

            mapping_text = self.txt_xa_mapping.get("1.0", "end")
            xa_mapping = parse_xa_mapping(mapping_text)
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            if not xa_mapping and not thon_mapping:
                messagebox.showerror("Thiếu bảng tra", "Bảng tra Thôn/Xã → Mã xã (ở Tab 5) đang trống.")
                return

            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            self.txt_log.delete("1.0", "end")

            checkpoint_task_id = None
            files_da_xong = None
            from app.services import checkpoint_service as _ckpt
            checkpoint_task_id = _ckpt.tinh_task_id("tong_hop_bao_cao_pdf", str(source), "")
            data_cu = _ckpt.doc_checkpoint(checkpoint_task_id)
            if data_cu and data_cu.get("trang_thai") == "DANG_XU_LY":
                so_da_xong = len(data_cu.get("da_hoan_thanh", []))
                so_tong = len(data_cu.get("danh_sach_file", []))
                if messagebox.askyesno(
                    "Phát hiện tác vụ dở dang",
                    f"Phát hiện đợt tổng hợp báo cáo TRƯỚC ĐÓ chưa hoàn thành (đã đọc {so_da_xong}/{so_tong} "
                    f"file, có thể do phần mềm bị đóng đột ngột).\n\n"
                    f"Bấm CÓ để TIẾP TỤC từ chỗ dở dang.\nBấm KHÔNG để BẮT ĐẦU LẠI TỪ ĐẦU."):
                    files_da_xong = set(data_cu.get("da_hoan_thanh", []))
                else:
                    _ckpt.xoa_checkpoint(checkpoint_task_id)

            self.btn_report.config(state="disabled")
            self._enable_pause_cancel("control_report", self.btn_report_pause, self.btn_report_cancel)
            self.pb_report.config(value=0, maximum=100)
            self._report_start_time = time.time()

            from app.services import task_manager_service as _tm
            self.task_info_report = _tm.TaskInfo(
                ten_chuc_nang="8a. Tổng hợp báo cáo ra Excel", thu_muc_dau_ra="",
                trang_thai=_tm.TrangThaiTask.DANG_KHOI_TAO)

            def update_report_progress(idx, total, fname):
                elapsed = time.time() - self._report_start_time
                self.pb_report.config(value=idx, maximum=total)
                self.lbl_report_progress.config(text=f"Đang đọc {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_report_progress(idx, total, fname))
                self.task_info_report.cap_nhat_tien_do(tien_do_toan_bo=idx,
                                                       trang_thai=_tm.TrangThaiTask.DANG_DOC_FILE)
                self.task_info_report.tong_so_luong = total
                append_tab_log_row(8, "8a. Tổng hợp báo cáo ra Excel", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_DOC_FILE")

            def worker():
                try:
                    perf_log_path = None
                    if getattr(self.app_config, "word_hien_thi_thong_ke_hieu_nang", True):
                        perf_log_path = os.path.join(get_standard_logs_dir(), "LOG_HIEU_NANG_XU_LY.csv")
                    rows, error_files = compile_report_from_pdfs(
                        source, xa_mapping, thon_mapping, use_ocr, ocr_dpi, self.log,
                        debug=self.var_debug_content.get(), control=self.control_report, progress_cb=progress_cb,
                        toc_do_xu_ly=getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"),
                        so_file_moi_dot=getattr(self.app_config, "word_so_file_moi_dot", 100),
                        nghi_giua_dot_giay=getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0),
                        perf_log_path=perf_log_path,
                        checkpoint_task_id=checkpoint_task_id, files_da_xong=files_da_xong)
                    export_report_from_pdfs(output_path, rows, error_files)

                    self.task_info_report.so_luong_thanh_cong = len(rows)
                    self.task_info_report.so_luong_loi = len(error_files)
                    self.task_info_report.trang_thai = _tm.TrangThaiTask.THANH_CONG

                    self.after(0, lambda: self.lbl_report_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._report_start_time)}s)"))
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {len(rows)} thửa đất được tổng hợp | {len(error_files)} file lỗi/thiếu thông tin")
                    self.log(f"Đã lưu báo cáo: {output_path}")
                    self.after(0, lambda: self.mark_workflow_step("bao_cao", "Đã hoàn thành"))
                    messagebox.showinfo("Hoàn tất tổng hợp báo cáo",
                                         f"Đã tổng hợp: {len(rows)} thửa đất\nFile lỗi: {len(error_files)}\n\n"
                                         f"Báo cáo lưu tại:\n{output_path}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_report.config(state="normal")
                    self._disable_pause_cancel(self.btn_report_pause, self.btn_report_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 8b: Thống kê tiến độ từ Excel bất kỳ -------------------

        def pick_stat_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel", filetypes=[("Excel files", "*.xlsx *.xlsm")])
            if p:
                self.var_stat_excel.set(p)

        def run_stat_clicked(self):
            excel_path = self.var_stat_excel.get().strip()
            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel hợp lệ.")
                return

            col_status = self.var_stat_col_status.get().strip() or None
            col_gcn = self.var_stat_col_gcn.get().strip() or None
            if not col_status and not col_gcn:
                messagebox.showerror("Thiếu thông tin", "Vui lòng nhập ít nhất 1 trong 2 cột: Trạng thái thực hiện hoặc GCN.")
                return

            try:
                header_row = int(self.var_stat_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            target_total = None
            raw_target = self.var_stat_target_total.get().strip()
            if raw_target:
                try:
                    target_total = float(raw_target)
                except ValueError:
                    messagebox.showerror("Lỗi", "Tổng khối lượng được giao phải là số.")
                    return

            done_values = [v.strip() for v in self.var_stat_done_values.get().split(",")]
            gcn_values = [v.strip() for v in self.var_stat_gcn_values.get().split(",")]

            self.txt_log.delete("1.0", "end")
            self.btn_stat.config(state="disabled")
            self.pb_stat.config(value=0, maximum=100)
            self._stat_start_time = time.time()

            def update_stat_progress(idx, total, mota):
                elapsed = time.time() - self._stat_start_time
                self.pb_stat.config(value=idx, maximum=total)
                self.lbl_stat_progress.config(text=f"Đang quét {idx}/{total}: {mota} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, mota):
                self.after(0, lambda: update_stat_progress(idx, total, mota))
                append_tab_log_row(8, "8b. Thống kê tiến độ từ Excel", mota,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XU_LY")

            def worker():
                try:
                    result = analyze_excel_progress(
                        excel_path=excel_path,
                        sheet_name=self.var_stat_sheet.get().strip() or None,
                        header_row=header_row,
                        col_count=self.var_stat_col_count.get().strip() or None,
                        col_status=col_status, done_values=done_values,
                        col_gcn=col_gcn, gcn_done_values=gcn_values,
                        target_total=target_total, log_cb=self.log, progress_cb=progress_cb,
                    )
                    self.after(0, lambda: self.lbl_stat_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._stat_start_time)}s)"))

                    self.log("=" * 60)
                    self.log(f"Tổng số dòng có dữ liệu: {result['total_rows']}")
                    self.log(f"Tổng khối lượng được giao (mẫu số tính %): {result['target_total']}")
                    if result["col_status_used"]:
                        self.log("--- Tiến độ thực hiện ---")
                        self.log(f"  Đã thực hiện: {result['done_count']}")
                        self.log(f"  Chưa thực hiện: {result['not_done_count']}")
                        self.log(f"  Bỏ trống trạng thái: {result['status_blank_count']}")
                        self.log(f"  Tỷ lệ đã thực hiện: {result['percent_done']}%")
                    if result["col_gcn_used"]:
                        self.log("--- Giấy chứng nhận (GCN) ---")
                        self.log(f"  Đã nhập GCN: {result['gcn_done_count']}")
                        self.log(f"  Chưa nhập GCN: {result['gcn_notdone_count']}")
                        self.log(f"  Bỏ trống GCN: {result['gcn_blank_count']}")
                        self.log(f"  Tỷ lệ đã nhập GCN: {result['percent_gcn_done']}%")

                    save_path = filedialog.asksaveasfilename(
                        title="Lưu file thống kê tại", defaultextension=".xlsx",
                        filetypes=[("Excel files", "*.xlsx")], initialfile="THONG_KE_TIEN_DO.xlsx")
                    if save_path:
                        export_progress_report(save_path, result, {"excel_path": excel_path})
                        self.log(f"\nĐã lưu file thống kê: {save_path}")

                    messagebox.showinfo("Hoàn tất thống kê", "Đã thống kê xong, xem chi tiết ở khung Nhật ký.")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_stat.config(state="normal")

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- Khối lượng và tiến độ thực hiện -------------------

        def _refresh_prog_tree(self, results):
            for item in self.tree_prog.get_children():
                self.tree_prog.delete(item)
            for r in results:
                self.tree_prog.insert("", "end", values=(r["stt"], r["maxa"], r["to"], r["thua"],
                                                          r["file"], r["status"], r["note"]))

        def _parse_prog_int(self, var, label):
            raw = var.get().strip()
            if not raw:
                return 0
            try:
                return int(float(raw))
            except ValueError:
                messagebox.showerror("Lỗi", f"{label} phải là số.")
                return None

        def run_progress_update_clicked(self):
            total = self._parse_prog_int(self.var_prog_total, "Tổng khối lượng")
            if total is None:
                return
            done_manual = self._parse_prog_int(self.var_prog_done, "Đã thực hiện")
            if done_manual is None:
                return
            in_progress = self._parse_prog_int(self.var_prog_inprogress, "Đang thực hiện")
            if in_progress is None:
                return

            source = self.prog_picker.get_source()
            auto_count = self.var_prog_auto_count.get()

            if not source:
                # Không chọn nguồn PDF -> chỉ tính theo số nhập thủ công
                chua, ty_le = compute_progress(total, done_manual, in_progress)
                self.var_prog_remaining.set(str(chua))
                self.var_prog_percent.set(f"{ty_le}%")
                self._refresh_prog_tree([])
                self.prog_scan_results = []
                self.log(f"Đã cập nhật tiến độ (nhập thủ công): Tổng {total} | Đã {done_manual} | "
                         f"Đang {in_progress} | Chưa {chua} | Tỷ lệ {ty_le}%")
                return

            detect_filename = self.var_prog_detect_filename.get()
            detect_content = self.var_prog_detect_content.get()
            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            self.txt_log.delete("1.0", "end")
            self._enable_pause_cancel("control_prog", self.prog_pause, self.prog_cancel)

            def worker():
                try:
                    results = scan_completed_pdfs(source, xa_mapping, thon_mapping, detect_filename,
                                                  detect_content, use_ocr, ocr_dpi, self.log,
                                                  control=self.control_prog)
                    self.prog_scan_results = results
                    done_count = sum(1 for r in results if r["status"] == "DA_NHAN_DIEN")
                    effective_done = len(results) if auto_count else done_manual
                    chua, ty_le = compute_progress(total, effective_done, in_progress)

                    def update_ui():
                        self.var_prog_remaining.set(str(chua))
                        self.var_prog_percent.set(f"{ty_le}%")
                        if auto_count:
                            self.var_prog_done.set(str(len(results)))
                        self._refresh_prog_tree(results)

                    self.after(0, update_ui)
                    self.log("=" * 60)
                    self.log(f"Đã quét {len(results)} file PDF ({done_count} nhận diện được Tờ/Thửa, "
                             f"{len(results) - done_count} cần kiểm tra).")
                    self.log(f"Tổng {total} | Đã {effective_done} | Đang {in_progress} | Chưa {chua} | "
                             f"Tỷ lệ hoàn thành {ty_le}%")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self._disable_pause_cancel(self.prog_pause, self.prog_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def show_progress_chart(self):
            if not HAS_MATPLOTLIB:
                messagebox.showerror("Thiếu thư viện",
                                     "Cần cài: pip install matplotlib\nSau đó khởi động lại phần mềm.")
                return
            try:
                total = int(float(self.var_prog_total.get().strip() or "0"))
                done = int(float(self.var_prog_done.get().strip() or "0"))
                in_progress = int(float(self.var_prog_inprogress.get().strip() or "0"))
            except ValueError:
                messagebox.showerror("Lỗi", "Vui lòng nhập số hợp lệ và bấm 'Cập nhật tiến độ' trước.")
                return
            chua, ty_le = compute_progress(total, done, in_progress)

            if total <= 0:
                messagebox.showinfo("Chưa có dữ liệu", "Vui lòng nhập Tổng khối lượng và bấm 'Cập nhật tiến độ' trước.")
                return

            win = tk.Toplevel(self)
            win.title("Biểu đồ kết quả thực hiện")
            win.geometry("900x480")

            fig = Figure(figsize=(9, 4.5), dpi=100)

            ax1 = fig.add_subplot(1, 2, 1)
            labels = ["Đã thực hiện", "Đang thực hiện", "Chưa thực hiện"]
            values = [done, in_progress, chua]
            colors = ["#2e7d32", "#f57c00", "#c62828"]
            non_zero = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
            if non_zero:
                ax1.pie([v for _, v, _ in non_zero], labels=[l for l, _, _ in non_zero],
                       colors=[c for _, _, c in non_zero], autopct="%1.1f%%", startangle=90,
                       wedgeprops={"width": 0.4})
            ax1.set_title(f"Tỷ lệ hoàn thành: {ty_le}%")

            ax2 = fig.add_subplot(1, 2, 2)
            cats = ["Tổng", "Đã TH", "Đang TH", "Chưa TH"]
            vals2 = [total, done, in_progress, chua]
            bars = ax2.bar(cats, vals2, color=["#1a237e", "#2e7d32", "#f57c00", "#c62828"])
            ax2.set_title("Khối lượng thực hiện")
            for b, v in zip(bars, vals2):
                ax2.text(b.get_x() + b.get_width() / 2, v, str(v), ha="center", va="bottom", fontsize=9)

            fig.tight_layout()
            canvas = FigureCanvasTkAgg(fig, master=win)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True)

        def export_progress_report_clicked(self):
            try:
                total = int(float(self.var_prog_total.get().strip() or "0"))
                done = int(float(self.var_prog_done.get().strip() or "0"))
                in_progress = int(float(self.var_prog_inprogress.get().strip() or "0"))
            except ValueError:
                messagebox.showerror("Lỗi", "Vui lòng nhập số hợp lệ.")
                return
            if total <= 0:
                messagebox.showinfo("Chưa có dữ liệu", "Vui lòng nhập Tổng khối lượng và bấm 'Cập nhật tiến độ' trước.")
                return
            chua, ty_le = compute_progress(total, done, in_progress)

            save_dir = filedialog.askdirectory(title="Chọn thư mục lưu báo cáo tiến độ")
            if not save_dir:
                return
            stamp = time.strftime("%Y%m%d_%H%M%S")
            path = os.path.join(save_dir, f"BAO_CAO_TIEN_DO_THUC_HIEN_{stamp}.xlsx")

            completed_list = [r for r in self.prog_scan_results if r["status"] == "DA_NHAN_DIEN"]
            check_list = [r for r in self.prog_scan_results if r["status"] != "DA_NHAN_DIEN"]

            try:
                export_execution_progress_report(path, total, done, in_progress, chua, ty_le,
                                                 completed_list, check_list)
                self.log(f"✓ Đã xuất báo cáo tiến độ: {path}")
                messagebox.showinfo("Hoàn tất", f"Đã xuất báo cáo tiến độ:\n{path}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không xuất được báo cáo: {e}")

        # ------------------- MỤC 8c: Điền vào Excel mẫu có sẵn -------------------

        COL8C_CONFIG_FILENAME = "Cau_hinh_cot_bao_cao.json"

        def _validate_excel_column(self, col):
            """Kiểm tra 1 chuỗi có phải tên cột Excel hợp lệ (A-Z, AA-ZZ...) không. Trả về (ok, col_chuan_hoa)."""
            c = (col or "").strip().upper()
            if c == "":
                return True, ""  # để trống = không ghi trường đó, hợp lệ
            import re as _re
            if _re.fullmatch(r"[A-Z]{1,3}", c):
                return True, c
            return False, c

        def get_col8c_map(self):
            """Đọc cấu hình cột từ giao diện, kiểm tra hợp lệ. Trả về (col_map, loi). loi=None nếu OK."""
            col_map = {}
            seen = {}
            for key, var in self.var_col8c.items():
                ok, c = self._validate_excel_column(var.get())
                if not ok:
                    return None, f"Cột '{var.get()}' không hợp lệ (chỉ dùng chữ cái A-Z, VD: B, H, AA)."
                if c:
                    if c in seen:
                        return None, f"Cột '{c}' bị dùng cho cả '{seen[c]}' và '{key}'. Mỗi cột chỉ nên gán 1 trường."
                    seen[c] = key
                    col_map[key] = c
            if not col_map.get("maxa") or not col_map.get("soto") or not col_map.get("sothua"):
                return None, "Bắt buộc phải có cột cho Mã xã, Số tờ, Số thửa (đây là khóa đối chiếu)."
            return col_map, None

        def reset_col8c_config(self):
            for k, v in self._default_col8c.items():
                self.var_col8c[k].set(v)
            self.lbl_col8c_status.config(text="✓ Đã khôi phục mặc định.", fg="#2e7d32")

        def _col8c_config_path(self):
            return os.path.join(get_app_dir(), "data", self.COL8C_CONFIG_FILENAME)

        def save_col8c_config(self):
            col_map, err = self.get_col8c_map()
            if err:
                messagebox.showerror("Cấu hình cột không hợp lệ", err)
                return
            # Lưu toàn bộ 12 trường (kể cả trống) để tải lại đúng
            data = {k: self.var_col8c[k].get().strip().upper() for k in self.var_col8c}
            try:
                path = self._col8c_config_path()
                os.makedirs(os.path.dirname(path), exist_ok=True)
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                self.lbl_col8c_status.config(text=f"✓ Đã lưu cấu hình cột.", fg="#2e7d32")
                self.log(f"✓ Đã lưu cấu hình cột báo cáo vào {path}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không lưu được cấu hình cột: {e}")

        def load_col8c_config(self):
            path = self._col8c_config_path()
            if not os.path.isfile(path):
                messagebox.showinfo("Chưa có", "Chưa có cấu hình cột đã lưu. Hãy 'Lưu cấu hình cột' trước.")
                return
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                for k in self.var_col8c:
                    if k in data:
                        self.var_col8c[k].set(data[k])
                self.lbl_col8c_status.config(text="✓ Đã tải cấu hình cột đã lưu.", fg="#2e7d32")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không đọc được cấu hình cột: {e}")

        def pick_wb8c_template(self):
            p = filedialog.askopenfilename(title="Chọn file Excel MẪU", filetypes=[("Excel files", "*.xlsx *.xlsm")])
            if p:
                self.var_wb8c_template.set(p)

        def pick_wb8c_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục xuất Excel kết quả")
            if p:
                self.var_wb8c_output.set(p)

        def _refresh_wb8c_tree(self, log_rows):
            for item in self.tree_wb8c.get_children():
                self.tree_wb8c.delete(item)
            for r in log_rows:
                self._insert_row_colored(self.tree_wb8c, (
                    r["file"], r["maxa"], r["ten"], r["cccd"], r.get("diachitt", ""),
                    r.get("so_to_goc", r["to"]), r["to"], r["thua"], r["dt"],
                    r.get("loai_dat", ""), r.get("nguon_so_to", ""), r.get("diem_tin_cay", ""),
                    r["status"], r["note"]), r["status"])

        def retry_wb8c_failed(self):
            rows = getattr(self, "_wb8c_last_log_rows", None)
            if not rows:
                messagebox.showinfo("Chưa có kết quả", "Chưa có lần chạy nào trước đó để lấy danh sách file lỗi.")
                return
            failed_paths = [r["path"] for r in rows
                            if ("LOI" in r["status"].upper() or "CAN_KIEM_TRA" in r["status"].upper())
                            and r.get("path") and os.path.isfile(r["path"])]
            if not failed_paths:
                messagebox.showinfo("Không có file lỗi", "Không có file lỗi/cần kiểm tra nào để chạy lại "
                                    "(hoặc các file đó không còn tồn tại trên đĩa).")
                return
            self.wb8c_picker.selected_files = failed_paths
            self.log(f"🔁 Chạy lại {len(failed_paths)} file lỗi/cần kiểm tra từ lần chạy trước...")
            self.run_wb8c_clicked(dry_run=False)

        def run_wb8c_clicked(self, dry_run=True):
            template_path = self.var_wb8c_template.get().strip()
            source = self.wb8c_picker.get_source()
            output_dir = self.var_wb8c_output.get().strip()
            sheet_name = self.var_wb8c_sheet.get().strip() or None

            if not template_path or not os.path.isfile(template_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel MẪU hợp lệ.")
                return
            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF nguồn.")
                return
            if not dry_run and not output_dir:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục xuất Excel kết quả.")
                return
            if not HAS_PYPDF:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pypdf")
                return

            try:
                header_row = int(self.var_wb8c_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu ghi dữ liệu phải là số.")
                return

            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300
            clear_old = self.var_wb8c_clear_old.get()
            infer_cccd = self.var_wb8c_infer_cccd.get()
            no_overwrite = self.var_wb8c_no_overwrite.get()
            warn_dup = self.var_wb8c_warn_dup.get()

            # Đọc cấu hình cột động, kiểm tra hợp lệ trước khi chạy
            col_map, col_err = self.get_col8c_map()
            if col_err:
                messagebox.showerror("Cấu hình cột không hợp lệ", col_err)
                return

            if not dry_run and clear_old:
                used_cols_str = ", ".join(col_map.values())
                if not messagebox.askyesno("Xác nhận xóa dữ liệu cũ",
                                            f"Bạn đã tick 'Xóa dữ liệu cũ' — toàn bộ nội dung hiện có trong các cột "
                                            f"{used_cols_str} từ dòng bắt đầu sẽ bị XÓA trước khi ghi mới "
                                            "(file gốc không bị ảnh hưởng vì kết quả lưu ra file MỚI). Tiếp tục?"):
                    return

            self.txt_log.delete("1.0", "end")
            self._refresh_wb8c_tree([])

            checkpoint_task_id = None
            files_da_xong = None
            if not dry_run:
                from app.services import checkpoint_service as _ckpt
                checkpoint_task_id = _ckpt.tinh_task_id("dien_excel_mau_tu_pdf", str(source), output_dir)
                data_cu = _ckpt.doc_checkpoint(checkpoint_task_id)
                if data_cu and data_cu.get("trang_thai") == "DANG_XU_LY":
                    so_da_xong = len(data_cu.get("da_hoan_thanh", []))
                    so_tong = len(data_cu.get("danh_sach_file", []))
                    if messagebox.askyesno(
                        "Phát hiện tác vụ dở dang",
                        f"Phát hiện đợt điền Excel MẪU TRƯỚC ĐÓ chưa hoàn thành (đã đọc {so_da_xong}/{so_tong} "
                        f"file, có thể do phần mềm bị đóng đột ngột).\n\n"
                        f"Bấm CÓ để TIẾP TỤC từ chỗ dở dang.\nBấm KHÔNG để BẮT ĐẦU LẠI TỪ ĐẦU."):
                        files_da_xong = set(data_cu.get("da_hoan_thanh", []))
                    else:
                        _ckpt.xoa_checkpoint(checkpoint_task_id)

            self.btn_wb8c_dryrun.config(state="disabled")
            self.btn_wb8c_run.config(state="disabled")
            self._enable_pause_cancel("control_wb8c", self.btn_wb8c_pause, self.btn_wb8c_cancel)
            self.pb_wb8c.config(value=0, maximum=100)
            self._wb8c_start_time = time.time()

            from app.services import task_manager_service as _tm2
            self.task_info_wb8c = _tm2.TaskInfo(
                ten_chuc_nang="8c. Điền Excel MẪU từ PDF", thu_muc_dau_ra=output_dir,
                trang_thai=_tm2.TrangThaiTask.DANG_KHOI_TAO)

            def update_wb8c_progress(idx, total, fname):
                elapsed = time.time() - self._wb8c_start_time
                self.pb_wb8c.config(value=idx, maximum=total)
                self.lbl_wb8c_progress.config(text=f"Đang đọc {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_wb8c_progress(idx, total, fname))
                self.task_info_wb8c.cap_nhat_tien_do(tien_do_toan_bo=idx,
                                                     trang_thai=_tm2.TrangThaiTask.DANG_DOC_FILE)
                self.task_info_wb8c.tong_so_luong = total
                append_tab_log_row(8, "8c. Điền Excel MẪU từ PDF", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_DOC_FILE")

            def worker():
                try:
                    output_path = None
                    if not dry_run:
                        stamp = time.strftime("%Y%m%d_%H%M")
                        output_path = os.path.join(output_dir, f"BAO_CAO_TONG_HOP_TU_PDF_{stamp}.xlsx")

                    debug_text_dir = None
                    if self.var_wb8c_debug_text.get():
                        debug_text_dir = os.path.join(get_app_data_dir(), "Logs", "pdf_text_debug")

                    perf_log_path = None
                    if getattr(self.app_config, "word_hien_thi_thong_ke_hieu_nang", True):
                        perf_log_path = os.path.join(get_standard_logs_dir(), "LOG_HIEU_NANG_XU_LY.csv")
                    log_rows = fill_excel_template_from_pdfs(
                        template_path, source, sheet_name, header_row,
                        xa_mapping, thon_mapping, use_ocr, ocr_dpi,
                        clear_old, infer_cccd, output_path, dry_run, self.log, control=self.control_wb8c,
                        col_map=col_map, no_overwrite=no_overwrite, warn_duplicate=warn_dup,
                        debug_text_dir=debug_text_dir,
                        threshold_cao=getattr(self.app_config, "nguong_tin_cay_cao", 90),
                        threshold_thap=getattr(self.app_config, "nguong_tin_cay_thap", 70),
                        progress_cb=progress_cb,
                        toc_do_xu_ly=getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"),
                        so_file_moi_dot=getattr(self.app_config, "word_so_file_moi_dot", 100),
                        nghi_giua_dot_giay=getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0),
                        perf_log_path=perf_log_path,
                        checkpoint_task_id=checkpoint_task_id, files_da_xong=files_da_xong)

                    self._wb8c_last_log_rows = log_rows
                    for _r in log_rows:
                        self.task_info_wb8c.ghi_nhan_ket_qua(_r.get("status", ""))
                    self.task_info_wb8c.trang_thai = _tm2.TrangThaiTask.THANH_CONG
                    n_failed = sum(1 for r in log_rows if "LOI" in r["status"].upper() or "CAN_KIEM_TRA" in r["status"].upper())
                    self.after(0, lambda: self.btn_wb8c_retry.config(state="normal" if n_failed else "disabled"))
                    self.after(0, lambda: self._refresh_wb8c_tree(log_rows))
                    self.after(0, lambda: self.lbl_wb8c_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._wb8c_start_time)}s)"))
                    if debug_text_dir:
                        self.log(f"Đã xuất text debug từng file vào: {debug_text_dir}")

                    import time
                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    log_dir = output_dir or os.path.dirname(template_path)
                    csv_path = os.path.join(log_dir, f"LOG_TONG_HOP_PDF_VAO_EXCEL_{stamp}.csv")
                    write_wb8c_log_csv(csv_path, log_rows)

                    check_path = None
                    if self.var_wb8c_export_check.get():
                        check_path = os.path.join(log_dir, f"DANH_SACH_CAN_KIEM_TRA_{stamp}.csv")
                        n_check = write_wb8c_check_csv(check_path, log_rows)
                        self.log(f"Đã xuất danh sách cần kiểm tra ({n_check} file): {check_path}")

                    ok_count = sum(1 for r in log_rows if r["status"] in ("DA_GHI_EXCEL", "CHAY_THU"))
                    err_count = len(log_rows) - ok_count

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {ok_count} OK | {err_count} cần kiểm tra")
                    self.log(f"File log CSV: {csv_path}")
                    if output_path:
                        self.log(f"File Excel kết quả: {output_path}")

                    msg = f"OK: {ok_count}\nCần kiểm tra: {err_count}\n\nLog CSV: {csv_path}"
                    if check_path:
                        msg += f"\nDanh sách cần kiểm tra: {check_path}"
                    if output_path:
                        msg += f"\nExcel kết quả: {output_path}"
                    messagebox.showinfo("Chạy thử hoàn tất" if dry_run else "Tổng hợp hoàn tất", msg)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_wb8c_dryrun.config(state="normal")
                    self.btn_wb8c_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_wb8c_pause, self.btn_wb8c_cancel)

            threading.Thread(target=worker, daemon=True).start()


        def pick_tong_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel tổng đã chỉnh sửa",
                                           filetypes=[("Excel files", "*.xlsx *.xlsm *.xls")])
            if p:
                self.var_tong_excel.set(p)
                try:
                    wb = openpyxl.load_workbook(p, read_only=True)
                    if wb.sheetnames:
                        self.var_tong_sheet.set(wb.sheetnames[0])
                except Exception:
                    pass

        def pick_tong_report_template(self):
            p = filedialog.askopenfilename(title="Chọn file Excel báo cáo mẫu",
                                           filetypes=[("Excel files", "*.xlsx *.xlsm *.xls")])
            if p:
                self.var_tong_report_template.set(p)

        def pick_tong_pdf_zip(self):
            p = filedialog.askopenfilename(title="Chọn file nén PDF (.zip)",
                                           filetypes=[("Zip files", "*.zip")])
            if p:
                self.var_tong_pdf_zip.set(p)

        def pick_tong_pdf_rar(self):
            p = filedialog.askopenfilename(title="Chọn file nén PDF (.rar)",
                                           filetypes=[("Rar files", "*.rar")])
            if p:
                self.var_tong_pdf_zip.set(p)

        def run_tong_reconcile_clicked(self, dry_run=True):
            excel_tong = self.var_tong_excel.get().strip()
            sheet_tong = self.var_tong_sheet.get().strip()
            template_report = self.var_tong_report_template.get().strip()
            sheet_report = self.var_tong_report_sheet.get().strip() or "Chưa GCN"
            zip_path = self.var_tong_pdf_zip.get().strip()
            source = self.tong_pdf_picker.get_source()

            if not excel_tong or not os.path.isfile(excel_tong):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel tổng đã chỉnh sửa.")
                return
            if not sheet_tong:
                messagebox.showerror("Thiếu thông tin", "Vui lòng nhập tên sheet Excel tổng.")
                return
            if not template_report or not os.path.isfile(template_report):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel báo cáo mẫu.")
                return
            if not source and not zip_path:
                messagebox.showerror("Thiếu thông tin",
                                     "Vui lòng chọn thư mục PDF HOẶC file nén .zip cần đối chiếu.")
                return
            try:
                header_row_tong = int(self.var_tong_header_row.get().strip())
                header_row_report = int(self.var_tong_report_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Sai định dạng", "Dòng tiêu đề/Dòng bắt đầu ghi phải là số nguyên.")
                return
            try:
                dientich_dung_sai = float(self.var_tong_dientich_dung_sai.get().strip().replace(",", "."))
            except ValueError:
                dientich_dung_sai = 1.0

            col_map_tong = {k: v.get().strip() for k, v in self.var_col_tong.items()}
            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))

            self.txt_log.delete("1.0", "end")
            self.btn_tong_dryrun.config(state="disabled")
            self.btn_tong_run.config(state="disabled")
            self._enable_pause_cancel("control_tong", self.btn_tong_pause, self.btn_tong_cancel)
            self.pb_tong.config(value=0, maximum=100)
            self._tong_progress_start_time = time.time()

            def update_tong_progress(idx, total, fname):
                elapsed = time.time() - self._tong_progress_start_time
                self.pb_tong.config(value=idx, maximum=total)
                self.lbl_tong_progress.config(text=f"Đang đối chiếu {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_tong_progress(idx, total, fname))
                append_tab_log_row(8, "8d. Đối chiếu Excel tổng với PDF", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_DOI_CHIEU")

            def worker():
                _tong_start_str = time.strftime("%Y-%m-%d %H:%M:%S")
                try:
                    pdf_source = source
                    if zip_path:
                        self.log(f"Đang giải nén file nén PDF: {zip_path}")
                        extract_fn = extract_pdfs_from_rar if zip_path.lower().endswith(".rar") else extract_pdfs_from_zip
                        extract_dir, pdf_files = extract_fn(zip_path, log_cb=self.log)
                        if not pdf_files:
                            raise RuntimeError("Không giải nén được PDF nào từ file nén đã chọn.")
                        pdf_source = extract_dir

                    result = reconcile_excel_tong_with_pdfs(
                        excel_tong, sheet_tong, header_row_tong, col_map_tong,
                        pdf_source, xa_mapping, thon_mapping, True, 300, self.log, control=self.control_tong,
                        doc_so_sanh=self.var_tong_doc_so_sanh.get(), dientich_dung_sai=dientich_dung_sai,
                        progress_cb=progress_cb)
                    self.after(0, lambda: self.lbl_tong_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._tong_progress_start_time)}s)"))

                    self.after(0, lambda: self._refresh_tong_tree(result["so_sanh_rows"]))

                    if not dry_run:
                        col_map_report = {"maxa": "B", "ten": "H", "cccd": "I", "ngaysinh": "J", "gioitinh": "K",
                                          "diachitt": "L", "soto": "V", "sothua": "W", "diachithua": "X",
                                          "dientich": "Y", "mucdich": "Z", "dientichloai1": "AA"}
                        stamp = time.strftime("%Y%m%d_%H%M")
                        save_dir = self.var_wb8c_output.get().strip() or os.path.dirname(template_report)
                        output_path = os.path.join(
                            save_dir, f"BAO_CAO_TU_EXCEL_TONG_KHOP_PDF_CO_SO_SANH_{stamp}.xlsx")
                        report_log_rows, output_path = write_report_with_so_sanh(
                            template_report, sheet_report, header_row_report, col_map_report,
                            result, output_path, self.log, clear_old_data=self.var_tong_clear_old.get())

                        log_dir = save_dir
                        write_log_doi_chieu_excel_tong_va_pdf(
                            os.path.join(log_dir, f"LOG_DOI_CHIEU_EXCEL_TONG_VA_PDF_{stamp}.csv"), result)
                        write_log_so_sanh_pdf_va_excel_tong(
                            os.path.join(log_dir, f"LOG_SO_SANH_PDF_VA_EXCEL_TONG_{stamp}.csv"), result["so_sanh_rows"])
                        write_log_ghi_bao_cao_tu_excel_tong(
                            os.path.join(log_dir, f"LOG_GHI_BAO_CAO_TU_EXCEL_TONG_{stamp}.csv"), report_log_rows)
                        write_danh_sach_can_kiem_tra_tong(
                            os.path.join(log_dir, f"DANH_SACH_CAN_KIEM_TRA_{stamp}.csv"), result)
                        self.log(f"\nĐã xuất báo cáo: {output_path}")

                        archive_note = ""
                        if self.var_tong_luu_tru.get():
                            archive_path = archive_report(output_path, {
                                "excel_tong_nguon": excel_tong, "file_bao_cao_mau": template_report,
                                "thu_muc_ho_so_nguon": pdf_source if isinstance(pdf_source, str) else "(nhiều file)",
                                "sheet_excel_tong": sheet_tong, "sheet_bao_cao": sheet_report,
                                "so_ho_so_da_quet": len(result["matched"]) + len(result["pdf_khong_co_trong_excel"]),
                                "so_ho_so_khop": len(result["matched"]),
                                "so_ho_so_da_cap_nhat": len(result["matched"]),
                                "so_ho_so_can_kiem_tra": len(result["pdf_khong_co_trong_excel"]) + len(result["trung_khoa"]),
                                "thoi_gian_bat_dau": _tong_start_str, "trang_thai": "DA_LUU_TRU",
                                "ghi_chu": "",
                            }, log_cb=self.log)
                            archive_note = f"\nĐã lưu trữ: {archive_path}"

                        messagebox.showinfo("Hoàn tất",
                                            f"Đã khớp {len(result['matched'])} hồ sơ.\n\nFile báo cáo:\n{output_path}"
                                            + archive_note)
                    else:
                        messagebox.showinfo("Chạy thử hoàn tất",
                                            f"Khớp: {len(result['matched'])}\n"
                                            f"PDF không có trong Excel tổng: {len(result['pdf_khong_co_trong_excel'])}\n"
                                            f"Excel tổng không có PDF: {len(result['excel_khong_co_pdf'])}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_tong_dryrun.config(state="normal")
                    self.btn_tong_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_tong_pause, self.btn_tong_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def _refresh_tong_tree(self, so_sanh_rows):
            for item in self.tree_tong.get_children():
                self.tree_tong.delete(item)
            for sr in so_sanh_rows:
                self._insert_row_colored(self.tree_tong, (
                    sr["file"], sr["truong"], sr["gia_tri_pdf"], sr["gia_tri_excel_tong"],
                    sr["ket_qua"], sr["co_chinh_sua"]), sr["ket_qua"])

        def open_report_history_window(self):
            win = tk.Toplevel(self)
            win.title("Lịch sử báo cáo đã lưu trữ")
            win.geometry("1100x520")

            row_search = tk.Frame(win); row_search.pack(fill="x", padx=8, pady=(8, 4))
            tk.Label(row_search, text="Tìm theo tên/ngày:").pack(side="left")
            var_search = tk.StringVar()
            tk.Entry(row_search, textvariable=var_search, width=30).pack(side="left", padx=(4, 6))
            tk.Button(row_search, text="Tìm", command=lambda: refresh_list(var_search.get().strip())).pack(side="left")
            tk.Button(row_search, text="Làm mới danh sách", command=lambda: refresh_list("")).pack(side="left", padx=(6, 0))

            cols_hist = ("ma", "ten", "so_khop", "so_can_kt", "hoan_thanh", "trang_thai")
            headers_hist = ["Mã báo cáo", "Tên file báo cáo", "Số hồ sơ khớp", "Cần kiểm tra",
                           "Thời gian hoàn thành", "Trạng thái"]
            widths_hist = [130, 320, 100, 90, 160, 120]
            tree_hist = ttk.Treeview(win, columns=cols_hist, show="headings", height=18)
            for c, h, w in zip(cols_hist, headers_hist, widths_hist):
                tree_hist.heading(c, text=h)
                tree_hist.column(c, width=w, anchor="w")
            vsb_hist = tk.Scrollbar(win, orient="vertical", command=tree_hist.yview)
            tree_hist.configure(yscrollcommand=vsb_hist.set)
            tree_hist.pack(side="left", fill="both", expand=True, padx=(8, 0), pady=(0, 8))
            vsb_hist.pack(side="left", fill="y", pady=(0, 8), padx=(0, 8))

            history_map = {}

            def refresh_list(keyword):
                for item in tree_hist.get_children():
                    tree_hist.delete(item)
                history_map.clear()
                history = read_report_history()
                kw = keyword.lower()
                for row in reversed(history):  # moi nhat len dau
                    if kw and kw not in row.get("TenBaoCao", "").lower() and kw not in row.get("ThoiGianHoanThanh", "").lower():
                        continue
                    item_id = tree_hist.insert("", "end", values=(
                        row.get("MaBaoCao", ""), row.get("TenBaoCao", ""), row.get("SoHoSoKhop", ""),
                        row.get("SoHoSoCanKiemTra", ""), row.get("ThoiGianHoanThanh", ""), row.get("TrangThai", "")))
                    history_map[item_id] = row

            def get_selected_path():
                sel = tree_hist.selection()
                if not sel:
                    messagebox.showinfo("Chưa chọn", "Vui lòng chọn 1 báo cáo trong danh sách trước.")
                    return None
                row = history_map.get(sel[0])
                if not row:
                    return None
                return row.get("DuongDanFile", "")

            def action_open_file():
                path = get_selected_path()
                if not path:
                    return
                if not os.path.isfile(path):
                    messagebox.showerror("Không tìm thấy file", f"File không còn tồn tại:\n{path}")
                    return
                try:
                    if sys.platform == "win32":
                        os.startfile(path)
                    elif sys.platform == "darwin":
                        subprocess.Popen(["open", path])
                    else:
                        subprocess.Popen(["xdg-open", path])
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Không mở được file: {e}")

            def action_open_folder():
                path = get_selected_path()
                if not path:
                    return
                folder = os.path.dirname(path)
                if not os.path.isdir(folder):
                    messagebox.showerror("Không tìm thấy thư mục", f"Thư mục không còn tồn tại:\n{folder}")
                    return
                try:
                    if sys.platform == "win32":
                        os.startfile(folder)
                    elif sys.platform == "darwin":
                        subprocess.Popen(["open", folder])
                    else:
                        subprocess.Popen(["xdg-open", folder])
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Không mở được thư mục: {e}")

            def action_copy_path():
                path = get_selected_path()
                if not path:
                    return
                win.clipboard_clear()
                win.clipboard_append(path)
                messagebox.showinfo("Đã sao chép", "Đã sao chép đường dẫn file vào clipboard.")

            row_btns = tk.Frame(win); row_btns.pack(side="right", fill="y", padx=(0, 8), pady=8)
            tk.Button(row_btns, text="Mở file", width=20, command=action_open_file).pack(pady=3)
            tk.Button(row_btns, text="Mở thư mục chứa", width=20, command=action_open_folder).pack(pady=3)
            tk.Button(row_btns, text="Sao chép đường dẫn", width=20, command=action_copy_path).pack(pady=3)

            refresh_list("")
try:
    from app.ui.tab_word import WordTabMixin
    _UI_WORD_MODULE_OK = True
except Exception:
    _UI_WORD_MODULE_OK = False

    class WordTabMixin:
        """Bản sao lưu trữ tại chỗ - dùng khi không import được app/ui/tab_word.py (VD thiếu
        thư mục app/ khi đóng gói sai). Nội dung GIỐNG HỆT bản đã tách, để phần mềm không bao giờ vỡ."""

        def _build_tab_word(self, parent, pad):
            frm_word = tk.LabelFrame(parent, labelwidget=tk.Label(
                parent, text="  7a. Chuyển đổi Word/PDF hàng loạt  ", font=("Segoe UI", 10, "bold"),
                fg="#1a237e", bg="#c5cae9", relief="raised", bd=1))
            frm_word.pack(fill="x", **pad)
            tk.Label(frm_word,
                     text="Chuyển PDF sang .docx. Chế độ 'Tự nhận diện' (khuyến nghị): PDF có lớp chữ dùng bố cục "
                          "chuẩn (giữ bảng biểu thật); PDF scan tự OCR ra văn bản CÓ THỂ CHỈNH SỬA (không chỉ chèn ảnh).",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(4, 6))

            row_in = tk.Frame(frm_word)
            row_in.pack(fill="x", padx=6, pady=3)
            tk.Label(row_in, text="Nguồn PDF:", width=28, anchor="w").pack(side="left")
            self.word_input_picker = SourcePicker(
                row_in, "Nguồn PDF", [("PDF files", "*.pdf")], file_label="PDF")
            self.word_input_picker.pack(side="left", fill="x", expand=True)

            row_out = tk.Frame(frm_word)
            row_out.pack(fill="x", padx=6, pady=3)
            tk.Label(row_out, text="Thư mục lưu file Word:", width=28, anchor="w").pack(side="left")
            tk.Entry(row_out, textvariable=self.var_word_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_out, text="Chọn...", command=self.pick_word_output).pack(side="left")

            row_mode = tk.Frame(frm_word)
            row_mode.pack(fill="x", padx=6, pady=3)
            tk.Label(row_mode, text="Chế độ chuyển đổi PDF→Word:", width=28, anchor="w").pack(side="left")
            self.var_word_mode = tk.StringVar(value="Tự nhận diện")
            ttk.Combobox(row_mode, textvariable=self.var_word_mode, state="readonly", width=16,
                        values=["Tự nhận diện", "Nhanh", "Giữ bố cục", "OCR bản quét"]).pack(side="left")
            tk.Label(row_mode, text="   Hậu tố tên file:").pack(side="left", padx=(12, 4))
            self.var_word_suffix = tk.StringVar(value="")
            tk.Entry(row_mode, textvariable=self.var_word_suffix, width=14).pack(side="left")
            self.var_word_no_overwrite = tk.BooleanVar(value=True)
            tk.Checkbutton(row_mode, text="Không ghi đè (tự thêm _1, _2...)",
                           variable=self.var_word_no_overwrite).pack(side="left", padx=(12, 0))

            row_opt = tk.Frame(frm_word)
            row_opt.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_opt, text="Gộp tất cả vào 1 file .docx duy nhất (ngoài các file riêng lẻ)",
                           variable=self.var_word_merge).pack(side="left")
            tk.Label(row_opt, text="   Tên file gộp:").pack(side="left")
            tk.Entry(row_opt, textvariable=self.var_word_merged_name, width=20).pack(side="left")

            row_run = tk.Frame(frm_word)
            row_run.pack(fill="x", padx=6, pady=(4, 3))
            tk.Button(row_run, text="📄 XUẤT WORD HÀNG LOẠT", font=("Arial", 11, "bold"),
                      bg="#283593", fg="white", command=self.run_word_export_clicked).pack(side="left")
            self.btn_word_export_pause, self.btn_word_export_cancel = self._add_pause_cancel(row_run, "control_word_export")
            self.btn_word_export_pause.pack(side="left", padx=(8, 3))
            self.btn_word_export_cancel.pack(side="left", padx=3)

            row_progress_word = tk.Frame(frm_word)
            row_progress_word.pack(fill="x", padx=6, pady=(0, 6))
            self.pb_word_export = ttk.Progressbar(row_progress_word, orient="horizontal", mode="determinate", length=400)
            self.pb_word_export.pack(side="left", padx=(0, 10))
            self.lbl_word_export_progress = tk.Label(row_progress_word, text="", fg="#555", anchor="w")
            self.lbl_word_export_progress.pack(side="left", fill="x", expand=True)

            # ================= 7b: ĐÃ GỘP VÀO 7c (CÁCH 2) — không hiển thị riêng nữa =================
            # (Nguồn dữ liệu Excel + cấu hình cột trước đây ở mục 7b độc lập, nay dồn vào phần
            # "CÁCH 2" của 7c bên dưới để tùy biến ngay tại nơi dùng, không cần chuyển qua lại 2 mục.
            # Các biến self.var_wx_* vẫn giữ nguyên và được dùng trực tiếp trong 7c.)

            frm_mm = styled_labelframe(parent, text="7c. Xuất Word THEO MẪU CÓ SẴN (mail-merge, giữ nguyên bố cục gốc)")
            frm_mm.pack(fill="x", **pad)
            tk.Label(frm_mm,
                     text="Dùng đúng file mẫu Word bạn cung cấp (có sẵn 2 mẫu: Đơn đăng ký / Thông báo xác nhận), "
                          "tự thay các token name/id/location/diachithua/numberthua/numberto/Sdientich/mucdich/"
                          "core/year bằng dữ liệu thực (location=địa chỉ ở cột 'Địa chỉ', diachithua=địa chỉ ở "
                          "cột 'Địa chỉ thửa đất' NẾU có cấu hình - dùng khi file Excel có 2 cột địa chỉ khác "
                          "nhau: thường trú và thửa đất), GIỮ NGUYÊN toàn bộ bố cục/định dạng của file mẫu gốc.",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_m1 = tk.Frame(frm_mm)
            row_m1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_m1, text="File mẫu (.docx):", width=18, anchor="w").pack(side="left")
            tk.Entry(row_m1, textvariable=self.var_mailmerge_template).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_m1, text="Chọn...", command=self.pick_mailmerge_template).pack(side="left")

            row_m2 = tk.Frame(frm_mm)
            row_m2.pack(fill="x", padx=6, pady=(0, 6))
            tk.Label(row_m2, text="Mẫu có sẵn:", width=18, anchor="w").pack(side="left")
            tk.Button(row_m2, text="Đơn đăng ký đất đai",
                      command=lambda: self.use_bundled_template("mau_don_dang_ky.docx")).pack(side="left", padx=3)
            tk.Button(row_m2, text="Thông báo xác nhận",
                      command=lambda: self.use_bundled_template("mau_thong_bao_xac_nhan.docx")).pack(side="left", padx=3)

            row_m3 = tk.Frame(frm_mm)
            row_m3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_m3, text="Mẫu đặt tên file:", width=18, anchor="w").pack(side="left")
            tk.Entry(row_m3, textvariable=self.var_mailmerge_filename_template).pack(side="left", fill="x", expand=True, padx=4)

            row_m4 = tk.Frame(frm_mm)
            row_m4.pack(fill="x", padx=6, pady=3)
            tk.Label(row_m4, text="Thư mục lưu kết quả:", width=18, anchor="w").pack(side="left")
            tk.Entry(row_m4, textvariable=self.var_mailmerge_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_m4, text="Chọn...", command=self.pick_mailmerge_output).pack(side="left")

            tk.Label(frm_mm, text="Nguồn dữ liệu — CÁCH 1: đọc trực tiếp từ file/thư mục PDF",
                     font=("Arial", 9, "bold"), anchor="w").pack(anchor="w", padx=6, pady=(8, 0))
            self.mailmerge_pdf_picker = SourcePicker(
                frm_mm, "Chọn 1 thư mục PDF, HOẶC chọn 1/nhiều file PDF bất kỳ",
                filetypes=[("PDF files", "*.pdf")], file_label="file PDF")
            self.mailmerge_pdf_picker.pack(fill="x", padx=6, pady=3)
            row_m_run1 = tk.Frame(frm_mm)
            row_m_run1.pack(fill="x", padx=6, pady=(0, 8))
            self.btn_mm_pdf = tk.Button(row_m_run1, text="🔀 XUẤT THEO MẪU (TỪ PDF)", font=("Arial", 11, "bold"),
                                         bg="#004d40", fg="white", command=self.run_mailmerge_from_pdf_clicked)
            self.btn_mm_pdf.pack(side="left")
            self.btn_mm_pdf_pause, self.btn_mm_pdf_cancel = self._add_pause_cancel(row_m_run1, "control_mm_pdf")
            self.btn_mm_pdf_pause.pack(side="left", padx=(8, 3))
            self.btn_mm_pdf_cancel.pack(side="left", padx=3)
            tk.Label(row_m_run1, text="Dùng chung cấu hình OCR + bảng tra Thôn/Xã ở Tab 5.",
                     fg="#555").pack(side="left", padx=10)

            row_mm_pdf_progress = tk.Frame(frm_mm); row_mm_pdf_progress.pack(fill="x", padx=6, pady=(0, 6))
            self.pb_mm_pdf = ttk.Progressbar(row_mm_pdf_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_mm_pdf.pack(side="left", padx=(0, 10))
            self.lbl_mm_pdf_progress = tk.Label(row_mm_pdf_progress, text="", fg="#555", anchor="w")
            self.lbl_mm_pdf_progress.pack(side="left", fill="x", expand=True)

            tk.Label(frm_mm, text="Nguồn dữ liệu — CÁCH 2: lấy từ file Excel bất kỳ (tự chọn cột cho từng trường)",
                     font=("Arial", 9, "bold"), anchor="w").pack(anchor="w", padx=6, pady=(4, 0))
            row_m_x1 = tk.Frame(frm_mm)
            row_m_x1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_m_x1, text="File Excel:", width=18, anchor="w").pack(side="left")
            tk.Entry(row_m_x1, textvariable=self.var_wx_excel).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_m_x1, text="Chọn...", command=self.pick_wx_excel).pack(side="left")

            row_m_x2 = tk.Frame(frm_mm)
            row_m_x2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_m_x2, text="Sheet (trống=đầu tiên):", width=18, anchor="w").pack(side="left")
            tk.Entry(row_m_x2, textvariable=self.var_wx_sheet, width=15).pack(side="left", padx=(0, 14))
            tk.Label(row_m_x2, text="Dòng bắt đầu dữ liệu:").pack(side="left")
            tk.Entry(row_m_x2, textvariable=self.var_wx_header_row, width=6).pack(side="left")

            row_m_x3 = tk.Frame(frm_mm)
            row_m_x3.pack(fill="x", padx=6, pady=3)
            for label, var in [("STT:", self.var_wx_col_stt), ("Mã xã:", self.var_wx_col_maxa),
                                ("Tờ:", self.var_wx_col_to), ("Thửa:", self.var_wx_col_thua)]:
                tk.Label(row_m_x3, text=label, anchor="w").pack(side="left")
                tk.Entry(row_m_x3, textvariable=var, width=6).pack(side="left", padx=(0, 10))

            row_m_x4 = tk.Frame(frm_mm)
            row_m_x4.pack(fill="x", padx=6, pady=3)
            for label, var in [("Tên chủ:", self.var_wx_col_ten), ("Địa chỉ:", self.var_wx_col_diachi),
                                ("Diện tích:", self.var_wx_col_dt), ("CCCD:", self.var_wx_col_id)]:
                tk.Label(row_m_x4, text=label, anchor="w").pack(side="left")
                tk.Entry(row_m_x4, textvariable=var, width=6).pack(side="left", padx=(0, 10))

            row_m_x4b = tk.Frame(frm_mm)
            row_m_x4b.pack(fill="x", padx=6, pady=3)
            for label, var in [("Mục đích SD:", self.var_wx_col_mucdich), ("Nguồn gốc SD:", self.var_wx_col_core),
                                ("Thời hạn SD:", self.var_wx_col_year)]:
                tk.Label(row_m_x4b, text=label, anchor="w").pack(side="left")
                tk.Entry(row_m_x4b, textvariable=var, width=6).pack(side="left", padx=(0, 10))

            row_m_run2 = tk.Frame(frm_mm)
            row_m_run2.pack(fill="x", padx=6, pady=(3, 8))
            self.btn_mm_excel = tk.Button(row_m_run2, text="🔀 XUẤT THEO MẪU (TỪ EXCEL)", font=("Arial", 11, "bold"),
                                           bg="#33691e", fg="white", command=self.run_mailmerge_from_excel_clicked)
            self.btn_mm_excel.pack(side="left")
            self.btn_mm_excel_pause, self.btn_mm_excel_cancel = self._add_pause_cancel(row_m_run2, "control_mm_excel")
            self.btn_mm_excel_pause.pack(side="left", padx=(8, 3))
            self.btn_mm_excel_cancel.pack(side="left", padx=3)

            row_mm_excel_progress = tk.Frame(frm_mm); row_mm_excel_progress.pack(fill="x", padx=6, pady=(0, 8))
            self.pb_mm_excel = ttk.Progressbar(row_mm_excel_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_mm_excel.pack(side="left", padx=(0, 10))
            self.lbl_mm_excel_progress = tk.Label(row_mm_excel_progress, text="", fg="#555", anchor="w")
            self.lbl_mm_excel_progress.pack(side="left", fill="x", expand=True)

            # ================= 7d: Xuất Word hàng loạt (Excel tổng + mẫu GT/TBXN, đặc tả đầy đủ) =================
            frm_wb = styled_labelframe(
                parent, text="7d. Xuất Word hàng loạt — Excel tổng + mẫu Đơn đăng ký / Thông báo xác nhận (khuyến nghị)")
            frm_wb.pack(fill="both", expand=True, **pad)
            tk.Label(frm_wb,
                     text="Mỗi dòng Excel → 1 (hoặc nhiều) file Word riêng, giữ nguyên định dạng/bố cục mẫu gốc. "
                          "Tên file: CHUACOGIAY_[Mã xã]_[Số tờ]_[Số thửa]_GT.docx / _TBXN.docx.\n"
                          "Nếu máy có Microsoft Word (qua pywin32) sẽ dùng Word thật để giữ định dạng chính xác nhất; "
                          "nếu không, tự dùng phương án dự phòng (chỉ áp dụng cho mẫu .docx).",
                     fg="#555", justify="left", wraplength=880).pack(anchor="w", padx=6, pady=(6, 6))

            row_wb1 = tk.Frame(frm_wb)
            row_wb1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb1, text="① File Excel tổng:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_wb1, textvariable=self.var_wb_excel).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_wb1, text="Chọn...", command=self.pick_wb_excel).pack(side="left")

            row_wb2 = tk.Frame(frm_wb)
            row_wb2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb2, text="② Mẫu Đơn đăng ký:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_wb2, textvariable=self.var_wb_template_gt).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_wb2, text="Chọn...", command=self.pick_wb_template_gt).pack(side="left")
            tk.Button(row_wb2, text="Dùng mẫu có sẵn",
                      command=lambda: self.use_bundled_template_for(self.var_wb_template_gt, "mau_don_dang_ky.docx")
                      ).pack(side="left", padx=4)

            row_wb3 = tk.Frame(frm_wb)
            row_wb3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb3, text="③ Mẫu Thông báo xác nhận:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_wb3, textvariable=self.var_wb_template_tbxn).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_wb3, text="Chọn...", command=self.pick_wb_template_tbxn).pack(side="left")
            tk.Button(row_wb3, text="Dùng mẫu có sẵn",
                      command=lambda: self.use_bundled_template_for(self.var_wb_template_tbxn, "mau_thong_bao_xac_nhan.docx")
                      ).pack(side="left", padx=4)

            row_wb4 = tk.Frame(frm_wb)
            row_wb4.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb4, text="④ Thư mục xuất file:", width=20, anchor="w").pack(side="left")
            tk.Entry(row_wb4, textvariable=self.var_wb_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_wb4, text="Chọn...", command=self.pick_wb_output).pack(side="left")

            row_wb5 = tk.Frame(frm_wb)
            row_wb5.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb5, text="⑤ Tên sheet (trống=đầu tiên):").pack(side="left")
            tk.Entry(row_wb5, textvariable=self.var_wb_sheet, width=15).pack(side="left", padx=(4, 20))
            tk.Label(row_wb5, text="⑥ Dòng bắt đầu dữ liệu:").pack(side="left")
            tk.Entry(row_wb5, textvariable=self.var_wb_header_row, width=6).pack(side="left", padx=4)

            row_wb_cols1 = tk.Frame(frm_wb)
            row_wb_cols1.pack(fill="x", padx=6, pady=(8, 3))
            tk.Label(row_wb_cols1, text="Cột dữ liệu:", font=("Arial", 9, "bold")).pack(side="left")
            for label, var in [("Mã xã:", self.var_wb_col_maxa), ("Tên chủ:", self.var_wb_col_name),
                                ("CCCD:", self.var_wb_col_id), ("Địa chỉ:", self.var_wb_col_location),
                                ("Địa chỉ thửa đất (tùy chọn):", self.var_wb_col_diachithua)]:
                tk.Label(row_wb_cols1, text=label).pack(side="left", padx=(10, 2))
                tk.Entry(row_wb_cols1, textvariable=var, width=5).pack(side="left")

            row_wb_cols2 = tk.Frame(frm_wb)
            row_wb_cols2.pack(fill="x", padx=6, pady=3)
            for label, var in [("Tờ:", self.var_wb_col_to), ("Thửa:", self.var_wb_col_thua),
                                ("Diện tích:", self.var_wb_col_dt), ("Mục đích:", self.var_wb_col_mucdich),
                                ("Nguồn gốc:", self.var_wb_col_core), ("Thời hạn:", self.var_wb_col_year)]:
                tk.Label(row_wb_cols2, text=label).pack(side="left", padx=(10, 2))
                tk.Entry(row_wb_cols2, textvariable=var, width=5).pack(side="left")

            row_wb_cols3 = tk.Frame(frm_wb)
            row_wb_cols3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_wb_cols3, text="Cấu hình cột số văn bản (tùy chọn):", font=("Arial", 9, "bold")).pack(side="left")
            for label, var in [("Số TB:", self.var_wb_col_numbertb), ("Số BB:", self.var_wb_col_numberbb),
                                ("Số CV:", self.var_wb_col_numbercv), ("Số KH:", self.var_wb_col_numberkh),
                                ("Số VB chung:", self.var_wb_col_sovanban), ("Ký hiệu:", self.var_wb_col_kyhieu),
                                ("Ngày VB:", self.var_wb_col_ngayvanban)]:
                tk.Label(row_wb_cols3, text=label).pack(side="left", padx=(10, 2))
                tk.Entry(row_wb_cols3, textvariable=var, width=5).pack(side="left")

            tk.Label(row_wb_cols1,
                     text="(Nếu Excel chưa có cột Mã xã, để trống ô 'Mã xã' — phần mềm sẽ tự suy theo tên thôn/xã "
                          "trong cột Địa chỉ)", fg="#777").pack(side="left", padx=10)

            row_wb_opt = tk.Frame(frm_wb)
            row_wb_opt.pack(fill="x", padx=6, pady=(8, 3))
            tk.Label(row_wb_opt, text="⑦ Tùy chọn:", font=("Arial", 9, "bold")).pack(side="left")
            tk.Checkbutton(row_wb_opt, text="Xuất mẫu Đơn đăng ký (GT)", variable=self.var_wb_export_gt).pack(side="left", padx=6)
            tk.Checkbutton(row_wb_opt, text="Xuất mẫu Thông báo xác nhận (TBXN)",
                           variable=self.var_wb_export_tbxn).pack(side="left", padx=6)
            tk.Checkbutton(row_wb_opt, text="Gộp 2 mẫu vào chung 1 file/dòng",
                           variable=self.var_wb_combine).pack(side="left", padx=6)

            # ===== Cấu hình hậu tố tên file (tùy chọn, không ép cứng GT/TBXN) =====
            frm_suffix = styled_labelframe(parent, text="🏷 Cấu hình hậu tố tên file")
            frm_suffix.pack(fill="x", **pad)
            self._register_advanced_widget(frm_suffix, fill="x", **pad)
            tk.Checkbutton(frm_suffix, text="Thêm hậu tố tùy chỉnh vào tên file (tắt = dùng mặc định GT/TBXN như cũ)",
                           variable=self.var_suffix_enabled, font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 4))

            row_sf1 = tk.Frame(frm_suffix)
            row_sf1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_sf1, text="Chọn nhanh:", width=14, anchor="w").pack(side="left")
            ttk.Combobox(row_sf1, textvariable=self.var_suffix_combobox, values=SUFFIX_PRESETS,
                        width=14, state="readonly").pack(side="left", padx=(0, 20))
            tk.Label(row_sf1, text="Hậu tố tự nhập:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_sf1, textvariable=self.var_suffix_manual, width=20).pack(side="left")

            row_sf2 = tk.Frame(frm_suffix)
            row_sf2.pack(fill="x", padx=6, pady=3)
            tk.Checkbutton(row_sf2, text="Ưu tiên hậu tố nhập tay",
                           variable=self.var_suffix_prioritize_manual).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_sf2, text="Lấy hậu tố từ cột Excel:",
                           variable=self.var_suffix_use_excel).pack(side="left")
            tk.Entry(row_sf2, textvariable=self.var_wb_col_hauto, width=6).pack(side="left", padx=(4, 16))
            tk.Checkbutton(row_sf2, text="Viết hoa hậu tố", variable=self.var_suffix_uppercase).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_sf2, text="Bỏ dấu tiếng Việt", variable=self.var_suffix_remove_diacritics).pack(side="left")

            row_sf3 = tk.Frame(frm_suffix)
            row_sf3.pack(fill="x", padx=6, pady=3)
            tk.Label(row_sf3, text="Mẫu tên file:", width=14, anchor="w").pack(side="left")
            tk.Entry(row_sf3, textvariable=self.var_suffix_filename_template, width=55).pack(side="left", padx=(0, 6))
            tk.Label(row_sf3, text="Biến: {maxa}{soto}{sothua}{hoten}{cccd}{thon}{xacu}{hautofile}{ngay}{stt}",
                    fg="#777").pack(side="left")

            row_sf4 = tk.Frame(frm_suffix)
            row_sf4.pack(fill="x", padx=6, pady=(3, 8))
            tk.Button(row_sf4, text="👁 Xem trước tên file", command=self.preview_suffix_filename).pack(side="left", padx=(0, 10))
            tk.Label(row_sf4, textvariable=self.var_suffix_preview, font=("Consolas", 9, "bold"),
                    fg="#1a237e").pack(side="left")

            row_wb_run = tk.Frame(frm_wb)
            row_wb_run.pack(fill="x", padx=6, pady=(8, 8))
            self.btn_wb_dryrun = tk.Button(row_wb_run, text="⑧ CHẠY THỬ (chỉ kiểm tra + log CSV, KHÔNG xuất file)",
                                            font=("Arial", 10, "bold"), bg="#455a64", fg="white",
                                            command=lambda: self.run_word_batch_clicked(dry_run=True))
            self.btn_wb_dryrun.pack(side="left", padx=(0, 6))
            self.btn_wb_run = tk.Button(row_wb_run, text="⑨ XUẤT WORD HÀNG LOẠT",
                                         font=("Arial", 10, "bold"), bg="#2e7d32", fg="white",
                                         command=lambda: self.run_word_batch_clicked(dry_run=False))
            self.btn_wb_run.pack(side="left")
            self.btn_wb_pause, self.btn_wb_cancel = self._add_pause_cancel(row_wb_run, "control_wb")
            self.btn_wb_pause.pack(side="left", padx=(12, 3))
            self.btn_wb_cancel.pack(side="left", padx=3)

            row_wb_progress = tk.Frame(frm_wb); row_wb_progress.pack(fill="x", padx=6, pady=(0, 3))
            self.pb_wb = ttk.Progressbar(row_wb_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_wb.pack(side="left", padx=(0, 10))
            self.lbl_wb_progress = tk.Label(row_wb_progress, text="", fg="#555", anchor="w")
            self.lbl_wb_progress.pack(side="left", fill="x", expand=True)

            row_wb_perf = tk.Frame(frm_wb); row_wb_perf.pack(fill="x", padx=6, pady=(0, 8))
            self.lbl_wb_perf = tk.Label(row_wb_perf, text="", fg="#1565c0", anchor="w",
                                        font=("Consolas", 9))
            self.lbl_wb_perf.pack(side="left", fill="x", expand=True)

            tk.Label(frm_wb, text="⑩ Bảng trạng thái:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 2))
            cols = ("stt", "ten", "maxa", "to", "thua", "filename", "status", "note")
            headers = ["STT", "Tên chủ", "Mã xã", "Số tờ", "Số thửa", "Tên file xuất", "Trạng thái", "Ghi chú"]
            widths = [40, 130, 60, 50, 55, 220, 130, 220]
            self.tree_wb = ttk.Treeview(frm_wb, columns=cols, show="headings", height=10)
            for c, h, w in zip(cols, headers, widths):
                self.tree_wb.heading(c, text=h)
                self.tree_wb.column(c, width=w, anchor="w")
            vsb = tk.Scrollbar(frm_wb, orient="vertical", command=self.tree_wb.yview)
            self.tree_wb.configure(yscrollcommand=vsb.set)
            self.tree_wb.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

            # ================= 7e: Chuyển Word sang PDF hàng loạt =================
            frm_w2p = styled_labelframe(parent, text="7e. Chuyển Word sang PDF hàng loạt")
            frm_w2p.pack(fill="both", expand=True, **pad)
            tk.Label(frm_w2p,
                     text="Nhận file .doc, .docx, .docm — mỗi file xuất ra 1 file PDF cùng tên (VD "
                          "CHUACOGIAY_02140_29_199.docx → CHUACOGIAY_02140_29_199.pdf), giữ nguyên định dạng, "
                          "font chữ, căn lề, hình ảnh, chữ ký, con dấu. KHÔNG làm thay đổi file Word gốc.\n"
                          "Ưu tiên dùng Microsoft Word (qua pywin32) nếu máy có; nếu không, tự dùng LibreOffice "
                          "(nếu có cài) làm phương án dự phòng.",
                     fg="#555", justify="left", wraplength=1000).pack(anchor="w", padx=6, pady=(6, 6))

            row_w1 = tk.Frame(frm_w2p)
            row_w1.pack(fill="x", padx=6, pady=3)
            tk.Label(row_w1, text="① Thư mục chứa file Word:", width=24, anchor="w").pack(side="left")
            tk.Entry(row_w1, textvariable=self.var_w2p_input).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_w1, text="Chọn...", command=self.pick_w2p_input).pack(side="left")

            row_w2 = tk.Frame(frm_w2p)
            row_w2.pack(fill="x", padx=6, pady=3)
            tk.Label(row_w2, text="② Thư mục xuất file PDF:", width=24, anchor="w").pack(side="left")
            tk.Entry(row_w2, textvariable=self.var_w2p_output).pack(side="left", fill="x", expand=True, padx=4)
            tk.Button(row_w2, text="Chọn...", command=self.pick_w2p_output).pack(side="left")

            row_w3 = tk.Frame(frm_w2p)
            row_w3.pack(fill="x", padx=6, pady=(6, 3))
            tk.Checkbutton(row_w3, text="③ Quét cả thư mục con", variable=self.var_w2p_subfolders).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_w3, text="④ Ghi đè file PDF nếu đã tồn tại", variable=self.var_w2p_overwrite).pack(side="left", padx=(0, 16))
            tk.Checkbutton(row_w3, text="⑤ Giữ nguyên cấu trúc thư mục con khi xuất", variable=self.var_w2p_keepstruct).pack(side="left")

            row_w_run = tk.Frame(frm_w2p)
            row_w_run.pack(fill="x", padx=6, pady=(8, 4))
            self.btn_w2p_dryrun = tk.Button(row_w_run, text="⑥ CHẠY THỬ (chỉ kiểm tra, KHÔNG xuất PDF)",
                                             font=("Arial", 10, "bold"), bg="#455a64", fg="white",
                                             command=lambda: self.run_w2p_clicked(dry_run=True))
            self.btn_w2p_dryrun.pack(side="left", padx=(0, 6))
            self.btn_w2p_run = tk.Button(row_w_run, text="⑦ CHUYỂN SANG PDF",
                                          font=("Arial", 10, "bold"), bg="#2e7d32", fg="white",
                                          command=lambda: self.run_w2p_clicked(dry_run=False))
            self.btn_w2p_run.pack(side="left", padx=(0, 6))
            self.btn_w2p_pause = tk.Button(row_w_run, text="⏸ Tạm dừng", state="disabled",
                                            command=self.toggle_pause_w2p)
            self.btn_w2p_pause.pack(side="left", padx=(12, 6))
            self.btn_w2p_cancel = tk.Button(row_w_run, text="⏹ Hủy", state="disabled",
                                             command=self.cancel_w2p)
            self.btn_w2p_cancel.pack(side="left")

            row_w2p_progress = tk.Frame(frm_w2p); row_w2p_progress.pack(fill="x", padx=6, pady=(0, 4))
            self.pb_w2p = ttk.Progressbar(row_w2p_progress, orient="horizontal", mode="determinate", length=400)
            self.pb_w2p.pack(side="left", padx=(0, 10))
            self.lbl_w2p_progress = tk.Label(row_w2p_progress, text="", fg="#555", anchor="w")
            self.lbl_w2p_progress.pack(side="left", fill="x", expand=True)

            tk.Label(frm_w2p, text="⑧ Bảng trạng thái:", font=("Arial", 9, "bold")).pack(anchor="w", padx=6, pady=(6, 2))
            cols_w2p = ("stt", "word", "pdf", "status", "note")
            headers_w2p = ["STT", "Tên file Word", "Tên file PDF xuất ra", "Trạng thái", "Ghi chú lỗi"]
            widths_w2p = [40, 220, 220, 150, 260]
            self.tree_w2p = ttk.Treeview(frm_w2p, columns=cols_w2p, show="headings", height=10)
            for c, h, w in zip(cols_w2p, headers_w2p, widths_w2p):
                self.tree_w2p.heading(c, text=h)
                self.tree_w2p.column(c, width=w, anchor="w")
            vsb_w2p = tk.Scrollbar(frm_w2p, orient="vertical", command=self.tree_w2p.yview)
            self.tree_w2p.configure(yscrollcommand=vsb_w2p.set)
            self.tree_w2p.pack(side="left", fill="both", expand=True, padx=(6, 0), pady=(0, 8))
            vsb_w2p.pack(side="left", fill="y", pady=(0, 8), padx=(0, 6))

        def pick_word_input(self):
            p = filedialog.askdirectory(title="Chọn thư mục chứa PDF nguồn")
            if p:
                self.var_word_input.set(p)

        def pick_word_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu file Word")
            if p:
                self.var_word_output.set(p)

        def run_word_export_clicked(self):
            source = self.word_input_picker.get_source()
            output = self.var_word_output.get().strip()
            merge_flag = self.var_word_merge.get()
            merged_name = self.var_word_merged_name.get().strip() or "TONG_HOP.docx"
            mode_map = {"Tự nhận diện": "tu_nhan_dien", "Nhanh": "nhanh",
                       "Giữ bố cục": "giu_bo_cuc", "OCR bản quét": "ocr"}
            mode = mode_map.get(self.var_word_mode.get(), "tu_nhan_dien")
            suffix = self.var_word_suffix.get().strip()
            no_overwrite = self.var_word_no_overwrite.get()

            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục hoặc (các) file PDF nguồn.")
                return
            if not output:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu file Word.")
                return
            if not HAS_PDF2DOCX:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install pdf2docx")
                return
            if merge_flag and not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install docxcompose python-docx")
                return

            self.txt_log.delete("1.0", "end")
            self._enable_pause_cancel("control_word_export", self.btn_word_export_pause, self.btn_word_export_cancel)
            self.pb_word_export.config(value=0, maximum=100)
            self.lbl_word_export_progress.config(text="Đang chuẩn bị...")
            self._word_export_start_time = time.time()

            def update_progress_ui(idx, total, fname):
                elapsed = time.time() - self._word_export_start_time
                if idx > 1:
                    avg = elapsed / (idx - 1)
                    remaining = avg * (total - idx + 1)
                    eta_text = f" | Còn khoảng {int(remaining // 60)} phút {int(remaining % 60)} giây"
                else:
                    eta_text = ""
                self.pb_word_export.config(value=idx, maximum=total)
                self.lbl_word_export_progress.config(
                    text=f"Đang xử lý {idx}/{total}: {fname}  (đã chạy {int(elapsed)}s{eta_text})")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_progress_ui(idx, total, fname))
                append_tab_log_row(7, "7a. Chuyển đổi Word/PDF hàng loạt", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XU_LY")

            def worker():
                try:
                    log_rows = batch_export_word_smart(
                        source, output, mode, suffix, no_overwrite, self.log,
                        control=self.control_word_export, progress_cb=progress_cb)
                    n_ok = sum(1 for r in log_rows if r["status"] == "THANH_CONG")
                    n_check = sum(1 for r in log_rows if r["status"].startswith("CAN_KIEM_TRA"))
                    n_loi = sum(1 for r in log_rows if r["status"] == "LOI_CHUYEN_PDF")
                    self.after(0, lambda: self.lbl_word_export_progress.config(
                        text=f"✓ Hoàn tất — {len(log_rows)} file ({int(time.time() - self._word_export_start_time)}s)"))

                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    csv_path = os.path.join(output, f"LOG_CHUYEN_DOI_WORD_PDF_{stamp}.csv")
                    write_log_chuyen_doi_word_pdf(csv_path, log_rows)

                    try:
                        std_log = _create_standard_run_log(
                            get_standard_logs_dir(), "CHUYEN_DOI_XUAT_HO_SO", run_id=stamp)
                        for r in log_rows:
                            std_log.add(action="PDF_SANG_WORD", source_file=r.get("file_in", ""),
                                       output_file=r.get("file_out", ""), status=r["status"],
                                       message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    merged_path = None
                    docx_paths = [r["path_out"] for r in log_rows if r["path_out"]]
                    if merge_flag and docx_paths:
                        self.log("Đang gộp tất cả vào 1 file .docx...")
                        master = DocxDocument(docx_paths[0])
                        composer = DocxComposer(master)
                        for p in docx_paths[1:]:
                            composer.doc.add_page_break()
                            composer.append(DocxDocument(p))
                        merged_path = os.path.join(output, merged_name)
                        composer.save(merged_path)
                        self.log(f"✓ Đã gộp thành: {merged_path}")

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {n_ok} thành công, {n_check} cần kiểm tra, {n_loi} lỗi.")
                    self.log(f"Log chi tiết: {csv_path}")
                    if merged_path:
                        self.log(f"File gộp: {merged_path}")
                    messagebox.showinfo("Hoàn tất xuất Word",
                                         f"Thành công: {n_ok}\nCần kiểm tra: {n_check}\nLỗi: {n_loi}\n\n"
                                         f"Log: {csv_path}" + (f"\nFile gộp: {merged_path}" if merged_path else ""))
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self._disable_pause_cancel(self.btn_word_export_pause, self.btn_word_export_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 7b: Xuất Word từ Excel bất kỳ -------------------

        def pick_wx_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel", filetypes=[("Excel files", "*.xlsx *.xlsm")])
            if p:
                self.var_wx_excel.set(p)

        def pick_wx_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_wx_output.set(p)

        def run_wx_clicked(self):
            excel_path = self.var_wx_excel.get().strip()
            output_folder = self.var_wx_output.get().strip()
            template_text = self.txt_wx_template.get("1.0", "end").rstrip("\n")
            filename_template = self.var_wx_filename_template.get().strip() or "HoSo_{stt}"

            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel hợp lệ.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return
            if not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install python-docx docxcompose")
                return

            try:
                header_row = int(self.var_wx_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            col_map = {
                "stt": self.var_wx_col_stt.get().strip(),
                "maxa": self.var_wx_col_maxa.get().strip(),
                "to": self.var_wx_col_to.get().strip(),
                "thua": self.var_wx_col_thua.get().strip(),
                "ten": self.var_wx_col_ten.get().strip(),
                "dia_chi": self.var_wx_col_diachi.get().strip(),
                "dt": self.var_wx_col_dt.get().strip(),
                "id": self.var_wx_col_id.get().strip(),
                "mucdich": self.var_wx_col_mucdich.get().strip(),
                "core": self.var_wx_col_core.get().strip(),
                "year": self.var_wx_col_year.get().strip(),
            }

            self.txt_log.delete("1.0", "end")
            self.btn_wx.config(state="disabled")
            self._enable_pause_cancel("control_wx", self.btn_wx_pause, self.btn_wx_cancel)

            def worker():
                try:
                    docx_paths, merged_path = export_word_from_excel(
                        excel_path=excel_path,
                        sheet_name=self.var_wx_sheet.get().strip() or None,
                        header_row=header_row,
                        col_map=col_map,
                        template_text=template_text,
                        output_folder=output_folder,
                        combine_into_one=self.var_wx_merge.get(),
                        filename_template=filename_template,
                        log_cb=self.log,
                        control=self.control_wx,
                    )
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã xuất {len(docx_paths)} file .docx từ Excel.")
                    if merged_path:
                        self.log(f"File gộp: {merged_path}")
                    messagebox.showinfo("Hoàn tất xuất Word từ Excel",
                                         f"Đã xuất {len(docx_paths)} file .docx.\n" +
                                         (f"File gộp: {merged_path}" if merged_path else ""))
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_wx.config(state="normal")
                    self._disable_pause_cancel(self.btn_wx_pause, self.btn_wx_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 7c: Xuất Word theo mẫu có sẵn (mail-merge) -------------------

        def pick_mailmerge_template(self):
            p = filedialog.askopenfilename(title="Chọn file mẫu .docx", filetypes=[("Word files", "*.docx")])
            if p:
                self.var_mailmerge_template.set(p)

        def use_bundled_template(self, filename):
            path = os.path.join(get_base_dir(), "assets", filename)
            if os.path.isfile(path):
                self.var_mailmerge_template.set(path)
            else:
                messagebox.showerror("Không tìm thấy mẫu", f"Không tìm thấy file mẫu đi kèm: {filename}")

        def pick_mailmerge_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục lưu kết quả")
            if p:
                self.var_mailmerge_output.set(p)

        def run_mailmerge_from_pdf_clicked(self):
            template_path = self.var_mailmerge_template.get().strip()
            output_folder = self.var_mailmerge_output.get().strip()
            source = self.mailmerge_pdf_picker.get_source()
            filename_template = self.var_mailmerge_filename_template.get().strip() or "HoSo_{stt}"

            if not template_path or not os.path.isfile(template_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file mẫu .docx hợp lệ.")
                return
            if not source:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục HOẶC chọn file PDF nguồn dữ liệu.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return
            if not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install python-docx")
                return

            mapping_text = self.txt_xa_mapping.get("1.0", "end")
            xa_mapping = parse_xa_mapping(mapping_text)
            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            use_ocr = self.var_use_ocr.get()
            try:
                ocr_dpi = int(self.var_ocr_dpi.get().strip())
            except ValueError:
                ocr_dpi = 300

            self.txt_log.delete("1.0", "end")
            self.btn_mm_pdf.config(state="disabled")
            self._enable_pause_cancel("control_mm_pdf", self.btn_mm_pdf_pause, self.btn_mm_pdf_cancel)
            self.pb_mm_pdf.config(value=0, maximum=100)
            self._mm_pdf_start_time = time.time()

            def update_mm_pdf_progress(idx, total, ten):
                elapsed = time.time() - self._mm_pdf_start_time
                self.pb_mm_pdf.config(value=idx, maximum=total)
                self.lbl_mm_pdf_progress.config(text=f"Đang xử lý {idx}/{total}: {ten} (đã chạy {int(elapsed)}s)")

            def progress_cb_pdf_read(idx, total, fname):
                self.after(0, lambda: update_mm_pdf_progress(idx, total, f"(đọc PDF) {fname}"))
                append_tab_log_row(7, "7c. Xuất Word THEO MẪU (từ PDF) - đọc PDF", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_DOC_FILE")

            def progress_cb_merge(idx, total, ten):
                self.after(0, lambda: update_mm_pdf_progress(idx, total, f"(tạo Word) {ten}"))
                append_tab_log_row(7, "7c. Xuất Word THEO MẪU (từ PDF) - tạo Word", ten,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XUAT_FILE")

            def worker():
                try:
                    rows, error_files = compile_report_from_pdfs(
                        source, xa_mapping, thon_mapping, use_ocr, ocr_dpi, self.log,
                        debug=self.var_debug_content.get(), control=self.control_mm_pdf,
                        progress_cb=progress_cb_pdf_read,
                        toc_do_xu_ly=getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"),
                        so_file_moi_dot=getattr(self.app_config, "word_so_file_moi_dot", 100),
                        nghi_giua_dot_giay=getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0))
                    self.log(f"\nĐã đọc {len(rows)} thửa đất từ PDF, {len(error_files)} file lỗi.")
                    created, errors = batch_mail_merge_from_pdf_or_excel(
                        rows, template_path, output_folder, filename_template, self.log,
                        control=self.control_mm_pdf, progress_cb=progress_cb_merge)
                    self.after(0, lambda: self.lbl_mm_pdf_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._mm_pdf_start_time)}s)"))
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã tạo {len(created)} file Word theo mẫu | {len(errors)} lỗi")
                    messagebox.showinfo("Hoàn tất", f"Đã tạo {len(created)} file Word theo mẫu.\nLưu tại: {output_folder}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_mm_pdf.config(state="normal")
                    self._disable_pause_cancel(self.btn_mm_pdf_pause, self.btn_mm_pdf_cancel)

            threading.Thread(target=worker, daemon=True).start()

        def run_mailmerge_from_excel_clicked(self):
            template_path = self.var_mailmerge_template.get().strip()
            output_folder = self.var_mailmerge_output.get().strip()
            excel_path = self.var_wx_excel.get().strip()
            filename_template = self.var_mailmerge_filename_template.get().strip() or "HoSo_{stt}"

            if not template_path or not os.path.isfile(template_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file mẫu .docx hợp lệ.")
                return
            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel ở phần 'CÁCH 2' phía trên.")
                return
            if not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục lưu kết quả.")
                return
            if not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install python-docx")
                return

            try:
                header_row = int(self.var_wx_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu (phần CÁCH 2) phải là số.")
                return

            col_map = {
                "stt": self.var_wx_col_stt.get().strip(),
                "maxa": self.var_wx_col_maxa.get().strip(),
                "to": self.var_wx_col_to.get().strip(),
                "thua": self.var_wx_col_thua.get().strip(),
                "ten": self.var_wx_col_ten.get().strip(),
                "dia_chi": self.var_wx_col_diachi.get().strip(),
                "dt": self.var_wx_col_dt.get().strip(),
                "id": self.var_wx_col_id.get().strip(),
                "mucdich": self.var_wx_col_mucdich.get().strip(),
                "core": self.var_wx_col_core.get().strip(),
                "year": self.var_wx_col_year.get().strip(),
            }

            self.txt_log.delete("1.0", "end")
            self.btn_mm_excel.config(state="disabled")
            self._enable_pause_cancel("control_mm_excel", self.btn_mm_excel_pause, self.btn_mm_excel_cancel)
            self.pb_mm_excel.config(value=0, maximum=100)
            self._mm_excel_start_time = time.time()

            def update_mm_excel_progress(idx, total, ten):
                elapsed = time.time() - self._mm_excel_start_time
                self.pb_mm_excel.config(value=idx, maximum=total)
                self.lbl_mm_excel_progress.config(text=f"Đang xử lý {idx}/{total}: {ten} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, ten):
                self.after(0, lambda: update_mm_excel_progress(idx, total, ten))
                append_tab_log_row(7, "7c. Xuất Word THEO MẪU (từ Excel)", ten,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XUAT_FILE")

            def worker():
                try:
                    rows = load_excel_rows_generic(excel_path, self.var_wx_sheet.get().strip() or None,
                                                    header_row, col_map)
                    self.log(f"Đã đọc {len(rows)} dòng từ Excel.")
                    created, errors = batch_mail_merge_from_pdf_or_excel(
                        rows, template_path, output_folder, filename_template, self.log,
                        control=self.control_mm_excel, progress_cb=progress_cb)
                    self.after(0, lambda: self.lbl_mm_excel_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._mm_excel_start_time)}s)"))
                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: đã tạo {len(created)} file Word theo mẫu | {len(errors)} lỗi")
                    messagebox.showinfo("Hoàn tất", f"Đã tạo {len(created)} file Word theo mẫu.\nLưu tại: {output_folder}")
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_mm_excel.config(state="normal")
                    self._disable_pause_cancel(self.btn_mm_excel_pause, self.btn_mm_excel_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 7d: Xuất Word hàng loạt (Excel tổng + mẫu GT/TBXN) -------------------

        def pick_wb_excel(self):
            p = filedialog.askopenfilename(title="Chọn file Excel tổng",
                                            filetypes=[("Excel files", "*.xlsx *.xlsm *.xls")])
            if p:
                self.var_wb_excel.set(p)

        def pick_wb_template_gt(self):
            p = filedialog.askopenfilename(title="Chọn mẫu Đơn đăng ký đất đai",
                                            filetypes=[("Word files", "*.docx *.doc")])
            if p:
                self.var_wb_template_gt.set(p)

        def pick_wb_template_tbxn(self):
            p = filedialog.askopenfilename(title="Chọn mẫu Thông báo xác nhận",
                                            filetypes=[("Word files", "*.docx *.doc")])
            if p:
                self.var_wb_template_tbxn.set(p)

        def pick_wb_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục xuất file")
            if p:
                self.var_wb_output.set(p)

        def use_bundled_template_for(self, var, filename):
            path = os.path.join(get_base_dir(), "assets", filename)
            if os.path.isfile(path):
                var.set(path)
            else:
                messagebox.showerror("Không tìm thấy mẫu", f"Không tìm thấy file mẫu đi kèm: {filename}")

        def _refresh_wb_tree(self, results):
            for item in self.tree_wb.get_children():
                self.tree_wb.delete(item)
            for r in results:
                self._insert_row_colored(self.tree_wb,
                    (r["stt"], r["ten"], r["maxa"], r["to"], r["thua"], r["filename"], r["status"], r["note"]),
                    r["status"])

        def preview_suffix_filename(self):
            """Xem trước tên file sẽ xuất ra dựa trên cấu hình hậu tố hiện tại (dữ liệu mẫu minh họa)."""
            hauto, nguon = resolve_file_suffix(
                "GT",
                manual_suffix=self.var_suffix_manual.get(),
                combobox_suffix=self.var_suffix_combobox.get(),
                excel_value="(giá trị mẫu ở cột Excel)" if self.var_suffix_use_excel.get() else "",
                use_excel_col=self.var_suffix_use_excel.get(),
                prioritize_manual=self.var_suffix_prioritize_manual.get(),
                uppercase=self.var_suffix_uppercase.get(),
                remove_diacritics_opt=self.var_suffix_remove_diacritics.get(),
            )
            try:
                fname = render_output_filename(
                    self.var_suffix_filename_template.get(), maxa="02140", to="29", thua="199",
                    ten="Nguyễn Văn A", cccd="001234567890", thon="Nà Dường", xacu="Văn Lang cũ",
                    hauto=hauto, stt="1")
                self.var_suffix_preview.set(f"{fname}.docx   (nguồn hậu tố: {nguon})")
            except ValueError as e:
                self.var_suffix_preview.set(f"⚠ Lỗi mẫu tên file: {e}")

        def run_word_batch_clicked(self, dry_run=True):
            excel_path = self.var_wb_excel.get().strip()
            template_gt = self.var_wb_template_gt.get().strip()
            template_tbxn = self.var_wb_template_tbxn.get().strip()
            output_folder = self.var_wb_output.get().strip()
            export_gt = self.var_wb_export_gt.get()
            export_tbxn = self.var_wb_export_tbxn.get()
            combine = self.var_wb_combine.get()

            if not excel_path or not os.path.isfile(excel_path):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn file Excel tổng hợp lệ.")
                return
            if not export_gt and not export_tbxn:
                messagebox.showerror("Thiếu thông tin", "Vui lòng tick ít nhất 1 mẫu để xuất (GT hoặc TBXN).")
                return
            if export_gt and (not template_gt or not os.path.isfile(template_gt)):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn mẫu Đơn đăng ký hợp lệ.")
                return
            if export_tbxn and (not template_tbxn or not os.path.isfile(template_tbxn)):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn mẫu Thông báo xác nhận hợp lệ.")
                return
            if not dry_run and not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục xuất file.")
                return
            if not HAS_DOCXCOMPOSE:
                messagebox.showerror("Thiếu thư viện", "Cần cài: pip install python-docx docxcompose")
                return

            try:
                header_row = int(self.var_wb_header_row.get().strip())
            except ValueError:
                messagebox.showerror("Lỗi", "Dòng bắt đầu dữ liệu phải là số.")
                return

            col_map = {
                "stt": "",
                "maxa": self.var_wb_col_maxa.get().strip(),
                "name": self.var_wb_col_name.get().strip(),
                "id": self.var_wb_col_id.get().strip(),
                "location": self.var_wb_col_location.get().strip(),
                "diachithua": self.var_wb_col_diachithua.get().strip(),
                "numberto": self.var_wb_col_to.get().strip(),
                "numberthua": self.var_wb_col_thua.get().strip(),
                "Sdientich": self.var_wb_col_dt.get().strip(),
                "mucdich": self.var_wb_col_mucdich.get().strip(),
                "core": self.var_wb_col_core.get().strip(),
                "year": self.var_wb_col_year.get().strip(),
                "numbertb": self.var_wb_col_numbertb.get().strip(),
                "numberbb": self.var_wb_col_numberbb.get().strip(),
                "numbercv": self.var_wb_col_numbercv.get().strip(),
                "numberkh": self.var_wb_col_numberkh.get().strip(),
                "sovanban": self.var_wb_col_sovanban.get().strip(),
                "kyhieu": self.var_wb_col_kyhieu.get().strip(),
                "ngayvanban": self.var_wb_col_ngayvanban.get().strip(),
                "hauto": self.var_wb_col_hauto.get().strip(),
            }

            thon_mapping = parse_xa_mapping(self.txt_thon_mapping.get("1.0", "end"))
            xa_mapping = parse_xa_mapping(self.txt_xa_mapping.get("1.0", "end"))

            suffix_config = None
            if self.var_suffix_enabled.get():
                suffix_config = {
                    "enabled": True,
                    "combobox": self.var_suffix_combobox.get(),
                    "manual": self.var_suffix_manual.get(),
                    "prioritize_manual": self.var_suffix_prioritize_manual.get(),
                    "use_excel_col": self.var_suffix_use_excel.get(),
                    "uppercase": self.var_suffix_uppercase.get(),
                    "remove_diacritics": self.var_suffix_remove_diacritics.get(),
                    "filename_template": self.var_suffix_filename_template.get().strip() or DEFAULT_FILENAME_TEMPLATE,
                }

            self.txt_log.delete("1.0", "end")
            self._refresh_wb_tree([])

            checkpoint_task_id = None
            stt_da_xong = None
            if not dry_run:
                from app.services import checkpoint_service as _ckpt
                checkpoint_task_id = _ckpt.tinh_task_id("xuat_word_hang_loat", excel_path, output_folder)
                data_cu = _ckpt.doc_checkpoint(checkpoint_task_id)
                if data_cu and data_cu.get("trang_thai") == "DANG_XU_LY":
                    so_da_xong = len(data_cu.get("da_hoan_thanh", []))
                    so_tong = len(data_cu.get("danh_sach_file", []))
                    muon_tiep_tuc = messagebox.askyesno(
                        "Phát hiện tác vụ dở dang",
                        f"Phát hiện đợt xuất Word TRƯỚC ĐÓ chưa hoàn thành (đã xong {so_da_xong}/{so_tong} "
                        f"dòng, có thể do phần mềm bị đóng đột ngột).\n\n"
                        f"Bấm CÓ để TIẾP TỤC từ chỗ dở dang (không xử lý lại các dòng đã xong).\n"
                        f"Bấm KHÔNG để BẮT ĐẦU LẠI TỪ ĐẦU.")
                    if muon_tiep_tuc:
                        stt_da_xong = set(data_cu.get("da_hoan_thanh", []) + data_cu.get("file_loi", []) +
                                          data_cu.get("can_kiem_tra", []))
                    else:
                        _ckpt.xoa_checkpoint(checkpoint_task_id)

            self.btn_wb_dryrun.config(state="disabled")
            self.btn_wb_run.config(state="disabled")
            self._enable_pause_cancel("control_wb", self.btn_wb_pause, self.btn_wb_cancel)
            self.pb_wb.config(value=0, maximum=100)
            self.lbl_wb_perf.config(text="")
            self._wb_start_time = time.time()

            from app.services import task_manager_service as _tm
            self.task_info_wb = _tm.TaskInfo(
                ten_chuc_nang="7d. Xuất Word hàng loạt (Excel tổng + mẫu)",
                thu_muc_dau_ra=output_folder, trang_thai=_tm.TrangThaiTask.DANG_KHOI_TAO)

            def update_wb_progress(idx, total, ten):
                elapsed = time.time() - self._wb_start_time
                self.pb_wb.config(value=idx, maximum=total)
                self.lbl_wb_progress.config(text=f"Đang xử lý {idx}/{total}: {ten} (đã chạy {int(elapsed)}s)")

            def update_wb_perf(stats):
                parts = [f"CPU: {stats.get('cpu_percent', 0):.0f}%", f"RAM: {stats.get('ram_mb', 0):.0f} MB"]
                so_word = stats.get("so_tien_trinh_word")
                if so_word is not None:
                    parts.append(f"Số WINWORD.EXE: {so_word}")
                self.lbl_wb_perf.config(text="  |  ".join(parts))

            def perf_stats_cb(stats):
                self.after(0, lambda: update_wb_perf(stats))

            def progress_cb(idx, total, ten):
                self.after(0, lambda: update_wb_progress(idx, total, ten))
                self.task_info_wb.cap_nhat_tien_do(tien_do_toan_bo=idx,
                                                   trang_thai=_tm.TrangThaiTask.DANG_XUAT_FILE)
                self.task_info_wb.tong_so_luong = total
                append_tab_log_row(7, "7d. Xuất Word hàng loạt (Excel tổng + mẫu)", ten,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XUAT_FILE")

            def worker():
                try:
                    rows = read_excel_rows_smart(excel_path, self.var_wb_sheet.get().strip() or None,
                                                  header_row, col_map, log_cb=self.log)
                    self.log(f"Đã đọc {len(rows)} dòng từ Excel.")

                    if HAS_WIN32COM and not dry_run:
                        self.log("Đang mở 1 phiên Microsoft Word DÙNG CHUNG cho toàn bộ đợt xuất (nhanh hơn "
                                 "mở/tắt Word từng hồ sơ)...")
                        perf_log_path = None
                        if getattr(self.app_config, "word_hien_thi_thong_ke_hieu_nang", True):
                            perf_log_path = os.path.join(get_standard_logs_dir(), "LOG_HIEU_NANG_WORD.csv")
                        with WordCOMSession() as word_session:
                            results = run_word_batch_export(
                                rows, thon_mapping, xa_mapping, template_gt, template_tbxn,
                                export_gt, export_tbxn, combine, output_folder, dry_run, self.log,
                                control=self.control_wb, word_session=word_session,
                                commune_rows=self.commune_config_rows, suffix_config=suffix_config,
                                progress_cb=progress_cb,
                                toc_do_xu_ly=getattr(self.app_config, "word_toc_do_xu_ly", "can_bang"),
                                so_file_moi_dot=getattr(self.app_config, "word_so_file_moi_dot", 100),
                                nghi_giua_dot_giay=getattr(self.app_config, "word_nghi_giua_dot_giay", 2.0),
                                perf_log_path=perf_log_path,
                                gioi_han_ram_mb=getattr(self.app_config, "word_gioi_han_ram_mb", 0),
                                perf_stats_cb=perf_stats_cb,
                                checkpoint_task_id=checkpoint_task_id, stt_da_xong=stt_da_xong)
                    else:
                        results = run_word_batch_export(
                            rows, thon_mapping, xa_mapping, template_gt, template_tbxn,
                            export_gt, export_tbxn, combine, output_folder, dry_run, self.log,
                            control=self.control_wb, commune_rows=self.commune_config_rows,
                            suffix_config=suffix_config, progress_cb=progress_cb,
                            checkpoint_task_id=checkpoint_task_id, stt_da_xong=stt_da_xong)

                    self.after(0, lambda: self._refresh_wb_tree(results))
                    self.after(0, lambda: self.lbl_wb_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._wb_start_time)}s)"))

                    for _r in results:
                        self.task_info_wb.ghi_nhan_ket_qua(_r.get("status", ""))
                    self.task_info_wb.trang_thai = _tm.TrangThaiTask.THANH_CONG

                    import time
                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    log_dir = output_folder or os.path.dirname(excel_path)
                    csv_path = os.path.join(log_dir, f"LOG_XUAT_WORD_{stamp}.csv")
                    write_word_batch_log_csv(csv_path, results)

                    try:
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "XUAT_WORD_HANG_LOAT", run_id=stamp)
                        for r in results:
                            std_log.add(action="EXPORT_WORD", source_file=excel_path,
                                       output_file=r.get("filename", ""), ma_xa=r.get("maxa"),
                                       so_to=r.get("to"), so_thua=r.get("thua"),
                                       status=r["status"], message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    ok_count = sum(1 for r in results if r["status"] in ("OK", "OK (chạy thử - chưa xuất)"))
                    err_count = sum(1 for r in results if r["status"] in ("LỖI", "CẦN KIỂM TRA"))

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {ok_count} OK | {err_count} cần kiểm tra/lỗi")
                    self.log(f"File log CSV: {csv_path}")

                    msg = f"Đã xử lý xong.\nOK: {ok_count}\nCần kiểm tra/Lỗi: {err_count}\n\nLog CSV: {csv_path}"
                    self.after(0, lambda: self.mark_workflow_step("xuat_word", "Đã chạy thử" if dry_run else "Đã hoàn thành"))
                    messagebox.showinfo("Chạy thử hoàn tất" if dry_run else "Xuất Word hoàn tất", msg)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_wb_dryrun.config(state="normal")
                    self.btn_wb_run.config(state="normal")
                    self._disable_pause_cancel(self.btn_wb_pause, self.btn_wb_cancel)

            threading.Thread(target=worker, daemon=True).start()

        # ------------------- MỤC 7e: Chuyển Word sang PDF hàng loạt -------------------

        def pick_w2p_input(self):
            p = filedialog.askdirectory(title="Chọn thư mục chứa file Word")
            if p:
                self.var_w2p_input.set(p)

        def pick_w2p_output(self):
            p = filedialog.askdirectory(title="Chọn thư mục xuất file PDF")
            if p:
                self.var_w2p_output.set(p)

        def _refresh_w2p_tree(self, results):
            for item in self.tree_w2p.get_children():
                self.tree_w2p.delete(item)
            for r in results:
                self._insert_row_colored(self.tree_w2p,
                    (r["stt"], r["word"], r["pdf"], r["status"], r["note"]), r["status"])

        def toggle_pause_w2p(self):
            if self.control_w2p.is_paused():
                self.control_w2p.resume()
                self.btn_w2p_pause.config(text="⏸ Tạm dừng")
                self.log("▶ Đã tiếp tục.")
            else:
                self.control_w2p.pause()
                self.btn_w2p_pause.config(text="▶ Tiếp tục")
                self.log("⏸ Đã tạm dừng — bấm 'Tiếp tục' để chạy tiếp, hoặc 'Hủy' để dừng hẳn.")

        def cancel_w2p(self):
            if messagebox.askyesno("Xác nhận hủy", "Hủy tác vụ đang chạy? Các file chưa xử lý sẽ bị bỏ qua."):
                self.control_w2p.cancel()
                self.log("⏹ Đang hủy...")

        def run_w2p_clicked(self, dry_run=True):
            source = self.var_w2p_input.get().strip()
            output_folder = self.var_w2p_output.get().strip()

            if not source or not os.path.isdir(source):
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục chứa file Word hợp lệ.")
                return
            if not dry_run and not output_folder:
                messagebox.showerror("Thiếu thông tin", "Vui lòng chọn thư mục xuất file PDF.")
                return

            include_sub = self.var_w2p_subfolders.get()
            overwrite = self.var_w2p_overwrite.get()
            keep_struct = self.var_w2p_keepstruct.get()

            self.txt_log.delete("1.0", "end")
            self._refresh_w2p_tree([])
            self.control_w2p.reset()
            self.btn_w2p_dryrun.config(state="disabled")
            self.btn_w2p_run.config(state="disabled")
            self.btn_w2p_pause.config(state="normal", text="⏸ Tạm dừng")
            self.btn_w2p_cancel.config(state="normal")
            self.pb_w2p.config(value=0, maximum=100)
            self._w2p_start_time = time.time()

            def update_w2p_progress(idx, total, fname):
                elapsed = time.time() - self._w2p_start_time
                self.pb_w2p.config(value=idx, maximum=total)
                self.lbl_w2p_progress.config(text=f"Đang xử lý {idx}/{total}: {fname} (đã chạy {int(elapsed)}s)")

            def progress_cb(idx, total, fname):
                self.after(0, lambda: update_w2p_progress(idx, total, fname))
                append_tab_log_row(7, "7e. Chuyển Word sang PDF hàng loạt", fname,
                                   tien_do=f"{idx}/{total}", trang_thai="DANG_XUAT_FILE")

            def worker():
                try:
                    if HAS_WIN32COM and not dry_run:
                        self.log("Đang mở 1 phiên Microsoft Word DÙNG CHUNG cho toàn bộ đợt chuyển đổi "
                                 "(nhanh hơn mở/tắt Word từng file)...")
                        with WordCOMSession() as word_session:
                            results = batch_word_to_pdf(
                                source, output_folder or source, include_sub, overwrite, keep_struct,
                                dry_run, self.log, control=self.control_w2p, word_session=word_session,
                                progress_cb=progress_cb)
                    else:
                        results = batch_word_to_pdf(
                            source, output_folder or source, include_sub, overwrite, keep_struct,
                            dry_run, self.log, control=self.control_w2p, progress_cb=progress_cb)

                    self.after(0, lambda: self._refresh_w2p_tree(results))
                    self.after(0, lambda: self.lbl_w2p_progress.config(
                        text=f"✓ Hoàn tất ({int(time.time() - self._w2p_start_time)}s)"))

                    stamp = time.strftime("%Y%m%d_%H%M%S")
                    log_dir = output_folder or source
                    csv_path = os.path.join(log_dir, f"LOG_WORD_TO_PDF_{stamp}.csv")
                    write_word_to_pdf_log_csv(csv_path, results)

                    try:
                        std_log = _create_standard_run_log(get_standard_logs_dir(), "CHUYEN_WORD_PDF", run_id=stamp)
                        for r in results:
                            std_log.add(action="WORD_TO_PDF", source_file=r.get("word", ""),
                                       output_file=r.get("pdf", ""), status=r["status"],
                                       message=r.get("note", ""))
                        std_log_path = std_log.save()
                        self.log(f"File log chuẩn (Logs chung): {std_log_path}")
                    except Exception as e:
                        self.log(f"⚠ Không ghi được log chuẩn (không ảnh hưởng log CSV chính): {e}")

                    ok_count = sum(1 for r in results if r["status"].startswith("OK"))
                    other_count = len(results) - ok_count

                    self.log("\n" + "=" * 60)
                    self.log(f"TỔNG KẾT: {ok_count} OK | {other_count} bỏ qua/lỗi")
                    self.log(f"File log CSV: {csv_path}")

                    msg = f"OK: {ok_count}\nBỏ qua/Lỗi: {other_count}\n\nLog CSV: {csv_path}"
                    self.after(0, lambda: self.mark_workflow_step("chuyen_pdf", "Đã chạy thử" if dry_run else "Đã hoàn thành"))
                    messagebox.showinfo("Chạy thử hoàn tất" if dry_run else "Chuyển đổi hoàn tất", msg)
                except Exception as e:
                    self.log(f"LỖI: {e}")
                    messagebox.showerror("Lỗi xử lý", str(e))
                finally:
                    self.btn_w2p_dryrun.config(state="normal")
                    self.btn_w2p_run.config(state="normal")
                    self.btn_w2p_pause.config(state="disabled", text="⏸ Tạm dừng")
                    self.btn_w2p_cancel.config(state="disabled")

            threading.Thread(target=worker, daemon=True).start()



class App(tk.Tk, SettingsTabMixin, DupTabMixin, MergeWordTabMixin, SplitTabMixin, ExcelTabMixin, ReconcileTabMixin, WorkflowTabMixin, ContentTabMixin, PdfEditTabMixin, ReportTabMixin, WordTabMixin):
    def __init__(self):
        super().__init__()
        self.title(f"Tiện ích CSDL Đất Đai - SỸ LAND {RELEASE_LABEL} (build {APP_VERSION})")
        self.geometry("1200x820")
        self.minsize(900, 600)
        try:
            self.state("zoomed")  # mở full màn hình sẵn trên Windows để thấy hết các tiện ích
        except tk.TclError:
            pass
        self._set_window_icon()
        apply_default_fonts(self)

        # Nạp cấu hình phần mềm (data/app_settings.json) - dùng module config_loader nếu có,
        # thiếu thì dùng AppConfig mặc định (không lỗi, không treo).
        self.app_config = load_app_config_safe()

        self.var_folder1 = tk.StringVar()
        self.var_folder2 = tk.StringVar()
        self.var_excel = tk.StringVar()
        self.var_output_matched = tk.StringVar()
        self.var_output_unmatched = tk.StringVar()
        self.var_sheet = tk.StringVar()
        self.var_header_row = tk.StringVar(value="5")
        self.var_col_stt = tk.StringVar(value="A")
        self.var_col_maxa = tk.StringVar(value="B")
        self.var_col_to = tk.StringVar(value="V")
        self.var_col_thua = tk.StringVar(value="W")
        self.var_col_ten = tk.StringVar(value="H")
        self.var_col_files = tk.StringVar(value="AX")
        self.var_col_dt = tk.StringVar(value="Y")
        self.var_dt_tolerance = tk.StringVar(value="5")

        self.var_crit_maxa = tk.BooleanVar(value=True)
        self.var_crit_to = tk.BooleanVar(value=True)
        self.var_crit_thua = tk.BooleanVar(value=True)
        self.var_crit_dt = tk.BooleanVar(value=False)
        self.var_crit_ten = tk.BooleanVar(value=False)

        self.var_rename_folder = tk.StringVar()
        self.var_name_template = tk.StringVar(value="CHUACOGIAY_{maxa}_{to}_{thua}_{loai}")
        self.var_rename_content_fallback = tk.BooleanVar(value=True)

        # --- Mục 5: Đổi tên theo nội dung PDF ---
        self.var_pdfcontent_folder = tk.StringVar()
        self.var_use_ocr = tk.BooleanVar(value=True)
        self.var_ocr_dpi = tk.StringVar(value="300")
        self.var_ren_execute_now = tk.BooleanVar(value=False)
        self.var_debug_content = tk.BooleanVar(value=False)
        self.var_content_name_template = tk.StringVar(value="CHUACOGIAY_{maxa}_{to}_{thua}_{loai}")

        # --- Chế độ OCR chuyên dụng cho PDF scan ---
        self.var_ocr_matrix = tk.StringVar(value="4")
        self.var_ocr_mode = tk.StringVar(value=getattr(self.app_config, "ocr_mode", "Chuẩn"))
        self.var_ocr_suffix = tk.StringVar(value="GT")
        self.var_ocr_output = tk.StringVar()
        self.var_ocr_debug = tk.BooleanVar(value=False)
        self.last_ocr_rename_result = None
        self.commune_config_rows = []  # cấu hình địa bàn đã nạp (list CommuneConfigRow), rỗng nếu chưa nạp

        # Cờ theo dõi "đã chạy thử chưa" trong phiên làm việc - dùng cho cài đặt "Bắt buộc chạy thử"
        self._dry_run_done_ocr = False
        self._dry_run_done_dup = False
        self._dry_run_done_reconcile = False
        self._dry_run_done_pdfedit = False
        self._dry_run_done_compare_folders = False

        # --- Mục 8: Tổng hợp báo cáo từ PDF ---
        self.var_report_folder = tk.StringVar()
        self.var_report_output = tk.StringVar()

        # --- Mục 8b: Thống kê tiến độ từ Excel bất kỳ ---
        self.var_stat_excel = tk.StringVar()
        self.var_stat_sheet = tk.StringVar()
        self.var_stat_header_row = tk.StringVar(value="2")
        self.var_stat_col_count = tk.StringVar(value="A")
        self.var_stat_col_status = tk.StringVar(value="")
        self.var_stat_done_values = tk.StringVar(value="Đã thực hiện, Đã nộp, Xong")
        self.var_stat_col_gcn = tk.StringVar(value="")
        self.var_stat_gcn_values = tk.StringVar(value="Đã nhập, Có GCN")
        self.var_stat_target_total = tk.StringVar(value="")

        # --- Mục 8b-bis: Khối lượng và tiến độ thực hiện ---
        self.var_prog_total = tk.StringVar(value="")
        self.var_prog_done = tk.StringVar(value="")
        self.var_prog_inprogress = tk.StringVar(value="")
        self.var_prog_remaining = tk.StringVar(value="0")
        self.var_prog_percent = tk.StringVar(value="0%")
        self.var_prog_auto_count = tk.BooleanVar(value=True)
        self.var_prog_detect_filename = tk.BooleanVar(value=True)
        self.var_prog_detect_content = tk.BooleanVar(value=False)
        self.prog_scan_results = []  # kết quả quét gần nhất (list dict) - dùng khi xuất báo cáo

        # --- Mục 8c: Điền vào Excel mẫu có sẵn ---
        self.var_wb8c_template = tk.StringVar()
        self.var_wb8c_output = tk.StringVar()
        self.var_wb8c_sheet = tk.StringVar(value="Chưa GCN")
        self.var_wb8c_header_row = tk.StringVar(value="5")
        self.var_wb8c_clear_old = tk.BooleanVar(value=False)
        self.var_wb8c_infer_cccd = tk.BooleanVar(value=True)
        self.var_wb8c_no_overwrite = tk.BooleanVar(value=False)
        self.var_wb8c_warn_dup = tk.BooleanVar(value=True)
        self.var_wb8c_export_check = tk.BooleanVar(value=True)
        self.var_wb8c_debug_text = tk.BooleanVar(value=False)
        # Cấu hình cột động (12 trường) - giá trị mặc định khớp sheet "Chưa GCN"
        _default_cols = {"maxa": "B", "ten": "H", "cccd": "I", "ngaysinh": "J", "gioitinh": "K",
                         "diachitt": "L", "soto": "V", "sothua": "W", "diachithua": "X",
                         "dientich": "Y", "mucdich": "Z", "dientichloai1": "AA"}
        self.var_col8c = {k: tk.StringVar(value=v) for k, v in _default_cols.items()}
        self._default_col8c = dict(_default_cols)

        # --- Mục 6: Tách / Gộp PDF ---
        self.var_split_input = tk.StringVar()
        self.var_split_output = tk.StringVar()
        self.var_split_page_spec = tk.StringVar(value="1")
        self.var_split_suffix = tk.StringVar(value="_trang1")
        self.var_merge_output = tk.StringVar()
        self.var_group_input = tk.StringVar()
        self.var_group_output = tk.StringVar()
        self.var_group_size = tk.StringVar(value="2")

        # --- Mục 7: Xuất Word ---
        self.var_word_input = tk.StringVar()
        self.var_word_output = tk.StringVar()
        self.var_word_merge = tk.BooleanVar(value=False)
        self.var_word_merged_name = tk.StringVar(value="TONG_HOP.docx")

        # --- Mục 7b: Xuất Word từ Excel bất kỳ ---
        self.var_wx_excel = tk.StringVar()
        self.var_wx_sheet = tk.StringVar()
        self.var_wx_header_row = tk.StringVar(value="5")
        self.var_wx_col_stt = tk.StringVar(value="A")
        self.var_wx_col_maxa = tk.StringVar(value="B")
        self.var_wx_col_to = tk.StringVar(value="W")
        self.var_wx_col_thua = tk.StringVar(value="X")
        self.var_wx_col_ten = tk.StringVar(value="H")
        self.var_wx_col_diachi = tk.StringVar(value="Y")
        self.var_wx_col_dt = tk.StringVar(value="Z")
        self.var_wx_filename_template = tk.StringVar(value="HoSo_{maxa}_{to}_{thua}")
        self.var_wx_output = tk.StringVar()
        self.var_wx_col_id = tk.StringVar(value="I")
        self.var_wx_col_mucdich = tk.StringVar(value="Z")
        self.var_wx_col_core = tk.StringVar(value="AB")
        self.var_wx_col_year = tk.StringVar(value="AD")
        self.var_wx_merge = tk.BooleanVar(value=False)

        # --- Mục 7c: Xuất Word theo mẫu có sẵn (mail-merge) ---
        self.var_mailmerge_template = tk.StringVar()
        self.var_mailmerge_filename_template = tk.StringVar(value="HoSo_{maxa}_{to}_{thua}")
        self.var_mailmerge_output = tk.StringVar()

        # --- Mục 7d: Xuất Word hàng loạt (Excel tổng + mẫu GT/TBXN) ---
        self.var_wb_excel = tk.StringVar()
        self.var_wb_template_gt = tk.StringVar()
        self.var_wb_template_tbxn = tk.StringVar()
        self.var_wb_output = tk.StringVar()
        self.var_wb_sheet = tk.StringVar()
        self.var_wb_header_row = tk.StringVar(value="5")
        self.var_wb_col_maxa = tk.StringVar(value="")
        self.var_wb_col_name = tk.StringVar(value="H")
        self.var_wb_col_id = tk.StringVar(value="I")
        self.var_wb_col_location = tk.StringVar(value="X")
        self.var_wb_col_diachithua = tk.StringVar(value="")
        self.var_wb_col_to = tk.StringVar(value="V")
        self.var_wb_col_thua = tk.StringVar(value="W")
        self.var_wb_col_dt = tk.StringVar(value="Y")
        self.var_wb_col_mucdich = tk.StringVar(value="Z")
        self.var_wb_col_core = tk.StringVar(value="AB")
        self.var_wb_col_year = tk.StringVar(value="AD")
        self.var_wb_col_numbertb = tk.StringVar(value="")
        self.var_wb_col_numberbb = tk.StringVar(value="")
        self.var_wb_col_numbercv = tk.StringVar(value="")
        self.var_wb_col_numberkh = tk.StringVar(value="")
        self.var_wb_col_sovanban = tk.StringVar(value="")
        self.var_wb_col_kyhieu = tk.StringVar(value="")
        self.var_wb_col_ngayvanban = tk.StringVar(value="")

        # --- Cấu hình hậu tố tên file tùy chọn (Tab 7d) ---
        self.var_suffix_enabled = tk.BooleanVar(value=False)
        self.var_suffix_combobox = tk.StringVar(value="GT")
        self.var_suffix_manual = tk.StringVar(value="")
        self.var_suffix_prioritize_manual = tk.BooleanVar(value=False)
        self.var_suffix_use_excel = tk.BooleanVar(value=False)
        self.var_wb_col_hauto = tk.StringVar(value="")
        self.var_suffix_uppercase = tk.BooleanVar(value=False)
        self.var_suffix_remove_diacritics = tk.BooleanVar(value=False)
        self.var_suffix_filename_template = tk.StringVar(value=DEFAULT_FILENAME_TEMPLATE)
        self.var_suffix_preview = tk.StringVar(value="")
        self.var_wb_export_gt = tk.BooleanVar(value=True)
        self.var_wb_export_tbxn = tk.BooleanVar(value=True)
        self.var_wb_combine = tk.BooleanVar(value=False)

        # --- Mục 7e: Chuyển Word sang PDF hàng loạt ---
        self.var_w2p_input = tk.StringVar()
        self.var_w2p_output = tk.StringVar()
        self.var_w2p_subfolders = tk.BooleanVar(value=False)
        self.var_w2p_overwrite = tk.BooleanVar(value=False)
        self.var_w2p_keepstruct = tk.BooleanVar(value=True)
        self.control_w2p = TaskControl()

        # --- Mục 9: Gộp file Word hàng loạt ---
        self.var_merge_word_folder = tk.StringVar()
        self.var_merge_word_pagebreak = tk.BooleanVar(value=True)

        # --- Mục 10: Lọc file trùng giữa 2 thư mục ---
        self.var_dup_reference = tk.StringVar()
        self.var_dup_compare = tk.StringVar()
        self.var_dup_output = tk.StringVar()

        # --- Mục 11: Lọc Excel tổng theo PDF & hoàn thiện báo cáo ---
        self.var_rec_excel = tk.StringVar()
        self.var_rec_pdf_folder = tk.StringVar()
        self.var_rec_output = tk.StringVar()
        self.var_rec_sheet = tk.StringVar()
        self.var_rec_header_row = tk.StringVar(value="5")
        self.var_rec_col_maxa = tk.StringVar(value="B")
        self.var_rec_col_to = tk.StringVar(value="V")
        self.var_rec_col_thua = tk.StringVar(value="W")
        self.var_rec_col_stt = tk.StringVar(value="A")
        self.var_rec_subfolders = tk.BooleanVar(value=False)
        self.var_rec_complete = tk.BooleanVar(value=True)
        self.var_rec_overwrite = tk.BooleanVar(value=False)
        self.var_rec_renumber = tk.BooleanVar(value=True)

        # --- NHÓM C: So sánh thư mục Word/PDF ---
        self.var_cmp_folder_a = tk.StringVar()
        self.var_cmp_folder_b = tk.StringVar()
        self.var_cmp_output = tk.StringVar()
        self.var_cmp_subfolders = tk.BooleanVar(value=False)
        self.var_cmp_use_maxa = tk.BooleanVar(value=True)
        self.var_cmp_fallback_thua = tk.BooleanVar(value=True)
        self.var_cmp_move_a = tk.BooleanVar(value=True)
        self.var_cmp_move_b = tk.BooleanVar(value=True)
        self.var_cmp_copy_mode = tk.BooleanVar(value=False)

        # --- Mục 12: Chỉnh sửa PDF hàng loạt ---
        self.pdfedit_files = []              # list đường dẫn đầy đủ
        self.pdfedit_rotations = {}          # {file_path: {page_index0based: goc_xoay_them}}
        self.pdfedit_current_file = None
        self.pdfedit_current_page = 0        # 0-based
        self.pdfedit_total_pages = 0
        self.pdfedit_zoom = 1.0
        self.pdfedit_photo_ref = None
        self.var_pdfedit_subfolders = tk.BooleanVar(value=False)
        self.var_pdfedit_output = tk.StringVar()
        self.var_pdfedit_range = tk.StringVar(value="1-5")
        self.var_pdfedit_enhance_on = tk.BooleanVar(value=False)
        self.var_pdfedit_grayscale = tk.BooleanVar(value=False)
        self.var_pdfedit_contrast_on = tk.BooleanVar(value=False)
        self.var_pdfedit_contrast_val = tk.StringVar(value="1.5")
        self.var_pdfedit_sharpen = tk.BooleanVar(value=False)
        self.var_pdfedit_denoise = tk.BooleanVar(value=False)
        self.var_pdfedit_threshold_on = tk.BooleanVar(value=False)
        self.var_pdfedit_threshold_val = tk.StringVar(value="150")
        self.var_pdfedit_deskew = tk.BooleanVar(value=False)
        self.var_pdfedit_ocr_after = tk.BooleanVar(value=False)
        self.var_pdfedit_crop_border = tk.BooleanVar(value=False)
        self.var_pdfedit_normalize_a4 = tk.BooleanVar(value=False)

        self._build_ui()
        if getattr(self.app_config, "update_check_enabled", True):
            self.after(800, self._check_for_update_async)

    def _set_window_icon(self):
        base_dir = get_base_dir()
        ico_path = os.path.join(base_dir, "assets", "app_icon.ico")
        png_path = os.path.join(base_dir, "assets", "logo_icon_512.png")
        try:
            if os.path.isfile(ico_path):
                self.iconbitmap(ico_path)
                return
        except Exception:
            pass
        try:
            if os.path.isfile(png_path):
                img = tk.PhotoImage(file=png_path)
                self.iconphoto(True, img)
                self._icon_ref = img  # giữ tham chiếu tránh bị garbage-collect
        except Exception:
            pass

    def _build_banner(self):
        base_dir = get_base_dir()
        icon_path = os.path.join(base_dir, "assets", "logo_icon_512.png")
        outer = tk.Frame(self, bg="#f5f7f6")
        outer.pack(fill="x", side="top")

        header_row = tk.Frame(outer, bg="#f5f7f6")
        header_row.pack(fill="x")

        left_group = tk.Frame(header_row, bg="#f5f7f6")
        left_group.pack(side="left", padx=(10, 0), pady=6)

        # Chỉ icon logo là ẢNH (không chứa chữ nướng sẵn) — mọi CHỮ đều do chính Tkinter vẽ trực
        # tiếp bằng font hệ thống (Segoe UI đã set toàn ứng dụng), tránh lỗi hiển thị dấu tiếng Việt
        # có thể xảy ra khi chữ bị "nướng" sẵn vào ảnh PNG lúc đóng gói.
        try:
            from PIL import ImageTk, Image as PILImage
            if os.path.isfile(icon_path):
                img = PILImage.open(icon_path).resize((52, 52), PILImage.LANCZOS)
                self._logo_icon_ref = ImageTk.PhotoImage(img)
                tk.Label(left_group, image=self._logo_icon_ref, bg="#f5f7f6").pack(side="left", padx=(0, 10))
        except Exception:
            pass

        text_col = tk.Frame(left_group, bg="#f5f7f6")
        text_col.pack(side="left")
        tk.Label(text_col, text="SỸ LAND", font=("Segoe UI", 18, "bold"),
                 fg="#0d4434", bg="#f5f7f6", anchor="w", justify="left").pack(anchor="w")
        tk.Label(text_col,
                 text="Tiện ích CSDL Đất đai • Đổi tên hồ sơ • Tách/Gộp PDF • Xuất Word/PDF • Tổng hợp báo cáo",
                 font=("Segoe UI", 9), fg="#556055", bg="#f5f7f6", anchor="w", justify="left").pack(anchor="w")

        # Thanh phiên bản / nút cập nhật đặt bên PHẢI, cùng hàng với logo (header gọn, không chiếm thêm dòng)
        self._build_update_row(header_row)

        tk.Frame(outer, bg="#0d4434", height=3).pack(fill="x")

    def _build_update_row(self, parent):
        row = tk.Frame(parent, bg="#f5f7f6")
        row.pack(side="right", fill="y", padx=14)
        self.btn_open_download_page = tk.Button(
            row, text="⬇ Mở trang tải bản mới", font=("Segoe UI", 8), command=self.open_download_page,
            relief="flat", bg="#e8ede9", fg="#0d4434", cursor="hand2", state="disabled")
        self.btn_open_download_page.pack(side="right", pady=(0, 4), padx=(0, 6))
        tk.Button(row, text="🔄 Kiểm tra cập nhật", font=("Segoe UI", 8),
                  command=lambda: self._check_for_update_async(force=True),
                  relief="flat", bg="#e8ede9", fg="#0d4434", cursor="hand2").pack(side="right", pady=(0, 4))
        self.lbl_update_status = tk.Label(row, text="", font=("Segoe UI", 8), fg="#789", bg="#f5f7f6")
        self.lbl_update_status.pack(side="right", padx=8, pady=(0, 4))
        tk.Label(row, text=f"Phiên bản: v{APP_VERSION}", font=("Segoe UI", 8), fg="#789",
                 bg="#f5f7f6").pack(side="right", pady=(0, 4))

    def _insert_row_colored(self, tree, values, status_text):
        """
        Chèn 1 dòng vào Treeview với MÀU theo trạng thái - đúng mục V/VIII tài liệu "CHỈNH SỬA
        GIAO DIỆN TAB 7, TAB 8...": Xanh lá = thành công, Vàng = cần kiểm tra, Đỏ = lỗi,
        Xám = chờ xử lý/bỏ qua. Khớp theo CHỨA CHUỖI CON (không cần khớp chính xác) để dùng
        chung được cho MỌI bảng kết quả trong phần mềm dù mỗi nơi đặt tên trạng thái hơi khác
        nhau (VD "THANH_CONG"/"OK", "LOI_DOC_PDF"/"LOI_CHUYEN_PDF", "CAN_KIEM_TRA_BANG"...).
        Tự động tạo tag màu 1 lần duy nhất cho mỗi Treeview (idempotent - gọi nhiều lần không sao).
        """
        if not getattr(tree, "_da_cau_hinh_mau_trang_thai", False):
            tree.tag_configure("mau_thanh_cong", background="#c8e6c9")
            tree.tag_configure("mau_can_kiem_tra", background="#fff9c4")
            tree.tag_configure("mau_loi", background="#ffcdd2")
            tree.tag_configure("mau_bo_qua", background="#e0e0e0")
            tree._da_cau_hinh_mau_trang_thai = True

        s = (status_text or "").upper()
        if "LOI" in s or "ERROR" in s or "THIEU_DU_LIEU" in s:
            tag = "mau_loi"
        elif "CAN_KIEM_TRA" in s or s == "KHAC":
            tag = "mau_can_kiem_tra"
        elif "BO_QUA" in s or "DA_DUNG_BOI_NGUOI_DUNG" in s or "SKIP" in s:
            tag = "mau_bo_qua"
        elif ("THANH_CONG" in s or s == "OK" or "DA_TAO" in s or "DA_XOAY" in s
              or "DA_LAM_THANG" in s or s == "GIONG_NHAU"):
            tag = "mau_thanh_cong"
        else:
            tag = ""
        tree.insert("", "end", values=values, tags=(tag,) if tag else ())

    def _register_advanced_widget(self, widget, **pack_kwargs):
        """
        Đăng ký 1 widget là "chỉ dành cho chế độ Nâng cao" - đúng tính năng "Chế độ Cơ bản/Nâng
        cao" (ẩn bớt tùy chọn kỹ thuật/gỡ lỗi ít dùng cho người mới). Gọi ngay SAU KHI tạo widget
        (widget đã .pack() bình thường theo layout) - hàm này chỉ GHI NHỚ cách pack lại
        (pack_kwargs) để dùng khi cần hiện lại, rồi áp dụng NGAY trạng thái ẩn/hiện hiện tại.

        pack_kwargs: đúng các tham số đã dùng khi .pack() widget này (side, fill, padx, pady...)
        - dùng để pack lại chính xác khi chuyển sang chế độ Nâng cao.
        """
        if not hasattr(self, "_advanced_only_widgets"):
            self._advanced_only_widgets = []
        self._advanced_only_widgets.append((widget, pack_kwargs))
        is_advanced = getattr(getattr(self, "app_config", None), "che_do_nang_cao", False)
        if not is_advanced:
            widget.pack_forget()

    def _duong_dan_token_ban_quyen(self):
        return os.path.join(get_app_data_dir(), "license_token.dat")

    def _lay_trang_thai_ban_quyen(self):
        """Trả về (trang_thai, license_data) - đọc token cục bộ đã mã hóa, tính trạng thái hiện tại."""
        from app.services import license_service as _ls
        data = _ls.load_local_token(self._duong_dan_token_ban_quyen())
        return _ls.tinh_trang_thai(data), data

    def _ap_dung_gioi_han_ban_quyen(self):
        """
        Khóa/mở các Tab NGHIỆP VỤ CHÍNH theo trạng thái bản quyền - đúng mục XVII tài liệu "LỆNH
        BỔ SUNG PHẦN CÀI ĐẶT VÀ QUẢN LÝ BẢN QUYỀN": khi CHƯA kích hoạt/ĐÃ hết hạn/BỊ khóa, chỉ Tab
        Cài đặt (để đăng ký/kích hoạt/kiểm tra môi trường/liên hệ hỗ trợ) còn dùng được - KHÔNG
        đóng hẳn phần mềm (đúng yêu cầu "không được làm mất khả năng mở phần mềm để kích hoạt").
        Gọi lại hàm này SAU KHI kích hoạt thành công để mở lại ngay, không cần khởi động lại phần mềm.
        """
        trang_thai, _ = self._lay_trang_thai_ban_quyen()
        so_tab = self.notebook.index("end")
        idx_tab_cai_dat = so_tab - 1  # Tab Cài đặt luôn là tab CUỐI CÙNG

        hop_le = trang_thai in ("DA_KICH_HOAT", "DANG_DUNG_THU", "SAP_HET_HAN")

        for i in range(so_tab):
            if i == idx_tab_cai_dat:
                continue
            try:
                self.notebook.tab(i, state=("normal" if hop_le else "disabled"))
            except Exception:
                pass

        if hop_le:
            self.lbl_license_banner.pack_forget()
        else:
            thong_bao = {
                "CHUA_KICH_HOAT": "⚠ Phần mềm CHƯA ĐƯỢC KÍCH HOẠT - vui lòng đăng ký tài khoản và nhập key do Nguyễn Sỹ cấp ở Tab Cài đặt.",
                "HET_HAN": "⚠ Bản quyền ĐÃ HẾT HẠN - vui lòng liên hệ Nguyễn Sỹ để gia hạn (Tab Cài đặt).",
                "BI_KHOA": "⚠ Bản quyền ĐÃ BỊ KHÓA - vui lòng liên hệ Nguyễn Sỹ (Tab Cài đặt).",
            }.get(trang_thai, "⚠ Phần mềm chưa sẵn sàng sử dụng đầy đủ - vui lòng vào Tab Cài đặt.")
            self.lbl_license_banner.config(text=thong_bao, fg="white", bg="#c62828")
            self.lbl_license_banner.pack(fill="x")
            try:
                self.notebook.select(idx_tab_cai_dat)
            except Exception:
                pass

    def _apply_advanced_mode_visibility(self, is_advanced):
        """
        Ẩn/hiện TẤT CẢ widget đã đăng ký qua `_register_advanced_widget()` theo chế độ hiện tại.
        Gọi khi: khởi động phần mềm (theo cấu hình đã lưu), hoặc khi người dùng đổi checkbox
        "Chế độ NÂNG CAO" ở Tab Cài đặt (áp dụng NGAY, không cần khởi động lại phần mềm).
        """
        for widget, pack_kwargs in getattr(self, "_advanced_only_widgets", []):
            try:
                if is_advanced:
                    widget.pack(**pack_kwargs)
                else:
                    widget.pack_forget()
            except Exception:
                pass  # widget co the da bi huy (VD dong cua so con) - bo qua an toan

    def _add_pause_cancel(self, parent, control_attr):
        """
        Tạo 1 TaskControl mới lưu vào self.<control_attr>, cùng 2 nút "⏸ Tạm dừng" / "⏹ Hủy"
        (ban đầu vô hiệu hóa - self._enable_pause_cancel bật lên khi tác vụ bắt đầu chạy).
        Trả về (btn_pause, btn_cancel) để đặt vào layout.
        """
        control = TaskControl()
        setattr(self, control_attr, control)

        btn_pause = tk.Button(parent, text="⏸ Tạm dừng", state="disabled")
        btn_cancel = tk.Button(parent, text="⏹ Hủy", state="disabled")

        def toggle():
            if control.is_paused():
                control.resume()
                btn_pause.config(text="⏸ Tạm dừng")
                self.log("▶ Đã tiếp tục.")
            else:
                control.pause()
                btn_pause.config(text="▶ Tiếp tục")
                self.log("⏸ Đã tạm dừng — bấm 'Tiếp tục' để chạy tiếp, hoặc 'Hủy' để dừng hẳn.")

        def cancel():
            if messagebox.askyesno("Xác nhận hủy", "Hủy tác vụ đang chạy? Phần chưa xử lý sẽ bị bỏ qua."):
                control.cancel()
                self.log("⏹ Đang hủy...")

        btn_pause.config(command=toggle)
        btn_cancel.config(command=cancel)
        return btn_pause, btn_cancel

    def _enable_pause_cancel(self, control_attr, btn_pause, btn_cancel):
        getattr(self, control_attr).reset()
        btn_pause.config(state="normal", text="⏸ Tạm dừng")
        btn_cancel.config(state="normal")

    def _disable_pause_cancel(self, btn_pause, btn_cancel):
        btn_pause.config(state="disabled", text="⏸ Tạm dừng")
        btn_cancel.config(state="disabled")

    def _check_require_dry_run(self, dry_run_flag_attr, do_real_run, operation_label):
        """
        Áp dụng cài đặt 'Bắt buộc chạy thử trước khi chạy thật' (Tab Cài đặt). Nếu đang bật cài
        đặt này, đây là 1 lần CHẠY THẬT, và khu vực thao tác này CHƯA từng chạy thử trong phiên
        làm việc hiện tại -> hiện cảnh báo mạnh, để người dùng tự quyết định (không chặn cứng,
        vì có thể người dùng đã tự kiểm tra kỹ theo cách khác).
        Trả về True nếu được phép tiếp tục, False nếu người dùng chọn hủy.
        """
        if not do_real_run:
            return True
        if not getattr(self.app_config, "require_dry_run", False):
            return True
        if getattr(self, dry_run_flag_attr, False):
            return True
        return messagebox.askyesno(
            "⚠ Chưa chạy thử",
            f"Cài đặt hiện đang YÊU CẦU chạy thử trước khi chạy thật (Tab ⚙ Cài đặt).\n\n"
            f"Bạn chưa bấm 'Chạy thử' cho thao tác '{operation_label}' trong phiên làm việc này.\n\n"
            f"Bạn có chắc chắn muốn CHẠY THẬT LUÔN, bỏ qua bước chạy thử không?")

    def _maybe_auto_backup(self, file_paths, operation_label):
        """
        Áp dụng cài đặt 'Tự động sao lưu trước các thao tác rủi ro' (Tab Cài đặt). Nếu đang bật,
        tự động copy các file trong file_paths vào Backup/BACKUP_<thời điểm> TRƯỚC khi thực hiện
        đổi tên/di chuyển/chỉnh sửa thật. Không chặn thao tác nếu backup lỗi - chỉ ghi log cảnh báo,
        vì mục đích là AN TOÀN THÊM chứ không phải điều kiện bắt buộc.
        """
        if not getattr(self.app_config, "auto_backup", False):
            return
        if not file_paths:
            return
        run_id = time.strftime("%Y%m%d_%H%M%S")
        backup_base = os.path.join(get_app_data_dir(), "Backup")
        try:
            folder, done, errors = create_backup(file_paths, backup_base, run_id)
            self.log(f"🗄 Đã tự động sao lưu {len(done)} file vào: {folder}  (trước khi '{operation_label}')")
            if errors:
                self.log(f"⚠ Lỗi tự động sao lưu {len(errors)} file (thao tác vẫn tiếp tục): "
                         f"{[e[0] for e in errors][:3]}")
        except Exception as e:
            self.log(f"⚠ Tự động sao lưu thất bại (thao tác vẫn tiếp tục): {e}")

    def _make_scrollable(self, parent):
        """Bọc `parent` (1 tab) trong Canvas + Scrollbar dọc, trả về frame bên trong để chứa nội dung tab."""
        canvas = tk.Canvas(parent, highlightthickness=0)
        scrollbar = tk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas)

        inner_window = canvas.create_window((0, 0), window=inner, anchor="nw")
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(inner_window, width=e.width))
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        # Chỉ cuộn khi con trỏ chuột đang ở TRÊN canvas này (tránh xung đột giữa các tab khác nhau)
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _on_mousewheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

        return inner

    def _build_ui(self):
        pad = {"padx": 8, "pady": 5}

        self._build_banner()

        self.frm_license_banner = tk.Frame(self)
        self.frm_license_banner.pack(fill="x", padx=0, pady=0)
        self.lbl_license_banner = tk.Label(self.frm_license_banner, text="", font=("Segoe UI", 9, "bold"),
                                           pady=4)
        # (chỉ .pack() nhãn này khi thực sự cần cảnh báo - xem _ap_dung_gioi_han_ban_quyen)

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=6, pady=(6, 0))

        tab_workflow = tk.Frame(self.notebook)
        tab_excel = tk.Frame(self.notebook)
        tab_content = tk.Frame(self.notebook)
        tab_report = tk.Frame(self.notebook)
        tab_pdftools = tk.Frame(self.notebook)      # gộp: Tách/Gộp PDF + Chỉnh sửa PDF hàng loạt
        tab_word = tk.Frame(self.notebook)           # gộp: Gộp file Word (trước) + Xuất Word hàng loạt (sau)
        tab_reconcile_all = tk.Frame(self.notebook)  # gộp: Lọc file trùng + Lọc Excel theo PDF & hoàn thiện
        tab_settings = tk.Frame(self.notebook)

        self.notebook.add(tab_workflow, text="★ Quy trình xử lý hồ sơ")
        self.notebook.add(tab_excel, text="1-4. Đối chiếu & đổi tên (theo Excel)")
        self.notebook.add(tab_content, text="5. Đổi tên theo nội dung PDF")
        self.notebook.add(tab_pdftools, text="6. Xử lý PDF (Tách/Gộp • Xoay • Xem trước • Cải thiện)")
        self.notebook.add(tab_word, text="7. Chuyển đổi và xuất hồ sơ")
        self.notebook.add(tab_report, text="8. Tổng hợp và hoàn thiện báo cáo")
        self.notebook.add(tab_reconcile_all, text="9. Đối chiếu và hoàn thiện báo cáo")
        self.notebook.add(tab_settings, text="⚙ Cài đặt")

        self._build_tab_workflow(self._make_scrollable(tab_workflow), pad, {
            "content": tab_content, "pdftools": tab_pdftools, "word": tab_word,
            "report": tab_report, "reconcile": tab_reconcile_all,
        })
        self._build_tab_excel(self._make_scrollable(tab_excel), pad)
        self._build_tab_content(self._make_scrollable(tab_content), pad)
        pdftools_scrollable = self._make_scrollable(tab_pdftools)
        self._build_tab_split(pdftools_scrollable, pad)
        self._build_tab_pdfedit(pdftools_scrollable, pad)
        # Tab gộp 7+9: ưu tiên hiển thị "Gộp file Word" (Tab 9 cũ) TRƯỚC, rồi mới đến "Xuất Word" (Tab 7 cũ)
        word_scrollable = self._make_scrollable(tab_word)
        self._build_tab_merge_word(word_scrollable, pad)
        self._build_tab_word(word_scrollable, pad)
        self._build_tab_report(self._make_scrollable(tab_report), pad)
        reconcile_scrollable = self._make_scrollable(tab_reconcile_all)
        self._build_tab_dup(reconcile_scrollable, pad)
        self._build_tab_reconcile(reconcile_scrollable, pad)
        self._build_tab_settings(self._make_scrollable(tab_settings), pad)
        self._idx_tab_settings = 7  # dùng khi cần chuyển focus tới Tab Cài đặt (kích hoạt bản quyền)
        self._ap_dung_gioi_han_ban_quyen()

        frm_log = styled_labelframe(self, text="Kết quả / Nhật ký (dùng chung cho mọi tab)")
        frm_log.pack(fill="both", expand=False, padx=8, pady=(6, 0))
        self.txt_log = scrolledtext.ScrolledText(frm_log, font=("Consolas", 9), height=12)
        self.txt_log.pack(fill="both", expand=True)

        tk.Label(self, text="Phần mềm thuộc Nguyễn Sỹ  |  SĐT/Zalo: 0972560335  |  Email: minhsybk@gmail.com",
                 font=("Segoe UI", 8), fg="#666", pady=3).pack(fill="x", padx=8, pady=(2, 4))


    def log(self, msg):
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")
        self.update_idletasks()

    # ------------------- CHẾ ĐỘ OCR CHUYÊN DỤNG (PDF scan) -------------------


    # ------------------- MỤC 8: Tổng hợp báo cáo từ PDF -------------------


    # ------------------- MỤC 6: Tách / Gộp PDF -------------------


    # ------------------- MỤC 7: Xuất Word hàng loạt -------------------


    # ------------------- MỤC 9: Gộp file Word hàng loạt -------------------

    # ------------------- MỤC 10: Lọc & di chuyển file trùng -------------------

    # ------------------- MỤC 11: Lọc Excel tổng theo PDF & hoàn thiện báo cáo -------------------


    # ------------------- MỤC 12: Chỉnh sửa PDF hàng loạt -------------------


    # ------------------- TỰ ĐỘNG KIỂM TRA / CẬP NHẬT PHIÊN BẢN -------------------

    def _check_for_update_async(self, force=False):
        """
        Kiểm tra phiên bản mới từ latest_version.json (online). Mặc định (force=False, dùng lúc
        khởi động) chỉ kiểm tra tối đa 1 lần/ngày để không làm chậm phần mềm. Bấm nút "Kiểm tra
        cập nhật" thủ công (force=True) luôn kiểm tra ngay bất kể đã kiểm tra hôm nay chưa.
        Mọi lỗi (mất mạng, server lỗi...) đều bị BỎ QUA ÂM THẦM - không hiện lỗi khó hiểu, không
        làm phần mềm treo hay bị ảnh hưởng.
        """
        if hasattr(self, "lbl_update_status"):
            self.lbl_update_status.config(text="Đang kiểm tra...")

        def worker():
            info = None
            try:
                if not force and not should_check_update_today():
                    self.after(0, lambda: self._set_update_status("Đã kiểm tra hôm nay."))
                    return
                check_url = get_update_check_url()
                if not check_url:
                    self.after(0, lambda: self._set_update_status(
                        "Chưa cấu hình nơi kiểm tra cập nhật (xem update_check_config.json)."))
                    return
                info = fetch_latest_version_info(check_url, timeout=UPDATE_CHECK_TIMEOUT_SECONDS)
                mark_update_checked_today()
            except Exception:
                info = None  # im lặng bỏ qua - không được làm phần mềm lỗi/treo vì kiểm tra cập nhật

            def finish():
                if not info or not info.get("latest_version"):
                    self._set_update_status("Không kiểm tra được cập nhật (không có mạng?).")
                    return
                remote_v = info["latest_version"]
                if parse_version_tuple(remote_v) > parse_version_tuple(APP_VERSION):
                    self._set_update_status(f"Có bản mới v{remote_v}!")
                    self.latest_update_info = info
                    if hasattr(self, "btn_open_download_page"):
                        self.btn_open_download_page.config(state="normal")
                    self.show_update_dialog(info)
                else:
                    self._set_update_status("Đang dùng bản mới nhất.")

            self.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()

    def _set_update_status(self, text):
        if hasattr(self, "lbl_update_status"):
            self.lbl_update_status.config(text=text)

    def show_update_dialog(self, info):
        """
        Hiện hộp thoại thông báo có bản mới. Chỉ THÔNG BÁO + mở link tải bằng trình duyệt khi
        người dùng đồng ý - KHÔNG tự tải/tự ghi đè file .exe đang chạy, KHÔNG tự đóng phần mềm,
        KHÔNG tự chạy file cài đặt mới khi chưa được xác nhận.
        """
        remote_v = info.get("latest_version", "?")
        notes = info.get("release_notes") or []
        notes_text = "\n".join(f"• {n}" for n in notes) if notes else "(không có ghi chú)"
        download_url = (info.get("download_url") or "").strip()
        force_update = bool(info.get("force_update", False))

        win = tk.Toplevel(self)
        win.title("Có phiên bản mới")
        win.geometry("480x360")
        win.transient(self)
        win.grab_set()

        tk.Label(win, text="Đã có phiên bản SỸ LAND mới. Anh có muốn tải về không?",
                 font=("Segoe UI", 12, "bold"), fg="#0d4434", wraplength=440,
                 justify="left").pack(pady=(16, 4), padx=16, fill="x")
        tk.Label(win, text=f"Phiên bản mới: v{remote_v}   (đang dùng: v{APP_VERSION})",
                 font=("Segoe UI", 9), fg="#555").pack()
        if info.get("release_date"):
            tk.Label(win, text=f"Ngày phát hành: {info['release_date']}",
                     font=("Segoe UI", 9), fg="#777").pack()
        if force_update:
            tk.Label(win, text="⚠ Đây là bản cập nhật quan trọng, khuyến nghị cập nhật sớm.",
                     font=("Segoe UI", 9, "bold"), fg="#c62828").pack(pady=(4, 0))

        tk.Label(win, text="Nội dung cập nhật:", font=("Segoe UI", 9, "bold"), anchor="w").pack(
            fill="x", padx=16, pady=(14, 2))
        txt = scrolledtext.ScrolledText(win, height=8, font=("Segoe UI", 9), wrap="word")
        txt.pack(fill="both", expand=True, padx=16)
        txt.insert("1.0", notes_text)
        txt.config(state="disabled")

        row_btn = tk.Frame(win)
        row_btn.pack(fill="x", pady=14, padx=16)

        def on_download():
            win.destroy()
            if download_url:
                import webbrowser
                webbrowser.open(download_url)
                self._set_update_status(f"Đã mở trang tải v{remote_v} bằng trình duyệt.")
            else:
                messagebox.showinfo("Thiếu đường dẫn tải",
                                     "Không có đường dẫn tải (download_url) trong thông tin phiên bản mới.")

        def on_skip():
            win.destroy()
            self._set_update_status(f"Đã bỏ qua v{remote_v} (đang dùng v{APP_VERSION}).")

        tk.Button(row_btn, text="✓ Tải về ngay", font=("Segoe UI", 10, "bold"),
                  bg="#2e7d32", fg="white", command=on_download).pack(side="left", padx=4)
        tk.Button(row_btn, text="Để sau", font=("Segoe UI", 10), command=on_skip).pack(side="left", padx=4)

    def open_download_page(self):
        """Mở trang tải bản mới nhất đã biết (nút 'Mở trang tải bản mới' ở khu vực Giới thiệu/header)."""
        info = getattr(self, "latest_update_info", None)
        if not info or not info.get("download_url"):
            messagebox.showinfo("Chưa có bản mới",
                                 "Chưa phát hiện phiên bản mới nào. Bấm 'Kiểm tra cập nhật' trước.")
            return
        import webbrowser
        webbrowser.open(info["download_url"])


def apply_default_fonts(root):
    """
    Áp dụng font mặc định cho TOÀN BỘ ứng dụng để tránh lỗi hiển thị chữ tiếng Việt
    (dấu bị mất/vỡ) trên một số máy Windows dùng font hệ thống khác nhau.
    Ưu tiên "Segoe UI" (font chuẩn Windows, hỗ trợ tiếng Việt tốt), dự phòng "Arial",
    cuối cùng mới dùng font mặc định của Tk nếu không có font nào ở trên.
    """
    preferred = ["Segoe UI", "Arial", "DejaVu Sans", "TkDefaultFont"]
    try:
        available = set(tkfont.families(root))
    except Exception:
        available = set()

    chosen = next((f for f in preferred if f in available), None)
    if not chosen:
        return  # giữ nguyên font mặc định của hệ thống nếu không tìm được font nào phù hợp

    size = 9
    try:
        for name in ("TkDefaultFont", "TkTextFont", "TkMenuFont", "TkHeadingFont",
                     "TkCaptionFont", "TkSmallCaptionFont", "TkIconFont", "TkTooltipFont"):
            f = tkfont.nametofont(name, root=root)
            f.configure(family=chosen, size=size)
    except Exception:
        pass

    # Áp cho các widget tk cổ điển (Label/Button/Entry/Text...) không tự theo TkDefaultFont
    root.option_add("*Font", (chosen, size))


class LoginWindow(tk.Tk):
    """Màn hình đăng nhập hiện trước khi vào phần mềm chính. Đăng nhập đúng -> self.success=True."""

    def __init__(self):
        super().__init__()
        self.success = False
        apply_default_fonts(self)
        self.title("Đăng nhập - Tiện ích CSDL Đất Đai")
        self.geometry("400x350")
        self.resizable(False, False)
        try:
            self._set_icon()
        except Exception:
            pass

        self.protocol("WM_DELETE_WINDOW", self._on_close)

        frm = tk.Frame(self, padx=24, pady=24)
        frm.pack(fill="both", expand=True)

        tk.Label(frm, text="TIỆN ÍCH CSDL ĐẤT ĐAI", font=("Poppins" if self._font_ok("Poppins") else "Arial", 14, "bold"),
                 fg="#0f4c3a").pack(pady=(0, 4))
        tk.Label(frm, text="SỸ LAND", font=("Arial", 10), fg="#666").pack(pady=(0, 18))

        tk.Label(frm, text="Email:", anchor="w").pack(fill="x")
        self.var_id = tk.StringVar(value="")
        entry_id = tk.Entry(frm, textvariable=self.var_id)
        entry_id.pack(fill="x", pady=(2, 12))

        tk.Label(frm, text="Mật khẩu:", anchor="w").pack(fill="x")
        self.var_pw = tk.StringVar(value="")
        entry_pw = tk.Entry(frm, textvariable=self.var_pw, show="•")
        entry_pw.pack(fill="x", pady=(2, 4))

        self.lbl_error = tk.Label(frm, text="", fg="#c62828")
        self.lbl_error.pack(fill="x", pady=(4, 8))

        btn = tk.Button(frm, text="Đăng nhập", font=("Arial", 10, "bold"), bg="#0f4c3a", fg="white",
                         command=self._try_login)
        btn.pack(fill="x", pady=(6, 0))
        tk.Button(frm, text="Đăng ký tài khoản mới", command=self._open_register).pack(fill="x", pady=(7, 0))
        tk.Label(frm, text="Tài khoản dùng chung với website SỸ LAND", fg="#66736c").pack(pady=(8, 0))

        entry_id.bind("<Return>", lambda e: entry_pw.focus_set())
        entry_pw.bind("<Return>", lambda e: self._try_login())
        entry_id.focus_set()

        self.update_idletasks()
        w, h = 400, 350
        x = (self.winfo_screenwidth() - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _font_ok(self, name):
        try:
            return name in tkfont.families()
        except Exception:
            return False

    def _set_icon(self):
        base_dir = get_base_dir()
        ico_path = os.path.join(base_dir, "assets", "app_icon.ico")
        if os.path.isfile(ico_path):
            self.iconbitmap(ico_path)

    def _try_login(self):
        self.lbl_error.config(text="Đang xác thực…", fg="#455a64")
        self.update_idletasks()
        try:
            from app.services import account_service
            self.auth_session = account_service.sign_in(self.var_id.get(), self.var_pw.get())
            self.success = True
            self.destroy()
        except Exception as exc:
            self.lbl_error.config(text=str(exc), fg="#c62828")
            self.var_pw.set("")

    def _open_register(self):
        win = tk.Toplevel(self)
        win.title("Đăng ký tài khoản SỸ LAND")
        win.geometry("420x390")
        win.resizable(False, False)
        frm = tk.Frame(win, padx=24, pady=20)
        frm.pack(fill="both", expand=True)
        values = {"name": tk.StringVar(), "email": tk.StringVar(value=self.var_id.get().strip()), "password": tk.StringVar(), "confirm": tk.StringVar()}
        for label, key, show in [("Họ và tên", "name", ""), ("Email", "email", ""), ("Mật khẩu (ít nhất 8 ký tự)", "password", "•"), ("Xác nhận mật khẩu", "confirm", "•")]:
            tk.Label(frm, text=label, anchor="w").pack(fill="x")
            tk.Entry(frm, textvariable=values[key], show=show).pack(fill="x", pady=(2, 10))
        status = tk.Label(frm, text="", fg="#c62828", wraplength=350, justify="left")
        status.pack(fill="x", pady=(2, 8))

        def submit():
            if values["password"].get() != values["confirm"].get():
                status.config(text="Mật khẩu xác nhận chưa khớp.")
                return
            status.config(text="Đang đăng ký…", fg="#455a64")
            win.update_idletasks()
            try:
                from app.services import account_service
                result = account_service.sign_up(values["email"].get(), values["password"].get(), values["name"].get())
                self.var_id.set(values["email"].get().strip().lower())
                self.var_pw.set("")
                if result.get("session") or result.get("access_token"):
                    status.config(text="Đăng ký thành công. Bạn có thể đăng nhập ngay.", fg="#2e7d32")
                else:
                    status.config(text="Đã đăng ký. Hãy kiểm tra email để xác nhận tài khoản rồi đăng nhập.", fg="#2e7d32")
            except Exception as exc:
                status.config(text=str(exc), fg="#c62828")

        tk.Button(frm, text="Tạo tài khoản", bg="#0f4c3a", fg="white", command=submit).pack(fill="x")

    def _on_close(self):
        self.success = False
        self.destroy()


if __name__ == "__main__":
    login = LoginWindow()
    login.mainloop()
    if login.success:
        app = App()
        app.mainloop()
